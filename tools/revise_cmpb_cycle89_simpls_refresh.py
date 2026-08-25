from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
IN_DIR = ROOT / "artifacts/CMPB_rewrite_20260824_cycle88"
OUT_DIR = ROOT / "artifacts/CMPB_rewrite_20260824_cycle89"
MAIN_IN = IN_DIR / "fastPLS_CMPB_main_cycle88_0.99.23_20260824.docx"
SUPP_IN = IN_DIR / "fastPLS_CMPB_supplement_cycle88_0.99.23_20260824.docx"
MAIN_OUT = OUT_DIR / "fastPLS_CMPB_main_cycle89_0.99.24_20260824.docx"
SUPP_OUT = OUT_DIR / "fastPLS_CMPB_supplement_cycle89_0.99.24_20260824.docx"


def replace_exact(document, old, new):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == old]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph, found {len(matches)}: {old[:100]}")
    paragraph = matches[0]
    style = paragraph.style
    paragraph.text = new
    paragraph.style = style


def set_cell(cell, text, size=5):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)


OUT_DIR.mkdir(parents=True, exist_ok=True)

main = Document(MAIN_IN)
replace_exact(
    main,
    "The public pls() interface selects PLS family, component count, solver, backend, and, for classification, argmax or latent-space linear discriminant analysis (LDA). pls.single.cv() selects settings within one cross-validation layer; pls.double.cv() uses nested cross-validation. The audited 0.99.23 interface exposes PLS modelling, prediction, evaluation, cross-validation, score plotting, and standalone truncated SVD; PCA and nearest-neighbour classifiers are not package APIs. Requested estimators are never silently substituted.",
    "The public pls() interface selects PLS family, component count, solver, backend, and, for classification, argmax or latent-space linear discriminant analysis (LDA). pls.single.cv() selects settings within one cross-validation layer; pls.double.cv() uses nested cross-validation. The audited 0.99.24 interface exposes PLS modelling, prediction, evaluation, cross-validation, score plotting, and standalone truncated SVD; PCA and nearest-neighbour classifiers are not package APIs. Requested estimators are never silently substituted.",
)
replace_exact(
    main,
    "SIMPLS follows de Jong's sequential score, loading, orthogonalization, and rank-one deflation equations [11]. The innovation is computational: fastPLS retains deflation products, latent quantities, coefficients, and predictions incrementally, so all requested component counts are snapshots of one maximal path rather than independent refits. Each direction is still extracted from the current deflated state, and deterministic IRLBA comparisons are evaluated against prespecified numerical tolerances.",
    "SIMPLS follows de Jong's sequential score, loading, orthogonalization, and rank-one deflation equations [11]. The innovation is computational: fastPLS retains deflation products, latent quantities, coefficients, and predictions incrementally, so all requested component counts are snapshots of one maximal path rather than independent refits. Every component invokes a fresh rank-one direction calculation on the current deflated cross-covariance: IRLBA starts a new iterative solve, whereas rSVD draws a new oversampled sketch using the base seed plus the zero-based component index. Candidate-direction blocks, cross-component warm starts, and adaptive refresh were rejected during development and are not used by the public CPU, CUDA, or Metal paths. Retained optimizations cache rank-one deflation products, update coefficient and fitted-value paths incrementally, cache cross-products when shape-appropriate, use compact prediction factors, and support implicit cross-covariance products."
)
replace_exact(
    main,
    "Code and benchmark outputs are available at https://github.com/tkcaccia/fastPLS; reusable components are at https://github.com/tkcaccia/kodama-cpp. Historical quantitative results remain tied to fastPLS 0.99.6, base commit 6e50bd318f20289101f6b723953830aefa8b95d6, and source-archive SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85; they are not relabelled as reruns. The reviewed interface and definitive multi-seed rSVD qualification use fastPLS 0.99.23. Its exact source archive fastPLS_0.99.23.tar.gz has SHA-256 bd8bd6b0dd85219cd9dbe25b429cea02d0a6633df490618d3a3ad73f177fd5e8. Analysis-specific scripts and archive digests are reported in Supplementary Table S15.",
    "Code and benchmark outputs are available at https://github.com/tkcaccia/fastPLS; reusable components are at https://github.com/tkcaccia/kodama-cpp. Historical quantitative results remain tied to fastPLS 0.99.6, base commit 6e50bd318f20289101f6b723953830aefa8b95d6, and source-archive SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85; they are not relabelled as reruns. The reviewed interface is fastPLS 0.99.24. Its exact source archive fastPLS_0.99.24.tar.gz has SHA-256 f27fa7d3c96d6d8a550857bab23146be5b4cefe2868ce17d4bfcc0fe7c87b5db. The definitive CPU/CUDA multi-seed rSVD qualification was generated with 0.99.23 before the documentation and Metal direction-refresh correction; the active CPU/CUDA numerical path and qualified controls are unchanged in 0.99.24. Analysis-specific scripts and archive digests are reported in Supplementary Table S15."
)
main.save(MAIN_OUT)

supp = Document(SUPP_IN)
replace_exact(
    supp,
    "This supplement distinguishes the reviewed fastPLS 0.99.23 interface and multi-seed rSVD qualification from the 0.99.6 source archive used for historical quantitative benchmarks. Table S15 maps each analysis to its result archive and records an exact package commit only when run metadata captured it. A later package version, result date, or manuscript commit is never treated as evidence of a historical computational SHA, and historical values are not relabelled as 0.99.23 reruns.",
    "This supplement distinguishes the reviewed fastPLS 0.99.24 interface from the 0.99.6 source archive used for historical quantitative benchmarks and the 0.99.23 archive used for the definitive CPU/CUDA multi-seed rSVD qualification. Table S15 maps each analysis to its result archive and records an exact package commit only when run metadata captured it. A later package version, result date, or manuscript commit is never treated as evidence of a historical computational SHA, and historical values are not relabelled as later-version reruns."
)
replace_exact(
    supp,
    "Incremental deflation, coefficient, and prediction updates reorganize numerical work while retaining sequential SIMPLS orthogonalization and deflation. The former one-vector warm-start and adaptive randomized refresh were rejected by formal validation and are not used by the public algorithm.",
    "Incremental deflation, coefficient, and prediction updates reorganize numerical work while retaining sequential SIMPLS orthogonalization and deflation. The public CPU, CUDA, and Metal paths all request one newly computed direction from the current deflated operator for every component. The preceding latent direction is never inserted into the next solver start, and candidate blocks are never consumed across deflation steps. Retained optimizations are cached rank-one deflation products, incremental coefficient and fitted-value updates, conditional cross-product caching, compact prediction, implicit cross-covariance products, and reusable allocation workspaces. Cross-component warm starts, multi-direction block refresh, and adaptive refresh policies were rejected and removed from the public execution path."
)
replace_exact(
    supp,
    "Current benchmark workflows record repository state, benchmark-script checksum, package version, source-archive SHA-256, compiler, BLAS/LAPACK, thread settings, accelerator libraries, seeds, rSVD controls, and data/split identifiers. Table S15 maps each analysis to its exact evidence archive. NMR and ImageNet used fastPLS_0.99.6.tar.gz (SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85) from base commit 6e50bd318f20289101f6b723953830aefa8b95d6; these benchmark values are not relabelled as later-version reruns. The reviewed 0.99.23 source archive has SHA-256 bd8bd6b0dd85219cd9dbe25b429cea02d0a6633df490618d3a3ad73f177fd5e8 and generated the definitive multi-seed rSVD qualification.",
    "Current benchmark workflows record repository state, benchmark-script checksum, package version, source-archive SHA-256, compiler, BLAS/LAPACK, thread settings, accelerator libraries, seeds, rSVD controls, and data/split identifiers. Table S15 maps each analysis to its exact evidence archive. NMR and ImageNet used fastPLS_0.99.6.tar.gz (SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85) from base commit 6e50bd318f20289101f6b723953830aefa8b95d6; these benchmark values are not relabelled as later-version reruns. The 0.99.23 archive (SHA-256 bd8bd6b0dd85219cd9dbe25b429cea02d0a6633df490618d3a3ad73f177fd5e8) generated the definitive multi-seed CPU/CUDA rSVD qualification. The reviewed 0.99.24 archive (SHA-256 f27fa7d3c96d6d8a550857bab23146be5b4cefe2868ce17d4bfcc0fe7c87b5db) standardizes fresh-per-component direction extraction across CPU, CUDA, and Metal and adds backend-invariance tests and fitted-model rule diagnostics."
)
replace_exact(
    supp,
    "The ledger never infers a historical commit from a package version or result date. Where a run did not record its Git SHA, the source status is explicitly not recoverable. SHA-256 prefixes identify immutable result archives; full digests are retained in the machine-readable ledger. The 0.99.23 release qualification is documented in publication_release_0.99.23/rsvd_qualification and does not alter historical result provenance.",
    "The ledger never infers a historical commit from a package version or result date. Where a run did not record its Git SHA, the source status is explicitly not recoverable. SHA-256 prefixes identify immutable result archives; full digests are retained in the machine-readable ledger. The 0.99.23 solver-control qualification is documented in publication_release_0.99.23/rsvd_qualification and does not alter historical result provenance; the 0.99.24 archive is the reviewed public implementation."
)

# Clarify that rSVD oversampling is not cross-component direction reuse.
storage = supp.tables[1]
for row in storage.rows:
    if row.cells[0].text == "SIMPLS":
        set_cell(
            row.cells[3],
            "Cached deflation products and incremental prediction path; each rSVD direction is approximate but freshly computed, with no cross-component candidate reuse",
        )

# Add the reviewed release to the provenance ledger without rewriting the
# historical 0.99.23 qualification entry.
ledger = supp.tables[18]
cells = ledger.add_row().cells
values = (
    "A20",
    "Public SIMPLS direction-refresh rule; CPU/CUDA/Metal dispatch tests",
    "fastPLS_0.99.24.tar.gz; tests/testthat/test-simpls-direction-refresh.R",
    "0.99.24",
    "exact source archive; CPU, native Metal, and CUDA runtime checks",
    "tests/testthat/test-simpls-direction-refresh.R",
    "f27fa7d3c96d",
)
for cell, value in zip(cells, values):
    set_cell(cell, value)

supp.save(SUPP_OUT)
print(MAIN_OUT)
print(SUPP_OUT)
