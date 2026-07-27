#!/usr/bin/env python3
"""Add the authoritative route-level float32 capability and evidence tables."""

import csv
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle64"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle65"
EVIDENCE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle65_20260726"
)
SOURCE_EVIDENCE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle13_20260725"
)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle64_0.99.6_20260726.docx"
SUPP_SOURCE = (
    SOURCE / "fastPLS_CMPB_supplement_cycle64_0.99.6_20260726.docx"
)
MAIN_OUTPUT = OUTPUT / "fastPLS_CMPB_main_cycle65_0.99.6_20260726.docx"
SUPP_OUTPUT = (
    OUTPUT / "fastPLS_CMPB_supplement_cycle65_0.99.6_20260726.docx"
)

CAPABILITY = SOURCE_EVIDENCE / "float32_capability_matrix.csv"
PAIRED = SOURCE_EVIDENCE / "float32_float64_paired_resources.csv"
AGREEMENT = SOURCE_EVIDENCE / "float32_float64_controlled_agreement.csv"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_paragraph(document, prefix, text):
    paragraph = find_paragraph(document, prefix)
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def remove_old_capability_table(document):
    expected = [
        "Model",
        "Backend",
        "Solver",
        "Status",
        "Evidence",
        "Observed limitation",
        "Automatic warning",
    ]
    for table in document.tables:
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header == expected:
            table._element.getparent().remove(table._element)
            return
    raise RuntimeError("Old float32 capability table not found")


def format_table(table, font_size=5.0, first_left=True):
    table.style = "Table"
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for row_index, row in enumerate(table.rows):
        row_properties = row._tr.get_or_add_trPr()
        keep_row = OxmlElement("w:cantSplit")
        row_properties.append(keep_row)
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if first_left and column_index == 0
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    if row_index == 0:
                        run.bold = True


def add_table(document, headers, rows, font_size=5.0):
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    format_table(table, font_size=font_size)
    return table


def fmt(value, digits=3):
    if pd.isna(value):
        return "NR"
    return f"{float(value):.{digits}f}"


def revise_main():
    document = Document(MAIN_SOURCE)

    methods = replace_paragraph(
        document,
        "Standard R numeric matrices store",
        (
            "Standard R numeric matrices store each value in double precision "
            "(float64; eight bytes), whereas matrices supplied through the "
            "float package use four-byte float32 values. Float64 is the "
            "numerical reference. Float32 is not a package-wide capability: "
            "each PLS-family, backend, solver, and classifier route is classified "
            "as validated, experimental, hybrid, unavailable, or failed. "
            "Compiled CPU and CUDA rSVD PLS-SVD have the broadest current "
            "validation. Metal and float32 IRLBA are experimental; accelerator "
            "OPLS and nonlinear kernel PLS retain host stages; Metal LDA uses "
            "the compiled CPU float32 solver after score projection; and the "
            "Windows fallback is restricted to CPU rSVD PLS-SVD, SIMPLS, and "
            "linear kernel PLS with argmax. Public calls warn once for "
            "experimental, hybrid, or measured-risk routes and stop before "
            "allocation for unavailable combinations. The authoritative route "
            "matrix and paired precision/resource measurements are reported in "
            "Supplementary Tables S60-S62."
        ),
    )
    methods.paragraph_format.keep_together = True

    results = replace_paragraph(
        document,
        "Single-precision inputs provide",
        (
            "Single precision was route-dependent rather than uniformly "
            "beneficial. Across the paired MetRef and PRISM tests, float32 "
            "reduced stored input size by approximately one half, but "
            "float32/float64 runtime ratios ranged from 0.17 to 14.18 and "
            "incremental host-RSS changes ranged from a 336-MB reduction to a "
            "34-MB increase. PLS-SVD preserved MetRef accuracy on CPU and CUDA; "
            "SIMPLS and linear kernel PLS were two percentage points lower in "
            "the current matched runs. PRISM RMSD differences were small "
            "(absolute differences <=0.00170), although several float32 routes "
            "were slower. Small controlled tests showed decoded-label agreement "
            "of 1.0 and regression prediction correlations above 0.9999999999 "
            "for all four families on CPU, CUDA, and Metal, but these tests do "
            "not validate large routes. GPU-baseline-corrected workspace memory "
            "was not recorded in the paired MetRef/PRISM archive; sampled GPU "
            "use is therefore reported separately and is not relabelled as "
            "workspace memory (Supplementary Tables S60-S62)."
        ),
    )
    results.paragraph_format.keep_together = True

    replace_paragraph(
        document,
        "Float32 approximately halves raw input storage",
        (
            "Float32 halves raw representation size but is not a general speed, "
            "peak-memory, or numerical-equivalence advantage. The authoritative "
            "route matrix separates numerical validation from execution "
            "residency and platform availability. In particular, moderate "
            "PLS-SVD routes were stable, SIMPLS-derived classification was "
            "precision-sensitive, accelerator OPLS and nonlinear kernel PLS "
            "were hybrid, and extreme-response NMR exposed failed or "
            "impractical routes. The public interface now warns or errors "
            "consistently before allocation, and confirmatory analysis should "
            "use float64 CPU IRLBA unless the selected float32 route has been "
            "validated on the target matrix shape and held-out task."
        ),
    )

    replace_paragraph(
        document,
        "fastPLS combines a shape-aware",
        (
            "fastPLS combines a shape-aware accelerated sequential SIMPLS "
            "implementation with memory-aware PLS-SVD, compiled validation, "
            "compact prediction, route-conditional float32 execution, and "
            "latent-space LDA. Compiled float64 CPU is the reference route. "
            "NVIDIA CUDA, Apple Metal, float32 storage, and rSVD are reported "
            "only for the model/backend combinations that passed the stated "
            "numerical checks; experimental, hybrid, unavailable, and failed "
            "routes remain visible rather than being pooled into a general "
            "acceleration claim. The GPL-3 R package calls reusable C++ "
            "components maintained with the MIT-licensed kodama-cpp codebase."
        ),
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - route-level float32 capability"
    )
    document.save(MAIN_OUTPUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)
    remove_old_capability_table(document)
    replace_paragraph(
        document,
        "Implementation support and numerical validation are distinct.",
        (
            "Implementation support and numerical validation are distinct. "
            "The earlier family-by-backend summary has been superseded by the "
            "single authoritative route matrix in Table S60, which includes "
            "family, kernel scope, backend, prediction endpoint, residency, "
            "solver support, platform constraints, Windows behavior, "
            "extreme-response status, and public API action. Status has one of "
            "five meanings: validated, experimental, hybrid, unavailable, or "
            "failed. 'Hybrid' denotes a callable route with material host and "
            "device stages; it does not imply numerical failure."
        ),
    )
    replace_paragraph(
        document,
        "Warnings are evaluated before model workspaces",
        (
            "Warnings and errors are evaluated before model workspaces are "
            "allocated and depend on the requested estimator, backend, solver, "
            "classifier, kernel, platform, response dimension, and component "
            "count rather than dataset identity. Unavailable combinations stop. "
            "Experimental and hybrid combinations warn once. For numeric "
            "responses with q >= 10,000 and at least 50 components, PLS-SVD "
            "warns about the measured performance/memory risk; SIMPLS, OPLS, "
            "and kernel PLS are labelled failed for that validation regime and "
            "emit a numerical/performance warning."
        ),
    )
    replace_paragraph(
        document,
        "Table S36.",
        (
            "The previous Table S36 has been withdrawn to avoid two competing "
            "capability summaries. Tables S60-S62 are authoritative."
        ),
    )
    replace_paragraph(
        document,
        "Measured anchors. Float32 approximately halved",
        (
            "The paired evidence now separates raw input storage, pre-fit "
            "baseline RSS, absolute peak RSS, incremental host RSS, sampled GPU "
            "use, runtime, predictive metric, and controlled prediction "
            "agreement. GPU-baseline-corrected workspace memory was not recorded "
            "for the MetRef/PRISM precision runs and is reported as NR rather "
            "than inferred."
        ),
    )

    capability = pd.read_csv(CAPABILITY)
    paired = pd.read_csv(PAIRED)
    agreement = pd.read_csv(AGREEMENT)

    document.add_heading(
        "S44. Authoritative float32 route capability and paired evidence",
        level=1,
    )
    document.add_paragraph(
        (
            "Table S60 is the single capability authority for the public PLS "
            "interface. The matrix describes current implementation and "
            "validation scope, not a promise that every future matrix shape "
            "will be numerically equivalent. Table S61 reports paired "
            "real-data metric, runtime, and memory measurements. Table S62 "
            "reports complete prediction agreement from the controlled route "
            "screen. Float64 remains the confirmatory reference."
        ),
        style="First Paragraph",
    )

    rows = []
    for _, row in capability.iterrows():
        rows.append(
            [
                row["family"],
                row["kernel_scope"],
                row["backend"],
                row["endpoint"].replace("classification: ", "class: "),
                row["status"],
                row["execution_residency"],
                row["supported_solver"],
                row["public_api_behavior"],
                row["windows_status"],
                row["extreme_response_status"],
            ]
        )
    add_table(
        document,
        [
            "Family",
            "Kernel",
            "Backend",
            "Endpoint",
            "Status",
            "Residency",
            "Solver",
            "API",
            "Windows",
            "Extreme q",
        ],
        rows,
        font_size=4.2,
    )
    document.add_paragraph(
        (
            "Table S60. Authoritative float32 capability matrix. Extreme q "
            "denotes numeric responses with q >= 10,000 and at least 50 "
            "components. The machine-readable table additionally records "
            "evidence, limitations, platform constraints, and exact warning/"
            "error behavior."
        ),
        style="Caption",
    )

    performance_rows = []
    for _, row in paired.iterrows():
        performance_rows.append(
            [
                row["dataset"],
                row["family"],
                row["backend"],
                row["metric_name"],
                f'{fmt(row["metric_float64"], 5)} / {fmt(row["metric_float32"], 5)}',
                fmt(row["metric_delta_float32_minus_float64"], 5),
                f'{fmt(row["total_time_sec_float64"], 3)} / {fmt(row["total_time_sec_float32"], 3)}',
                fmt(row["time_ratio_float32_over_float64"], 2),
                f'{fmt(row["input_storage_mb_float64"], 1)} / {fmt(row["input_storage_mb_float32"], 1)}',
            ]
        )
    add_table(
        document,
        [
            "Dataset",
            "Family",
            "Backend",
            "Metric",
            "f64 / f32",
            "Delta",
            "Time f64 / f32 (s)",
            "f32/f64",
            "Input f64 / f32 (MB)",
        ],
        performance_rows,
        font_size=4.8,
    )
    document.add_paragraph(
        (
            "Table S61a. Paired float64/float32 predictive metric, total "
            "fitting-plus-prediction time, and stored input size. Values are "
            "medians of three isolated runs; conversion occurred before timing."
        ),
        style="Caption",
    )

    memory_rows = []
    for _, row in paired.iterrows():
        memory_rows.append(
            [
                row["dataset"],
                row["family"],
                row["backend"],
                f'{fmt(row["baseline_host_rss_mb_float64"], 1)} / {fmt(row["baseline_host_rss_mb_float32"], 1)}',
                f'{fmt(row["peak_host_rss_mb_float64"], 1)} / {fmt(row["peak_host_rss_mb_float32"], 1)}',
                f'{fmt(row["incremental_host_rss_mb_float64"], 1)} / {fmt(row["incremental_host_rss_mb_float32"], 1)}',
                fmt(row["incremental_host_rss_delta_mb"], 1),
                f'{fmt(row["sampled_peak_gpu_used_mb_float64"], 0)} / {fmt(row["sampled_peak_gpu_used_mb_float32"], 0)}',
                "NR / NR",
            ]
        )
    add_table(
        document,
        [
            "Dataset",
            "Family",
            "Backend",
            "Baseline RSS f64/f32",
            "Peak RSS f64/f32",
            "Incremental RSS f64/f32",
            "Incremental delta",
            "Sampled GPU f64/f32",
            "GPU workspace f64/f32",
        ],
        memory_rows,
        font_size=4.7,
    )
    document.add_paragraph(
        (
            "Table S61b. Paired process-memory measurements (MB). Incremental "
            "RSS is absolute peak minus the immediately pre-fit baseline. "
            "Sampled GPU values are process-level device use and include CUDA "
            "context/runtime state. GPU workspace values are NR because a "
            "paired pre-fit GPU baseline was not archived."
        ),
        style="Caption",
    )

    agreement_rows = []
    for _, row in agreement.iterrows():
        agreement_rows.append(
            [
                row["dataset"],
                row["method"],
                row["backend"],
                row["status"],
                fmt(row["float64_metric"], 7),
                fmt(row["float32_metric"], 7),
                fmt(row["metric_delta_float32_minus_float64"], 7),
                fmt(row["prediction_agreement"], 10),
                fmt(row["relative_prediction_error"], 8),
            ]
        )
    add_table(
        document,
        [
            "Controlled task",
            "Family",
            "Backend",
            "Status",
            "f64 metric",
            "f32 metric",
            "Delta",
            "Prediction agreement",
            "Relative prediction error",
        ],
        agreement_rows,
        font_size=4.7,
    )
    document.add_paragraph(
        (
            "Table S62. Controlled float64/float32 prediction agreement. "
            "Classification agreement is decoded-label agreement; regression "
            "agreement is prediction correlation. These small tests verify "
            "route mechanics but do not replace the real-data route warnings "
            "in Table S60."
        ),
        style="Caption",
    )

    document.core_properties.title = (
        "fastPLS CMPB supplement - route-level float32 capability"
    )
    document.save(SUPP_OUTPUT)


def copy_evidence():
    EVIDENCE.mkdir(parents=True, exist_ok=True)
    for path in (CAPABILITY, PAIRED, AGREEMENT):
        shutil.copy2(path, EVIDENCE / path.name)
    with (EVIDENCE / "README.md").open("w", encoding="utf-8") as handle:
        handle.write(
            "# Float32 route-level capability revision\n\n"
            "The capability matrix is authoritative. The paired resource table "
            "reports baseline and incremental host RSS separately. Sampled GPU "
            "memory includes context/runtime state; baseline-corrected GPU "
            "workspace memory was not archived and is therefore `NR`.\n"
        )


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    copy_evidence()
    revise_main()
    revise_supplement()
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)


if __name__ == "__main__":
    main()
