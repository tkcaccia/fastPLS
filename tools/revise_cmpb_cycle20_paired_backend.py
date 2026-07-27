#!/usr/bin/env python3

from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle19"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle20"
EVIDENCE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle20_20260725"
)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle19_0.99.6_20260725.docx"
SUPP_SOURCE = (
    SOURCE / "fastPLS_CMPB_supplement_cycle19_0.99.6_20260725.docx"
)
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle19_20260725.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle20_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle20_0.99.6_20260725.docx"
RESPONSE_OUT = (
    OUT / "fastPLS_CMPB_response_to_reviewers_cycle20_20260725.docx"
)

PAIRED_CSV = EVIDENCE / "paired_backend_selected_summary.csv"
WIDE_CSV = EVIDENCE / "paired_backend_selected_wide.csv"
FIGURE = EVIDENCE / "plots" / "paired_backend_performance_all_datasets.png"

spec = spec_from_file_location(
    "cycle19_helpers",
    ROOT / "tools" / "revise_cmpb_cycle19_predictive_uncertainty.py",
)
c19 = module_from_spec(spec)
spec.loader.exec_module(c19)
c16 = c19.c16


DATASET_LABELS = {
    "MetRef": "metref",
    "CCLE": "ccle",
    "TCGA-BRCA": "tcga_brca",
    "TCGA-HNSC methyl.": "tcga_hnsc_methylation",
    "TCGA-HNSC methylation": "tcga_hnsc_methylation",
    "GTEx v8": "gtex_v8",
    "TCGA Pan-Cancer": "tcga_pan_cancer",
    "Retina": "retina",
    "Tabula Muris": "tabula",
    "CIFAR-100": "cifar100",
    "CBMC CITE-seq": "cbmc_citeseq",
    "PRISM": "prism",
    "NMR": "nmr",
}
DISPLAY_DATASET = {
    "metref": "MetRef",
    "ccle": "CCLE",
    "tcga_brca": "TCGA-BRCA",
    "tcga_hnsc_methylation": "TCGA-HNSC methylation",
    "gtex_v8": "GTEx v8",
    "tcga_pan_cancer": "TCGA Pan-Cancer",
    "retina": "Retina",
    "tabula": "Tabula Muris",
    "cifar100": "CIFAR-100",
    "cbmc_citeseq": "CBMC CITE-seq",
    "prism": "PRISM",
    "nmr": "NMR",
}
METHOD_COLUMNS = {
    1: "plssvd",
    2: "simpls",
    3: "opls",
    4: "kernelpls",
}
METHOD_LABELS = {
    "plssvd": "PLS-SVD",
    "simpls": "SIMPLS",
    "opls": "OPLS",
    "kernelpls": "kernel PLS",
}


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def format_metric(row):
    point = float(row["point_estimate"])
    lower = float(row["ci_lower"])
    upper = float(row["ci_upper"])
    if row["metric_name"] == "accuracy":
        return f"Acc={point:.3f} [{lower:.3f},{upper:.3f}]"
    if point >= 100:
        return f"RMSD={point:.1f} [{lower:.0f},{upper:.0f}]"
    if point >= 0.01:
        return f"RMSD={point:.3f} [{lower:.3f},{upper:.3f}]"
    return f"RMSD={point:.2e} [{lower:.2e},{upper:.2e}]"


def format_backend_row(row):
    if row["status"] != "ok":
        return f'{row["engine"]}: {row["status"]}'
    iqr = float(row["total_time_sec_q75"]) - float(
        row["total_time_sec_q25"]
    )
    gpu = (
        f'{float(row["gpu_mem_mb_median"]):.0f}'
        if pd.notna(row["gpu_mem_mb_median"])
        else "—"
    )
    return (
        f'{row["engine"]} {format_metric(row)}; '
        f't={float(row["total_time_sec_median"]):.3f}s '
        f'(IQR {iqr:.3f}); '
        f'H/G={float(row["host_rss_mb_median"]):.0f}/{gpu}; '
        f'n={int(row["n_runs"])}; OK'
    )


def paired_cell(rows):
    ok = rows[rows["status"] == "ok"]
    if ok.empty:
        return "CPU/CUDA: not evaluated in NMR protocol"
    k = int(ok["effective_ncomp"].dropna().iloc[0])
    selection = str(ok["selection_status"].iloc[0])
    marker = "" if selection == "interior tested value" else "†"
    lines = [f"k={k}{marker}"]
    for engine in ("CPU", "CUDA"):
        hit = rows[rows["engine"] == engine]
        if hit.empty:
            lines.append(f"{engine}: missing")
        else:
            lines.append(format_backend_row(hit.iloc[0]))
    return "\n".join(lines)


def update_main_table(table, paired):
    for row in table.rows[1:]:
        dataset = DATASET_LABELS.get(row.cells[0].text.strip())
        if dataset is None:
            continue
        for column, method in METHOD_COLUMNS.items():
            rows = paired[
                paired["dataset"].eq(dataset)
                & paired["method_panel"].eq(method)
            ]
            row.cells[column].text = paired_cell(rows)
    c16.style_table(table, font_size=4.15)
    c16.prevent_row_splitting(table)
    repeat_header(table.rows[0])


def revise_main(paired):
    document = Document(MAIN_SOURCE)

    design = c16.find_paragraph(document, "Within each dataset")
    design.text = design.text.replace(
        "CPU/CUDA rSVD comparisons within a PLS family used the same "
        "preprocessing, response, component count, and argmax or regression "
        "prediction rule and are therefore estimator matched.",
        "CPU and CUDA rSVD rows within a PLS family used the same "
        "preprocessing, response, component count, and argmax or regression "
        "prediction rule and are therefore estimator matched. Both backends "
        "were retained in the primary summary, including their execution "
        "status and separate host/device memory measurements; no fastest-row "
        "filter was applied.",
    )

    results = c16.find_paragraph(document, "Table 1 and Figure 2")
    results.text = (
        "Table 1 and Figure 2 show both matched CPU and CUDA backends for the "
        "complete twelve-task biomedical benchmark at the best "
        "training-validation component value within each prespecified grid and "
        "PLS family. A dagger identifies a lower or upper tested-grid boundary "
        "or a PLS-SVD response-rank limit; these values are not claimed as "
        "global optima. All 46 evaluated CPU/CUDA pairs completed, while NMR "
        "OPLS and kernel PLS are marked as not evaluated rather than omitted. "
        "The predictive intervals overlapped for every pair; 26 of 36 "
        "classification pairs had identical observed accuracy, and the largest "
        "difference was 2.27 percentage points for TCGA-BRCA OPLS. Backend "
        "choice affected computation more strongly. CPU was faster in 31 of 46 "
        "pairs, primarily on small or modest matrix problems, whereas CUDA was "
        "faster in 15 pairs. The largest CUDA advantages occurred on CIFAR-100 "
        "(up to 14.55-fold); on NMR, CUDA accelerated selected PLS-SVD and "
        "SIMPLS by 2.72- and 5.39-fold, respectively. CUDA peak host RSS ranged "
        "from 0.98 to 1.82 times the CPU value and required 138-3,414 MB of "
        "sampled GPU memory. Thus the paired display exposes when acceleration "
        "is offset by launch, transfer, or memory overhead."
    )

    table_caption = c16.find_paragraph(document, "Table 1.")
    table_caption.text = (
        "Table 1. Paired CPU/CUDA biomedical workflow benchmark at the best "
        "training-validation value within each evaluated component grid. Each "
        "cell retains both matched backends at the same family-specific "
        "component count and reports the outer-test metric with 95% interval, "
        "median total fitting-plus-prediction time with run IQR, peak host/GPU "
        "memory in MB, completed runs, and execution status. †Lower or upper "
        "tested-grid boundary, or response-rank limit for PLS-SVD; these "
        "entries are not global optima. Accuracy uses Wilson intervals and "
        "RMSD uses 10,000-resample held-out-sample bootstrap intervals, both "
        "conditional on the fixed outer split. NMR OPLS and kernel PLS are "
        "explicitly labelled not evaluated."
    )

    figure_caption = c16.find_paragraph(document, "Figure 2.")
    figure_caption.text = (
        "Figure 2. Matched CPU and CUDA outer-test performance at the best "
        "training-validation value within each evaluated grid for the twelve "
        "biomedical benchmark tasks. Circles and triangles denote CPU and CUDA, "
        "respectively; segments join rows with identical data, model family, "
        "component count, and prediction rule. Error bars are 95% Wilson score "
        "intervals for accuracy and 10,000-resample held-out-sample bootstrap "
        "intervals for multivariate RMSD. NE denotes a workflow not evaluated "
        "under the NMR protocol. Exact backend-specific runtime and memory "
        "values appear in Table 1. ImageNet remains in the separate matched "
        "retrieval protocol in Figure 4."
    )
    c16.replace_picture_before_caption(
        document,
        "Figure 2.",
        FIGURE,
        Inches(7.15),
    )
    update_main_table(document.tables[0], paired)

    document.core_properties.title = (
        "fastPLS CMPB manuscript - paired backend analysis cycle 20"
    )
    document.save(MAIN_OUT)


def compact_metric(value, dataset):
    if pd.isna(value):
        return "—"
    if dataset in ("cbmc_citeseq",):
        return f"{value:.1f}"
    if dataset in ("nmr",):
        return f"{value:.2e}"
    return f"{value:.4f}"


def supplement_rows(wide):
    rows = []
    for _, row in wide.iterrows():
        dataset = str(row["dataset"])
        status = "CPU/CUDA OK"
        rows.append(
            (
                DISPLAY_DATASET[dataset],
                METHOD_LABELS[str(row["method_panel"])],
                int(row["effective_ncomp"]),
                (
                    f'{compact_metric(row["point_estimate.CPU"], dataset)} / '
                    f'{compact_metric(row["point_estimate.CUDA"], dataset)}'
                ),
                (
                    f'{row["total_time_sec_median.CPU"]:.3f} / '
                    f'{row["total_time_sec_median.CUDA"]:.3f}'
                ),
                f'{row["time_speedup_cpu_over_cuda"]:.2f}',
                (
                    f'{row["host_rss_mb_median.CPU"]:.0f} / '
                    f'{row["host_rss_mb_median.CUDA"]:.0f}'
                ),
                f'{row["gpu_mem_mb_median.CUDA"]:.0f}',
                status,
            )
        )
    for method in ("opls", "kernelpls"):
        rows.append(
            (
                "NMR",
                METHOD_LABELS[method],
                "—",
                "—",
                "—",
                "—",
                "—",
                "—",
                "not evaluated",
            )
        )
    return rows


def revise_supplement(wide):
    document = Document(SUPP_SOURCE)
    heading = document.add_heading(
        "S27. Matched CPU/CUDA selected-setting audit",
        level=1,
    )
    heading.paragraph_format.page_break_before = True
    document.add_paragraph(
        "The primary benchmark no longer filters each family to its fastest "
        "completed backend. CPU and CUDA were compared at identical outer split, "
        "preprocessing, response, family-specific training-selected component "
        "count, and prediction rule. All 46 evaluated pairs completed. NMR OPLS "
        "and kernel PLS were outside the prespecified NMR protocol and are "
        "retained as not evaluated. The CPU/CUDA predictive intervals overlapped "
        "for every completed pair. CPU was faster in 31 pairs and CUDA in 15. "
        "A speedup above one in Table S32 indicates that CUDA was faster; a "
        "value below one indicates that CPU was faster. CUDA used 138-3,414 MB "
        "of sampled device memory and generally increased peak host RSS on the "
        "small and medium tasks. The machine-readable file "
        "paired_backend_selected_summary.csv retains backend-specific quartiles, "
        "confidence intervals, run counts, and statuses."
    )
    caption = document.add_paragraph(
        "Table S32. Matched CPU/CUDA results at each family-specific "
        "training-selected component count. Metric, time, and host memory are "
        "shown as CPU / CUDA. Speedup is CPU time divided by CUDA time; values "
        "above one favor CUDA. GPU memory is the sampled CUDA peak in MB."
    )
    caption.style = "Caption"
    table = c16.add_table(
        document,
        [
            "Dataset",
            "Family",
            "k",
            "Metric\nCPU/CUDA",
            "Total s\nCPU/CUDA",
            "Speedup",
            "Host MB\nCPU/CUDA",
            "GPU MB",
            "Status",
        ],
        supplement_rows(wide),
        font_size=4.9,
    )
    c16.prevent_row_splitting(table)
    repeat_header(table.rows[0])

    document.core_properties.title = (
        "fastPLS CMPB supplement - paired backend analysis cycle 20"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    document.add_heading(
        "21. The fastest-backend summary concealed paired information",
        level=1,
    )
    document.add_paragraph(
        "Reviewer comment: Table 1 retained only the fastest CPU or CUDA row "
        "for each family, which could conceal backend disagreement, failures, "
        "and memory trade-offs. Both matched backend results should be visible "
        "in the main analysis or an immediately interpretable paired figure."
    )
    document.add_paragraph(
        "Response: Corrected. We removed the fastest-row filter from the primary "
        "presentation. Table 1 now shows CPU and CUDA together in every cell at "
        "the identical family-specific training-selected component count, with "
        "backend-specific predictive metric and 95% interval, total time, host "
        "RSS, GPU memory, run count, and status. Figure 2 is now a paired "
        "CPU/CUDA figure: circles and triangles are joined only when data, "
        "preprocessing, family, component count, and prediction rule match. All "
        "46 evaluated pairs completed; NMR OPLS and kernel PLS are visibly "
        "labelled not evaluated. Every paired predictive interval overlapped, "
        "but the computational trade-offs were substantial: CPU was faster in "
        "31 pairs, CUDA in 15, the maximum CUDA speedup was 14.55-fold, and "
        "CUDA used 138-3,414 MB sampled device memory. We also reran the CPU NMR "
        "counterparts at the revised selected settings (PLS-SVD k=5 and SIMPLS "
        "k=50) with three isolated repetitions. Supplementary Section S27 and "
        "Table S32 provide the compact paired audit, and the complete "
        "machine-readable table retains quartiles and statuses."
    )
    document.core_properties.title = (
        "fastPLS CMPB response - paired backend analysis cycle 20"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    paired = pd.read_csv(PAIRED_CSV)
    wide = pd.read_csv(WIDE_CSV)
    revise_main(paired)
    revise_supplement(wide)
    revise_response()
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
