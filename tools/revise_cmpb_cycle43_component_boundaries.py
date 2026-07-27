#!/usr/bin/env python3

import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle42"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle43"
SUMMARY = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle43_20260726"
    / "component_boundary_summary.csv"
)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle42_0.99.6_20260726.docx"
SUPP_SOURCE = (
    SOURCE / "fastPLS_CMPB_supplement_cycle42_0.99.6_20260726.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle43_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle43_0.99.6_20260726.docx"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_text_in_paragraph(paragraph, old, new):
    if old not in paragraph.text:
        raise RuntimeError(f"Text not found: {old[:100]}")
    text = paragraph.text.replace(old, new)
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    if style is not None:
        inserted.style = style
    if text:
        inserted.add_run(text)
    return inserted


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def prevent_row_splitting(table):
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_width(cell, width_inches):
    width = int(width_inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def add_table_after(paragraph, headers, rows, widths):
    container = paragraph._parent
    table = container.add_table(
        rows=1, cols=len(headers), width=Inches(sum(widths))
    )
    table.autofit = False
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = str(value)
    paragraph._p.addnext(table._tbl)
    repeat_header(table.rows[0])
    prevent_row_splitting(table)
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for item in cell.paragraphs:
                item.paragraph_format.space_before = Pt(0)
                item.paragraph_format.space_after = Pt(0)
                item.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if column_index == 0
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in item.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(6.0)
                    if row_index == 0:
                        run.bold = True
    return table


def summary_rows():
    labels = {
        "plssvd": "PLS-SVD",
        "simpls": "SIMPLS",
        "opls": "OPLS",
        "kernelpls": "Kernel PLS",
        "all": "All families",
    }
    rows = []
    with SUMMARY.open(newline="") as handle:
        for row in csv.DictReader(handle):
            rows.append(
                (
                    labels[row["family"]],
                    row["evaluated"],
                    row["interior"],
                    row["lower_boundary"],
                    row["upper_boundary"],
                    row["response_rank_limit"],
                    row["not_evaluated"],
                )
            )
    return rows


def revise_main():
    document = Document(MAIN_SOURCE)

    methods = find_paragraph(
        document, "Within each dataset, methods used identical fixed outer splits."
    )
    replace_text_in_paragraph(
        methods,
        (
            "The selected model was then evaluated once on the untouched outer "
            "test partition."
        ),
        (
            "The selected model was then evaluated once on the untouched outer "
            "test partition. Selection is reported as the best value within the "
            "prespecified grid. A lower-bound value is left-censored, an "
            "upper-bound value is right-censored, and a PLS-SVD value at the "
            "response-rank ceiling is structurally constrained; none is interpreted "
            "as a global optimum."
        ),
    )

    results = find_paragraph(
        document, "Table 1 and Figure 2 show both matched CPU and CUDA backends"
    )
    replace_text_in_paragraph(
        results,
        (
            "A dagger identifies a lower or upper tested-grid boundary or a "
            "PLS-SVD response-rank limit; these values are not claimed as global "
            "optima."
        ),
        (
            "A dagger identifies a lower or upper tested-grid boundary or a "
            "PLS-SVD response-rank limit; these values are not claimed as global "
            "optima. Counting each dataset-family selection once, 24 of 46 "
            "(52.2%) occurred at a tested-grid boundary (20 upper and four lower), "
            "nine (19.6%) PLS-SVD selections reached the response-rank ceiling, "
            "and only 13 (28.3%) were interior. Thus 33 of 46 selected settings "
            "were boundary- or rank-constrained."
        ),
    )

    discussion_anchor = find_paragraph(
        document, "The direct family-speed experiment also bounds the SIMPLS claim."
    )
    insert_after(
        discussion_anchor,
        (
            "Component selection is an additional limitation. More than half of "
            "the dataset-family choices occurred at an evaluated-grid endpoint, "
            "and a further nine PLS-SVD choices were capped by response rank. "
            "Accordingly, Table 1 compares reproducible workflows at prespecified "
            "training-selected operating points; it does not establish globally "
            "optimal latent dimensionality. A wider or adaptive grid could change "
            "the selected component count and absolute runtime for endpoint cases, "
            "although the paired CPU/CUDA contrasts at a fixed A remain valid. "
            "Future predictive studies should extend endpoint grids or use nested "
            "cross-validation when component selection is a primary inferential "
            "objective."
        ),
        discussion_anchor.style,
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - quantified component-grid boundaries"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)

    conventions = find_paragraph(
        document, "Two uses of component grids are distinguished."
    )
    replace_text_in_paragraph(
        conventions,
        (
            "Training-only model selection reports the best value within the "
            "evaluated grid."
        ),
        (
            "Training-only model selection reports the best value within the "
            "evaluated grid. Endpoint selections are treated as censored with "
            "respect to the unobserved continuation of the prediction curve."
        ),
    )

    boundary = find_paragraph(
        document,
        "Component counts are reported as training-selected values within "
        "prespecified grids.",
    )
    boundary_text = (
        "Component counts are reported as training-selected values within "
        "prespecified grids. Across the 46 evaluated dataset-family selections, "
        "20 were at an upper tested-grid boundary, four were at a lower boundary, "
        "nine PLS-SVD values were constrained by response rank, and 13 were "
        "interior. By family, boundary counts were 8/12 for SIMPLS, 8/11 for "
        "OPLS, and 8/11 for kernel PLS; PLS-SVD had 9/12 response-rank-limited "
        "values and three interior selections. Endpoint and rank-limited values "
        "are retained because they are the reproducible operating points selected "
        "from the prespecified grids, but they do not establish global optima. "
        "NMR PLS-SVD selected an interior value of 5 components; NMR SIMPLS "
        "selected the smallest value in its one-standard-error set, 50. With only "
        "five paired splits, the error differences among 50, 75, and 100 "
        "components were not resolved."
    )
    style = boundary.style
    boundary.clear()
    boundary.style = style
    boundary.add_run(boundary_text)

    table29_caption = find_paragraph(
        document,
        "Table S29. Component-selection status for every dataset and PLS family",
    )
    replace_text_in_paragraph(
        table29_caption,
        (
            "Table S29. Component-selection status for every dataset and PLS "
            "family shown in Table 1."
        ),
        (
            "Table S29. Component-selection status for every dataset and PLS "
            "family shown in Table 1. Lower and upper endpoints are left- and "
            "right-censored, respectively; response-rank limits are structural "
            "PLS-SVD ceilings."
        ),
    )

    s26 = find_paragraph(
        document, "S26. Predictive uncertainty on the prespecified outer test sets"
    )
    caption = s26.insert_paragraph_before(
        (
            "Table S38. Frequency of component-selection status by PLS family. "
            "Each dataset-family choice is counted once; CPU and CUDA rows sharing "
            "the same selected component count are not duplicated. NE denotes not "
            "evaluated."
        )
    )
    caption.style = table29_caption.style
    add_table_after(
        caption,
        (
            "Family",
            "Evaluated",
            "Interior",
            "Lower",
            "Upper",
            "Rank limit",
            "NE",
        ),
        summary_rows(),
        (1.25, 0.75, 0.75, 0.70, 0.70, 0.85, 0.55),
    )

    document.core_properties.title = (
        "fastPLS CMPB supplement - quantified component-grid boundaries"
    )
    document.save(SUPP_OUT)


def main():
    if not SUMMARY.exists():
        raise FileNotFoundError(SUMMARY)
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
