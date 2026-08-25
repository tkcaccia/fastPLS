from pathlib import Path
from copy import deepcopy

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle95"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle96"
OUTPUT.mkdir(parents=True, exist_ok=True)


def replace_paragraph(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def find_paragraph(document, phrase):
    matches = [p for p in document.paragraphs if phrase in p.text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph containing {phrase!r}; found {len(matches)}")
    return matches[0]


def insert_after(paragraph, text, style="Body Text"):
    new = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addnext(new._p)
    return new


main_source = SOURCE / "fastPLS_CMPB_main_cycle95_0.99.25_20260825.docx"
main_output = OUTPUT / "fastPLS_CMPB_main_cycle96_0.99.25_20260825.docx"
main = Document(main_source)

api = find_paragraph(main, "The public pls() interface selects PLS family")
replace_paragraph(
    api,
    "The public pls() interface selects PLS family, component count, solver, backend, and, for "
    "classification, argmax or latent-space linear discriminant analysis (LDA). pls.single.cv() "
    "selects settings within one cross-validation layer; pls.double.cv() uses nested cross-validation. "
    "The audited 0.99.25 namespace explicitly exports pls(), pls.single.cv(), pls.double.cv(), "
    "evaluate(), plot.permutation(), ViP(), fastsvd(), fastcor(), fastPLS_backend(), has_cuda(), and "
    "has_metal(); registered predict() and plot() methods use the standard R generics. The supported "
    "model families are PLS-SVD, SIMPLS, OPLS, and kernel PLS; the public classifiers are argmax and "
    "LDA; and rSVD is the default solver, whereas IRLBA is CPU-only. PCA and nearest-neighbour "
    "classifiers are not package APIs. The legacy lda_ridge argument is deprecated and ignored, emits "
    "a warning when supplied, and is excluded from cross-validation tuning because LDA uses a fixed "
    "scale-normalized Cholesky fallback sequence. Requested estimators and unavailable backends are "
    "never silently substituted. Route status and host/device residency are defined in Supplementary "
    "Tables S1 and S9."
)
main.save(main_output)


supp_source = SOURCE / "fastPLS_CMPB_supplement_cycle95_0.99.25_20260825.docx"
supp_output = OUTPUT / "fastPLS_CMPB_supplement_cycle96_0.99.25_20260825.docx"
supp = Document(supp_source)

scope = find_paragraph(supp, "The double-precision CPU backend is the broadest reference implementation")
insert_after(
    scope,
    "The 0.99.25 namespace explicitly exports pls(), pls.single.cv(), pls.double.cv(), evaluate(), "
    "plot.permutation(), ViP(), fastsvd(), fastcor(), fastPLS_backend(), has_cuda(), and has_metal(). "
    "Registered predict() and plot() methods use the standard R generics. No PCA function or PCA class "
    "method is exported. Public model families are PLS-SVD, SIMPLS, OPLS, and kernel PLS; public "
    "classification heads are argmax and LDA. rSVD is the default approximate solver with qualified "
    "CPU/CUDA controls, while IRLBA is available only on CPU. The legacy lda_ridge argument is "
    "deprecated and ignored: supplying it warns, and it is absent from cross-validation tuning records. "
    "LDA instead uses the fixed relative-ridge sequence 10^-8, 10^-6, 10^-5, 10^-4, 10^-3, and 10^-2, "
    "advancing only after Cholesky failure."
)

residency = supp.tables[0]
for values in (
    (
        "Solver dispatch",
        "IRLBA deterministic; rSVD qualified approximate",
        "rSVD qualified approximate; IRLBA unavailable",
        "rSVD experimental/hybrid; IRLBA unavailable",
        "Unavailable solver/backend combinations stop explicitly",
    ),
    (
        "Classification heads",
        "Argmax and compiled LDA",
        "Argmax and CUDA LDA",
        "Argmax; Metal projection + CPU LDA",
        "Metal LDA is hybrid; lda_ridge is deprecated and ignored",
    ),
):
    cloned_row = deepcopy(residency.rows[-1]._tr)
    residency._tbl.append(cloned_row)
    cells = residency.rows[-1].cells
    for cell, value in zip(cells, values):
        paragraph = cell.paragraphs[0]
        replace_paragraph(paragraph, value)
        for extra in list(cell.paragraphs[1:]):
            cell._tc.remove(extra._p)

supp.save(supp_output)

print(main_output)
print(supp_output)
