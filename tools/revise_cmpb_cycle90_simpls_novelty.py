from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
IN_DIR = ROOT / "artifacts/CMPB_rewrite_20260824_cycle89"
OUT_DIR = ROOT / "artifacts/CMPB_rewrite_20260824_cycle90"
MAIN_IN = IN_DIR / "fastPLS_CMPB_main_cycle89_0.99.24_20260824.docx"
SUPP_IN = IN_DIR / "fastPLS_CMPB_supplement_cycle89_0.99.24_20260824.docx"
MAIN_OUT = OUT_DIR / "fastPLS_CMPB_main_cycle90_0.99.25_20260824.docx"
SUPP_OUT = OUT_DIR / "fastPLS_CMPB_supplement_cycle90_0.99.25_20260824.docx"


def replace_exact(document, old, new):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == old]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph, found {len(matches)}: {old[:100]}")
    paragraph = matches[0]
    style = paragraph.style
    paragraph.text = new
    paragraph.style = style


def set_cell(cell, text, size=5):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)


OUT_DIR.mkdir(parents=True, exist_ok=True)

main = Document(MAIN_IN)
replace_exact(
    main,
    "Methods: The implementation reuses deflation products, latent quantities, coefficients, and predictions along one maximal component path, with compact prediction and optional implicit cross-covariance products. Deterministic float64 CPU IRLBA supported estimator-matched validation against de Jong SIMPLS and independent R software. Approximate rSVD and accelerator routes were audited separately under prespecified criteria detailed in the Supplement.",
    "Methods: The implementation uses compiled, shape-dependent execution: cached deflation products and cross-products, incremental coefficient and fitted-value updates, compact latent prediction, and optional implicit cross-covariance products. The standard sequential component path is also returned by reference SIMPLS software and is not claimed as novel. Deterministic float64 CPU IRLBA supported estimator-matched validation against de Jong SIMPLS and independent R software. Approximate rSVD and accelerator routes were audited separately under prespecified criteria detailed in the Supplement.",
)
replace_exact(
    main,
    "Among PLS formulations, SIMPLS constructs sequential components without explicitly deflating the predictor matrix [11]. PLS-SVD provides a one-shot cross-covariance comparator [10], while OPLS and kernel PLS extend the framework to orthogonal filtering and nonlinear relations [12-14]. Existing R implementations do not jointly provide an optimized SIMPLS component path, memory-aware prediction, and compiled CPU and accelerator execution.",
    "Among PLS formulations, SIMPLS constructs sequential components without explicitly deflating the predictor matrix [11]. PLS-SVD provides a one-shot cross-covariance comparator [10], while OPLS and kernel PLS extend the framework to orthogonal filtering and nonlinear relations [12-14]. Reference software such as pls::simpls.fit already returns coefficient and fitted-value arrays for every component prefix. The unmet computational need is therefore not path generation itself, but shape-aware control of intermediate products, dense output storage, prediction, and accelerator execution.",
)
replace_exact(
    main,
    "We present fastPLS, whose principal methodological contribution is an accelerated execution of de Jong SIMPLS. Sequential deflation, orthogonalization, and component definitions are retained, while previously computed products, coefficients, and predictions are reused through the component path. Low-rank solvers, implicit products, float32, cross-validation, LDA, and CPU/CUDA/Metal execution are supporting options around this core estimator. OPLS and kernel PLS reuse the same optimized SIMPLS engine. The GPL-3 R package uses reusable components from the MIT-licensed kodama-cpp project.",
    "We present fastPLS, whose principal methodological contribution is a compiled, shape-dependent execution of de Jong SIMPLS. Sequential deflation, orthogonalization, and component definitions are retained, while avoidable intermediate products and dense outputs are reduced according to matrix shape and requested output. Low-rank solvers, implicit products, float32, cross-validation, LDA, and CPU/CUDA/Metal execution are supporting options around this core estimator. OPLS and kernel PLS reuse the same optimized SIMPLS engine. The GPL-3 R package uses reusable components from the MIT-licensed kodama-cpp project.",
)
replace_exact(main, "2.2 Accelerated SIMPLS and related PLS models", "2.2 Shape-aware SIMPLS execution and related PLS models")
replace_exact(
    main,
    "SIMPLS follows de Jong's sequential score, loading, orthogonalization, and rank-one deflation equations [11]. The innovation is computational: fastPLS retains deflation products, latent quantities, coefficients, and predictions incrementally, so all requested component counts are snapshots of one maximal path rather than independent refits. Every component invokes a fresh rank-one direction calculation on the current deflated cross-covariance: IRLBA starts a new iterative solve, whereas rSVD draws a new oversampled sketch using the base seed plus the zero-based component index. Candidate-direction blocks, cross-component warm starts, and adaptive refresh were rejected during development and are not used by the public CPU, CUDA, or Metal paths. Retained optimizations cache rank-one deflation products, update coefficient and fitted-value paths incrementally, cache cross-products when shape-appropriate, use compact prediction factors, and support implicit cross-covariance products.",
    "SIMPLS follows de Jong's sequential score, loading, orthogonalization, and rank-one deflation equations [11]. As in pls::simpls.fit, one fit supplies the standard sequential path through all components up to the requested maximum; fastPLS does not claim this path construction as novel. Its contribution is computational and shape dependent: rank-one deflation products are reused; coefficient and fitted-value quantities are updated incrementally; X-transpose-X is cached only when its setup and storage can be amortized; compact prediction retains latent factors instead of dense coefficient and prediction paths; and implicit products avoid an explicit predictor-response cross-covariance when that matrix is limiting. Every component invokes a fresh rank-one direction calculation on the current deflated state. Candidate-direction blocks, cross-component warm starts, and adaptive refresh were rejected and are not used by the public CPU, CUDA, or Metal paths.",
)
replace_exact(
    main,
    "Algorithm 1 summarizes the component path; detailed de Jong mapping is in the Supplement.",
    "Algorithm 1 summarizes the unchanged estimator and the retained execution updates. The minimally optimized compiled baseline, asymptotic costs, storage terms, and detailed de Jong mapping are provided in Supplementary Table S2 and Section S12.1.",
)
replace_exact(
    main,
    "Algorithm 1. Accelerated SIMPLS path. Direction extraction uses deterministic IRLBA or approximate rSVD; score construction, orthogonalization, and deflation follow de Jong [11].",
    "Algorithm 1. Shape-aware SIMPLS execution. Direction extraction uses deterministic IRLBA or approximate rSVD; score construction, orthogonalization, and deflation follow de Jong [11].",
)

algorithm = main.tables[0].cell(0, 0)
algorithm.text = algorithm.text.replace(
    "3.6 Append rₐ, qₐ, and vₐ; update Bₐ = R₁:ₐQ₁:ₐᵀ and Ŷₐ = Ŷₐ₋₁ + tₐqₐᵀ.",
    "3.6 Append rₐ, qₐ, and vₐ; update Bₐ = Bₐ₋₁ + rₐqₐᵀ and Ŷₐ = Ŷₐ₋₁ + tₐqₐᵀ.",
).replace(
    "Output: one maximal component path with requested prefixes, rather than independently refitted models.",
    "Output: the standard sequential component path; requested outputs are retained as dense snapshots or compact latent factors.",
)

replace_exact(
    main,
    "Results are organized around the accelerated SIMPLS contribution. Primary evidence comprises deterministic estimator comparison and the matched IRLBA comparison with independent implementations. rSVD, precision, and hardware results are reported separately with their numerical-audit status; the NMR and ImageNet large-scale analyses are exploratory feasibility studies rather than confirmatory evidence.",
    "Results are organized around the shape-dependent SIMPLS execution contribution. Primary evidence comprises deterministic estimator comparison and the matched IRLBA comparison with independent implementations. rSVD, precision, and hardware results are reported separately with their numerical-audit status; the NMR and ImageNet large-scale analyses are exploratory feasibility studies rather than confirmatory evidence.",
)
replace_exact(main, "3.1 Accelerated SIMPLS compared with independent R implementations", "3.1 Shape-aware SIMPLS compared with independent R implementations")
replace_exact(
    main,
    "Same-code ablations on CIFAR-100, MetRef, PRISM, and Retina produced zero endpoint-metric differences and identical classifications. Compact prediction reduced incremental RSS by up to 77.7% and time by up to 1.24-fold; implicit cross-covariance products reduced RSS by up to 70.6% but were faster only in the high-response PRISM regime. This shape dependence motivates adaptive internal routing rather than universal activation (Supplementary Table S7a).",
    "Same-code ablations on CIFAR-100, MetRef, PRISM, and Retina produced zero endpoint-metric differences and identical classifications. The effects were not uniformly faster: cached deflation products changed time by 0.94-1.02-fold, conditional X-transpose-X caching by 0.99-1.00-fold, incremental coefficients by 0.99-1.07-fold, and compact prediction by 0.97-1.24-fold. Compact prediction reduced incremental RSS by up to 77.7%. Implicit cross-covariance products ranged from a 0.065-fold slowdown to a 6.24-fold gain and reduced RSS by up to 70.6%, with the speed gain confined to the high-response PRISM regime. These results support shape-dependent routing, not universal acceleration (Supplementary Tables S2 and S7a).",
)
replace_exact(
    main,
    "A direct matched-shape timing study further separated estimator choice from implementation cost. On one CPU, the SIMPLS/PLS-SVD time ratio ranged from 1.00 to 3.84; on CUDA it ranged from 0.92 to 0.98 across five synthetic matrix regimes. Thus, the accelerated SIMPLS path approached one-shot PLS-SVD runtime on the tested CUDA shapes, without implying that the two PLS estimators are statistically identical (Supplementary Table S7b).",
    "A direct matched-shape timing study further separated estimator choice from implementation cost. On one CPU, the SIMPLS/PLS-SVD time ratio ranged from 1.00 to 3.84; on CUDA it ranged from 0.92 to 0.98 across five synthetic matrix regimes. Thus, shape-aware SIMPLS approached one-shot PLS-SVD runtime on the tested CUDA shapes, without implying that the two PLS estimators are statistically identical or that SIMPLS is universally faster (Supplementary Table S7b).",
)
replace_exact(
    main,
    "The principal contribution of fastPLS is an accelerated SIMPLS execution path, not a new PLS estimator. Deterministic validation showed that reuse of sequential quantities met the prespecified deterministic numerical tolerances against de Jong SIMPLS, while the external comparison showed lower runtime and, for large tasks, lower memory without changing matched accuracy.",
    "The principal contribution of fastPLS is a shape-dependent SIMPLS execution and storage layer, not a new PLS estimator and not the standard availability of all component prefixes. Deterministic validation showed that the retained execution changes met the prespecified numerical tolerances against de Jong SIMPLS, while the external comparison showed lower runtime and, for large tasks, lower memory without changing matched accuracy.",
)
replace_exact(
    main,
    "rSVD, implicit products, float32, CUDA, and Metal are optional implementation mechanisms around accelerated SIMPLS. Qualified NMR controls met the prespecified approximate-route tolerances, but rSVD remains stochastic and CPU IRLBA remains the deterministic reference. The million-sample ImageNet route demonstrated feasibility with current package code, while its hybrid residency, single run, noncanonical split, and lack of an estimator-matched large-scale control preclude a general accelerator or accuracy claim.",
    "rSVD, implicit products, float32, CUDA, and Metal are optional implementation mechanisms around shape-aware SIMPLS. Qualified NMR controls met the prespecified approximate-route tolerances, but rSVD remains stochastic and CPU IRLBA remains the deterministic reference. The million-sample ImageNet route demonstrated feasibility with current package code, while its hybrid residency, single run, noncanonical split, and lack of an estimator-matched large-scale control preclude a general accelerator or accuracy claim.",
)
replace_exact(
    main,
    "The public pls() interface selects PLS family, component count, solver, backend, and, for classification, argmax or latent-space linear discriminant analysis (LDA). pls.single.cv() selects settings within one cross-validation layer; pls.double.cv() uses nested cross-validation. The audited 0.99.24 interface exposes PLS modelling, prediction, evaluation, cross-validation, score plotting, and standalone truncated SVD; PCA and nearest-neighbour classifiers are not package APIs. Requested estimators are never silently substituted.",
    "The public pls() interface selects PLS family, component count, solver, backend, and, for classification, argmax or latent-space linear discriminant analysis (LDA). pls.single.cv() selects settings within one cross-validation layer; pls.double.cv() uses nested cross-validation. The audited 0.99.25 interface exposes PLS modelling, prediction, evaluation, cross-validation, score plotting, and standalone truncated SVD; PCA and nearest-neighbour classifiers are not package APIs. Requested estimators are never silently substituted.",
)
replace_exact(
    main,
    "Code and benchmark outputs are available at https://github.com/tkcaccia/fastPLS; reusable components are at https://github.com/tkcaccia/kodama-cpp. Historical quantitative results remain tied to fastPLS 0.99.6, base commit 6e50bd318f20289101f6b723953830aefa8b95d6, and source-archive SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85; they are not relabelled as reruns. The reviewed interface is fastPLS 0.99.24. Its exact source archive fastPLS_0.99.24.tar.gz has SHA-256 f27fa7d3c96d6d8a550857bab23146be5b4cefe2868ce17d4bfcc0fe7c87b5db. The definitive CPU/CUDA multi-seed rSVD qualification was generated with 0.99.23 before the documentation and Metal direction-refresh correction; the active CPU/CUDA numerical path and qualified controls are unchanged in 0.99.24. Analysis-specific scripts and archive digests are reported in Supplementary Table S15.",
    "Code and benchmark outputs are available at https://github.com/tkcaccia/fastPLS; reusable components are at https://github.com/tkcaccia/kodama-cpp. Historical quantitative results remain tied to fastPLS 0.99.6, base commit 6e50bd318f20289101f6b723953830aefa8b95d6, and source-archive SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85; they are not relabelled as reruns. The reviewed interface is fastPLS 0.99.25. Its exact source archive fastPLS_0.99.25.tar.gz has SHA-256 74e134ef22d591c3ac5b06910880ea56dac1970b5277cc6c1c25d3dcff3d9e58. The definitive CPU/CUDA multi-seed rSVD qualification was generated with 0.99.23 before the documentation and Metal direction-refresh correction; the active CPU/CUDA numerical path and qualified controls are unchanged in 0.99.25. Analysis-specific scripts and archive digests are reported in Supplementary Table S15.",
)
main.save(MAIN_OUT)

supp = Document(SUPP_IN)
replace_exact(
    supp,
    "Several requested component counts are extracted from one maximal path.",
    "As in pls::simpls.fit, one fit supplies the standard sequential path through all components up to the requested maximum; this is not claimed as a fastPLS novelty.",
)
replace_exact(
    supp,
    "Incremental deflation, coefficient, and prediction updates reorganize numerical work while retaining sequential SIMPLS orthogonalization and deflation. The public CPU, CUDA, and Metal paths all request one newly computed direction from the current deflated operator for every component. The preceding latent direction is never inserted into the next solver start, and candidate blocks are never consumed across deflation steps. Retained optimizations are cached rank-one deflation products, incremental coefficient and fitted-value updates, conditional cross-product caching, compact prediction, implicit cross-covariance products, and reusable allocation workspaces. Cross-component warm starts, multi-direction block refresh, and adaptive refresh policies were rejected and removed from the public execution path.",
    "A minimally optimized compiled baseline forms the cross-covariance explicitly, performs a fresh leading-direction solve and the standard de Jong score/loading/orthogonalization/deflation update for every component, and reconstructs requested dense coefficients and fitted values from retained latent factors. fastPLS changes this execution, not the estimator: it caches the deflation row product; updates coefficients and fitted values incrementally; conditionally caches X-transpose-X; retains compact latent factors when dense outputs are unnecessary; and uses implicit cross-covariance products when predictor-response storage is limiting. The public CPU, CUDA, and Metal paths request one newly computed direction from the current deflated state for every component. Cross-component warm starts, multi-direction block refresh, and adaptive refresh policies were rejected and removed.",
)
replace_exact(
    supp,
    "This supplement distinguishes the reviewed fastPLS 0.99.24 interface from the 0.99.6 source archive used for historical quantitative benchmarks and the 0.99.23 archive used for the definitive CPU/CUDA multi-seed rSVD qualification. Table S15 maps each analysis to its result archive and records an exact package commit only when run metadata captured it. A later package version, result date, or manuscript commit is never treated as evidence of a historical computational SHA, and historical values are not relabelled as later-version reruns.",
    "This supplement distinguishes the reviewed fastPLS 0.99.25 interface from the 0.99.6 source archive used for historical quantitative benchmarks and the 0.99.23 archive used for the definitive CPU/CUDA multi-seed rSVD qualification. Table S15 maps each analysis to its result archive and records an exact package commit only when run metadata captured it. A later package version, result date, or manuscript commit is never treated as evidence of a historical computational SHA, and historical values are not relabelled as later-version reruns.",
)
replace_exact(
    supp,
    "Current benchmark workflows record repository state, benchmark-script checksum, package version, source-archive SHA-256, compiler, BLAS/LAPACK, thread settings, accelerator libraries, seeds, rSVD controls, and data/split identifiers. Table S15 maps each analysis to its exact evidence archive. NMR and ImageNet used fastPLS_0.99.6.tar.gz (SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85) from base commit 6e50bd318f20289101f6b723953830aefa8b95d6; these benchmark values are not relabelled as later-version reruns. The 0.99.23 archive (SHA-256 bd8bd6b0dd85219cd9dbe25b429cea02d0a6633df490618d3a3ad73f177fd5e8) generated the definitive multi-seed CPU/CUDA rSVD qualification. The reviewed 0.99.24 archive (SHA-256 f27fa7d3c96d6d8a550857bab23146be5b4cefe2868ce17d4bfcc0fe7c87b5db) standardizes fresh-per-component direction extraction across CPU, CUDA, and Metal and adds backend-invariance tests and fitted-model rule diagnostics.",
    "Current benchmark workflows record repository state, benchmark-script checksum, package version, source-archive SHA-256, compiler, BLAS/LAPACK, thread settings, accelerator libraries, seeds, rSVD controls, and data/split identifiers. Table S15 maps each analysis to its exact evidence archive. NMR and ImageNet used fastPLS_0.99.6.tar.gz (SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85) from base commit 6e50bd318f20289101f6b723953830aefa8b95d6; these benchmark values are not relabelled as later-version reruns. The 0.99.23 archive (SHA-256 bd8bd6b0dd85219cd9dbe25b429cea02d0a6633df490618d3a3ad73f177fd5e8) generated the definitive multi-seed CPU/CUDA rSVD qualification. The reviewed 0.99.25 archive (SHA-256 74e134ef22d591c3ac5b06910880ea56dac1970b5277cc6c1c25d3dcff3d9e58) standardizes fresh-per-component direction extraction across CPU, CUDA, and Metal, adds backend-invariance diagnostics, and defines the shape-dependent execution contribution without claiming component-path generation as novel.",
)
replace_exact(
    supp,
    "The ledger never infers a historical commit from a package version or result date. Where a run did not record its Git SHA, the source status is explicitly not recoverable. SHA-256 prefixes identify immutable result archives; full digests are retained in the machine-readable ledger. The 0.99.23 solver-control qualification is documented in publication_release_0.99.23/rsvd_qualification and does not alter historical result provenance; the 0.99.24 archive is the reviewed public implementation.",
    "The ledger never infers a historical commit from a package version or result date. Where a run did not record its Git SHA, the source status is explicitly not recoverable. SHA-256 prefixes identify immutable result archives; full digests are retained in the machine-readable ledger. The 0.99.23 solver-control qualification is documented in publication_release_0.99.23/rsvd_qualification and does not alter historical result provenance; the 0.99.25 archive is the reviewed public implementation.",
)

# Table S2: define the baseline and give explicit asymptotic terms.
storage = supp.tables[1]
for row in storage.rows:
    label = row.cells[0].text
    if label == "SIMPLS":
        values = (
            "Minimally optimized compiled SIMPLS",
            "Form S: O(npq); then A fresh direction solves and O(A[np+nq+pq]+pA²) sequential work, excluding solver cost",
            "S and bases: O(pq+(2p+q)A); dense requested prefixes: O(|C|(pq+nq))",
            "Reference execution baseline; the standard path through components 1,...,A is not a fastPLS novelty",
        )
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)

new = storage.add_row().cells
values = (
    "fastPLS shape-aware SIMPLS",
    "Incremental path: O(Apq+Anq) output updates; compact prediction: O(n_test A[p+q]); implicit operator pair of width l: O(n[p+q]l)",
    "Compact factors: O((p+q)A); implicit work arrays: O((n+p+q)l), without O(pq) S; optional XᵀX cache: O(p²)",
    "Selects cached, compact, or implicit execution by shape; trades setup/operator work for lower repeated work or storage",
)
for cell, value in zip(new, values):
    set_cell(cell, value)

ledger = supp.tables[18]
for row in ledger.rows:
    if row.cells[0].text == "A20":
        values = (
            "A20",
            "Public SIMPLS direction rule and shape-aware execution documentation",
            "fastPLS_0.99.25.tar.gz; tests/testthat/test-simpls-direction-refresh.R",
            "0.99.25",
            "exact source archive; CPU, native Metal, and CUDA runtime checks",
            "tests/testthat/test-simpls-direction-refresh.R; benchmark/SIMPLS_DEJONG_MAPPING.md",
            "74e134ef22d5",
        )
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)

replace_exact(
    supp,
    "Each ablation changed one internal execution feature while keeping the data, split, SIMPLS estimator, deterministic IRLBA solver, component count, seed, and prediction head fixed. Three isolated runs were used per configuration. Speed-up is reference time divided by optimized time; positive RSS reduction denotes a lower fit-window incremental peak. Results show that memory reduction and speed are separate and matrix-shape dependent.",
    "Each ablation changed one internal execution feature relative to the minimally optimized compiled SIMPLS baseline while keeping the data, split, estimator, deterministic IRLBA solver, component count, seed, and prediction head fixed. Three isolated runs were used per configuration. Speed-up is baseline time divided by optimized time; positive RSS reduction denotes a lower fit-window incremental peak. Cached deflation removes repeated formation of the deflation row but not its O(pq) order. Incremental coefficients replace O(pq sum(C)) prefix reconstruction by O(Apq); fitted-value updates analogously replace O(nq sum(C)) by O(Anq). Compact prediction replaces dense prefix storage O(|C|pq) by latent storage O((p+q)A). Conditional X-transpose-X caching pays O(np²) setup and O(p²) storage to reduce repeated sample-space work. Implicit products remove O(pq) cross-covariance storage but repeat O(n[p+q]l) operator work per sketch. The measured ranges therefore quantify shape-dependent trade-offs rather than universal speedups.",
)

supp.save(SUPP_OUT)
print(MAIN_OUT)
print(SUPP_OUT)
