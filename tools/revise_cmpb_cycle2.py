"""Create cycle-2 CMPB drafts from the cycle-1 evidence-based documents."""

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Desktop/fastPLS/CMPB_rewrite_20260724_cycle1")
OUT = Path("/Users/stefano/Desktop/fastPLS/CMPB_rewrite_20260724_cycle2")
OUT.mkdir(parents=True, exist_ok=True)
FIGURES = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github/benchmark_results/review_nmr_20260724/plots")

MAIN_SRC = ROOT / "fastPLS_CMPB_main_cycle1_0.99.6_20260724.docx"
SUPP_SRC = ROOT / "fastPLS_CMPB_supplement_cycle1_0.99.6_20260724.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle2_0.99.7_20260724.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle2_0.99.7_20260724.docx"
REVIEW_OUT = OUT / "fastPLS_CMPB_independent_reviewer_report_cycle2_20260724.docx"


def paragraph_after(paragraph, text, style=None):
    node = OxmlElement("w:p")
    paragraph._p.addnext(node)
    inserted = Paragraph(node, paragraph._parent)
    if style:
        inserted.style = style
    inserted.add_run(text)
    return inserted


def replace_paragraph(doc, startswith, replacement):
    for p in doc.paragraphs:
        if p.text.startswith(startswith):
            p.clear()
            p.add_run(replacement)
            return p
    raise RuntimeError(f"No paragraph starts with: {startswith}")


def insert_after_paragraph(doc, startswith, text, style=None):
    for p in doc.paragraphs:
        if p.text.startswith(startswith):
            return paragraph_after(p, text, style=style)
    raise RuntimeError(f"No paragraph starts with: {startswith}")


def add_caption(doc, text):
    p = doc.add_paragraph()
    try:
        p.style = "Caption"
    except KeyError:
        pass
    p.add_run(text)


def revise_main():
    copy2(MAIN_SRC, MAIN_OUT)
    doc = Document(MAIN_OUT)
    replace_paragraph(
        doc,
        "Methods: fastPLS provides",
        "Methods: fastPLS provides PLS-SVD, SIMPLS, orthogonal PLS (OPLS), and kernel PLS through one R interface. "
        "For deterministic IRLBA direction extraction, compiled SIMPLS reproduces de Jong's sequential estimator while avoiding repeated model fits for component prefixes and reusing deflation and prediction quantities. "
        "For rSVD, the same SIMPLS framework uses an explicitly approximate low-rank direction solver, with optional state reuse confined to that approximation. "
        "The package combines bundled CPU IRLBA, randomized SVD (rSVD), implicit cross-covariance products, compact prediction, compiled validation, and CPU, NVIDIA CUDA, and Apple Metal backends. "
        "Selected float32 execution paths are evaluated separately from the double-precision reference. Performance was evaluated using fixed data partitions, total fitting-plus-prediction time, predictive metrics, and peak memory."
    )
    replace_paragraph(
        doc,
        "Results: On controlled synthetic matrices",
        "Results: With deterministic IRLBA, fastPLS SIMPLS agreed with pls::simpls.fit on controlled synthetic matrices to numerical precision and on preprocessed MetRef data across 2-18 components (prediction correlation 1.000; maximum predictive-subspace angle 0.0027 degrees). "
        "rSVD was evaluated separately as an approximate alternative. In multivariate NMR, training-only validation selected 100 components; held-out SIMPLS-rSVD achieved RMSD 0.000861 on CPU and 0.000805 on CUDA, while CUDA reduced model fit plus prediction from 19.55 to 3.09 seconds. "
        "fastPLS also processed a million-sample ImageNet/DINOv2 post-feature-extraction stress test. Under the documented interfaces and resource limits, the evaluated external R workflows did not complete that task."
    )
    replace_paragraph(
        doc,
        "The SIMPLS estimator follows",
        "The SIMPLS estimator follows de Jong's sequential construction [11]. At component k, the dominant direction of the current cross-covariance state is converted to a score and loading pair, orthogonalized against the preceding SIMPLS basis, and used in the standard rank-one deflation. "
        "When IRLBA is requested, every deflated update is computed by the requested iterative solver; no randomized direction is substituted. The rSVD route instead approximates the leading direction and may reuse a randomized workspace across deflations. "
        "In both cases, deflation terms, latent quantities, and prediction prefixes are retained incrementally, avoiding independent refits for each requested component count."
    )
    replace_paragraph(
        doc,
        "Across completed benchmarks, the numerical backend",
        "Across completed benchmarks, the numerical backend had less influence on prediction than the PLS family, component count, or classification head. "
        "The compiled SIMPLS path avoids separate fits for each component prefix. Deterministic IRLBA retains the de Jong estimator at every sequential update; rSVD is a separately labelled approximation whose predictive behaviour is benchmarked rather than assumed to be identical."
    )
    replace_paragraph(
        doc,
        "A controlled numerical check directly compared",
        "A controlled numerical check compared compiled deterministic IRLBA SIMPLS with pls::simpls.fit on fixed well-conditioned, ill-conditioned, and rank-deficient multivariate regression matrices (ntrain=180, p=60, q=4, five components). "
        "Prediction correlation was 1.000, relative prediction and coefficient errors were at most 1.15x10^-12 and 7.62x10^-13, respectively, and the maximum predictive-subspace angle was 2.6x10^-6 degrees. "
        "On identically preprocessed MetRef data, 2, 5, 10, and 18 component fits also had prediction correlation 1.000, label agreement 1.000, and maximum subspace angle 0.0027 degrees. "
        "The corresponding rSVD comparisons are reported as approximate low-rank fits, not as estimator-equivalence evidence."
    )
    nmr_anchor = replace_paragraph(
        doc,
        "NMR represented the extreme multivariate-response setting",
        "NMR represented the extreme multivariate-response setting (1,200 training and 321 held-out spectra; p=13,000; q=28,355). "
        "The predictor water region between 4.6 and 4.8 ppm was set to zero in both training and test predictors before any fit. A fixed 20% inner split of the training spectra selected components from 10, 25, 50, 75, and 100 using validation RMSD, leaving the held-out test spectra untouched. "
        "Validation RMSD decreased from 0.001137 at 10 to 0.000894 at 100 components, which was selected. On the held-out test set, double-precision SIMPLS-rSVD achieved RMSD 0.000861 and R2/Q2 0.9926 on CPU; CUDA achieved RMSD 0.000805 and R2/Q2 0.9935. "
        "Model fitting plus prediction required 19.55 s on CPU and 3.09 s on CUDA. Peak host RSS was 3,144 MB and 3,470 MB, respectively, and the sampled CUDA compute-applications peak was 3,432 MB. Per-spectrum and per-response error distributions, as well as global and zoomed observed-versus-predicted spectra, are provided in the Supplementary Material."
    )
    insert_after_paragraph(
        doc,
        "NMR represented the extreme multivariate-response setting",
        "The matched CPU/CUDA analysis should be interpreted as a single-run, fixed-split computational validation, not as an uncertainty estimate of biological generalization. The representative spectra were selected mechanically as the held-out spectrum whose RMSD was closest to the held-out median; they were not selected by visual inspection."
    )
    replace_paragraph(
        doc,
        "fastPLS extends established PLS algorithms",
        "fastPLS extends established PLS algorithms through implementation rather than a new statistical objective. PLS-SVD remains a direct one-shot approach when the response rank supports the desired path. "
        "SIMPLS is more flexible but sequential: with deterministic IRLBA, fastPLS preserves de Jong's component and deflation definitions, while rSVD trades exact direction extraction for an explicitly benchmarked approximation. "
        "The computational gain derives from compiled execution, incremental component-prefix handling, cached quantities, and memory-aware prediction rather than from relabelling an approximate estimator as exact SIMPLS."
    )
    doc.save(MAIN_OUT)


def revise_supplement():
    copy2(SUPP_SRC, SUPP_OUT)
    doc = Document(SUPP_OUT)
    doc.add_heading("S12. Corrected NMR validation and SIMPLS agreement", level=1)
    doc.add_paragraph(
        "NMR predictor preprocessing and selection. The NMR task comprised 1,200 training and 321 held-out spectra, 13,000 predictors, and 28,355 numeric responses. "
        "The predictor water region from 4.6 to 4.8 ppm was zeroed in both the training and held-out predictor matrices. A fixed seed (123) selected 20% of training spectra as an inner validation set. "
        "SIMPLS-rSVD models with 10, 25, 50, 75, and 100 components were fitted independently. Validation RMSD was 0.001137, 0.000998, 0.000929, 0.000913, and 0.000894, respectively; 100 components were selected."
    )
    doc.add_paragraph(
        "Held-out NMR results. At 100 components, CPU SIMPLS-rSVD achieved R2/Q2=0.9926, RMSD=0.000861, MAE=0.000132, and median per-spectrum RMSD=0.000517. CUDA achieved R2/Q2=0.9935, RMSD=0.000805, MAE=0.000133, and median per-spectrum RMSD=0.000518. "
        "Measured fit-plus-prediction time was 19.55 s on CPU and 3.09 s on CUDA. Maximum process RSS was 3,144 MB and 3,470 MB. CUDA compute-applications memory sampled at 0.2-second intervals peaked at 3,432 MB."
    )
    doc.add_paragraph(
        "Estimator agreement. On fixed synthetic data, deterministic IRLBA SIMPLS matched pls::simpls.fit to numerical precision. On fixed, preprocessed MetRef data, the deterministic IRLBA route had prediction correlation and decoded-label agreement of 1.000 at 2, 5, 10, and 18 components; the maximum principal angle of the predictive subspaces was 0.0027 degrees. rSVD is not used as estimator-equivalence evidence because it is an approximate singular-direction solver."
    )
    for name, caption in [
        ("nmr_spectrum_full.png", "Figure S12. Observed and predicted full held-out NMR spectrum. The spectrum was selected mechanically by median per-spectrum RMSD; the model used 100 components selected from the training-only validation split."),
        ("nmr_spectrum_zoom.png", "Figure S13. Low-ppm zoom of the representative held-out NMR spectrum in Figure S12."),
        ("nmr_per_spectrum_rmsd.png", "Figure S14. Distribution of held-out per-spectrum RMSD for the matched CPU and CUDA SIMPLS-rSVD fits at 100 components."),
        ("nmr_speed_memory.png", "Figure S15. Total fit-plus-prediction time, peak host resident memory, and sampled CUDA compute-applications memory for the matched NMR fits.")
    ]:
        doc.add_picture(str(FIGURES / name), width=Inches(6.2))
        add_caption(doc, caption)
    doc.save(SUPP_OUT)


def create_review():
    doc = Document()
    doc.add_heading("Independent reviewer report: cycle 2 update", level=0)
    doc.add_paragraph("Manuscript: fastPLS: scalable partial least squares with compiled CPU and accelerator backends for high-dimensional biomedical data")
    doc.add_heading("Recommendation", level=1)
    doc.add_paragraph("Major revision")
    doc.add_heading("Assessment after cycle 2", level=1)
    doc.add_paragraph(
        "This revision materially improves the manuscript. It corrects an important ambiguity: deterministic IRLBA SIMPLS is now distinguished from rSVD, which is appropriately presented as an approximate direction solver. "
        "The authors also provide a transparent NMR validation with a training-only component selection procedure, matched fixed-split CPU/CUDA results, per-spectrum error distributions, and mechanistically selected spectrum overlays."
    )
    doc.add_heading("Remaining major comments", level=1)
    for text in [
        "The updated NMR evidence is strong but single-run. At least three isolated repetitions, including distributional timing and memory results, are still needed before precise speedup claims are generalized.",
        "The main multi-dataset benchmark has not yet been regenerated with the corrected IRLBA SIMPLS dispatch. All final comparative tables must identify requested and executed estimator/backend, use a precision-matched float64 primary comparison, and include dispersion and failures.",
        "The manuscript still describes a number of SIMPLS acceleration mechanisms together. Provide a compact ablation showing which mechanisms are active for deterministic IRLBA versus rSVD, and quantify their independent effects on at least one real moderate and one large data set.",
        "The float32 contribution remains numerical/storage compatibility rather than demonstrated efficiency. Broader family-by-backend precision validation is required, or float32 should be substantially narrowed in the title, abstract, and conclusions.",
        "A residency/complexity table is still required to distinguish fully device-resident PLS-SVD/SIMPLS operations from hybrid CUDA OPLS, nonlinear kernel PLS, and Metal stages."
    ]:
        doc.add_paragraph(text, style="List Number")
    doc.add_heading("Minor comments", level=1)
    for text in [
        "Define R2 and Q2 in conventional notation and state whether the reported NMR R2 is fitted or held-out predictive R2.",
        "The NMR spectrum figure should label the response spectrum and explicitly state that the water mask applies to predictors, not the response.",
        "Use an immutable release tag rather than a working commit hash in the final data-and-code statement.",
        "Retain the explicit statement that ImageNet is a non-biomedical computational stress test in the abstract as well as the methods."
    ]:
        doc.add_paragraph(text, style="List Bullet")
    doc.save(REVIEW_OUT)


revise_main()
revise_supplement()
create_review()
print(MAIN_OUT)
print(SUPP_OUT)
print(REVIEW_OUT)
