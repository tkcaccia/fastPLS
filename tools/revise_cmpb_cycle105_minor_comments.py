from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle104"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle105"
OUTPUT.mkdir(parents=True, exist_ok=True)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle104_0.99.25_20260825.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle104_0.99.25_20260825.docx"
MAIN_OUTPUT = OUTPUT / "fastPLS_CMPB_main_cycle105_0.99.25_20260825.docx"
SUPP_OUTPUT = OUTPUT / "fastPLS_CMPB_supplement_cycle105_0.99.25_20260825.docx"

FIGURE_ROOT = ROOT / "benchmark_results" / "frozen_release_0.99.25" / "figures"
MAIN_FIGURES = {
    "Figure 2.": (FIGURE_ROOT / "Figure_2_frozen_external_simpls.png", 6.45),
    "Figure 3.": (FIGURE_ROOT / "Figure_3_frozen_cpu_cuda.png", 6.45),
    "Figure 4.": (FIGURE_ROOT / "Figure_4_frozen_nmr.png", 6.45),
}
SUPP_FIGURES = {
    "Figure S1.": (FIGURE_ROOT / "Figure_S1_frozen_scaling.png", 6.45),
    "Figure S3.": (FIGURE_ROOT / "Figure_5_frozen_imagenet.png", 6.45),
}


def paragraph_by_prefix(document, prefix):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}; found {len(matches)}")
    return matches[0]


def insert_after(document, paragraph, text="", style=None):
    new = document.add_paragraph(style=style or paragraph.style)
    new.text = text
    paragraph._p.addnext(new._p)
    return new


def insert_before(document, paragraph, text="", style=None):
    new = document.add_paragraph(style=style or paragraph.style)
    new.text = text
    paragraph._p.addprevious(new._p)
    return new


def replace_in_paragraphs(document, replacements):
    for paragraph in document.paragraphs:
        if "benchmark_results/frozen_release" in paragraph.text or "provenance/frozen" in paragraph.text:
            continue
        value = paragraph.text
        for old, new in replacements:
            value = value.replace(old, new)
        if value != paragraph.text:
            paragraph.text = value


def replace_everywhere(document, old, new):
    for paragraph in document.paragraphs:
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old in paragraph.text:
                        paragraph.text = paragraph.text.replace(old, new)


def set_cell(cell, text):
    cell.text = str(text)


def replace_table_rows(table, rows):
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    for row, values in zip(table.rows, rows):
        if len(values) != len(row.cells):
            raise RuntimeError("Table column count mismatch")
        for cell, value in zip(row.cells, values):
            set_cell(cell, value)


def replace_preceding_figure(document, caption_prefix, image_path, width):
    if not image_path.is_file():
        raise FileNotFoundError(image_path)
    caption = paragraph_by_prefix(document, caption_prefix)
    previous = caption._p.getprevious()
    while previous is not None and previous.tag.endswith("}p"):
        if previous.xpath(".//w:drawing"):
            previous.getparent().remove(previous)
            break
        if "".join(previous.itertext()).strip():
            break
        previous = previous.getprevious()
    figure = document.add_paragraph()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.keep_with_next = True
    figure.paragraph_format.space_before = 0
    figure.paragraph_format.space_after = 3
    figure.add_run().add_picture(str(image_path), width=Inches(width))
    caption._p.addprevious(figure._p)


def format_code(paragraph):
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(5)
    for run in paragraph.runs:
        run.font.name = "Courier New"
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(35, 35, 35)
    ppr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill", "F2F2F2")
    ppr.append(shading)


def format_tables(document, size=6.4):
    for table in document.tables:
        table.autofit = True
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = 1
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1
                    for run in paragraph.runs:
                        run.font.size = Pt(size)
                        if row_index == 0:
                            run.bold = True


def update_main():
    document = Document(MAIN_SOURCE)

    paragraph_by_prefix(document, "Background and objective:").text = (
        "Background and objective: High-dimensional biomedical partial least-squares (PLS) "
        "workflows, including multivariate NMR prediction and repeated validation, can be limited "
        "by the sequential cost and storage of SIMPLS, de Jong's sequential PLS algorithm. We "
        "developed fastPLS to reduce this burden while retaining the component equations."
    )
    paragraph_by_prefix(document, "Methods:").text = (
        "Methods: Central evidence used archived fastPLS 0.99.25 (Git commit 7887401b09e2). The "
        "implementation combines compiled shape-dependent execution, incremental updates, compact "
        "latent prediction, and optional implicit cross-covariance products. We evaluated SIMPLS "
        "and PLS-SVD, which derives components from a singular value decomposition of the predictor-"
        "response cross-covariance. Classification used maximum-score (argmax) decoding or linear "
        "discriminant analysis. Fixed-control implicitly restarted Lanczos bidiagonalization (IRLBA) "
        "provided the deterministic numerical reference; approximate randomized SVD (rSVD) and "
        "accelerator routes were evaluated separately under predefined criteria."
    )

    paragraph_by_prefix(document, "The public pls() interface").text = (
        "The public model-fitting interface selects the PLS family, component count, low-rank solver, "
        "execution backend, and classification head. Single and nested cross-validation use the same "
        "dispatch contract. Unsupported combinations stop before fitting, requested estimators are not "
        "silently replaced, and fitted objects record solver diagnostics and host/device residency "
        "(Supplementary Tables S1 and S9). Complete API, deprecation, and compatibility details are "
        "provided in the package documentation and Supplementary Section S1."
    )
    paragraph_by_prefix(document, "The architecture separates preprocessing").text = (
        "The architecture separates preprocessing, the PLS estimator, low-rank direction extraction, "
        "and CPU, NVIDIA CUDA, or Apple Metal execution (Figure 1). Benchmark rows record both the "
        "requested and executed estimators, and any mismatch is treated as a failed run."
    )

    paragraph_by_prefix(document, "Direction extraction is modular").text = (
        "Direction extraction is modular rather than the principal estimator contribution. float64 CPU "
        "fitting supports fixed-control IRLBA [15] as the deterministic numerical reference and "
        "approximate rSVD [16], which uses a Gaussian range sketch, orthonormalization, power iterations, "
        "and a reduced decomposition. IRLBA is an iterative truncated solver, not an exact dense SVD. "
        "Archived release 0.99.25 uses rSVD oversampling 20 and two power iterations for the evaluated "
        "CPU/CUDA routes. Explicit unqualified overrides warn and are recorded in model diagnostics. "
        "Controls, seeds, validation thresholds, and route-specific status are consolidated in "
        "Supplementary Sections S13-S14 and Tables S8-S9."
    )

    paragraph_by_prefix(document, "The controlled study completed").text = (
        "The controlled study completed all 486 CPU/CUDA runs. Automatic rSVD met tolerance at all "
        "sample-size, predictor-size, requested-prefix, and cross-covariance-size points, but not "
        "throughout the retained-component, class-count, rank, or CUDA response-dimension sweeps; "
        "73 of 276 automatic-route runs were excluded from validated speed claims. Among routes that "
        "met tolerance, CUDA first exceeded CPU rSVD at n = 5,000 and p = 2,000 and was 3.04-5.97-fold "
        "faster across the 2-64 MiB cross-covariance sweep. CPU implicit products reduced incremental "
        "RSS by 29.7-47.5%, but were slower below 32 MiB and approximately time-neutral at 32-64 MiB. "
        "A 56-run Metal diagnostic completed without execution failure, but all 20 Metal rSVD routes "
        "were numerically discordant and excluded from validated performance claims (Supplementary "
        "Tables S7c-S7d; Figure S1)."
    )
    paragraph_by_prefix(document, "Hardware acceleration remained").text = (
        "Hardware acceleration remained route and shape dependent. Archived-release selected CPU/CUDA "
        "SIMPLS-rSVD workflows were compared on identical splits and component counts (Figure 3; "
        "Supplementary Table S11). CPU and CUDA accuracies differed by at most 0.0003 across MetRef, "
        "Retina, and CIFAR-100; CUDA was slower on the two smaller tasks but reduced CIFAR-100 total "
        "time from 8.07 to 0.724 s. The multi-seed numerical validation met the stated tolerances in "
        "585/585 CPU and 40/40 CUDA checks for oversampling 20 and two power iterations. Rejected "
        "controls and discordant Metal rSVD routes are excluded from validated performance claims. "
        "CPU IRLBA remains the deterministic numerical reference. float32 halves stored input "
        "representation but did not uniformly improve runtime, process memory, or agreement "
        "(Supplementary Tables S8-S9)."
    )

    paragraph_by_prefix(document, "NMR comprised").text = (
        "NMR comprised 1,200 training and 321 held-out spectra, with 13,000 predictors and 28,355 "
        "responses. The 4.6-4.8 ppm predictor interval was defined a priori from chemical-shift labels "
        "and set to zero in both training and test predictors before any training-only split, suppressing "
        "residual water resonance that could dominate latent directions. No held-out response information "
        "was used, and response spectra remained unmasked so prediction and metrics covered the complete "
        "target. Training-only selection used five paired splits and family-specific grids up to 300 "
        "components. The one-standard-error rule retained five PLS-SVD and 50 SIMPLS components, both "
        "interior to their evaluated grids (complete grids in Supplementary Section S17)."
    )

    paragraph_by_prefix(document, "The storage route and hardware").text = (
        "The storage route and hardware should likewise follow matrix shape. Implicit cross-covariance "
        "products are useful when explicitly storing the predictor-response cross-product or deflated "
        "intermediates is the memory bottleneck. In the controlled CPU study they reduced incremental "
        "RSS by 29.7-47.5%, but were slower below 32 MiB and approximately time-neutral at 32-64 MiB. "
        "CUDA first surpassed CPU rSVD at n = 5,000 in the sample-size sweep and p = 2,000 in the "
        "predictor-size sweep. These are hardware-specific crossovers because transfer, context creation, "
        "synchronization, aspect ratio, and device memory remain decisive. Metal was disadvantaged by "
        "host-assisted stages and dispatch overhead; its discordant rSVD routes were excluded from valid "
        "accelerator comparisons."
    )
    paragraph_by_prefix(document, "Solver choice separates").text = (
        "Solver choice separates exploratory acceleration from confirmatory analysis. rSVD with "
        "oversampling 20 and two power iterations is the primary approximate route, but scientifically "
        "consequential analyses should inspect case diagnostics and compare several predefined seeds. "
        "Material seed variation, diagnostic failure, ill-conditioning, or a requirement for deterministic "
        "reproducibility indicates CPU IRLBA. Boundary component counts remain best within the evaluated "
        "grid rather than optima. Nonlinear kernel PLS requires an n-by-n Gram matrix; the matrix alone "
        "occupies approximately 0.75, 4.66, and 18.63 GiB in float64 at n = 10,000, 25,000, and 50,000, "
        "before copies and workspaces. Unsupported combinations stop explicitly, whereas experimental "
        "or unqualified approximate routes warn and record diagnostics. Table 1 summarizes these decisions."
    )

    credit = paragraph_by_prefix(document, "Dupe Ojo:")
    credit.text = credit.text.replace(
        " This draft allocation must be confirmed by all authors before submission.", ""
    )
    acknowledgements = paragraph_by_prefix(document, "Acknowledgements")
    funding_heading = insert_before(document, acknowledgements, "Funding", style="Heading 1")
    insert_after(
        document,
        funding_heading,
        "This research received no specific grant from funding agencies in the public, commercial, "
        "or not-for-profit sectors. Computing resources are acknowledged separately below.",
        style="Normal",
    )
    availability = paragraph_by_prefix(document, "Code and benchmark outputs")
    availability.text = availability.text.replace(
        "A persistent archive identifier will replace the checksum-only review object before acceptance.",
        "The source archive, checksum, benchmark manifest, and public build record identify the evaluated "
        "release; a persistent repository deposit will accompany the accepted article."
    )

    recommendation = document.tables[1]
    recommendation.rows[6].cells[2].text = (
        "Use CPU IRLBA for confirmatory analysis. Use rSVD for exploratory acceleration or routine "
        "analysis only when case diagnostics and multi-seed sensitivity are satisfactory; retain controls "
        "and numerical status with the result."
    )
    recommendation.rows[9].cells[1].text = (
        "The float64 Gram matrix alone is 8n-squared bytes: about 0.75 GiB at 10,000, 4.66 GiB at "
        "25,000, and 18.63 GiB at 50,000 samples."
    )

    references = paragraph_by_prefix(document, "References")
    last_reference = document.paragraphs[-1]
    insert_after(
        document,
        last_reference,
        "[36] Johnson J, Douze M, Jegou H. Billion-scale similarity search with GPUs. IEEE Trans Big "
        "Data. 2021;7:535-547. https://doi.org/10.1109/TBDATA.2019.2921572.",
        style=last_reference.style,
    )

    replace_in_paragraphs(document, [
        ("frozen fastPLS", "archived fastPLS"),
        ("Frozen release", "Archived release"),
        ("frozen release", "archived release"),
        ("Frozen-release", "Archived-release"),
        ("frozen-release", "archived-release"),
        ("definitive audit", "numerical validation"),
        ("audited separately", "evaluated separately"),
        ("quarantined", "excluded from validated performance claims"),
        ("reviewed interface", "evaluated interface"),
        ("review object", "archived source"),
        ("Frozen central", "Archived central"),
        ("frozen CPU", "archived CPU"),
        ("frozen controlled", "archived controlled"),
        ("numerical-audit status", "numerical-validation status"),
    ])
    replace_everywhere(document, " MB", " MiB")
    replace_everywhere(document, "(MB)", "(MiB)")
    replace_everywhere(document, " GB", " GiB")

    for prefix, (path, width) in MAIN_FIGURES.items():
        replace_preceding_figure(document, prefix, path, width)
    format_tables(document, size=6.5)
    document.save(MAIN_OUTPUT)


def update_supplement():
    document = Document(SUPP_SOURCE)

    paragraph_by_prefix(document, "All central quantitative evidence").text = (
        "Central quantitative evidence used archived fastPLS 0.99.25, Git commit "
        "7887401b09e25f54a546a253c255741cb1ab48e5, from source archive SHA-256 "
        "604f74d72f5e9540f7378efd37f2ca6ed87ccb9e73b912211331a8cd97233481. The archive, "
        "scripts, input checksums, session information, generated tables, and file ledger are indexed in "
        "Table S15. The deposited NMR workflow is retained only as explicitly labelled historical context."
    )
    paragraph_by_prefix(document, "Route-level evidence is consolidated").text = (
        "Route-level evidence is consolidated rather than repeated. Table S1 defines residency; Tables "
        "S7-S7b cover deterministic validation, execution ablation, and PLS-SVD/SIMPLS shape comparison; "
        "Tables S7c-S8 and Figure S1 report controlled scaling and approximate-solver status; Table S9 "
        "defines float32 capability; Tables S10-S13 report software, backend, NMR, and ImageNet results; "
        "and Table S15 records provenance. Table S6 defines the evidence-status terminology used throughout."
    )

    paragraph_by_prefix(document, "S10. Reporting conventions").text = "S10. Evidence-status terminology"
    paragraph_by_prefix(document, "Completed, failed").text = (
        "Table S6 defines one hierarchy for execution, residency, and numerical evidence. A route can "
        "execute successfully yet remain experimental or numerically discordant. Only routes meeting "
        "the stated numerical tolerances contribute to validated speed claims; unavailable and failed "
        "routes remain visible."
    )
    paragraph_by_prefix(document, "S11. Authoritative evidence map").text = "S11. Evidence organization"
    paragraph_by_prefix(document, "This compact supplement contains").text = (
        "Each central result maps to one table or figure in Sections S12-S20. Expanded component paths, "
        "raw replicates, and machine-level diagnostics are indexed in the repository rather than repeated "
        "here."
    )
    paragraph_by_prefix(document, "Table S6.").text = (
        "Table S6. Evidence-status glossary used in the manuscript, figures, and Supplementary tables."
    )
    replace_table_rows(document.tables[8], [
        ("Term", "Meaning", "Reporting consequence"),
        ("Available / unavailable", "The requested route can / cannot be built and executed on the platform.", "Unavailable routes stop before model allocation."),
        ("Native / hybrid", "All principal stages remain on the named backend / one or more stages execute on the host.", "Residency is reported separately from numerical status."),
        ("Deterministic reference", "Fixed-input, fixed-control CPU IRLBA or dense-LAPACK validation route.", "Used for confirmatory numerical comparison; not called an exact dense solver when IRLBA is used."),
        ("Qualified approximate", "The rSVD route met every stated tolerance in the specified validation panel.", "Controls, seed, version, and status accompany each result."),
        ("Concordant / discordant", "Paired predictions and endpoint metrics do / do not meet the stated agreement tolerances.", "Discordant routes are excluded from validated speed claims."),
        ("Validated", "The family/backend/endpoint combination has direct numerical evidence within the stated scope.", "Claims are limited to that scope and precision."),
        ("Experimental / hybrid", "The route executes but evidence is incomplete or includes host-assisted stages.", "Warnings and diagnostics are retained; confirmatory claims are avoided."),
        ("Rejected / failed", "Controls failed numerical criteria / execution did not complete.", "The row remains visible and cannot support performance claims."),
    ])

    paragraph_by_prefix(document, "rSVD uses a Gaussian").text = (
        "rSVD uses a Gaussian range sketch, oversampling, power iterations, and a recorded seed. "
        "Qualification required relative Frobenius prediction error <= 0.05, prediction correlation "
        ">= 0.99, each score/projection/loading subspace angle <= 10 degrees, decoded-label agreement "
        ">= 0.99, and absolute predictive-metric difference <= 0.01. These predefined values are "
        "computational non-inferiority screens, not clinical or exact-equivalence margins. Release "
        "validation used seeds 1, 7, 19, 43, and 123 and required every check to meet tolerance. The "
        "oversampling-20, two-power configuration met 585/585 CPU and 40/40 CUDA checks in that panel; "
        "the broader controlled shape study identified out-of-tolerance regimes, which are reported "
        "separately and excluded from validated speed claims. Metal remains unqualified."
    )
    paragraph_by_prefix(document, "A historical single-seed development audit").text = (
        "A historical single-seed development evaluation met 101/117 tolerance checks at oversampling "
        "10 with one power iteration. Archived-release validation therefore used oversampling 20, two "
        "power iterations, and seeds 1, 7, 19, 43, and 123. All 585 CPU checks and all 40 CUDA checks "
        "met the predefined tolerances; rejected CUDA 10/1 and 10/2 controls met 0/40 and 27/40 checks. "
        "Component-level outputs are in benchmark_results/frozen_release_0.99.25/"
        "simpls_estimator_preservation/ and rsvd_cuda_reliability_final/. These are solver-control "
        "reliability screens, not estimator-equivalence evidence."
    )
    paragraph_by_prefix(document, "Table S8.").text = (
        "Table S8. rSVD numerical-validation summary for fastPLS 0.99.25. Release rows use seeds 1, 7, "
        "19, 43, and 123; rejected alternatives were evaluated over the same seeds."
    )
    paragraph_by_prefix(document, "Table S7. Definitive").text = paragraph_by_prefix(
        document, "Table S7. Definitive"
    ).text.replace("Definitive", "Primary")
    paragraph_by_prefix(document, "Table S9. Definitive").text = paragraph_by_prefix(
        document, "Table S9. Definitive"
    ).text.replace("Definitive", "Primary")
    for row in document.tables[14].rows[1:]:
        row.cells[-1].text = f"v0.99.25; {row.cells[-1].text}"

    paragraph_by_prefix(document, "The definitive repeated comparison").text = paragraph_by_prefix(
        document, "The definitive repeated comparison"
    ).text.replace("The definitive", "The repeated")
    paragraph_by_prefix(document, "Table S10c.").text = paragraph_by_prefix(
        document, "Table S10c."
    ).text.replace("Values are MB", "Values are MiB")
    paragraph_by_prefix(document, "Table S10d.").text = paragraph_by_prefix(
        document, "Table S10d."
    ).text.replace("Values are MB", "Values are MiB")
    paragraph_by_prefix(document, "This frozen-release comparison").text = paragraph_by_prefix(
        document, "This frozen-release comparison"
    ).text.replace("This frozen-release", "This archived-release")
    paragraph_by_prefix(document, "Table S10f.").text = (
        "Table S10f. Archived-release single-thread CPU end-to-end comparison with IKPLS. Time and IQR "
        "are fit-plus-prediction seconds; RSS values are MiB. fastPLS rSVD rows use version 0.99.25, "
        "oversampling 20, two power iterations, seed 123, and case-specific diagnostics."
    )
    paragraph_by_prefix(document, "Reproducibility. Frozen CPU results").text = paragraph_by_prefix(
        document, "Reproducibility. Frozen CPU results"
    ).text.replace("Reproducibility. Frozen CPU", "Reproducibility. Archived CPU").replace(
        "frozen environment", "archived environment"
    )
    for row in document.tables[21].rows[1:]:
        if "fastPLS" in row.cells[1].text:
            if "rSVD" in row.cells[1].text:
                row.cells[1].text = "fastPLS SIMPLS/rSVD\n(v0.99.25; os20/p2/s123; qualified)"
            else:
                row.cells[1].text = "fastPLS SIMPLS/IRLBA\n(v0.99.25; deterministic)"

    paragraph_by_prefix(document, "This focused frozen-release analysis").text = (
        "This focused archived-release analysis uses identical stored splits and component counts for "
        "paired CPU/CUDA SIMPLS-rSVD routes: MetRef 22, Retina 20, and CIFAR-100 100. Every row uses "
        "fastPLS 0.99.25, oversampling 20, two power iterations, seed 123, case-specific diagnostics, "
        "and three repetitions. Host memory is the post-prediction minus pre-fit RSS snapshot; device "
        "memory includes runtime/context allocation."
    )
    paragraph_by_prefix(document, "Table S11.").text = (
        "Table S11. Archived-release paired CPU/CUDA SIMPLS-rSVD results. Every row uses fastPLS "
        "0.99.25, oversampling 20, two power iterations, seed 123, and qualified-approximate status; "
        "time is median (IQR) over three runs."
    )
    table_s11 = document.tables[22]
    table_s11.rows[0].cells[6].text = "Host RSS delta MiB"
    table_s11.rows[0].cells[7].text = "GPU delta MiB"
    for row in table_s11.rows[1:]:
        backend = "CUDA" if "CUDA" in row.cells[-1].text else "CPU"
        row.cells[-1].text = f"{backend}; v0.99.25; os20/p2/s123; qualified"

    paragraph_by_prefix(document, "The family-selected paired analysis").text = (
        "The family-selected paired analysis compares CPU IRLBA, CPU rSVD, and CUDA rSVD at five "
        "PLS-SVD or 50 SIMPLS components. The 4.6-4.8 ppm interval was defined a priori from predictor "
        "chemical-shift labels and set to zero in both training and test predictors before any "
        "training-only split to suppress residual water resonance; no response coordinate was masked, "
        "and no held-out response information informed preprocessing. The deposited 165-component "
        "workflow remains historical context. Figure 4 displays AMI-0030-9 (index 38), selected as the "
        "held-out sample closest to the median 50-component SIMPLS CUDA/rSVD per-spectrum RMSD."
    )
    paragraph_by_prefix(document, "Table S12.").text = (
        "Table S12. Archived-release NMR evidence. Approximate rows use fastPLS 0.99.25, oversampling "
        "20, two power iterations, seed 123, and case-specific diagnostics; deterministic rows use CPU "
        "IRLBA. Family-selected prediction, the fixed-100-component implementation comparison, and the "
        "deposited historical workflow answer different questions."
    )
    table_s12 = document.tables[23]
    table_s12.rows[0].cells[6].text = "Delta host MiB"
    table_s12.rows[0].cells[7].text = "Delta GPU MiB"
    for row in table_s12.rows[1:]:
        status = row.cells[-1].text
        if "rSVD" in status:
            row.cells[-1].text = "v0.99.25; rSVD os20/p2/s123; qualified"
        elif "deterministic" in status:
            row.cells[-1].text = "v0.99.25; CPU IRLBA; deterministic"

    image_paragraph = paragraph_by_prefix(document, "The historical archive contained")
    image_paragraph.text = image_paragraph.text.replace("Frozen fastPLS", "Archived fastPLS")
    faiss = insert_after(
        document,
        image_paragraph,
        "A separate repository benchmark uses FAISS [36] to compare nearest-neighbour retrieval on raw "
        "embeddings, PCA scores, and PLS scores, including transformation, index construction, and query "
        "time. It is an external retrieval analysis rather than a fastPLS classifier and is not pooled "
        "with the SIMPLS-LDA classification experiment below.",
    )
    paragraph_by_prefix(document, "Table S13.").text = (
        "Table S13. Historical archived-release ImageNet downstream classification. Every row uses "
        "fastPLS 0.99.25, CUDA rSVD oversampling 20, two power iterations, seed 123, and qualified-"
        "approximate controls. One shared fit supplies all component prefixes; the 1,000-component row "
        "is a boundary stress point, not a selected optimum. All rows are single-run, partially "
        "reproducible feasibility estimates."
    )
    table_s13 = document.tables[24]
    table_s13.rows[0].cells[6].text = "Host RSS MiB"
    table_s13.rows[0].cells[7].text = "GPU MiB"
    for row in table_s13.rows[1:]:
        row.cells[-1].text = "v0.99.25; os20/p2/s123; qualified"

    paragraph_by_prefix(document, "All central rows resolve").text = (
        "All central rows resolve to archived fastPLS 0.99.25, Git commit "
        "7887401b09e25f54a546a253c255741cb1ab48e5, and source archive SHA-256 "
        "604f74d72f5e9540f7378efd37f2ca6ed87ccb9e73b912211331a8cd97233481. The "
        "machine-readable ledger contains SHA-256 values for 185 evidence files. The deposited "
        "165-component NMR workflow is the sole historical context row in the central ledger."
    )
    paragraph_by_prefix(document, "Table S15.").text = (
        "Table S15. Archived-release central analysis provenance. Archive paths are relative to the "
        "repository root; full per-file checksums are in provenance/frozen_evidence_file_ledger.csv."
    )

    example = paragraph_by_prefix(document, "benchmark/synthetic_end_to_end_example.R")
    example.text = (
        "benchmark/synthetic_end_to_end_example.R reproduces the complete unrestricted workflow. The "
        "minimal user example below shows deterministic fitting, the primary rSVD route, prediction, "
        "diagnostic inspection, CUDA dispatch, and informative handling of an unavailable backend."
    )
    code = insert_after(
        document,
        example,
        'remotes::install_github("tkcaccia/fastPLS", ref = "7887401b09e25f54a546a253c255741cb1ab48e5")\n'
        'library(fastPLS)\n'
        'data(breast)\n\n'
        'fit_det <- pls(breast$X_train, breast$y_train, ncomp = 1:5,\n'
        '               method = "simpls", svd.method = "irlba", backend = "cpu")\n'
        'fit_rsvd <- pls(breast$X_train, breast$y_train, ncomp = 1:5,\n'
        '                method = "simpls", svd.method = "rsvd", backend = "cpu",\n'
        '                oversample = 20, power = 2, seed = 123)\n'
        'fit_rsvd$diagnostics[c("solver", "status", "approximation_audited")]\n'
        'fit_rsvd$diagnostics$rsvd[c("oversample", "power", "case_audit")]\n'
        'pred <- predict(fit_rsvd, breast$X_test)\n'
        'pred_5 <- pred$Ypred[["ncomp=5"]]\n'
        'evaluate(predicted = pred_5, observed = breast$y_test)\n\n'
        'if (has_cuda()) {\n'
        '  fit_cuda <- pls(breast$X_train, breast$y_train, ncomp = 1:5,\n'
        '                  method = "simpls", svd.method = "rsvd", backend = "cuda")\n'
        '  list(residency = fit_cuda$gpu_resident, diagnostics = fit_cuda$diagnostics)\n'
        '} else {\n'
        '  try(pls(breast$X_train, breast$y_train, backend = "cuda"))\n'
        '}',
        style="Normal",
    )
    format_code(code)

    paragraph_by_prefix(document, "Full component paths").text = (
        "Full component paths, raw run rows, sensitivity analyses, and machine-level diagnostics are "
        "indexed in benchmark/MANUSCRIPT_EVIDENCE_ARCHIVE.md. Their CSV, RDS, PDF, PNG, log, and "
        "session-information files remain available for reproducibility without duplicating the "
        "definitive tables in this supplement."
    )
    paragraph_by_prefix(document, "The latest public BiocStaging matrix").text = (
        "The public BiocStaging build for fastPLS 0.99.25 used commit "
        "7887401b09e25f54a546a253c255741cb1ab48e5 and completed successfully "
        "(https://github.com/r-universe/biocstaging/actions/runs/32822192644). Hosted macOS jobs establish "
        "compilation and CPU fallback, not Metal runtime execution. Linux CUDA and macOS Metal runtime "
        "validation therefore remains tied to dedicated hardware reports. Unit tests verify informative "
        "failure before model allocation for unavailable backend/solver combinations. Public package "
        "metadata: https://biocstaging.r-universe.dev/fastPLS."
    )
    paragraph_by_prefix(document, "References are numbered").text = (
        "References are numbered as in the main manuscript. The final mapping is Retina [25], Tabula "
        "Muris [26], PRISM [27], ImageNet [28], DINOv2 [29], CIFAR-100 [30], UNI [31], Prov-GigaPath "
        "[32], Improved Kernel PLS [33], IKPLS [34], fast cross-product validation [35], and FAISS [36]."
    )

    replace_in_paragraphs(document, [
        ("frozen fastPLS", "archived fastPLS"),
        ("Frozen release", "Archived release"),
        ("frozen release", "archived release"),
        ("Frozen-release", "Archived-release"),
        ("frozen-release", "archived-release"),
        ("definitive", "primary"),
        ("authoritative", "primary"),
        ("review-cycle", "development-stage"),
        ("reviewed Mac build", "evaluated Mac build"),
        ("quarantined", "excluded from validated performance claims"),
        ("Frozen central", "Archived central"),
        ("frozen CPU", "archived CPU"),
        ("frozen controlled", "archived controlled"),
        ("frozen 0.99.25 archive", "archived 0.99.25 release"),
        ("numerical-audit", "numerical-validation"),
        ("definitive tables", "primary tables"),
    ])
    replace_everywhere(document, " MB", " MiB")
    replace_everywhere(document, "(MB)", "(MiB)")
    replace_everywhere(document, " GB", " GiB")

    for prefix, (path, width) in SUPP_FIGURES.items():
        replace_preceding_figure(document, prefix, path, width)
    format_tables(document, size=6.2)
    document.save(SUPP_OUTPUT)


def copy_figures():
    target = OUTPUT / "figures"
    target.mkdir(exist_ok=True)
    for source in set(path for path, _ in MAIN_FIGURES.values()) | set(
        path for path, _ in SUPP_FIGURES.values()
    ):
        copy2(source, target / source.name)


if __name__ == "__main__":
    update_main()
    update_supplement()
    copy_figures()
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)
