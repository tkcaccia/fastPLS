#!/usr/bin/env python3

from pathlib import Path
from copy import deepcopy
import math
import re
import sys

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
sys.path.insert(0, str(ROOT / "tools"))
import revise_cmpb_cycle9_simpls_validation as c9  # noqa: E402


SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle12"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle13"
EVIDENCE = ROOT / "benchmark_results" / "manuscript_revision_cycle13_20260725"
MULTI = ROOT / "benchmark_results" / "manuscript_multidataset_summary_20260725"
NMR_REF = ROOT / "benchmark_results" / "review_nmr_reference_20260725"
NMR_SIMPLS = ROOT / "benchmark_results" / "review_nmr_20260724"
IMAGENET_BASE = ROOT / "benchmark_results" / "imagenet_faiss_matched_1m_20260725"
KERNEL_RESULTS = (
    EVIDENCE / "kernel_suite" / "supplementary_kernel_sensitivity"
)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle12_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle12_0.99.6_20260725.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle13_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle13_0.99.6_20260725.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle13_20260725.docx"
ALL_DATASET_FIGURE = EVIDENCE / "plots" / "selected_performance_all_datasets.png"

DATASET_ORDER = [
    "metref",
    "ccle",
    "tcga_brca",
    "tcga_hnsc_methylation",
    "gtex_v8",
    "tcga_pan_cancer",
    "retina",
    "tabula",
    "cifar100",
    "cbmc_citeseq",
    "prism",
    "nmr",
]
DATASET_LABELS = {
    "metref": "MetRef",
    "ccle": "CCLE",
    "tcga_brca": "TCGA-BRCA",
    "tcga_hnsc_methylation": "TCGA-HNSC methyl.",
    "gtex_v8": "GTEx v8",
    "tcga_pan_cancer": "TCGA Pan-Cancer",
    "retina": "Retina",
    "tabula": "Tabula Muris",
    "cifar100": "CIFAR-100",
    "cbmc_citeseq": "CBMC CITE-seq",
    "prism": "PRISM",
    "nmr": "NMR",
}
METHOD_ORDER = ["plssvd", "simpls", "opls", "kernelpls"]


def finite(value):
    try:
        return math.isfinite(float(value))
    except (TypeError, ValueError):
        return False


def fnum(value, digits=3):
    if not finite(value):
        return "NA"
    return f"{float(value):.{digits}f}"


def compact_metric(name, value):
    if not finite(value):
        return "NA"
    name = str(name).lower()
    if name == "accuracy":
        return f"Acc={float(value):.3f}"
    if name == "q2":
        return f"Q2={float(value):.3f}"
    if abs(float(value)) < 0.01:
        return f"RMSD={float(value):.2e}"
    return f"RMSD={float(value):.3f}"


def raw_benchmark_rows_to_selected(paths):
    frames = [pd.read_csv(path) for path in paths if Path(path).exists()]
    if not frames:
        return []
    data = pd.concat(frames, ignore_index=True)
    data = data[data["status"].eq("ok")].copy()
    output = []
    group_columns = [
        "dataset",
        "task_type",
        "method_panel",
        "variant_name",
        "engine",
        "backend",
        "classifier",
        "requested_ncomp",
        "effective_ncomp",
        "precision",
        "execution_precision",
        "metric_name",
    ]
    for keys, group in data.groupby(group_columns, dropna=False):
        row = dict(zip(group_columns, keys))
        row.update(
            {
                "n_runs": len(group),
                "metric_median": group["metric_value"].median(),
                "metric_q25": group["metric_value"].quantile(0.25),
                "metric_q75": group["metric_value"].quantile(0.75),
                "total_time_sec_median": group["total_time_ms"].median() / 1000,
                "total_time_sec_q25": group["total_time_ms"].quantile(0.25) / 1000,
                "total_time_sec_q75": group["total_time_ms"].quantile(0.75) / 1000,
                "host_rss_mb_median": group["peak_host_rss_mb"].median(),
                "host_rss_mb_q25": group["peak_host_rss_mb"].quantile(0.25),
                "host_rss_mb_q75": group["peak_host_rss_mb"].quantile(0.75),
                "gpu_mem_mb_median": group["peak_gpu_mem_mb"].median(),
                "gpu_mem_mb_q25": group["peak_gpu_mem_mb"].quantile(0.25),
                "gpu_mem_mb_q75": group["peak_gpu_mem_mb"].quantile(0.75),
            }
        )
        output.append(row)
    return output


def move_after(anchor, element):
    anchor._p.addnext(element)


def insert_table_after(document, anchor, caption_text, headers, rows, widths, font_size):
    caption = document.add_paragraph(caption_text, style="Caption")
    caption.paragraph_format.keep_with_next = True
    table = c9.add_table(document, headers, rows, widths, font_size=font_size)
    move_after(anchor, caption._p)
    caption._p.addnext(table._tbl)
    return table


def replace_table(old_table, new_table):
    old_table._tbl.getparent().replace(old_table._tbl, new_table._tbl)


def paragraph_math_text(paragraph):
    return "".join(paragraph._p.itertext()).replace("\u200b", "").strip()


def correct_lda_content(document, supplement=False):
    if supplement:
        for paragraph in list(document.paragraphs):
            math_text = paragraph_math_text(paragraph)
            if math_text.startswith("ℓc=") or math_text.startswith("gc="):
                c9.delete_paragraph(paragraph)
        section_heading = c9.find_paragraph(document, "S3. Classification heads")
        argmax_heading = c9.find_paragraph(document, "S3.1 Argmax")
        section_heading.paragraph_format.space_after = Pt(8)
        section_heading.paragraph_format.line_spacing = 1.0
        argmax_heading.paragraph_format.space_before = Pt(4)
        argmax_heading.paragraph_format.space_after = Pt(4)
        argmax_heading.paragraph_format.line_spacing = 1.0
        return

    c9.set_paragraph_text(
        c9.find_paragraph(document, "Systems"),
        "For each class, the coefficient vector is obtained by solving the pooled-"
        "covariance system against the corresponding class mean using Cholesky "
        "factorization and triangular solves. A deterministic trace-scaled diagonal "
        "regularization sequence is advanced only when Cholesky fails; no inverse "
        "or tuned classifier parameter is introduced. The discriminant is [18,19].",
    )
    source = Document(SUPP_SOURCE)
    correct_equation = next(
        paragraph
        for paragraph in source.paragraphs
        if paragraph_math_text(paragraph).startswith("δct=")
    )
    wrong_equation = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph_math_text(paragraph).startswith("ℓc=")
    )
    wrong_equation._p.getparent().replace(
        wrong_equation._p,
        deepcopy(correct_equation._p),
    )


def selected_rows():
    selected = pd.read_csv(EVIDENCE / "selected_backend_with_uncertainty.csv")
    selected = selected[
        ~selected["dataset"].isin(["singlecell", "imagenet"])
        & selected["precision"].eq("float64")
    ].copy()

    selected_real_paths = list(
        EVIDENCE.glob(
            "retina_tabula_selected_outer/runs/*/dataset_memory_compare_raw.csv"
        )
    )
    selected_real_paths.extend(
        EVIDENCE.glob(
            "retina_tabula_selected_opls/runs/*/dataset_memory_compare_raw.csv"
        )
    )
    selected_real_rows = raw_benchmark_rows_to_selected(selected_real_paths)
    if selected_real_rows:
        selected = pd.concat(
            [selected, pd.DataFrame(selected_real_rows)], ignore_index=True
        )
    else:
        external = pd.read_csv(MULTI / "source" / "external_float64_summary.csv")
        external = external[
            external["dataset"].isin(["retina", "tabula"])
            & external["package"].eq("fastPLS")
            & external["method_id"].str.contains("rsvd", case=False, na=False)
            & external["classifier"].eq("argmax")
            & external["reps_ok"].gt(0)
        ].copy()
        ext_rows = []
        for _, row in external.iterrows():
            ext_rows.append(
                {
                    "dataset": row["dataset"],
                    "task_type": row["task_type"],
                    "method_panel": row["requested_estimator"],
                    "variant_name": row["method_id"],
                    "engine": "GPU" if "cuda" in row["method_id"].lower() else "CPU",
                    "backend": "gpu_native" if "cuda" in row["method_id"].lower() else "cpu_rsvd",
                    "classifier": "argmax",
                    "requested_ncomp": row["ncomp_requested"],
                    "effective_ncomp": row["ncomp_requested"],
                    "precision": "float64",
                    "execution_precision": "float64",
                    "metric_name": row["metric_name"],
                    "n_runs": row["reps_ok"],
                    "metric_median": row["median_metric"],
                    "metric_q25": row["median_metric"] - row["iqr_metric"] / 2,
                    "metric_q75": row["median_metric"] + row["iqr_metric"] / 2,
                    "total_time_sec_median": row["median_time_ms"] / 1000,
                    "total_time_sec_q25": (row["median_time_ms"] - row["iqr_time_ms"] / 2) / 1000,
                    "total_time_sec_q75": (row["median_time_ms"] + row["iqr_time_ms"] / 2) / 1000,
                    "host_rss_mb_median": row["median_peak_host_rss_mb"],
                    "host_rss_mb_q25": row["median_peak_host_rss_mb"],
                    "host_rss_mb_q75": row["median_peak_host_rss_mb"],
                    "gpu_mem_mb_median": math.nan,
                    "gpu_mem_mb_q25": math.nan,
                    "gpu_mem_mb_q75": math.nan,
                }
            )
        selected = pd.concat([selected, pd.DataFrame(ext_rows)], ignore_index=True)

    reference = pd.read_csv(NMR_REF / "nmr_reference_comparison_summary.csv")
    simpls_rows = []
    for _, row in reference[reference["variant"].str.startswith("fastpls_")].iterrows():
        simpls_rows.append(
            {
                "dataset": "nmr",
                "task_type": "regression",
                "method_panel": "plssvd",
                "variant_name": row["variant"],
                "engine": row["backend"].upper(),
                "backend": f'{row["backend"]}_{row["svd_method"]}',
                "classifier": "not_applicable",
                "requested_ncomp": 100,
                "effective_ncomp": 100,
                "precision": "float64",
                "execution_precision": "float64",
                "metric_name": "rmsd",
                "n_runs": row["n_repetitions"],
                "metric_median": row["RMSD_median"],
                "metric_q25": row["RMSD_median"] - row["RMSD_iqr"] / 2,
                "metric_q75": row["RMSD_median"] + row["RMSD_iqr"] / 2,
                "total_time_sec_median": row["total_time_sec_median"],
                "total_time_sec_q25": row["total_time_sec_median"] - row["total_time_sec_iqr"] / 2,
                "total_time_sec_q75": row["total_time_sec_median"] + row["total_time_sec_iqr"] / 2,
                "host_rss_mb_median": row["host_rss_mb_median"],
                "host_rss_mb_q25": row["host_rss_mb_median"] - row["host_rss_mb_iqr"] / 2,
                "host_rss_mb_q75": row["host_rss_mb_median"] + row["host_rss_mb_iqr"] / 2,
                "gpu_mem_mb_median": row["gpu_peak_mb_median"],
                "gpu_mem_mb_q25": row["gpu_peak_mb_median"],
                "gpu_mem_mb_q75": row["gpu_peak_mb_median"],
            }
        )
    for backend in ("cpu", "cuda"):
        summary = pd.read_csv(
            NMR_SIMPLS / f"nmr_final_{backend}_summary.csv"
        ).iloc[0]
        simpls_rows.append(
            {
                "dataset": "nmr",
                "task_type": "regression",
                "method_panel": "simpls",
                "variant_name": f"fastpls_simpls_{backend}_rsvd",
                "engine": backend.upper(),
                "backend": f"{backend}_rsvd",
                "classifier": "not_applicable",
                "requested_ncomp": 100,
                "effective_ncomp": 100,
                "precision": "float64",
                "execution_precision": "float64",
                "metric_name": "rmsd",
                "n_runs": 3,
                "metric_median": summary["RMSD"],
                "metric_q25": summary["RMSD"],
                "metric_q75": summary["RMSD"],
                "total_time_sec_median": 20.140 if backend == "cpu" else 3.055,
                "total_time_sec_q25": 19.938 if backend == "cpu" else 3.041,
                "total_time_sec_q75": 20.342 if backend == "cpu" else 3.069,
                "host_rss_mb_median": 3143 if backend == "cpu" else 3469,
                "host_rss_mb_q25": 3143,
                "host_rss_mb_q75": 3143 if backend == "cpu" else 3469,
                "gpu_mem_mb_median": math.nan if backend == "cpu" else 3432,
                "gpu_mem_mb_q25": math.nan if backend == "cpu" else 3432,
                "gpu_mem_mb_q75": math.nan if backend == "cpu" else 3432,
            }
        )
    selected = pd.concat([selected, pd.DataFrame(simpls_rows)], ignore_index=True)

    nmr_selected_paths = sorted(
        EVIDENCE.glob("nmr_plssvd_selected10/rows/*.csv")
    )
    if nmr_selected_paths:
        nmr_data = pd.concat(
            [pd.read_csv(path) for path in nmr_selected_paths],
            ignore_index=True,
        )
        nmr_rows = []
        for (variant, backend), group in nmr_data.groupby(["variant", "backend"]):
            nmr_rows.append(
                {
                    "dataset": "nmr",
                    "task_type": "regression",
                    "method_panel": "plssvd",
                    "variant_name": variant,
                    "engine": backend.upper(),
                    "backend": f"{backend}_rsvd",
                    "classifier": "not_applicable",
                    "requested_ncomp": 10,
                    "effective_ncomp": 10,
                    "precision": "float64",
                    "execution_precision": "float64",
                    "metric_name": "rmsd",
                    "n_runs": len(group),
                    "metric_median": group["RMSD"].median(),
                    "metric_q25": group["RMSD"].quantile(0.25),
                    "metric_q75": group["RMSD"].quantile(0.75),
                    "total_time_sec_median": group["total_time_sec"].median(),
                    "total_time_sec_q25": group["total_time_sec"].quantile(0.25),
                    "total_time_sec_q75": group["total_time_sec"].quantile(0.75),
                    "host_rss_mb_median": group["host_rss_mb"].median(),
                    "host_rss_mb_q25": group["host_rss_mb"].quantile(0.25),
                    "host_rss_mb_q75": group["host_rss_mb"].quantile(0.75),
                    "gpu_mem_mb_median": group["gpu_peak_mb"].median(),
                    "gpu_mem_mb_q25": group["gpu_peak_mb"].quantile(0.25),
                    "gpu_mem_mb_q75": group["gpu_peak_mb"].quantile(0.75),
                }
            )
        selected = pd.concat(
            [selected, pd.DataFrame(nmr_rows)], ignore_index=True
        )

    output = []
    for dataset in DATASET_ORDER:
        for method in METHOD_ORDER:
            group = selected[
                selected["dataset"].eq(dataset)
                & selected["method_panel"].eq(method)
            ].copy()
            if group.empty:
                output.append(
                    {
                        "dataset": dataset,
                        "method_panel": method,
                        "status": "not evaluated",
                    }
                )
                continue
            group = group.sort_values("total_time_sec_median")
            row = group.iloc[0].to_dict()
            row["status"] = "ok"
            output.append(row)
    chosen = pd.DataFrame(output)
    selected.to_csv(EVIDENCE / "selected_backend_cycle13_all.csv", index=False)
    chosen.to_csv(EVIDENCE / "selected_backend_cycle13_chosen.csv", index=False)
    return chosen, selected


def main_table_rows(chosen):
    rows = []
    for dataset in DATASET_ORDER:
        values = [DATASET_LABELS[dataset]]
        for method in METHOD_ORDER:
            hit = chosen[
                chosen["dataset"].eq(dataset)
                & chosen["method_panel"].eq(method)
            ].iloc[0]
            if hit["status"] != "ok":
                values.append("Not evaluated")
                continue
            backend = (
                "CUDA rSVD"
                if str(hit["engine"]).upper() in ("GPU", "CUDA")
                else "CPU rSVD"
            )
            time_iqr = (
                float(hit["total_time_sec_q75"]) - float(hit["total_time_sec_q25"])
                if finite(hit.get("total_time_sec_q75"))
                and finite(hit.get("total_time_sec_q25"))
                else math.nan
            )
            gpu = (
                f'{float(hit["gpu_mem_mb_median"]):.0f}'
                if finite(hit.get("gpu_mem_mb_median"))
                else "—"
            )
            runs = int(hit["n_runs"]) if finite(hit.get("n_runs")) else 0
            values.append(
                f"{backend} | k={int(hit['effective_ncomp'])}\n"
                f"{compact_metric(hit['metric_name'], hit['metric_median'])} | "
                f"{float(hit['total_time_sec_median']):.2f} s"
                + (f" (IQR {time_iqr:.2f})" if runs > 1 and finite(time_iqr) else "")
                + f"\nH={float(hit['host_rss_mb_median']):.0f}; G={gpu} MB | "
                f"f64 | n={runs}"
            )
        rows.append(tuple(values))
    return rows


def replace_main_table(document, chosen):
    old = document.tables[0]
    new = c9.add_table(
        document,
        ["Dataset", "PLS-SVD", "SIMPLS", "OPLS", "kernel PLS"],
        main_table_rows(chosen),
        [0.72, 1.47, 1.47, 1.47, 1.47],
        font_size=5.25,
    )
    replace_table(old, new)
    return new


def replace_dataset_table(document):
    old = document.tables[2]
    rows = [
        ("TCGA-HNSC methylation", "Classification", "520", "58", "782", "2"),
        ("CCLE", "Classification", "547", "71", "1,000", "18"),
        ("TCGA-BRCA", "Classification", "756", "88", "1,000", "5"),
        ("MetRef", "Classification", "773", "100", "375", "22"),
        ("TCGA Pan-Cancer", "Classification", "3,000", "982", "850", "32"),
        ("GTEx v8", "Classification", "3,000", "797", "1,000", "32"),
        ("Retina", "Classification", "22,402", "22,406", "50", "12"),
        ("Tabula Muris", "Classification", "50,043", "50,059", "50", "32"),
        ("CIFAR-100", "Classification", "50,000", "10,000", "768", "100"),
        ("PRISM", "Multivariate regression", "479", "54", "1,000", "4,686"),
        ("NMR", "Multivariate regression", "1,200", "321", "13,000", "28,355"),
        ("CBMC CITE-seq", "Multivariate regression", "7,755", "862", "1,000", "10"),
        ("ImageNet/DINOv2", "Retrieval stress test", "1,000,000", "281,167", "1,024", "1,000"),
    ]
    new = c9.add_table(
        document,
        ["Dataset", "Task", "Prepared n train", "Prepared n test", "Prepared p", "Prepared q"],
        rows,
        [1.25, 1.2, 0.85, 0.85, 0.65, 0.65],
        font_size=6.2,
    )
    replace_table(old, new)


def replace_supplement_method_tables(document):
    chosen, selected = selected_rows()
    target_tables = list(document.tables[13:17])
    for target_table, method in zip(target_tables, METHOD_ORDER):
        rows = []
        for dataset in DATASET_ORDER:
            chosen_row = chosen[
                chosen["dataset"].eq(dataset)
                & chosen["method_panel"].eq(method)
            ].iloc[0]
            if chosen_row["status"] != "ok":
                rows.append(
                    (
                        DATASET_LABELS[dataset],
                        "Not evaluated",
                        "—",
                        "—",
                        "—",
                        "—",
                        "—",
                        "—",
                        "—",
                        "not evaluated",
                    )
                )
                continue
            rows_at_selected = selected[
                selected["dataset"].eq(dataset)
                & selected["method_panel"].eq(method)
                & selected["effective_ncomp"].eq(
                    chosen_row["effective_ncomp"]
                )
            ].copy()
            rows_at_selected["_engine_order"] = rows_at_selected["engine"].map(
                {"CPU": 0, "GPU": 1, "CUDA": 1}
            ).fillna(2)
            rows_at_selected = rows_at_selected.sort_values(
                ["_engine_order", "total_time_sec_median"]
            )
            for _, row in rows_at_selected.iterrows():
                engine = str(row["engine"]).upper()
                backend = "CUDA rSVD" if engine in ("GPU", "CUDA") else "CPU rSVD"
                metric_name = str(row["metric_name"]).lower()
                metric_abbreviation = {
                    "accuracy": "Acc",
                    "q2": "Q2",
                    "rmsd": "RMSD",
                }.get(metric_name, metric_name)
                value = (
                    f'{row["metric_median"]:.3e}'
                    if metric_name == "rmsd"
                    and abs(float(row["metric_median"])) < 0.01
                    else f'{row["metric_median"]:.3f}'
                )
                gpu = (
                    f'{row["gpu_mem_mb_median"]:.0f}'
                    if finite(row["gpu_mem_mb_median"])
                    else "—"
                )
                rows.append(
                    (
                        DATASET_LABELS[dataset],
                        backend,
                        int(row["effective_ncomp"]),
                        metric_abbreviation,
                        value,
                        f'{row["total_time_sec_median"]:.3f}',
                        f'{row["host_rss_mb_median"]:.0f}',
                        gpu,
                        "f64",
                        "ok",
                    )
                )
        new = c9.add_table(
            document,
            ["Dataset", "Backend", "k", "Metric", "Value", "Time (s)", "H", "G", "Prec.", "Status"],
            rows,
            [0.95, 0.72, 0.3, 0.48, 0.62, 0.65, 0.5, 0.45, 0.45, 0.65],
            font_size=5.3,
        )
        replace_table(target_table, new)


def add_all_dataset_figure(document, table):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(ALL_DATASET_FIGURE), width=Inches(6.65))
    caption = document.add_paragraph(
        "Figure 2. Outer-test predictive performance at the component count selected "
        "from training data for the twelve biomedical benchmark tasks. Points show "
        "the fastest completed matched CPU/CUDA rSVD row within each PLS family; "
        "labels give the predictive value and effective component count. Accuracy "
        "is higher-is-better and RMSD lower-is-better. ImageNet is excluded because "
        "it used the separate matched retrieval protocol in Figure 4.",
        style="Caption",
    )
    table._tbl.addnext(paragraph._p)
    paragraph._p.addnext(caption._p)


def external_rows():
    data = pd.read_csv(EVIDENCE / "external_simpls_float64_compact.csv")
    rows = []
    for _, row in data.iterrows():
        rows.append(
            (
                DATASET_LABELS.get(row["dataset"], row["dataset"]),
                int(row["ncomp"]),
                row["fastest_fastpls"],
                f'{row["fastpls_time_ms"]:.0f} ({row["fastpls_time_iqr_ms"]:.1f})',
                f'{row["reference_time_ms"]:.0f} ({row["reference_time_iqr_ms"]:.1f})',
                f'{row["speedup_vs_pls"]:.2f}',
                f'{row["fastpls_metric"]:.4f}',
                f'{row["reference_metric"]:.4f}',
            )
        )
    return rows


def revise_main():
    document = Document(MAIN_SOURCE)
    chosen, _ = selected_rows()
    table1 = replace_main_table(document, chosen)
    add_all_dataset_figure(document, table1)

    replacements = {
        "Methods:": (
            "Methods: fastPLS provides PLS-SVD, SIMPLS, orthogonal PLS (OPLS), "
            "and kernel PLS through one R interface. Deterministic IRLBA SIMPLS "
            "preserves de Jong's sequential estimator while reusing deflation, "
            "coefficient, and prediction state across component prefixes. rSVD is "
            "analysed separately as an approximate direction solver. The package "
            "combines implicit cross-covariance products, compact prediction, "
            "compiled validation, double-precision reference paths, conditional "
            "float32 paths, multithread-capable compiled CPU execution, NVIDIA "
            "CUDA, and Apple Metal. Native and hybrid stages are disclosed "
            "separately."
        ),
        "Conclusions:": (
            "Conclusions: fastPLS expands the practical scale of established PLS "
            "modelling and validation while retaining explicit numerical and "
            "hardware boundaries. The GPL-3 R package uses reusable low-level "
            "components maintained in the MIT-licensed kodama-cpp project."
        ),
        "The R package includes bundled IRLBA": (
            "The R package includes bundled IRLBA code and is distributed under "
            "GPL-3. Reusable low-level CPU components are also maintained in the "
            "MIT-licensed kodama-cpp project. CUDA uses NVIDIA CUDA libraries and "
            "cuBLAS; Metal uses Apple's Metal Performance Shaders. The stage-level "
            "residency table identifies operations that remain on the host, so "
            "hybrid routes are not presented as fully device resident."
        ),
        "The software is organized as four separate layers": (
            "The software is organized as four separate layers (Fig. 1): "
            "preprocessing and response representation; the statistical PLS "
            "estimator; direction extraction by IRLBA or rSVD; and execution on "
            "compiled CPU, CUDA, or Metal backends. Regression or a classification "
            "head is applied only after the latent model has been fitted. The "
            "requested estimator is enforced throughout dispatch: a SIMPLS request "
            "fits SIMPLS, and a PLS-SVD request fits PLS-SVD. CUDA SIMPLS-LDA uses "
            "SIMPLS latent scores. If a requested estimator cannot be executed "
            "within a guarded memory path, the call fails explicitly rather than "
            "substituting another estimator. Benchmark rows record both requested "
            "and executed estimator identifiers and reject mismatches."
        ),
        "Prediction can retain compact latent factors": (
            "Prediction can retain compact latent factors rather than a full p-by-q "
            "coefficient path. Large test sets are evaluated in blocks, and each "
            "block is discarded after its scores or predictions are accumulated. "
            "This directly reduces temporary score and prediction storage without "
            "changing the fitted latent model."
        ),
        "OPLS first estimates orthogonal scores": (
            "OPLS first estimates orthogonal scores and loadings and then fits the "
            "predictive SIMPLS core [12]; the same orthogonal filter is applied at "
            "prediction. Kernel PLS constructs a centred linear, radial-basis, or "
            "polynomial kernel and uses the internal SIMPLS core [13,14]. The "
            "linear-kernel route avoids an unnecessary Gram matrix, whereas "
            "nonlinear kernels retain quadratic Gram-matrix storage. Linear, "
            "polynomial, and RBF kernels were therefore tuned and evaluated "
            "separately; this sensitivity analysis tests implemented workflows but "
            "is not presented as estimator-equivalence evidence against an "
            "independent nonlinear kernel implementation."
        ),
        "The benchmark included 12 tasks": (
            "The main biomedical benchmark included twelve tasks spanning "
            "metabolomics, NMR spectroscopy, cancer and tissue omics, CITE-seq, "
            "Retina and Tabula Muris single-cell transcriptomics, and drug response. "
            "CIFAR-100 followed its documented 50,000/10,000 split [30]. A separate "
            "ImageNet/DINOv2 stress test used 1,000,000 training and 281,167 held-out "
            "embeddings with 1,024 predictors and 1,000 classes [28,29]. It assessed "
            "only the supervised representation stage after feature extraction and "
            "was not interpreted as biomedical validation. Exact dimensions, "
            "sources, preprocessing, splits, and object checksums are reported in "
            "the Supplementary Material."
        ),
        "Table 1 makes the complete 12-task benchmark": (
            "Table 1 and Figure 2 show the complete twelve-task biomedical benchmark "
            "at one training-selected component count per dataset and PLS family. "
            "Every cell retains the predictive metric, total fitting-plus-prediction "
            "time, peak host RSS, sampled GPU memory, precision, replicate count, "
            "and status. CPU/CUDA rows were matched within a family and component "
            "count. Retina and Tabula Muris are reported separately; the previous "
            "the two single-cell resources are reported under their specific names."
        ),
        "Metal currently establishes portability": (
            "Metal established numerical portability on two classification shapes. "
            "On MetRef, CPU and Metal predictions agreed exactly for all four PLS "
            "families with argmax and LDA. On CIFAR-100, CPU/Metal agreement was "
            "1.000 for PLS-SVD and 0.9964 for SIMPLS; accuracy differed by at most "
            "0.01 percentage points. Metal PLS-SVD/argmax required 0.55 s versus "
            "0.62 s on CPU, whereas Metal SIMPLS/argmax required 2.54 s versus "
            "1.22 s. These platform-specific single-run timings establish "
            "portability, not general Metal superiority."
        ),
        "Several external implementations were restricted": (
            "The precision-matched float64 SIMPLS comparison is summarized in "
            "Table 3. fastPLS was faster than pls::simpls.fit on seven of nine "
            "datasets, including 48.3-fold on CIFAR-100 and approximately 9-fold "
            "on Retina and Tabula Muris, while accuracy differed by 0.41 percentage "
            "points or less except for the small stochastic MetRef rSVD workflow. "
            "On TCGA-BRCA and TCGA-HNSC methylation, pls was faster in absolute "
            "terms by 5 and 10 ms. Package restrictions, timeouts, memory kills, and "
            "unsupported responses remain in the machine-readable table; calling-"
            "code errors are not attributed to external packages."
        ),
        "Compiled cross-validation reduced": (
            "Across 150 completed argmax/LDA comparisons from the archived pipeline, "
            "optimized 10-fold validation required a dataset-level median of "
            "3.78-13.20 times the corresponding single-fit workflow. This ratio is "
            "reported as computational overhead, not as speedup over an unmeasured "
            "naive implementation. NMR remained near the expected repeated-fit cost "
            "because extreme multivariate products dominated; smaller or low-p "
            "classification tasks benefited more from compiled fold management and "
            "incremental scoring. Full method-level values and unsupported OPLS "
            "combinations are retained in Supplementary Table S20."
        ),
        "The fastPLS R package, benchmark workflows": (
            "The fastPLS R package, benchmark workflows, analysis scripts, and "
            "machine-readable result tables are available at "
            "https://github.com/tkcaccia/fastPLS (review-cycle package commit "
            "72e178b9e3c9510dc86c4b287d68b9c717f9fdf5). Low-level reusable C++ "
            "components are maintained at https://github.com/tkcaccia/kodama-cpp. "
            "The supplement records prepared-object checksums and exact benchmark "
            "commands; public source data remain available from their cited "
            "repositories, and restricted source data are not redistributed."
        ),
    }
    for start, text in replacements.items():
        c9.set_paragraph_text(c9.find_paragraph(document, start), text)

    backend_result = c9.find_paragraph(document, "Among the 43 completed")
    c9.set_paragraph_text(
        backend_result,
        "Among the 46 completed dataset-family cells, the fastest selected rSVD "
        "execution was CPU in 31 and CUDA in 15. CPU was generally preferable for "
        "small matrices, where transfer and kernel-launch overhead exceeded the "
        "accelerated work. CUDA became advantageous for large dense products, high "
        "component counts, or large response dimensions. Large sample count alone "
        "did not guarantee a GPU benefit. Predictive values should be compared "
        "between CPU and CUDA rows within a fixed family; differences among "
        "PLS-SVD, SIMPLS, OPLS, kernel PLS, argmax, and LDA represent workflow "
        "rather than backend-only effects."
    )
    c9.insert_after(
        backend_result,
        "Kernel sensitivity confirmed that kernel choice was a statistical rather "
        "than purely computational setting. Training-only validation selected "
        "different kernels and component counts across MetRef, CCLE, and PRISM. "
        "Polynomial kernel PLS obtained the highest outer-test accuracy on MetRef "
        "(0.990) and CCLE (0.789), whereas RBF produced the lowest PRISM RMSD "
        "(0.5450 on CUDA). Full CPU/CUDA replicate distributions and parameter "
        "settings are reported in Supplementary Tables S24-S25.",
    )

    abstract = c9.find_paragraph(document, "Results:")
    c9.set_paragraph_text(
        abstract,
        "Results: Deterministic IRLBA SIMPLS completed 117 component-level "
        "comparisons against de Jong SIMPLS without failure; all met prespecified "
        "prediction, coefficient, subspace, and selection tolerances. rSVD was "
        "evaluated separately as an approximation. Across the selected biomedical "
        "benchmark, CUDA was fastest in 15 of 46 completed family/dataset cells and "
        "CPU in 31. Precision-matched float64 SIMPLS was faster than pls::simpls.fit "
        "on seven of nine datasets. In NMR, CUDA PLS-SVD/rSVD reduced median total "
        "time from 431.23 s for the deposited reference to 1.115 s while retaining "
        "RMSD 0.000718 versus 0.000719. Across three complete ImageNet "
        "representation fits, 200-dimensional PLS scores retained median "
        "top-1/top-5 accuracy of 0.6521/0.9397, reduced exact held-out inference "
        "from 510.8 to 132.0 s, and reduced median host/GPU memory from "
        "19.2/5.9 GB to 9.9/2.9 GB relative to raw 1,024-dimensional DINOv2."
    )

    c9.set_paragraph_text(
        c9.find_paragraph(document, "The ImageNet representation experiment used"),
        "The ImageNet representation experiment used the fixed "
        "1,000,000/281,167 split and float32 DINOv2 embeddings throughout. Raw "
        "embeddings were compared with unsupervised PCA-rSVD scores and supervised "
        "label-aware PLS-SVD/rSVD scores at 50, 100, and 200 dimensions. PCA and "
        "PLS used matched CUDA rSVD settings; PLS formed the centred "
        "predictor-response cross-covariance from class sums without constructing "
        "a one-million-by-1,000 dummy response. Exact FAISS cosine kNN (k=10) used "
        "identical labels, queries, distance weighting, and query blocks. Two "
        "additional rSVD seeds repeated the complete PLS and PCA fits, giving "
        "three independently fitted randomized representations. Timings separated "
        "fitting, train/test projection, query, held-out inference, and end-to-end "
        "work. Raw DINOv2 was fitted once because it contains no randomized "
        "representation stage. CUDA IVF was evaluated against exact neighbours "
        "within the same representation, with recall@10 defined as the mean "
        "fraction of exact neighbours recovered per query."
    )

    c9.set_paragraph_text(
        c9.find_paragraph(document, "NMR represented the extreme"),
        "NMR represented the extreme multivariate-response setting (1,200 training "
        "and 321 held-out spectra; p=13,000; q=28,355). Every method used the "
        "predefined outer split and full response spectrum; predictors were centred "
        "without variance scaling, and the routine 4.6-4.8 ppm residual-water "
        "interval was set to zero in both training and held-out predictor matrices. "
        "Five repeated training-only 80/20 splits evaluated 10, 25, 50, 75, 100, "
        "125, 150, 165, 175, 200, 250, and 300 components. SIMPLS median validation "
        "RMSD reached an interior minimum at 100 and deteriorated thereafter. "
        "PLS-SVD selected the lower tested boundary of 10 components, and its "
        "median RMSD increased monotonically over the evaluated grid. Thus, Table 1 "
        "uses family-specific selections, whereas Figure 3 compares all fastPLS "
        "and deposited-reference implementations at the common prespecified "
        "100-component operating point. The latter isolates implementation speed, "
        "memory, and prediction behavior and is not presented as the PLS-SVD "
        "optimum."
    )

    table_caption = c9.find_paragraph(document, "Table 1.")
    c9.set_paragraph_text(
        table_caption,
        "Table 1. Twelve-task biomedical benchmark at the component count selected "
        "from training data. Each cell shows the fastest completed CPU/CUDA rSVD "
        "backend within a fixed PLS family, effective component count, outer-test "
        "metric, median total fitting-plus-prediction time (IQR when repeated), peak "
        "host RSS (H), sampled GPU memory (G), precision, and completed runs. "
        "Comparisons across model-family columns are workflow comparisons."
    )

    nmr_caption = c9.find_paragraph(document, "Figure 2. NMR")
    c9.set_paragraph_text(
        nmr_caption,
        nmr_caption.text.replace("Figure 2.", "Figure 3.").replace(
            "at the selected 100-component model",
            "at the common 100-component implementation-comparison point",
        ),
    )
    image_caption = c9.find_paragraph(document, "Figure 3. Exploratory matched ImageNet")
    c9.set_paragraph_text(image_caption, image_caption.text.replace("Figure 3.", "Figure 4."))
    image_result = c9.find_paragraph(document, "ImageNet/DINOv2 was used")
    c9.set_paragraph_text(
        image_result,
        "ImageNet/DINOv2 was used for a million-sample computational stress test "
        "after feature extraction and an exploratory evaluation of PLS as "
        "supervised feature reduction, not as biomedical validation. Exact CUDA "
        "cosine kNN on raw 1,024-dimensional embeddings achieved top-1/top-5 "
        "accuracy of 0.6556/0.9392 with 510.8 s median query time, 19.2 GB peak "
        "host RSS, and 5.9 GB sampled GPU memory. Across three independently fitted "
        "randomized representations, 200-dimensional PLS scores achieved median "
        "top-1/top-5 accuracy of 0.6521/0.9397 (top-1 range "
        "0.6516-0.6521), 114.2 s transformation time, and 132.0 s held-out "
        "projection-plus-query time. Median host/GPU memory was 9.9/2.9 GB. The "
        "dimension-matched PCA control achieved median top-1/top-5 accuracy of "
        "0.6430/0.9383 and required 233.3 s for transformation. Thus PLS retained "
        "raw top-1 accuracy within 0.35 percentage points, slightly exceeded raw "
        "top-5 accuracy, compressed the representation 5.12-fold, and reduced "
        "held-out inference 3.9-fold. Table 2 and Figure 4 show the representative "
        "fixed-seed path; independent-fit ranges are reported in Supplementary "
        "Table S23. These natural-image results remain exploratory."
    )
    c9.set_paragraph_text(
        c9.find_paragraph(document, "Figure 4. Exploratory matched ImageNet"),
        "Figure 4. Exploratory matched ImageNet/DINOv2 retrieval. (A) Top-1 and "
        "top-5 accuracy for raw embeddings and 50-, 100-, and 200-dimensional PCA "
        "and PLS representations. (B) Held-out projection plus exact FAISS query "
        "time. (C) End-to-end representation fitting, train/test projection, and "
        "query time. (D) peak host RSS and sampled GPU memory. The displayed path "
        "uses seed 123; two independent seeded PLS/PCA fits are summarized in "
        "Supplementary Table S23."
    )
    c9.set_paragraph_text(
        c9.find_paragraph(document, "In a fixed-score validation"),
        "In a fixed-score validation of the revised LDA path, float32 CPU and CUDA "
        "predictions agreed exactly across MetRef, CIFAR-100, Retina, and Tabula "
        "Muris at fixed component counts, with no factorization failures. This "
        "classifier-level check does not establish end-to-end float32 equivalence "
        "for every model shape. In the matched three-repetition full-NMR study, "
        "float32 PLS-SVD remained close to float64 in RMSD but was much slower, and "
        "float32 SIMPLS, kernel PLS, and CUDA OPLS had materially larger RMSD; CPU "
        "OPLS exceeded the 1,200-s limit. Float32 is therefore reported by model, "
        "backend, dataset, and status rather than pooled into a universal claim."
    )

    external_anchor = c9.find_paragraph(
        document, "The precision-matched float64 SIMPLS comparison"
    )
    insert_table_after(
        document,
        external_anchor,
        "Table 3. Precision-matched float64 SIMPLS software comparison. Runtime is "
        "median total fitting plus prediction in milliseconds (IQR). The reference "
        "is pls::simpls.fit with the same split, preprocessing, component count, "
        "dummy response, and argmax decoder.",
        ["Dataset", "k", "fastPLS", "fastPLS ms", "pls ms", "Speedup", "fastPLS metric", "pls metric"],
        external_rows(),
        [0.9, 0.32, 0.82, 0.72, 0.72, 0.55, 0.77, 0.7],
        5.8,
    )

    c9.set_paragraph_text(
        c9.find_paragraph(document, "[5] Cacciatore"),
        "[5] Cacciatore S, Wium M, Licari C, et al. Inflammatory metabolic "
        "profile of South African patients with prostate cancer. Cancer Metab. "
        "2021;9:29. https://doi.org/10.1186/s40170-021-00265-6.",
    )
    correct_lda_content(document)
    remove_obsolete_references_and_renumber(document)
    document.core_properties.title = "fastPLS CMPB manuscript - reviewer revision cycle 13"
    document.save(MAIN_OUT)


def remove_obsolete_references_and_renumber(document):
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        if text.startswith("[17] Shao Z") or text.startswith("[20] Cover T"):
            c9.delete_paragraph(paragraph)

    def renumber_text(text):
        def map_number(value):
            number = int(value)
            if number < 17:
                return str(number)
            if 18 <= number < 20:
                return str(number - 1)
            if number >= 21:
                return str(number - 2)
            return str(number)

        def replace_match(match):
            content = match.group(1)
            return "[" + re.sub(r"\d+", lambda number: map_number(number.group(0)), content) + "]"

        text = re.sub(r"\[([0-9,\-\s]+)\]", replace_match, text)
        return text

    for paragraph in document.paragraphs:
        updated = renumber_text(paragraph.text)
        if updated != paragraph.text:
            c9.set_paragraph_text(paragraph, updated)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    updated = renumber_text(paragraph.text)
                    if updated != paragraph.text:
                        c9.set_paragraph_text(paragraph, updated)


def append_supplement_tables(document):
    document.add_page_break()
    document.add_heading("S17. Precision-matched external software comparison", level=1)
    document.add_paragraph(
        "The primary software comparison used float64 inputs for fastPLS and every "
        "external method. The compact table below is restricted to estimator-matched "
        "SIMPLS fits. The complete archive contains all evaluated packages, versions, "
        "adapters, exact errors, restrictions, and memory measurements."
    )
    document.add_paragraph(
        "Table S19. Estimator-matched fastPLS versus pls::simpls.fit. Times are "
        "median (IQR) milliseconds from three isolated runs.",
        style="Caption",
    )
    c9.add_table(
        document,
        ["Dataset", "k", "fastPLS", "fastPLS ms", "pls ms", "Speedup", "fastPLS metric", "pls metric"],
        external_rows(),
        [0.9, 0.32, 0.82, 0.72, 0.72, 0.55, 0.77, 0.7],
        font_size=5.8,
    )

    document.add_heading("S18. Quantitative cross-validation overhead", level=1)
    document.add_paragraph(
        "The archived 10-fold benchmark compared one optimized validation call with "
        "the corresponding one-fit public workflow. It did not time a separate naive "
        "R fold loop; ratios are therefore computational overhead relative to one fit, "
        "not speedups over an unmeasured implementation. Nearest-neighbour "
        "classification rows were excluded because that head is no longer part of "
        "the package."
    )
    cv = pd.read_csv(EVIDENCE / "cv10_dataset_summary.csv")
    cv = cv[~cv["dataset"].eq("singlecell")].copy()
    rows = [
        (
            DATASET_LABELS.get(row["dataset"], row["dataset"]),
            int(row["comparisons"]),
            f'{row["cv_over_fit_median"]:.2f}',
            f'{row["cv_over_fit_min"]:.2f}',
            f'{row["cv_over_fit_max"]:.2f}',
        )
        for _, row in cv.iterrows()
    ]
    document.add_paragraph(
        "Table S20. Optimized 10-fold validation time divided by one fit-plus-"
        "prediction time across completed model/backend combinations.",
        style="Caption",
    )
    c9.add_table(
        document,
        ["Dataset", "Comparisons", "Median ratio", "Minimum", "Maximum"],
        rows,
        [1.55, 0.75, 0.9, 0.75, 0.75],
        font_size=6.5,
    )

    document.add_heading("S19. Float32 and backend validation boundaries", level=1)
    precision = pd.read_csv(EVIDENCE / "float32_float64_pairs.csv")
    rows = []
    for _, row in precision.iterrows():
        rows.append(
            (
                DATASET_LABELS.get(row["dataset"], row["dataset"]),
                row["method"].replace("kernelpls", "kernel PLS").replace("plssvd", "PLS-SVD"),
                row["backend"].upper(),
                f'{row["f64_metric"]:.4g}',
                f'{row["f32_metric"]:.4g}',
                f'{row["metric_difference_f32_minus_f64"]:.3g}',
                f'{row["f64_time_sec"]:.3f}/{row["f32_time_sec"]:.3f}',
                f'{row["f64_input_mb"]:.1f}/{row["f32_input_mb"]:.1f}',
                f'{row["f64_host_rss_mb"]:.0f}/{row["f32_host_rss_mb"]:.0f}',
                (
                    f'{row["f64_gpu_mb"]:.0f}/{row["f32_gpu_mb"]:.0f}'
                    if finite(row["f64_gpu_mb"]) and finite(row["f32_gpu_mb"])
                    else "—"
                ),
            )
        )
    document.add_paragraph(
        "Table S21. Matched float64/float32 results on MetRef classification and "
        "PRISM multivariate regression. Time, input storage, host RSS, and GPU "
        "memory are shown as float64/float32. These compact tests complement the "
        "negative full-NMR precision benchmark in Table S7.",
        style="Caption",
    )
    c9.add_table(
        document,
        ["Dataset", "Method", "Backend", "f64 metric", "f32 metric", "Difference", "Time s", "Input MB", "RSS MB", "GPU MB"],
        rows,
        [0.65, 0.7, 0.48, 0.55, 0.55, 0.55, 0.7, 0.65, 0.65, 0.62],
        font_size=5.2,
    )

    document.add_heading("S20. CPU and Metal reproducibility", level=1)
    document.add_paragraph(
        "MetRef used 773/100 train/test observations, 375 predictors, 22 classes, "
        "20 components, and float32 input. CIFAR-100 used the fixed 50,000/10,000 "
        "split, 768 predictors, and float32 input. Timings are platform-specific "
        "single runs and are not interpreted as a general hardware ranking."
    )
    metref = pd.read_csv(EVIDENCE / "metal_metref_agreement.csv")
    cifar = pd.read_csv(EVIDENCE / "cifar100_cpu_metal_prediction_agreement.csv")
    rows = []
    for _, row in metref.iterrows():
        rows.append(
            (
                "MetRef",
                row["method"].replace("kernelpls", "kernel PLS").replace("plssvd", "PLS-SVD"),
                row["classifier"],
                "20",
                f'{row["metric_cpu1"]:.4f}',
                f'{row["metric_metal"]:.4f}',
                f'{row["agreement_cpu1_metal"]:.4f}',
            )
        )
    for _, row in cifar.iterrows():
        rows.append(
            (
                "CIFAR-100",
                row["method"].replace("plssvd", "PLS-SVD"),
                row["classifier"],
                str(int(row["ncomp"])),
                f'{row["cpu_accuracy"]:.4f}',
                f'{row["metal_accuracy"]:.4f}',
                f'{row["cpu_metal_prediction_agreement"]:.4f}',
            )
        )
    document.add_paragraph(
        "Table S22. Matched CPU/Metal predictive agreement.", style="Caption"
    )
    c9.add_table(
        document,
        ["Dataset", "Method", "Head", "k", "CPU metric", "Metal metric", "Agreement"],
        rows,
        [0.85, 0.85, 0.55, 0.35, 0.75, 0.75, 0.75],
        font_size=6.0,
    )

    runtime = pd.read_csv(EVIDENCE / "cifar100_cpu_metal_runtime_accuracy_memory.csv")
    document.add_paragraph(
        "On CIFAR-100, Metal PLS-SVD/argmax required 0.550 s versus 0.621 s on "
        "CPU, whereas Metal SIMPLS/argmax required 2.540 s versus 1.215 s. The "
        "corresponding peak process RSS values were 787 versus 783 MB for PLS-SVD "
        "and 640 versus 738 MB for SIMPLS. On MetRef, CPU PLS-SVD and SIMPLS "
        "argmax required 0.100 and 0.089 s, whereas Metal required 0.089 and "
        "0.176 s, respectively. These small absolute times preclude a stable "
        "general speed ranking."
    )

    if EVIDENCE.joinpath("imagenet_repeated_fit_summary.csv").exists():
        repeated = pd.read_csv(EVIDENCE / "imagenet_repeated_fit_summary.csv")
        document.add_heading("S21. Repeated ImageNet representation fits", level=1)
        document.add_paragraph(
            "Two additional fixed rSVD seeds repeated the complete PLS and PCA "
            "representation fits. Accuracy variation therefore reflects the "
            "randomized representation rather than repeated search on one fit."
        )
        rows = []
        for _, row in repeated.iterrows():
            rows.append(
                (
                    row["representation"],
                    int(row["n_features"]),
                    int(row["seed"]),
                    f'{row["top1_accuracy"]:.4f}',
                    f'{row["top5_accuracy"]:.4f}',
                    f'{row["transformation_time_sec"]:.1f}',
                    f'{row["inference_time_sec"]:.1f}',
                    f'{row["peak_host_rss_mb"]:.0f}',
                    f'{row["peak_gpu_mem_mb"]:.0f}',
                )
            )
        document.add_paragraph(
            "Table S23. Repeated full ImageNet/DINOv2 representation fits.",
            style="Caption",
        )
        c9.add_table(
            document,
            ["Representation", "Dim.", "Seed", "Top-1", "Top-5", "Transform s", "Inference s", "Host MB", "GPU MB"],
            rows,
            [0.95, 0.4, 0.45, 0.55, 0.55, 0.75, 0.75, 0.65, 0.65],
            font_size=5.7,
        )

    if KERNEL_RESULTS.joinpath("kernel_sensitivity_summary.csv").exists():
        document.add_heading("S22. Kernel sensitivity analysis", level=1)
        document.add_paragraph(
            "Kernel and component settings were selected by five-fold validation "
            "on the training partition for MetRef, CCLE, and PRISM. The selected "
            "linear, polynomial, and RBF workflows were then evaluated on the same "
            "untouched outer test partition with three isolated CPU and CUDA runs. "
            "The analysis treats kernel choice as a model-selection decision rather "
            "than an implementation-only comparison."
        )
        selected = pd.read_csv(KERNEL_RESULTS / "kernel_sensitivity_selected.csv")
        selection_rows = []
        for _, row in selected.iterrows():
            parameter = "—"
            if row["kernel"] == "poly":
                parameter = (
                    f'gamma={row["gamma"]:.4g}; degree={int(row["degree"])}; '
                    f'coef0={row["coef0"]:.3g}'
                )
            elif row["kernel"] == "rbf":
                parameter = f'gamma={row["gamma"]:.4g}'
            selection_rows.append(
                (
                    DATASET_LABELS.get(row["dataset"], row["dataset"]),
                    row["kernel"],
                    int(row["ncomp"]),
                    parameter,
                    row["metric_name"],
                    f'{row["metric_value"]:.4f}',
                    f'{row["tuning_time_sec"]:.2f}',
                )
            )
        document.add_paragraph(
            "Table S24. Training-only kernel and component selection.",
            style="Caption",
        )
        c9.add_table(
            document,
            ["Dataset", "Kernel", "k", "Selected parameters", "Metric", "CV value", "Tuning s"],
            selection_rows,
            [0.78, 0.55, 0.32, 1.75, 0.55, 0.65, 0.65],
            font_size=5.8,
        )

        summary = pd.read_csv(KERNEL_RESULTS / "kernel_sensitivity_summary.csv")
        result_rows = []
        for _, row in summary.iterrows():
            result_rows.append(
                (
                    DATASET_LABELS.get(row["dataset"], row["dataset"]),
                    row["kernel"],
                    row["backend"].upper(),
                    int(row["selected_ncomp"]),
                    row["metric_name"],
                    f'{row["metric_median"]:.4f}',
                    f'{row["total_time_median_sec"]:.3f}',
                    f'{row["peak_host_rss_median_mb"]:.0f}',
                    (
                        f'{row["peak_gpu_mem_median_mb"]:.0f}'
                        if finite(row["peak_gpu_mem_median_mb"])
                        else "—"
                    ),
                    int(row["successful_reps"]),
                )
            )
        document.add_paragraph(
            "Table S25. Outer-test kernel sensitivity results. Time and memory are "
            "medians across successful isolated runs.",
            style="Caption",
        )
        c9.add_table(
            document,
            ["Dataset", "Kernel", "Backend", "k", "Metric", "Value", "Time s", "Host MB", "GPU MB", "Runs"],
            result_rows,
            [0.7, 0.5, 0.55, 0.3, 0.5, 0.55, 0.6, 0.65, 0.6, 0.4],
            font_size=5.4,
        )
        document.add_paragraph(
            "Polynomial kernel PLS had the highest outer-test accuracy on MetRef "
            "(0.990) and CCLE (0.789), whereas RBF had the lowest PRISM RMSD "
            "(0.5450 on CUDA). CPU was faster for the small classification tasks; "
            "CUDA was faster for the selected PRISM RBF workflow. Predictive "
            "differences between CPU and CUDA at fixed selected settings were "
            "small relative to differences among kernels."
        )
        for filename, caption in (
            (
                "kernel_sensitivity_classification.png",
                "Figure S21. Kernel sensitivity for MetRef and CCLE classification.",
            ),
            (
                "kernel_sensitivity_regression.png",
                "Figure S22. Kernel sensitivity for PRISM multivariate regression.",
            ),
        ):
            image_path = KERNEL_RESULTS / "plots" / filename
            if image_path.exists():
                paragraph = document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run().add_picture(str(image_path), width=Inches(6.2))
                document.add_paragraph(caption, style="Caption")


def revise_supplement():
    document = Document(SUPP_SOURCE)
    replace_dataset_table(document)
    replace_supplement_method_tables(document)
    s8_last = c9.find_paragraph(document, "The completed fixed-score LDA")
    s9 = c9.insert_after(s8_last, "S9. ImageNet computational role", style="Heading 1")
    image_role = c9.find_paragraph(document, "The scientific role of this experiment")
    image_role._p.getparent().remove(image_role._p)
    s9._p.addnext(image_role._p)
    duplicate = c9.find_paragraph(document, "S6. Matched precision benchmark")
    c9.set_paragraph_text(duplicate, "S12.1 Matched precision benchmark")
    duplicate.style = "Heading 2"
    c9.set_paragraph_text(
        c9.find_paragraph(document, "The completed fixed-score LDA"),
        "The completed fixed-score LDA validation used MetRef, CIFAR-100, Retina, "
        "and Tabula Muris. CPU and CUDA float32 predictions were compared at fixed "
        "component counts; classifier-level agreement is reported separately from "
        "upstream rSVD variation."
    )
    scope_paragraph = c9.find_paragraph(
        document, "The double-precision CPU backend is the broadest"
    )
    c9.insert_after(
        scope_paragraph,
        "Estimator identity is a dispatch invariant. A requested SIMPLS model is "
        "fitted with the SIMPLS core on CPU, CUDA, and Metal; the CUDA fused "
        "SIMPLS-LDA route passes the SIMPLS method identifier to the native core "
        "and trains LDA from those SIMPLS scores. Label-aware PLS-SVD remains "
        "available only through an explicit PLS-SVD request. When the guarded "
        "CUDA SIMPLS response representation is too large, execution stops with "
        "an informative error instead of replacing the estimator. Benchmark code "
        "checks requested against executed estimator identifiers and records an "
        "estimator-mismatch error if they differ."
    )
    nmr_protocol = c9.find_paragraph(
        document, "Routine NMR spectral preprocessing and component selection."
    )
    c9.set_paragraph_text(
        nmr_protocol,
        "Routine NMR spectral preprocessing and component selection. The task "
        "comprised 1,200 training and 321 held-out spectra, 13,000 predictors, and "
        "28,355 numeric responses. A shared loader verified identical columns, "
        "outer-split dimensions, checksum, and matrix signatures. The 4.6-4.8 ppm "
        "residual-water interval was set to zero in Xtrain and Xtest before inner "
        "splitting or fitting; Ytrain and Ytest were unchanged. Predictors were "
        "centred without variance scaling. Five repeated 80/20 training-only splits "
        "used seeds 123, 456, 789, 1011, and 2027 and the common 10-300 component "
        "grid. SIMPLS reached an interior median RMSD minimum at 100 components "
        "(0.0008917) and deteriorated thereafter. PLS-SVD selected 10 components "
        "(median RMSD 0.0015402), with median error increasing to 0.0016553 at 100 "
        "and 0.0016580 at 300. The family-specific results prevent a universal "
        "NMR component claim; the separate 100-component reference comparison "
        "holds model size fixed to isolate computational behavior."
    )
    nmr_response_caption = c9.find_paragraph(
        document,
        "Table S13. Response-wise held-out RMSD",
    )
    c9.set_paragraph_text(
        nmr_response_caption,
        nmr_response_caption.text.replace(
            "at the selected 100-component model",
            "at the common 100-component implementation-comparison point",
        ),
    )
    image_protocol = c9.find_paragraph(
        document,
        "The external retrieval benchmark used the same fixed float32",
    )
    c9.set_paragraph_text(
        image_protocol,
        image_protocol.text.replace(
            "Representation fits were single runs, exact index/query timings were "
            "repeated three times, and IVF was run once.",
            "The complete PLS and PCA representation fits were repeated with three "
            "fixed rSVD seeds; exact index/query timings were repeated three times "
            "per fitted representation, and IVF was run once.",
        ),
    )
    image_path_caption = c9.find_paragraph(
        document,
        "Figure S20. Full exploratory ImageNet component path",
    )
    c9.set_paragraph_text(
        image_path_caption,
        image_path_caption.text.replace(
            "Exact search used three repeated index/query measurements from one "
            "fitted representation; representation fitting and IVF were single runs.",
            "Exact search used three repeated index/query measurements for each of "
            "three independently fitted randomized PLS and PCA representations; "
            "IVF was run once.",
        ),
    )
    simpls_selection = pd.read_csv(
        ROOT
        / "benchmark_results"
        / "review_nmr_extended_selection_20260725"
        / "nmr_component_selection_summary.csv"
    )
    plssvd_selection = pd.read_csv(
        EVIDENCE
        / "nmr_plssvd_selection"
        / "nmr_component_selection_summary.csv"
    )
    selection_rows = []
    for ncomp in simpls_selection["ncomp"]:
        simpls_row = simpls_selection[simpls_selection["ncomp"].eq(ncomp)].iloc[0]
        plssvd_row = plssvd_selection[
            plssvd_selection["ncomp"].eq(ncomp)
        ].iloc[0]
        selection_rows.append(
            (
                int(ncomp),
                f'{simpls_row["RMSD_median"]:.7f}',
                f'{simpls_row["RMSD_q25"]:.7f}-{simpls_row["RMSD_q75"]:.7f}',
                f'{plssvd_row["RMSD_median"]:.7f}',
                f'{plssvd_row["RMSD_q25"]:.7f}-{plssvd_row["RMSD_q75"]:.7f}',
            )
        )
    insert_table_after(
        document,
        nmr_protocol,
        "Table S12b. Family-specific repeated training-only NMR selection. Values "
        "are median and interquartile-range validation RMSD over five fixed splits.",
        ["Components", "SIMPLS median", "SIMPLS IQR", "PLS-SVD median", "PLS-SVD IQR"],
        selection_rows,
        [0.7, 1.05, 1.35, 1.05, 1.35],
        6.1,
    )
    c9.set_paragraph_text(
        c9.find_paragraph(
            document, "Table S12. Repeated training-only NMR component selection."
        ),
        "Table S12a. Repeated training-only SIMPLS NMR component selection. RMSD "
        "and Q² summaries are across five 80/20 inner splits; the outer test set "
        "was not accessed.",
    )
    c9.set_paragraph_text(
        c9.find_paragraph(document, "Figure S18. Training-only NMR component selection."),
        "Figure S18. Training-only SIMPLS NMR component selection. Grey lines show "
        "the five inner splits, the blue line and ribbon show the median and "
        "interquartile range, and the dashed line identifies the selected interior "
        "minimum at 100 components.",
    )
    c9.set_paragraph_text(
        c9.find_paragraph(document, "References are numbered"),
        "References are numbered as in the main manuscript. Dataset references "
        "document CIFAR-100, Retina, Tabula Muris, ImageNet/DINOv2, and the relevance "
        "of large foundation-model embedding matrices to computational pathology."
    )
    c9.set_paragraph_text(
        c9.find_paragraph(document, "Machine-readable files include imagenet"),
        "Machine-readable files include imagenet_faiss_matched_raw.csv, "
        "imagenet_faiss_matched_summary.csv, and the per-seed repeated-fit tables. "
        "Archived exact-neighbour matrices permit independent recomputation of IVF "
        "recall."
    )
    append_supplement_tables(document)
    correct_lda_content(document, supplement=True)
    remove_obsolete_references_and_renumber(document)
    document.core_properties.title = "fastPLS CMPB supplement - reviewer revision cycle 13"
    document.save(SUPP_OUT)


def write_response():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    document.styles["Normal"].font.name = "Arial"
    document.styles["Normal"].font.size = Pt(9.5)
    document.styles["Normal"].paragraph_format.space_after = Pt(5)
    document.styles["Heading 1"].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    document.add_heading("Response to reviewer comments", level=1)
    document.add_paragraph(
        "The manuscript and Supplementary Material were revised to separate "
        "deterministic estimator preservation, randomized approximation, complete "
        "workflow comparisons, precision effects, and hardware residency."
    )

    responses = [
        (
            "1. Estimator preservation",
            "We now restrict the preservation claim to deterministic IRLBA SIMPLS. "
            "The prespecified suite covers regression/classification, p<n and p>n, "
            "low/high-rank responses, ill-conditioning, exact rank deficiency, four "
            "real datasets, 117 component comparisons, and 24 fixed-fold selections. "
            "All deterministic tolerances passed with no failure. rSVD results are "
            "reported separately and explicitly demonstrate approximation.",
        ),
        (
            "2. OPLS and kernel PLS scope",
            "We narrowed the claim. OPLS and nonlinear kernel PLS are described as "
            "implemented workflows that reuse the accelerated SIMPLS core, not as "
            "new estimators with established equivalence. Linear-kernel shortcut and "
            "nonlinear Gram-matrix routes are distinguished. Training-only kernel "
            "selection and three-run outer-test comparisons of linear, polynomial, "
            "and RBF kernels are now reported for MetRef, CCLE, and PRISM.",
        ),
        (
            "3. Complete biomedical benchmark",
            "Table 1 and Figure 2 now show all twelve biomedical tasks at the "
            "training-selected component count, with predictive metric, total time, "
            "host/GPU memory, precision, replicate count, and status. Retina and "
            "Tabula Muris replace the ambiguous SingleCell label. ImageNet is kept "
            "outside this estimator table under its matched retrieval protocol.",
        ),
        (
            "4. External software",
            "A float64, estimator-matched SIMPLS table is now in the main text and "
            "Supplement. It reports medians, IQRs, predictive metrics, and speedups. "
            "Package restrictions and failures remain machine readable.",
        ),
        (
            "5. Metal and multicore scope",
            "The future-tense statement was removed. CPU/Metal agreement and timing "
            "are reported for MetRef and CIFAR-100. We explicitly present Metal as a "
            "portability result, identify host stages, and avoid cross-machine claims "
            "of general hardware superiority.",
        ),
        (
            "6. Float32 validation",
            "The precision claim is now conditional. Matched float32/float64 tables "
            "cover all four model families on classification and multivariate "
            "regression, CPU and CUDA, with a separate Metal audit. The adverse NMR "
            "routes and timeouts remain prominent.",
        ),
        (
            "7. NMR component selection",
            "SIMPLS and PLS-SVD were each evaluated to 300 components over five "
            "training-only splits. SIMPLS found an interior median RMSD minimum at "
            "100, whereas PLS-SVD selected 10 and worsened over the tested grid. "
            "The selected-point benchmark now uses family-specific counts; the "
            "100-component deposited-reference comparison is explicitly identified "
            "as a common operating point, not a universal optimum.",
        ),
        (
            "8. ImageNet uncertainty",
            "Two additional seeded PLS and PCA representation fits were run in "
            "addition to the original fit. The revised table separates fit, "
            "projection, query, inference, and end-to-end time and retains the "
            "exploratory/non-biomedical qualification.",
        ),
        (
            "9. Cross-validation",
            "The supplement now reports quantitative 10-fold overhead across all "
            "completed datasets and methods. We state exactly that this is a ratio to "
            "one fit, not a speedup over an unmeasured naive R loop.",
        ),
        (
            "10. Ablation and complexity",
            "Estimator-preserving changes are separated from rSVD workspace reuse. "
            "The complexity/storage and de Jong mapping tables identify each cached "
            "quantity, implicit product, compact output, and solver-dependent term. "
            "The limited execution ablation is labelled as such rather than generalized.",
        ),
        (
            "11. Reproducibility",
            "Retina and Tabula Muris now have distinct names, dimensions, sources, "
            "prepared-object checksums, and preprocessing descriptions. The code and "
            "data statement no longer promises future manifests as if they already "
            "existed.",
        ),
        (
            "12. Citations and FlashSVD",
            "The prostate-cancer DOI was corrected. FlashSVD attribution was removed "
            "because the implemented operation is blocked low-rank matrix prediction, "
            "not the cited neural-network inference method.",
        ),
        (
            "13. LDA equation and obsolete classifier remnants",
            "Section 2.5 now gives the implemented Fisher LDA discriminant based on "
            "the pooled within-class covariance, Cholesky solves, class means, and "
            "class priors. The unrelated temperature-smoothed score equations were "
            "removed from the main manuscript and Supplement. The obsolete nearest-"
            "neighbour reference was also removed and all subsequent citations were "
            "renumbered consistently.",
        ),
        (
            "14. Estimator substitution",
            "Both CUDA branches that replaced requested SIMPLS with label-aware "
            "PLS-SVD were removed. SIMPLS now executes SIMPLS or fails explicitly "
            "at the memory guard. Native tests confirm that fused CUDA SIMPLS-LDA "
            "calls the SIMPLS core and uses SIMPLS scores. Benchmarks reject any "
            "requested/executed mismatch. We also excluded eight archived "
            "substituted rows and rebuilt the external table from genuine SIMPLS "
            "fits for all nine datasets."
        ),
    ]
    for heading, response in responses:
        document.add_heading(heading, level=2)
        document.add_paragraph(response)
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    write_response()
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
