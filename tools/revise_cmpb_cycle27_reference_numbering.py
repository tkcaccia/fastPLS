#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle26"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle27"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle26_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle26_0.99.6_20260725.docx"
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle26_20260725.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle27_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle27_0.99.6_20260725.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle27_20260725.docx"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def insert_after(paragraph, text):
    node = OxmlElement("w:p")
    paragraph._p.addnext(node)
    new_paragraph = Paragraph(node, paragraph._parent)
    new_paragraph.style = paragraph.style
    new_paragraph.add_run(text)
    return new_paragraph


def revise_main():
    document = Document(MAIN_SOURCE)
    introduction = find_paragraph(
        document, "This matrix regime is increasingly relevant"
    )
    introduction.text = (
        "This matrix regime is increasingly relevant after foundation-model "
        "feature extraction. Computational pathology models such as UNI and "
        "Prov-GigaPath transform very large collections of whole-slide-image "
        "tiles into high-dimensional representations for supervised downstream "
        "analyses [31,32]. Although ImageNet is not biomedical, a large matrix "
        "of precomputed DINOv2 embeddings [28,29] provides a reproducible test "
        "of the same post-extraction computational stage: fitting and predicting "
        "from many observations represented by dense foundation-model features. "
        "We use it to assess scalability of this stage, not to infer that "
        "natural-image accuracy transfers to pathology or to claim clinical "
        "validity."
    )

    paragraph = find_paragraph(
        document, "The main biomedical benchmark included twelve tasks"
    )
    paragraph.text = (
        "The main biomedical benchmark included twelve tasks spanning MetRef "
        "metabolomics [20], NMR spectroscopy [7], CBMC CITE-seq [21], GTEx [22], "
        "CCLE [23], TCGA cancer cohorts [24], Retina [25] and Tabula Muris [26] "
        "single-cell transcriptomics, and PRISM drug response [27]. CIFAR-100 "
        "followed its documented 50,000/10,000 split [30]. A separate "
        "exploratory ImageNet/DINOv2 stress test used a pooled archive of "
        "1,281,167 precomputed 1,024-dimensional embeddings from 1,000 classes "
        "[28,29]. The archive did not retain an authoritative canonical "
        "train/validation flag. With seed 123, 1,000,000 row indices were sampled "
        "without replacement for development training and the complementary "
        "281,167 rows formed a development holdout. This is not the standard "
        "ImageNet train/validation split. It assessed only the supervised "
        "representation stage after feature extraction and was not interpreted "
        "as biomedical or external predictive validation. Exact dimensions, "
        "checksums, split construction, and limitations are reported in the "
        "Supplementary Material."
    )

    external = find_paragraph(
        document, "External comparisons use independent R implementations"
    )
    external.text = external.text.replace(
        "including functions from pls,",
        "including functions from pls [19],",
    )

    references = {
        25: (
            "[25] Macosko EZ, Basu A, Satija R, et al. Highly parallel "
            "genome-wide expression profiling of individual cells using "
            "nanoliter droplets. Cell. 2015;161:1202-1214. "
            "https://doi.org/10.1016/j.cell.2015.05.002."
        ),
        26: (
            "[26] Tabula Muris Consortium. Single-cell transcriptomics of 20 "
            "mouse organs creates a Tabula Muris. Nature. 2018;562:367-372. "
            "https://doi.org/10.1038/s41586-018-0590-4."
        ),
        27: (
            "[27] Corsello SM, et al. Discovering the anticancer potential of "
            "non-oncology drugs by systematic viability profiling. Nat Cancer. "
            "2020;1:235-248. https://doi.org/10.1038/s43018-019-0018-6."
        ),
        28: (
            "[28] Deng J, et al. ImageNet: a large-scale hierarchical image "
            "database. Proc IEEE CVPR. 2009:248-255. "
            "https://doi.org/10.1109/CVPR.2009.5206848."
        ),
        29: (
            "[29] Oquab M, et al. DINOv2: learning robust visual features "
            "without supervision. Trans Mach Learn Res. 2024."
        ),
        30: (
            "[30] Krizhevsky A, Hinton G. Learning multiple layers of features "
            "from tiny images. Technical report. University of Toronto; 2009. "
            "https://www.cs.toronto.edu/~kriz/learning-features-2009-TR.pdf."
        ),
        31: (
            "[31] Chen RJ, Ding T, Lu MY, et al. Towards a general-purpose "
            "foundation model for computational pathology. Nat Med. "
            "2024;30:850-862. https://doi.org/10.1038/s41591-024-02857-3."
        ),
    }
    for number, text in references.items():
        reference = find_paragraph(document, f"[{number}]")
        reference.text = text
    last_reference = find_paragraph(document, "[31]")
    insert_after(
        last_reference,
        "[32] Xu H, Usuyama N, Bagga J, et al. A whole-slide foundation "
        "model for digital pathology from real-world data. Nature. "
        "2024;630:181-188. https://doi.org/10.1038/s41586-024-07441-w.",
    )
    document.core_properties.title = (
        "fastPLS CMPB manuscript - corrected reference numbering"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)
    images = find_paragraph(
        document, "CIFAR-100 and ImageNet were evaluated"
    )
    images.text = images.text.replace(
        "partition [29]", "partition [30]"
    ).replace(
        "labels [27,28]", "labels [28,29]"
    )
    pathology = find_paragraph(
        document, "ImageNet is included as a large-scale"
    )
    pathology.text = pathology.text.replace("[30,31]", "[31,32]")
    pathology_role = find_paragraph(
        document, "The scientific role of this experiment is computational."
    )
    pathology_role.text = pathology_role.text.replace("[30,31]", "[31,32]")
    note = find_paragraph(
        document, "References are numbered as in the main manuscript."
    )
    note.text = (
        "References are numbered as in the main manuscript. The final mapping "
        "is Retina [25], Tabula Muris [26], PRISM [27], ImageNet [28], DINOv2 "
        "[29], CIFAR-100 [30], UNI [31], and Prov-GigaPath [32]. These sources "
        "document the single-cell benchmarks and the biomedical relevance of "
        "large foundation-model embedding matrices."
    )
    document.core_properties.title = (
        "fastPLS CMPB supplement - corrected reference numbering"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    document.add_heading("28. Reference numbering was inconsistent", level=1)
    document.add_paragraph(
        "Reviewer comment: Two references were numbered 17, citation 32 was "
        "absent, and citations for ImageNet, DINOv2, CIFAR-100, and pathology "
        "therefore resolved to incorrect sources."
    )
    document.add_paragraph(
        "Response: Corrected. The unused FlashSVD and candidate-nearest-neighbour "
        "references had been removed without consistently updating every "
        "downstream citation. We also replaced the unrelated single-cell source "
        "with separate Retina and Tabula Muris references. We rebuilt the "
        "bibliography as a continuous 1-32 sequence and synchronized the main "
        "manuscript and supplement. The final mapping is Retina [25], Tabula "
        "Muris [26], PRISM [27], ImageNet [28], DINOv2 [29], CIFAR-100 [30], "
        "UNI [31], and Prov-GigaPath [32]. An automated DOCX audit now "
        "rejects duplicate or missing reference numbers, unresolved citations, "
        "and drift in these mappings."
    )
    document.core_properties.title = (
        "fastPLS CMPB response - corrected reference numbering"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    revise_response()
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
