#!/usr/bin/env python3

from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Inches


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle23"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle24"
EVIDENCE = (
    ROOT / "benchmark_results" / "simpls_multidataset_ablation_20260725"
)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle23_0.99.6_20260725.docx"
SUPP_SOURCE = (
    SOURCE / "fastPLS_CMPB_supplement_cycle23_0.99.6_20260725.docx"
)
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle23_20260725.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle24_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle24_0.99.6_20260725.docx"
RESPONSE_OUT = (
    OUT / "fastPLS_CMPB_response_to_reviewers_cycle24_20260725.docx"
)
EFFECTS = EVIDENCE / "simpls_multidataset_ablation_effects.csv"
FIGURE = EVIDENCE / "simpls_optimization_ablation.png"

spec = spec_from_file_location(
    "cycle23_helpers",
    ROOT / "tools" / "revise_cmpb_cycle23_cv_comparator.py",
)
c23 = module_from_spec(spec)
spec.loader.exec_module(c23)
c16 = c23.c16


def insert_after(paragraph, text, style=None):
    node = OxmlElement("w:p")
    paragraph._p.addnext(node)
    new = Paragraph(node, paragraph._parent)
    if style is not None:
        new.style = style
    elif paragraph.style is not None:
        new.style = paragraph.style
    new.add_run(text)
    return new


def replace_all(document, old, new):
    for paragraph in document.paragraphs:
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)


def repeat_header(row):
    table_properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    table_properties.append(header)


def effect_rows(data):
    labels = {
        "cached_XtX": "Cached X'X",
        "incremental_coefficients": "Incremental coefficients",
        "cached_deflation_products": "Cached deflation products",
        "compact_prediction": "Compact prediction",
        "matrix_free": "Matrix-free products",
    }
    dataset_labels = {
        "metref": "MetRef",
        "retina": "Retina",
        "prism": "PRISM",
        "cifar100": "CIFAR-100",
    }
    rows = []
    for _, x in data.iterrows():
        applicable = bool(x["optimization_applicable"])
        rows.append(
            (
                dataset_labels[x["dataset"]],
                f'{int(x["n_train"])}/{int(x["p"])}/{int(x["q"])}/'
                f'{int(x["ncomp"])}',
                labels[x["optimization"]],
                "yes" if applicable else "inactive by shape rule",
                f'{x["reference_time_sec"]:.3f}/'
                f'{x["optimized_time_sec"]:.3f}',
                f'{x["speedup"]:.2f}'
                if applicable
                else "not interpreted",
                f'{x["reference_incremental_rss_mb"]:.1f}/'
                f'{x["optimized_incremental_rss_mb"]:.1f}',
                f'{x["rss_reduction_pct"]:.1f}'
                if applicable
                else "not interpreted",
                f'{x["prediction_agreement_min"]:.6f}',
            )
        )
    return rows


def revise_main():
    document = Document(MAIN_SOURCE)
    design = c16.find_paragraph(document, "Estimator preservation and randomized")
    insert_after(
        design,
        "A controlled implementation ablation then isolated cached X'X, "
        "incremental coefficient-path updates, cached deflation products, "
        "compact prediction, and matrix-free cross-covariance products. "
        "Deterministic CPU SIMPLS/IRLBA was used so execution effects were not "
        "confounded with randomized approximation. MetRef, Retina, PRISM, and "
        "CIFAR-100 represented contrasting n, p, and q regimes. Each "
        "reference/optimized pair used the same stored split, full component "
        "path, scaling, seed, and prediction rule in three isolated processes. "
        "RSS was recorded after data loading and garbage collection and sampled "
        "only during fitting and prediction; the reported increment is the "
        "fit-window peak minus that pre-fit baseline.",
        style="Body Text",
    )

    validation = c16.find_paragraph(
        document, "A formal estimator-preservation study"
    )
    insert_after(
        validation,
        "The multi-dataset ablation completed all 120 isolated runs, and every "
        "optimized/reference pair produced identical predictions (maximum "
        "regression difference 4.9 x 10^-15). Effects depended on matrix shape. "
        "For PRISM (q=4,686), compact prediction reduced incremental RSS by "
        "77.7% and improved runtime 1.24-fold; matrix-free products reduced RSS "
        "by 70.6% and improved runtime 6.24-fold. On CIFAR-100, compact "
        "prediction reduced RSS by 4.2% with unchanged runtime, whereas "
        "matrix-free products reduced RSS by 5.8% but were 15.4-fold slower. "
        "Cached deflation reduced PRISM RSS by 18.0% with a 1.02-fold runtime "
        "effect. Incremental coefficient updates and the shape-gated cached "
        "X'X route were approximately neutral in the evaluated settings. These "
        "results identify compact prediction and matrix-free products as "
        "memory-enabling, shape-dependent mechanisms rather than universal "
        "speedups (Supplementary Table S34 and Figure S23).",
        style="Normal",
    )
    document.core_properties.title = (
        "fastPLS CMPB manuscript - multi-dataset SIMPLS ablation"
    )
    document.save(MAIN_OUT)


def revise_supplement(data):
    document = Document(SUPP_SOURCE)
    old = c16.find_paragraph(document, "Execution ablation.")
    old.text = (
        "Approximate-solver ablation. The earlier targeted rSVD experiment "
        "tested solver workspace reuse on MetRef and a CIFAR-100 subset. It is "
        "retained as an approximate-solver sensitivity analysis and is not "
        "used to attribute deterministic SIMPLS implementation speed. The "
        "controlled deterministic implementation ablation is reported in "
        "Section S30."
    )

    replace_all(document, "Tables S20 and S32", "Tables S20 and S33")
    replace_all(document, "Section S29 and Table S32", "Section S29 and Table S33")
    cv_caption = c16.find_paragraph(
        document, "Table S32. Matched ten-fold SIMPLS cross-validation"
    )
    cv_caption.text = cv_caption.text.replace("Table S32.", "Table S33.")

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(
        "S30. Multi-dataset SIMPLS implementation ablation", level=1
    )
    document.add_paragraph(
        "The ablation used deterministic CPU SIMPLS/IRLBA to hold the "
        "statistical estimator fixed. Five mechanisms were changed one at a "
        "time: (i) cached versus direct X'X score normalization, (ii) "
        "incremental versus prefix-recomputed coefficient paths, (iii) cached "
        "versus inline deflation products, (iv) compact latent prediction "
        "versus storage of the complete p x q x k coefficient cube, and (v) "
        "implicit matrix products versus explicit X'Y. MetRef, Retina, PRISM, "
        "and CIFAR-100 were evaluated at 22, 20, 5, and 50 components, "
        "respectively. Every call requested the complete path from component 1 "
        "to k. Three isolated runs were made per configuration. Runtime is "
        "fitting plus prediction. Host RSS was sampled after a synchronized "
        "pre-fit baseline; incremental RSS is the sampled fit-window peak minus "
        "that baseline. Dataset loading is therefore excluded from the memory "
        "increment."
    )
    document.add_paragraph(
        "The cached X'X production condition requires k >= 20, p <= n, "
        "n >= 8p, and p <= 512; only Retina met that condition. Its other rows "
        "are retained as inactive controls and are not interpreted as effects. "
        "The NMR full coefficient-cube reference was not launched because "
        "13,000 x 28,355 x 50 double-precision coefficients would require "
        "137.3 GiB before other workspaces. This analytically demonstrates why "
        "compact prediction is required there but does not convert an "
        "unexecuted reference into a measured speedup."
    )
    table_caption = document.add_paragraph(
        "Table S34. Controlled deterministic SIMPLS implementation ablation. "
        "Dimensions are n/p/q/k. Times and incremental RSS are reference / "
        "optimized medians over three isolated runs. Speedup is reference time "
        "divided by optimized time; values above one favor the optimization. "
        "RSS reduction is relative to the reference. The machine-readable "
        "effects table also reports paired runtime and RSS IQRs.",
        style="Caption",
    )
    table_caption.paragraph_format.page_break_before = True
    table = c16.add_table(
        document,
        [
            "Dataset",
            "n/p/q/k",
            "Mechanism",
            "Active",
            "Time ref/opt s",
            "Speedup",
            "RSS ref/opt MB",
            "RSS reduction %",
            "Prediction agreement",
        ],
        effect_rows(data),
        font_size=4.7,
    )
    repeat_header(table.rows[0])
    c16.prevent_row_splitting(table)
    document.add_paragraph(
        "All 120 runs completed. Classification label agreement was 1.000000 "
        "for every pair; the largest absolute regression prediction difference "
        "was 4.9 x 10^-15. Compact prediction had its clearest effect for "
        "high-q PRISM, reducing incremental RSS by 77.7% and runtime by 19.3%. "
        "Matrix-free products were also favorable for PRISM but slower for "
        "MetRef, Retina, and especially CIFAR-100. Cached deflation reduced "
        "PRISM RSS by 18.0% but had little runtime effect. The remaining "
        "mechanisms were approximately neutral at the evaluated shapes. Thus "
        "the ablation supports conditional memory and runtime benefits, not a "
        "claim that every optimization accelerates every dataset."
    )
    document.add_picture(str(FIGURE), width=Inches(5.75))
    figure_caption = document.add_paragraph(
        "Figure S23. Median runtime and incremental host-memory effects of the "
        "SIMPLS ablation. Runtime is log2(reference / optimized), so positive "
        "values favor optimization; memory is the reduction in "
        "baseline-corrected peak RSS. Cached X'X appears only where its "
        "production shape rule was active.",
        style="Caption",
    )
    figure_caption.paragraph_format.keep_together = True
    document.core_properties.title = (
        "fastPLS CMPB supplement - multi-dataset SIMPLS ablation"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    document.add_heading(
        "25. The implementation ablation was insufficient", level=1
    )
    document.add_paragraph(
        "Reviewer comment: The manuscript attributed speed to cached X'X, "
        "incremental coefficients, cached deflation products, compact "
        "prediction, and matrix-free operations without quantifying each "
        "mechanism across datasets."
    )
    document.add_paragraph(
        "Response: Agreed. We added a controlled ablation on four real datasets "
        "with contrasting matrix shapes. Each mechanism was switched off one "
        "at a time under deterministic SIMPLS/IRLBA, with identical stored "
        "splits, complete component paths, scaling, seeds, and prediction "
        "rules. Three isolated processes were used per configuration, and "
        "baseline-corrected fit-window RSS was sampled externally. All 120 "
        "runs completed and every paired prediction agreed exactly (maximum "
        "regression difference 4.9 x 10^-15). Compact prediction reduced PRISM "
        "RSS by 77.7% and improved runtime 1.24-fold. Matrix-free products "
        "reduced PRISM RSS by 70.6% and improved runtime 6.24-fold, but were "
        "15.4-fold slower on CIFAR-100 for only a 5.8% RSS reduction. Cached "
        "deflation reduced PRISM RSS by 18.0% with little runtime change; "
        "incremental coefficient updates and cached X'X were approximately "
        "neutral in the evaluated settings. We therefore revised the text to "
        "state that these mechanisms are shape-dependent and memory-enabling, "
        "not universal speedups. Methods, Results, Supplementary Section S30, "
        "Table S34, Figure S23, unit tests, scripts, and machine-readable tables "
        "were added."
    )
    document.core_properties.title = (
        "fastPLS CMPB response - multi-dataset SIMPLS ablation"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    data = pd.read_csv(EFFECTS)
    revise_main()
    revise_supplement(data)
    revise_response()
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
