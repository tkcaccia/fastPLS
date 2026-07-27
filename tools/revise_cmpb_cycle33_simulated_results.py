#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle32"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle33"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle32_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle32_0.99.6_20260725.docx"
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle32_20260725.docx"
)

MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle33_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle33_0.99.6_20260726.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle33_20260726.docx"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def revise_main():
    document = Document(MAIN_SOURCE)
    methods = find_paragraph(
        document,
        "Estimator preservation and randomized approximation were evaluated",
    )
    methods.text += (
        " The exact synthetic generator and its numerical results are reported "
        "together in Supplementary Section S6; task dimensions and detailed "
        "agreement endpoints are in Tables S8-S11."
    )

    results = find_paragraph(
        document,
        "A formal estimator-preservation study compared accelerated",
    )
    results.text += (
        " Synthetic generation settings and the corresponding synthetic-only "
        "discrepancy ranges are summarized together in Supplementary Section "
        "S6."
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - simulated-data methods and results"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)
    heading = find_paragraph(document, "S6. Simulated data")
    heading.text = "S6. Synthetic datasets and corresponding numerical results"

    regression = find_paragraph(
        document,
        "Regression simulations use a shared low-rank latent model.",
    )
    regression.text = (
        "The simulated datasets were used only for the formal SIMPLS "
        "estimator-preservation and rSVD reliability analyses; no separate "
        "n/p/q performance-scaling sweep is claimed. Five multivariate "
        "regression regimes and three dummy-response classification regimes "
        "were generated with seeds 101, 202, and 303. Regression data used "
        "independent standard-normal latent scores and Gaussian predictor and "
        "response loading matrices. Predictor and response matrices were "
        "formed by multiplying the latent-score matrix by the transposed "
        "loading matrices and adding independent Gaussian noise with standard "
        "deviation 0.05. The regimes covered p < n, p > n, low- and high-rank Y, near-collinear "
        "predictors, and exact rank deficiency. Near-collinearity was created "
        "by making two loading columns linear combinations of the first "
        "loading column plus noise with standard deviation 1 × 10⁻⁷; exact rank "
        "deficiency was created by duplicating ten predictor columns. Training "
        "and held-out rows were generated in one draw and separated before "
        "model fitting. Exact n, p, q, latent rank, and component grids are "
        "listed in Table S8."
    )

    equation = Paragraph(regression._p.getnext(), regression._parent)
    equation.text = (
        "Classification regimes used balanced shuffled class labels and "
        "Gaussian class-specific latent prototypes. Each latent observation "
        "was its class prototype plus Gaussian noise with standard deviation "
        "0.45; predictors were obtained through a Gaussian loading matrix with "
        "additional noise of standard deviation 0.08, and responses were "
        "one-hot class indicators. The three regimes covered p < n with a "
        "low-rank response, p > n with a higher-rank response, and "
        "near-collinear predictors. Every compared implementation received "
        "identical matrices, splits, component grids, folds, and seeds."
    )

    obsolete_regression = find_paragraph(
        document,
        "where  contains decreasing signal strengths",
    )
    obsolete_regression.text = (
        "The corresponding deterministic results comprised 117 "
        "component-level comparisons. All passed the prespecified tolerances. "
        "Within the synthetic classification regimes, the worst relative "
        "prediction and coefficient errors were 2.55 × 10⁻⁸ and 4.91 × "
        "10⁻⁸; within synthetic regression they were 1.09 × 10⁻⁵ and 1.13 "
        "× 10⁻⁵. Maximum score-subspace angles were 2.09 × 10⁻⁶ degrees for "
        "classification and 0.00143 degrees for regression. The component "
        "selected by fixed five-fold validation agreed with pls::simpls.fit "
        "for all eight synthetic tasks. Detailed regime definitions, "
        "deterministic discrepancies, and component-selection curves are "
        "reported in Tables S8, S10, and S11 and Figures S16-S17."
    )

    obsolete_classification = find_paragraph(
        document,
        "Classification simulations generate normalized class-specific",
    )
    obsolete_classification.text = (
        "For approximate rSVD, oversampling 10 with one power iteration passed "
        "101 of 117 component-level checks, whereas two power iterations "
        "passed 117 of 117. The focused CUDA audit passed all evaluated points "
        "with either four power iterations at oversampling 10 or oversampling "
        "20. These results are reported in Tables S27-S28 and are treated as "
        "solver-reliability evidence rather than estimator-equivalence "
        "evidence."
    )

    document.core_properties.title = (
        "fastPLS CMPB supplement - simulated-data methods and results"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    heading = document.add_heading(
        "33. Simulated-data procedures lacked corresponding results",
        level=1,
    )
    heading.paragraph_format.keep_with_next = True
    comment = document.add_paragraph(
        "Reviewer comment: The simulated-data section contains generation "
        "procedures but no corresponding results. Either include the results "
        "or remove this section."
    )
    comment.paragraph_format.keep_with_next = True
    response = document.add_paragraph(
        "Response: Corrected. The previous Supplementary Section S6 described "
        "an obsolete generic n/p/q sweep that was not part of the current "
        "reported analysis. We removed that unsupported description and "
        "replaced it with the exact five regression and three classification "
        "generators used in the formal estimator-preservation and rSVD "
        "reliability studies. The section now reports the corresponding "
        "synthetic-only prediction, coefficient, subspace, component-selection, "
        "and rSVD pass/fail results and points directly to Tables S8, S10-S11, "
        "and S27-S28 and Figures S16-S17. The main Methods and Results now "
        "cross-reference this combined methods-and-results section. No older "
        "synthetic scaling result is presented as evidence for the current "
        "implementation."
    )
    response.paragraph_format.keep_together = True
    document.core_properties.title = (
        "fastPLS CMPB response to reviewers - simulated-data results"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    revise_response()


if __name__ == "__main__":
    main()
