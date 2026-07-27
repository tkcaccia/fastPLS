#!/usr/bin/env python3

from pathlib import Path
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle56"
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle57"
MAIN_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_main_cycle56_0.99.6_20260726.docx"
)
SUPP_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_supplement_cycle56_0.99.6_20260726.docx"
)
MAIN_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_main_cycle57_0.99.6_20260726.docx"
)
SUPP_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_supplement_cycle57_0.99.6_20260726.docx"
)

FIGURE_DIR = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle57_20260726"
)
EXTERNAL_FIGURE = FIGURE_DIR / "external_single_cpu_accuracy_time_memory.png"
INTERNAL_FIGURE = FIGURE_DIR / "internal_backend_solver_speedups.png"
TECHNICAL_FIGURE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle52_20260726"
    / "main_precision_solver_classifier.png"
)
NMR_FIGURES = [
    (
        "A",
        ROOT
        / "benchmark_results"
        / "review_nmr_20260724"
        / "plots"
        / "nmr_spectrum_full.png",
    ),
    (
        "B",
        ROOT
        / "benchmark_results"
        / "review_nmr_20260724"
        / "plots"
        / "nmr_spectrum_zoom.png",
    ),
    (
        "C",
        ROOT
        / "benchmark_results"
        / "nmr_reference_metrics_20260726"
        / "nmr_reference_predictive.png",
    ),
    (
        "D",
        ROOT
        / "benchmark_results"
        / "nmr_reference_metrics_20260726"
        / "nmr_reference_resources.png",
    ),
]
IMAGENET_FIGURE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle54_20260726"
    / "imagenet_lda_extended_main.png"
)


def require_files(paths):
    missing = [str(path) for path in paths if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing required files:\n" + "\n".join(missing))


def remove_between(start_paragraph, end_paragraph):
    current = start_paragraph._p.getnext()
    while current is not None and current is not end_paragraph._p:
        following = current.getnext()
        current.getparent().remove(current)
        current = following


def move_before(element, anchor_paragraph):
    anchor_paragraph._p.addprevious(element)


def add_paragraph_before(
    document,
    anchor,
    text="",
    style=None,
    alignment=None,
):
    paragraph = document.add_paragraph(text, style=style)
    if alignment is not None:
        paragraph.alignment = alignment
    if style == "Caption":
        paragraph.paragraph_format.keep_together = True
    move_before(paragraph._p, anchor)
    return paragraph


def add_picture_before(document, anchor, path, width):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=width)
    move_before(paragraph._p, anchor)
    return paragraph


def add_page_break_before(document, anchor):
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    move_before(paragraph._p, anchor)
    return paragraph


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def add_nmr_panel_before(document, anchor):
    table = document.add_table(rows=2, cols=2)
    table.autofit = False
    remove_table_borders(table)
    for index, (label, image_path) in enumerate(NMR_FIGURES):
        cell = table.cell(index // 2, index % 2)
        cell.width = Inches(3.05)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        label_paragraph = cell.paragraphs[0]
        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        label_run = label_paragraph.add_run(label)
        label_run.bold = True
        label_run.font.size = Pt(9)
        image_paragraph = cell.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.add_run().add_picture(
            str(image_path),
            width=Inches(2.80),
        )
    move_before(table._tbl, anchor)
    return table


def rewrite_main():
    document = Document(MAIN_SOURCE)
    results = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("3. Results")
    )
    discussion = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("4. Discussion")
    )
    remove_between(results, discussion)

    add_paragraph_before(
        document,
        discussion,
        (
            "We first evaluated whether the computational implementation of "
            "fastPLS improved the practical use of SIMPLS relative to "
            "independent R software under a controlled single-CPU setting. "
            "We then quantified the additional acceleration obtained from "
            "four-thread CPU execution and the CUDA and Metal backends, and "
            "examined when randomized SVD (rSVD) was preferable to the "
            "deterministic IRLBA route. The final analyses focus on two "
            "settings that extend the scale of PLS modelling: multivariate "
            "NMR prediction and million-sample ImageNet embeddings. Formal "
            "estimator-preservation tests, numerical audits, implementation "
            "ablations, component paths, and complete backend tables are "
            "reported in the Supplementary Material."
        ),
        style="First Paragraph",
    )

    add_paragraph_before(
        document,
        discussion,
        "3.1 Single-CPU comparison with independent R implementations",
        style="Heading 2",
    )
    add_paragraph_before(
        document,
        discussion,
        (
            "The primary software comparison used double-precision inputs, "
            "fixed outer splits, and one effective BLAS thread for fastPLS "
            "and all external packages (Figure 2). Deterministic fastPLS "
            "SIMPLS with argmax decoding produced the same test accuracy as "
            "pls::simpls.fit on all nine completed classification datasets. "
            "It was faster on seven datasets, including 4.23-fold on "
            "CIFAR-100, 8.65-fold on Retina, and 8.90-fold on Tabula Muris. "
            "When the same latent-space LDA workflow was compared with "
            "plsgenomics PLS-LDA, accuracy was identical on all eight "
            "completed datasets and fastPLS was faster on six, including "
            "6.44-fold on Retina and 6.77-fold on Tabula Muris. The remaining "
            "independent implementations provide broader workflow context "
            "but are not interpreted as estimator-matched comparisons when "
            "their algorithms or classification rules differ."
        ),
        style="First Paragraph",
    )
    add_paragraph_before(
        document,
        discussion,
        (
            "Peak host memory in Figure 2 is the absolute resident set size "
            "of the complete R process and therefore measures end-to-end "
            "feasibility rather than isolated algorithmic workspace. "
            "fastPLS did not reduce this quantity uniformly on small datasets, "
            "where package and process overhead dominate. Its advantage "
            "became clear on the larger matched SIMPLS tasks: peak RSS was "
            "1.69 GB rather than 13.09 GB on CIFAR-100 and 0.68 GB rather "
            "than 2.20 GB on Tabula Muris. Baseline-corrected memory audits "
            "and method-specific failures are reported in the Supplementary "
            "Material."
        ),
        style="Body Text",
    )
    add_picture_before(
        document,
        discussion,
        EXTERNAL_FIGURE,
        width=Inches(6.15),
    )
    add_paragraph_before(
        document,
        discussion,
        (
            "Figure 2. SIMPLS classification workflows in fastPLS and "
            "independent R packages. All calculations used matched float64 "
            "inputs, fixed training/test splits, and one effective BLAS "
            "thread. Panels show outer-test accuracy (A), total fitting plus "
            "prediction time in seconds (B), and absolute peak process RSS "
            "in MB (C). NE denotes a method that was not evaluated or could "
            "not complete the dataset. Absolute RSS includes the R process, "
            "loaded data, and libraries and should not be interpreted as "
            "isolated workspace allocation."
        ),
        style="Caption",
    )

    add_paragraph_before(
        document,
        discussion,
        "3.2 Internal acceleration and low-rank solver choice",
        style="Heading 2",
    )
    add_paragraph_before(
        document,
        discussion,
        (
            "Backend acceleration depended on the amount and shape of the "
            "linear algebra rather than sample size alone (Figure 3A-C). "
            "Across 44 matched non-NMR CPU-CUDA comparisons, CUDA was faster "
            "in 13. The largest gains occurred on CIFAR-100: 8.90-fold for "
            "PLS-SVD, 14.50-fold for SIMPLS, 4.39-fold for OPLS, and "
            "14.55-fold for kernel PLS, with accuracy differences no larger "
            "than 0.11 percentage points. CUDA also accelerated several "
            "CBMC CITE-seq and GTEx v8 configurations, whereas CPU execution "
            "remained preferable for small matrices because device setup and "
            "transfer costs were not amortized. Metal was faster than the "
            "matched CPU route only for the two CIFAR-100 configurations "
            "tested, and its current metric differences require the numerical "
            "qualification reported in the Supplementary Material. A "
            "four-thread request on MetRef produced only 1.00- to 1.24-fold "
            "speed-ups because the reference BLAS used in that experiment "
            "did not provide a scalable multithreaded matrix-multiplication "
            "path; these results should not be generalized to optimized "
            "multithreaded BLAS installations."
        ),
        style="First Paragraph",
    )
    add_paragraph_before(
        document,
        discussion,
        (
            "The benefit of rSVD over IRLBA was likewise workload-dependent "
            "(Figure 3D). For the nine classification datasets, matched "
            "CPU SIMPLS runs were 1.00- to 1.45-fold faster with rSVD "
            "(median, 1.04-fold). Eight datasets differed by no more than "
            "0.21 percentage points in accuracy, whereas MetRef differed by "
            "4.0 percentage points, illustrating that rSVD is an approximate "
            "solver rather than a deterministic replacement. The advantage "
            "was much larger for the NMR cross-covariance problem: at 100 "
            "components, rSVD reduced CPU SIMPLS time from 436.3 to 19.6 s "
            "(22.3-fold) with unchanged displayed RMSD. We therefore use "
            "IRLBA for deterministic estimator-matched validation and rSVD "
            "for the principal performance benchmark, with fixed "
            "oversampling, power-iteration, and seed settings and with "
            "numerical audit criteria reported in the Supplementary "
            "Material."
        ),
        style="Body Text",
    )
    add_paragraph_before(
        document,
        discussion,
        (
            "Single-precision inputs provide a separate memory-oriented "
            "capability. Matched float32 storage reduced the input-matrix "
            "footprint by approximately one half in the validated routes, "
            "but did not provide a universal runtime or peak-RSS advantage "
            "because conversions, numerical fallbacks, and backend residency "
            "differ by model family. Precision-specific agreement and the "
            "validated CPU, CUDA, and Metal combinations are therefore "
            "reported in the Supplementary Material rather than pooled with "
            "the double-precision software comparison."
        ),
        style="Body Text",
    )
    add_picture_before(
        document,
        discussion,
        INTERNAL_FIGURE,
        width=Inches(6.35),
    )
    add_paragraph_before(
        document,
        discussion,
        (
            "Figure 3. Internal acceleration and solver regimes in fastPLS. "
            "Values above one indicate faster execution by the accelerated "
            "route. (A) CUDA relative to matched CPU rSVD for the four PLS "
            "families; NMR is excluded because it is analysed separately. "
            "(B) Metal relative to matched CPU rSVD; daggers identify "
            "absolute predictive-metric differences greater than 0.005. "
            "(C) four-thread-request relative to one-thread execution on "
            "MetRef. (D) CPU rSVD relative to CPU IRLBA for SIMPLS; the NMR "
            "point uses the fixed 100-component comparison. Complete "
            "accuracy, memory, uncertainty, and numerical-audit results are "
            "provided in the Supplementary Material."
        ),
        style="Caption",
    )

    add_paragraph_before(
        document,
        discussion,
        "3.3 Multivariate NMR prediction",
        style="Heading 2",
    )
    add_paragraph_before(
        document,
        discussion,
        (
            "The NMR task contained 1,200 training spectra, 321 held-out "
            "spectra, 13,000 NOESY predictor bins, and 28,355 diffusion-edited "
            "response intensities. The 4.6-4.8 ppm water interval was set to "
            "zero in both training and test predictors as part of the common "
            "preprocessing protocol. Component selection was performed using "
            "training data only. The one-standard-error rule selected five "
            "components for PLS-SVD and 50 components for SIMPLS; the latter "
            "was the parsimonious choice within a shallow 50-100-component "
            "validation region rather than a uniquely identified optimum."
        ),
        style="First Paragraph",
    )
    add_paragraph_before(
        document,
        discussion,
        (
            "At these family-specific selected settings, CUDA PLS-SVD/rSVD "
            "completed fitting and prediction in 0.898 s with RMSD 0.001043 "
            "and Q2 0.9892, whereas CUDA SIMPLS/rSVD required 1.971 s and "
            "improved RMSD to 0.000759 and Q2 to 0.9943. This comparison "
            "therefore separates the faster low-component PLS-SVD model from "
            "the more accurate higher-component SIMPLS model."
        ),
        style="Body Text",
    )
    fixed_nmr_paragraph = add_paragraph_before(
        document,
        discussion,
        (
            "A second, fixed 100-component benchmark isolated computational "
            "differences at the prediction setting used for the deposited "
            "reference workflow (Figure 4). The deposited fastsimpls "
            "PLS-SVD/IRLBA implementation required 431.2 s and 6.10 GB peak "
            "host RSS, with RMSD 0.000719 and Q2 0.99484. fastPLS "
            "PLS-SVD/rSVD required 16.3 s on CPU and 1.12 s with CUDA, "
            "producing RMSD 0.000729 and 0.000718, respectively. fastPLS "
            "SIMPLS/rSVD required 20.1 s on CPU and 3.06 s with CUDA, with "
            "RMSD 0.000861 and 0.000805. The corresponding predicted spectra "
            "retained correlations of at least 0.9982 with the observations. "
            "Peak host RSS for fastPLS ranged from 2.96 to 3.47 GB; peak "
            "device allocation was 0.66 GB for CUDA PLS-SVD and 3.43 GB for "
            "CUDA SIMPLS. Because this historical contrast changes solver, "
            "implementation, and hardware together, it is interpreted as a "
            "workflow comparison; matched solver and backend contrasts are "
            "reported separately in the Supplementary Material."
        ),
        style="Body Text",
    )
    fixed_nmr_paragraph.paragraph_format.keep_together = True
    add_page_break_before(document, discussion)
    add_nmr_panel_before(document, discussion)
    add_paragraph_before(
        document,
        discussion,
        (
            "Figure 4. NMR prediction and comparison with the deposited "
            "reference workflow. Observed and predicted diffusion-edited "
            "spectra are shown over the full chemical-shift range (A) and "
            "the 1.7-0.5 ppm region (B). Predictive metrics (C) and total "
            "time, peak host RSS, and peak GPU allocation (D) compare the "
            "100-component deposited fastsimpls PLS-SVD/IRLBA workflow with "
            "fastPLS PLS-SVD/rSVD and SIMPLS/rSVD on CPU and CUDA. The "
            "4.6-4.8 ppm water interval was removed consistently during "
            "preprocessing and metric calculation."
        ),
        style="Caption",
    )

    add_paragraph_before(
        document,
        discussion,
        "3.4 ImageNet-scale supervised representation",
        style="Heading 2",
    )
    add_paragraph_before(
        document,
        discussion,
        (
            "ImageNet/DINOv2 was used as a computational stress test and as "
            "an exploratory evaluation of PLS as supervised feature "
            "extraction, not as evidence of biomedical predictive validity. "
            "The fixed benchmark used 1,000,000 training embeddings and "
            "281,167 held-out embeddings with 1,024 features and 1,000 "
            "classes. CUDA SIMPLS/rSVD was evaluated from 100 to 1,000 "
            "components with argmax and latent-space LDA. LDA outperformed "
            "argmax at every component count: top-1 accuracy increased from "
            "0.6270 to 0.7793 at 100 components and from 0.7995 to 0.8093 at "
            "1,000 components. CPU and CUDA predictions agreed to displayed "
            "precision, while CUDA reduced total LDA time from 218.8 to "
            "14.5 s at 100 components and from 2,199.7 to 316.1 s at 1,000 "
            "components (Figure 5)."
        ),
        style="First Paragraph",
    )
    add_paragraph_before(
        document,
        discussion,
        (
            "An independent FAISS benchmark compared nearest-neighbour "
            "retrieval on raw 1,024-dimensional DINOv2 embeddings, "
            "unsupervised PCA scores, and supervised PLS scores using the "
            "same query set. Raw-feature retrieval achieved top-1 and top-5 "
            "accuracy of 0.6556 and 0.9392. With 200 PLS components, these "
            "values were 0.6516 and 0.9397, corresponding to a 5.12-fold "
            "representation reduction, a 0.40-percentage-point top-1 loss, "
            "and approximately fourfold lower projection-plus-query time. "
            "PCA provided the unsupervised dimensionality-reduction control. "
            "These single-run retrieval measurements are exploratory; their "
            "role is to show that supervised PLS scores can retain most "
            "neighbourhood information while reducing representation and "
            "query costs."
        ),
        style="Body Text",
    )
    add_picture_before(
        document,
        discussion,
        IMAGENET_FIGURE,
        width=Inches(6.35),
    )
    add_paragraph_before(
        document,
        discussion,
        (
            "Figure 5. ImageNet-scale SIMPLS classification. CUDA "
            "SIMPLS/rSVD was fitted to 1,000,000 DINOv2 embeddings and "
            "evaluated on 281,167 held-out embeddings over 100-1,000 "
            "components. Panels report top-1 accuracy, total fitting plus "
            "prediction time, peak host RSS, and peak GPU allocation for "
            "argmax and latent-space LDA. Measurements are single exploratory "
            "runs on a fixed, non-standard split; exact data provenance, "
            "top-5 accuracy, retrieval controls, and component-level values "
            "are provided in the Supplementary Material."
        ),
        style="Caption",
    )

    document.save(MAIN_OUTPUT)


def append_supplement():
    document = Document(SUPP_SOURCE)
    add_heading = document.add_paragraph(
        "S40. Consolidated technical validation",
        style="Heading 1",
    )
    add_heading.paragraph_format.page_break_before = True
    document.add_paragraph(
        (
            "The principal manuscript now emphasizes the comparative "
            "software benchmark, backend scaling, and the NMR and ImageNet "
            "applications. The technical evidence supporting those claims "
            "is retained in this Supplement: estimator-preservation and "
            "de Jong mapping (Sections S13-S14), precision validation "
            "(Sections S19 and S32), classifier agreement (Section S20), "
            "rSVD numerical reliability (Section S23), direct PLS-SVD versus "
            "SIMPLS shape experiments (Section S24), cross-validation "
            "comparisons (Section S29), optimization ablations (Section S30), "
            "OPLS and kernel-PLS validation (Section S34), and backend "
            "residency and Metal validation (Sections S36-S38). Figure S40 "
            "provides a compact visual index of the precision, solver, and "
            "classifier evidence; the underlying tables remain authoritative."
        ),
        style="First Paragraph",
    )
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(
        str(TECHNICAL_FIGURE),
        width=Inches(6.35),
    )
    document.add_paragraph(
        (
            "Figure S40. Consolidated numerical and workflow validation. "
            "Panels summarize matched float32-versus-float64 behaviour, "
            "deterministic IRLBA and approximate rSVD evidence, classifier "
            "agreement, and backend-dependent execution. Detailed values, "
            "failure criteria, tolerances, and route-specific limitations are "
            "reported in Sections S13, S19-S20, S23-S24, S29-S34, and S36-S38."
        ),
        style="Caption",
    )
    document.save(SUPP_OUTPUT)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    require_files(
        [
            MAIN_SOURCE,
            SUPP_SOURCE,
            EXTERNAL_FIGURE,
            INTERNAL_FIGURE,
            TECHNICAL_FIGURE,
            IMAGENET_FIGURE,
            *[path for _, path in NMR_FIGURES],
        ]
    )
    shutil.copy2(MAIN_SOURCE, OUTPUT_DIR / MAIN_SOURCE.name)
    shutil.copy2(SUPP_SOURCE, OUTPUT_DIR / SUPP_SOURCE.name)
    rewrite_main()
    append_supplement()
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)


if __name__ == "__main__":
    main()
