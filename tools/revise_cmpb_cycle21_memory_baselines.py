#!/usr/bin/env python3

from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle20"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle21"
EVIDENCE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle21_20260725"
)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle20_0.99.6_20260725.docx"
SUPP_SOURCE = (
    SOURCE / "fastPLS_CMPB_supplement_cycle20_0.99.6_20260725.docx"
)
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle20_20260725.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle21_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle21_0.99.6_20260725.docx"
RESPONSE_OUT = (
    OUT / "fastPLS_CMPB_response_to_reviewers_cycle21_20260725.docx"
)

PAIRED_CSV = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle20_20260725"
    / "paired_backend_selected_summary.csv"
)
MEMORY_CSV = EVIDENCE / "selected_memory_baseline_summary.csv"

spec = spec_from_file_location(
    "cycle20_helpers",
    ROOT / "tools" / "revise_cmpb_cycle20_paired_backend.py",
)
c20 = module_from_spec(spec)
spec.loader.exec_module(c20)
c16 = c20.c16


def memory_value(value):
    return "—" if pd.isna(value) else f"{float(value):.0f}"


def memory_triplet(row, prefix):
    if prefix == "gpu" and row["engine"] == "CPU":
        return "—/—/—"
    if prefix == "host":
        values = (
            row["baseline_host_rss_mb_median"],
            row["process_peak_host_rss_mb_median"],
            row["incremental_host_rss_mb_median"],
        )
    else:
        values = (
            row["baseline_gpu_mem_mb_median"],
            row["peak_gpu_mem_mb_median"],
            row["incremental_gpu_mem_mb_median"],
        )
    return "/".join(memory_value(value) for value in values)


def format_backend_row(row):
    if row["status"] != "ok":
        return f'{row["engine"]}: {row["status"]}'
    iqr = float(row["total_time_sec_q75"]) - float(
        row["total_time_sec_q25"]
    )
    return (
        f'{row["engine"]} {c20.format_metric(row)}; '
        f't={float(row["total_time_sec_median"]):.3f}s '
        f'(IQR {iqr:.3f}); n={int(row["n_runs"])}; OK\n'
        f'H B/P/Δ={memory_triplet(row, "host")}; '
        f'G B/P/Δ={memory_triplet(row, "gpu")}'
    )


def paired_cell(rows):
    ok = rows[rows["status"] == "ok"]
    if ok.empty:
        return "CPU/CUDA: not evaluated in NMR protocol"
    k = int(ok["effective_ncomp"].dropna().iloc[0])
    selection = str(ok["selection_status"].iloc[0])
    marker = "" if selection == "interior tested value" else "†"
    lines = [f"k={k}{marker}"]
    for engine in ("CPU", "CUDA"):
        hit = rows[rows["engine"] == engine]
        lines.append(
            f"{engine}: missing"
            if hit.empty
            else format_backend_row(hit.iloc[0])
        )
    return "\n".join(lines)


def update_main_table(table, paired):
    for row in table.rows[1:]:
        dataset = c20.DATASET_LABELS.get(row.cells[0].text.strip())
        if dataset is None:
            continue
        for column, method in c20.METHOD_COLUMNS.items():
            rows = paired[
                paired["dataset"].eq(dataset)
                & paired["method_panel"].eq(method)
            ]
            row.cells[column].text = paired_cell(rows)
    c16.style_table(table, font_size=3.85)
    c16.prevent_row_splitting(table)
    c20.repeat_header(table.rows[0])


def memory_summary(memory):
    host_fraction = (
        memory["incremental_host_rss_mb_median"]
        / memory["process_peak_host_rss_mb_median"]
    )
    cuda = memory[memory["engine"] == "CUDA"]
    return {
        "host_fraction_median": 100 * host_fraction.median(),
        "host_increment_min": memory[
            "incremental_host_rss_mb_median"
        ].min(),
        "host_increment_max": memory[
            "incremental_host_rss_mb_median"
        ].max(),
        "gpu_baseline_max": cuda["baseline_gpu_mem_mb_median"].max(),
        "gpu_increment_min": cuda["incremental_gpu_mem_mb_median"].min(),
        "gpu_increment_max": cuda["incremental_gpu_mem_mb_median"].max(),
    }


def revise_main(paired, memory):
    document = Document(MAIN_SOURCE)
    values = memory_summary(memory)

    design = c16.find_paragraph(document, "Within each dataset")
    design.text = design.text.replace(
        "Both backends were retained in the primary summary, including their "
        "execution status and separate host/device memory measurements; no "
        "fastest-row filter was applied.",
        "Both backends were retained in the primary summary, including their "
        "execution status and baseline-corrected host/device memory "
        "measurements; no fastest-row filter was applied. Immediately after "
        "data and libraries were loaded and garbage collection completed, the "
        "benchmark recorded process RSS and process-specific GPU memory. A "
        "synchronized sampler then covered fitting and prediction. We report "
        "the pre-fit baseline, the absolute isolated-process peak, and the "
        "incremental fit-window peak above baseline. Absolute RSS remains a "
        "feasibility measure; the increment better isolates workflow memory."
    )

    results = c16.find_paragraph(document, "Table 1 and Figure 2")
    results.text = results.text.replace(
        "CUDA peak host RSS ranged from 0.98 to 1.82 times the CPU value and "
        "required 138-3,414 MB of sampled GPU memory. Thus the paired display "
        "exposes when acceleration is offset by launch, transfer, or memory "
        "overhead.",
        (
            "The memory audit separates data and runtime overhead from fitting "
            "allocations. Incremental host RSS ranged from "
            f'{values["host_increment_min"]:.0f} to '
            f'{values["host_increment_max"]:.0f} MB and represented a median '
            f'{values["host_fraction_median"]:.1f}% of the absolute process '
            "peak. All measured CUDA pre-fit process baselines were "
            f'{values["gpu_baseline_max"]:.0f} MB because CUDA contexts were '
            "initialized during fitting; incremental device memory "
            f'ranged from {values["gpu_increment_min"]:.0f} to '
            f'{values["gpu_increment_max"]:.0f} MB. Thus absolute host RSS '
            "documents whether a workflow fits on the machine, whereas the "
            "baseline-corrected values expose algorithmic and prediction "
            "workspace costs."
        ),
    )

    caption = c16.find_paragraph(document, "Table 1.")
    caption.text = (
        "Table 1. Paired CPU/CUDA biomedical workflow benchmark at the best "
        "training-validation value within each evaluated component grid. Each "
        "cell retains both matched backends at the same family-specific "
        "component count and reports the outer-test metric with 95% interval, "
        "median total fitting-plus-prediction time with run IQR, completed "
        "runs, and memory in MB. H and G denote host and process-specific GPU "
        "memory; B/P/Δ denotes the immediately pre-fit baseline, absolute "
        "isolated-process peak, and incremental fit-window peak above baseline, "
        "respectively. †Lower or upper tested-grid boundary, or response-rank "
        "limit for PLS-SVD; these entries are not global optima. Accuracy uses "
        "Wilson intervals and RMSD uses 10,000-resample held-out-sample "
        "bootstrap intervals, conditional on the fixed outer split. NMR OPLS "
        "and kernel PLS are explicitly labelled not evaluated."
    )
    update_main_table(document.tables[0], paired)

    document.core_properties.title = (
        "fastPLS CMPB manuscript - baseline-corrected memory cycle 21"
    )
    document.save(MAIN_OUT)


def compact_memory_rows(memory):
    rows = []
    for _, row in memory.iterrows():
        rows.append(
            (
                c20.DISPLAY_DATASET[str(row["dataset"])],
                c20.METHOD_LABELS[str(row["method_panel"])],
                row["engine"],
                int(row["effective_ncomp"]),
                memory_triplet(row, "host"),
                memory_triplet(row, "gpu"),
                int(row["n_runs"]),
                "OK",
            )
        )
    for method in ("opls", "kernelpls"):
        for engine in ("CPU", "CUDA"):
            rows.append(
                (
                    "NMR",
                    c20.METHOD_LABELS[method],
                    engine,
                    "—",
                    "—",
                    "—",
                    "—",
                    "not evaluated",
                )
            )
    return rows


def revise_supplement(memory):
    document = Document(SUPP_SOURCE)
    methods = c16.find_paragraph(document, "Each fit is run")
    methods.text = (
        "Each fit is run in an isolated R process. Total time is fitting plus "
        "prediction and excludes package installation and dataset acquisition. "
        "After data and libraries are loaded, garbage collection is completed "
        "and host RSS and process-specific GPU memory are recorded immediately "
        "before fitting. The R process then signals a synchronized external "
        "sampler, which polls host RSS and PID-matched nvidia-smi device memory "
        "every 0.02 s throughout fitting and prediction. /usr/bin/time records "
        "the absolute maximum RSS over the complete isolated process. "
        "Incremental host memory is the maximum of fit-window samples and "
        "post-fit/post-prediction snapshots minus the pre-fit baseline; "
        "incremental GPU memory is the PID-specific sampled peak minus its "
        "pre-fit baseline. Negative differences are truncated at zero. Thus "
        "absolute process RSS measures whole-workflow feasibility, whereas the "
        "baseline-corrected increment estimates fitting and prediction "
        "workspace. Device work is synchronized before each timed section ends."
    )

    audit = c16.find_paragraph(
        document, "The primary benchmark no longer filters"
    )
    audit.text = audit.text.replace(
        "CUDA used 138-3,414 MB of sampled device memory and generally "
        "increased peak host RSS on the small and medium tasks. The "
        "machine-readable file paired_backend_selected_summary.csv retains "
        "backend-specific quartiles, confidence intervals, run counts, and "
        "statuses.",
        "Because absolute RSS includes loaded data, libraries, and benchmark "
        "infrastructure, Table S32 now reports baseline, absolute peak, and "
        "incremental peak separately for host and GPU memory. The "
        "machine-readable files selected_memory_baseline_summary.csv and "
        "selected_memory_raw.csv retain all snapshots, sampled peaks, "
        "quartiles, run counts, and statuses.",
    )

    caption = c16.find_paragraph(document, "Table S32.")
    caption.text = (
        "Table S32. Matched CPU/CUDA memory audit at each family-specific "
        "training-selected component count. B/P/Δ denotes pre-fit baseline, "
        "absolute isolated-process peak, and incremental fit-window peak above "
        "baseline, in MB. Host and GPU values are process-specific."
    )
    old_table = document.tables[-1]
    new_table = c16.add_table(
        document,
        [
            "Dataset",
            "Family",
            "Backend",
            "k",
            "Host B/P/Δ\nMB",
            "GPU B/P/Δ\nMB",
            "Runs",
            "Status",
        ],
        compact_memory_rows(memory),
        font_size=5.0,
    )
    c16.prevent_row_splitting(new_table)
    c20.repeat_header(new_table.rows[0])
    old_table._tbl.getparent().replace(old_table._tbl, new_table._tbl)

    document.core_properties.title = (
        "fastPLS CMPB supplement - baseline-corrected memory cycle 21"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    document.add_heading(
        "22. Absolute process RSS did not isolate algorithmic memory",
        level=1,
    )
    document.add_paragraph(
        "Reviewer comment: Reported host RSS included R, loaded data, "
        "libraries, and benchmark infrastructure. Baseline RSS immediately "
        "before fitting and incremental peak RSS should be reported alongside "
        "absolute RSS, with analogous GPU baseline correction."
    )
    document.add_paragraph(
        "Response: Corrected by new measurement rather than post-hoc "
        "interpretation. Every completed selected-setting CPU and CUDA workflow "
        "was rerun three times in an isolated process. After loading data and "
        "libraries and completing garbage collection, the process recorded host "
        "RSS and PID-specific GPU memory and synchronized with an external "
        "0.02-s sampler before fitting began. We now retain (i) the pre-fit "
        "baseline, (ii) absolute whole-process peak RSS, (iii) the sampled "
        "fit/prediction-window peak, and (iv) the incremental peak above "
        "baseline. GPU baseline, peak, and increment use the same PID-specific "
        "definition. Table 1 reports compact H and G B/P/Δ triplets; "
        "Supplementary Section S27 and Table S32 provide all backend rows. The "
        "Methods and Results now explicitly state that absolute RSS measures "
        "feasibility, while the increment is the closer estimate of workflow "
        "workspace. Machine-readable raw and summarized files are supplied."
    )
    document.core_properties.title = (
        "fastPLS CMPB response - baseline-corrected memory cycle 21"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    paired = pd.read_csv(PAIRED_CSV)
    memory = pd.read_csv(MEMORY_CSV)
    memory["engine"] = memory["engine"].replace({"GPU": "CUDA"})
    join = [
        "dataset",
        "method_panel",
        "engine",
        "effective_ncomp",
    ]
    paired = paired.merge(
        memory[
            join
            + [
                "process_peak_host_rss_mb_median",
                "baseline_host_rss_mb_median",
                "fit_window_peak_host_rss_mb_median",
                "incremental_host_rss_mb_median",
                "baseline_gpu_mem_mb_median",
                "peak_gpu_mem_mb_median",
                "incremental_gpu_mem_mb_median",
            ]
        ],
        on=join,
        how="left",
        validate="one_to_one",
    )
    evaluated = paired[paired["status"] == "ok"]
    if evaluated["incremental_host_rss_mb_median"].isna().any():
        raise RuntimeError("Missing host-memory rerun for an evaluated row")
    if evaluated.loc[
        evaluated["engine"] == "CUDA",
        "incremental_gpu_mem_mb_median",
    ].isna().any():
        raise RuntimeError("Missing GPU-memory rerun for a CUDA row")
    revise_main(paired, memory)
    revise_supplement(memory)
    revise_response()
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
