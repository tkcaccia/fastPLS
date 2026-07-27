from pathlib import Path
import csv

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import revise_cmpb_cycle9_simpls_validation as c9


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle9"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle10"
RESULTS = ROOT / "benchmark_results" / "review_nmr_extended_selection_20260725"
PLOTS = RESULTS / "plots"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle9_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle9_0.99.6_20260725.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle10_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle10_0.99.6_20260725.docx"
RESPONSE_OUT = OUT / "response_NMR_component_selection_20260725.docx"


def read_csv(path):
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def add_figure(document, image, width, caption):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image), width=Inches(width))
    cap = document.add_paragraph(caption, style="Caption")
    cap.paragraph_format.keep_together = True
    return cap


def revise_main():
    document = Document(MAIN_SOURCE)

    abstract = c9.find_paragraph(document, "Results:")
    c9.set_paragraph_text(
        abstract,
        "Results: In a prespecified validation spanning eight synthetic regimes "
        "across three seeds and four real datasets, deterministic IRLBA SIMPLS "
        "completed 117 component-level comparisons against de Jong SIMPLS without "
        "failure. All met the stated tolerances; the maximum relative prediction "
        "error was 1.09×10−5, the maximum latent-subspace angle was 0.00145°, and "
        "fixed five-fold component selection agreed in all 12 tested tasks. rSVD "
        "was evaluated separately as an approximate solver. In multivariate NMR, "
        "five repeated training-only splits over an extended 10–300 component grid "
        "selected 100 components by minimum median validation RMSD, with error "
        "increasing thereafter. Across three isolated runs, held-out SIMPLS-rSVD "
        "median total time was 20.14 s (IQR 0.40 s) on CPU and 3.06 s (IQR 0.03 s) "
        "on CUDA, with RMSD of 0.000861 and 0.000805, respectively."
    )

    nmr = c9.find_paragraph(document, "NMR represented the extreme")
    c9.set_paragraph_text(
        nmr,
        "NMR represented the extreme multivariate-response setting (1,200 training "
        "and 321 held-out spectra; p=13,000; q=28,355). The same protocol was used "
        "for every method: the predefined outer split and full multivariate response "
        "target were retained, predictors were centred without variance scaling, "
        "and the routine 4.6–4.8 ppm residual-water interval was set to zero in both "
        "training and held-out predictor matrices before fitting. Component selection "
        "used only the 1,200 training spectra. Five repeated 80/20 inner splits "
        "(seeds 123, 456, 789, 1011, and 2027) evaluated 10, 25, 50, 75, 100, 125, "
        "150, 165, 175, 200, 250, and 300 components. Median validation RMSD reached "
        "an interior minimum at 100 components (0.0008917), then increased to "
        "0.0009392 at 125, 0.0010048 at 165, and 0.0011272 at 300 components. "
        "Split-specific minima ranged from 25 to 100 components, and this variability "
        "is reported rather than suppressed. The held-out comparison therefore used "
        "100 components for all methods. The deposited fastsimpls PLS-SVD/IRLBA "
        "reference required median 431.23 s and 6,101 MB peak host RSS and achieved "
        "RMSD 0.0007194. fastPLS PLS-SVD/rSVD required 16.32 s and 2,964 MB on CPU "
        "and 1.115 s, 3,338 MB host RSS, and 664 MB GPU memory on CUDA, with RMSD "
        "0.0007292 and 0.0007183, respectively. CPU and CUDA SIMPLS-rSVD required "
        "20.14 and 3.06 s, with RMSD 0.000861 and 0.000805. Across the 28,355 "
        "response coordinates, median response-wise RMSD ranged only from "
        "8.18×10−5 to 8.39×10−5 across the reference and fastPLS methods."
    )

    fixed_split = c9.find_paragraph(document, "The matched CPU/CUDA analysis")
    c9.set_paragraph_text(
        fixed_split,
        "The predefined outer split was retained because the scientific objective "
        "was an exact computational and predictive comparison with the deposited "
        "analysis, not a new population-level performance estimate. Uncertainty in "
        "component selection was quantified across five independent training-only "
        "inner splits, whereas all outer-test spectra remained untouched until the "
        "component count was fixed. Held-out uncertainty is shown as the full "
        "per-spectrum and response-wise error distributions rather than inferred "
        "from repeated reuse of the test set."
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - extended NMR component selection"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)

    routine = c9.find_paragraph(
        document, "Routine NMR spectral preprocessing and component selection."
    )
    c9.set_paragraph_text(
        routine,
        "Routine NMR spectral preprocessing and component selection. The task "
        "comprised 1,200 training and 321 held-out spectra, 13,000 predictors, and "
        "28,355 numeric responses. A shared protocol loader verified identical "
        "predictor and response columns, predefined outer-split dimensions, input "
        "file checksum, and matrix signatures for every method. The residual-water "
        "interval from 4.6 to 4.8 ppm was set to zero in Xtrain and Xtest before "
        "inner splitting or fitting; Ytrain and Ytest were not modified. Predictors "
        "were centred without variance scaling. Component selection accessed only "
        "Xtrain and Ytrain and used five repeated 80/20 splits with seeds 123, 456, "
        "789, 1011, and 2027. The prespecified grid extended beyond the previous "
        "boundary to 10, 25, 50, 75, 100, 125, 150, 165, 175, 200, 250, and 300. "
        "Median validation RMSD was minimized at the interior value of 100 "
        "components (0.0008917) and deteriorated thereafter."
    )

    old_caption = c9.find_paragraph(
        document, "Figure S12. Observed and predicted full held-out NMR spectrum."
    )
    c9.set_paragraph_text(
        old_caption,
        "Figure S12. Observed and predicted full held-out NMR spectrum. The spectrum "
        "was selected mechanically by median per-spectrum RMSD; the model used 100 "
        "components selected by repeated training-only validation over an extended "
        "10–300 component grid."
    )

    document.add_page_break()
    document.add_heading(
        "S14. Extended NMR component selection and response-wise validation", level=1
    )
    document.add_paragraph(
        "One maximal SIMPLS-rSVD model was fitted per inner split, and requested "
        "component prefixes were evaluated in response blocks. This preserved the "
        "same global coefficient path while avoiding a validation-prediction cube. "
        "Selection was based on the minimum median RMSD across the five splits. "
        "The split-specific minima were 100, 75, 25, 75, and 50 components, "
        "demonstrating genuine split uncertainty; nevertheless, the aggregate curve "
        "had a clear interior minimum at 100 and deteriorated at every larger grid "
        "region."
    )

    component_rows = []
    for row in read_csv(RESULTS / "nmr_component_selection_summary.csv"):
        component_rows.append(
            (
                row["ncomp"],
                f'{float(row["RMSD_median"]):.7f}',
                f'{float(row["RMSD_q25"]):.7f}',
                f'{float(row["RMSD_q75"]):.7f}',
                f'{float(row["RMSD_min"]):.7f}',
                f'{float(row["RMSD_max"]):.7f}',
                f'{float(row["Q2_median"]):.4f}',
            )
        )
    caption = document.add_paragraph(
        "Table S12. Repeated training-only NMR component selection. RMSD and Q² "
        "summaries are across five 80/20 inner splits; the outer test set was not "
        "accessed.",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True
    c9.add_table(
        document,
        ["Components", "Median RMSD", "Q25", "Q75", "Minimum", "Maximum", "Median Q²"],
        component_rows,
        [0.75, 0.95, 0.8, 0.8, 0.85, 0.85, 0.85],
        font_size=7.4,
    )
    add_figure(
        document,
        PLOTS / "nmr_component_selection_repeated.png",
        6.6,
        "Figure S18. Training-only NMR component selection. Grey lines show the "
        "five inner splits, the blue line and ribbon show the median and "
        "interquartile range, and the dashed line identifies the selected interior "
        "minimum at 100 components.",
    )

    document.add_paragraph(
        "The outer split was intentionally fixed to reproduce the deposited "
        "reference comparison. For all five final methods, the input checksum, "
        "preprocessing rule, centred predictor matrices, full multivariate target, "
        "outer training/test allocation, and 100-component count were identical. "
        "The full distribution of errors was retained at two levels: 321 "
        "per-spectrum RMSD values and 28,355 response-wise RMSD values. This avoids "
        "reducing the scientific validation to one representative spectrum."
    )
    response_rows = []
    for row in read_csv(PLOTS / "nmr_responsewise_error_summary.csv"):
        response_rows.append(
            (
                row["method"],
                f'{float(row["RMSD_mean"]):.3g}',
                f'{float(row["RMSD_median"]):.3g}',
                f'{float(row["RMSD_q25"]):.3g}',
                f'{float(row["RMSD_q75"]):.3g}',
                f'{float(row["RMSD_q95"]):.3g}',
                f'{float(row["RMSD_q99"]):.3g}',
                f'{float(row["RMSD_max"]):.3g}',
            )
        )
    caption = document.add_paragraph(
        "Table S13. Response-wise held-out RMSD across 28,355 NMR response "
        "coordinates at the selected 100-component model.",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True
    c9.add_table(
        document,
        ["Method", "Mean", "Median", "Q25", "Q75", "Q95", "Q99", "Maximum"],
        response_rows,
        [1.45, 0.68, 0.68, 0.62, 0.62, 0.62, 0.62, 0.72],
        font_size=7.2,
    )
    add_figure(
        document,
        PLOTS / "nmr_responsewise_rmsd.png",
        6.6,
        "Figure S19. Distribution of response-wise RMSD over 28,355 held-out NMR "
        "response coordinates. The display is truncated at the pooled 99th "
        "percentile; complete quantiles and maxima are reported in Table S13.",
    )

    document.add_paragraph(
        "Scripts, raw results, and the protocol manifest are provided in benchmark/ "
        "and benchmark_results/review_nmr_extended_selection_20260725/."
    )
    document.core_properties.title = (
        "fastPLS CMPB supplement - extended NMR component selection"
    )
    document.save(SUPP_OUT)


def write_response():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(11)
    document.styles["Normal"].paragraph_format.space_after = Pt(7)
    document.styles["Heading 1"].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    document.add_heading("Response to reviewer: NMR component selection", level=1)
    reviewer = document.add_paragraph()
    reviewer.add_run("Reviewer comment. ").bold = True
    reviewer.add_run(
        "The selected value of 100 components was the upper boundary of the "
        "reported inner-validation grid while RMSD was still decreasing. Extend "
        "the training-only grid until a plateau or deterioration is observed, "
        "confirm identical preprocessing and targets for every method, and report "
        "response-wise errors and uncertainty over repeated splits."
    )

    response = document.add_paragraph()
    response.add_run("Response. ").bold = True
    response.add_run(
        "We agree. We replaced the single 10–100 inner split with five repeated "
        "training-only 80/20 splits and extended the grid to 300 components, "
        "including the 165-component value used in the earlier study. The median "
        "validation RMSD reached an interior minimum at 100 components (0.0008917), "
        "then increased to 0.0009392 at 125, 0.0010048 at 165, 0.0010630 at 200, "
        "and 0.0011272 at 300. Thus, 100 is no longer selected because it is a "
        "boundary. Split-specific minima were 100, 75, 25, 75, and 50 components "
        "and are reported explicitly."
    )
    document.add_paragraph(
        "A shared protocol loader now verifies the input checksum, matrix "
        "dimensions and signatures, predictor/response column identity, predefined "
        "outer split, full multivariate response target, and centring rule. The "
        "routine 4.6–4.8 ppm residual-water interval is set to zero in both Xtrain "
        "and Xtest before any split or fit; Ytrain and Ytest are unchanged. The same "
        "processed matrices and selected 100-component count are passed to the "
        "deposited fastsimpls reference and every fastPLS CPU/CUDA method."
    )
    document.add_paragraph(
        "We retained the distribution of RMSD over all 321 held-out spectra and "
        "added response-wise RMSD and MAE over all 28,355 response coordinates. "
        "Median response-wise RMSD ranged from 8.18×10−5 to 8.39×10−5 across the "
        "deposited reference and fastPLS methods. The predefined outer test split "
        "is retained to preserve exact comparability with the deposited analysis; "
        "selection uncertainty is quantified across repeated training-only splits, "
        "and the held-out distributions are reported without repeatedly tuning on "
        "or reusing the test set."
    )
    document.add_paragraph(
        "Changes appear in the main NMR Results section and Supplementary Section "
        "S14, Tables S12–S13, and Figures S18–S19. Scripts, raw split-level results, "
        "the protocol manifest, and plots are stored under benchmark/ and "
        "benchmark_results/review_nmr_extended_selection_20260725/."
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    write_response()
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
