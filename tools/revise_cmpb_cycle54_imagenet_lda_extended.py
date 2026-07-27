#!/usr/bin/env python3

import math
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle53"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle54"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle53_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle53_0.99.6_20260726.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle54_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle54_0.99.6_20260726.docx"

PIPELINE4 = (
    ROOT
    / "benchmark_results"
    / "manuscript_multidataset_summary_20260725"
    / "source"
    / "imagenet_pipeline4_summary.csv"
)
RETRIEVAL = (
    ROOT
    / "benchmark_results"
    / "imagenet_faiss_matched_1m_20260725"
    / "imagenet_faiss_matched_main_table.csv"
)
MAIN_FIGURE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle54_20260726"
    / "imagenet_lda_extended_main.png"
)
RETRIEVAL_FIGURE = (
    ROOT
    / "benchmark_results"
    / "imagenet_faiss_matched_1m_20260725"
    / "plots"
    / "imagenet_matched_retrieval.png"
)


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def remove_main_imagenet_table(document):
    caption = find_paragraph(
        document,
        "Table 1. Exploratory matched ImageNet/DINOv2 retrieval",
    )
    sibling = caption._p.getnext()
    while sibling is not None and sibling.tag != qn("w:tbl"):
        sibling = sibling.getnext()
    if sibling is None:
        raise RuntimeError("ImageNet table following Table 1 caption was not found")
    sibling.getparent().remove(sibling)
    caption._p.getparent().remove(caption._p)


def replace_figure(document, caption_prefix, image_path, caption_text):
    caption = find_paragraph(document, caption_prefix)
    previous = caption._p.getprevious()
    if previous is not None and previous.xpath(".//w:drawing"):
        previous.getparent().remove(previous)

    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_with_next = True
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(5.80))
    caption._p.addprevious(image_paragraph._p)

    caption.text = caption_text
    caption.style = "Caption"


def set_cell_margins(cell, top=55, start=65, bottom=55, end=65):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_table(table, widths, font_size=6.2):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    width_dxa = int(sum(widths) * 1440)

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)

    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            width = widths[col_index]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            if row_index == 0:
                shade_cell(cell, "D9EAF7")
            elif row_index % 2 == 0:
                shade_cell(cell, "F4F7F9")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
                    run.bold = row_index == 0


def add_table(document, headers, rows, widths, font_size=6.2):
    table = document.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table"
    for col_index, header in enumerate(headers):
        table.cell(0, col_index).text = header
    for row_index, values in enumerate(rows, start=1):
        for col_index, value in enumerate(values):
            table.cell(row_index, col_index).text = str(value)
    format_table(table, widths, font_size=font_size)
    return table


def f4(value):
    return "NA" if not math.isfinite(float(value)) else f"{float(value):.4f}"


def f1(value):
    return "NA" if not math.isfinite(float(value)) else f"{float(value):.1f}"


def integer(value):
    return "NA" if not math.isfinite(float(value)) else f"{float(value):,.0f}"


def prepare_data():
    pipeline4 = pd.read_csv(PIPELINE4)
    classification = pipeline4[
        pipeline4["classifier"].isin(["argmax", "lda"])
        & pipeline4["backend"].isin(["cpu", "cuda"])
        & (pipeline4["status"] == "ok")
        & pipeline4["ncomp"].isin(range(100, 1001, 100))
    ].copy()
    if len(classification) != 40:
        raise RuntimeError(
            f"Expected 40 completed ImageNet rows, found {len(classification)}"
        )
    retrieval = pd.read_csv(RETRIEVAL)
    return classification, retrieval


def revise_main(classification):
    document = Document(MAIN_SOURCE)
    remove_main_imagenet_table(document)

    methods_anchor = find_paragraph(
        document,
        "The main biomedical benchmark included twelve tasks",
    )
    insert_paragraph_after(
        methods_anchor,
        (
            "The ImageNet stress test fitted float64 SIMPLS/rSVD at 100, 200, "
            "... , 1,000 components on the same 1,000,000/281,167 development "
            "split. CPU and CUDA workflows were evaluated with both direct "
            "argmax decoding and pooled-covariance LDA on the SIMPLS scores. "
            "Each point recorded top-1 accuracy, complete fitting-plus-"
            "prediction time, peak process RSS, and sampled process-specific "
            "GPU memory. The independent feature-reduction control compared "
            "raw DINOv2, PCA, and PLS representations using exact CUDA FAISS "
            "cosine kNN; transformation and held-out query time were included. "
            "These ImageNet measurements remain exploratory because the split "
            "is noncanonical and component counts were not selected by nested "
            "validation."
        ),
        style="Body Text",
    )

    result = find_paragraph(
        document,
        "ImageNet/DINOv2 was used as a million-sample",
    )
    result.text = (
        "ImageNet/DINOv2 was used as a million-sample post-extraction stress "
        "test and exploratory supervised-reduction analysis, not as biomedical "
        "or external predictive validation (Figure 3; Supplementary Tables "
        "S47-S49). Across the complete 100-1,000-component SIMPLS/rSVD path, "
        "LDA produced higher top-1 accuracy than argmax at every component "
        "count. At 100 components, accuracy was 0.7793 for LDA and 0.6270 for "
        "argmax; at 1,000 components it was 0.8093 and 0.7995, respectively. "
        "CPU and CUDA predictions agreed to the displayed precision. The "
        "measured CUDA LDA workflow required 14.5 s at 100 components and "
        "316.1 s at 1,000 components, compared with 218.8-2,199.7 s for CPU "
        "LDA. These are complete workflow measurements, not classifier-only "
        "timings, and each component point is a single exploratory run."
    )
    retrieval_result = insert_paragraph_after(
        result,
        (
            "The independent FAISS analysis addressed the separate question of "
            "supervised feature reduction. Raw 1,024-dimensional embeddings "
            "gave top-1/top-5 accuracy 0.6556/0.9392. The 200-dimensional PLS "
            "representation gave 0.6516/0.9397, corresponding to 5.12-fold "
            "compression, a 0.40-percentage-point top-1 loss, and approximately "
            "fourfold lower held-out projection-plus-query time. PCA remained "
            "the unsupervised dimension-matched control. The small top-5 "
            "difference is descriptive and is not interpreted as improvement."
        ),
        style="Body Text",
    )

    replace_figure(
        document,
        "Figure 3. Exploratory matched ImageNet/DINOv2 retrieval",
        MAIN_FIGURE,
        (
            "Figure 3. Exploratory ImageNet/DINOv2 analysis on the fixed "
            "noncanonical 1,000,000/281,167 development split. (A) SIMPLS/rSVD "
            "top-1 accuracy with argmax and LDA over 100-1,000 components; CPU "
            "and CUDA predictions agreed. (B) Complete CPU and CUDA workflow "
            "time. (C) Host RSS and process-specific GPU memory for CUDA runs. "
            "(D) Exact-FAISS top-1/top-5 accuracy for raw DINOv2, PCA, and PLS "
            "representations. Classifier points are single runs; FAISS queries "
            "used three repeats. Values are in Supplementary Tables S47-S49."
        ),
    )

    abstract = find_paragraph(document, "Results:")
    abstract.text += (
        " In the exploratory million-sample ImageNet stress test, SIMPLS-LDA "
        "outperformed argmax across all 100-1,000-component points, reaching "
        "top-1 accuracy 0.8093 at 1,000 components."
    )

    precision_result = find_paragraph(
        document,
        "Float32 provided a clear storage benefit",
    )
    precision_result.text = precision_result.text.replace(
        "Figure 5A; Table 2",
        "Figure 5A; Table 1",
    )
    precision_caption = find_paragraph(
        document,
        "Table 2. Numerical and execution choices in fastPLS",
    )
    precision_caption.text = precision_caption.text.replace(
        "Table 2.",
        "Table 1.",
        1,
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - extended ImageNet LDA component path"
    )
    document.save(MAIN_OUT)


def add_supplement_tables(document, classification, retrieval):
    document.add_page_break()
    document.add_heading(
        "S39. Extended ImageNet classification and retrieval", level=1
    )
    document.add_paragraph(
        "The classification stress test used the same float64 SIMPLS/rSVD "
        "model family, centring, split, and component grid for CPU and CUDA. "
        "Argmax decoded the response scores directly; LDA used pooled "
        "within-class covariance in the SIMPLS score space. CPU and CUDA "
        "accuracies agreed to the displayed precision. Top-5 scores were not "
        "retained by this classifier-path campaign; top-5 accuracy was measured "
        "in the separate FAISS representation experiment."
    )

    retrieval_rows = []
    labels = {
        "raw_dinov2": "Raw DINOv2",
        "pca_scores": "PCA/rSVD",
        "pls_scores": "PLS-SVD/rSVD",
    }
    for _, row in retrieval.iterrows():
        retrieval_rows.append(
            (
                labels[row["feature_space"]],
                integer(row["n_features"]),
                f"{float(row['compression_ratio']):.2f}",
                f4(row["top1_accuracy"]),
                f4(row["top5_accuracy"]),
                f1(row["transformation_time_sec"]),
                f"{f1(row['query_time_median_sec'])} "
                f"({f1(row['query_time_iqr_sec'])})",
                f1(row["inference_time_median_sec"]),
                f"{integer(row['peak_host_rss_mb'])}/"
                f"{integer(row['peak_gpu_mem_mb'])}",
                f4(row["ivf_neighbour_recall_at_10"]),
            )
        )
    caption = document.add_paragraph(
        "Table S47. Matched ImageNet/DINOv2 feature-reduction results. Exact "
        "CUDA cosine kNN used k=10. Transformation includes fitting and train/"
        "holdout projection; query is median (IQR) over three exact FAISS runs. "
        "H/G denotes peak host RSS and sampled process-level GPU memory in MB.",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True
    add_table(
        document,
        [
            "Representation",
            "Dim.",
            "1024/dim.",
            "Top-1",
            "Top-5",
            "Transform (s)",
            "Query (s)",
            "Inference (s)",
            "H/G (MB)",
            "IVF R@10",
        ],
        retrieval_rows,
        [0.92, 0.38, 0.48, 0.45, 0.45, 0.62, 0.68, 0.68, 0.78, 0.50],
        font_size=5.8,
    )

    grouped = {}
    for _, row in classification.iterrows():
        grouped[(int(row["ncomp"]), row["classifier"], row["backend"])] = row
    path_rows = []
    memory_rows = []
    for ncomp in range(100, 1001, 100):
        ca = grouped[(ncomp, "argmax", "cuda")]
        cl = grouped[(ncomp, "lda", "cuda")]
        pa = grouped[(ncomp, "argmax", "cpu")]
        pl = grouped[(ncomp, "lda", "cpu")]
        path_rows.append(
            (
                ncomp,
                f4(ca["accuracy"]),
                f4(cl["accuracy"]),
                f"{100 * (cl['accuracy'] - ca['accuracy']):.2f}",
                f1(pa["total_fit_predict_sec"]),
                f1(ca["total_fit_predict_sec"]),
                f1(pl["total_fit_predict_sec"]),
                f1(cl["total_fit_predict_sec"]),
            )
        )
        memory_rows.append(
            (
                ncomp,
                f1(pa["peak_host_rss_mb"] / 1024),
                f1(ca["peak_host_rss_mb"] / 1024),
                f1(pl["peak_host_rss_mb"] / 1024),
                f1(cl["peak_host_rss_mb"] / 1024),
                f1(ca["peak_gpu_compute_apps_mb"] / 1024),
                f1(cl["peak_gpu_compute_apps_mb"] / 1024),
            )
        )

    caption = document.add_paragraph(
        "Table S48. ImageNet SIMPLS/rSVD classification over the extended "
        "component grid. Accuracy is identical across CPU and CUDA to the "
        "displayed precision. Times are complete fitting plus prediction in "
        "seconds. Each point is one exploratory run.",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True
    add_table(
        document,
        [
            "A",
            "Argmax acc.",
            "LDA acc.",
            "LDA gain (pp)",
            "CPU argmax (s)",
            "CUDA argmax (s)",
            "CPU LDA (s)",
            "CUDA LDA (s)",
        ],
        path_rows,
        [0.35, 0.72, 0.72, 0.72, 0.82, 0.82, 0.82, 0.82],
        font_size=6.1,
    )

    caption = document.add_paragraph(
        "Table S49. ImageNet SIMPLS/rSVD process-level memory over the extended "
        "component grid. Host values are absolute peak RSS; GPU values are "
        "sampled process-specific peaks and include CUDA context and runtime "
        "state. Units are GiB.",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True
    add_table(
        document,
        [
            "A",
            "CPU argmax host",
            "CUDA argmax host",
            "CPU LDA host",
            "CUDA LDA host",
            "CUDA argmax GPU",
            "CUDA LDA GPU",
        ],
        memory_rows,
        [0.35, 0.90, 0.90, 0.86, 0.86, 0.90, 0.90],
        font_size=6.1,
    )

    figure_paragraph = document.add_paragraph()
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.paragraph_format.keep_with_next = True
    figure_paragraph.add_run().add_picture(
        str(RETRIEVAL_FIGURE), width=Inches(6.25)
    )
    document.add_paragraph(
        "Figure S39. Detailed exploratory ImageNet raw/PCA/PLS retrieval "
        "analysis. Exact search used three repeated index/query measurements "
        "from one fitted representation; representation fitting and IVF were "
        "single runs. The 50/100/200 grid was an exploratory compression "
        "control, whereas the main Figure 3 reports the extended "
        "100-1,000-component SIMPLS classifier path.",
        style="Caption",
    )


def revise_supplement(classification, retrieval):
    document = Document(SUPP_SOURCE)
    add_supplement_tables(document, classification, retrieval)
    document.core_properties.title = (
        "fastPLS CMPB supplement - extended ImageNet LDA component path"
    )
    document.save(SUPP_OUT)


def main():
    for path in (
        MAIN_SOURCE,
        SUPP_SOURCE,
        PIPELINE4,
        RETRIEVAL,
        MAIN_FIGURE,
        RETRIEVAL_FIGURE,
    ):
        if not path.exists():
            raise FileNotFoundError(path)

    OUT.mkdir(parents=True, exist_ok=True)
    classification, retrieval = prepare_data()
    revise_main(classification)
    revise_supplement(classification, retrieval)
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
