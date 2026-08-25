from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle96"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle97"
OUTPUT.mkdir(parents=True, exist_ok=True)


def replace_paragraph(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text


def find_paragraph(document, phrase):
    matches = [p for p in document.paragraphs if phrase in p.text]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph containing {phrase!r}; found {len(matches)}"
        )
    return matches[0]


def insert_after(paragraph, text, style="Body Text"):
    new = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addnext(new._p)
    return new


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "A6A6A6")


def format_table(table, widths):
    table.autofit = False
    table.style = "Table"
    table.alignment = 1
    set_table_borders(table)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(int(w.inches * 1440) for w in widths)))
    tbl_w.set(qn("w:type"), "dxa")
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell.width = widths[col_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                set_cell_shading(cell, "1F4E78")
            elif row_index % 2 == 0:
                set_cell_shading(cell, "EAF2F8")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if col_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7.4)
                    if row_index == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_header(table.rows[0])


source = SOURCE / "fastPLS_CMPB_main_cycle96_0.99.25_20260825.docx"
output = OUTPUT / "fastPLS_CMPB_main_cycle97_0.99.25_20260825.docx"
document = Document(source)

p1 = find_paragraph(document, "The principal contribution of fastPLS is a shape-dependent")
p2 = find_paragraph(document, "rSVD, implicit products, float32, CUDA, and Metal")
p3 = find_paragraph(document, "OPLS and kernel PLS demonstrate deterministic reuse")

replace_paragraph(
    p1,
    "The computational results support a shape-dependent choice rather than one universally "
    "preferred PLS family. On one CPU, PLS-SVD was as fast as or faster than SIMPLS across the "
    "five matched synthetic shapes (SIMPLS/PLS-SVD time ratio 1.00-3.84), as expected for a "
    "one-shot decomposition. On the qualified CUDA shapes, SIMPLS approached or marginally "
    "exceeded PLS-SVD speed (ratio 0.92-0.98), while retaining sequential orthogonalization and "
    "support for component counts not restricted by response rank. Family selection should "
    "nevertheless be based on training-only predictive validation, because PLS-SVD and SIMPLS "
    "are different estimators. Compact prediction matters most when the test set, response "
    "dimension, or number of requested prefixes is large: it reduced incremental RSS by up to "
    "77.7% by avoiding dense coefficient and fitted-response paths, but offered little benefit "
    "when those outputs were intrinsically small."
)
p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

replace_paragraph(
    p2,
    "The storage route and hardware should likewise follow matrix shape. Implicit "
    "cross-covariance products are useful when explicitly storing X-transpose-Y or its "
    "deflated intermediates is the memory bottleneck. In the controlled CPU study they reduced "
    "incremental RSS by 29.7-47.5%, but were slower when the explicit cross-covariance occupied "
    "less than 32 MB and approximately time-neutral at 32-64 MB; the large speed gain occurred "
    "only in the high-response PRISM regime. CUDA first surpassed CPU rSVD at n = 5,000 in the "
    "sample-size sweep and p = 2,000 in the predictor-size sweep. These are empirical crossover "
    "points on the tested hardware, not portable thresholds: transfer, context creation, "
    "synchronization, matrix aspect ratio, and available device memory remain decisive. Metal "
    "was disadvantaged by host-assisted stages, conversion and dispatch overhead, and the small "
    "integrated-GPU test system; moreover, all 20 Metal rSVD diagnostic routes were numerically "
    "discordant and were therefore excluded from speed claims rather than interpreted as valid "
    "accelerator comparisons."
)
p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

replace_paragraph(
    p3,
    "Solver choice separates exploratory acceleration from confirmatory analysis. The qualified "
    "rSVD defaults (oversampling 20 and two power iterations) should be checked across several "
    "prespecified seeds when a result is scientifically consequential; users should compare "
    "predictions, endpoint metrics, and, where relevant, latent subspaces, without selecting the "
    "seed that performs best on held-out test labels. Material seed variation, diagnostic "
    "failure, ill-conditioning, very small samples, or a requirement for deterministic "
    "reproducibility are indications to use CPU IRLBA. When a selected component count reaches "
    "the upper grid boundary, the result is the best within that grid rather than an optimum; "
    "the training-only grid should be extended if rank and computation permit, or a one-standard-"
    "error rule and the boundary limitation should be reported. Nonlinear kernel PLS requires an "
    "n-by-n Gram matrix: the matrix alone occupies approximately 0.8, 5.0, and 20.0 GB in float64 "
    "at n = 10,000, 25,000, and 50,000, before copies and workspaces. It should therefore be "
    "restricted to moderate n or replaced by a linear route or an externally approximated "
    "kernel. Unsupported backend/solver combinations stop explicitly; experimental or "
    "unqualified approximate routes warn and record diagnostics, and fastPLS does not silently "
    "replace the requested estimator. Table 1 summarizes these practical decisions."
)
p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

caption = insert_after(
    p3,
    "Table 1. Practical interpretation of the computational evidence and recommended fastPLS route.",
    style="Caption",
)
caption.paragraph_format.keep_with_next = True

rows = [
    ("Question", "Observed trigger or constraint", "Recommended action"),
    (
        "PLS family",
        "CPU matched shapes favored one-shot PLS-SVD; qualified CUDA shapes placed sequential SIMPLS near parity. PLS-SVD is response-rank limited.",
        "Use training-only prediction to choose the estimator. Prefer PLS-SVD for a low-rank, one-shot path; use SIMPLS when sequential components or components beyond response rank are required.",
    ),
    (
        "Compact prediction",
        "Large test n, large q, or many requested prefixes; up to 77.7% lower incremental RSS in the ablation.",
        "Retain latent factors and block predictions. Dense coefficient, score, loading, or fitted paths should be requested only when scientifically needed.",
    ),
    (
        "Implicit products",
        "Memory pressure from X-transpose-Y; slower below 32 MB and near time-neutral at 32-64 MB in the controlled CPU sweep.",
        "Use primarily to reduce memory for large p by q products, especially high q. Keep the explicit route for small cross-covariance matrices unless profiling supports otherwise.",
    ),
    (
        "CUDA",
        "First qualified crossover at n = 5,000 or p = 2,000 in separate one-factor sweeps on the tested NVIDIA system.",
        "Benchmark the intended shape including transfer and synchronization. Interpret speed only when paired predictions meet numerical tolerances.",
    ),
    (
        "Metal",
        "Host-assisted stages and all 20 diagnostic rSVD routes were numerically discordant on the tested Apple system.",
        "Treat Metal rSVD as experimental; use deterministic CPU execution for confirmatory work until the route is qualified.",
    ),
    (
        "IRLBA or rSVD",
        "Need for deterministic reproducibility, diagnostic failure, material seed variation, ill-conditioning, or small problems where approximation saves little time.",
        "Use CPU IRLBA for confirmation. Use qualified rSVD controls for accelerated exploratory or production workflows and retain solver diagnostics.",
    ),
    (
        "rSVD seeds",
        "Scientifically consequential approximate fit or a new matrix regime.",
        "Run several prespecified seeds and compare predictions, endpoint metrics, and subspaces. Do not choose a seed using held-out test performance; switch to IRLBA if conclusions vary.",
    ),
    (
        "Boundary component",
        "The selected component is the largest value evaluated or is constrained by rank.",
        "Report 'best within the evaluated grid'; extend the training-only grid where feasible or use a prespecified parsimonious rule such as one standard error.",
    ),
    (
        "Nonlinear kernel PLS",
        "The float64 Gram matrix alone is 8n-squared bytes: about 0.8 GB at 10,000, 5.0 GB at 25,000, and 20.0 GB at 50,000 samples.",
        "Restrict to moderate n after accounting for copies and workspaces. For larger n, use linear PLS or an externally validated kernel approximation.",
    ),
    (
        "Unsupported route",
        "Unavailable backend/solver/family combination or unqualified approximate controls.",
        "Unavailable combinations produce an error. Experimental or unqualified routes warn and record diagnostics; no silent estimator substitution is permitted.",
    ),
]

table = document.add_table(rows=len(rows), cols=3)
for i, values in enumerate(rows):
    for j, value in enumerate(values):
        table.cell(i, j).text = value
format_table(table, [Inches(1.12), Inches(2.57), Inches(2.81)])
caption._p.addnext(table._tbl)

document.save(output)
print(output)
