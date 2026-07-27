#!/usr/bin/env python3

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle50"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle51"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle50_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle50_0.99.6_20260726.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle51_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle51_0.99.6_20260726.docx"

EXTERNAL_FIGURE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle51_20260726"
    / "external_simpls_argmax_lda_main.png"
)
COMPUTATIONAL_FIGURE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle49_20260726"
    / "main_selected_computational_performance.png"
)
SHAPE_FIGURE = (
    ROOT
    / "benchmark_results"
    / "simpls_vs_plssvd_shapes_20260726"
    / "simpls_vs_plssvd_shapes_runtime_ratio.png"
)


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def remove_element(element):
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def replace_text(paragraph, old, new):
    if old in paragraph.text:
        paragraph.text = paragraph.text.replace(old, new)


def clear_runs(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def replace_picture(paragraph, path, width, parent):
    if not isinstance(paragraph, Paragraph):
        paragraph = Paragraph(paragraph, parent)
    clear_runs(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))


def find_next_table(paragraph):
    node = paragraph._p.getnext()
    while node is not None:
        if node.tag == qn("w:tbl"):
            return node
        node = node.getnext()
    raise RuntimeError(f"No table follows paragraph: {paragraph.text[:80]}")


def append_figure(document, heading, body, image, caption):
    document.add_page_break()
    document.add_paragraph(heading, style="Heading 2")
    document.add_paragraph(body, style="First Paragraph")
    document.add_page_break()
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(image), width=Inches(5.25))
    document.add_paragraph(caption, style="Caption")


def replace_cell_text_preserve(cell, text):
    text_nodes = list(cell._tc.iter(qn("w:t")))
    if not text_nodes:
        cell.paragraphs[0].add_run(text)
        return
    text_nodes[0].text = text
    for text_node in text_nodes[1:]:
        text_node.text = ""


def revise_main():
    if not EXTERNAL_FIGURE.exists():
        raise FileNotFoundError(EXTERNAL_FIGURE)

    document = Document(MAIN_SOURCE)

    # The direct family-speed figure is already reported as Supplementary Figure S24.
    speed_heading = find_paragraph(
        document,
        "3.1.1 Direct PLS-SVD versus SIMPLS speed comparison",
    )
    speed_body = find_paragraph(
        document,
        "The direct matched experiment showed a crossover",
    )
    speed_caption = find_paragraph(
        document,
        "Figure 2. Direct matched total-time comparison",
    )
    speed_picture = speed_caption._p.getprevious()
    for element in (
        speed_heading._p,
        speed_body._p,
        speed_picture,
        speed_caption._p,
    ):
        remove_element(element)

    benchmark_summary = find_paragraph(
        document,
        "Supplementary Tables S14-S17 report matched CPU and CUDA results",
    )
    benchmark_summary.text = benchmark_summary.text.replace(
        "Supplementary Figure S25 visualizes these results and, in a separate "
        "row, matched CPU/Metal validation for the four datasets evaluated on "
        "Apple hardware.",
        "Supplementary Figure S25 visualizes these results and, in a separate "
        "row, matched CPU/Metal validation for the four datasets evaluated on "
        "Apple hardware. Supplementary Figure S38 summarizes selected-point "
        "runtime and process-level memory across the datasets other than NMR "
        "and ImageNet.",
    )

    # Renumber the retained main-text figures after removing the old Figure 2.
    nmr_body = find_paragraph(
        document,
        "A separate analysis fixed the deposited Nature Communications",
    )
    replace_text(nmr_body, "(Figure 3)", "(Figure 2)")
    nmr_caption = find_paragraph(document, "Figure 3. NMR benchmark")
    replace_text(nmr_caption, "Figure 3.", "Figure 2.")

    imagenet_body = find_paragraph(
        document,
        "ImageNet/DINOv2 was used as a million-sample",
    )
    replace_text(imagenet_body, "Figure 4", "Figure 3")
    imagenet_caption = find_paragraph(
        document,
        "Figure 4. Exploratory matched ImageNet/DINOv2 retrieval",
    )
    replace_text(imagenet_caption, "Figure 4.", "Figure 3.")

    # Replace the previous cross-dataset figure with the external SIMPLS comparison.
    external_heading = find_paragraph(
        document,
        "3.4 Cross-dataset computational performance",
    )
    external_heading.text = "3.4 Comparison with independent R implementations"
    external_body = find_paragraph(
        document,
        "At each family-specific component count selected using training data",
    )
    external_body.text = (
        "The primary software comparison used float64 inputs, identical fixed "
        "outer splits, matched component counts, and three isolated fitting-plus-"
        "prediction runs (Figure 4). Deterministic fastPLS SIMPLS/argmax and "
        "pls::simpls.fit produced identical accuracy on all nine datasets; "
        "fastPLS was faster on seven, including 4.23-fold on CIFAR-100, 8.65-fold "
        "on Retina, and 8.90-fold on Tabula Muris. Across the eight datasets for "
        "which both LDA workflows completed, fastPLS SIMPLS/LDA and plsgenomics "
        "PLS-LDA also produced identical accuracy. fastPLS was faster on six, "
        "including 6.44-fold on Retina and 6.77-fold on Tabula Muris, whereas "
        "plsgenomics was faster by 11-13 ms on the two smallest absolute-time "
        "comparisons. Representative workflows from mdatools, plsdepot, pcv, "
        "chemometrics, mixOmics, and spls provide broader context. These latter "
        "rows differ in estimator or decoder and are therefore workflow, not "
        "implementation-only, comparisons."
    )
    old_cross_caption = find_paragraph(
        document,
        "Figure 5. Selected-point computational performance",
    )
    external_picture = old_cross_caption._p.getprevious()
    replace_picture(
        external_picture,
        EXTERNAL_FIGURE,
        width=6.45,
        parent=document,
    )
    old_cross_caption.text = (
        "Figure 4. SIMPLS classification workflows in fastPLS and independent "
        "R packages. (A) Outer-test accuracy. (B) Total fitting plus prediction "
        "time; cell labels report seconds and color uses a logarithmic scale. "
        "fastPLS argmax and LDA use deterministic CPU IRLBA. Each available "
        "workflow used float64 inputs, the same fixed split and component count "
        "within a dataset, and three isolated runs. NE denotes not evaluated. "
        "Only fastPLS argmax versus pls::simpls.fit and fastPLS LDA versus "
        "plsgenomics PLS-LDA are estimator-and-decoder matched; the other rows "
        "are contextual workflow comparisons."
    )

    # The deterministic two-method table is retained in the Supplement as Table S19.
    cross_validation_heading = find_paragraph(
        document,
        "3.6 External software and cross-validation",
    )
    cross_validation_heading.text = "3.6 Cross-validation performance"
    old_external_paragraph = find_paragraph(
        document,
        "The primary float64 estimator-matched SIMPLS comparison",
    )
    remove_element(old_external_paragraph._p)
    old_table_caption = find_paragraph(
        document,
        "Table 2. Estimator-matched float64 SIMPLS comparison",
    )
    remove_element(find_next_table(old_table_caption))
    remove_element(old_table_caption._p)

    discussion_speed = find_paragraph(
        document,
        "The direct family-speed experiment also bounds the SIMPLS claim",
    )
    discussion_speed.text = discussion_speed.text.replace(
        "balanced Linux-CPU case",
        "balanced Intel Core i7-13700 CPU case",
    )
    discussion_speed.text = discussion_speed.text.replace(
        "The direct family-speed experiment",
        "The direct family-speed experiment in Supplementary Figure S24",
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - external SIMPLS comparison"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    if not COMPUTATIONAL_FIGURE.exists():
        raise FileNotFoundError(COMPUTATIONAL_FIGURE)
    if not SHAPE_FIGURE.exists():
        raise FileNotFoundError(SHAPE_FIGURE)

    document = Document(SUPP_SOURCE)

    # Use hardware descriptors rather than local machine nicknames or OS labels.
    speed_caption = find_paragraph(
        document,
        "Figure S24. Matched SIMPLS/PLS-SVD total-time ratio",
    )
    replace_picture(
        speed_caption._p.getprevious(),
        SHAPE_FIGURE,
        width=6.45,
        parent=document,
    )
    speed_caption.text = speed_caption.text.replace(
        "Absolute CPU times should not be compared between Chiamaka Linux and "
        "Apple M3.",
        "Absolute CPU times should not be compared between the Intel Core "
        "i7-13700 workstation and Apple M3.",
    )
    speed_table_caption = find_paragraph(
        document,
        "Table S37. Direct matched PLS-family speed comparison",
    )
    speed_table = find_next_table(speed_table_caption)
    for cell in speed_table.iter(qn("w:tc")):
        for text_node in cell.iter(qn("w:t")):
            if text_node.text == "Linux CPU":
                text_node.text = "Intel i7-13700 CPU"
            elif text_node.text == "Linux CUDA":
                text_node.text = "RTX 5060 Ti CUDA"

    environment_table = document.tables[4]
    replace_cell_text_preserve(
        environment_table.cell(1, 1),
        "Intel Core i7-13700 (16 cores, 24 logical CPUs); 32,526,480 kB RAM; "
        "NVIDIA GeForce RTX 5060 Ti, 16,311 MiB VRAM",
    )
    replace_cell_text_preserve(
        environment_table.cell(8, 1),
        "Reference BLAS 3.10.0; LAPACK 3.10.0",
    )

    external_section = find_paragraph(
        document,
        "S17. Precision-matched external software comparison",
    )
    external_text = find_paragraph(
        document,
        "The primary software comparison uses float64 inputs",
    )
    external_text.text += (
        " Main-text Figure 4 additionally compares deterministic fastPLS "
        "SIMPLS with argmax and LDA against representative independent "
        "workflows from pls, plsgenomics, mdatools, plsdepot, pcv, "
        "chemometrics, mixOmics, and spls. Exact method-level values are "
        "distributed in external_simpls_argmax_lda_main.csv; comparisons that "
        "change estimator or prediction head are labelled as workflow "
        "comparisons rather than implementation-only speed tests."
    )

    append_figure(
        document,
        "S38. Cross-dataset selected-point computational performance",
        (
            "At each family-specific component count selected using training "
            "data, CPU and CUDA rSVD paths were compared using the same split, "
            "model family, component count, prediction rule, and float64 inputs. "
            "NMR and ImageNet are excluded because their dedicated analyses "
            "separate predictive selection, fixed-complexity computation, and "
            "representation retrieval. Across the remaining 44 paired family "
            "comparisons, CPU was faster in 31 and CUDA in 13. CUDA was faster "
            "for every family on CIFAR-100 and for three of four families on "
            "GTEx v8, TCGA Pan-Cancer, and CBMC CITE-seq. Absolute host RSS and "
            "sampled GPU process memory describe end-to-end feasibility rather "
            "than isolated algorithm workspaces."
        ),
        COMPUTATIONAL_FIGURE,
        (
            "Figure S38. Selected-point computational performance across 11 "
            "datasets (NMR and ImageNet excluded). Points and bars are medians "
            "and interquartile ranges from three isolated runs at each family-"
            "specific training-selected component count. Colors denote PLS "
            "families; circles denote CPU and triangles CUDA. (A) Fitting plus "
            "prediction time. (B) Absolute peak process RSS. (C) Sampled process "
            "GPU memory including CUDA context. Axes are logarithmic; paths use "
            "float64 rSVD. Memory is process-level, not workspace-only."
        ),
    )

    document.core_properties.title = (
        "fastPLS CMPB supplementary material - reorganized figures"
    )
    document.save(SUPP_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
