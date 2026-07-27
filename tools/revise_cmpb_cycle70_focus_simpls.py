#!/usr/bin/env python3
"""Refocus the CMPB manuscript on accelerated SIMPLS as the primary contribution."""

from pathlib import Path
import re

from docx import Document

from revise_cmpb_cycle67_consolidate_evidence import (
    find_paragraph,
    normalize_submission_terminology,
    replace_paragraph,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = (
    ROOT
    / "artifacts"
    / "CMPB_rewrite_20260726_cycle69"
    / "fastPLS_CMPB_main_cycle69_0.99.6_20260726.docx"
)
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle70"
OUTPUT = OUTPUT_DIR / "fastPLS_CMPB_main_cycle70_0.99.6_20260726.docx"


def word_count(text):
    return len(re.findall(r"\b[\w'-]+\b", text))


def revise_front_matter(document):
    replace_paragraph(
        document,
        "fastPLS:",
        (
            "fastPLS: accelerated SIMPLS for high-dimensional biomedical data "
            "with compiled CPU and accelerator backends"
        ),
    )
    replace_paragraph(
        document,
        "Background and objective:",
        (
            "Background and objective: SIMPLS is an established partial least "
            "squares (PLS) estimator, but its sequential component extraction can "
            "become costly for long component paths, large multivariate responses, "
            "and repeated validation. We developed fastPLS primarily to accelerate "
            "SIMPLS without changing its statistical definition."
        ),
    )
    replace_paragraph(
        document,
        "Methods:",
        (
            "Methods: The accelerated SIMPLS path reuses deflation products, latent "
            "quantities, coefficients, and predictions across components and "
            "requested model prefixes. Compact prediction and optional implicit "
            "cross-covariance products reduce avoidable storage. Deterministic IRLBA "
            "and approximate randomized SVD (rSVD) are interchangeable direction "
            "solvers; compiled CPU is the reference implementation, with qualified "
            "NVIDIA CUDA and Apple Metal routes. Estimator preservation, runtime, "
            "memory, prediction agreement, and predictive performance were evaluated "
            "against de Jong SIMPLS and independent R software. NMR and ImageNet "
            "embeddings provided extreme-response and million-sample stress tests."
        ),
    )
    replace_paragraph(
        document,
        "Results:",
        (
            "Results: Deterministic fastPLS SIMPLS met all prespecified tolerances "
            "in 117 component-level comparisons with de Jong SIMPLS. In matched "
            "single-CPU comparisons, it was faster than pls::simpls.fit on seven of "
            "nine datasets with identical accuracy. On NMR, selected CUDA SIMPLS "
            "achieved RMSD 0.000759 and reduced matched CPU time 5.94-fold. In the "
            "exploratory ImageNet analysis, CUDA SIMPLS-LDA processed 1,000,000 "
            "training and 281,167 held-out embeddings and reached top-1 accuracy "
            "0.8093 at 1,000 components. Supporting analyses showed that rSVD, "
            "float32, and accelerator benefits were route and matrix-shape specific; "
            "OPLS and kernel PLS extensions passed their separate deterministic "
            "reliability study."
        ),
    )
    replace_paragraph(
        document,
        "Conclusions:",
        (
            "Conclusions: fastPLS makes sequential SIMPLS substantially more "
            "practical for large biomedical workflows while preserving the "
            "deterministic estimator. The additional solvers, model families, and "
            "hardware routes extend this core contribution but require "
            "route-specific numerical qualification."
        ),
    )
    replace_paragraph(
        document,
        "Keywords:",
        (
            "Keywords: partial least squares; accelerated SIMPLS; randomized SVD; "
            "CUDA; Metal; metabolomics"
        ),
    )


def revise_introduction(document):
    replace_paragraph(
        document,
        "PLS-SVD extracts directions",
        (
            "Among PLS formulations, SIMPLS constructs sequential components without "
            "explicitly deflating the predictor matrix [11]. PLS-SVD provides a "
            "one-shot cross-covariance comparator [10], while OPLS and kernel PLS "
            "extend the framework to orthogonal filtering and nonlinear relations "
            "[12-14]. Existing R implementations do not jointly provide an optimized "
            "SIMPLS component path, memory-aware prediction, and compiled CPU and "
            "accelerator execution."
        ),
    )
    replace_paragraph(
        document,
        "We present fastPLS,",
        (
            "We present fastPLS, whose principal methodological contribution is an "
            "accelerated execution of de Jong SIMPLS. Sequential deflation, "
            "orthogonalization, and component definitions are retained, while "
            "previously computed products, coefficients, and predictions are reused "
            "through the component path. Low-rank solvers, implicit products, "
            "float32, cross-validation, LDA, and CPU/CUDA/Metal execution are "
            "supporting options around this core estimator. OPLS and kernel PLS reuse "
            "the same optimized SIMPLS engine. The GPL-3 R package uses reusable "
            "components from the MIT-licensed kodama-cpp project."
        ),
    )


def revise_methods(document):
    replace_paragraph(
        document,
        "2.2 PLS algorithms and accelerated SIMPLS",
        "2.2 Accelerated SIMPLS and related PLS models",
    )
    replace_paragraph(
        document,
        "PLS-SVD operates on the centred cross-covariance",
        (
            "The accelerated method targets SIMPLS; PLS-SVD is retained as a "
            "one-shot comparator. Both operate on the centred cross-covariance"
        ),
    )
    replace_paragraph(
        document,
        "SIMPLS follows de Jong's",
        (
            "SIMPLS follows de Jong's sequential score, loading, orthogonalization, "
            "and rank-one deflation equations [11]. The innovation is computational: "
            "fastPLS retains deflation products, latent quantities, coefficients, "
            "and predictions incrementally, so all requested component counts are "
            "snapshots of one maximal path rather than independent refits. Each "
            "direction is still extracted from the current deflated state, preserving "
            "the deterministic estimator when IRLBA is used."
        ),
    )
    replace_paragraph(
        document,
        "A shape-aware route can cache",
        (
            "A shape-aware route can cache X-transpose-X for score normalization and "
            "loading calculations. Compact latent factors and blocked prediction "
            "avoid retaining dense coefficient or prediction paths. Optional implicit "
            "products and rSVD change how directions are computed or stored, not the "
            "SIMPLS deflation equations."
        ),
    )
    replace_paragraph(
        document,
        "OPLS applies an orthogonal predictor filter",
        (
            "OPLS and kernel PLS are secondary extensions that reuse the accelerated "
            "SIMPLS core. OPLS first applies an orthogonal predictor filter [12]. "
            "Linear kernel PLS dispatches to the linear route; RBF and polynomial "
            "kernels form an n-by-n Gram matrix [13,14], restricting nonlinear models "
            "to moderate sample sizes."
        ),
    )
    replace_paragraph(
        document,
        "float64 CPU fitting supports bundled deterministic IRLBA",
        (
            "Direction extraction is modular rather than the principal estimator "
            "contribution. Float64 CPU fitting supports deterministic IRLBA [15] and "
            "approximate rSVD [16]. rSVD uses a Gaussian range sketch, "
            "orthonormalization, optional power iterations, and a reduced "
            "decomposition. CUDA and Metal support rSVD; approximate results were "
            "audited by prediction, subspace, label, and metric criteria "
            "(Supplementary Table S8)."
        ),
    )
    replace_paragraph(
        document,
        "Estimator preservation, OPLS/kernel reliability",
        (
            "The primary validation tested whether accelerated deterministic SIMPLS "
            "preserved de Jong's estimator and improved runtime. rSVD approximation, "
            "backend portability, float32, OPLS/kernel extensions, cross-validation, "
            "and repeated outer partitions were evaluated separately as supporting "
            "analyses. Detailed designs, thresholds, provenance, and complete results "
            "are in Supplementary Tables S3-S15."
        ),
    )


def revise_results(document):
    replace_paragraph(
        document,
        "We first compared fastPLS SIMPLS",
        (
            "Results are organized around the accelerated SIMPLS contribution: "
            "estimator preservation, comparison with independent implementations, "
            "and performance across matrix shapes. Solver, precision, and hardware "
            "analyses qualify the implementation, while NMR and ImageNet demonstrate "
            "large-response and large-sample use cases."
        ),
    )
    replace_paragraph(
        document,
        "Deterministic SIMPLS met all prespecified",
        (
            "The primary estimator-preservation study found that deterministic "
            "fastPLS SIMPLS met all prespecified tolerances in 117 component-level "
            "comparisons with de Jong SIMPLS. Separately, OPLS and kernel PLS "
            "extensions passed all 66 setting/task comparisons and all 1,540 "
            "fold-component fits in their deterministic float64 CPU reliability "
            "study (Supplementary Table S7)."
        ),
    )
    replace_paragraph(
        document,
        "3.1 Comparison with independent R implementations",
        "3.1 Accelerated SIMPLS compared with independent R implementations",
    )
    replace_paragraph(
        document,
        "Of 44 non-NMR CPU-CUDA pairs",
        (
            "Hardware acceleration was a supporting, route-dependent result rather "
            "than the central claim. Of 44 non-NMR CPU-CUDA pairs, 28 met numerical "
            "criteria and CUDA was faster in seven, with eligible speed-up up to "
            "8.90-fold. Sixteen discordant routes were excluded. Six of 12 CPU-Metal "
            "pairs were concordant and none was faster with Metal (Figure 3; "
            "Supplementary Table S11)."
        ),
    )
    replace_paragraph(
        document,
        "CPU SIMPLS rSVD was only",
        (
            "The choice of direction solver modified the speed-accuracy trade-off of "
            "accelerated SIMPLS. CPU rSVD was 1.00-1.45-fold faster than IRLBA across "
            "nine classification tasks, but MetRef accuracy differed by 4.0 "
            "percentage points. On NMR at 100 components, rSVD reduced CPU time from "
            "436.3 to 19.6 s. IRLBA therefore remains the deterministic reference and "
            "rSVD an audited approximate option (Supplementary Table S8)."
        ),
    )


def revise_discussion(document):
    replace_paragraph(
        document,
        "fastPLS reorganizes established PLS computation",
        (
            "The principal contribution of fastPLS is an accelerated SIMPLS "
            "execution path, not a new PLS estimator. Deterministic validation showed "
            "that reuse of sequential quantities preserved de Jong SIMPLS, while the "
            "external comparison showed lower runtime and, for large tasks, lower "
            "memory without changing matched accuracy."
        ),
    )
    replace_paragraph(
        document,
        "Benefits depend on matrix shape",
        (
            "rSVD, implicit products, float32, CUDA, and Metal should be interpreted "
            "as optional implementation mechanisms around accelerated SIMPLS. Their "
            "benefits depend on matrix shape and numerical agreement: CPU is "
            "preferable for many small tasks, selected concordant CUDA routes benefit "
            "large dense or extreme-response problems, and Metal demonstrated "
            "portability but no speed advantage here."
        ),
    )
    replace_paragraph(
        document,
        "NMR demonstrates practical",
        (
            "OPLS and kernel PLS demonstrate reuse of the accelerated core, whereas "
            "NMR and ImageNet test its feasible scale. ImageNet establishes "
            "million-sample execution and supervised compression, not biomedical "
            "validity. Limitations include finite component grids, conditional "
            "uncertainty, quadratic nonlinear kernels, approximate rSVD, and "
            "route-specific precision or accelerator disagreement."
        ),
    )
    replace_paragraph(
        document,
        "fastPLS makes established PLS estimators",
        (
            "fastPLS makes sequential SIMPLS and its validation feasible across "
            "larger biomedical matrix regimes by reusing component-path computation "
            "and reducing avoidable storage. OPLS, kernel PLS, low-rank solvers, and "
            "accelerator backends extend this core contribution; deterministic "
            "float64 CPU SIMPLS remains the confirmatory reference."
        ),
    )


def audit(document):
    abstract = " ".join(
        find_paragraph(document, prefix).text
        for prefix in (
            "Background and objective:",
            "Methods:",
            "Results:",
            "Conclusions:",
        )
    )
    introduction = next(
        i
        for i, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == "1. Introduction"
    )
    references = next(
        i
        for i, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == "References"
    )
    main_text = " ".join(
        paragraph.text for paragraph in document.paragraphs[introduction:references]
    )
    abstract_words = word_count(abstract)
    main_words = word_count(main_text)
    if abstract_words > 350:
        raise RuntimeError(f"Abstract has {abstract_words} words")
    if main_words > 3500:
        raise RuntimeError(f"Main text has {main_words} words")
    required = (
        "principal methodological contribution",
        "accelerated SIMPLS",
        "supporting options",
    )
    joined = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for phrase in required:
        if phrase not in joined:
            raise RuntimeError(f"Missing focus phrase: {phrase}")
    print(f"Abstract words: {abstract_words}")
    print(f"Main-text words: {main_words}")


def main():
    document = Document(SOURCE)
    revise_front_matter(document)
    revise_introduction(document)
    revise_methods(document)
    revise_results(document)
    revise_discussion(document)
    normalize_submission_terminology(document)
    audit(document)
    document.core_properties.title = (
        "fastPLS CMPB manuscript - accelerated SIMPLS focus"
    )
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
