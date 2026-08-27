from pathlib import Path

from docx import Document


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE_DIR = ROOT / "artifacts/CMPB_rewrite_20260826_cycle111"
OUTDIR = ROOT / "artifacts/CMPB_rewrite_20260826_cycle112"


MAIN_REPLACEMENTS = {
    "Methods: Central evidence used archived fastPLS 0.99.25 (Git commit 7887401b09e2). The implementation combines compiled shape-dependent execution, incremental updates, compact latent prediction, and optional implicit cross-covariance products. We evaluated SIMPLS and PLS-SVD, which derives components from a singular value decomposition of the predictor-response cross-covariance. Classification used maximum-score (argmax) decoding or linear discriminant analysis. Fixed-control implicitly restarted Lanczos bidiagonalization (IRLBA) provided the deterministic numerical reference; approximate randomized SVD (rSVD) and accelerator routes were evaluated separately under predefined criteria.":
    "Methods: Central evidence used archived fastPLS 0.99.25 (Git commit 7887401b09e2). The implementation combines compiled incremental SIMPLS updates, reuse of sequential deflation quantities, compact latent prediction, and optional implicit cross-covariance products. We evaluated SIMPLS and PLS-SVD, which derives components from a singular value decomposition of the predictor-response cross-covariance. Classification used maximum-score (argmax) decoding or linear discriminant analysis. Fixed-control implicitly restarted Lanczos bidiagonalization (IRLBA) provided the deterministic numerical reference; approximate randomized SVD (rSVD) and accelerator routes were evaluated separately under predefined criteria.",

    "We present fastPLS, whose principal methodological contribution is a compiled, shape-dependent execution of de Jong SIMPLS. Sequential deflation, orthogonalization, and component definitions are retained, while avoidable intermediate products and dense outputs are reduced according to matrix shape and requested output. The R interface also provides PLS-SVD, OPLS, kernel PLS, compact latent prediction, optional implicit predictor-response products, compiled single and nested validation, and explicit solver and backend diagnostics. Low-rank solvers, float32 execution, linear discriminant analysis, and CPU, NVIDIA CUDA, and Apple Metal routes support the central SIMPLS implementation. The benchmark therefore separates numerical validation against pls::simpls.fit, comparison with available R PLS workflows, and a distinct cross-language comparison with IKPLS. The GPL-3 R package uses reusable components from the MIT-licensed kodama-cpp project.":
    "We present fastPLS, whose principal methodological contribution is an accelerated compiled execution of de Jong SIMPLS. Sequential deflation, orthogonalization, and component definitions are retained, while rank-one deflation products and previously computed sequential quantities are reused, coefficient and fitted-value updates are accumulated incrementally, and compact latent prediction avoids unnecessary dense coefficient and prediction paths. Optional caching and implicit predictor-response products provide secondary workload-specific memory optimizations. The R interface also provides PLS-SVD, OPLS, kernel PLS, compiled single and nested validation, and explicit solver and backend diagnostics. Low-rank solvers, float32 execution, linear discriminant analysis, and CPU, NVIDIA CUDA, and Apple Metal routes support the central SIMPLS implementation. The benchmark therefore separates numerical validation against pls::simpls.fit, comparison with available R PLS workflows, and a distinct cross-language comparison with IKPLS. The GPL-3 R package uses reusable components from the MIT-licensed kodama-cpp project.",

    "2.2 Shape-aware SIMPLS execution and related PLS models":
    "2.2 Accelerated SIMPLS execution and related PLS models",

    "SIMPLS follows de Jong's sequential score, loading, orthogonalization, and rank-one deflation equations [11]. As in pls::simpls.fit, one fit supplies the standard sequential path through all components up to the requested maximum; fastPLS does not claim this path construction as novel. Its contribution is computational and shape dependent: rank-one deflation products are reused; coefficient and fitted-value quantities are updated incrementally; X-transpose-X is cached only when its setup and storage can be amortized; compact prediction retains latent factors instead of dense coefficient and prediction paths; and implicit products avoid an explicit predictor-response cross-covariance when that matrix is limiting. Every component invokes a fresh rank-one direction calculation on the current deflated state. Candidate-direction blocks, cross-component warm starts, and adaptive refresh were rejected and are not used by the public CPU, CUDA, or Metal paths.":
    "SIMPLS follows de Jong's sequential score, loading, orthogonalization, and rank-one deflation equations [11]. As in pls::simpls.fit, one fit supplies the standard sequential path through all components up to the requested maximum; fastPLS does not claim this path construction as novel. Its main computational contribution is the incremental reuse of work along that path: rank-one deflation products are reused, coefficient and fitted-value quantities are accumulated as each component is added, and compact prediction retains latent factors instead of dense coefficient and prediction paths. Caching X-transpose-X when its setup can be amortized and using implicit products when the predictor-response cross-covariance is limiting are secondary workload-specific optimizations. Every component invokes a fresh rank-one direction calculation on the current deflated state. Candidate-direction blocks, cross-component warm starts, and adaptive refresh were rejected and are not used by the public CPU, CUDA, or Metal paths.",

    "A shape-aware route can cache X-transpose-X for score normalization and loading calculations. Compact latent factors and blocked prediction avoid retaining dense coefficient or prediction paths. Optional implicit products and rSVD change how directions are computed or stored, not the SIMPLS deflation equations.":
    "The accelerated path reuses sequential deflation quantities and updates prediction-related quantities incrementally. Compact latent factors and blocked prediction avoid retaining dense coefficient or prediction paths. X-transpose-X caching, optional implicit products, and rSVD affect the computational route or storage, not the SIMPLS deflation equations.",

    "Algorithm 1. Shape-aware SIMPLS execution. Direction extraction uses deterministic IRLBA or approximate rSVD; score construction, orthogonalization, and deflation follow de Jong [11].":
    "Algorithm 1. Accelerated incremental SIMPLS execution. Direction extraction uses deterministic IRLBA or approximate rSVD; score construction, orthogonalization, and deflation follow de Jong [11].",

    "Results are organized around the shape-dependent SIMPLS execution contribution. We first compare complete single-CPU classification workflows across independent R implementations, then separate numerically concordant CPU, CUDA, and Metal execution. NMR provides the principal biomedical high-response case study, while ImageNet is retained as a qualified foundation-model-scale feasibility analysis. Exact estimator validation, approximate-solver qualification, route diagnostics, and complete result tables remain in the Supplementary Material.":
    "Results are organized around the accelerated incremental SIMPLS execution and its memory-saving prediction path. We first compare complete single-CPU classification workflows across independent R implementations, then separate numerically concordant CPU, CUDA, and Metal execution. NMR provides the principal biomedical high-response case study, while ImageNet is retained as a qualified foundation-model-scale feasibility analysis. Exact estimator validation, approximate-solver qualification, route diagnostics, and complete result tables remain in the Supplementary Material.",

    "The computational results support a shape-dependent choice rather than one universally preferred PLS family or implementation. On one CPU, PLS-SVD was as fast as or faster than SIMPLS across the five matched synthetic shapes (SIMPLS/PLS-SVD time ratio 1.00-3.84), as expected for a one-shot decomposition. On the qualified CUDA shapes, SIMPLS approached or marginally exceeded PLS-SVD speed (ratio 0.92-0.98), while retaining sequential orthogonalization and support for component counts not restricted by response rank. Family selection should nevertheless be based on training-only predictive validation, because PLS-SVD and SIMPLS are different estimators. Compact prediction matters most when the test set, response dimension, or number of requested prefixes is large: it reduced incremental RSS by up to 77.7% by avoiding dense coefficient and fitted-response paths, but offered little benefit when those outputs were intrinsically small. In the broader R-package panel, fastPLS had the lowest observed total time on seven of nine classification datasets, although the matched minimal-output comparison with pls::simpls.fit showed smaller, dataset-dependent differences. IKPLS was faster in the separate single-thread cross-language experiment, emphasizing that fastPLS contributes an R-native de Jong SIMPLS workflow, multivariate-response storage controls, nested validation, multiple PLS families, and route diagnostics rather than universal superiority over every PLS implementation.":
    "The principal gain arises from reusing sequential SIMPLS quantities, updating coefficients and fitted values incrementally, and predicting from compact latent factors rather than from retained dense paths. The benefit increases when the response dimension, test set, or number of requested component prefixes is large: compact prediction reduced incremental RSS by up to 77.7%, but offered little benefit when those outputs were intrinsically small. PLS-family and hardware choices remain workload dependent. On one CPU, PLS-SVD was as fast as or faster than SIMPLS across the five matched synthetic matrices (SIMPLS/PLS-SVD time ratio 1.00-3.84), as expected for a one-shot decomposition. On the qualified CUDA cases, SIMPLS approached or marginally exceeded PLS-SVD speed (ratio 0.92-0.98), while retaining sequential orthogonalization and support for component counts not restricted by response rank. Family selection should nevertheless use training-only predictive validation because PLS-SVD and SIMPLS are different estimators. In the broader R-package panel, fastPLS had the lowest observed total time on seven of nine classification datasets, although the matched minimal-output comparison with pls::simpls.fit showed smaller, dataset-dependent differences. IKPLS was faster in the separate single-thread cross-language experiment, emphasizing that fastPLS contributes an R-native de Jong SIMPLS workflow, multivariate-response storage controls, nested validation, multiple PLS families, and route diagnostics rather than universal superiority over every PLS implementation.",

    "The storage route and hardware should likewise follow matrix shape. Implicit cross-covariance products are useful when explicitly storing the predictor-response cross-product or deflated intermediates is the memory bottleneck. In the controlled CPU study they reduced incremental RSS by 29.7-47.5%, but were slower below 32 MiB and approximately time-neutral at 32-64 MiB. CUDA first surpassed CPU rSVD at n = 5,000 in the sample-size sweep and p = 2,000 in the predictor-size sweep. These are hardware-specific crossovers because transfer, context creation, synchronization, aspect ratio, and device memory remain decisive. Metal was disadvantaged by host-assisted stages and dispatch overhead; its discordant rSVD routes were excluded from valid accelerator comparisons.":
    "Optional storage routes and hardware acceleration provide secondary, workload-specific gains. Implicit cross-covariance products are useful when explicitly storing the predictor-response cross-product or deflated intermediates is the memory bottleneck. In the controlled CPU study they reduced incremental RSS by 29.7-47.5%, but were slower below 32 MiB and approximately time-neutral at 32-64 MiB. CUDA first surpassed CPU rSVD at n = 5,000 in the sample-size sweep and p = 2,000 in the predictor-size sweep. These hardware-specific crossovers depend on transfer, context creation, synchronization, aspect ratio, and device memory. Metal was disadvantaged by host-assisted stages and dispatch overhead; its discordant rSVD routes were excluded from valid accelerator comparisons.",
}


SUPPLEMENT_REPLACEMENTS = {
    "For tall matrices with moderate , a shape-aware path caches  and the initial cross-covariance so that score normalization and loadings avoid repeated observation-space products.":
    "When its setup and storage can be amortized, the implementation caches the predictor cross-product and initial predictor-response cross-covariance so that score normalization and loadings avoid repeated observation-space products.",

    "These development-stage ablations predate the frozen central rerun and are retained only to document how the execution design was developed; they do not support release-level speed claims. Each ablation changed one internal feature relative to a minimally optimized compiled SIMPLS baseline while holding data, split, estimator, deterministic IRLBA solver, component count, seed, and prediction head fixed. Three isolated runs were used per configuration. The measured ranges quantify shape-dependent trade-offs rather than universal speed-ups.":
    "These development-stage ablations predate the frozen central rerun and are retained only to document how the execution design was developed; they do not support release-level speed claims. Each ablation changed one internal feature relative to a minimally optimized compiled SIMPLS baseline while holding data, split, estimator, deterministic IRLBA solver, component count, seed, and prediction head fixed. Three isolated runs were used per configuration. The measured ranges quantify workload-dependent trade-offs rather than universal speed-ups.",
}


TABLE_CELL_REPLACEMENTS = {
    "Yes; shape-dependent and primarily memory-oriented": "Yes; optional and primarily memory-oriented",
    "fastPLS shape-aware SIMPLS": "fastPLS accelerated SIMPLS",
    "Selects cached, compact, or implicit execution by shape; trades setup/operator work for lower repeated work or storage":
    "Reuses sequential quantities and supports compact prediction; optional caching and implicit products trade setup or operator work for lower repeated work or storage",
}


def replace_paragraphs(document: Document, replacements: dict[str, str]) -> None:
    for old, new in replacements.items():
        matches = [paragraph for paragraph in document.paragraphs if paragraph.text == old]
        if len(matches) != 1:
            raise RuntimeError(f"Expected one paragraph match, found {len(matches)} for: {old[:90]}")
        matches[0].text = new


def replace_table_cells(document: Document, replacements: dict[str, str]) -> None:
    counts = {key: 0 for key in replacements}
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text in replacements:
                    cell.text = replacements[cell.text]
                    counts[cell.text] = counts.get(cell.text, 0)
    # Cell replacement uses cell.text assignment, so validate from final text instead.
    final_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
    for old, new in replacements.items():
        if old in final_cells or new not in final_cells:
            raise RuntimeError(f"Table-cell replacement failed: {old}")


def main() -> None:
    OUTDIR.mkdir(parents=True, exist_ok=True)

    main_doc = Document(SOURCE_DIR / "fastPLS_CMPB_main_cycle111_0.99.25_20260826.docx")
    replace_paragraphs(main_doc, MAIN_REPLACEMENTS)
    main_doc.save(OUTDIR / "fastPLS_CMPB_main_cycle112_0.99.25_20260826.docx")

    supplement = Document(SOURCE_DIR / "fastPLS_CMPB_supplement_cycle111_0.99.25_20260826.docx")
    replace_paragraphs(supplement, SUPPLEMENT_REPLACEMENTS)
    replace_table_cells(supplement, TABLE_CELL_REPLACEMENTS)
    supplement.save(OUTDIR / "fastPLS_CMPB_supplement_cycle112_0.99.25_20260826.docx")


if __name__ == "__main__":
    main()
