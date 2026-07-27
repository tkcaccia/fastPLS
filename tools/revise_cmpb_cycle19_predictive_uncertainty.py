#!/usr/bin/env python3

from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle18"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle19"
EVIDENCE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle19_20260725"
)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle18_0.99.6_20260725.docx"
SUPP_SOURCE = (
    SOURCE / "fastPLS_CMPB_supplement_cycle18_0.99.6_20260725.docx"
)
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle18_20260725.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle19_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle19_0.99.6_20260725.docx"
RESPONSE_OUT = (
    OUT / "fastPLS_CMPB_response_to_reviewers_cycle19_20260725.docx"
)

UNCERTAINTY_CSV = EVIDENCE / "selected_predictive_uncertainty.csv"
IMAGENET_CSV = EVIDENCE / "imagenet_predictive_uncertainty.csv"
FIGURE = EVIDENCE / "plots" / "selected_performance_all_datasets_with_ci.png"

spec = spec_from_file_location(
    "cycle18_helpers",
    ROOT / "tools" / "revise_cmpb_cycle18_separate_nmr_questions.py",
)
c18 = module_from_spec(spec)
spec.loader.exec_module(c18)
c16 = c18.c16


DATASET_LABELS = {
    "MetRef": "metref",
    "CCLE": "ccle",
    "TCGA-BRCA": "tcga_brca",
    "TCGA-HNSC methyl.": "tcga_hnsc_methylation",
    "TCGA-HNSC methylation": "tcga_hnsc_methylation",
    "GTEx v8": "gtex_v8",
    "TCGA Pan-Cancer": "tcga_pan_cancer",
    "Retina": "retina",
    "Tabula Muris": "tabula",
    "CIFAR-100": "cifar100",
    "CBMC CITE-seq": "cbmc_citeseq",
    "PRISM": "prism",
    "NMR": "nmr",
}
METHOD_COLUMNS = {
    1: "plssvd",
    2: "simpls",
    3: "opls",
    4: "kernelpls",
}


def format_ci(point, lower, upper):
    if point >= 100:
        return f"[{lower:.0f}, {upper:.0f}]"
    if point >= 0.01:
        return f"[{lower:.3f}, {upper:.3f}]"
    return f"[{lower:.2e}, {upper:.2e}]"


def format_estimate(point):
    if point >= 100:
        return f"{point:.1f}"
    if point >= 0.01:
        return f"{point:.4f}"
    return f"{point:.3e}"


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def update_main_benchmark_table(table, uncertainty):
    lookup = uncertainty.set_index(["dataset", "method_panel"])
    for row in table.rows[1:]:
        dataset = DATASET_LABELS.get(row.cells[0].text.strip())
        if dataset is None:
            continue
        for column, method in METHOD_COLUMNS.items():
            key = (dataset, method)
            if key not in lookup.index:
                continue
            record = lookup.loc[key]
            cell = row.cells[column]
            lines = cell.text.splitlines()
            if len(lines) < 2 or " | " not in lines[1]:
                continue
            metric, timing = lines[1].split(" | ", 1)
            interval = format_ci(
                record["point_estimate"],
                record["ci_lower"],
                record["ci_upper"],
            )
            lines[1] = f"{metric} {interval} | {timing}"
            cell.text = "\n".join(lines)
    c16.style_table(table, font_size=4.75)
    c16.prevent_row_splitting(table)
    repeat_header(table.rows[0])


def update_imagenet_table(table, imagenet):
    lookup = imagenet.set_index(["representation", "dimension"])
    for row in table.rows[1:]:
        representation = row.cells[0].text.strip()
        dimension = int(row.cells[1].text.replace(",", ""))
        key = (representation, dimension)
        if key not in lookup.index:
            continue
        record = lookup.loc[key]
        row.cells[3].text = (
            f'{record["top1"]:.4f}\n'
            f'[{record["top1_lower"]:.4f}, {record["top1_upper"]:.4f}]'
        )
        row.cells[4].text = (
            f'{record["top5"]:.4f}\n'
            f'[{record["top5_lower"]:.4f}, {record["top5_upper"]:.4f}]'
        )
    c16.style_table(table, font_size=4.8)
    c16.prevent_row_splitting(table)
    repeat_header(table.rows[0])


def revise_main(uncertainty, imagenet):
    document = Document(MAIN_SOURCE)

    design = c16.find_paragraph(document, "Within each dataset")
    design.text = (
        "Within each dataset, methods used identical fixed outer splits. For "
        "the main selected-point summary, five-fold cross-validation on the "
        "training partition selected the component count separately for each "
        "PLS family using accuracy for classification and RMSD for multivariate "
        "regression; NMR used the repeated training-only selection described "
        "below. The selected model was then evaluated once on the untouched "
        "outer test partition. CPU/CUDA rSVD comparisons within a PLS family "
        "used the same preprocessing, response, component count, and argmax or "
        "regression prediction rule and are therefore estimator matched. "
        "Comparisons across PLS families, LDA heads, or external packages are "
        "complete-workflow comparisons and are labelled separately. "
        "Component-path figures remain supplementary and are not interpreted "
        "as unbiased performance after test-set selection. Accuracy was the "
        "proportion of correctly decoded labels; top-5 accuracy was the "
        "proportion for which the observed label occurred among the five "
        "highest class scores. Classification uncertainty was quantified by "
        "two-sided 95% Wilson score intervals from the held-out correct/error "
        "counts. For multivariate regression, two-sided 95% percentile "
        "intervals were obtained from 10,000 held-out-sample bootstrap "
        "resamples; each resample retained the complete response-error profile "
        "of a sampled test observation before RMSD was recomputed. These "
        "intervals are conditional on the prespecified outer split and do not "
        "include variation caused by drawing another training/test split. "
        "Runtime comprised fitting and prediction. Peak resident host memory "
        "and peak CUDA memory were recorded separately. Successful repeated "
        "runs were summarized by the median and IQR; these computational "
        "replicates quantify timing dispersion, not predictive uncertainty. "
        "Failures, timeouts, and operating-system kills were retained. For NMR "
        "SIMPLS, mean RMSD and its standard error were calculated across the "
        "same five paired training-only splits. The smallest component count "
        "whose mean was no more than one standard error above the minimum mean "
        "was selected."
    )

    results = c16.find_paragraph(document, "Table 1 and Figure 2")
    results.text = (
        "Table 1 and Figure 2 show the complete twelve-task biomedical "
        "benchmark at the best training-validation component value within each "
        "prespecified grid and PLS family. A dagger identifies a lower or upper "
        "tested-grid boundary or a PLS-SVD response-rank limit; these values are "
        "not claimed as global optima. Every cell retains the predictive "
        "metric with its 95% interval, total fitting-plus-prediction time, peak "
        "host RSS, sampled GPU memory, precision, replicate count, and status. "
        "The intervals were widest where test sets were small. On CCLE "
        "(n_test=71), the four accuracy intervals overlapped broadly, so their "
        "point estimates do not establish a ranking. TCGA-HNSC methylation "
        "produced observed accuracy 1.000 for every family, but the Wilson "
        "lower bound was 0.938 because only 58 observations were held out. In "
        "contrast, the large Retina, Tabula Muris, and CIFAR-100 test sets gave "
        "narrow intervals. CPU/CUDA rows were matched within a family and "
        "component count, and Retina and Tabula Muris are reported separately."
    )

    table_caption = c16.find_paragraph(document, "Table 1.")
    table_caption.text += (
        " Brackets give 95% Wilson score intervals for accuracy and "
        "10,000-resample held-out-sample bootstrap intervals for RMSD. These "
        "predictive intervals are conditional on the fixed outer split; timing "
        "IQR quantifies only computational-run dispersion."
    )

    figure_caption = c16.find_paragraph(document, "Figure 2.")
    figure_caption.text += (
        " Error bars are 95% Wilson score intervals for classification and "
        "10,000-resample held-out-sample bootstrap intervals for multivariate "
        "RMSD; they are conditional on the prespecified outer split."
    )
    c16.replace_picture_before_caption(
        document,
        "Figure 2.",
        FIGURE,
        Inches(7.15),
    )
    update_main_benchmark_table(document.tables[0], uncertainty)

    imagenet_results = c16.find_paragraph(
        document, "ImageNet/DINOv2 was used"
    )
    imagenet_results.text = (
        "ImageNet/DINOv2 was used for a million-sample computational stress "
        "test after feature extraction and an exploratory evaluation of PLS as "
        "supervised feature reduction, not as biomedical validation. Exact "
        "CUDA cosine kNN on raw 1,024-dimensional embeddings achieved "
        "top-1/top-5 accuracy of 0.6556/0.9392 with Wilson 95% intervals "
        "0.6538-0.6574 and 0.9383-0.9401, respectively, 510.8 s median query "
        "time, 19.2 GB peak host RSS, and 5.9 GB sampled GPU memory. Across "
        "three independently fitted randomized representations, 200-dimensional "
        "PLS scores achieved median top-1/top-5 accuracy of 0.6521/0.9397 "
        "(top-1 range 0.6516-0.6521), 114.2 s transformation time, and 132.0 s "
        "held-out projection-plus-query time. For the representative seed-123 "
        "fit shown in Table 2, the corresponding intervals were "
        "0.6498-0.6534 and 0.9388-0.9406. Median host/GPU memory was 9.9/2.9 "
        "GB. The dimension-matched PCA control achieved median top-1/top-5 "
        "accuracy of 0.6430/0.9383 and required 233.3 s for transformation. "
        "Thus PLS compressed the representation 5.12-fold and reduced held-out "
        "inference 3.9-fold, with a small top-1 loss and overlapping top-5 "
        "intervals relative to raw features. Table 2 and Figure 4 show the "
        "representative fixed-seed path; independent-fit ranges are reported in "
        "Supplementary Table S23. The uncertainty intervals are conditional on "
        "the fixed split, and these natural-image results remain exploratory."
    )

    imagenet_caption = c16.find_paragraph(document, "Table 2.")
    imagenet_caption.text += (
        " Brackets below top-1 and top-5 estimates are 95% Wilson score "
        "intervals conditional on this fixed test set."
    )
    update_imagenet_table(document.tables[2], imagenet)

    document.core_properties.title = (
        "fastPLS CMPB manuscript - predictive uncertainty cycle 19"
    )
    document.save(MAIN_OUT)


def uncertainty_rows(uncertainty):
    labels = {value: key for key, value in DATASET_LABELS.items()}
    family = {
        "plssvd": "PLS-SVD",
        "simpls": "SIMPLS",
        "opls": "OPLS",
        "kernelpls": "kernel PLS",
    }
    rows = []
    for _, record in uncertainty.iterrows():
        rows.append(
            (
                labels[record["dataset"]],
                family[record["method_panel"]],
                int(record["effective_ncomp"]),
                record["metric_name"].upper()
                if record["metric_name"] == "rmsd"
                else "Accuracy",
                format_estimate(record["point_estimate"]),
                format_ci(
                    record["point_estimate"],
                    record["ci_lower"],
                    record["ci_upper"],
                ),
                int(record["n_test"]),
                (
                    "Wilson"
                    if record["metric_name"] == "accuracy"
                    else "Sample bootstrap"
                ),
            )
        )
    return rows


def imagenet_rows(imagenet):
    rows = []
    for _, record in imagenet.iterrows():
        rows.append(
            (
                record["representation"],
                int(record["dimension"]),
                f'{record["top1"]:.4f}',
                f'[{record["top1_lower"]:.4f}, {record["top1_upper"]:.4f}]',
                f'{record["top5"]:.4f}',
                f'[{record["top5_lower"]:.4f}, {record["top5_upper"]:.4f}]',
                int(record["n_test"]),
            )
        )
    return rows


def revise_supplement(uncertainty, imagenet):
    document = Document(SUPP_SOURCE)
    heading = document.add_heading(
        "S26. Predictive uncertainty on the prespecified outer test sets",
        level=1,
    )
    heading.paragraph_format.page_break_before = True
    document.add_paragraph(
        "Predictive uncertainty was evaluated separately from computational "
        "repeatability. For classification, two-sided 95% Wilson score "
        "intervals were calculated from the number of correct predictions and "
        "the held-out sample size. Wilson intervals remain informative at "
        "boundary estimates such as 100% accuracy, unlike a direct "
        "nonparametric bootstrap of identical correct predictions. For "
        "multivariate regression, 10,000 bootstrap samples were drawn with "
        "replacement from held-out observations. The complete response-error "
        "vector for each sampled observation was retained and the global RMSD "
        "was recomputed, thereby preserving dependence among response "
        "variables within an observation. All intervals are conditional on the "
        "predefined outer split and quantify finite-test-set uncertainty; they "
        "do not estimate variability from repeating preprocessing, model "
        "selection, or the training/test split. Runtime IQRs elsewhere in the "
        "manuscript describe repeated execution of a fixed workflow and must "
        "not be interpreted as predictive confidence intervals."
    )
    caption = document.add_paragraph(
        "Table S30. Predictive estimates and 95% intervals for every "
        "family-specific selected workflow in Table 1. Wilson intervals are "
        "used for accuracy; multivariate RMSD uses a 10,000-resample "
        "held-out-sample percentile bootstrap."
    )
    caption.style = "Caption"
    table = c16.add_table(
        document,
        [
            "Dataset",
            "Family",
            "k",
            "Metric",
            "Estimate",
            "95% CI",
            "n_test",
            "Method",
        ],
        uncertainty_rows(uncertainty),
        font_size=5.0,
    )
    c16.prevent_row_splitting(table)
    repeat_header(table.rows[0])

    document.add_paragraph(
        "The widest classification intervals occurred in CCLE, TCGA-BRCA, "
        "MetRef, and TCGA-HNSC methylation because their held-out sets contained "
        "58-100 observations. Consequently, method point estimates in these "
        "datasets should not be ranked without considering interval overlap. "
        "Intervals were much narrower in CIFAR-100 and the two single-cell "
        "tasks because their test sets were substantially larger."
    )
    caption = document.add_paragraph(
        "Table S31. Wilson 95% score intervals for the representative seed-123 "
        "ImageNet/DINOv2 retrieval paths in Table 2. These intervals describe "
        "finite-test-set uncertainty on the fixed 281,167-image test set; "
        "variation across independently fitted randomized representations is "
        "reported separately in Table S23."
    )
    caption.style = "Caption"
    table = c16.add_table(
        document,
        [
            "Representation",
            "Dim.",
            "Top-1",
            "Top-1 95% CI",
            "Top-5",
            "Top-5 95% CI",
            "n_test",
        ],
        imagenet_rows(imagenet),
        font_size=5.2,
    )
    c16.prevent_row_splitting(table)
    repeat_header(table.rows[0])

    document.core_properties.title = (
        "fastPLS CMPB supplement - predictive uncertainty cycle 19"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    document.add_heading(
        "20. Predictive uncertainty was inadequate for small test sets",
        level=1,
    )
    document.add_paragraph(
        "Reviewer comment: Several test sets contain only 58-100 observations, "
        "yet accuracy was presented without confidence intervals. Repeated "
        "computational runs do not quantify predictive uncertainty. Bootstrap "
        "intervals, repeated outer splits, or nested cross-validation should be "
        "provided for the principal predictive claims."
    )
    document.add_paragraph(
        "Response: Corrected. We now report two-sided 95% Wilson score intervals "
        "for every classification estimate in the twelve-task benchmark and "
        "the ImageNet analysis. Wilson intervals were chosen because they remain "
        "informative for boundary results: for example, observed accuracy 1.000 "
        "on the 58-sample TCGA-HNSC test set now has interval 0.938-1.000. For "
        "CBMC CITE-seq, PRISM, and NMR multivariate regression, we generated or "
        "reused the selected-model held-out prediction matrices and calculated "
        "10,000-resample percentile bootstrap intervals by resampling complete "
        "test-observation error profiles. Table 1, Figure 2, and Table 2 now "
        "show predictive intervals; Supplementary Tables S30 and S31 provide "
        "the complete numerical results and methods. The Methods and Results "
        "now explicitly distinguish these intervals from the IQR across "
        "repeated timing runs. We also state the limitation precisely: the new "
        "intervals quantify finite-test-set uncertainty conditional on the "
        "prespecified outer split and do not include training-split or "
        "model-selection variability. Accordingly, claims for small datasets "
        "were softened wherever intervals overlap."
    )
    document.core_properties.title = (
        "fastPLS CMPB response - predictive uncertainty cycle 19"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    uncertainty = pd.read_csv(UNCERTAINTY_CSV)
    imagenet = pd.read_csv(IMAGENET_CSV)
    revise_main(uncertainty, imagenet)
    revise_supplement(uncertainty, imagenet)
    revise_response()
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
