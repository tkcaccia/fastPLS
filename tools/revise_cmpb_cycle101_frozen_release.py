from copy import deepcopy
from pathlib import Path
from shutil import copy2

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle100"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle101"
OUTPUT.mkdir(parents=True, exist_ok=True)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle100_0.99.25_20260825.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle100_0.99.25_20260825.docx"
MAIN_OUTPUT = OUTPUT / "fastPLS_CMPB_main_cycle101_0.99.25_20260825.docx"
SUPP_OUTPUT = OUTPUT / "fastPLS_CMPB_supplement_cycle101_0.99.25_20260825.docx"

RESULTS = ROOT / "benchmark_results" / "frozen_release_0.99.25"
FIGURES = RESULTS / "figures"
VERSION = "0.99.25"
COMMIT = "7887401b09e25f54a546a253c255741cb1ab48e5"
ARCHIVE_SHA = "604f74d72f5e9540f7378efd37f2ca6ed87ccb9e73b912211331a8cd97233481"


def paragraph_by_prefix(document, prefix):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}; found {len(matches)}")
    return matches[0]


def replace_paragraph(paragraph, text):
    paragraph.text = text


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def replace_figure(document, caption_prefix, image_path, width):
    caption = paragraph_by_prefix(document, caption_prefix)
    previous = caption._p.getprevious()
    searched = 0
    while previous is not None and searched < 8:
        prior = previous.getprevious()
        if previous.xpath(".//w:drawing") or previous.xpath(".//w:pict"):
            previous.getparent().remove(previous)
            break
        if "".join(previous.itertext()).strip():
            break
        previous = prior
        searched += 1
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(image_path), width=Inches(width))
    caption._p.addprevious(p._p)
    caption.paragraph_format.keep_together = True


def set_cell_text(cell, value, size=6.5):
    cell.text = str(value)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.size = Pt(size)


def replace_table_rows(table, rows, size=6.5):
    template = deepcopy(table.rows[1]._tr if len(table.rows) > 1 else table.rows[0]._tr)
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)
    for values in rows:
        new_tr = deepcopy(template)
        table._tbl.append(new_tr)
        row = table.rows[-1]
        if len(row.cells) != len(values):
            raise RuntimeError(f"Table expects {len(row.cells)} values; got {len(values)}")
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value, size=size)


def remove_table(table):
    table._tbl.getparent().remove(table._tbl)


def fnum(x, digits=3):
    if pd.isna(x):
        return "NA"
    return f"{float(x):.{digits}f}"


def update_main():
    d = Document(MAIN_SOURCE)
    replace_paragraph(paragraph_by_prefix(d, "Methods:"),
        "Methods: All central evidence was regenerated with frozen fastPLS 0.99.25 "
        f"from Git commit {COMMIT[:12]} and source archive SHA-256 {ARCHIVE_SHA}. "
        "The implementation uses compiled shape-dependent execution, incremental updates, "
        "compact latent prediction, and optional implicit cross-covariance products. Fixed-control "
        "CPU IRLBA was the deterministic numerical reference; approximate rSVD and accelerator "
        "routes were audited separately under prespecified criteria.")
    replace_paragraph(paragraph_by_prefix(d, "Results:"),
        "Results: Frozen-release fastPLS SIMPLS met all 117 prespecified deterministic numerical "
        "tolerance checks. In 108 repeated single-CPU comparisons with pls::simpls.fit, accuracy "
        "was identical; fastPLS was faster on five of nine datasets under ordinary public workflows, "
        "with a largest speed-up of 4.85-fold. For the 13,000 by 28,355 NMR problem at 50 SIMPLS "
        "components, CPU IRLBA, CPU rSVD, and CUDA rSVD required 692.21, 152.09, and 3.91 s, "
        "respectively, with RMSD 0.0007561, 0.0007560, and 0.0007561. The exploratory ImageNet "
        "stress test completed 1,000,000-sample fitting and blocked prediction of 281,167 held-out "
        "embeddings within 9.9 GB peak host RSS.")
    replace_paragraph(paragraph_by_prefix(d, "Direction extraction is modular"),
        "Direction extraction is modular rather than the principal estimator contribution. float64 "
        "CPU fitting supports fixed-control IRLBA [15] as the deterministic numerical reference and "
        "approximate rSVD [16], which uses a Gaussian range sketch, orthonormalization, power "
        "iterations, and a reduced decomposition. IRLBA is an iterative truncated solver, not an "
        "exact dense SVD. Frozen release 0.99.25 uses rSVD oversampling 20 and two power iterations, "
        "the CPU/CUDA controls that met every prespecified check across five seeds. Explicit "
        "unqualified overrides warn and are recorded in model diagnostics. Exact controls, seeds, "
        "audit thresholds, and route-specific qualification are consolidated in Supplementary "
        "Sections S13-S14 and Tables S8-S9.")
    replace_paragraph(paragraph_by_prefix(d, "Two comparisons answered"),
        "Two comparisons answered different questions. Numerical-kernel validation compared "
        "deterministic float64 fastPLS SIMPLS with de Jong SIMPLS in pls::simpls.fit. A separate "
        "frozen-release CPU software comparison used fastPLS 0.99.25 and IKPLS 6.1.2 NumPy "
        "Algorithms 1 and 2 on identical stored splits, externally training-centred float64 predictors, "
        "centred one-hot responses, component counts, and final held-out predictions. Each route ran "
        "three times in a fresh process with one effective thread. Because IKPLS and SIMPLS are "
        "different estimators and retain different state, this is an end-to-end software comparison, "
        "not an estimator-matched benchmark. An earlier JAX/CUDA development comparison was not "
        "reproduced from the frozen environment and is excluded from central evidence.")
    replace_paragraph(paragraph_by_prefix(d, "Twelve tasks covered"),
        "The broader benchmark design covered twelve tasks spanning metabolomics, NMR, CITE-seq, "
        "tissue and cancer omics, single-cell transcriptomics, drug response, and CIFAR-100 "
        "[7,20-27,30]. The frozen central reruns reported here comprise deterministic SIMPLS "
        "validation and external comparison, controlled scaling and solver qualification, selected "
        "CPU/CUDA routes, NMR, and ImageNet. Methods used identical stored splits and training-only "
        "component grids within each analysis. Classification used accuracy and Wilson intervals; "
        "multivariate regression used RMSD, independent-test Q² relative to the training-response "
        "mean, and held-out bootstrap intervals. Runtime included fitting and prediction. Absolute "
        "host RSS was the isolated-process high-water mark; baseline-corrected peak RSS was peak "
        "minus the pre-fit baseline and remains a complete-process rather than algorithmic-workspace "
        "measure. Formula-based dense-object sizes were reported separately. GPU increments likewise "
        "include runtime-context allocation.")
    replace_paragraph(paragraph_by_prefix(d, "The CPU comparison completed"),
        "The frozen CPU comparison completed all 36 planned runs. IKPLS NumPy Algorithm 2 was "
        "fastest: median totals were 0.00059 s on Breast, 0.00298 s on MetRef, and 0.218 s on "
        "CIFAR-100, compared with 0.005, 0.033, and 3.585 s for fastPLS rSVD. Breast accuracy was "
        "identical (94.29%). IKPLS reached 75.0% on MetRef and 70.95% on CIFAR-100; fastPLS rSVD "
        "reached 77.0% and 72.13%, while deterministic fastPLS IRLBA reached 77.0% and 70.77%. "
        "CIFAR-100 complete-process peak RSS was 584 MB for IKPLS Algorithm 2 and 1,178 MB for "
        "fastPLS rSVD, including different language runtimes and allocators.")
    remove_paragraph(paragraph_by_prefix(d, "On the NVIDIA RTX 5060 Ti, the cold end-to-end CUDA comparison"))
    replace_paragraph(paragraph_by_prefix(d, "Repeated outer partitions showed"),
        "Component counts are described as best within their prespecified training-selection grids, "
        "not as unconstrained optima. Historical repeated-partition sensitivity analyses are retained "
        "only as development records in the repository and do not support frozen-release main-text claims.")
    replace_paragraph(paragraph_by_prefix(d, "Same-code ablations"),
        "The frozen controlled study isolated matrix shape while keeping rSVD controls fixed. "
        "Automatic routes completed every requested run, but numerical qualification depended on "
        "retained components, class count, cross-covariance rank, and the CUDA response-dimension "
        "sweep. Runtime and memory claims therefore exclude every route outside tolerance rather than "
        "treating successful execution as numerical validation.")
    replace_paragraph(paragraph_by_prefix(d, "Deterministic fastPLS SIMPLS met"),
        "Deterministic fastPLS SIMPLS met the prespecified numerical tolerances in all 117 "
        "component-level comparisons with de Jong SIMPLS. Approximate rSVD was evaluated separately, "
        "and only fully audited controls support numerical-agreement claims. Development-stage OPLS, "
        "kernel-PLS, same-code ablation, and direct PLS-SVD/SIMPLS shape studies remain explicitly "
        "archival and do not support the frozen central performance claims.")
    replace_paragraph(paragraph_by_prefix(d, "The strict comparison completed"),
        "The strict comparison completed all 108 planned runs: nine datasets, two output profiles, "
        "two implementations, and three fresh-process repetitions. Accuracy was identical for every "
        "pair. With minimum common prediction outputs, fastPLS was faster on four datasets and "
        "pls::simpls.fit on five; the largest fastPLS advantage was 1.48-fold on GTEx v8. Under "
        "ordinary public workflows, fastPLS was faster on five datasets, including 2.39-fold on "
        "CIFAR-100, 3.35-fold on Retina, and 4.85-fold on Tabula Muris. Corresponding accuracies were "
        "0.8739 (8,739/10,000; Wilson 95% CI 0.8672-0.8803), 0.9678 "
        "(21,684/22,406; 0.9654-0.9700), and 0.8006 (40,077/50,059; 0.7971-0.8041). These timing "
        "profiles answer different questions and are not pooled (Figure 2; Supplementary Tables "
        "S10a-S10d).")
    replace_paragraph(paragraph_by_prefix(d, "Workflow gains partly reflected"),
        "Workflow gains partly reflected output policy. On CIFAR-100, the compact fastPLS fit object "
        "was 1.39 MB versus 7,777.99 MB for ordinary pls::simpls.fit. Median complete-process peak "
        "RSS was 2,493.57 versus 13,415.25 MB; from comparable 1,906.80 and 1,906.62 MB pre-fit "
        "baselines, corrected peak increments were 586.77 versus 11,508.84 MB. The theoretical "
        "largest retained dense object was the 0.586-MB final coefficient matrix for fastPLS and a "
        "3,814.70-MB fitted or residual response path for pls::simpls.fit. These values describe the "
        "specified complete workflow, not isolated algorithmic workspace. When both methods retained "
        "complete coefficient paths, fit objects were 59.30 and 58.60 MB, corrected increments were "
        "69.22 and 127.13 MB, and the speed-up was 1.21-fold. The broader package panel remains a "
        "workflow comparison with implementation-specific outputs (Supplementary Tables S10c-S10e).")
    replace_paragraph(paragraph_by_prefix(d, "A direct matched-shape timing study"),
        "Development-stage same-code ablations and direct PLS-SVD/SIMPLS shape comparisons helped "
        "design the final routing policy, but were not rerun from the frozen archive and no longer "
        "support main-text speed claims. They remain explicitly labelled archival analyses in the "
        "Supplementary repository index.")
    replace_paragraph(paragraph_by_prefix(d, "Hardware acceleration remained"),
        "Hardware acceleration remained route and shape dependent. Frozen-release selected CPU/CUDA "
        "SIMPLS-rSVD workflows were compared on identical splits and component counts (Figure 3; "
        "Supplementary Table S11). CPU and CUDA accuracies differed by at most 0.0003 across MetRef, "
        "Retina, and CIFAR-100; CUDA was slower on the two smaller tasks but reduced CIFAR-100 total "
        "time from 8.07 to 0.724 s. The broader 0.01/0.99 multi-seed audit qualified oversampling 20 "
        "with two power iterations in 585/585 CPU and 40/40 CUDA checks; rejected controls and all "
        "discordant Metal rSVD routes remain quarantined. CPU IRLBA remains the deterministic "
        "numerical reference. float32 halves stored input representation but did not uniformly improve "
        "runtime, process memory, or agreement (Supplementary Tables S8-S9).")
    replace_paragraph(paragraph_by_prefix(d, "Figure 2."),
        "Figure 2. Repeated deterministic float64 single-CPU SIMPLS public workflows with fastPLS "
        "0.99.25 and pls 2.9.0. Panels show median total fitting-plus-prediction time, baseline-corrected "
        "complete-process peak RSS, and held-out argmax accuracy. Error bars are IQRs from three fresh "
        "processes per method-dataset pair. Splits and component counts were identical, and accuracy "
        "was identical for every pair. Full values and the separate minimum-output comparison are in "
        "Supplementary Tables S10a-S10d.")
    replace_paragraph(paragraph_by_prefix(d, "Figure 3."),
        "Figure 3. Selected frozen-release CPU and CUDA SIMPLS-rSVD workflows. Panels show total "
        "runtime, post-prediction minus pre-fit host-RSS snapshots, and paired held-out accuracy for "
        "MetRef, Retina, and CIFAR-100. Points and IQRs summarize three runs. Every route used "
        "oversampling 20, two power iterations, seed 123, the same split, and the same component "
        "count. Host-memory values are snapshots rather than isolated workspace peaks; CUDA values "
        "include runtime/context allocation.")
    replace_paragraph(paragraph_by_prefix(d, "At the family-selected settings"),
        "At the family-selected settings, PLS-SVD at five components achieved RMSD 0.001043 and "
        "Q² 0.98916 across CPU IRLBA, CPU rSVD, and CUDA rSVD. SIMPLS at 50 components achieved "
        "RMSD 0.00075608, 0.00075595, and 0.00075606 and Q² 0.994299, 0.994301, and 0.994299, "
        "respectively. Figure 4 displays held-out sample AMI-0030-9 (index 38), selected by the "
        "prespecified rule of closest per-spectrum RMSD to the median under 50-component SIMPLS "
        "CUDA rSVD; it was not chosen for visual concordance.")
    replace_paragraph(paragraph_by_prefix(d, "A separate matched solver/backend analysis"),
        "A matched family-selected analysis held split, precision, and component count fixed. At five "
        "PLS-SVD components, CPU IRLBA, CPU rSVD, and CUDA rSVD required 262.86, 4.57, and 0.422 s. "
        "At 50 SIMPLS components, corresponding times were 692.21, 152.09, and 3.91 s. All rSVD "
        "rows used oversampling 20, two power iterations, and seed 123 and met the prespecified "
        "approximate-route tolerances. A second implementation-only table fixes both families at 100 "
        "components and is reported separately in the Supplement because changing family and "
        "component count answers a different question.")
    replace_paragraph(paragraph_by_prefix(d, "Figure 4."),
        "Figure 4. Frozen-release NMR prediction and computation. Panels A-C compare family-selected "
        "held-out RMSD, total time, and host-memory measurements; the deposited 165-component "
        "PLS-SVD/IRLBA workflow is historical scientific context, not a matched implementation. "
        "Panel D shows per-spectrum RMSD for all 321 held-out spectra. Panels E-F overlay observed "
        "and predicted intensities for AMI-0030-9 (index 38), the sample closest to the median "
        "SIMPLS CUDA/rSVD error, over the full response range and 1.7-0.5 ppm. Predictor water-region "
        "columns (4.6-4.8 ppm) were zeroed in training and test predictors; all 28,355 response "
        "coordinates remained in every metric. rSVD used oversampling 20, two power iterations, and seed 123.")
    replace_paragraph(paragraph_by_prefix(d, "Exploratory ImageNet experiment 1"),
        "The exploratory ImageNet classification experiment used frozen fastPLS 0.99.25 from the "
        f"SHA-256-identified archive ({ARCHIVE_SHA[:12]}...) on 1,000,000 training and 281,167 held-out "
        "1,024-dimensional embeddings. The route was label-aware float32 SIMPLS with CUDA rSVD and "
        "argmax or native CUDA LDA; sequential SIMPLS deflation and score projection remained "
        "host-resident, so it was hybrid. rSVD used oversampling 20, two power iterations, and seed 123.")
    replace_paragraph(paragraph_by_prefix(d, "A single shared component-path fit"),
        "One shared fit supplied 100-1,000-component prefixes; 1,000 was a boundary stress point, "
        "not a selected optimum. At 1,000 components, argmax achieved top-1/top-5 accuracy "
        "0.79980/0.95613 and LDA achieved 0.80938/0.93929. LDA top-1 represented 227,571/281,167 "
        "correct predictions (Wilson 95% CI 0.80792-0.81083). Argmax fitting and blocked prediction "
        "required 947.99 and 1,202.07 s; LDA required 955.06 and 270.27 s. Prediction used 5,000-row "
        "blocks. Peak host RSS was 9,884 MB for argmax and 9,718 MB for LDA; peak device used memory "
        "was 782 and 4,792 MB. Because the split was noncanonical and measurements were single runs, "
        "these are exploratory feasibility estimates.")
    replace_paragraph(paragraph_by_prefix(d, "Figure 5."),
        "Figure 5. Exploratory frozen-release ImageNet/DINOv2 SIMPLS stress test. Seed 123 assigned "
        "1,000,000 rows from the pooled 1,281,167-row embedding archive to training and 281,167 to "
        "a noncanonical holdout. Top-1 and top-5 accuracy are shown for argmax and LDA across "
        "100-1,000 requested components from one shared fit; 1,000 is a boundary stress point, not an "
        "optimum. Runtime separates fitting from 5,000-row blocked prediction. Memory values include "
        "runtime/context allocation. rSVD used oversampling 20, two power iterations, and seed 123; "
        "all values are single-run exploratory measurements.")
    replace_paragraph(paragraph_by_prefix(d, "The IKPLS comparison places"),
        "The frozen IKPLS comparison places the contribution in the high-performance PLS landscape. "
        "Improved Kernel PLS was substantially faster in the tested single-thread CPU workflows, "
        "whereas fastPLS offered an R-native de Jong SIMPLS path, multivariate-response storage "
        "controls, nested validation, multiple PLS families, and route diagnostics. This does not "
        "establish universal superiority of either software. The cross-language result is interpreted "
        "as end-to-end workflow evidence, separate from deterministic estimator validation.")
    replace_paragraph(paragraph_by_prefix(d, "Code and benchmark outputs are available"),
        "Code and benchmark outputs are available at https://github.com/tkcaccia/fastPLS; reusable "
        "components are at https://github.com/tkcaccia/kodama-cpp. All central quantitative evidence "
        f"was regenerated with fastPLS {VERSION}, Git commit {COMMIT}, from execution archive "
        f"fastPLS_0.99.25.tar.gz (SHA-256 {ARCHIVE_SHA}). The documented archive was generated from "
        "the same commit and has SHA-256 ecddd795e67990ffaf5fb2d690414c7494497261159e4ba9099d2912360c961b. "
        "The frozen bundle contains scripts, manifests, generated tables, session information, the "
        "vignette, reference manual, and a machine-readable file ledger. The deposited 165-component "
        "NMR workflow remains explicitly labelled historical context. A persistent archive identifier "
        "will replace the checksum-only review object before acceptance.")

    replace_figure(d, "Figure 2.", FIGURES / "Figure_2_frozen_external_simpls.png", 6.55)
    replace_figure(d, "Figure 3.", FIGURES / "Figure_3_frozen_cpu_cuda.png", 6.55)
    replace_figure(d, "Figure 4.", FIGURES / "Figure_4_frozen_nmr.png", 6.55)
    replace_figure(d, "Figure 5.", FIGURES / "Figure_5_frozen_imagenet.png", 6.55)
    d.save(MAIN_OUTPUT)


def external_rows(profile):
    x = pd.read_csv(RESULTS / "external_simpls" / "external_simpls_timing_summary.csv")
    x = x[x.comparison_profile == profile]
    rows = []
    for dataset in sorted(x.dataset.unique()):
        a = x[(x.dataset == dataset) & (x.implementation == "fastpls")].iloc[0]
        b = x[(x.dataset == dataset) & (x.implementation == "pls")].iloc[0]
        rows.append((dataset, profile.replace("_", " "),
                     f"{a.median_total_sec:.3f} ({a.iqr_total_sec:.3f})",
                     f"{b.median_total_sec:.3f} ({b.iqr_total_sec:.3f})",
                     f"{b.median_total_sec/a.median_total_sec:.2f}",
                     f"{a.median_accuracy:.4f} / {b.median_accuracy:.4f}",
                     f"{int(a.repetitions_completed)} / {int(b.repetitions_completed)}"))
    return rows


def memory_rows(profile):
    x = pd.read_csv(RESULTS / "external_simpls" / "external_simpls_timing_summary.csv")
    x = x[x.comparison_profile == profile]
    rows = []
    for dataset in sorted(x.dataset.unique()):
        a = x[(x.dataset == dataset) & (x.implementation == "fastpls")].iloc[0]
        b = x[(x.dataset == dataset) & (x.implementation == "pls")].iloc[0]
        rows.append((dataset,
                     f"{a.median_process_peak_rss_mb:.1f} / {b.median_process_peak_rss_mb:.1f}",
                     f"{a.median_prefit_process_rss_mb:.1f} / {b.median_prefit_process_rss_mb:.1f}",
                     f"{a.median_baseline_corrected_peak_increment_mb:.1f} / {b.median_baseline_corrected_peak_increment_mb:.1f}",
                     f"{a.theoretical_largest_retained_name}: {a.theoretical_largest_retained_mb:.2f} / "
                     f"{b.theoretical_largest_retained_name}: {b.theoretical_largest_retained_mb:.2f}"))
    return rows


def update_supplement():
    d = Document(SUPP_SOURCE)
    # Separate the frozen central CUDA environment from archival Metal diagnostics.
    for table in d.tables:
        for row in table.rows:
            for cell in row.cells:
                if "fastPLS 0.99.6. See Table S15" in cell.text:
                    set_cell_text(
                        cell,
                        "R 4.6.0; frozen central analyses used fastPLS 0.99.25 from commit "
                        f"{COMMIT[:12]}. See Table S15 for the archive checksum and analysis mapping.",
                        size=5.5,
                    )
                elif "fastPLS 0.99.6; package commit recorded by each Metal run" in cell.text:
                    set_cell_text(
                        cell,
                        "R 4.6.0; Metal diagnostics used development source states and are archival, "
                        "not frozen central evidence. See Table S15.",
                        size=5.5,
                    )
    replace_paragraph(paragraph_by_prefix(d, "This supplement distinguishes"),
        f"All central quantitative evidence in this supplement was regenerated with frozen fastPLS "
        f"{VERSION}, Git commit {COMMIT}, from execution archive SHA-256 {ARCHIVE_SHA}. The exact "
        "archive, scripts, input checksums, session information, generated tables, and per-file ledger "
        "are indexed in Table S15. Older results are excluded from central claims; the deposited NMR "
        "workflow is retained only as explicitly labelled historical scientific context.")
    replace_paragraph(paragraph_by_prefix(d, "S12.1 Execution ablation"),
        "S12.1 Archival execution ablation and direct PLS-family timing")
    replace_paragraph(paragraph_by_prefix(d, "Each ablation changed"),
        "These development-stage ablations predate the frozen central rerun and are retained only to "
        "document how the execution design was developed; they do not support release-level speed "
        "claims. Each ablation changed one internal feature relative to a minimally optimized compiled "
        "SIMPLS baseline while holding data, split, estimator, deterministic IRLBA solver, component "
        "count, seed, and prediction head fixed. Three isolated runs were used per configuration. "
        "The measured ranges quantify shape-dependent trade-offs rather than universal speed-ups.")
    replace_paragraph(paragraph_by_prefix(d, "For a linear operator M"),
        "For a linear operator M and target rank r, rSVD draws a Gaussian matrix, forms a range "
        "sketch, applies alternating products with M and its transpose, orthonormalizes the range, and "
        "decomposes the reduced matrix. Frozen release 0.99.25 uses oversampling 20 and two power "
        "iterations as the effective CPU/CUDA default. This configuration met all prespecified checks "
        "in the five-seed qualification panel; qualification applies to the controls and audited panel, "
        "not every future fit. Explicit unqualified controls warn and are stored in diagnostics.")
    replace_paragraph(paragraph_by_prefix(d, "Current benchmark workflows record"),
        "Every central benchmark used the frozen 0.99.25 execution archive and recorded its SHA-256, "
        "package version, input checksum, benchmark-script identity, compiler, BLAS/LAPACK, thread "
        "settings, accelerator libraries, seed, rSVD controls, repetition, and status. The NMR and "
        "ImageNet analyses were rerun from that archive. Table S15 and the machine-readable ledger "
        "map each claim to its exact frozen result directory.")
    replace_paragraph(paragraph_by_prefix(d, "The historical single-seed CPU audit"),
        "A historical single-seed development audit found 101/117 tolerance checks at oversampling "
        "10 with one power iteration. Frozen-release qualification therefore used oversampling 20, "
        "two power iterations, and seeds 1, 7, 19, 43, and 123. All 585 CPU checks and all 40 CUDA "
        "checks met the prespecified tolerances; rejected CUDA 10/1 and 10/2 controls met 0/40 and "
        "27/40 checks. Authoritative component-level outputs are in benchmark_results/"
        "frozen_release_0.99.25/simpls_estimator_preservation/ and rsvd_cuda_reliability_final/. "
        "These are solver-control reliability screens, not estimator-equivalence evidence.")
    replace_paragraph(paragraph_by_prefix(d, "Figure S1."),
        "Figure S1. Frozen-release controlled one-factor SIMPLS scaling. Median total time and "
        "baseline-corrected complete-process peak RSS are shown across sample, predictor, response, "
        "component, requested-prefix, rank, class-count, and cross-covariance-size sweeps. Open points "
        "were outside prespecified numerical tolerance and are excluded from speed claims. Each point "
        "summarizes three isolated runs; all 486 requested runs completed.")
    replace_paragraph(paragraph_by_prefix(d, "This comparison is not estimator matched"),
        "This frozen-release comparison is not estimator matched. IKPLS 6.1.2 implements "
        "Dayal-MacGregor Improved Kernel PLS, whereas fastPLS implements de Jong SIMPLS with "
        "deterministic IRLBA or approximate rSVD. The common CPU contract was float64 input, "
        "externally applied training centring, identical centred one-hot responses, identical splits and "
        "component counts, final held-out prediction, three fresh-process repetitions, and one effective "
        "thread. RSS includes language-runtime allocation and is a workflow-feasibility measurement.")
    replace_paragraph(paragraph_by_prefix(d, "Table S10f."),
        "Table S10f. Frozen-release single-thread CPU end-to-end comparison with IKPLS. Time and IQR "
        "are fit-plus-prediction seconds; RSS values are MB. rSVD used oversampling 20, two power "
        "iterations, and seed 123.")
    remove_paragraph(paragraph_by_prefix(d, "Table S10g."))
    replace_paragraph(paragraph_by_prefix(d, "Reproducibility. CPU results"),
        "Reproducibility. Frozen CPU results are in benchmark_results/frozen_release_0.99.25/"
        "ikpls_cross_language_cpu/. The benchmark used fastPLS 0.99.25, IKPLS 6.1.2, float64, three "
        "fresh processes, and fixed component counts of 10 (Breast), 22 (MetRef), and 50 "
        "(CIFAR-100). The earlier JAX/CUDA development comparison was not reproducible from the "
        "frozen environment and is not retained as central evidence.")
    replace_paragraph(paragraph_by_prefix(d, "Every row uses the component count selected"),
        "This focused frozen-release analysis uses identical stored splits and component counts for "
        "paired CPU/CUDA SIMPLS-rSVD routes: MetRef 22, Retina 20, and CIFAR-100 100. All rows use "
        "oversampling 20, two power iterations, seed 123, and three repetitions. Host memory is the "
        "post-prediction minus pre-fit RSS snapshot, not a peak workspace measure; device memory "
        "includes runtime/context allocation.")
    replace_paragraph(paragraph_by_prefix(d, "Table S11."),
        "Table S11. Frozen-release paired CPU/CUDA SIMPLS-rSVD results. Values use identical splits, "
        "component counts, and qualified controls; time is median (IQR) over three runs.")
    replace_paragraph(paragraph_by_prefix(d, "The paired backend analysis changes"),
        "The family-selected paired analysis compares CPU IRLBA, CPU rSVD, and CUDA rSVD at five "
        "PLS-SVD or 50 SIMPLS components. The deposited 165-component workflow uses the original "
        "centring-only protocol and remains historical context. Predictor columns strictly between "
        "4.6 and 4.8 ppm were set to zero in both training and test predictor matrices; no response "
        "coordinate was removed from metrics. Main-text Figure 4 displays AMI-0030-9 (index 38), "
        "selected as the held-out sample closest to the median 50-component SIMPLS CUDA/rSVD "
        "per-spectrum RMSD.")
    replace_paragraph(paragraph_by_prefix(d, "Table S12."),
        "Table S12. Frozen-release NMR evidence. Family-selected prediction, the separate fixed-100-"
        "component implementation comparison, and the deposited historical workflow answer different "
        "questions. fastPLS host memory is post-fit minus pre-fit RSS; historical memory is a measured "
        "peak increment and is not directly equivalent.")
    replace_paragraph(paragraph_by_prefix(d, "The pooled archive contained"),
        "The pooled archive contained 1,281,167 precomputed DINOv2 embeddings with 1,024 features "
        "and 1,000 classes. Seed 123 assigned 1,000,000 rows to training and 281,167 to a noncanonical "
        "holdout. Frozen fastPLS 0.99.25 used label-aware float32 SIMPLS, CUDA rSVD oversampling 20, "
        "two power iterations, seed 123, and argmax or CUDA LDA. Sequential deflation and score "
        "projection were host-resident. Prediction used 5,000-row blocks and online metric accumulation; "
        "all 20 requested component/head rows succeeded. The separate FAISS experiment has a different "
        "retrieval objective and is not pooled with classification.")
    replace_paragraph(paragraph_by_prefix(d, "Table S13."),
        "Table S13. Frozen-release exploratory ImageNet classification. One shared fit supplies all "
        "component prefixes for each head; fit time is therefore repeated only to expose the common "
        "cost. The 1,000-component row is a boundary stress point, not a selected optimum. All rows "
        "are single-run feasibility estimates.")
    replace_paragraph(paragraph_by_prefix(d, "S19. Repeated outer-partition uncertainty"),
        "S19. Archival repeated outer-partition sensitivity analysis")
    replace_paragraph(paragraph_by_prefix(d, "MetRef, GTEx v8, and Retina used"),
        "This development-stage sensitivity analysis was not regenerated from the frozen 0.99.25 "
        "archive and is excluded from central release claims. It is retained to document variability "
        "observed during method development. MetRef, GTEx v8, and Retina used ten stratified 80/20 "
        "outer partitions with five-fold training-only component selection; NMR used five random "
        "80/20 partitions and three-fold selection. The empirical ranges are descriptive, not "
        "confidence intervals for the frozen-release analyses.")
    replace_paragraph(paragraph_by_prefix(d, "Table S14."),
        "Table S14. Archival repeated-partition predictive dispersion and selection sensitivity. "
        "These development results are not part of the frozen 0.99.25 central evidence.")
    replace_paragraph(paragraph_by_prefix(d, "The ledger never infers"),
        f"All central rows resolve to frozen fastPLS {VERSION}, Git commit {COMMIT}, and execution "
        f"archive SHA-256 {ARCHIVE_SHA}. The machine-readable file ledger contains SHA-256 values for "
        "185 evidence files. Within the central provenance ledger, the deposited 165-component NMR "
        "workflow is the sole historical context row and is never relabelled as frozen-release "
        "evidence. Other development analyses retained elsewhere in this supplement are explicitly "
        "marked archival and excluded from central claims.")
    replace_paragraph(paragraph_by_prefix(d, "Table S15."),
        "Table S15. Frozen-release central analysis provenance. Archive paths are relative to the "
        "repository root; full per-file checksums are in provenance/frozen_evidence_file_ledger.csv.")

    # Frozen deterministic comparison and memory tables.
    replace_table_rows(d.tables[16], external_rows("estimator_kernel"), size=6.2)
    replace_table_rows(d.tables[18], memory_rows("estimator_kernel"), size=5.8)
    replace_table_rows(d.tables[19], memory_rows("complete_workflow"), size=5.8)

    # Frozen CPU IKPLS comparison.
    ik = pd.read_csv(RESULTS / "ikpls_cross_language_cpu" / "ikpls_cross_language_summary.csv")
    ik_rows = []
    for _, r in ik.iterrows():
        label = r.implementation.replace("IKPLS_numpy_alg", "IKPLS NumPy Algorithm ")
        label = label.replace("fastPLS_cpu_irlba", "fastPLS SIMPLS / IRLBA")
        label = label.replace("fastPLS_cpu_rsvd", "fastPLS SIMPLS / rSVD")
        ik_rows.append((r.dataset.title() if r.dataset != "cifar100" else "CIFAR-100", label,
                        f"{100*r.accuracy:.2f}", f"{r.median_total_sec:.6f}",
                        f"{r.iqr_total_sec:.6f}", f"{r.median_peak_rss_mb:.1f}",
                        f"{r.median_incremental_peak_rss_mb:.1f}"))
    replace_table_rows(d.tables[21], ik_rows, size=6.0)
    remove_table(d.tables[22])

    # Focused selected CPU/CUDA table.
    sb = pd.read_csv(RESULTS / "selected_backend" / "selected_backend_summary.csv")
    sb_rows = []
    for _, r in sb.iterrows():
        sb_rows.append((r.dataset.replace("cifar100", "CIFAR-100").replace("metref", "MetRef").replace("retina", "Retina"),
                        "SIMPLS", int(r.ncomp), "fixed paired point", f"accuracy {r.accuracy:.4f}",
                        f"{r.median_total_time_sec:.3f} ({r.time_iqr_sec:.3f})", f"{r.median_rss_after_prediction_mb-r.median_rss_before_fit_mb:.1f}",
                        f"{r.median_gpu_after_prediction_mb-r.median_gpu_before_fit_mb:.1f}" if r.backend == "cuda" else "NA",
                        f"{r.backend.upper()}; success"))
    replace_table_rows(d.tables[22], sb_rows, size=6.0)
    for cell, value in zip(d.tables[22].rows[0].cells,
                           ("Dataset", "Family", "A", "Setting", "Metric", "Time median (IQR), s",
                            "Host RSS delta MB", "GPU delta MB", "Status")):
        set_cell_text(cell, value, size=6.0)

    # NMR frozen routes plus historical context.
    ns = pd.read_csv(RESULTS / "nmr" / "derived" / "nmr_frozen_route_summary.csv")
    nrows = []
    for _, r in ns.iterrows():
        analysis = "family selected" if int(r.ncomp) in (5, 50) else "fixed 100 components"
        fam = "PLS-SVD" if r.family == "plssvd" else "SIMPLS"
        impl = f"{r.backend.upper()} / {r.solver.upper()}"
        metric = f"RMSD {r.RMSD_median:.7f}; Q² {r.Q2_median:.6f}"
        host = max(0.0, r.after_fit_rss_mb_median-r.baseline_rss_mb_median)
        nrows.append((analysis, fam, impl, int(r.ncomp), metric,
                      f"{r.total_time_sec_median:.3f}", f"post-fit {host:.1f}", "not captured",
                      "success; rSVD 20/2/123" if r.solver == "rsvd" else "success; deterministic"))
    nrows.append(("historical deposited workflow", "PLS-SVD", "deposited IRLBA", 165,
                  "RMSD 0.0007099", "447.601", "peak 3605.5", "NA",
                  "historical context; different workflow/hardware"))
    replace_table_rows(d.tables[23], nrows, size=5.7)

    # ImageNet frozen 20-row path.
    im = pd.read_csv(RESULTS / "imagenet" / "imagenet_all_results.csv")
    irows = []
    for _, r in im.iterrows():
        irows.append(("classification", r.classifier.upper(), int(r.ncomp),
                      f"{r.top1_accuracy:.5f}", f"{r.top5_accuracy:.5f}",
                      f"{r.total_time_sec:.2f}", f"{r.rss_peak_predict_mb:.0f}",
                      f"{r.gpu_peak_predict_mb:.0f}", r.audit_status))
    replace_table_rows(d.tables[24], irows, size=5.8)

    # Frozen provenance ledger (plus the explicitly historical comparator).
    prov = [
        ("F01", "Execution archive", "publication_release_0.99.25/", VERSION, COMMIT[:12], "frozen_release_manifest.tsv", ARCHIVE_SHA[:12]),
        ("F02", "Regression contracts", "benchmark_results/frozen_release_0.99.25/regression_contract_*", VERSION, COMMIT[:12], "run_release_regression_contract.R", ARCHIVE_SHA[:12]),
        ("F03", "SIMPLS estimator validation", "benchmark_results/frozen_release_0.99.25/simpls_estimator_preservation", VERSION, COMMIT[:12], "run_simpls_estimator_preservation.R", ARCHIVE_SHA[:12]),
        ("F04", "Controlled scaling", "benchmark_results/frozen_release_0.99.25/controlled_scaling", VERSION, COMMIT[:12], "controlled_scaling/run_grid.R", ARCHIVE_SHA[:12]),
        ("F05", "External SIMPLS comparison", "benchmark_results/frozen_release_0.99.25/external_simpls", VERSION, COMMIT[:12], "external_simpls_timing/run_benchmark.R", ARCHIVE_SHA[:12]),
        ("F06", "IKPLS CPU comparison", "benchmark_results/frozen_release_0.99.25/ikpls_cross_language_cpu", VERSION, COMMIT[:12], "ikpls_cross_language/run_benchmark.py", ARCHIVE_SHA[:12]),
        ("F07", "Selected CPU/CUDA", "benchmark_results/frozen_release_0.99.25/selected_backend", VERSION, COMMIT[:12], "run_selected_backend.R", ARCHIVE_SHA[:12]),
        ("F08", "NMR", "benchmark_results/frozen_release_0.99.25/nmr", VERSION, COMMIT[:12], "run_nmr_frozen.R", ARCHIVE_SHA[:12]),
        ("F09", "ImageNet", "benchmark_results/frozen_release_0.99.25/imagenet", VERSION, COMMIT[:12], "benchmark_imagenet_qualified_top5_path.R", ARCHIVE_SHA[:12]),
        ("H01", "Deposited NMR workflow", "benchmark_results/manuscript_revision_cycle64_20260726", "historical", "deposited source", "FastPLS (1).R", "historical"),
    ]
    replace_table_rows(d.tables[26], prov, size=5.4)

    replace_figure(d, "Figure S1.", FIGURES / "Figure_S1_frozen_scaling.png", 6.55)
    d.save(SUPP_OUTPUT)


def copy_assets():
    target = OUTPUT / "figures"
    target.mkdir(exist_ok=True)
    for path in [
        FIGURES / "Figure_2_frozen_external_simpls.png",
        FIGURES / "Figure_3_frozen_cpu_cuda.png",
        FIGURES / "Figure_4_frozen_nmr.png",
        FIGURES / "Figure_5_frozen_imagenet.png",
        FIGURES / "Figure_S1_frozen_scaling.png",
    ]:
        copy2(path, target / path.name)
    (OUTPUT / "README.md").write_text(
        "# Cycle 101\n\nCentral manuscript and supplementary evidence regenerated with frozen "
        f"fastPLS {VERSION}, commit {COMMIT}, execution archive SHA-256 {ARCHIVE_SHA}.\n",
        encoding="utf-8",
    )


if __name__ == "__main__":
    update_main()
    update_supplement()
    copy_assets()
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)
