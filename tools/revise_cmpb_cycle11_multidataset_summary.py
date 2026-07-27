from pathlib import Path
import math

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import revise_cmpb_cycle9_simpls_validation as c9


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle10"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle11"
RESULTS = ROOT / "benchmark_results" / "manuscript_multidataset_summary_20260725"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle10_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle10_0.99.6_20260725.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle11_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle11_0.99.6_20260725.docx"
RESPONSE_OUT = OUT / "response_multidataset_summary_20260725.docx"

METHODS = ["plssvd", "simpls", "opls", "kernelpls"]
METHOD_LABELS = {
    "plssvd": "PLS-SVD",
    "simpls": "SIMPLS",
    "opls": "OPLS",
    "kernelpls": "kernel PLS",
}
DATASET_LABELS = {
    "metref": "MetRef",
    "ccle": "CCLE",
    "tcga_brca": "TCGA-BRCA",
    "tcga_hnsc_methylation": "TCGA-HNSC methyl.",
    "gtex_v8": "GTEx v8",
    "tcga_pan_cancer": "TCGA Pan-Cancer",
    "singlecell": "SingleCell",
    "cifar100": "CIFAR-100",
    "cbmc_citeseq": "CBMC CITE-seq",
    "prism": "PRISM",
    "nmr": "NMR",
    "imagenet": "ImageNet",
}


def finite(value):
    try:
        return math.isfinite(float(value))
    except (TypeError, ValueError):
        return False


def fmt_num(value, digits=3):
    if not finite(value):
        return "NA"
    value = float(value)
    if abs(value) < 0.001 and value != 0:
        return f"{value:.2e}"
    return f"{value:.{digits}f}"


def fmt_memory(value):
    if not finite(value):
        return "—"
    return f"{float(value):.0f}"


def backend_label(value):
    value = str(value).lower()
    if "cuda" in value or "gpu" in value:
        return "CUDA rSVD"
    if "metal" in value:
        return "Metal rSVD"
    if "irlba" in value:
        return "CPU IRLBA"
    return "CPU rSVD"


def metric_label(value):
    value = str(value).lower()
    return {"accuracy": "Acc", "rmsd": "RMSD", "q2": "Q²"}.get(value, value)


def table_cell(row):
    status = str(row.get("execution_status", "missing"))
    if not int(row.get("n_success", 0) or 0):
        return f"Not completed\n{status.replace('_', ' ')}"
    star = "*" if status.startswith("exploratory") else ""
    k = int(float(row["effective_ncomp"])) if finite(row.get("effective_ncomp")) else "NA"
    gpu = fmt_memory(row.get("peak_gpu_mem_mb"))
    precision = str(row.get("execution_precision", row.get("precision", "float64")))
    precision = "f32" if "32" in precision else "f64"
    return (
        f"{backend_label(row.get('backend'))} | k={k}{star}\n"
        f"{metric_label(row.get('metric_name'))}={fmt_num(row.get('metric_value'))} | "
        f"{fmt_num(row.get('total_time_sec'), 2)} s\n"
        f"H={fmt_memory(row.get('peak_host_rss_mb'))}; G={gpu} MB | {precision} | {status}"
    )


def move_after(anchor, element):
    anchor._p.addnext(element)


def add_main_table(document, anchor, main_rows):
    rows = []
    for dataset in DATASET_LABELS:
        values = [DATASET_LABELS[dataset]]
        for method in METHODS:
            hit = main_rows[
                main_rows["dataset"].astype(str).eq(dataset)
                & main_rows["method_panel"].astype(str).eq(method)
            ]
            values.append(table_cell(hit.iloc[0].to_dict()) if len(hit) else "Not evaluated")
        rows.append(values)

    caption = document.add_paragraph(
        "Table 1. Multi-dataset summary at the component count selected from "
        "training data. Each cell reports the fastest completed CPU/CUDA rSVD "
        "backend for that fixed PLS family and component count, followed by the "
        "outer-test predictive metric, total fitting-plus-prediction time, peak "
        "host RSS (H), sampled peak GPU memory (G), execution precision, and "
        "status. Within-cell CPU/CUDA comparisons are estimator matched; comparisons "
        "across columns are workflow comparisons between different model families. "
        "ImageNet (*) is a single-run exploratory stress point selected from the "
        "reported outer grid rather than by training-only validation and is not used "
        "for estimator-performance claims.",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True
    table = c9.add_table(
        document,
        ["Dataset", "PLS-SVD", "SIMPLS", "OPLS", "kernel PLS"],
        rows,
        [1.03, 1.36, 1.36, 1.36, 1.36],
        font_size=6.25,
    )
    move_after(anchor, caption._p)
    caption._p.addnext(table._tbl)


def add_supplement_tables(document, full):
    document.add_page_break()
    document.add_heading(
        "S15. Complete selected-point multi-dataset benchmark", level=1
    )
    document.add_paragraph(
        "The tables below retain both CPU and CUDA rSVD rows, all failed or "
        "unavailable executions, precision, and resource measurements. Component "
        "counts were selected using training data only, except for the explicitly "
        "labelled exploratory ImageNet stress point. Argmax decoding was used for "
        "classification and the prespecified regression metric for numeric "
        "responses. Consequently, backend comparisons within one model family are "
        "estimator matched. Results involving LDA, external packages, or different "
        "PLS families are complete-workflow comparisons and are interpreted "
        "separately."
    )
    table_number = 14
    for method in METHODS:
        part = full[full["method_panel"].astype(str).eq(method)].copy()
        rows = []
        for _, row in part.iterrows():
            status = str(row.get("execution_status", ""))
            rows.append(
                (
                    DATASET_LABELS.get(str(row["dataset"]), str(row["dataset"])),
                    backend_label(row.get("backend")) if int(row.get("n_success", 0) or 0) else "—",
                    int(float(row["effective_ncomp"])) if finite(row.get("effective_ncomp")) else "—",
                    metric_label(row.get("metric_name")),
                    fmt_num(row.get("metric_value")),
                    fmt_num(row.get("total_time_sec"), 2),
                    fmt_memory(row.get("peak_host_rss_mb")),
                    fmt_memory(row.get("peak_gpu_mem_mb")),
                    "f32" if "32" in str(row.get("execution_precision")) else "f64",
                    status,
                )
            )
        caption = document.add_paragraph(
            f"Table S{table_number}. {METHOD_LABELS[method]} selected-point "
            "benchmark. Time is total fitting plus prediction; H and G are peak "
            "host RSS and sampled GPU memory in MB.",
            style="Caption",
        )
        caption.paragraph_format.keep_with_next = True
        c9.add_table(
            document,
            ["Dataset", "Backend", "k", "Metric", "Value", "Time (s)", "H", "G", "Prec.", "Status"],
            rows,
            [1.18, 0.78, 0.42, 0.55, 0.58, 0.62, 0.55, 0.55, 0.48, 0.94],
            font_size=6.4,
        )
        table_number += 1
    document.add_paragraph(
        "Machine-readable versions are provided as "
        "multidataset_selected_backend_summary.csv, "
        "multidataset_selected_main_rows.csv, and "
        "multidataset_selected_main_wide.csv. Full component paths remain in the "
        "supplementary benchmark archive and are not used to select the displayed "
        "outer-test point."
    )


def revise_main(main_rows):
    document = Document(MAIN_SOURCE)
    methods = c9.find_paragraph(document, "Within each dataset, methods used identical")
    c9.set_paragraph_text(
        methods,
        "Within each dataset, methods used identical fixed outer splits. For the "
        "main selected-point summary, five-fold cross-validation on the training "
        "partition selected the component count separately for each PLS family "
        "using accuracy for classification and RMSD for multivariate regression; "
        "NMR used the repeated training-only selection described below. The selected "
        "model was then evaluated once on the untouched outer test partition. "
        "CPU/CUDA rSVD comparisons within a PLS family used the same preprocessing, "
        "response, component count, and argmax or regression prediction rule and "
        "are therefore estimator matched. Comparisons across PLS families, LDA "
        "heads, or external packages are complete-workflow comparisons and are "
        "labelled separately. Component-path figures remain supplementary and are "
        "not interpreted as unbiased performance after test-set selection. Accuracy "
        "was the proportion of correctly decoded labels; top-5 accuracy was the "
        "proportion for which the observed label occurred among the five highest "
        "class scores. Runtime comprised fitting and prediction. Peak resident host "
        "memory and peak CUDA memory were recorded separately. Successful repeated "
        "runs were summarized by the median; failures, timeouts, and operating-system "
        "kills were retained."
    )

    result_intro = c9.find_paragraph(document, "Across completed benchmarks")
    c9.set_paragraph_text(
        result_intro,
        "Table 1 makes the complete 12-task benchmark visible at one prespecified "
        "selected point per dataset and model family. The selected component count "
        "was determined from training data before outer-test evaluation, and every "
        "cell retains predictive performance, total time, host memory, GPU memory, "
        "precision, and execution status. The only exception is ImageNet, which is "
        "explicitly labelled as a single-run exploratory stress test because an "
        "independent full-scale training-only component search was not available."
    )

    backend = c9.find_paragraph(document, "CPU execution was generally preferable")
    completed = main_rows[main_rows["n_success"].fillna(0).gt(0)]
    cpu_count = completed["backend"].astype(str).str.contains("cpu", case=False).sum()
    gpu_count = completed["backend"].astype(str).str.contains("cuda|gpu", case=False).sum()
    c9.set_paragraph_text(
        backend,
        f"Among the {len(completed)} completed dataset-family cells, the fastest "
        f"selected rSVD execution was CPU in {cpu_count} and CUDA in {gpu_count}. "
        "CPU was generally preferable for small matrices, where transfer and kernel "
        "launch overhead exceeded the accelerated work. CUDA became advantageous "
        "for large dense products, high component counts, or large response "
        "dimensions. Large sample count alone did not guarantee a GPU benefit. "
        "Predictive values should be compared between CPU and CUDA rows within a "
        "fixed family; differences between PLS-SVD, SIMPLS, OPLS, kernel PLS, "
        "argmax, and LDA represent workflow differences rather than backend-only "
        "effects."
    )
    add_main_table(document, backend, main_rows)
    document.core_properties.title = (
        "fastPLS CMPB manuscript - complete selected-point benchmark"
    )
    document.save(MAIN_OUT)


def revise_supplement(full):
    document = Document(SUPP_SOURCE)
    add_supplement_tables(document, full)
    document.core_properties.title = (
        "fastPLS CMPB supplement - complete selected-point benchmark"
    )
    document.save(SUPP_OUT)


def write_response(main_rows):
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(11)
    document.styles["Normal"].paragraph_format.space_after = Pt(7)
    document.styles["Heading 1"].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    document.add_heading(
        "Response to reviewer: complete multi-dataset benchmark", level=1
    )
    reviewer = document.add_paragraph()
    reviewer.add_run("Reviewer comment. ").bold = True
    reviewer.add_run(
        "The complete multi-dataset benchmark is not visible in the main Results. "
        "Add a main-text summary at the best training-selected component count for "
        "each dataset and method, displaying predictive metric, total time, host "
        "memory, GPU memory where applicable, precision, and execution status. "
        "Distinguish estimator-matched from workflow comparisons."
    )
    response = document.add_paragraph()
    response.add_run("Response. ").bold = True
    response.add_run(
        "We agree. We added main-text Table 1, which covers all 12 tasks and all "
        "four model families. Component count was selected separately for each "
        "family by training-only cross-validation before outer-test evaluation; "
        "NMR used five repeated training-only splits. Every table cell now reports "
        "the selected count, predictive metric, total fitting-plus-prediction time, "
        "peak host RSS, sampled peak GPU memory, precision, and status. The "
        "ImageNet stress point is marked exploratory because no independent "
        "full-scale component-selection run was available."
    )
    document.add_paragraph(
        "We also added Supplementary Section S15 and Tables S14-S17, which retain "
        "both CPU and CUDA rows rather than only the fastest backend, as well as "
        "all failures and unavailable methods. The Methods and Results now define "
        "CPU/CUDA comparisons within the same PLS family, component count, response, "
        "and argmax/regression rule as estimator matched. Comparisons across PLS "
        "families, LDA heads, or external packages are explicitly described as "
        "complete-workflow comparisons and are not used as estimator-only speed "
        "claims. Full component paths remain supplementary."
    )
    document.add_paragraph(
        f"The revised main summary contains "
        f"{int(main_rows['n_success'].fillna(0).gt(0).sum())} completed "
        "dataset-family cells; missing or non-prespecified combinations remain "
        "visible with their execution status rather than being dropped."
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    full = pd.read_csv(RESULTS / "multidataset_selected_backend_summary.csv")
    main_rows = pd.read_csv(RESULTS / "multidataset_selected_main_rows.csv")
    revise_main(main_rows)
    revise_supplement(full)
    write_response(main_rows)
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
