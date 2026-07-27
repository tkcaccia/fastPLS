#!/usr/bin/env python3

from pathlib import Path
from shutil import copy2
import subprocess

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle15"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle16"
EVIDENCE = (
    ROOT / "benchmark_results" / "manuscript_revision_cycle16_20260725"
)
NMR_DIR = EVIDENCE / "nmr_plssvd_extended_lower_grid"
OLD_SELECTED = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle13_20260725"
    / "selected_backend_cycle13_chosen.csv"
)
SIMPLS_SELECTION = (
    ROOT
    / "benchmark_results"
    / "review_nmr_extended_selection_20260725"
    / "nmr_component_selection_summary.csv"
)
PLOT_SCRIPT = ROOT / "tools" / "plot_cmpb_cycle16_component_boundaries.R"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle15_0.99.6_20260725.docx"
SUPP_SOURCE = (
    SOURCE / "fastPLS_CMPB_supplement_cycle15_0.99.6_20260725.docx"
)
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle15_20260725.docx"
)

MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle16_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle16_0.99.6_20260725.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle16_20260725.docx"

ALL_DATASET_FIGURE = EVIDENCE / "plots" / "selected_performance_all_datasets.png"
NMR_SELECTION_FIGURE = (
    EVIDENCE / "plots" / "nmr_component_selection_plssvd_simpls.png"
)

DATASET_LABELS = {
    "metref": "MetRef",
    "ccle": "CCLE",
    "tcga_brca": "TCGA-BRCA",
    "tcga_hnsc_methylation": "TCGA-HNSC methyl.",
    "gtex_v8": "GTEx v8",
    "tcga_pan_cancer": "TCGA Pan-Cancer",
    "retina": "Retina",
    "tabula": "Tabula Muris",
    "cifar100": "CIFAR-100",
    "cbmc_citeseq": "CBMC CITE-seq",
    "prism": "PRISM",
    "nmr": "NMR",
}
METHOD_LABELS = {
    "plssvd": "PLS-SVD",
    "simpls": "SIMPLS",
    "opls": "OPLS",
    "kernelpls": "kernel PLS",
}

GRIDS = {
    "metref": [2, 5, 10, 22, 50, 100],
    "ccle": [2, 5, 10, 18, 50, 100],
    "tcga_brca": [2, 5, 10, 20, 50],
    "tcga_hnsc_methylation": [2, 5, 10, 20, 50],
    "gtex_v8": [2, 5, 10, 20, 32, 50, 100],
    "tcga_pan_cancer": [2, 5, 10, 20, 50, 100],
    "retina": [2, 5, 10, 20, 50],
    "tabula": [2, 5, 10, 20, 50],
    "cifar100": [2, 5, 10, 20, 50, 100, 200],
    "cbmc_citeseq": [2, 5, 10, 20, 50],
    "prism": [2, 5, 10, 20, 50, 100],
    "nmr": [1, 2, 3, 5, 7, 10, 25, 50, 75, 100, 125, 150, 165, 175, 200, 250, 300],
}

RESPONSE_RANK_LIMITS = {
    ("metref", "plssvd"): 21,
    ("ccle", "plssvd"): 17,
    ("tcga_hnsc_methylation", "plssvd"): 1,
    ("gtex_v8", "plssvd"): 31,
    ("tcga_pan_cancer", "plssvd"): 31,
    ("retina", "plssvd"): 11,
    ("tabula", "plssvd"): 31,
    ("cifar100", "plssvd"): 99,
    ("cbmc_citeseq", "plssvd"): 10,
}

METHOD_ORDER = ["plssvd", "simpls", "opls", "kernelpls"]
DATASET_ORDER = [
    "metref",
    "ccle",
    "tcga_brca",
    "tcga_hnsc_methylation",
    "gtex_v8",
    "tcga_pan_cancer",
    "retina",
    "tabula",
    "cifar100",
    "cbmc_citeseq",
    "prism",
    "nmr",
]


def find_paragraph(document, start):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(start):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {start}")


def set_paragraph(document, start, text):
    paragraph = find_paragraph(document, start)
    paragraph.text = text
    return paragraph


def set_cell_margins(cell, top=45, start=55, bottom=45, end=55):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table, font_size=5.8):
    table.style = "Table"
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if row_index == 0
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
                    run.bold = row_index == 0


def prevent_row_splitting(table):
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def fill_table(table, columns, rows, font_size=5.8):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for index, label in enumerate(columns):
        table.rows[0].cells[index].text = str(label)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    style_table(table, font_size=font_size)


def add_table(document, columns, rows, font_size=5.8):
    table = document.add_table(rows=1, cols=len(columns))
    for index, label in enumerate(columns):
        table.rows[0].cells[index].text = str(label)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    style_table(table, font_size=font_size)
    return table


def replace_picture_before_caption(document, caption_start, image, width):
    paragraphs = document.paragraphs
    caption = find_paragraph(document, caption_start)
    caption_index = next(
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph._p is caption._p
    )
    for index in range(caption_index - 1, -1, -1):
        paragraph = paragraphs[index]
        if paragraph._p.xpath(".//w:drawing"):
            for child in list(paragraph._p):
                paragraph._p.remove(child)
            run = paragraph.add_run()
            run.add_picture(str(image), width=width)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
    raise RuntimeError(f"Image preceding caption not found: {caption_start}")


def selection_status(dataset, method, component):
    if pd.isna(component):
        return "not evaluated"
    component = int(component)
    cap = RESPONSE_RANK_LIMITS.get((dataset, method))
    if cap is not None and component == cap:
        return "response-rank limit"
    grid = GRIDS[dataset]
    if component == min(grid):
        return "lower tested-grid boundary"
    if component == max(grid):
        return "upper tested-grid boundary"
    return "interior tested value"


def build_evidence():
    selected = pd.read_csv(OLD_SELECTED)
    heldout_files = sorted((NMR_DIR / "heldout" / "rows").glob("*.csv"))
    heldout = pd.concat(
        [pd.read_csv(path) for path in heldout_files], ignore_index=True
    )
    nmr_mask = (
        selected["dataset"].eq("nmr")
        & selected["method_panel"].eq("plssvd")
    )
    selected.loc[nmr_mask, "variant_name"] = "fastpls_plssvd_cuda_rsvd"
    selected.loc[nmr_mask, "engine"] = "CUDA"
    selected.loc[nmr_mask, "backend"] = "cuda_rsvd"
    selected.loc[nmr_mask, "requested_ncomp"] = 5
    selected.loc[nmr_mask, "effective_ncomp"] = 5
    selected.loc[nmr_mask, "n_runs"] = len(heldout)
    selected.loc[nmr_mask, "metric_median"] = heldout["RMSD"].median()
    selected.loc[nmr_mask, "metric_q25"] = heldout["RMSD"].quantile(0.25)
    selected.loc[nmr_mask, "metric_q75"] = heldout["RMSD"].quantile(0.75)
    selected.loc[nmr_mask, "total_time_sec_median"] = heldout[
        "total_time_sec"
    ].median()
    selected.loc[nmr_mask, "total_time_sec_q25"] = heldout[
        "total_time_sec"
    ].quantile(0.25)
    selected.loc[nmr_mask, "total_time_sec_q75"] = heldout[
        "total_time_sec"
    ].quantile(0.75)
    selected.loc[nmr_mask, "host_rss_mb_median"] = heldout[
        "host_rss_mb"
    ].median()
    selected.loc[nmr_mask, "host_rss_mb_q25"] = heldout[
        "host_rss_mb"
    ].quantile(0.25)
    selected.loc[nmr_mask, "host_rss_mb_q75"] = heldout[
        "host_rss_mb"
    ].quantile(0.75)
    selected.loc[nmr_mask, "gpu_mem_mb_median"] = heldout[
        "gpu_peak_mb"
    ].median()
    selected.loc[nmr_mask, "gpu_mem_mb_q25"] = heldout[
        "gpu_peak_mb"
    ].quantile(0.25)
    selected.loc[nmr_mask, "gpu_mem_mb_q75"] = heldout[
        "gpu_peak_mb"
    ].quantile(0.75)

    selected["selection_status"] = [
        selection_status(row.dataset, row.method_panel, row.effective_ncomp)
        for row in selected.itertuples()
    ]
    selected.to_csv(
        EVIDENCE / "selected_backend_cycle16_chosen.csv", index=False
    )

    status_rows = []
    for dataset in DATASET_ORDER:
        row = {
            "dataset": dataset,
            "evaluated_grid": ",".join(str(x) for x in GRIDS[dataset]),
        }
        for method in METHOD_ORDER:
            hit = selected[
                selected["dataset"].eq(dataset)
                & selected["method_panel"].eq(method)
            ]
            if hit.empty or pd.isna(hit.iloc[0]["effective_ncomp"]):
                row[method] = "not evaluated"
            else:
                component = int(hit.iloc[0]["effective_ncomp"])
                status = hit.iloc[0]["selection_status"]
                row[method] = f"k={component}; {status}"
        status_rows.append(row)
    status = pd.DataFrame(status_rows)
    status.to_csv(EVIDENCE / "component_selection_status.csv", index=False)

    summary = pd.DataFrame(
        {
            "measure": [
                "selected_ncomp",
                "validation_RMSD_median",
                "validation_RMSD_q25",
                "validation_RMSD_q75",
                "heldout_RMSD_median",
                "heldout_total_time_sec_median",
                "heldout_host_rss_mb_median",
                "heldout_gpu_peak_mb_median",
                "heldout_repetitions",
            ],
            "value": [
                5,
                0.001507,
                0.001487,
                0.001547,
                heldout["RMSD"].median(),
                heldout["total_time_sec"].median(),
                heldout["host_rss_mb"].median(),
                heldout["gpu_peak_mb"].median(),
                len(heldout),
            ],
        }
    )
    summary.to_csv(EVIDENCE / "nmr_plssvd_cycle16_summary.csv", index=False)
    return selected, status


def table1_boundary_markers(document, selected):
    table = document.tables[0]
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    method_by_column = {
        index: {
            "PLS-SVD": "plssvd",
            "SIMPLS": "simpls",
            "OPLS": "opls",
            "kernel PLS": "kernelpls",
        }.get(label)
        for index, label in enumerate(headers)
    }
    dataset_by_label = {value: key for key, value in DATASET_LABELS.items()}
    for row in table.rows[1:]:
        dataset = dataset_by_label.get(row.cells[0].text.strip())
        if dataset is None:
            continue
        for column, method in method_by_column.items():
            if method is None:
                continue
            cell = row.cells[column]
            hit = selected[
                selected["dataset"].eq(dataset)
                & selected["method_panel"].eq(method)
            ]
            if hit.empty:
                continue
            record = hit.iloc[0]
            if pd.isna(record["effective_ncomp"]):
                continue
            if dataset == "nmr" and method == "plssvd":
                cell.text = (
                    "CUDA rSVD | k=5\n"
                    "RMSD=1.04e-03 | 0.90 s (IQR 0.01)\n"
                    "H=3289; G=582 MB | f64 | n=3"
                )
            if record["selection_status"] != "interior tested value":
                component = int(record["effective_ncomp"])
                cell.text = cell.text.replace(
                    f"k={component}", f"k={component}\u2020", 1
                )
    style_table(table, font_size=5.3)


def nmr_selection_table_rows():
    plssvd = pd.read_csv(
        NMR_DIR / "results" / "nmr_component_selection_summary.csv"
    ).set_index("ncomp")
    simpls = pd.read_csv(SIMPLS_SELECTION).set_index("ncomp")
    components = sorted(set(plssvd.index) | set(simpls.index))
    rows = []
    for component in components:
        if component in simpls.index:
            s = simpls.loc[component]
            simpls_median = f'{s["RMSD_median"]:.7f}'
            simpls_iqr = f'{s["RMSD_q25"]:.7f}-{s["RMSD_q75"]:.7f}'
        else:
            simpls_median = "\u2014"
            simpls_iqr = "\u2014"
        p = plssvd.loc[component]
        rows.append(
            (
                component,
                simpls_median,
                simpls_iqr,
                f'{p["RMSD_median"]:.7f}',
                f'{p["RMSD_q25"]:.7f}-{p["RMSD_q75"]:.7f}',
            )
        )
    return rows


def revise_main(selected):
    document = Document(MAIN_SOURCE)
    table1_boundary_markers(document, selected)
    # Keep each row of the four-panel NMR figure with its panel labels.
    prevent_row_splitting(document.tables[1])
    replace_picture_before_caption(
        document,
        "Figure 2.",
        ALL_DATASET_FIGURE,
        Inches(6.65),
    )
    set_paragraph(
        document,
        "Table 1 and Figure 2 show",
        "Table 1 and Figure 2 show the complete twelve-task biomedical "
        "benchmark at the best training-validation component value within each "
        "prespecified grid and PLS family. A dagger identifies a lower or upper "
        "tested-grid boundary or a PLS-SVD response-rank limit; these values are "
        "not claimed as global optima. Every cell retains the predictive metric, "
        "total fitting-plus-prediction time, peak host RSS, sampled GPU memory, "
        "precision, replicate count, and status. CPU/CUDA rows were matched "
        "within a family and component count. Retina and Tabula Muris are "
        "reported separately.",
    )
    set_paragraph(
        document,
        "Table 1.",
        "Table 1. Twelve-task biomedical workflow benchmark at the best "
        "training-validation value within each evaluated component grid. Each "
        "cell shows the completed CPU/CUDA rSVD workflow, effective component "
        "count, outer-test metric, median total fitting-plus-prediction time, "
        "peak host RSS, sampled GPU memory, precision, and completed runs. "
        "\u2020Lower or upper tested-grid boundary, or response-rank limit for "
        "PLS-SVD; these entries are not global optima.",
    )
    set_paragraph(
        document,
        "Figure 2.",
        "Figure 2. Outer-test predictive performance at the best "
        "training-validation value within each evaluated grid for the twelve "
        "biomedical benchmark tasks. Points show the fastest completed matched "
        "CPU/CUDA rSVD row within each PLS family. An asterisk marks a tested-grid "
        "boundary or response-rank limit; these values are not interpreted as "
        "global optima. Accuracy is higher-is-better and RMSD lower-is-better. "
        "ImageNet is excluded because it used the separate matched retrieval "
        "protocol in Figure 4.",
    )
    set_paragraph(
        document,
        "NMR represented the extreme",
        "NMR represented the extreme multivariate-response setting (1,200 "
        "training and 321 held-out spectra; p=13,000; q=28,355). Every method "
        "used the predefined outer split and full response spectrum; predictors "
        "were centred without variance scaling, and the routine 4.6-4.8 ppm "
        "residual-water interval was set to zero in both training and held-out "
        "predictor matrices. Five repeated training-only 80/20 splits evaluated "
        "1, 2, 3, 5, 7, 10, 25, 50, 75, 100, 125, 150, 165, 175, 200, 250, "
        "and 300 components for PLS-SVD. Its median validation RMSD reached an "
        "interior minimum at 5 components (0.001507), compared with 0.001540 at "
        "the former lower boundary of 10. SIMPLS retained its interior minimum "
        "at 100 components (0.0008917) over the 10-300 grid. The five-component "
        "CUDA PLS-SVD model achieved held-out RMSD 0.001043 with median total "
        "time 0.898 s across three repetitions. These are the best values within "
        "the evaluated grids, not assertions of global optimality. Figure 3 "
        "continues to compare implementations at the common prespecified "
        "100-component operating point to isolate computational behavior.",
    )
    document.core_properties.title = (
        "fastPLS CMPB manuscript - component-boundary revision cycle 16"
    )
    document.save(MAIN_OUT)


def revise_supplement(status):
    document = Document(SUPP_SOURCE)
    set_paragraph(
        document,
        "Two uses of component grids are distinguished.",
        "Two uses of component grids are distinguished. Benchmark trajectory "
        "figures evaluate prespecified component counts on a fixed test set and "
        "are descriptive. Training-only model selection reports the best value "
        "within the evaluated grid. Values at the lower or upper tested boundary, "
        "and PLS-SVD values constrained by response rank, are labelled explicitly "
        "and are never called optima. pls.single.cv() supports training-set "
        "selection, whereas pls.double.cv() provides an outer held-out layer when "
        "an unbiased estimate after tuning is required.",
    )
    set_paragraph(
        document,
        "Routine NMR spectral preprocessing and component selection.",
        "Routine NMR spectral preprocessing and component selection. The task "
        "comprised 1,200 training and 321 held-out spectra, 13,000 predictors, "
        "and 28,355 numeric responses. A shared loader verified identical "
        "columns, outer-split dimensions, checksum, and matrix signatures. The "
        "4.6-4.8 ppm residual-water interval was set to zero in Xtrain and Xtest "
        "before inner splitting or fitting; Ytrain and Ytest were unchanged. "
        "Predictors were centred without variance scaling. Five repeated 80/20 "
        "training-only splits used seeds 123, 456, 789, 1011, and 2027. The "
        "PLS-SVD grid was extended to 1, 2, 3, 5, and 7 below the former lower "
        "boundary of 10 while retaining the original values through 300. PLS-SVD "
        "reached an interior median RMSD minimum at 5 components (0.001507; "
        "IQR 0.001487-0.001547). SIMPLS reached an interior minimum at 100 "
        "components (0.0008917) over its 10-300 grid. The family-specific results "
        "prevent a universal NMR component claim; the separate 100-component "
        "reference comparison holds model size fixed to isolate computational "
        "behavior.",
    )
    set_paragraph(
        document,
        "Table S12b.",
        "Table S12b. Family-specific repeated training-only NMR selection. Values "
        "are median and interquartile-range validation RMSD over five fixed "
        "splits. PLS-SVD includes the extended lower grid; SIMPLS was evaluated "
        "over the prespecified 10-300 grid.",
    )
    fill_table(
        document.tables[5],
        [
            "Components",
            "SIMPLS median",
            "SIMPLS IQR",
            "PLS-SVD median",
            "PLS-SVD IQR",
        ],
        nmr_selection_table_rows(),
        font_size=5.6,
    )
    replace_picture_before_caption(
        document,
        "Figure S18.",
        NMR_SELECTION_FIGURE,
        Inches(6.65),
    )
    set_paragraph(
        document,
        "Figure S18.",
        "Figure S18. Repeated training-only NMR component selection. Grey lines "
        "show the five fixed inner splits; coloured lines and ribbons show the "
        "median and interquartile range. (A) The extended 1-300 PLS-SVD grid "
        "identifies an interior minimum at 5 components. The logarithmic x-axis "
        "makes the added low-component values visible. (B) SIMPLS retains an "
        "interior minimum at 100 components over the 10-300 grid.",
    )

    document.add_page_break()
    document.add_heading(
        "S25. Boundary-aware interpretation of component selection", level=1
    )
    document.add_paragraph(
        "Component counts are reported as the best training-validation values "
        "within prespecified grids. 'Upper' and 'lower' identify tested-grid "
        "boundaries; 'response-rank limit' identifies the maximum feasible "
        "PLS-SVD dimension imposed by the encoded response rank. None of these "
        "boundary labels establishes a global optimum. NMR PLS-SVD and SIMPLS "
        "have interior minima after extending the PLS-SVD grid below 10."
    )
    document.add_paragraph(
        "Table S29. Component-selection status for every dataset and PLS family "
        "shown in Table 1.",
        style="Caption",
    )
    rows = []
    for _, row in status.iterrows():
        rows.append(
            (
                DATASET_LABELS[row["dataset"]],
                row["evaluated_grid"],
                row["plssvd"],
                row["simpls"],
                row["opls"],
                row["kernelpls"],
            )
        )
    add_table(
        document,
        ["Dataset", "Evaluated grid", "PLS-SVD", "SIMPLS", "OPLS", "kernel PLS"],
        rows,
        font_size=5.2,
    )
    document.core_properties.title = (
        "fastPLS CMPB supplement - component-boundary revision cycle 16"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    set_paragraph(
        document,
        "SIMPLS and PLS-SVD were each evaluated",
        "SIMPLS was evaluated over 10-300 components and retained an interior "
        "median RMSD minimum at 100. We extended the PLS-SVD grid below the "
        "former lower boundary of 10 to include 1, 2, 3, 5, and 7 components "
        "while preserving the same five training-only splits and maximal "
        "300-component fit. PLS-SVD now has an interior minimum at 5 components "
        "(median validation RMSD 0.001507). A three-repetition full-training "
        "held-out evaluation at 5 components gave RMSD 0.001043 and median total "
        "time 0.898 s. All other boundary or response-rank-limited selections "
        "are explicitly marked and described as best within the evaluated grid, "
        "not as optima.",
    )
    document.add_heading(
        "17. Component selection remained boundary limited", level=1
    )
    document.add_paragraph(
        "Reviewer comment: Several reported component counts occur at the "
        "maximum tested value and should not be called optima. The NMR PLS-SVD "
        "grid should extend below 10."
    )
    document.add_paragraph(
        "Response: Corrected. We replaced optimality language with 'best within "
        "the evaluated grid,' added boundary and response-rank labels to Table 1 "
        "and Supplementary Table S29, and extended NMR PLS-SVD to 1, 2, 3, 5, "
        "and 7 components below 10. The repeated training-only curve identifies "
        "an interior minimum at 5 components, which was then evaluated in three "
        "isolated held-out runs."
    )
    document.core_properties.title = (
        "fastPLS CMPB response - component-boundary revision cycle 16"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    selected, status = build_evidence()
    subprocess.run(["Rscript", str(PLOT_SCRIPT)], check=True)
    revise_main(selected)
    revise_supplement(status)
    revise_response()
    for source in (
        EVIDENCE / "selected_backend_cycle16_chosen.csv",
        EVIDENCE / "component_selection_status.csv",
        EVIDENCE / "nmr_plssvd_cycle16_summary.csv",
        ALL_DATASET_FIGURE,
        NMR_SELECTION_FIGURE,
    ):
        copy2(source, OUT / source.name)
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
