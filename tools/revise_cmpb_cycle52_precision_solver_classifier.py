#!/usr/bin/env python3

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle51"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle52"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle51_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle51_0.99.6_20260726.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle52_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle52_0.99.6_20260726.docx"

TRADEOFF_FIGURE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle52_20260726"
    / "main_precision_solver_classifier.png"
)


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def set_cell_margins(cell, top=70, start=85, bottom=70, end=85):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def format_backend_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.92, 1.80, 1.78, 2.05]
    for row_index, row in enumerate(table.rows):
        for col_index, cell in enumerate(row.cells):
            cell.width = Inches(widths[col_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_index == 0:
                shade_cell(cell, "D9EAF7")
            elif row_index % 2 == 0:
                shade_cell(cell, "F4F7F9")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if col_index == 0
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7.5)
                    if row_index == 0:
                        run.bold = True


def add_backend_table(document, before_paragraph):
    caption = before_paragraph.insert_paragraph_before(
        style="Caption"
    )
    caption.add_run(
        "Table 2. Numerical and execution choices in fastPLS. Performance "
        "values are measured contrasts; execution rows describe supported "
        "routes. CPU threading is delegated to the linked BLAS implementation. "
        "The primary workstation used reference BLAS, so this study does not "
        "claim a measured multicore speed-up."
    )

    rows = [
        (
            "Choice",
            "Measured evidence",
            "Execution scope",
            "Interpretation",
        ),
        (
            "float32",
            "Input storage fell from 2.60 to 1.36 MB on MetRef and from "
            "24.25 to 12.70 MB on PRISM (approximately 48%).",
            "Automatically selected by float-package inputs on supported CPU, "
            "CUDA, and Metal routes.",
            "A reduced-storage option; runtime, workspace memory, and numerical "
            "agreement remain method- and shape-dependent.",
        ),
        (
            "IRLBA",
            "Deterministic SIMPLS passed all 117 estimator-preservation "
            "comparisons and all fixed-fold selections.",
            "Compiled CPU reference; accelerator requests may retain a host "
            "decomposition stage.",
            "Preferred for confirmatory analysis and for ill-conditioned or "
            "numerically sensitive problems.",
        ),
        (
            "rSVD",
            "CPU SIMPLS speed-up was 1.00-1.45-fold on nine classification "
            "tasks and 22.26-fold on NMR at 100 components.",
            "Available on CPU, NVIDIA CUDA, and Apple Metal; sketch controls "
            "and seed are recorded.",
            "An approximate solver. Agreement must be audited across seeds or "
            "against IRLBA before confirmatory use.",
        ),
        (
            "Compiled CPU",
            "CPU was faster in 31 of 44 selected-point CPU/CUDA contrasts "
            "outside the dedicated NMR and ImageNet analyses.",
            "C++ with BLAS/LAPACK; an optional OpenBLAS build can use multiple "
            "CPU cores.",
            "Usually preferable for small or modest matrices. Multicore gain "
            "depends on the linked BLAS and was not isolated here.",
        ),
        (
            "CUDA",
            "CUDA was faster in 13 of 44 selected-point contrasts and in every "
            "PLS family on CIFAR-100; it also accelerated NMR PLS-SVD/SIMPLS.",
            "NVIDIA GPU through CUDA, cuBLAS, cuSOLVER, and cuRAND; core "
            "PLS-SVD/SIMPLS products and selected prediction/LDA stages are "
            "device-native.",
            "Best suited to sufficiently large dense products. Transfer, "
            "context, and host-assisted OPLS/nonlinear-kernel stages matter.",
        ),
        (
            "Metal",
            "Metal accelerated selected large float64 CIFAR-100 paths by up to "
            "4.35-fold, but was not uniformly faster.",
            "Apple GPU through Metal and Metal Performance Shaders with unified "
            "memory; some decomposition, OPLS, kernel, and validation stages "
            "remain host-assisted.",
            "A portable accelerator route whose benefit depends on shape, "
            "precision, and effective stage residency.",
        ),
        (
            "LDA head",
            "Accuracy improved in seven of eight deterministic SIMPLS tasks: "
            "median gain 2.64 percentage points; median total-time ratio 1.13.",
            "Latent-score discriminant head with compiled CPU and CUDA paths; "
            "the effective classifier backend is reported.",
            "Useful when latent class structure is not decoded well by argmax; "
            "selection must remain training-only.",
        ),
    ]

    table = document.add_table(rows=len(rows), cols=4)
    table.style = "Table"
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            table.cell(row_index, col_index).text = value
    format_backend_table(table)
    before_paragraph._p.addprevious(table._tbl)


def revise_main():
    if not TRADEOFF_FIGURE.exists():
        raise FileNotFoundError(TRADEOFF_FIGURE)

    document = Document(MAIN_SOURCE)

    precision_method = find_paragraph(
        document,
        "Double precision is the numerical reference.",
    )
    precision_method.text = (
        "Standard R numeric matrices store each value in double precision "
        "(float64; eight bytes), whereas matrices supplied through the float "
        "package store four-byte float32 values. Double precision is the "
        "numerical reference. Float inputs automatically select float32 "
        "arithmetic on supported CPU, CUDA, and Metal routes, but reduced "
        "precision is treated as a conditional storage-saving mode rather than "
        "a universal acceleration or equivalence claim. Compiled CPU linear "
        "algebra uses the BLAS/LAPACK linked to R; builds can optionally link "
        "OpenBLAS so eligible matrix operations use multiple CPU cores. CUDA "
        "uses NVIDIA cuBLAS, cuSOLVER, and cuRAND, while Metal uses Metal "
        "Performance Shaders on Apple GPUs. Requested and effective stage "
        "backends are retained because OPLS filtering, nonlinear-kernel "
        "construction, some reduced decompositions, and validation routes can "
        "remain host-assisted. The Apple portability campaign used an Apple M3 "
        "with eight CPU cores, ten GPU cores, and 8 GB unified memory. "
        "Method- and shape-based warnings remained active for precision-"
        "sensitive classification, extreme multivariate responses, and "
        "nonlinear kernels (Supplementary Sections S19-S20)."
    )

    external_methods = find_paragraph(
        document,
        "External comparisons use independent R implementations",
    )
    insert_paragraph_after(
        external_methods,
        (
            "Three matched contrasts were used to separate precision, solver, "
            "and classifier effects. Float32 and float64 runs held dataset, "
            "split, PLS family, backend, component count, and prediction rule "
            "fixed. CPU rSVD and IRLBA were compared within fastPLS SIMPLS "
            "using float64 inputs and argmax; nine classification tasks used "
            "three isolated runs, whereas the fixed-100-component NMR contrast "
            "was a single feasibility run. The classifier contrast held the "
            "deterministic CPU IRLBA SIMPLS fit, split, and component count "
            "fixed and changed only argmax versus latent-space LDA. These "
            "contrasts report total fitting plus prediction time."
        ),
        style="Body Text",
    )

    abstract_results = find_paragraph(document, "Results:")
    abstract_results.text += (
        " Matched float32 inputs reduced stored matrix size by approximately "
        "48% on MetRef and PRISM, although speed and process memory were not "
        "uniformly improved. Within CPU SIMPLS, rSVD was only 1.00-1.45-fold "
        "faster than IRLBA across nine classification tasks but 22.26-fold "
        "faster on the extreme-response NMR task at 100 components. "
        "Latent-space LDA improved accuracy in seven of eight deterministic "
        "SIMPLS comparisons (median gain 2.64 percentage points) with a median "
        "1.13-fold total-time ratio relative to argmax."
    )

    precision_heading = find_paragraph(
        document,
        "3.5 Precision and backend agreement",
    )
    precision_heading.text = "3.5 Precision, solver, classifier, and backend trade-offs"

    precision_results = find_paragraph(
        document,
        "Float32 support was not uniformly beneficial.",
    )
    precision_results.text = (
        "Float32 provided a clear storage benefit but not a universal "
        "computational benefit (Figure 5A; Table 2). It reduced stored inputs "
        "from 2.60 to 1.36 MB on MetRef and from 24.25 to 12.70 MB on PRISM "
        "(47.9% and 47.7%). PLS-SVD preserved MetRef accuracy, but float32 "
        "SIMPLS and linear kernel PLS differed by two to five percentage "
        "points. Several PRISM routes were slower, and full-response NMR was "
        "numerically unsafe for float32 SIMPLS, OPLS, and linear kernel PLS. "
        "Float32 is therefore a conditional reduced-storage mode, not a "
        "default replacement for float64."
    )

    solver_results = find_paragraph(
        document,
        "The expanded Apple M3 campaign completed",
    )
    solver_results.text = (
        "Solver benefit was shape-dependent (Figure 5B). In matched float64 CPU "
        "SIMPLS/argmax runs, rSVD accelerated nine classification tasks by "
        "1.00-1.45-fold (median 1.04-fold). At 100 NMR components it reduced "
        "time from 436.253 to 19.597 s (22.26-fold), with RMSD unchanged to the "
        "shown precision. The NMR gain reflects its 28,355-column response; "
        "rSVD remains approximate. IRLBA is the confirmatory reference, while "
        "rSVD is a CPU/CUDA/Metal workflow whose sketch controls, seed, and "
        "audit status must be reported."
    )

    backend_results = find_paragraph(
        document,
        "The NMR result reinforced the shape dependence.",
    )
    backend_results.text = (
        "The prediction head affected classification more strongly "
        "(Figure 5C-D). With deterministic float64 CPU IRLBA SIMPLS, LDA "
        "improved seven of eight datasets by a median 2.64 percentage points, "
        "including 11.0 on MetRef, 7.5 on Tabula Muris, 4.2 on CCLE, and 3.0 on "
        "GTEx v8. TCGA-HNSC decreased by 1.7 points. The median total-time ratio "
        "was 1.13 relative to argmax (range 0.89-1.40). LDA can therefore "
        "improve multiclass prediction at modest cost, but must be selected "
        "using training data."
    )

    cross_validation_heading = find_paragraph(
        document,
        "3.6 Cross-validation performance",
    )

    table_break = cross_validation_heading.insert_paragraph_before()
    table_break.add_run().add_break(WD_BREAK.PAGE)
    add_backend_table(document, cross_validation_heading)

    figure_break = cross_validation_heading.insert_paragraph_before()
    figure_break.add_run().add_break(WD_BREAK.PAGE)
    picture = cross_validation_heading.insert_paragraph_before()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(TRADEOFF_FIGURE), width=Inches(6.0))
    caption = cross_validation_heading.insert_paragraph_before(style="Caption")
    caption.add_run(
        "Figure 5. Precision, solver, and classifier trade-offs. "
        "(A) Stored float64 and float32 inputs for MetRef and PRISM. "
        "(B) IRLBA/rSVD total-time ratio for matched float64 CPU SIMPLS/argmax. "
        "Classification values are three-run medians; NMR is one fixed-100-"
        "component run, with RMSD 0.0007075140 for both solvers. "
        "(C) LDA minus argmax accuracy for deterministic float64 CPU IRLBA "
        "SIMPLS. (D) Corresponding total-time ratio; values are three-run "
        "medians. rSVD is approximate, so panel B is not estimator-equivalence "
        "evidence."
    )

    discussion_precision = find_paragraph(
        document,
        "Float32 approximately halves raw input storage",
    )
    discussion_classifier = insert_paragraph_after(
        discussion_precision,
        (
            "The three choices answer different computational or statistical "
            "needs. Float32 primarily reduces representation size; rSVD "
            "reduces direction-extraction cost when matrix shape makes the "
            "deterministic solve expensive; and LDA changes the decision rule "
            "without changing the PLS estimator. Their benefits should "
            "therefore not be collapsed into a single fastest configuration. "
            "Likewise, compiled CPU, CUDA, and Metal are alternative execution "
            "routes rather than different statistical models. Optional "
            "multithreaded OpenBLAS can parallelize eligible CPU kernels, but "
            "the present benchmark used reference BLAS and provides no direct "
            "thread-scaling estimate."
        ),
        style="Body Text",
    )
    discussion_classifier.paragraph_format.keep_with_next = False

    conclusion = find_paragraph(
        document,
        "fastPLS combines a shape-aware accelerated sequential SIMPLS",
    )
    conclusion.text = conclusion.text.replace(
        "conditional float32 execution, and CPU, CUDA, and Metal backends.",
        "conditional float32 execution, optional latent-space LDA, and compiled "
        "CPU, NVIDIA CUDA, and Apple Metal backends.",
    )
    conclusion.text = conclusion.text.replace(
        "These capabilities make established PLS workflows",
        "Float32 reduces stored inputs, rSVD can substantially accelerate "
        "extreme-response problems, and LDA can improve multiclass prediction, "
        "but each advantage is conditional on the data and numerical route. "
        "These capabilities make established PLS workflows",
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - precision, solver, and classifier evidence"
    )
    document.save(MAIN_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    shutil.copy2(SUPP_SOURCE, SUPP_OUT)
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
