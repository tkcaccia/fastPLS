#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle33"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle34"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle33_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle33_0.99.6_20260726.docx"
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle33_20260726.docx"
)

MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle34_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle34_0.99.6_20260726.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle34_20260726.docx"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def insert_after(paragraph, text, style=None):
    new_xml = paragraph._parent.add_paragraph()._p
    paragraph._p.addnext(new_xml)
    new_paragraph = Paragraph(new_xml, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    new_paragraph.add_run(text)
    return new_paragraph


def set_cell_text(cell, text, bold=False, size=7.2, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_cell_margins(cell, top=45, start=55, bottom=45, end=55):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def revise_main():
    document = Document(MAIN_SOURCE)

    methods_anchor = find_paragraph(
        document,
        "The main biomedical benchmark included twelve tasks",
    )
    insert_after(
        methods_anchor,
        (
            "Prepared real-data matrices were not redistributed with fastPLS, "
            "including datasets whose authoritative source is publicly "
            "downloadable. The repository redistributes only the GPL-compatible "
            "breast and colon examples, synthetic generators and generated "
            "synthetic results, and aggregate benchmark tables, figures, "
            "manifests, split indices, and checksums that contain no source "
            "matrices or participant-level records. Public sources are reacquired "
            "from the cited provider by an executable dataset-acquisition script. "
            "ImageNet, NMR, the historical CCLE 18Q2 release, and the "
            "release-specific PRISM task require user-authorized local files; "
            "the script validates and checksums those files and never substitutes "
            "a different release. Dataset-specific access conditions and commands "
            "are given in Supplementary Section S5.6."
        ),
        style="Body Text",
    )

    availability = find_paragraph(document, "The fastPLS R package")
    availability.text = (
        "The fastPLS R package, benchmark workflows, analysis scripts, "
        "machine-readable result tables, synthetic generators, and aggregate "
        "benchmark outputs are available at https://github.com/tkcaccia/fastPLS "
        "(review-cycle package commit "
        "72e178b9e3c9510dc86c4b287d68b9c717f9fdf5). Low-level reusable C++ "
        "components are maintained at https://github.com/tkcaccia/kodama-cpp. "
        "The package bundles only the GPL-compatible breast and colon examples; "
        "prepared real-data benchmark matrices are not redistributed. "
        "benchmark/DATA_ACQUISITION.md and "
        "benchmark/acquire_publication_datasets.R provide authoritative source "
        "links, exact release identifiers, executable public downloads, and "
        "checksum validation for user-authorized restricted files. The "
        "supplement reports the redistribution status of every dataset and "
        "benchmark object."
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - dataset redistribution and acquisition"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)
    anchor = find_paragraph(document, "PRISM matches molecular profiles")

    heading = insert_after(anchor, "S5.6 Redistribution and executable acquisition")
    heading.style = "Heading 2"
    heading.paragraph_format.keep_with_next = True

    prose = insert_after(
        heading,
        (
            "Redistribution status was determined separately from source access "
            "status. A publicly downloadable source was not treated as permission "
            "to redistribute the processed benchmark matrix under the fastPLS "
            "licence. Only the GPL-compatible breast and colon package examples, "
            "synthetic generators and generated synthetic results, and aggregate "
            "tables, figures, manifests, split indices, and checksums are "
            "redistributed. None of the prepared real-data matrices in Table S3 "
            "is bundled. The executable workflow "
            "benchmark/acquire_publication_datasets.R downloads authoritative "
            "public sources or validates user-supplied local files, writes paths, "
            "sizes, access classes, and checksums to acquisition_manifest.csv, "
            "and refuses silent release substitution. Exact commands and "
            "preprocessing contracts are documented in "
            "benchmark/DATA_ACQUISITION.md."
        ),
        style="Body Text",
    )

    caption = insert_after(
        prose,
        (
            "Table S3a. Redistribution and acquisition status. 'No' means that "
            "fastPLS does not redistribute the prepared real-data matrix, even "
            "when the upstream source is public. The acquisition command is run "
            "from the repository root; restricted validation requires the "
            "environment variable shown."
        ),
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True

    headers = (
        "Dataset/object",
        "Authoritative source",
        "Access class",
        "Prepared object redistributed?",
        "Executable route",
    )
    rows = [
        (
            "breast; colon",
            "mixOmics; plsgenomics",
            "GPL-compatible package data",
            "Yes",
            "data(breast); data(colon)",
        ),
        (
            "Synthetic and aggregate outputs",
            "fastPLS generators/results",
            "Project-generated",
            "Yes",
            "Run published benchmark scripts",
        ),
        (
            "MetRef",
            "KODAMA::MetRef",
            "Public R package",
            "No",
            "--dataset=metref",
        ),
        (
            "CIFAR-100",
            "University of Toronto",
            "Public download",
            "No",
            "--dataset=cifar100",
        ),
        (
            "CBMC CITE-seq",
            "GEO GSE100866",
            "Public repository",
            "No",
            "--dataset=cbmc_citeseq",
        ),
        (
            "Retina",
            "GEO GSE63472; openTSNE pca_50",
            "Public repositories",
            "No",
            "--dataset=retina",
        ),
        (
            "Tabula Muris",
            "ExperimentHub EH1617",
            "Public Bioconductor data",
            "No",
            "--dataset=tabula",
        ),
        (
            "GTEx v8",
            "GTEx open access via UCSC Xena",
            "Open expression/phenotype",
            "No",
            "--dataset=gtex_v8",
        ),
        (
            "TCGA tasks",
            "TCGA open access via UCSC Xena",
            "Open molecular/phenotype",
            "No",
            "--dataset=tcga_brca,tcga_hnsc_methylation,tcga_pan_cancer",
        ),
        (
            "CCLE",
            "DepMap/CCLE 18Q2",
            "Historical release",
            "No",
            "FASTPLS_CCLE_RDATA=... --dataset=ccle",
        ),
        (
            "PRISM",
            "DepMap PRISM 19Q4",
            "Release-controlled",
            "No",
            "FASTPLS_PRISM_RDATA=... --dataset=prism",
        ),
        (
            "NMR",
            "Figshare share; MTBLS242/395/424",
            "Source-study terms",
            "No",
            "FASTPLS_NMR_RDATA=... --dataset=nmr",
        ),
        (
            "ImageNet/DINOv2",
            "ImageNet authorized copy; local extraction",
            "Gated/non-commercial terms",
            "No",
            "FASTPLS_IMAGENET_RDATA=... --dataset=imagenet",
        ),
    ]

    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(0.88), Inches(1.45), Inches(1.05), Inches(0.98), Inches(2.14)]

    for idx, (cell, header, width) in enumerate(
        zip(table.rows[0].cells, headers, widths)
    ):
        cell.width = width
        set_cell_text(cell, header, bold=True, size=7.0, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(cell, "D9EAF7")
        set_cell_margins(cell)
    repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])

    for row_values in rows:
        row = table.add_row()
        prevent_row_split(row)
        for idx, (cell, value, width) in enumerate(zip(row.cells, row_values, widths)):
            cell.width = width
            alignment = (
                WD_ALIGN_PARAGRAPH.CENTER
                if idx in (2, 3)
                else WD_ALIGN_PARAGRAPH.LEFT
            )
            set_cell_text(cell, value, size=6.8, align=alignment)
            set_cell_margins(cell)

    caption._p.addnext(table._tbl)

    note = insert_after(
        caption,
        (
            "The public acquisition command is Rscript "
            "benchmark/acquire_publication_datasets.R --dataset=<id> "
            "--out=<directory>. For ImageNet, NMR, CCLE, and PRISM, the listed "
            "environment variable points to a user-authorized local file. The "
            "workflow records a checksum but does not copy that file into the "
            "repository. GTEx uses only open-access expression and phenotype "
            "data; protected sequence and full donor-level files are outside the "
            "benchmark. The ImageNet embeddings remain a derived object subject "
            "to the source-image terms and are not redistributed."
        ),
        style="Body Text",
    )
    table._tbl.addnext(note._p)

    document.core_properties.title = (
        "fastPLS CMPB supplement - dataset redistribution and acquisition"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    heading = document.add_heading(
        "34. Dataset redistribution and acquisition were insufficiently explicit",
        level=1,
    )
    heading.paragraph_format.keep_with_next = True
    comment = document.add_paragraph(
        "Reviewer comment: Clarify which datasets and benchmark objects can be "
        "redistributed and provide executable acquisition instructions for "
        "restricted datasets."
    )
    comment.paragraph_format.keep_with_next = True
    response = document.add_paragraph(
        "Response: Corrected. We now separate source accessibility from "
        "redistribution permission. fastPLS bundles only the GPL-compatible "
        "breast and colon examples and redistributes project-generated synthetic "
        "data and aggregate benchmark products that contain no source matrices "
        "or participant-level observations. No prepared real-data benchmark "
        "matrix is redistributed, including matrices derived from publicly "
        "downloadable sources. Supplementary Section S5.6 and Table S3a report "
        "the authoritative source, access class, redistribution status, and "
        "executable route for every dataset. We added "
        "benchmark/DATA_ACQUISITION.md and an executable "
        "benchmark/acquire_publication_datasets.R workflow. Public sources are "
        "downloaded from their provider and recorded in a checksum manifest. "
        "ImageNet, NMR, the historical CCLE 18Q2 release, and release-specific "
        "PRISM data require explicit user-authorized local paths; the script "
        "validates and checksums them without copying them into the repository. "
        "It also refuses silent substitution of newer releases. The main "
        "Methods and Data and code availability sections now state this policy."
    )
    response.paragraph_format.keep_together = True
    document.core_properties.title = (
        "fastPLS CMPB response to reviewers - dataset access and redistribution"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    revise_response()


if __name__ == "__main__":
    main()
