#!/usr/bin/env python3
"""Separate NMR model selection, backend benchmarking, and historical context."""

import csv
import hashlib
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle63"
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle64"
EVIDENCE_DIR = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle64_20260726"
)

MAIN_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_main_cycle63_0.99.6_20260726.docx"
)
SUPP_SOURCE = (
    SOURCE_DIR
    / "fastPLS_CMPB_supplement_cycle63_0.99.6_20260726.docx"
)
MAIN_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_main_cycle64_0.99.6_20260726.docx"
)
SUPP_OUTPUT = (
    OUTPUT_DIR
    / "fastPLS_CMPB_supplement_cycle64_0.99.6_20260726.docx"
)

OLD_LEDGER = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle63_20260726"
    / "analysis_commit_provenance.csv"
)
NEW_LEDGER = EVIDENCE_DIR / "analysis_commit_provenance.csv"
SUMMARY_SCRIPT = ROOT / "benchmark" / "summarize_nmr_separated_analyses.R"
MAIN_FIGURE = (
    EVIDENCE_DIR
    / "nmr_separated_predictive_and_backend_benchmark.png"
)
HISTORICAL_FIGURE = EVIDENCE_DIR / "nmr_historical_reference_165.png"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_paragraph(document, prefix, text):
    paragraph = find_paragraph(document, prefix)
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def insert_after(paragraph, text, style="Normal"):
    new_paragraph = paragraph._parent.add_paragraph(style=style)
    new_paragraph.add_run(text)
    paragraph._p.addnext(new_paragraph._p)
    return new_paragraph


def insert_picture_before(paragraph, image, width):
    picture = paragraph._parent.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(image), width=width)
    paragraph._p.addprevious(picture._p)
    return picture


def remove_table(table):
    table._element.getparent().remove(table._element)


def tree_digest(directory):
    digest = hashlib.sha256()
    files = sorted(
        path
        for path in directory.rglob("*")
        if path.is_file() and path.name != "analysis_commit_provenance.csv"
    )
    for path in files:
        relative = path.relative_to(directory).as_posix()
        digest.update(relative.encode("utf-8"))
        digest.update(b"\0")
        digest.update(path.read_bytes())
        digest.update(b"\0")
    return digest.hexdigest()


def write_updated_ledger():
    with OLD_LEDGER.open(newline="", encoding="utf-8") as handle:
        rows = list(csv.DictReader(handle))
    rows.append(
        {
            "analysis_id": "A16",
            "reported_output": (
                "Main Figure 4; Section S43; Tables S55-S59; Figure S42"
            ),
            "result_archive": (
                "benchmark_results/manuscript_revision_cycle64_20260726"
            ),
            "package_version": "0.99.6",
            "package_commit": "not recorded in archived metadata",
            "package_commit_status": "not recoverable; not inferred",
            "benchmark_script": (
                "benchmark/summarize_nmr_separated_analyses.R"
            ),
            "benchmark_script_commit": (
                "not recorded in archived metadata"
            ),
            "benchmark_script_md5_current_copy": hashlib.md5(
                SUMMARY_SCRIPT.read_bytes()
            ).hexdigest(),
            "kodama_cpp_commit": "not recorded in archived metadata",
            "data_split_provenance": (
                "Deposited 1200/321 NMR split; training-only component "
                "selection over five paired splits; selected-point uncertainty "
                "from 10000 held-out-sample bootstrap resamples; paired CPU/CUDA "
                "benchmark fixed within family, rSVD, float64, component count, "
                "preprocessing, split, and target; historical reference retained "
                "separately at its published 165 components"
            ),
            "result_archive_sha256": tree_digest(EVIDENCE_DIR),
        }
    )
    with NEW_LEDGER.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)
    return rows[-1]


def read_csv(name):
    with (EVIDENCE_DIR / name).open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def numeric_or_zero(value):
    if value in (None, "", "NA", "NaN", "nan"):
        return 0.0
    return float(value)


def add_small_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table"
    for index, value in enumerate(headers):
        table.rows[0].cells[index].text = str(value)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cell.width = widths[column_index]
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if column_index == 0
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(6)
                    if row_index == 0:
                        run.bold = True
    return table


def add_caption(document, text):
    paragraph = document.add_paragraph(style="Caption")
    paragraph.add_run(text)
    return paragraph


def revise_main():
    document = Document(MAIN_SOURCE)

    replace_paragraph(
        document,
        "The NMR task contained",
        (
            "The NMR task contained 1,200 training spectra, 321 held-out "
            "spectra, 13,000 NOESY predictor bins, and 28,355 diffusion-edited "
            "response intensities. The 4.6-4.8 ppm water interval was set to "
            "zero in both training and test predictors as routine spectral "
            "preprocessing; the response spectra were not masked. We separated "
            "predictive model selection from implementation benchmarking "
            "(Figure 4). In the predictive analysis, five paired training-only "
            "splits and the one-standard-error rule selected five components "
            "for PLS-SVD and 50 for SIMPLS. These are family-specific predictive "
            "settings, not a common-complexity comparison."
        ),
    )
    replace_paragraph(
        document,
        "At these family-specific selected settings",
        (
            "At the selected settings, CUDA PLS-SVD/rSVD achieved held-out "
            "RMSD 0.001043 (95% bootstrap interval 0.001000-0.001085) and "
            "Q2 0.9892, whereas CUDA SIMPLS/rSVD achieved RMSD 0.000759 "
            "(0.000665-0.000884) and Q2 0.9943. The median per-spectrum RMSD "
            "was 0.000926 for PLS-SVD and 0.000504 for SIMPLS; median "
            "response-wise RMSD was 0.000113 and 0.000083, respectively. "
            "Errors increased with response intensity in both families, but "
            "SIMPLS remained lower in every stratum: aggregate RMSD in the "
            "top 1% of training-defined response intensities was 0.00310 for "
            "SIMPLS and 0.00527 for PLS-SVD. The bootstrap intervals quantify "
            "held-out-sample uncertainty conditional on the fixed outer split."
        ),
    )
    implementation = replace_paragraph(
        document,
        "A second, fixed 100-component benchmark",
        (
            "The implementation analysis then compared CPU with CUDA separately "
            "within each family at its selected component count, holding the "
            "outer split, preprocessing, response target, float64 precision, "
            "rSVD controls, and component count fixed. For PLS-SVD at five "
            "components, total fitting-plus-prediction time decreased from "
            "2.301 s on CPU to 0.648 s on CUDA (3.55-fold), with prediction "
            "correlation 1.000000 and between-backend prediction RMSD "
            "3.79e-8. For SIMPLS at 50 components, time decreased from "
            "10.525 to 1.773 s (5.94-fold), with prediction correlation "
            "0.999981 and between-backend RMSD 6.14e-5. Incremental host RSS "
            "was 390/711 MB for CPU/CUDA PLS-SVD and 407/698 MB for CPU/CUDA "
            "SIMPLS; incremental process GPU memory was 590 and 3,414 MB for "
            "the CUDA routes. These paired contrasts estimate backend effects "
            "without changing model family, solver, precision, or complexity."
        ),
    )
    historical = insert_after(
        implementation,
        (
            "The deposited Nature Communications workflow is reported "
            "separately as historical context. Its scientific model used 165 "
            "components with the original centring-only protocol, so the "
            "previous forced 100-component rerun was only an equal-size "
            "sensitivity analysis and is no longer used as the central NMR "
            "comparison. At the published 165 components, the unchanged "
            "deposited fastsimpls PLS-SVD/IRLBA workflow required 447.6 s, "
            "used 3,605 MB incremental host RSS, and achieved RMSD 0.000710. "
            "Because comparisons with fastPLS simultaneously change solver, "
            "implementation, PLS family, or hardware, they remain contextual "
            "workflow comparisons in the Supplementary Material rather than "
            "backend-only effects."
        ),
        style="Body Text",
    )
    historical.paragraph_format.keep_together = True

    nmr_table = None
    for table in document.tables:
        if len(table.rows) == 2 and len(table.columns) == 2:
            if table.cell(0, 0).text.strip().startswith("A"):
                nmr_table = table
                break
    if nmr_table is None:
        raise RuntimeError("NMR figure table not found")
    remove_table(nmr_table)

    caption = find_paragraph(document, "Figure 4.")
    blank_before = caption._p.getprevious()
    if blank_before is not None and blank_before.tag.endswith("}p"):
        text = "".join(blank_before.itertext()).strip()
        if not text:
            blank_before.getparent().remove(blank_before)
    insert_picture_before(caption, MAIN_FIGURE, Inches(6.55))
    figure_caption = replace_paragraph(
        document,
        "Figure 4.",
        (
            "Figure 4. Separated NMR predictive and implementation analyses. "
            "(A) Training-only component paths and one-standard-error "
            "selections. (B) Held-out per-spectrum RMSD distributions at five "
            "PLS-SVD and 50 SIMPLS components. (C) Response-wise RMSD over "
            "28,355 spectral coordinates. (D) Aggregate RMSD across response "
            "coordinates stratified by mean absolute intensity in Ytrain. "
            "(E) Paired CPU/CUDA implementation benchmark within each family; "
            "split, preprocessing, target, rSVD, float64 precision, and "
            "component count are fixed. Host and GPU memory are increments "
            "above the pre-fit process baseline; GPU values include runtime "
            "context and are not workspace-only allocations."
        ),
    )
    figure_caption.paragraph_format.keep_with_next = False
    figure_caption.paragraph_format.keep_together = False
    figure_caption.paragraph_format.space_before = Pt(3)
    figure_caption.paragraph_format.space_after = Pt(2)
    figure_caption.paragraph_format.line_spacing = 1
    for run in figure_caption.runs:
        run.font.size = Pt(7.5)

    component_discussion = find_paragraph(
        document,
        "Component selection is an additional limitation.",
    )
    insert_after(
        component_discussion,
        (
            "The NMR application illustrates why predictive selection and "
            "computational attribution must be separated. The selected "
            "PLS-SVD model was faster because it used five components, whereas "
            "the selected 50-component SIMPLS model was more accurate across "
            "global, per-spectrum, response-wise, and intensity-stratified "
            "errors. Within-family paired CPU/CUDA comparisons then isolated "
            "backend effects and showed numerical agreement together with "
            "3.55- and 5.94-fold speed-ups. The 165-component deposited "
            "workflow answers a third, historical question and is therefore "
            "not used to attribute a single implementation effect."
        ),
        style="Body Text",
    )

    replace_paragraph(
        document,
        "The fastPLS R package, benchmark workflows",
        (
            "The fastPLS R package, benchmark workflows, analysis scripts, "
            "machine-readable result tables, synthetic generators, and "
            "aggregate benchmark outputs are available at "
            "https://github.com/tkcaccia/fastPLS. Low-level reusable C++ "
            "components are maintained at "
            "https://github.com/tkcaccia/kodama-cpp. Supplementary Table S41 "
            "and benchmark_results/manuscript_revision_cycle64_20260726/"
            "analysis_commit_provenance.csv map each reported analysis to its "
            "result archive, generating script, recorded package version, "
            "source identifier when captured, data/split provenance, and "
            "SHA-256 archive digest. Historical archives without a recorded "
            "Git SHA are marked as not recoverable and are not assigned a "
            "later commit retrospectively. Prepared real-data matrices are not "
            "redistributed; acquisition instructions and checksum validation "
            "are provided in benchmark/DATA_ACQUISITION.md and "
            "benchmark/acquire_publication_datasets.R."
        ),
    )

    document.save(MAIN_OUTPUT)


def append_provenance_row(document, ledger_row):
    target = None
    for table in document.tables:
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header[:2] == ["ID", "Reported output"]:
            target = table
            break
    if target is None:
        raise RuntimeError("Provenance table not found")
    row = target.add_row()
    values = [
        ledger_row["analysis_id"],
        ledger_row["reported_output"],
        Path(ledger_row["result_archive"]).name,
        Path(ledger_row["benchmark_script"]).name,
        ledger_row["package_version"],
        "NR",
        ledger_row["result_archive_sha256"][:12],
    ]
    for index, (cell, value) in enumerate(zip(row.cells, values)):
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if index in (1, 2, 3)
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(5)


def revise_supplement(ledger_row):
    document = Document(SUPP_SOURCE)

    replace_paragraph(
        document,
        "Matched fixed-complexity NMR contrasts.",
        (
            "Earlier fixed-complexity NMR sensitivity analysis. A distinct "
            "analysis imposed 100 components while holding float64 precision, "
            "rSVD controls, preprocessing, split, and prediction target fixed "
            "for the fastPLS rows. This is retained only as a sensitivity "
            "analysis. It is not the family-specific predictive analysis and "
            "does not reproduce the original deposited scientific model, which "
            "used 165 components. The deposited 100-component row additionally "
            "changes implementation and solver and therefore remains a "
            "descriptive composite workflow rather than an implementation-only "
            "contrast. Section S43 provides the corrected selected-component "
            "backend pairs and the historical 165-component context."
        ),
    )
    replace_paragraph(
        document,
        "Table S6b.",
        (
            "Table S6b. Historical fixed-100-component sensitivity table. "
            "The component count was imposed for equal-size exploration and "
            "was not the original deposited scientific setting. Time and "
            "memory are medians from three isolated runs; CUDA memory is the "
            "sampled process-level peak including context and runtime state. "
            "Because the deposited row differs in implementation and solver, "
            "the table must not be interpreted as an implementation-only, "
            "solver-only, family-only, or hardware-only ranking."
        ),
    )
    replace_paragraph(
        document,
        "The outer split was intentionally fixed",
        (
            "The outer split was intentionally fixed across the NMR analyses. "
            "The earlier 100-component analysis imposed a common model size for "
            "sensitivity and retained 321 per-spectrum and 28,355 response-wise "
            "errors, but it did not reproduce the original 165-component "
            "deposited model and is not used for family-specific predictive "
            "selection. Section S43 supersedes it for inference: family-specific "
            "components are selected from training data, error distributions "
            "are evaluated at those settings, and backend effects are paired "
            "within family at fixed solver, precision, component count, split, "
            "preprocessing, and target."
        ),
    )

    document.add_heading(
        "S43. Separated NMR predictive, implementation, and historical analyses",
        level=1,
    )
    document.add_paragraph(
        (
            "Three questions were analysed independently. First, predictive "
            "model selection used five paired training-only splits and the "
            "one-standard-error rule separately for PLS-SVD and SIMPLS. Second, "
            "the implementation benchmark changed only CPU versus CUDA within "
            "each family at its selected component count; split, preprocessing, "
            "target, rSVD controls, float64 precision, and component count were "
            "fixed. Third, the deposited Nature Communications function was "
            "rerun at the original 165 components with its original "
            "centring-only protocol. The historical comparison is contextual "
            "because solver, implementation, family, or hardware differ across "
            "rows."
        ),
        style="First Paragraph",
    )

    family = read_csv("nmr_family_selected_error_summary.csv")
    bootstrap = read_csv("nmr_family_selected_bootstrap_uncertainty.csv")
    boot_lookup = {
        (row["family"], row["metric"]): row for row in bootstrap
    }
    rows = []
    for row in family:
        rmsd = boot_lookup[(row["family"], "RMSD")]
        q2 = boot_lookup[(row["family"], "Q2")]
        rows.append(
            [
                row["family"],
                row["ncomp"],
                (
                    f'{float(rmsd["estimate"]):.7f} '
                    f'[{float(rmsd["lower"]):.7f}, '
                    f'{float(rmsd["upper"]):.7f}]'
                ),
                (
                    f'{float(q2["estimate"]):.5f} '
                    f'[{float(q2["lower"]):.5f}, '
                    f'{float(q2["upper"]):.5f}]'
                ),
                (
                    f'{float(row["sample_RMSD_median"]):.7f} '
                    f'[{float(row["sample_RMSD_q25"]):.7f}, '
                    f'{float(row["sample_RMSD_q75"]):.7f}]'
                ),
                (
                    f'{float(row["response_RMSD_median"]):.7f} '
                    f'[{float(row["response_RMSD_q25"]):.7f}, '
                    f'{float(row["response_RMSD_q75"]):.7f}]'
                ),
            ]
        )
    add_small_table(
        document,
        [
            "Family",
            "A",
            "Global RMSD [95% CI]",
            "Q2 [95% CI]",
            "Spectrum RMSD median [IQR]",
            "Response RMSD median [IQR]",
        ],
        rows,
    )
    add_caption(
        document,
        (
            "Table S55. Family-specific NMR predictive results. Components "
            "were selected using training data only. Intervals are 10,000-"
            "replicate held-out-sample percentile bootstrap intervals "
            "conditional on the fixed 321-spectrum outer test set."
        ),
    )

    intensity = read_csv(
        "nmr_family_selected_intensity_stratified_error.csv"
    )
    add_small_table(
        document,
        [
            "Family",
            "Training-intensity stratum",
            "Responses",
            "Aggregate RMSD",
            "Aggregate MAE",
            "Response RMSD median [IQR]",
        ],
        [
            [
                row["family"],
                row["intensity_stratum"],
                row["n_responses"],
                f'{float(row["aggregate_RMSD"]):.7f}',
                f'{float(row["aggregate_MAE"]):.7f}',
                (
                    f'{float(row["median_response_RMSD"]):.7f} '
                    f'[{float(row["response_RMSD_q25"]):.7f}, '
                    f'{float(row["response_RMSD_q75"]):.7f}]'
                ),
            ]
            for row in intensity
        ],
    )
    add_caption(
        document,
        (
            "Table S56. NMR intensity-stratified held-out errors at the "
            "family-selected component counts. Strata were defined before "
            "test evaluation from mean absolute Ytrain intensity: 0-50%, "
            "50-90%, 90-99%, and top 1%."
        ),
    )

    backend = read_csv("nmr_paired_backend_only_summary.csv")
    add_small_table(
        document,
        [
            "Family",
            "A",
            "Backend",
            "RMSD",
            "Time (s)",
            "Baseline host RSS",
            "Incremental host RSS",
            "Incremental GPU",
            "Prediction agreement",
        ],
        [
            [
                row["family"],
                row["effective_ncomp"],
                row["engine"],
                f'{float(row["metric_median"]):.7f}',
                f'{float(row["total_time_sec_median"]):.3f}',
                f'{float(row["baseline_host_rss_mb_median"]):.0f} MB',
                f'{float(row["incremental_host_rss_mb_median"]):.0f} MB',
                f'{float(row["incremental_gpu_mem_mb_median"]):.0f} MB',
                (
                    f'r={float(row["prediction_correlation"]):.6f}; '
                    f'RMSD={float(row["prediction_RMSD"]):.2e}'
                ),
            ]
            for row in backend
        ],
    )
    add_caption(
        document,
        (
            "Table S57. Paired NMR backend-only implementation benchmark. "
            "CPU and CUDA are compared within PLS family while rSVD, float64 "
            "precision, selected component count, split, preprocessing, target, "
            "and prediction workflow are held fixed. Resource values are "
            "medians from isolated runs. GPU memory includes context/runtime "
            "state and is not workspace-only memory."
        ),
    )

    historical = read_csv("nmr_historical_reference_165_summary.csv")
    keep = {
        "nature_fastsimpls_plssvd": "Deposited PLS-SVD/IRLBA",
        "cpp_plssvd_cpu_rsvd": "fastPLS CPU PLS-SVD/rSVD",
        "gpu_plssvd_rsvd": "fastPLS CUDA PLS-SVD/rSVD",
        "cpp_simpls_cpu_rsvd": "fastPLS CPU SIMPLS/rSVD",
        "gpu_simpls_rsvd": "fastPLS CUDA SIMPLS/rSVD",
    }
    historical_rows = []
    for row in historical:
        if row.get("variant_name") not in keep:
            continue
        historical_rows.append(
            [
                keep[row["variant_name"]],
                row.get("effective_ncomp", "165"),
                f'{float(row["global_rmsd"]):.7f}',
                f'{float(row["total_time_sec_median"]):.3f}',
                (
                    f'{float(row["incremental_peak_host_rss_mb_median"]):.0f} '
                    "MB"
                ),
                (
                    f'{numeric_or_zero(row.get("peak_gpu_mem_mb_median")):.0f} '
                    "MB"
                ),
            ]
        )
    add_small_table(
        document,
        [
            "Workflow",
            "A",
            "RMSD",
            "Total time (s)",
            "Incremental host RSS",
            "Peak GPU memory",
        ],
        historical_rows,
    )
    add_caption(
        document,
        (
            "Table S58. Historical NMR workflow context at the 165 components "
            "used in the original scientific analysis. The deposited function "
            "retains its original centring-only protocol. This table is not a "
            "causal backend benchmark because solver, implementation, family, "
            "or hardware differ across rows."
        ),
    )

    document.add_paragraph()
    document.add_picture(str(HISTORICAL_FIGURE), width=Inches(6.35))
    add_caption(
        document,
        (
            "Figure S42. Historical NMR workflows at the published "
            "165-component setting. The display provides scientific context "
            "for the deposited function and is intentionally separated from "
            "the paired backend-only comparison in Table S57."
        ),
    )
    file_table = add_small_table(
        document,
        ["Content", "Machine-readable file"],
        [
            [
                "Selected-component paths",
                "nmr_family_selected_component_paths.csv",
            ],
            [
                "Family-level error summary",
                "nmr_family_selected_error_summary.csv",
            ],
            [
                "Per-spectrum errors",
                "nmr_family_selected_per_spectrum_rmsd.csv",
            ],
            [
                "Response-wise errors",
                "nmr_family_selected_response_wise_error.csv",
            ],
            [
                "Intensity-stratified errors",
                "nmr_family_selected_intensity_stratified_error.csv",
            ],
            [
                "Bootstrap uncertainty",
                "nmr_family_selected_bootstrap_uncertainty.csv",
            ],
            [
                "Paired backend benchmark",
                "nmr_paired_backend_only_summary.csv",
            ],
            [
                "Historical 165-component context",
                "nmr_historical_reference_165_summary.csv",
            ],
        ],
    )
    page_break = document.add_page_break()
    file_table._tbl.addprevious(page_break._p)
    add_caption(
        document,
        (
            "Table S59. Machine-readable NMR error distributions and "
            "reproducibility files. The listed files, source prediction "
            "objects, generating note, and session information are archived "
            "under benchmark_results/"
            "manuscript_revision_cycle64_20260726/. The generating script is "
            "benchmark/summarize_nmr_separated_analyses.R."
        ),
    )

    replace_paragraph(
        document,
        "Table S41 is the authoritative mapping",
        (
            "Table S41 is the authoritative mapping between manuscript outputs "
            "and computational archives. A package commit identifies code "
            "loaded for computation; a benchmark-script identifier identifies "
            "orchestration code; and the archive SHA-256 identifies the result "
            "tree. These identifiers are not interchangeable. Historical "
            "analyses that recorded fastPLS 0.99.6 without an immutable source "
            "SHA remain marked as unavailable rather than being assigned a "
            "later commit retrospectively."
        ),
    )
    replace_paragraph(
        document,
        "Machine-readable ledger:",
        (
            "Machine-readable ledger: benchmark_results/"
            "manuscript_revision_cycle64_20260726/"
            "analysis_commit_provenance.csv. The A16 digest excludes the ledger "
            "file itself and covers the generated NMR result archive."
        ),
    )
    append_provenance_row(document, ledger_row)
    document.save(SUPP_OUTPUT)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    required = [
        MAIN_SOURCE,
        SUPP_SOURCE,
        OLD_LEDGER,
        SUMMARY_SCRIPT,
        MAIN_FIGURE,
        HISTORICAL_FIGURE,
    ]
    missing = [path for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("\n".join(str(path) for path in missing))
    ledger_row = write_updated_ledger()
    revise_main()
    revise_supplement(ledger_row)
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)
    print(NEW_LEDGER)


if __name__ == "__main__":
    main()
