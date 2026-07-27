#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle39"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle40"
EVIDENCE = ROOT / "benchmark_results" / "nmr_matched_contrasts_20260726"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle39_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle39_0.99.6_20260726.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle40_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle40_0.99.6_20260726.docx"

PANEL_C = EVIDENCE / "nmr_matched_prediction_contrasts.png"
PANEL_D = EVIDENCE / "nmr_matched_resource_contrasts.png"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_paragraph(document, prefix, text):
    paragraph = find_paragraph(document, prefix)
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def set_panel(cell, label, image):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = paragraph.add_run(f"{label}\n")
    label_run.bold = True
    label_run.font.size = Pt(10)
    paragraph.add_run().add_picture(str(image), width=Inches(3.55))


def replace_picture_before_caption(document, caption_start, image, width):
    paragraphs = document.paragraphs
    caption = find_paragraph(document, caption_start)
    caption_index = next(
        index
        for index, paragraph in enumerate(paragraphs)
        if paragraph._p is caption._p
    )
    for index in range(caption_index - 1, -1, -1):
        paragraph = paragraphs[index]
        if paragraph._p.xpath(".//w:drawing"):
            for child in list(paragraph._p):
                paragraph._p.remove(child)
            paragraph.add_run().add_picture(str(image), width=width)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
    raise RuntimeError(f"Image preceding caption not found: {caption_start}")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def prevent_row_splitting(table):
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def style_table(table, size=5.2):
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    repeat_header(table.rows[0])
    prevent_row_splitting(table)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(size)
                    if row_index == 0:
                        run.bold = True


def add_table_after(paragraph, headers, rows):
    document = paragraph._parent
    table = document.add_table(
        rows=1,
        cols=len(headers),
        width=Inches(7.0),
    )
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = str(value)
    paragraph._p.addnext(table._tbl)
    style_table(table)
    return table


def revise_main():
    document = Document(MAIN_SOURCE)

    replace_paragraph(
        document,
        "3.2.2 Fixed-complexity implementation benchmark",
        "3.2.2 Matched fixed-complexity contrasts",
    )
    replace_paragraph(
        document,
        "A separate benchmark fixed every implementation at 100 components",
        (
            "A separate analysis fixed fastPLS PLS-SVD and SIMPLS at 100 components, "
            "float64 precision, and the same one-power rSVD setting. It supports two "
            "matched questions. First, changing only hardware accelerated PLS-SVD "
            "from 16.323 s on CPU to 1.115 s on CUDA (14.64-fold) and SIMPLS from "
            "20.140 to 3.055 s (6.59-fold). RMSD changed from 0.0007292 to 0.0007183 "
            "for PLS-SVD and from 0.0008606 to 0.0008047 for SIMPLS. CUDA increased "
            "host RSS by 13% and 10%, respectively, and used 664 MB and 3,432 MB of "
            "sampled device memory. Second, changing only PLS family showed that "
            "PLS-SVD was 1.23-fold faster than SIMPLS on CPU and 2.74-fold faster on "
            "CUDA at this imposed component count. The deposited reference is shown "
            "only as a composite historical workflow: it used PLS-SVD with IRLBA on "
            "CPU and required 431.234 s, but comparison with fastPLS PLS-SVD/rSVD "
            "changes implementation and solver simultaneously. Its apparent "
            "26.42-fold difference from fastPLS CPU therefore cannot be attributed "
            "to either factor alone. Figure 3C-D and Supplementary Table S6a present "
            "the valid pairwise contrasts; Supplementary Table S6b retains the raw "
            "workflows for transparency. This operating point was imposed for "
            "computational comparison and is not the family-specific predictive "
            "selection."
        ),
    )

    figure_table = document.tables[1]
    set_panel(figure_table.cell(1, 0), "C", PANEL_C)
    set_panel(figure_table.cell(1, 1), "D", PANEL_D)

    replace_paragraph(
        document,
        "Figure 3. Fixed-complexity NMR implementation benchmark",
        (
            "Figure 3. Fixed-complexity NMR analysis at 100 components. This common "
            "component count was imposed and was not the family-specific predictive "
            "selection. (A) Observed and rSVD-based SIMPLS prediction of the held-out "
            "spectrum across the full chemical-shift range. (B) Enlarged 0.5-1.7 ppm "
            "region. (C) Per-spectrum RMSD shown as four matched contrasts: CPU versus "
            "CUDA within PLS-SVD, CPU versus CUDA within SIMPLS, PLS-SVD versus SIMPLS "
            "on CPU, and PLS-SVD versus SIMPLS on CUDA. Each contrast holds the solver, "
            "precision, component count, preprocessing, split, and all non-labelled "
            "factors fixed. (D) Corresponding runtime and host-RSS ratios. For backend "
            "contrasts, runtime is CPU/CUDA and RSS is CUDA/CPU; for family contrasts, "
            "both ratios are SIMPLS/PLS-SVD. The rSVD runs used oversampling 10, one "
            "power iteration, and seeds 124, 125, and 126; their audit status is "
            "workflow-only. The deposited PLS-SVD/IRLBA CPU reference is excluded from "
            "the matched panels because it changes implementation and solver together; "
            "its raw workflow result remains in Supplementary Table S6b. The "
            "representative spectrum was selected mechanically as the test spectrum "
            "whose RMSD was closest to the held-out median."
        ),
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - matched NMR contrasts"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)

    replace_paragraph(
        document,
        "Fixed-complexity NMR implementation benchmark.",
        (
            "Matched fixed-complexity NMR contrasts. A distinct analysis held "
            "component count (100), float64 precision, rSVD controls, preprocessing, "
            "split, and prediction target fixed. The hardware contrast compares CPU "
            "with CUDA separately within fastPLS PLS-SVD and SIMPLS. The family "
            "contrast compares PLS-SVD with SIMPLS separately on CPU and CUDA. These "
            "contrasts are summarized in Table S6a. The deposited fastsimpls reference "
            "used PLS-SVD with IRLBA on CPU; comparison with fastPLS PLS-SVD/rSVD "
            "therefore changes implementation and solver simultaneously and is "
            "retained only as a composite historical workflow in Table S6b. No "
            "implementation-only or solver-only effect is estimated from that row."
        ),
    )

    old_caption = find_paragraph(
        document,
        "Table S6. Fixed-complexity float64 NMR implementation comparison",
    )
    new_caption = old_caption.insert_paragraph_before(
        (
            "Table S6a. Prespecified matched contrasts in the fixed-100-component "
            "float64 NMR analysis. A differing factor is interpreted only where all "
            "listed fixed factors are shared. Runtime and memory use medians from "
            "three isolated runs. The rSVD setting was oversampling 10, one power "
            "iteration, and seeds 124/125/126; audit status is workflow-only."
        )
    )
    new_caption.style = "Caption"
    new_caption.paragraph_format.page_break_before = True
    rows = [
        (
            "Hardware within PLS-SVD",
            "Family=PLS-SVD; solver=rSVD; A=100; f64",
            "CPU vs CUDA",
            "16.323 vs 1.115 s; CUDA 14.64x faster",
            "0.0007292 vs 0.0007183",
            "Host 2,964 vs 3,338 MB; CUDA 664 MB",
        ),
        (
            "Hardware within SIMPLS",
            "Family=SIMPLS; solver=rSVD; A=100; f64",
            "CPU vs CUDA",
            "20.140 vs 3.055 s; CUDA 6.59x faster",
            "0.0008606 vs 0.0008047",
            "Host 3,143 vs 3,469 MB; CUDA 3,432 MB",
        ),
        (
            "Family on CPU",
            "Backend=CPU; solver=rSVD; A=100; f64",
            "PLS-SVD vs SIMPLS",
            "16.323 vs 20.140 s; PLS-SVD 1.23x faster",
            "0.0007292 vs 0.0008606",
            "Host 2,964 vs 3,143 MB",
        ),
        (
            "Family on CUDA",
            "Backend=CUDA; solver=rSVD; A=100; f64",
            "PLS-SVD vs SIMPLS",
            "1.115 vs 3.055 s; PLS-SVD 2.74x faster",
            "0.0007183 vs 0.0008047",
            "Host 3,338 vs 3,469 MB; GPU 664 vs 3,432 MB",
        ),
        (
            "Composite historical workflow",
            "A=100; f64; split and preprocessing fixed",
            "Deposited PLS-SVD/IRLBA CPU vs fastPLS PLS-SVD/rSVD CPU",
            "431.234 vs 16.323 s; 26.42x composite difference",
            "0.0007194 vs 0.0007292",
            "Host 6,101 vs 2,964 MB; not factor-attributable",
        ),
    ]
    add_table_after(
        new_caption,
        [
            "Contrast",
            "Fixed factors",
            "Differing factor/levels",
            "Time result",
            "RMSD",
            "Memory",
        ],
        rows,
    )
    replace_paragraph(
        document,
        "Table S6. Fixed-complexity float64 NMR implementation comparison",
        (
            "Table S6b. Raw fixed-complexity float64 NMR workflows at the imposed "
            "100-component operating point. Time and memory are medians from three "
            "isolated runs; prediction agreement is correlation with the deposited "
            "reference. The deposited row differs from fastPLS in both implementation "
            "and solver, so this table is descriptive and must not be read as an "
            "implementation-only, solver-only, family-only, or hardware-only ranking."
        ),
    )

    replace_picture_before_caption(
        document,
        "Figure S14. Distribution of held-out per-spectrum RMSD",
        PANEL_C,
        Inches(6.7),
    )
    replace_paragraph(
        document,
        "Figure S14. Distribution of held-out per-spectrum RMSD",
        (
            "Figure S14. Matched held-out per-spectrum RMSD contrasts at the common "
            "100-component point. Backend comparisons are made separately within "
            "PLS-SVD and SIMPLS; family comparisons are made separately on CPU and "
            "CUDA. The deposited reference is excluded because it changes solver and "
            "implementation simultaneously."
        ),
    )
    replace_picture_before_caption(
        document,
        "Figure S15. Total fit-plus-prediction time",
        PANEL_D,
        Inches(6.7),
    )
    replace_paragraph(
        document,
        "Figure S15. Total fit-plus-prediction time",
        (
            "Figure S15. Matched fixed-complexity NMR computational contrasts. "
            "Runtime and host-RSS ratios isolate backend within each family and family "
            "within each backend. Raw absolute time, host RSS, and sampled GPU memory "
            "for every workflow remain in Table S6b."
        ),
    )

    document.core_properties.title = (
        "fastPLS CMPB supplementary material - matched NMR contrasts"
    )
    document.save(SUPP_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    for path in (PANEL_C, PANEL_D):
        if not path.exists():
            raise RuntimeError(f"Missing NMR contrast figure: {path}")
    revise_main()
    revise_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
