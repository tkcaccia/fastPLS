#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
OUTPUT_DIR = (
    ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle58_review"
)
OUTPUT = (
    OUTPUT_DIR
    / "fastPLS_CMPB_fresh_reviewer_report_cycle58_20260726.docx"
)

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
RISK = RGBColor(155, 28, 28)
CAUTION = RGBColor(122, 90, 0)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def configure_styles(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = document.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = DARK_BLUE
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(10)

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(11)
    subtitle.font.color.rgb = MUTED
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(12)

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "Independent reviewer report | fastPLS"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED


def add_labelled_paragraph(document, label, text, color=None):
    paragraph = document.add_paragraph()
    label_run = paragraph.add_run(label)
    label_run.bold = True
    if color is not None:
        label_run.font.color.rgb = color
    paragraph.add_run(text)
    return paragraph


def add_major_comment(document, number, title, assessment, required):
    document.add_heading(f"{number}. {title}", level=2)
    document.add_paragraph(assessment)
    add_labelled_paragraph(
        document,
        "Required revision: ",
        required,
        color=RISK,
    )


def add_minor_comment(document, number, text):
    paragraph = document.add_paragraph()
    number_run = paragraph.add_run(f"{number}. ")
    number_run.bold = True
    paragraph.add_run(text)


def build_report():
    document = Document()
    configure_styles(document)

    title = document.add_paragraph(
        "Reviewer Report",
        style="Title",
    )
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    document.add_paragraph(
        (
            "Manuscript: “fastPLS: scalable partial least squares with "
            "compiled CPU and accelerator backends for high-dimensional "
            "biomedical data”"
        ),
        style="Subtitle",
    )
    add_labelled_paragraph(
        document,
        "Journal: ",
        "Computer Methods and Programs in Biomedicine",
    )
    add_labelled_paragraph(
        document,
        "Recommendation: ",
        "Major revision",
        color=RISK,
    )
    add_labelled_paragraph(
        document,
        "Review basis: ",
        (
            "Main manuscript and Supplementary Material supplied as cycle58, "
            "read as a new submission."
        ),
    )

    document.add_heading("Overall Assessment", level=1)
    document.add_paragraph(
        (
            "This manuscript describes an ambitious and potentially valuable "
            "software contribution for scaling established partial least "
            "squares methods to high-dimensional biomedical data. The work is "
            "strongest where it separates deterministic estimator validation "
            "from approximate workflow benchmarking. In particular, the "
            "component-level comparison with de Jong SIMPLS, the independent "
            "OPLS and kernel-PLS operator construction, explicit failure "
            "criteria, retained execution failures, and stage-level backend "
            "residency reporting are unusually careful for a computational "
            "software paper. The NMR application is scientifically relevant, "
            "and the ImageNet analysis is appropriately framed primarily as a "
            "stress test."
        )
    )
    document.add_paragraph(
        (
            "However, the evidential hierarchy is not yet fully consistent. "
            "The formally acceptable rSVD setting is not always the setting "
            "used for headline performance claims; the four-thread comparison "
            "does not establish multithreaded scaling; and Metal speed results "
            "include predictive discrepancies that prevent a clean "
            "backend-equivalence interpretation. The new setting-level OPLS "
            "and kernel-PLS validation is also not synchronized with the "
            "Abstract. These issues are addressable, but they affect the main "
            "claims and therefore require major revision."
        )
    )

    document.add_heading("Principal Strengths", level=1)
    strengths = [
        (
            "The deterministic SIMPLS validation is broad and reports "
            "predictions, coefficients, subspaces, selected components, and "
            "failures rather than accuracy alone."
        ),
        (
            "The independent OPLS and kernel-PLS references avoid reusing the "
            "package's operator-construction code. The expanded grid across "
            "orthogonal-component counts, RBF bandwidths, and polynomial "
            "forms is convincing for deterministic float64 CPU execution."
        ),
        (
            "The manuscript is transparent about noncanonical ImageNet "
            "splitting, conditional float32 support, hybrid accelerator "
            "stages, process-level memory, and component-grid boundaries."
        ),
        (
            "The NMR task provides a meaningful biomedical example involving "
            "an extreme multivariate response rather than only synthetic "
            "scaling matrices."
        ),
        (
            "The software comparison uses a clearly identified deterministic "
            "SIMPLS pair as the primary estimator-matched comparison."
        ),
    ]
    for index, strength in enumerate(strengths, 1):
        add_minor_comment(document, index, strength)

    document.add_heading("Major Comments", level=1)

    add_major_comment(
        document,
        1,
        "The rSVD validation setting and headline benchmark setting are not aligned",
        (
            "The Supplement reports that oversampling by 10 with one power "
            "iteration passed only 101 of 117 component-level checks, whereas "
            "two power iterations passed 117 of 117. Nevertheless, many "
            "principal workflow, family-speed, backend, and selected-point "
            "rows use the faster one-power setting. The Abstract first "
            "presents the two-power result as the accepted approximation and "
            "then reports timing advantages without consistently identifying "
            "which power setting generated each headline value. The four "
            "percentage-point MetRef difference further shows that this is "
            "not a cosmetic distinction."
        ),
        (
            "Either rerun the principal rSVD results with the prespecified "
            "two-power validated setting, or label every one-power result "
            "prominently as exploratory and remove it from confirmatory "
            "performance claims. Every main-text rSVD figure and table should "
            "state oversampling, power iterations, seed policy, and numerical "
            "audit status. A compact comparison of one- versus two-power "
            "runtime and prediction agreement on representative small, large, "
            "and extreme-response datasets is needed."
        ),
    )

    add_major_comment(
        document,
        2,
        "The CPU multithreading claim is not experimentally demonstrated",
        (
            "Figure 3 includes a 'four-thread request' panel based only on "
            "MetRef, while the Methods and Supplement state that reference "
            "BLAS was used and that no controlled thread-scaling experiment "
            "was conducted. The observed 1.00-1.24-fold ratios therefore do "
            "not quantify four-core acceleration and may mainly reflect a "
            "runtime request ignored by the linked BLAS."
        ),
        (
            "Remove this panel and the associated acceleration language, or "
            "perform a genuine thread-scaling experiment using a documented "
            "multithreaded BLAS with verified active thread counts. At minimum "
            "test 1, 2, and 4 threads across several matrix regimes and report "
            "accuracy, total time, and memory. Distinguish compiled CPU "
            "execution from demonstrated multicore scaling."
        ),
    )

    add_major_comment(
        document,
        3,
        "The accelerator comparison mixes speed with unresolved numerical disagreement",
        (
            "CUDA is faster in only 13 of 44 non-NMR matched comparisons, and "
            "Metal is faster in only two tested configurations. More "
            "importantly, the Metal CIFAR-100 results include accuracy changes "
            "of several percentage points for PLS-SVD and SIMPLS. Marking "
            "these cells with a dagger does not make them valid speed "
            "comparisons of the same numerical workflow. The Abstract and "
            "Conclusion still foreground accelerator backends despite these "
            "limitations."
        ),
        (
            "Separate backend performance into numerically concordant and "
            "discordant routes. Speed-up summaries should exclude or visibly "
            "quarantine routes that fail the stated metric/agreement "
            "tolerance. Provide a paired table of CPU, CUDA, and Metal "
            "prediction agreement alongside runtime and memory, and narrow "
            "the portability claim to the validated model/precision/shape "
            "combinations."
        ),
    )

    add_major_comment(
        document,
        4,
        "The external software figure combines matched and non-matched workflows too readily",
        (
            "The deterministic fastPLS SIMPLS versus pls::simpls.fit and "
            "fastPLS LDA versus plsgenomics comparisons are appropriate "
            "matched analyses. Figure 2 also displays many broader package "
            "workflows with different estimators, preprocessing conventions, "
            "response formulations, or classifiers. Although the text warns "
            "against estimator-level interpretation, the heatmap visually "
            "invites direct ranking. Absolute process RSS additionally "
            "includes package loading, data, and runtime state, so small-task "
            "memory comparisons are especially difficult to interpret."
        ),
        (
            "Visually separate the two estimator-matched comparisons from the "
            "contextual workflow survey. For matched pairs, report paired "
            "runtime ratios, predictive agreement with uncertainty, baseline "
            "RSS, incremental peak RSS, and execution status. Keep the wider "
            "package survey supplementary or clearly label it as a "
            "functionality/feasibility comparison rather than a speed ranking."
        ),
    )

    add_major_comment(
        document,
        5,
        "The newly expanded OPLS and kernel-PLS evidence is strong but not synchronized with the manuscript",
        (
            "Supplementary Section S41 now reports 66 of 66 endpoint passes, "
            "66 of 66 component-selection agreements, and 1,540 successful "
            "fold-component fits across OPLS north=1-3 and eight kernel "
            "settings. The Abstract still reports only 18 endpoint and 18 "
            "selection comparisons. The Methods describe the narrower study, "
            "and the Discussion refers generally to the evaluated kernels "
            "without acknowledging the expanded setting grid."
        ),
        (
            "Update the Abstract, Methods, Results, Discussion, and "
            "reproducibility ledger to describe the 66-setting/task study "
            "consistently. State clearly that it validates implementation "
            "reliability, not predictive superiority of any kernel. Preserve "
            "the important scope limitation to deterministic float64 CPU "
            "execution."
        ),
    )

    add_major_comment(
        document,
        6,
        "The NMR section still combines model selection and implementation benchmarking",
        (
            "The family-selected comparison uses five PLS-SVD components and "
            "50 SIMPLS components, whereas Figure 4 fixes all methods at 100 "
            "components and simultaneously changes PLS family, SVD solver, "
            "implementation, and hardware. The manuscript acknowledges this "
            "confounding, but the figure remains the central NMR performance "
            "display. Moreover, global RMSD and Q2 can be dominated by broad "
            "low-intensity regions in a 28,355-response spectrum."
        ),
        (
            "Present two explicitly separate analyses: predictive model "
            "selection at family-specific training-selected components, and "
            "a paired computational benchmark in which family, solver, "
            "precision, and component count are held fixed while only the "
            "implementation/backend changes. Report per-spectrum and "
            "response-wise error distributions, peak-region or intensity-"
            "stratified errors, and uncertainty for the family-selected "
            "models. Clarify why the deposited historical reference is "
            "evaluated at 100 components if the original scientific workflow "
            "used a different component count."
        ),
    )

    add_major_comment(
        document,
        7,
        "The ImageNet analysis remains exploratory and should not carry a quasi-confirmatory accuracy claim",
        (
            "The pooled archive lacks the canonical train/validation flag, "
            "the development holdout informed earlier tuning, and most "
            "ImageNet measurements are single runs. The manuscript states "
            "these limitations, yet the Abstract highlights top-1 accuracy "
            "0.8093 without an uncertainty interval or an external baseline "
            "under the same split. The PLS retrieval result is slightly worse "
            "in top-1 than raw DINOv2, and the small top-5 difference is not "
            "evidence of improvement."
        ),
        (
            "Retain ImageNet as a feasibility and supervised-compression case "
            "study, but move the isolated accuracy value out of the Abstract "
            "or label it explicitly as exploratory. Report repeat dispersion "
            "for transformation, fitting, projection, indexing, and query "
            "time. Clearly distinguish the SIMPLS-LDA experiment from the "
            "PLS-SVD/FAISS representation experiment and avoid interpreting "
            "the reused development holdout as an independent test set."
        ),
    )

    add_major_comment(
        document,
        8,
        "Float32 is presented as a package-wide capability despite route-specific failures",
        (
            "The manuscript is appropriately cautious, but the package-level "
            "description still suggests automatic float32 execution across "
            "all four PLS families and backends. The Supplement shows major "
            "slowdowns, accuracy/RMSD changes, unavailable Windows paths, and "
            "host-assisted routes. Halving the representation size does not "
            "by itself establish reduced peak memory or acceptable numerical "
            "behavior."
        ),
        (
            "Provide a single authoritative capability matrix listing every "
            "family/backend/classifier combination as validated, "
            "experimental, hybrid, unavailable, or failed. For supported "
            "float32 routes, report paired float64 differences in prediction, "
            "metric, runtime, baseline RSS, incremental RSS, and device "
            "memory. Ensure that unsupported or numerically unsafe routes "
            "fail or warn consistently in the public API."
        ),
    )

    add_major_comment(
        document,
        9,
        "Predictive uncertainty and component selection remain conditional on one outer split",
        (
            "Wilson and held-out bootstrap intervals quantify uncertainty "
            "conditional on a fixed test set, not variability from training "
            "sampling, preprocessing, or component selection. More than half "
            "of the family-specific choices occur at a grid boundary, and "
            "some PLS-SVD paths are rank constrained. Timing replicates do not "
            "address predictive uncertainty."
        ),
        (
            "For the principal biomedical claims, add repeated outer splits "
            "or nested cross-validation for a representative small, medium, "
            "and large classification task and for NMR if feasible. At "
            "minimum, report selection frequencies and predictive dispersion "
            "across repeated training partitions. Continue to describe "
            "boundary selections as best within the evaluated grid, not "
            "optimal."
        ),
    )

    add_major_comment(
        document,
        10,
        "The software provenance is not yet sufficient for exact reproduction of all headline analyses",
        (
            "The manuscript acknowledges that several historical archives "
            "record package version 0.99.6 without a recoverable Git SHA. This "
            "is transparent but leaves the reader unable to reconstruct the "
            "exact source for some principal results. The newly added S41 "
            "experiment also needs an explicit source-commit mapping. A "
            "future repository state cannot substitute for the code used to "
            "generate a reported number."
        ),
        (
            "Rerun the final headline tables and figures from one frozen, "
            "tagged commit and record package SHA, kodama-cpp SHA, compiler "
            "flags, BLAS/LAPACK, thread counts, CUDA/Metal libraries, dataset "
            "checksums, split checksum, command line, and random seeds in each "
            "result manifest. Archive the exact release and generated tables "
            "at a persistent DOI-bearing repository before publication."
        ),
    )

    add_major_comment(
        document,
        11,
        "The practical API and backend semantics need a concise user-facing validation",
        (
            "The manuscript describes a unified pls() interface, estimator "
            "enforcement, compiled cross-validation, compact prediction, and "
            "automatic precision dispatch. However, the Supplement is "
            "primarily numerical and does not provide a compact end-to-end "
            "usability assessment showing that identical public calls produce "
            "predictable model objects, warnings, errors, and predictions "
            "across platforms."
        ),
        (
            "Add a small public-API conformance suite covering regression and "
            "classification for all four families, argmax/LDA, CPU/CUDA/Metal "
            "availability, float64/float32 dispatch, unsupported routes, "
            "predict(), and both cross-validation functions. Report pass/fail "
            "status on Linux, macOS, and Windows CPU builds and on available "
            "accelerator systems."
        ),
    )

    add_major_comment(
        document,
        12,
        "The presentation is too long and internally repetitive for the central message",
        (
            "The Supplement is 124 pages and contains repeated validation "
            "summaries, overlapping capability tables, multiple component "
            "paths, and review-cycle material. This density makes it difficult "
            "to identify the definitive evidence. The Discussion also repeats "
            "the sentence beginning 'The computational gain derives from "
            "compiled execution'."
        ),
        (
            "Consolidate duplicate tables and move archival review-cycle "
            "material to the repository. Retain one definitive table each for "
            "estimator validation, rSVD reliability, backend residency, "
            "float32 capability, external comparison, selected-point "
            "performance, and provenance. Remove duplicated prose and ensure "
            "that every main-text claim maps to one authoritative "
            "Supplementary table or figure."
        ),
    )

    document.add_heading("Minor Comments", level=1)
    minor_comments = [
        (
            "The Abstract reports 18 OPLS/kernel endpoint comparisons, whereas "
            "Supplementary Section S41 reports 66. Synchronize all counts."
        ),
        (
            "Algorithm 1 uses K for requested component prefixes, but Section "
            "2.1 reserves K for the number of cross-validation folds. Use a "
            "different symbol for the set of requested prefixes."
        ),
        (
            "Remove the duplicated Discussion sentence: 'The computational "
            "gain derives from compiled execution...'."
        ),
        (
            "Use one notation for Q-squared and define whether the reported "
            "NMR value is global, response-weighted, or averaged over response "
            "columns."
        ),
        (
            "State in every figure caption whether time includes fitting, "
            "prediction, host-device transfers, synchronization, and returned "
            "model transfer; Figure 3 currently relies on the Methods for this."
        ),
        (
            "Figure 3B should not present a speed-up as favorable when the "
            "predictive metric fails the agreement tolerance. Use a distinct "
            "failed/discordant encoding rather than only a dagger."
        ),
        (
            "The statement that CPU builds can optionally use OpenBLAS is not "
            "evidence of multicore acceleration. Keep implementation "
            "capability separate from measured performance."
        ),
        (
            "Clarify whether the OPLS ncomp argument denotes predictive "
            "components or the total predictive-plus-orthogonal budget in the "
            "public API, cross-validation, and benchmark tables."
        ),
        (
            "For nonlinear kernel PLS, report the computational consequence "
            "of the O(n-squared) Gram matrix more prominently and give explicit "
            "practical sample-size limits for the tested hardware."
        ),
        (
            "The main Results should state the number of successful and failed "
            "external-package runs, not only 'completed datasets'."
        ),
        (
            "Provide confidence intervals or exact correct/test counts beside "
            "the principal classification accuracies in the main text."
        ),
        (
            "Clarify whether the NMR water interval was excluded from the "
            "response metric, zeroed only in predictors, or both. The current "
            "phrasing differs between text and Figure 4."
        ),
        (
            "State whether the representative spectrum in Figure 4 was chosen "
            "a priori, randomly, or by median RMSD."
        ),
        (
            "The ImageNet LDA and FAISS analyses use different estimators and "
            "objectives; label them as separate experiments in the figure and "
            "Results."
        ),
        (
            "The DINOv2 reference should include complete bibliographic "
            "information or a stable article URL/identifier."
        ),
        (
            "The availability statement should identify the exact release tag "
            "reviewed here rather than only repository URLs."
        ),
        (
            "Check all cross-references after consolidation; the manuscript "
            "now refers to Supplementary Table S41 for provenance while S41 "
            "is also a section number, which is legal but easy to misread."
        ),
        (
            "Perform a final language and notation edit for consistent "
            "PLS-SVD, rSVD, float32/float64, CPU/CUDA/Metal, component, fold, "
            "and prediction-head terminology."
        ),
    ]
    for index, comment in enumerate(minor_comments, 1):
        add_minor_comment(document, index, comment)

    document.add_heading("Recommendation", level=1)
    document.add_paragraph(
        (
            "The manuscript contains a publishable computational contribution, "
            "and I would encourage resubmission after major revision. The "
            "highest priorities are to align the rSVD settings used for "
            "validation and headline performance, remove or replace the "
            "unsupported multithreading comparison, separate numerically "
            "discordant accelerator routes from valid speed comparisons, "
            "synchronize the new OPLS/kernel validation throughout the paper, "
            "and freeze all final analyses to one reproducible source release."
        )
    )

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_report()
