from pathlib import Path

from docx import Document


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE_DIR = ROOT / "artifacts/CMPB_rewrite_20260826_cycle113"
OUTDIR = ROOT / "artifacts/CMPB_rewrite_20260826_cycle114"


MAIN_REPLACEMENTS = {
    "Numerical equivalence and software-level performance were assessed separately. Deterministic float64 fastPLS SIMPLS was first compared with the de Jong SIMPLS implementation in pls::simpls.fit to evaluate the numerical kernel. A separate archived-release CPU experiment compared fastPLS 0.99.25 with IKPLS 6.1.2 NumPy Algorithms 1 and 2 using identical stored splits, externally training-centred float64 predictors, centred one-hot responses, component counts, and final held-out predictions. Each route was executed three times in a fresh process with one effective thread. Because IKPLS and SIMPLS implement different estimators and retain different internal state, this experiment was interpreted as an end-to-end software comparison rather than an estimator-matched benchmark. An earlier JAX/CUDA development comparison could not be reproduced from the frozen environment and was therefore excluded from the central evidence.":
    "Numerical equivalence and software-level performance were assessed separately. Deterministic float64 fastPLS SIMPLS was first compared with the de Jong SIMPLS implementation in pls::simpls.fit to evaluate the numerical kernel. A separate archived-release CPU experiment compared fastPLS 0.99.25 with the two NumPy formulations provided by IKPLS 6.1.2. The score-explicit formulation, termed Algorithm 1 in IKPLS, computes and retains the training-score matrix, whereas the cross-product formulation, termed Algorithm 2, fits from the predictor cross-product and predictor-response cross-product without retaining training scores. Both formulations used identical stored splits, externally training-centred float64 predictors, centred one-hot responses, component counts, and final held-out predictions. Each route was executed three times in a fresh process with one effective thread. Because IKPLS and SIMPLS implement different estimators and retain different internal state, this experiment was interpreted as an end-to-end software comparison rather than an estimator-matched benchmark. An earlier JAX/CUDA development comparison could not be reproduced from the frozen environment and was therefore excluded from the central evidence.",

    "The archived CPU comparison completed all 36 planned runs. IKPLS NumPy Algorithm 2 was fastest: median totals were 0.00059 s on Breast, 0.00298 s on MetRef, and 0.218 s on CIFAR-100, compared with 0.005, 0.033, and 3.585 s for fastPLS rSVD. Breast accuracy was identical (94.29%). IKPLS reached 75.0% on MetRef and 70.95% on CIFAR-100; fastPLS rSVD reached 77.0% and 72.13%, while deterministic fastPLS IRLBA reached 77.0% and 70.77%. CIFAR-100 complete-process peak RSS was 584 MiB for IKPLS Algorithm 2 and 1,178 MiB for fastPLS rSVD, including different language runtimes and allocators.":
    "The archived CPU comparison completed all 36 planned runs. The IKPLS cross-product formulation was fastest, with median total times of 0.00059 s on Breast, 0.00298 s on MetRef, and 0.218 s on CIFAR-100, compared with 0.005, 0.033, and 3.585 s for fastPLS rSVD. Breast accuracy was identical (94.29%). IKPLS reached 75.0% on MetRef and 70.95% on CIFAR-100; fastPLS rSVD reached 77.0% and 72.13%, while deterministic fastPLS IRLBA reached 77.0% and 70.77%. On CIFAR-100, complete-process peak resident-set size was 584 MiB for the IKPLS cross-product formulation and 1,178 MiB for fastPLS rSVD, including their different language runtimes and memory allocators.",
}


SUPPLEMENT_REPLACEMENTS = {
    "This archived-release comparison is not estimator matched. IKPLS 6.1.2 implements Dayal-MacGregor Improved Kernel PLS, whereas fastPLS implements de Jong SIMPLS with deterministic IRLBA or approximate rSVD. The common CPU contract was float64 input, externally applied training centring, identical centred one-hot responses, identical splits and component counts, final held-out prediction, three fresh-process repetitions, and one effective thread. RSS includes language-runtime allocation and is a workflow-feasibility measurement.":
    "This archived-release comparison is not estimator matched. IKPLS 6.1.2 implements Dayal-MacGregor Improved Kernel PLS in two formulations: the score-explicit variant (IKPLS Algorithm 1) computes and retains the training-score matrix, whereas the cross-product variant (IKPLS Algorithm 2) fits from the predictor cross-product and predictor-response cross-product without retaining training scores. fastPLS instead implements de Jong SIMPLS with deterministic IRLBA or approximate rSVD. The common CPU contract comprised float64 input, externally applied training centring, identical centred one-hot responses, identical splits and component counts, final held-out prediction, three fresh-process repetitions, and one effective thread. Resident-set size includes language-runtime allocation and is therefore interpreted as a workflow-feasibility measurement.",

    "Table S10f. Archived-release single-thread CPU end-to-end comparison with IKPLS. Time and IQR are fit-plus-prediction seconds; RSS values are MiB. fastPLS rSVD rows use version 0.99.25, oversampling 20, two power iterations, seed 123, and case-specific diagnostics.":
    "Table S10f. Archived-release single-thread CPU end-to-end comparison with IKPLS. The score-explicit variant corresponds to IKPLS Algorithm 1 and the cross-product variant to IKPLS Algorithm 2. Time and IQR are fitting-plus-prediction seconds; resident-set-size values are MiB. fastPLS rSVD rows use version 0.99.25, oversampling 20, two power iterations, seed 123, and case-specific diagnostics.",
}


TABLE_LABELS = {
    "IKPLS NumPy Algorithm 1": "IKPLS score-explicit (Alg. 1)",
    "IKPLS NumPy Algorithm 2": "IKPLS cross-product (Alg. 2)",
}


def replace_paragraphs(document: Document, replacements: dict[str, str]) -> None:
    for old, new in replacements.items():
        matches = [paragraph for paragraph in document.paragraphs if paragraph.text == old]
        if len(matches) != 1:
            raise RuntimeError(f"Expected one paragraph match, found {len(matches)} for: {old[:90]}")
        matches[0].text = new


def replace_table_labels(document: Document) -> None:
    counts = {old: 0 for old in TABLE_LABELS}
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text in TABLE_LABELS:
                            old = run.text
                            run.text = TABLE_LABELS[old]
                            counts[old] += 1
    if counts != {"IKPLS NumPy Algorithm 1": 3, "IKPLS NumPy Algorithm 2": 3}:
        raise RuntimeError(f"Unexpected IKPLS table label counts: {counts}")


def main() -> None:
    OUTDIR.mkdir(parents=True, exist_ok=True)

    main_doc = Document(SOURCE_DIR / "fastPLS_CMPB_main_cycle113_0.99.25_20260826.docx")
    replace_paragraphs(main_doc, MAIN_REPLACEMENTS)
    main_doc.save(OUTDIR / "fastPLS_CMPB_main_cycle114_0.99.25_20260826.docx")

    supplement = Document(SOURCE_DIR / "fastPLS_CMPB_supplement_cycle113_0.99.25_20260826.docx")
    replace_paragraphs(supplement, SUPPLEMENT_REPLACEMENTS)
    replace_table_labels(supplement)
    supplement.save(OUTDIR / "fastPLS_CMPB_supplement_cycle114_0.99.25_20260826.docx")


if __name__ == "__main__":
    main()
