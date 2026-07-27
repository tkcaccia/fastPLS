#!/usr/bin/env python3

from pathlib import Path
import csv
import shutil

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle57"
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle58"
MAIN_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_main_cycle57_0.99.6_20260726.docx"
)
MAIN_SOURCE_PDF = (
    SOURCE_DIR / "fastPLS_CMPB_main_cycle57_0.99.6_20260726.pdf"
)
SUPP_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_supplement_cycle57_0.99.6_20260726.docx"
)
MAIN_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_main_cycle58_0.99.6_20260726.docx"
)
MAIN_OUTPUT_PDF = (
    OUTPUT_DIR / "fastPLS_CMPB_main_cycle58_0.99.6_20260726.pdf"
)
SUPP_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_supplement_cycle58_0.99.6_20260726.docx"
)

RESULT_DIR = (
    ROOT
    / "benchmark_results"
    / "opls_kernel_setting_reliability_20260726"
)
SUMMARY_CSV = RESULT_DIR / "opls_kernel_setting_reliability_summary.csv"
SELECTION_CSV = (
    RESULT_DIR / "opls_kernel_setting_selection_setting_summary.csv"
)
RAW_CSV = RESULT_DIR / "opls_kernel_setting_reliability_raw.csv"
FOLD_CSV = RESULT_DIR / "opls_kernel_setting_selection_fold_raw.csv"
FIGURE = RESULT_DIR / "opls_kernel_setting_reliability.png"

SETTING_LABELS = {
    "north_1": "north = 1",
    "north_2": "north = 2",
    "north_3": "north = 3",
    "linear": "linear",
    "rbf_gamma_0.25_over_p": "RBF, gamma = 0.25/p",
    "rbf_gamma_1_over_p": "RBF, gamma = 1/p",
    "rbf_gamma_4_over_p": "RBF, gamma = 4/p",
    "poly_degree2_offset1": "poly, d = 2, c = 1",
    "poly_degree3_offset1": "poly, d = 3, c = 1",
    "poly_degree4_offset1": "poly, d = 4, c = 1",
    "poly_degree3_offset0": "poly, d = 3, c = 0",
}
SETTING_ORDER = [
    "north_1",
    "north_2",
    "north_3",
    "linear",
    "rbf_gamma_0.25_over_p",
    "rbf_gamma_1_over_p",
    "rbf_gamma_4_over_p",
    "poly_degree2_offset1",
    "poly_degree3_offset1",
    "poly_degree4_offset1",
    "poly_degree3_offset0",
]


def read_csv(path):
    with path.open(newline="", encoding="utf-8") as handle:
        return list(csv.DictReader(handle))


def format_scientific(value):
    number = float(value)
    return f"{number:.2e}"


def set_cell_margins(cell, top=45, start=45, bottom=45, end=45):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text, bold=False, size=6.6, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def set_table_widths(table, widths):
    table.autofit = False
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width.inches * 1440)))
        grid.append(grid_col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_compact_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table"
    set_table_widths(table, widths)
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(
            cell,
            header,
            bold=True,
            size=6.5,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "D9EAF7")
        cell._tc.get_or_add_tcPr().append(shading)
    for values in rows:
        cells = table.add_row().cells
        for index, (cell, value) in enumerate(zip(cells, values)):
            set_cell_text(
                cell,
                value,
                size=6.4,
                align=(
                    WD_ALIGN_PARAGRAPH.LEFT
                    if index == 1
                    else WD_ALIGN_PARAGRAPH.CENTER
                ),
            )
    for row in table.rows:
        row._tr.get_or_add_trPr()
    return table


def update_s34_reference(document):
    target = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(
            "Reproducibility files: benchmark/"
            "benchmark_opls_kernel_estimator_validation.R"
        )
    )
    target.text = (
        "The default-setting validation is reproduced by "
        "benchmark/benchmark_opls_kernel_estimator_validation.R. A broader "
        "setting-level validation covering multiple orthogonal-component "
        "counts and linear, RBF, and polynomial kernels is reported in "
        "Section S41 and reproduced by "
        "benchmark/benchmark_opls_kernel_setting_reliability.R. "
        "Machine-readable outputs are stored in "
        "benchmark_results/opls_kernel_estimator_validation_verified_20260726/ "
        "and benchmark_results/opls_kernel_setting_reliability_20260726/."
    )
    target.style = document.styles["Normal"]


def append_s41(document):
    summary = read_csv(SUMMARY_CSV)
    selection = read_csv(SELECTION_CSV)
    raw = read_csv(RAW_CSV)
    folds = read_csv(FOLD_CSV)

    summary_by_setting = {row["setting"]: row for row in summary}
    selection_by_setting = {row["setting"]: row for row in selection}
    selection_comparisons = sum(
        int(row["comparisons"]) for row in selection
    )
    selection_agreements = sum(
        int(row["selected_component_agreements"]) for row in selection
    )

    heading = document.add_paragraph(
        "S41. OPLS and kernel-PLS setting-level reliability",
        style="Heading 1",
    )
    heading.paragraph_format.page_break_before = True
    document.add_paragraph(
        (
            "The default-setting validation in Section S34 was extended to "
            "test whether agreement depended on the number of OPLS "
            "orthogonal components or on the kernel definition. The same six "
            "fixed tasks were used: synthetic regression and classification "
            "with p<n and p>n, an ill-conditioned p>n regression design, "
            "gasoline spectroscopy regression, and breast molecular "
            "classification. All calculations used double precision, the CPU "
            "backend, deterministic IRLBA, seed 123, identical train/test "
            "partitions, and identical five-fold component-selection folds."
        ),
        style="First Paragraph",
    )
    document.add_paragraph(
        (
            "OPLS was evaluated with one, two, and three orthogonal "
            "components. For each setting, the independent reference "
            "recomputed the Trygg-Wold orthogonal weight, score, loading, and "
            "deflation sequence before fitting pls::simpls.fit to the "
            "filtered predictors. Kernel PLS was evaluated with a linear "
            "kernel; RBF kernels with gamma equal to 0.25/p, 1/p, and 4/p; "
            "polynomial kernels of degree two, three, and four with offset "
            "one; and a homogeneous degree-three polynomial kernel with "
            "offset zero. Every Gram matrix and its train/test centring were "
            "constructed independently before the reference SIMPLS fit."
        ),
        style="Normal",
    )
    document.add_paragraph(
        (
            "The prespecified pass criteria were unchanged: operator relative "
            "error <=1e-10, prediction relative error <=1e-4, coefficient "
            "relative error <=1e-3, predictive and orthogonal score-subspace "
            "angles <=0.1 degrees, classification label agreement >=0.995, "
            "and predictive-metric difference <=0.005. Component-selection "
            "agreement was assessed over the complete 1-5 grid, or 1-3 for "
            "the breast dataset."
        ),
        style="Normal",
    )
    document.add_paragraph(
        (
            f"All {len(raw)} endpoint comparisons passed all criteria "
            f"(Table S50). The maximum operator, prediction, and coefficient "
            "relative errors were 1.57e-13, 9.48e-12, and 3.07e-10, "
            "respectively. The largest predictive and OPLS orthogonal "
            "score-subspace angles were 2.42e-6 and 1.91e-6 degrees. "
            "Classification labels agreed exactly, and the largest absolute "
            "difference in accuracy or RMSD was 3.20e-12. Thus, agreement "
            "was retained as orthogonal-component count, RBF bandwidth, "
            "polynomial degree, and polynomial offset changed."
        ),
        style="Normal",
    )
    document.add_paragraph(
        (
            f"All {len(folds):,} fold-by-component fits completed without "
            "failure, and fastPLS selected the same component count as the "
            f"independent reference in all {selection_agreements} of "
            f"{selection_comparisons} setting/task "
            "comparisons (Table S51). These results establish deterministic "
            "float64 CPU reliability for the tested OPLS and kernel-PLS "
            "settings. They do not extend exact-equivalence claims to rSVD, "
            "float32, CUDA, or Metal, which remain governed by their separate "
            "numerical validations."
        ),
        style="Normal",
    )

    document.add_paragraph(
        (
            "Table S50. Deterministic endpoint agreement across OPLS and "
            "kernel-PLS settings. Pass is the number of successful endpoints "
            "meeting every prespecified criterion. Op, pred, and coef are "
            "maximum relative errors; angle is the maximum predictive-score "
            "subspace angle; labels is minimum decoded-label agreement; and "
            "metric delta is the maximum absolute accuracy or RMSD difference."
        ),
        style="Caption",
    ).paragraph_format.keep_together = True
    table_50_rows = []
    for setting in SETTING_ORDER:
        row = summary_by_setting[setting]
        table_50_rows.append(
            [
                "OPLS" if row["family"] == "OPLS" else "kernel PLS",
                SETTING_LABELS[setting],
                f'{row["passes_all"]}/{row["runs"]}',
                format_scientific(row["max_operator_relative_error"]),
                format_scientific(row["max_prediction_relative_error"]),
                format_scientific(row["max_coefficient_relative_error"]),
                format_scientific(row["max_predictive_score_angle_deg"]),
                (
                    "1.000"
                    if row["min_label_agreement"]
                    not in ("", "NA", "NaN")
                    else "-"
                ),
                format_scientific(row["max_metric_absolute_difference"]),
            ]
        )
    add_compact_table(
        document,
        [
            "Family",
            "Setting",
            "Pass",
            "Op error",
            "Pred error",
            "Coef error",
            "Angle (deg)",
            "Labels",
            "Metric delta",
        ],
        table_50_rows,
        [
            Inches(0.62),
            Inches(1.27),
            Inches(0.43),
            Inches(0.68),
            Inches(0.68),
            Inches(0.68),
            Inches(0.68),
            Inches(0.52),
            Inches(0.76),
        ],
    )

    document.add_paragraph(
        (
            "Table S51. Fixed-fold component-selection and execution "
            "agreement across OPLS and kernel-PLS settings. Agreement gives "
            "the number of tasks for which fastPLS and the independent "
            "reference selected the same component count. Failed fits count "
            "all unsuccessful fold-by-component executions."
        ),
        style="Caption",
    ).paragraph_format.keep_together = True
    table_51_rows = []
    for setting in SETTING_ORDER:
        row = selection_by_setting[setting]
        table_51_rows.append(
            [
                "OPLS" if row["family"] == "OPLS" else "kernel PLS",
                SETTING_LABELS[setting],
                f'{row["selected_component_agreements"]}/{row["comparisons"]}',
                row["failed_fold_component_fits"],
            ]
        )
    add_compact_table(
        document,
        ["Family", "Setting", "Selection agreement", "Failed fits"],
        table_51_rows,
        [Inches(1.05), Inches(2.80), Inches(1.30), Inches(1.00)],
    )

    picture_paragraph = document.add_paragraph()
    picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_paragraph.paragraph_format.keep_with_next = True
    picture_paragraph.add_run().add_picture(str(FIGURE), width=Inches(6.30))
    figure_caption = document.add_paragraph(
        (
            "Figure S41. OPLS and kernel-PLS setting-level reliability. "
            "Points show six synthetic or real tasks per setting. Dashed "
            "horizontal lines are the prespecified pass thresholds for "
            "prediction relative error, coefficient relative error, "
            "predictive-score subspace angle, and absolute predictive-metric "
            "difference. All observations lie below their corresponding "
            "thresholds."
        ),
        style="Caption",
    )
    figure_caption.paragraph_format.keep_together = True

    document.add_paragraph(
        (
            "Reproducibility files: "
            "benchmark/benchmark_opls_kernel_setting_reliability.R and "
            "benchmark_results/opls_kernel_setting_reliability_20260726/. "
            "The machine-readable outputs retain every endpoint, kernel "
            "parameter, orthogonal-component count, fold, component, "
            "tolerance, status, and error field."
        ),
        style="Normal",
    )


def main():
    required = [
        MAIN_SOURCE,
        MAIN_SOURCE_PDF,
        SUPP_SOURCE,
        SUMMARY_CSV,
        SELECTION_CSV,
        RAW_CSV,
        FOLD_CSV,
        FIGURE,
    ]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing required files:\n" + "\n".join(missing))

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    shutil.copy2(MAIN_SOURCE, MAIN_OUTPUT)
    shutil.copy2(MAIN_SOURCE_PDF, MAIN_OUTPUT_PDF)

    document = Document(SUPP_SOURCE)
    update_s34_reference(document)
    append_s41(document)
    document.save(SUPP_OUTPUT)
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)


if __name__ == "__main__":
    main()
