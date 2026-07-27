#!/usr/bin/env python3

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle48"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle49"
RESULTS = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle49_20260726"
)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle48_0.99.6_20260726.docx"
SUPP_SOURCE = (
    SOURCE / "fastPLS_CMPB_supplement_cycle48_0.99.6_20260726.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle49_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle49_0.99.6_20260726.docx"
FIGURE = RESULTS / "main_selected_computational_performance.png"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def insert_after(anchor, paragraph):
    anchor._p.addnext(paragraph._p)
    return paragraph


def revise_main():
    if not FIGURE.exists():
        raise FileNotFoundError(FIGURE)

    document = Document(MAIN_SOURCE)
    figure_4_caption = find_paragraph(document, "Figure 4.")

    heading = document.add_paragraph(
        "3.4 Cross-dataset computational performance",
        style="Heading 2",
    )
    paragraph = document.add_paragraph(
        (
            "At each family-specific component count selected using training data, "
            "the faster backend depended on matrix shape rather than PLS family "
            "alone (Figure 5). Across 44 paired family comparisons, CPU execution "
            "was faster in 31 and CUDA in 13. CUDA was faster for all four families "
            "on CIFAR-100, with a median 11.2-fold CPU/CUDA time ratio across "
            "families, and for three of four families on GTEx v8, TCGA Pan-Cancer, "
            "and CBMC CITE-seq. CPU was faster for all four families on the "
            "remaining seven datasets. Median absolute peak host RSS across rows "
            "was 489 MB for CPU and 819 MB for CUDA; sampled CUDA process memory "
            "ranged from 192 to 532 MB. These memory values characterize complete "
            "benchmark-process feasibility and include loaded data, libraries, "
            "and accelerator runtime state; they are not isolated workspace "
            "allocations."
        ),
        style="First Paragraph",
    )
    page_break = document.add_page_break()
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.add_run().add_picture(str(FIGURE), width=Inches(5.2))
    caption = document.add_paragraph(
        (
            "Figure 5. Selected-point computational performance across 11 datasets "
            "(NMR and ImageNet excluded). Points and bars are medians and "
            "interquartile ranges from three isolated runs at each family-specific "
            "training-selected component count. Colors denote PLS families; circles "
            "CPU and triangles CUDA. (A) Fitting plus prediction time. (B) Absolute "
            "peak process RSS. (C) Sampled process GPU memory including CUDA "
            "context. Axes are logarithmic; paths use float64 rSVD. Memory is "
            "process-level, not workspace-only."
        ),
        style="Caption",
    )

    anchor = figure_4_caption
    for item in (heading, paragraph, page_break, image_paragraph, caption):
        anchor = insert_after(anchor, item)

    precision_heading = find_paragraph(
        document,
        "3.4 Precision and backend agreement",
    )
    precision_heading.text = "3.5 Precision and backend agreement"
    external_heading = find_paragraph(
        document,
        "3.5 External software and cross-validation",
    )
    external_heading.text = "3.6 External software and cross-validation"

    document.core_properties.title = (
        "fastPLS CMPB manuscript - cross-dataset computational performance"
    )
    document.save(MAIN_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    shutil.copy2(SUPP_SOURCE, SUPP_OUT)
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
