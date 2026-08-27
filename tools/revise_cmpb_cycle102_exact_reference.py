from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle101"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle102"
OUTPUT.mkdir(parents=True, exist_ok=True)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle101_0.99.25_20260825.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle101_0.99.25_20260825.docx"
MAIN_OUTPUT = OUTPUT / "fastPLS_CMPB_main_cycle102_0.99.25_20260825.docx"
SUPP_OUTPUT = OUTPUT / "fastPLS_CMPB_supplement_cycle102_0.99.25_20260825.docx"


def paragraph_by_prefix(document, prefix):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}; found {len(matches)}")
    return matches[0]


def set_cell_text(cell, value, size=6.5):
    cell.text = str(value)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(size)


def append_table_row(table, values, size=6.5):
    template = deepcopy(table.rows[-1]._tr)
    table._tbl.append(template)
    row = table.rows[-1]
    if len(row.cells) != len(values):
        raise RuntimeError(f"Table expects {len(row.cells)} values; got {len(values)}")
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value, size=size)


def update_main():
    document = Document(MAIN_SOURCE)

    results = paragraph_by_prefix(document, "Results:")
    results.text = (
        "Results: An independent dense-LAPACK de Jong panel completed all 82 component-prefix "
        "comparisons without numerical failure. Outside an intentionally near-tied singular-value "
        "case, maximum held-out prediction error was 4.53 x 10^-15; the tied case retained a "
        "0.015-degree maximum subspace angle and 3.36 x 10^-4 relative prediction error. In 108 "
        "repeated single-CPU comparisons with pls::simpls.fit, accuracy was identical; fastPLS was "
        "faster on five of nine datasets under ordinary public workflows, with a largest speed-up "
        "of 4.85-fold. For the 13,000 by 28,355 NMR problem at 50 SIMPLS components, CPU IRLBA, "
        "CPU rSVD, and CUDA rSVD required 692.21, 152.09, and 3.91 s, respectively, with RMSD "
        "0.0007561, 0.0007560, and 0.0007561."
    )

    primary = paragraph_by_prefix(document, "The primary evidence tested")
    primary.text = (
        "Estimator preservation was first assessed against an independent de Jong implementation "
        "using dense LAPACK SVD on fixed well-conditioned, nearly tied, rank-deficient, collinear, "
        "p < n, p > n, high-response, effective-rank-boundary, regression, and dummy-response "
        "classification cases. IRLBA was then assessed separately as a deterministic iterative "
        "numerical reference against pls::simpls.fit. Approximate rSVD was evaluated in a third, "
        "separate qualification using prediction, subspace, decoded-label, and endpoint screens "
        "over seeds 1, 7, 19, 43, and 123. Exact definitions and results are in Supplementary "
        "Sections S12-S14 and Tables S7-S9."
    )

    validation = paragraph_by_prefix(document, "Deterministic fastPLS SIMPLS met")
    validation.text = (
        "The compiled dense-LAPACK SIMPLS path completed 82/82 component-prefix comparisons with "
        "the independent exact-reference updates and no convergence failure. Median and 95th-"
        "percentile relative held-out prediction errors were 7.62 x 10^-16 and 2.72 x 10^-15. "
        "All non-tied cases remained at floating-point precision; the deliberately near-tied case "
        "illustrated non-identifiability of individual singular vectors but retained close predictive "
        "and subspace agreement. Separately, IRLBA SIMPLS met the prespecified tolerances in all "
        "117 comparisons with pls::simpls.fit. rSVD was not used as exact estimator-preservation "
        "evidence."
    )

    algorithm = document.tables[0].cell(0, 0)
    algorithm.text = (
        "Input: centred X; numeric, indicator, or label-aware response Y; maximum component count "
        "A; requested component-count set C; solver.\n"
        "1. Define S₀ = XᵀY explicitly, or define operator products S(v) = Xᵀ(Yv) and "
        "Sᵀ(u) = Yᵀ(Xu); label-aware products use class sums.\n"
        "2. If p ≤ n and storage permits, cache G = XᵀX. Initialize R, Q, and V as empty; "
        "set fitted values Ŷ₀ = 0.\n"
        "3. For component a = 1, …, A:\n"
        "3.1 Extract the leading left singular direction rₐ in predictor space from state Sₐ₋₁ "
        "using a fresh IRLBA solve or a fresh oversampled rSVD sketch.\n"
        "3.2 Set tₐ = Xrₐ; divide rₐ and tₐ by ‖tₐ‖₂.\n"
        "3.3 Compute predictor and response loadings pₐ = Xᵀtₐ and qₐ = Yᵀtₐ.\n"
        "3.4 Orthogonalize vₐ = (I − VVᵀ)pₐ and normalize vₐ.\n"
        "3.5 Deflate Sₐ = Sₐ₋₁ − vₐ(vₐᵀSₐ₋₁), or update the equivalent implicit operator.\n"
        "3.6 Append rₐ, qₐ, and vₐ; update Bₐ = Bₐ₋₁ + rₐqₐᵀ and "
        "Ŷₐ = Ŷₐ₋₁ + tₐqₐᵀ.\n"
        "3.7 If a ∈ C, store only the requested coefficient/prediction snapshot or compact "
        "latent representation.\n"
        "Output: the standard sequential component path; requested outputs are retained as dense "
        "snapshots or compact latent factors."
    )
    for paragraph in algorithm.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(7.5)

    document.save(MAIN_OUTPUT)


def update_supplement():
    document = Document(SUPP_SOURCE)

    synthetic = paragraph_by_prefix(document, "Two prespecified synthetic designs")
    synthetic.text = (
        "Three prespecified synthetic designs answered separate questions. The independent dense-"
        "reference panel covered ten conditions and 82 component prefixes: well-conditioned, nearly "
        "tied leading singular values, rank-deficient X and Y, highly collinear predictors, p < n, "
        "p > n, high response dimension, a component requested at effective rank, multivariate "
        "regression, and dummy-response classification. A second formal comparison used five "
        "multivariate regression and three classification regimes with seeds 101, 202, and 303 for "
        "IRLBA and rSVD assessment. The separate controlled scaling design is reported in Section "
        "S12.2. Complete dimensions and generated matrices are reproducible from the repository scripts."
    )

    deterministic = paragraph_by_prefix(document, "The corresponding deterministic results comprised")
    deterministic.text = (
        "The dense-reference audit completed all 82 component prefixes without convergence failure. "
        "Across all cases, median and 95th-percentile coefficient errors were 8.50 x 10^-16 and "
        "3.39 x 10^-15; corresponding held-out prediction errors were 7.62 x 10^-16 and "
        "2.72 x 10^-15. Maximum score/loading/projection subspace angles outside the intentionally "
        "near-tied case were 2.42 x 10^-6 degrees, and classification label agreement was 1.000. "
        "For singular values separated by only 10^-12 relatively, individual basis vectors were "
        "not identifiable; prediction error was 3.36 x 10^-4 and the maximum subspace angle was "
        "0.015 degrees. Orthogonality and deflation residuals remained below 6.36 x 10^-15 and "
        "8.67 x 10^-16. Separately, the IRLBA comparison comprised 117 component-level endpoints, "
        "all meeting prespecified tolerances, with fixed-fold selection agreement in all tasks."
    )

    section = paragraph_by_prefix(document, "Deterministic estimator comparison and approximate-solver")
    section.text = (
        "Deterministic estimator comparison and approximate-solver agreement are separate questions. "
        "Table S7 first reports the independent dense-LAPACK de Jong reference panel, then the "
        "iterative IRLBA comparison with pls::simpls.fit and the independent OPLS/kernel checks. "
        "rSVD is excluded from estimator-preservation evidence and evaluated in Section S13. The "
        "dense panel reports distributions as well as maxima because worst values are dominated by "
        "the deliberately near-tied singular-value case."
    )

    caption = paragraph_by_prefix(document, "Table S7. Definitive deterministic estimator-validation summary")
    caption.text = (
        "Table S7. Definitive float64 CPU estimator-validation summary. Dense-reference results "
        "cover 82 component prefixes in ten conditions; the text reports median and 95th-percentile "
        "errors in addition to the maxima below. Angles are in degrees."
    )

    claim_table = document.tables[8]
    set_cell_text(claim_table.rows[1].cells[0],
                  "Compiled SIMPLS agrees with an independent dense-LAPACK de Jong reference")
    set_cell_text(claim_table.rows[1].cells[2],
                  "82 component prefixes across ten numerical conditions; full error distributions")

    validation_table = document.tables[9]
    existing_rows = [[cell.text for cell in row.cells] for row in validation_table.rows[1:]]
    for row in list(validation_table.rows)[1:]:
        validation_table._tbl.remove(row._tr)
    append_table_row(validation_table, [
        "SIMPLS / compiled dense LAPACK",
        "Independent de Jong updates / base LAPACK SVD",
        "82 / 82 completed",
        "Not a selection study",
        "3.36e-04 (4.53e-15 excluding tied case)",
        "3.70e-04 (5.24e-15 excluding tied case)",
        "0.0150 (2.42e-06 excluding tied case)",
        "0 decoded-label difference",
        "Machine precision except the non-identifiable near-tied basis"
    ])
    for row in existing_rows:
        append_table_row(validation_table, row)

    provenance = document.tables[26]
    append_table_row(provenance, [
        "F10",
        "Exact dense SIMPLS reference",
        "benchmark_results/frozen_release_0.99.25/simpls_exact_reference",
        "0.99.25",
        "Frozen package; post-freeze audit script 0ef55efdf44d",
        "benchmark_simpls_exact_reference.R",
        "0b38a7690db6"
    ])

    document.save(SUPP_OUTPUT)


if __name__ == "__main__":
    update_main()
    update_supplement()
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)
