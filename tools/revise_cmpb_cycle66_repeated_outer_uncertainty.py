#!/usr/bin/env python3
"""Add repeated outer-partition uncertainty and selection-stability evidence."""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle65"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle66"
EVIDENCE = ROOT / "benchmark_results" / "manuscript_revision_cycle66_20260726"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle65_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle65_0.99.6_20260726.docx"
MAIN_OUTPUT = OUTPUT / "fastPLS_CMPB_main_cycle66_0.99.6_20260726.docx"
SUPP_OUTPUT = OUTPUT / "fastPLS_CMPB_supplement_cycle66_0.99.6_20260726.docx"

SUMMARY = EVIDENCE / "repeated_outer_predictive_dispersion_summary.csv"
FREQUENCY = EVIDENCE / "repeated_outer_selection_frequency.csv"
FAILURES = EVIDENCE / "repeated_outer_failures.csv"
PREDICTION_FIGURE = EVIDENCE / "repeated_outer_predictive_dispersion.png"
SELECTION_FIGURE = EVIDENCE / "repeated_outer_selection_frequency.png"

METHOD_LABELS = {
    "plssvd": "PLS-SVD",
    "simpls": "SIMPLS",
    "opls": "OPLS",
    "kernelpls": "kernel PLS",
}
CLASSIFIER_LABELS = {
    "argmax": "argmax",
    "lda": "LDA",
    "regression": "regression",
}
DATASET_LABELS = {
    "metref": "MetRef",
    "gtex_v8": "GTEx v8",
    "retina": "Retina",
    "nmr": "NMR",
}


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_paragraph(document, prefix, text):
    paragraph = find_paragraph(document, prefix)
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def insert_after(paragraph, text="", style=None):
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    if style is not None:
        inserted.style = style
    if text:
        inserted.add_run(text)
    return inserted


def format_table(table, font_size=6.0):
    table.style = "Table"
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for row_index, row in enumerate(table.rows):
        row_properties = row._tr.get_or_add_trPr()
        row_properties.append(OxmlElement("w:cantSplit"))
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT if column_index == 0
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.bold = row_index == 0


def add_table(document, headers, rows, font_size=6.0):
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    format_table(table, font_size)
    return table


def fmt(value, digits=4):
    if pd.isna(value):
        return "NA"
    return f"{float(value):.{digits}f}"


def summarize_results(summary):
    def row(dataset, method, classifier):
        found = summary[
            (summary.dataset == dataset)
            & (summary.method == method)
            & (summary.classifier == classifier)
        ]
        if len(found) != 1:
            raise RuntimeError(
                f"Expected one summary row for {dataset}/{method}/{classifier}"
            )
        return found.iloc[0]

    metref_lda = [
        row("metref", method, "lda")
        for method in ("plssvd", "simpls", "opls", "kernelpls")
    ]
    retina_lda = [
        row("retina", method, "lda")
        for method in ("plssvd", "simpls", "opls", "kernelpls")
    ]
    gtex_argmax = [
        row("gtex_v8", method, "argmax")
        for method in ("plssvd", "simpls", "opls", "kernelpls")
    ]
    nmr = [
        row("nmr", method, "regression")
        for method in ("plssvd", "simpls")
    ]

    return (
        "Repeated outer partitions quantified variability from training sampling "
        "and component selection rather than uncertainty conditional on one "
        "holdout. Across ten stratified partitions, MetRef LDA accuracy ranged "
        f"from {min(x.metric_mean for x in metref_lda):.4f} to "
        f"{max(x.metric_mean for x in metref_lda):.4f} across PLS families, "
        f"with between-partition SDs of {min(x.metric_sd for x in metref_lda):.4f}-"
        f"{max(x.metric_sd for x in metref_lda):.4f}. Retina LDA accuracy was "
        f"{min(x.metric_mean for x in retina_lda):.4f}-"
        f"{max(x.metric_mean for x in retina_lda):.4f} "
        f"(SD {min(x.metric_sd for x in retina_lda):.4f}-"
        f"{max(x.metric_sd for x in retina_lda):.4f}), whereas GTEx v8 argmax "
        f"accuracy was {min(x.metric_mean for x in gtex_argmax):.4f}-"
        f"{max(x.metric_mean for x in gtex_argmax):.4f} "
        f"(SD {min(x.metric_sd for x in gtex_argmax):.4f}-"
        f"{max(x.metric_sd for x in gtex_argmax):.4f}). GTEx v8 LDA was not "
        "estimable because rare classes were absent from at least one inner "
        "training fold; all such runs remain reported as failures. Across five "
        f"NMR partitions, mean outer-test RMSD was {nmr[0].metric_mean:.6f} "
        f"(SD {nmr[0].metric_sd:.6f}) for PLS-SVD and "
        f"{nmr[1].metric_mean:.6f} (SD {nmr[1].metric_sd:.6f}) for SIMPLS. "
        "MetRef and GTEx v8 selections were predominantly upper-bound or "
        "response-rank constrained, whereas SIMPLS-derived retina selections "
        "varied across interior values. Component counts are therefore "
        "described throughout as best within the evaluated grid, never as "
        "unqualified optima."
    )


def revise_main(summary):
    document = Document(MAIN_SOURCE)

    anchor = find_paragraph(document, "Within each dataset, methods used identical")
    insert_after(
        anchor,
        (
            "Predictive uncertainty beyond a single outer split was assessed in "
            "representative small (MetRef), medium (GTEx v8), and large (Retina) "
            "classification tasks and in NMR regression. Classification used ten "
            "stratified 80/20 outer partitions; each outer-training partition used "
            "5-fold cross-validation to select components separately for every "
            "PLS family and for argmax and LDA. NMR used five random 80/20 outer "
            "partitions and 3-fold training-only selection for PLS-SVD and SIMPLS. "
            "Classification used deterministic float64 CPU IRLBA; NMR used "
            "fixed-seed float64 CUDA rSVD for feasibility and is explicitly an "
            "approximate-workflow analysis. We report selection frequencies, "
            "outer-test metric mean, SD, median, and empirical 2.5th-97.5th "
            "percentile range. These percentile ranges are descriptive across "
            "partitions and are not labelled confidence intervals."
        ),
        style="Body Text",
    )

    results_anchor = find_paragraph(
        document, "Formal deterministic reliability testing was completed"
    )
    insert_after(results_anchor, summarize_results(summary), style="Body Text")

    replace_paragraph(
        document,
        "Component selection is an additional limitation.",
        (
            "Component selection is an additional limitation. The initial "
            "selected-point benchmark used one prespecified outer split, and more "
            "than half of its dataset-family choices occurred at an evaluated-grid "
            "endpoint; nine additional PLS-SVD choices were capped by response "
            "rank. The repeated-partition analysis now quantifies how both "
            "prediction and selected dimensionality change when the training "
            "sample changes. It confirms stable accuracy and mainly interior "
            "SIMPLS-derived selections for Retina, greater predictive dispersion "
            "for MetRef, and persistent boundary censoring for MetRef and GTEx v8. "
            "These results reduce, but do not eliminate, uncertainty from training "
            "sampling: ten classification and five NMR partitions are not a "
            "substitute for application-specific external validation. All "
            "component counts are therefore reported as best within their "
            "evaluated training grid, and response-rank endpoints remain "
            "structurally constrained rather than globally optimal."
        ),
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - repeated outer-partition uncertainty"
    )
    document.save(MAIN_OUTPUT)


def revise_supplement(summary, frequency, failures):
    document = Document(SUPP_SOURCE)

    document.add_heading(
        "S45. Repeated outer-partition predictive uncertainty and component selection",
        level=1,
    )
    document.add_paragraph(
        (
            "Purpose and protocol. Wilson and held-out bootstrap intervals in "
            "Section S26 quantify uncertainty conditional on a fixed test set. "
            "This additional analysis instead perturbs the training sample, "
            "preprocessing estimates, inner folds, selected component count, and "
            "outer test observations together. MetRef, GTEx v8, and Retina "
            "represent small, medium, and large biomedical classification tasks. "
            "Ten fixed seeds generated stratified 80/20 outer partitions, followed "
            "by 5-fold training-only selection for PLS-SVD, SIMPLS, OPLS, and "
            "linear kernel PLS under argmax and LDA. NMR used five fixed seeds, "
            "random 80/20 partitions, 3-fold inner selection, and PLS-SVD and "
            "SIMPLS. Every final model was refitted on the full outer-training "
            "partition and evaluated once on its held-out partition."
        )
    )
    document.add_paragraph(
        (
            "Numerical scope. Classification used deterministic float64 CPU IRLBA "
            "to avoid conflating training-sample variability with randomized-SVD "
            "variation. NMR used float64 CUDA rSVD with oversampling 10, two power "
            "iterations, inner-fold seed 9101, and fit seed 123 because complete "
            "deterministic NMR repetition was not computationally practical. NMR "
            "therefore quantifies an approximate workflow. Reported 2.5th and "
            "97.5th percentiles are empirical ranges across 10 or five partitions, "
            "not formal confidence limits."
        )
    )

    document.add_paragraph(
        "Table S63. Prespecified repeated outer-partition protocol.",
        style="Caption",
    )
    add_table(
        document,
        [
            "Task",
            "Outer partitions",
            "Outer split",
            "Inner CV",
            "Families",
            "Classifier",
            "Backend / solver",
        ],
        [
            [
                "MetRef",
                "10",
                "Stratified 80/20",
                "5-fold",
                "PLS-SVD; SIMPLS; OPLS; kernel PLS",
                "argmax; LDA",
                "CPU / IRLBA / float64",
            ],
            [
                "GTEx v8",
                "10",
                "Stratified 80/20",
                "5-fold",
                "PLS-SVD; SIMPLS; OPLS; kernel PLS",
                "argmax; LDA",
                "CPU / IRLBA / float64",
            ],
            [
                "Retina",
                "10",
                "Stratified 80/20",
                "5-fold",
                "PLS-SVD; SIMPLS; OPLS; kernel PLS",
                "argmax; LDA",
                "CPU / IRLBA / float64",
            ],
            [
                "NMR",
                "5",
                "Random 80/20",
                "3-fold",
                "PLS-SVD; SIMPLS",
                "Regression",
                "CUDA / rSVD / float64",
            ],
        ],
        font_size=5.3,
    )

    document.add_paragraph(
        (
            "Table S64. Predictive dispersion and selected-component range across "
            "repeated outer partitions. Metric ranges are empirical partition "
            "quantiles. Boundary frequency is the fraction selected at the "
            "effective upper grid endpoint; PLS-SVD rank endpoints are marked."
        ),
        style="Caption",
    )
    rows = []
    for item in summary.itertuples():
        metric_digits = 6 if item.metric_name == "RMSD" else 4
        boundary = f"{100 * item.upper_boundary_frequency:.0f}%"
        if bool(item.rank_constrained_grid):
            boundary += " (rank)"
        rows.append(
            [
                DATASET_LABELS[item.dataset],
                METHOD_LABELS[item.method],
                CLASSIFIER_LABELS[item.classifier],
                str(int(item.n_outer_success)),
                (
                    f"{fmt(item.metric_mean, metric_digits)} "
                    f"({fmt(item.metric_sd, metric_digits)})"
                ),
                (
                    f"{fmt(item.metric_q025, metric_digits)}-"
                    f"{fmt(item.metric_q975, metric_digits)}"
                ),
                (
                    f"{int(item.selected_ncomp_min)}-"
                    f"{int(item.selected_ncomp_max)}"
                ),
                boundary,
            ]
        )
    add_table(
        document,
        [
            "Task",
            "Family",
            "Head",
            "n",
            "Metric mean (SD)",
            "Empirical 2.5%-97.5%",
            "Selected A range",
            "Upper boundary",
        ],
        rows,
        font_size=5.2,
    )

    document.add_paragraph(
        (
            "Table S65. Component-selection frequencies. Rows with zero frequency "
            "are omitted; values sum to one within task, family, and prediction "
            "head."
        ),
        style="Caption",
    )
    frequency_rows = [
        [
            DATASET_LABELS[item.dataset],
            METHOD_LABELS[item.method],
            CLASSIFIER_LABELS[item.classifier],
            str(int(item.selected_ncomp)),
            str(int(item.count)),
            f"{100 * item.frequency:.0f}%",
        ]
        for item in frequency.itertuples()
    ]
    add_table(
        document,
        ["Task", "Family", "Head", "Selected A", "Count", "Frequency"],
        frequency_rows,
        font_size=5.3,
    )

    failure_rows = []
    if len(failures):
        grouped = (
            failures.groupby(["dataset", "method", "classifier", "error"], dropna=False)
            .size()
            .reset_index(name="count")
        )
        for item in grouped.itertuples():
            reason = item.error
            if not isinstance(reason, str) or not reason.strip():
                reason = "Metric unavailable after completed execution"
            failure_rows.append(
                [
                    DATASET_LABELS.get(item.dataset, item.dataset),
                    METHOD_LABELS.get(item.method, item.method),
                    CLASSIFIER_LABELS.get(item.classifier, item.classifier),
                    str(int(item.count)),
                    reason,
                ]
            )
    document.add_paragraph(
        (
            "Table S66. Repeated-partition failures. Failures are retained rather "
            "than silently removed."
        ),
        style="Caption",
    )
    add_table(
        document,
        ["Task", "Family", "Head", "Count", "Reason"],
        failure_rows or [["None", "-", "-", "0", "-"]],
        font_size=5.5,
    )

    document.add_picture(str(PREDICTION_FIGURE), width=Inches(6.5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.paragraphs[-1].paragraph_format.keep_with_next = True
    caption = document.add_paragraph(
        (
            "Figure S43. Predictive dispersion across repeated outer partitions. "
            "Points are held-out partition results; boxes summarize the empirical "
            "distribution. Accuracy is shown for classification and RMSD for NMR."
        ),
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = False
    document.add_picture(str(SELECTION_FIGURE), width=Inches(6.5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.paragraphs[-1].paragraph_format.keep_with_next = True
    caption = document.add_paragraph(
        (
            "Figure S44. Training-only component-selection frequencies across "
            "repeated outer partitions. Endpoint selections identify censoring "
            "within the evaluated grid and are not interpreted as global optima."
        ),
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = False

    document.core_properties.title = (
        "fastPLS CMPB supplement - repeated outer-partition uncertainty"
    )
    document.save(SUPP_OUTPUT)


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    summary = pd.read_csv(SUMMARY)
    frequency = pd.read_csv(FREQUENCY)
    failures = pd.read_csv(FAILURES)
    required = {
        ("metref", "plssvd", "argmax"),
        ("gtex_v8", "plssvd", "argmax"),
        ("retina", "plssvd", "argmax"),
        ("nmr", "plssvd", "regression"),
        ("nmr", "simpls", "regression"),
    }
    observed = set(zip(summary.dataset, summary.method, summary.classifier))
    missing = required - observed
    if missing:
        raise RuntimeError(f"Repeated outer evidence is incomplete: {sorted(missing)}")
    revise_main(summary)
    revise_supplement(summary, frequency, failures)
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)


if __name__ == "__main__":
    main()
