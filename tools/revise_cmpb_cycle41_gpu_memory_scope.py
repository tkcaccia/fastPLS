#!/usr/bin/env python3

from pathlib import Path

from docx import Document


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle40"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle41"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle40_0.99.6_20260726.docx"
SUPP_SOURCE = (
    SOURCE / "fastPLS_CMPB_supplement_cycle40_0.99.6_20260726.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle41_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle41_0.99.6_20260726.docx"


def replace_text(document, old, new, expected=1):
    count = 0
    for paragraph in document.paragraphs:
        if old in paragraph.text:
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    count += 1
                    break
            else:
                text = paragraph.text.replace(old, new)
                style = paragraph.style
                paragraph.clear()
                paragraph.style = style
                paragraph.add_run(text)
                count += 1
    if count != expected:
        raise RuntimeError(
            f"Expected {expected} replacements, found {count}: {old[:90]}"
        )


def replace_table_text(document, old, new, expected=1):
    count = 0
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old in paragraph.text:
                        text = paragraph.text.replace(old, new)
                        style = paragraph.style
                        paragraph.clear()
                        paragraph.style = style
                        paragraph.add_run(text)
                        count += 1
    if count != expected:
        raise RuntimeError(
            f"Expected {expected} table replacements, found {count}: {old}"
        )


def revise_main():
    document = Document(MAIN_SOURCE)

    replace_text(
        document,
        (
            "including their execution status and baseline-corrected host/device "
            "memory measurements"
        ),
        (
            "including their execution status and process-level host/device memory "
            "measurements"
        ),
    )
    replace_text(
        document,
        (
            "We report the pre-fit baseline, the absolute isolated-process peak, "
            "and the incremental fit-window peak above baseline. Absolute RSS "
            "remains a feasibility measure; the increment better isolates workflow "
            "memory."
        ),
        (
            "We report the pre-fit baseline, the absolute isolated-process peak, "
            "and the fit-window increase above baseline. The host-memory increase "
            "partly separates fitting and prediction from the already loaded data "
            "and runtime. The GPU increase is instead an incremental process-level "
            "device footprint: because the CUDA context was not preinitialized, it "
            "also includes context creation, CUDA-library state, allocator pools, "
            "and persistent device objects, and must not be interpreted as isolated "
            "PLS or SVD workspace memory."
        ),
    )
    replace_text(
        document,
        (
            "The memory audit separates data and runtime overhead from fitting "
            "allocations. Incremental host RSS ranged from 2 to 1037 MB and "
            "represented a median 41.3% of the absolute process peak. All measured "
            "CUDA pre-fit process baselines were 0 MB because CUDA contexts were "
            "initialized during fitting; incremental device memory ranged from 192 "
            "to 3414 MB. Thus absolute host RSS documents whether a workflow fits "
            "on the machine, whereas the baseline-corrected values expose "
            "algorithmic and prediction workspace costs."
        ),
        (
            "The host-memory audit partly separates loaded data and runtime overhead "
            "from allocations made during fitting and prediction. Incremental host "
            "RSS ranged from 2 to 1037 MB and represented a median 41.3% of the "
            "absolute process peak. All measured CUDA pre-fit process baselines were "
            "0 MB because CUDA contexts were initialized inside the measured call; "
            "the resulting 192-3414 MB increments therefore include context and "
            "library initialization, allocator pools, data and model storage, and "
            "algorithm workspaces. They quantify the end-to-end incremental "
            "process-level GPU footprint, not isolated workspace memory. Absolute "
            "host RSS and the process-level GPU peak document feasibility on the "
            "benchmark machine."
        ),
    )
    replace_text(
        document,
        (
            "H and G denote host and process-specific GPU memory; B/P/Δ denotes the "
            "immediately pre-fit baseline, absolute isolated-process peak, and "
            "incremental fit-window peak above baseline, respectively."
        ),
        (
            "H and G denote host and process-specific GPU memory; B/P/Δ denotes the "
            "immediately pre-fit baseline, absolute isolated-process peak, and "
            "fit-window increase above baseline, respectively. For G, Δ includes "
            "CUDA context and runtime-library initialization because the context was "
            "not preinitialized; it is a process-level device footprint rather than "
            "isolated algorithmic workspace."
        ),
    )
    replace_text(
        document,
        (
            "and used 664 MB and 3,432 MB of sampled device memory."
        ),
        (
            "and reached sampled process-level GPU peaks of 664 MB and 3,432 MB, "
            "respectively; these values include CUDA context and runtime state."
        ),
    )
    replace_text(
        document,
        "H/G are peak host RSS and sampled GPU memory in MB.",
        (
            "H/G are peak host RSS and sampled process-level GPU memory in MB; "
            "the GPU value is not workspace-only."
        ),
    )
    replace_text(
        document,
        "(D) Peak host RSS and sampled GPU memory.",
        (
            "(D) Peak host RSS and sampled process-level GPU memory, including "
            "CUDA context and runtime state."
        ),
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - clarified GPU memory scope"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)

    replace_text(
        document,
        (
            "Incremental host memory is the maximum of fit-window samples and "
            "post-fit/post-prediction snapshots minus the pre-fit baseline; "
            "incremental GPU memory is the PID-specific sampled peak minus its "
            "pre-fit baseline. Negative differences are truncated at zero. Thus "
            "absolute process RSS measures whole-workflow feasibility, whereas the "
            "baseline-corrected increment estimates fitting and prediction workspace."
        ),
        (
            "Incremental host memory is the maximum of fit-window samples and "
            "post-fit/post-prediction snapshots minus the pre-fit baseline; "
            "incremental process-level GPU memory is the PID-specific sampled peak "
            "minus its pre-fit baseline. Negative differences are truncated at zero. "
            "CUDA contexts were not preinitialized, so a zero pre-fit GPU baseline "
            "does not remove context creation, CUDA-library handles, allocator pools, "
            "persistent device objects, data/model storage, or workspaces created by "
            "the measured call. The GPU increment therefore measures the end-to-end "
            "device footprint triggered by the workflow and is not an isolated "
            "algorithmic-workspace measurement. Isolating workspace would require "
            "CUDA-API instrumentation after persistent runtime state had been "
            "initialized, which was not performed."
        ),
    )
    replace_text(
        document,
        "sampled 1,578 MB versus 664 MB of GPU memory.",
        (
            "sampled process-level GPU peaks of 1,578 MB versus 664 MB, including "
            "CUDA context and runtime state."
        ),
    )
    replace_text(
        document,
        (
            "Raw absolute time, host RSS, and sampled GPU memory for every workflow "
            "remain in Table S6b."
        ),
        (
            "Raw absolute time, host RSS, and sampled process-level GPU peaks for "
            "every workflow remain in Table S6b; GPU values include context and "
            "runtime initialization."
        ),
    )
    replace_text(
        document,
        (
            "RSS denotes peak host resident set size and GPU memory is the sampled "
            "process peak."
        ),
        (
            "RSS denotes peak host resident set size and GPU memory is the sampled "
            "process-level peak, including CUDA context and runtime state."
        ),
    )
    replace_text(
        document,
        "H and G are peak host RSS and sampled GPU memory in MB.",
        (
            "H and G are peak host RSS and sampled process-level GPU memory in MB; "
            "G includes CUDA context and runtime state."
        ),
        expected=4,
    )
    replace_text(
        document,
        (
            "H/G are peak host/GPU memory in MB."
        ),
        (
            "H/G are peak host/process-level GPU memory in MB; G includes CUDA "
            "context and runtime state."
        ),
    )
    replace_text(
        document,
        (
            "Time, input storage, host RSS, and GPU memory are shown as "
            "float64/float32."
        ),
        (
            "Time, input storage, host RSS, and sampled process-level GPU memory are "
            "shown as float64/float32; GPU values include CUDA context and runtime "
            "state."
        ),
    )
    replace_text(
        document,
        (
            "Because absolute RSS includes loaded data, libraries, and benchmark "
            "infrastructure, Table S32 now reports baseline, absolute peak, and "
            "incremental peak separately for host and GPU memory."
        ),
        (
            "Because absolute RSS includes loaded data, libraries, and benchmark "
            "infrastructure, Table S32 reports baseline, absolute peak, and increase "
            "above baseline separately for host and GPU memory. The GPU increase is "
            "process-level and includes CUDA context and runtime initialization; it "
            "does not isolate solver or PLS workspace."
        ),
    )
    replace_text(
        document,
        (
            "Table S32. Matched CPU/CUDA memory audit at each family-specific "
            "training-selected component count. B/P/Δ denotes pre-fit baseline, "
            "absolute isolated-process peak, and incremental fit-window peak above "
            "baseline, in MB. Host and GPU values are process-specific."
        ),
        (
            "Table S32. Matched CPU/CUDA memory audit at each family-specific "
            "training-selected component count. B/P/Δ denotes pre-fit baseline, "
            "absolute isolated-process peak, and fit-window increase above baseline, "
            "in MB. Host and GPU values are process-specific. GPU Δ includes CUDA "
            "context and runtime-library initialization and is not an isolated "
            "workspace measurement."
        ),
    )
    replace_text(
        document,
        (
            "Table S6a. Prespecified matched contrasts in the fixed-100-component "
            "float64 NMR analysis. A differing factor is interpreted only where all "
            "listed fixed factors are shared. Runtime and memory use medians from "
            "three isolated runs."
        ),
        (
            "Table S6a. Prespecified matched contrasts in the fixed-100-component "
            "float64 NMR analysis. A differing factor is interpreted only where all "
            "listed fixed factors are shared. Runtime and memory use medians from "
            "three isolated runs. CUDA values are sampled process-level peaks that "
            "include context and runtime initialization, not workspace-only memory."
        ),
    )
    replace_text(
        document,
        (
            "Time and memory are medians from three isolated runs; prediction "
            "agreement is correlation with the deposited reference."
        ),
        (
            "Time and memory are medians from three isolated runs; CUDA memory is "
            "the sampled process-level peak including context and runtime state, and "
            "prediction agreement is correlation with the deposited reference."
        ),
    )
    replace_table_text(
        document,
        "GPU memory (MB)",
        "GPU peak (MB)",
    )

    document.core_properties.title = (
        "fastPLS CMPB supplementary material - clarified GPU memory scope"
    )
    document.save(SUPP_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
