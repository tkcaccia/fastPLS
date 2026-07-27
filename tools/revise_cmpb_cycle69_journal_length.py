#!/usr/bin/env python3
"""Compress the cycle-68 manuscript to the CMPB original-article word limits."""

from pathlib import Path
import re

from docx import Document

from revise_cmpb_cycle67_consolidate_evidence import (
    find_paragraph,
    normalize_submission_terminology,
    remove_paragraph,
    replace_paragraph,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = (
    ROOT
    / "artifacts"
    / "CMPB_rewrite_20260726_cycle68"
    / "fastPLS_CMPB_main_cycle68_0.99.6_20260726.docx"
)
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle69"
OUTPUT = OUTPUT_DIR / "fastPLS_CMPB_main_cycle69_0.99.6_20260726.docx"


def remove_by_prefix(document, prefixes):
    for prefix in prefixes:
        remove_paragraph(find_paragraph(document, prefix))


def word_count(text):
    return len(re.findall(r"\b[\w'-]+\b", text))


def revise_abstract(document):
    replace_paragraph(
        document,
        "Background and objective:",
        (
            "Background and objective: Partial least squares (PLS) combines "
            "supervised dimension reduction with prediction, but conventional "
            "implementations become restrictive for long component paths, large "
            "multivariate responses, and repeated validation. We developed fastPLS "
            "to extend established PLS models to these computational regimes."
        ),
    )
    replace_paragraph(
        document,
        "Methods:",
        (
            "Methods: fastPLS provides PLS-SVD, SIMPLS, orthogonal PLS, and kernel "
            "PLS through compiled CPU, NVIDIA CUDA, and Apple Metal routes. Its "
            "SIMPLS implementation reuses sequential deflation, coefficient, and "
            "prediction quantities; supports implicit cross-covariance products and "
            "compact prediction; and separates the estimator from deterministic "
            "IRLBA and approximate randomized SVD (rSVD). Deterministic float64 "
            "implementations were compared with independent references. Runtime, "
            "host/device memory, prediction agreement, and predictive performance "
            "were evaluated across biomedical datasets; NMR and ImageNet embeddings "
            "provided extreme-response and million-sample stress tests."
        ),
    )
    replace_paragraph(
        document,
        "Results:",
        (
            "Results: Deterministic SIMPLS met all prespecified tolerances in 117 "
            "component-level comparisons with de Jong SIMPLS. OPLS and kernel PLS "
            "met all endpoints in 66 setting/task comparisons and 1,540 "
            "fold-component fits. In matched single-CPU comparisons, fastPLS "
            "SIMPLS was faster than pls::simpls.fit on seven of nine datasets with "
            "identical accuracy. CUDA was faster in seven of 28 numerically "
            "concordant non-NMR routes, with speed-up up to 8.90-fold; no concordant "
            "Metal route was faster. On NMR, selected CUDA SIMPLS achieved RMSD "
            "0.000759 and reduced matched CPU time 5.94-fold. In the exploratory "
            "ImageNet analysis, CUDA SIMPLS-LDA processed 1,000,000 training and "
            "281,167 held-out embeddings and reached top-1 accuracy 0.8093 at 1,000 "
            "components. Float32 reduced stored input size but was not uniformly "
            "faster or less memory intensive."
        ),
    )
    replace_paragraph(
        document,
        "Conclusions:",
        (
            "Conclusions: fastPLS extends the feasible scale of established PLS "
            "workflows while exposing numerical, precision, and hardware limits. "
            "Compiled float64 CPU execution remains the confirmatory reference; "
            "rSVD and accelerator routes require route-specific numerical checks."
        ),
    )


def revise_introduction(document):
    replace_paragraph(
        document,
        "Partial least squares (PLS) regression combines",
        (
            "Partial least squares (PLS) combines supervised dimension reduction "
            "with prediction and is widely used when biomedical predictors are "
            "correlated, high dimensional, or accompanied by multivariate responses "
            "[1-3]. Applications include cancer metabolomics and prediction of "
            "multiple NMR spectra from one acquisition [4-7]."
        ),
    )
    replace_paragraph(
        document,
        "The computational demands of these applications",
        (
            "Increasing spectral resolution, single-cell sample counts, "
            "foundation-model embeddings, and repeated validation make PLS "
            "computationally demanding. KODAMA, for example, repeatedly fits "
            "cross-validated PLS discriminant models [8,9]. UNI and Prov-GigaPath "
            "illustrate the related post-extraction matrix regime in computational "
            "pathology [31,32]. The ImageNet/DINOv2 stress test assessed reproducible "
            "scalability, not biomedical validity [28,29]."
        ),
    )
    remove_by_prefix(
        document,
        ["This matrix regime is increasingly relevant after foundation-model"],
    )
    replace_paragraph(
        document,
        "PLS formulations address different objectives",
        (
            "PLS-SVD extracts directions from predictor-response "
            "cross-covariance [10], SIMPLS constructs sequential components without "
            "explicitly deflating the predictor matrix [11], OPLS removes structured "
            "predictor variation orthogonal to the response [12], and kernel PLS "
            "models nonlinear relations [13,14]. Existing R implementations do not "
            "jointly provide memory-aware paths, compiled validation, float32, and "
            "CPU and accelerator execution."
        ),
    )
    replace_paragraph(
        document,
        "Here we present fastPLS",
        (
            "We present fastPLS, which reorganizes SIMPLS execution to reuse "
            "sequential quantities and requested component prefixes while retaining "
            "the established estimator. A common interface separates PLS family, "
            "low-rank solver, prediction head, and compiled CPU, NVIDIA CUDA, or "
            "Apple Metal backend. The GPL-3 R package uses reusable components from "
            "the MIT-licensed kodama-cpp project."
        ),
    )


def revise_methods(document):
    replace_paragraph(
        document,
        "The public pls() function selects",
        (
            "The public pls() interface selects PLS family, component count, solver, "
            "backend, and, for classification, argmax or latent-space linear "
            "discriminant analysis (LDA). pls.single.cv() selects settings within "
            "one cross-validation layer; pls.double.cv() uses nested "
            "cross-validation. Requested estimators are never silently substituted."
        ),
    )
    replace_paragraph(
        document,
        "Notation is fixed throughout:",
        (
            "Let a index a component. A denotes the maximum retained component count, "
            "C the requested component-count set, and K denotes the number of "
            "cross-validation folds. Lowercase k is reserved for retrieval cutoffs."
        ),
    )
    replace_paragraph(
        document,
        "The software is organized as four separate layers",
        (
            "The architecture separates preprocessing, the PLS estimator, direction "
            "extraction by IRLBA or rSVD, and CPU/CUDA/Metal execution (Figure 1). "
            "Benchmark rows recorded requested and executed estimators and rejected "
            "mismatches."
        ),
    )
    replace_paragraph(
        document,
        "Figure 1. fastPLS architecture",
        (
            "Figure 1. fastPLS architecture separating response representation, PLS "
            "estimator, low-rank solver, backend, and prediction head. Accelerator "
            "routes may be native or explicitly reported as hybrid."
        ),
    )
    replace_paragraph(
        document,
        "The R package includes bundled IRLBA code",
        (
            "CPU routines use compiled C/C++ and R-linked BLAS/LAPACK; CUDA uses "
            "NVIDIA CUDA/cuBLAS and Metal uses Apple Metal Performance Shaders. The "
            "external comparison used one effective BLAS thread, so no multicore "
            "speed-up is claimed. Versions, build flags, residency, and thread "
            "settings are in Supplementary Tables S1 and S4a."
        ),
    )
    replace_paragraph(
        document,
        "Predictors are centred and optionally scaled",
        (
            "Predictors were centred and optionally scaled using training "
            "statistics; numeric responses were centred. Classification used a "
            "multivariate PLS-DA response, with class-wise products replacing dense "
            "one-hot matrices on large-class routes."
        ),
    )
    replace_paragraph(
        document,
        "The requested truncated decomposition is computed once",
        (
            "PLS-SVD decomposes the centred cross-covariance once at the largest "
            "valid requested rank and reuses prefixes. Its component count is "
            "bounded by cross-covariance rank and, for C centred classes, by C-1."
        ),
    )
    replace_paragraph(
        document,
        "The SIMPLS estimator follows de Jong",
        (
            "SIMPLS follows de Jong's sequential score, loading, orthogonalization, "
            "and rank-one deflation equations [11]. fastPLS retains deflation "
            "products, latent quantities, coefficients, and predictions "
            "incrementally, so several requested component counts are snapshots of "
            "one maximal path rather than independent refits. Each direction is "
            "nevertheless recomputed from the current deflated state by IRLBA or a "
            "fresh rSVD sketch."
        ),
    )
    replace_paragraph(
        document,
        "For tall matrices with moderate",
        (
            "A shape-aware route can cache X-transpose-X for score normalization and "
            "loading calculations. Compact latent factors and blocked prediction "
            "avoid retaining a dense coefficient or prediction path when it would "
            "dominate memory."
        ),
    )
    replace_paragraph(
        document,
        "Algorithm 1 maps these execution optimizations",
        "Algorithm 1 summarizes the component path; detailed de Jong mapping is in the Supplement.",
    )
    replace_paragraph(
        document,
        "Algorithm 1. Accelerated SIMPLS component path",
        (
            "Algorithm 1. Accelerated SIMPLS path. Direction extraction uses "
            "deterministic IRLBA or approximate rSVD; score construction, "
            "orthogonalization, and deflation follow de Jong [11]."
        ),
    )
    replace_paragraph(
        document,
        "OPLS first estimates orthogonal scores",
        (
            "OPLS applies an orthogonal predictor filter before the SIMPLS core "
            "[12]. Linear kernel PLS dispatches to the linear route; RBF and "
            "polynomial kernels form an n-by-n Gram matrix [13,14]. One float64 Gram "
            "matrix requires 8n-squared bytes before copies and workspaces, limiting "
            "nonlinear kernels to moderate n. Validation reached n=180; the larger "
            "memory-planning estimates in the Supplement are not empirical limits."
        ),
    )
    replace_paragraph(
        document,
        "Double-precision CPU fitting supports",
        (
            "Float64 CPU fitting supports bundled deterministic IRLBA [15] and "
            "approximate rSVD [16]. rSVD uses a Gaussian range sketch, "
            "orthonormalization, optional power iterations, and a reduced "
            "decomposition. CUDA and Metal support rSVD; important approximate "
            "results were audited by prediction, subspace, label, and metric "
            "criteria (Supplementary Table S8)."
        ),
    )
    replace_paragraph(
        document,
        "All products use the same training-derived centring",
        (
            "When explicit cross-covariance is large, identical centred/scaled "
            "operators are evaluated as X-transpose-(YV) and Y-transpose-(XU). "
            "Class-wise sums avoid dense indicator responses. Blocking changes "
            "storage, not the global operator."
        ),
    )
    replace_paragraph(
        document,
        "Prediction can retain compact latent factors",
        (
            "Compact latent prediction processes large test matrices in row blocks "
            "and releases temporary scores after each block."
        ),
    )
    replace_paragraph(
        document,
        "Standard R numeric matrices store each value",
        (
            "Standard R matrices use eight-byte float64 values; float-package inputs "
            "use four-byte float32 values. Float64 is the reference. Float32 support "
            "is route specific and classified as validated, experimental, hybrid, "
            "unavailable, or failed (Supplementary Table S9)."
        ),
    )
    replace_paragraph(
        document,
        "Regression returns continuous predictions",
        (
            "Regression returns continuous predictions. Classification uses argmax "
            "PLS-DA or LDA fitted to PLS scores. LDA uses pooled within-class "
            "covariance, Cholesky solves, class priors, and deterministic "
            "trace-scaled regularization only when factorization fails [17,18]."
        ),
    )
    remove_by_prefix(
        document,
        ["For each class, the coefficient vector is obtained"],
    )
    replace_paragraph(
        document,
        "Cross-validation keeps fold construction",
        (
            "Fold construction, fitting, prediction, and metric accumulation remain "
            "compiled where supported; grouped observations can be constrained to "
            "one fold. Hybrid OPLS, nonlinear-kernel, and Metal paths are identified "
            "explicitly."
        ),
    )
    replace_paragraph(
        document,
        "The main biomedical benchmark included twelve tasks",
        (
            "Twelve tasks covered metabolomics, NMR, CITE-seq, tissue and cancer "
            "omics, single-cell transcriptomics, drug response, and CIFAR-100 "
            "[7,20-27,30]. Methods used identical stored splits and training-selected "
            "component grids. CIFAR-100 followed its documented 50,000/10,000 split "
            "[30]. Classification used accuracy and Wilson intervals; "
            "multivariate regression used RMSD, Q2, and held-out bootstrap "
            "intervals. Runtime included fitting and prediction. Isolated-process "
            "baseline and peak host RSS and process-specific GPU memory were "
            "recorded; GPU increments include runtime context, not only workspace."
        ),
    )
    replace_paragraph(
        document,
        "The ImageNet stress test fitted",
        (
            "A separate exploratory ImageNet/DINOv2 analysis used a fixed "
            "1,000,000/281,167 development split of 1,024-dimensional embeddings. "
            "The split was noncanonical and previously informed component choices. "
            "SIMPLS argmax/LDA classification and raw/PCA/PLS FAISS retrieval were "
            "therefore treated as separate exploratory experiments."
        ),
    )
    replace_paragraph(
        document,
        "Within each dataset, methods used identical",
        (
            "Five-fold training-only selection was performed separately by PLS "
            "family; boundary selections were described as best within the grid. "
            "Accelerator speed-up was interpreted only when absolute predictive "
            "metric difference was at most 0.005 and paired prediction agreement at "
            "least 0.995. CUDA timing covered data transfer, execution, "
            "synchronization, and returned model or predictions."
        ),
    )
    replace_paragraph(
        document,
        "Predictive uncertainty beyond a single outer split",
        (
            "Estimator preservation, OPLS/kernel reliability, rSVD approximation, "
            "implementation ablations, external software, and repeated outer "
            "partitions were prespecified as separate analyses. Deterministic "
            "float64 IRLBA supported equivalence claims; rSVD supported only "
            "approximate workflow comparisons. Detailed designs, thresholds, data "
            "provenance, and complete results are in Supplementary Tables S3-S15."
        ),
    )
    remove_by_prefix(
        document,
        [
            "Prepared real-data matrices were not redistributed",
            "Supplementary Table S6 provides a claim-to-evidence map",
            "For CBMC CITE-seq specifically",
            "OPLS and kernel settings were prespecified explicitly",
            "External comparisons use independent R implementations",
            "Three matched contrasts were used",
            "Estimator preservation and randomized approximation were evaluated",
            "A second prespecified estimator-validation study",
            "A controlled implementation ablation",
            "A separate prespecified family-speed experiment",
            "The ImageNet representation experiment used",
        ],
    )


def revise_results(document):
    replace_paragraph(
        document,
        "We first evaluated whether the computational implementation",
        (
            "We first compared fastPLS SIMPLS with independent R software, then "
            "examined solver and backend effects, and finally evaluated NMR and "
            "ImageNet stress tests. Detailed numerical audits and complete paths are "
            "reported in Supplementary Tables S6-S15."
        ),
    )
    replace_paragraph(
        document,
        "Formal deterministic reliability testing",
        (
            "Deterministic SIMPLS met all prespecified tolerances in 117 "
            "component-level comparisons with de Jong SIMPLS. Across three OPLS and "
            "eight kernel settings, all 66 setting/task comparisons and all 1,540 "
            "fold-component fits met the specified criteria. These results establish "
            "reliability only for the evaluated deterministic float64 CPU settings "
            "(Supplementary Table S7)."
        ),
    )
    replace_paragraph(
        document,
        "Repeated outer partitions quantified",
        (
            "Repeated outer partitions showed that predictive variation exceeded "
            "timing variation and that several component selections remained "
            "boundary or rank constrained (Supplementary Table S14)."
        ),
    )
    replace_paragraph(
        document,
        "The primary software comparison used double-precision inputs",
        (
            "The float64 single-CPU comparison attempted 126 external-package runs: "
            "110 completed, 12 were package limitations, two timed out, and two "
            "errored. In the matched subset, fastPLS and pls::simpls.fit had "
            "identical accuracy on nine datasets; fastPLS was faster on seven, "
            "including 4.23-fold on CIFAR-100, 8.65-fold on Retina, and 8.90-fold on "
            "Tabula Muris. Matched accuracies were 0.8739 (8,739/10,000; Wilson 95% "
            "CI 0.8672-0.8803), 0.9678 (21,684/22,406; 0.9654-0.9700), and 0.8006 "
            "(40,077/50,059; 0.7971-0.8041), respectively. FastPLS LDA matched "
            "plsgenomics accuracy on eight datasets and was faster on six (Figure 2; "
            "Supplementary Table S10)."
        ),
    )
    replace_paragraph(
        document,
        "Peak host memory in Figure 2",
        (
            "Absolute process RSS was not uniformly reduced on small tasks. On "
            "CIFAR-100 it was 1.69 GB for fastPLS versus 13.09 GB for the matched "
            "reference, and on Tabula Muris 0.68 versus 2.20 GB."
        ),
    )
    replace_paragraph(
        document,
        "Figure 2. SIMPLS classification workflows",
        (
            "Figure 2. Matched float64 single-CPU SIMPLS workflows. Panels report "
            "accuracy, fitting-plus-prediction time, and absolute process RSS. NE "
            "denotes unavailable or incomplete runs."
        ),
    )
    replace_paragraph(
        document,
        "Backend performance was interpreted only after",
        (
            "Of 44 non-NMR CPU-CUDA pairs, 28 met the numerical criteria and CUDA "
            "was faster in seven, with eligible speed-up up to 8.90-fold. Sixteen "
            "discordant routes were excluded from speed claims. Six of 12 CPU-Metal "
            "pairs were concordant and none was faster with Metal. Accelerator "
            "benefit was therefore route and shape specific (Figure 3; "
            "Supplementary Table S11)."
        ),
    )
    replace_paragraph(
        document,
        "The benefit of rSVD over IRLBA",
        (
            "CPU SIMPLS rSVD was only 1.00-1.45-fold faster than IRLBA across nine "
            "classification tasks; MetRef accuracy differed by 4.0 percentage "
            "points. On NMR at 100 components, rSVD reduced CPU time from 436.3 to "
            "19.6 s. Thus IRLBA remains the deterministic reference and rSVD an "
            "audited approximate option (Supplementary Table S8)."
        ),
    )
    replace_paragraph(
        document,
        "Single precision was route-dependent",
        (
            "Float32 approximately halved stored inputs on MetRef and PRISM but did "
            "not uniformly improve runtime, incremental memory, or agreement "
            "(Supplementary Table S9)."
        ),
    )
    replace_paragraph(
        document,
        "Figure 3. Numerically qualified backend",
        (
            "Figure 3. Numerically qualified backend and solver comparisons. "
            "Discordant routes are excluded from speed-up summaries."
        ),
    )
    replace_paragraph(
        document,
        "The NMR task contained 1,200 training spectra",
        (
            "NMR comprised 1,200 training and 321 held-out spectra, with 13,000 "
            "predictors and 28,355 responses. Predictor columns between 4.6 and "
            "4.8 ppm were zeroed in training and test data; responses were unmasked. "
            "Training-only one-standard-error selection retained five PLS-SVD and "
            "50 SIMPLS components."
        ),
    )
    replace_paragraph(
        document,
        "At the selected settings",
        (
            "Selected CUDA PLS-SVD and SIMPLS achieved RMSD 0.001043 (95% bootstrap "
            "interval 0.001000-0.001085) and 0.000759 (0.000665-0.000884), "
            "respectively. SIMPLS also had lower median per-spectrum, response-wise, "
            "and high-intensity errors (Figure 4)."
        ),
    )
    replace_paragraph(
        document,
        "The implementation analysis then compared CPU with CUDA",
        (
            "Holding family, split, rSVD, float64, and component count fixed, CUDA "
            "reduced PLS-SVD time from 2.301 to 0.648 s and SIMPLS time from 10.525 "
            "to 1.773 s, with prediction correlations 1.000000 and 0.999981. CUDA "
            "device increments were 590 and 3,414 MB, including runtime context."
        ),
    )
    replace_paragraph(
        document,
        "The deposited Nature Communications workflow",
        (
            "The deposited 165-component historical workflow required 447.6 s and "
            "achieved RMSD 0.000710. Because estimator, solver, protocol, and "
            "hardware differ, this is contextual rather than an implementation-only "
            "comparison (Supplementary Table S12)."
        ),
    )
    replace_paragraph(
        document,
        "Figure 4. Separated NMR predictive",
        (
            "Figure 4. NMR predictive and implementation analyses: training-only "
            "component selection, held-out per-spectrum and response-wise errors, "
            "intensity-stratified RMSD, and matched CPU/CUDA resources. All 321 "
            "held-out spectra and all response coordinates contribute."
        ),
    )
    replace_paragraph(
        document,
        "ImageNet/DINOv2 was used as a computational stress test",
        (
            "In exploratory ImageNet experiment 1, CUDA SIMPLS/rSVD used 1,000,000 "
            "training and 281,167 held-out embeddings. LDA improved top-1 accuracy "
            "from 0.6270 to 0.7793 at 100 components and from 0.7995 to 0.8093 at "
            "1,000 components. At 1,000 components LDA classified "
            "227,535/281,167 observations correctly (Wilson 95% CI "
            "0.8078-0.8107); CUDA reduced total LDA time from 2,199.7 to 316.1 s "
            "(Figure 5)."
        ),
    )
    replace_paragraph(
        document,
        "Experiment 2, representation retrieval",
        (
            "Separate FAISS retrieval compared raw, PCA, and PLS representations. "
            "Raw DINOv2 top-1/top-5 accuracy was 0.6556/0.9392; 200-component PLS "
            "gave 0.6516/0.9397 with 5.12-fold compression and approximately "
            "fourfold lower projection-plus-query time. These single-run, "
            "noncanonical-holdout results show an exploratory compression trade-off, "
            "not an accuracy improvement (Supplementary Table S13)."
        ),
    )
    replace_paragraph(
        document,
        "Figure 5. ImageNet experiment 1",
        (
            "Figure 5. Exploratory ImageNet SIMPLS classification across 100-1,000 "
            "components, reporting accuracy, time, host RSS, and GPU memory for "
            "argmax and LDA."
        ),
    )


def revise_discussion_and_endmatter(document):
    replace_paragraph(
        document,
        "fastPLS accelerates established PLS estimators",
        (
            "fastPLS reorganizes established PLS computation rather than defining a "
            "new estimator. Deterministic SIMPLS, OPLS, and kernel-PLS validation "
            "supports the evaluated float64 CPU routes; rSVD remains approximate. "
            "The external comparison indicates that compiled sequential reuse can "
            "reduce time and memory without changing matched accuracy."
        ),
    )
    replace_paragraph(
        document,
        "Computational benefit was conditional",
        (
            "Benefits depend on matrix shape. CPU execution is preferable for many "
            "small tasks, while numerically concordant CUDA routes benefit selected "
            "large dense or extreme-response problems. Metal demonstrated "
            "portability but no speed advantage here. Float32 reduces representation "
            "size but requires route-specific validation."
        ),
    )
    replace_paragraph(
        document,
        "The NMR study separates predictive model selection",
        (
            "NMR demonstrates practical multivariate-response prediction and a "
            "matched CUDA benefit; ImageNet demonstrates million-sample feasibility "
            "and supervised compression, not biomedical validity. Principal "
            "limitations are finite component grids, conditional uncertainty, "
            "noncanonical ImageNet provenance, quadratic nonlinear kernels, and "
            "discordant accelerator routes."
        ),
    )
    remove_by_prefix(document, ["The remaining limitations are explicit"])
    replace_paragraph(
        document,
        "fastPLS combines a shape-aware accelerated sequential SIMPLS",
        (
            "fastPLS makes established PLS estimators and validation feasible across "
            "larger biomedical matrix regimes through compiled sequential reuse, "
            "memory-aware prediction, and qualified CPU or accelerator routes. "
            "Float64 CPU remains the confirmatory reference."
        ),
    )
    replace_paragraph(
        document,
        "This computational software study used public",
        (
            "This software study used public, simulated, or previously collected "
            "de-identified data; source-study ethics are reported in the cited "
            "publications."
        ),
    )
    replace_paragraph(
        document,
        "The fastPLS R package, benchmark workflows",
        (
            "Code and benchmark outputs are available at "
            "https://github.com/tkcaccia/fastPLS; reusable components are at "
            "https://github.com/tkcaccia/kodama-cpp. The reviewed snapshot is "
            "fastPLS 0.99.6, commit "
            "6e50bd318f20289101f6b723953830aefa8b95d6. A matching archival release "
            "tag is required before submission. Analysis-specific provenance is in "
            "Supplementary Table S15."
        ),
    )


def audit(document):
    abstract = " ".join(
        find_paragraph(document, prefix).text
        for prefix in (
            "Background and objective:",
            "Methods:",
            "Results:",
            "Conclusions:",
        )
    )
    introduction = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == "1. Introduction"
    )
    references = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == "References"
    )
    main_text = " ".join(
        paragraph.text for paragraph in document.paragraphs[introduction:references]
    )
    abstract_words = word_count(abstract)
    main_words = word_count(main_text)
    if abstract_words > 350:
        raise RuntimeError(f"Abstract has {abstract_words} words")
    if main_words > 3500:
        raise RuntimeError(f"Main text has {main_words} words")
    print(f"Abstract words: {abstract_words}")
    print(f"Main-text words (Introduction through availability): {main_words}")


def main():
    document = Document(SOURCE)
    revise_abstract(document)
    revise_introduction(document)
    revise_methods(document)
    revise_results(document)
    revise_discussion_and_endmatter(document)
    normalize_submission_terminology(document)
    audit(document)
    document.core_properties.title = "fastPLS CMPB manuscript - journal-length revision"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
