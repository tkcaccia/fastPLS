from copy import deepcopy
from pathlib import Path
from shutil import copy2
from posixpath import dirname, join, normpath
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from lxml import etree

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.table import _Row


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle98"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle99"
OUTPUT.mkdir(parents=True, exist_ok=True)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle98_0.99.25_20260825.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle98_0.99.25_20260825.docx"
MAIN_OUTPUT = OUTPUT / "fastPLS_CMPB_main_cycle99_0.99.25_20260825.docx"
SUPP_OUTPUT = OUTPUT / "fastPLS_CMPB_supplement_cycle99_0.99.25_20260825.docx"

EXTRACTED = OUTPUT / "source_images"
EXTRACTED.mkdir(exist_ok=True)
FIGURE_1 = EXTRACTED / "figure1_architecture.png"

FIGURES_MAIN = (
    ("Figure 1.", FIGURE_1, 6.35),
    (
        "Figure 2.",
        ROOT
        / "benchmark_results"
        / "external_simpls_timing_publication_20260825"
        / "external_simpls_timing_profiles.png",
        6.45,
    ),
    (
        "Figure 3.",
        ROOT
        / "benchmark_results"
        / "manuscript_revision_cycle62_20260726"
        / "accelerator_concordance_speedups.png",
        6.45,
    ),
    (
        "Figure 4.",
        ROOT
        / "benchmark_results"
        / "manuscript_revision_cycle80_20260727"
        / "nmr_qualified"
        / "nmr_qualified_main_figure.png",
        5.35,
    ),
    (
        "Figure 5.",
        ROOT
        / "benchmark_results"
        / "manuscript_revision_cycle80_20260727"
        / "imagenet_float32_simpls_lda_path"
        / "imagenet_float32_simpls_lda_main_figure.png",
        6.45,
    ),
)

FIGURES_SUPP = (
    (
        "Figure S1.",
        ROOT
        / "benchmark_results"
        / "controlled_scaling_publication_cuda_20260825"
        / "controlled_scaling_overview.png",
        6.45,
    ),
    (
        "Figure S2.",
        ROOT
        / "benchmark_results"
        / "manuscript_revision_cycle62_20260726"
        / "rsvd_workflow_speed_supp.png",
        6.45,
    ),
)


def extract_figure_1():
    source = (
        ROOT
        / "artifacts"
        / "CMPB_rewrite_20260808_cycle87"
        / "fastPLS_CMPB_main_cycle87_0.99.10_20260808.docx"
    )
    with ZipFile(source) as archive:
        FIGURE_1.write_bytes(archive.read("word/media/rId23.png"))


def paragraph_by_prefix(document, prefix):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}; found {len(matches)}")
    return matches[0]


def insert_embedded_figure(document, caption_prefix, image_path, width):
    if not image_path.is_file():
        raise FileNotFoundError(image_path)
    caption = paragraph_by_prefix(document, caption_prefix)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 3
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    caption._p.addprevious(paragraph._p)
    caption.paragraph_format.keep_together = True
    caption.paragraph_format.keep_with_next = False


def table_with_value(document, value):
    matches = []
    for table in document.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == value:
                matches.append((table, row))
    if len(matches) != 1:
        raise RuntimeError(f"Expected one table row beginning {value!r}; found {len(matches)}")
    return matches[0]


def move_a21_to_provenance(document):
    wrong_table, wrong_row = table_with_value(document, "A21")
    wrong_table._tbl.remove(wrong_row._tr)

    provenance, row_a20 = table_with_value(document, "A20")
    provenance_a22, row_a22 = table_with_value(document, "A22")
    if provenance._tbl is not provenance_a22._tbl:
        raise RuntimeError("A20 and A22 are not in the same provenance table")

    new_tr = deepcopy(row_a20._tr)
    row_a22._tr.addprevious(new_tr)
    new_row = _Row(new_tr, provenance)
    values = (
        "A21",
        "Controlled one-factor SIMPLS scaling",
        "benchmark_results/controlled_scaling_publication_cuda_20260825",
        "0.99.25",
        "exact source archive; isolated CPU/CUDA processes; Metal diagnostic on reviewed Mac build",
        "scripts/run_controlled_scaling.sh; benchmark/controlled_scaling/*.R",
        "74e134ef22d5",
    )
    if len(new_row.cells) != len(values):
        raise RuntimeError(
            f"Provenance table has {len(new_row.cells)} columns; expected {len(values)}"
        )
    for cell, value in zip(new_row.cells, values):
        text_nodes = cell._tc.xpath(".//w:t")
        if not text_nodes:
            raise RuntimeError("Expected formatted text in copied provenance cell")
        text_nodes[0].text = value
        for node in text_nodes[1:]:
            node.text = ""


def write_standalone_figures():
    destination = OUTPUT / "figures"
    destination.mkdir(exist_ok=True)
    for index, (_, path, _) in enumerate(FIGURES_MAIN, start=1):
        copy2(path, destination / f"Figure_{index}.png")
    for index, (_, path, _) in enumerate(FIGURES_SUPP, start=1):
        copy2(path, destination / f"Figure_S{index}.png")
    (destination / "README.md").write_text(
        "# Review figures\n\n"
        "These seven PNG files are the exact raster assets embedded in the Cycle 99 "
        "manuscript and supplementary DOCX files. Main Figures 1-5 and Supplementary "
        "Figures S1-S2 are numbered consistently with their document captions.\n",
        encoding="utf-8",
    )


def verify_embedded_images(path, expected):
    with ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml").decode("utf-8")
        embedded = document_xml.count("r:embed=")
        if embedded < expected:
            raise RuntimeError(f"{path.name}: found {embedded} embedded images; expected {expected}")


def strip_orphaned_document_images(path):
    relationship_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
    office_rel_ns = (
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    )
    with ZipFile(path) as archive:
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        active_ids = set(
            document_xml.xpath(
                ".//@r:embed",
                namespaces={"r": office_rel_ns},
            )
        )
        document_rels_name = "word/_rels/document.xml.rels"
        document_rels = etree.fromstring(archive.read(document_rels_name))
        for relationship in list(document_rels):
            if (
                relationship.get("Type", "").endswith("/image")
                and relationship.get("Id") not in active_ids
            ):
                document_rels.remove(relationship)

        updated_relationships = {
            document_rels_name: etree.tostring(
                document_rels, xml_declaration=True, encoding="UTF-8", standalone=True
            )
        }
        used_media = set()
        for member in archive.namelist():
            if not member.endswith(".rels"):
                continue
            xml = updated_relationships.get(member, archive.read(member))
            relationships = etree.fromstring(xml)
            source_dir = dirname(member.replace("/_rels/", "/"))
            for relationship in relationships.findall(f"{{{relationship_ns}}}Relationship"):
                if relationship.get("Type", "").endswith("/image"):
                    used_media.add(normpath(join(source_dir, relationship.get("Target"))))

        with NamedTemporaryFile(suffix=".docx", delete=False) as handle:
            temporary = Path(handle.name)
        try:
            with ZipFile(temporary, "w", compression=ZIP_DEFLATED) as output:
                for item in archive.infolist():
                    if item.filename.startswith("word/media/") and item.filename not in used_media:
                        continue
                    output.writestr(
                        item,
                        updated_relationships.get(item.filename, archive.read(item.filename)),
                    )
            temporary.replace(path)
        finally:
            temporary.unlink(missing_ok=True)


def main():
    extract_figure_1()
    write_standalone_figures()

    manuscript = Document(MAIN_SOURCE)
    for prefix, path, width in FIGURES_MAIN:
        insert_embedded_figure(manuscript, prefix, path, width)
    manuscript.save(MAIN_OUTPUT)
    strip_orphaned_document_images(MAIN_OUTPUT)
    verify_embedded_images(MAIN_OUTPUT, len(FIGURES_MAIN))

    supplement = Document(SUPP_SOURCE)
    move_a21_to_provenance(supplement)
    for prefix, path, width in FIGURES_SUPP:
        insert_embedded_figure(supplement, prefix, path, width)
    supplement.save(SUPP_OUTPUT)
    strip_orphaned_document_images(SUPP_OUTPUT)
    verify_embedded_images(SUPP_OUTPUT, len(FIGURES_SUPP))

    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)


if __name__ == "__main__":
    main()
