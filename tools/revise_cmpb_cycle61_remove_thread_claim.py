#!/usr/bin/env python3
"""Remove unsupported CPU thread-scaling claims and the Figure 3 thread panel."""

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle60"
DEST = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle61"
FIGURE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle61_20260726"
    / "internal_backend_solver_speedups_no_threads.png"
)

MAIN_SRC = SRC / "fastPLS_CMPB_main_cycle60_0.99.6_20260726.docx"
SUPP_SRC = SRC / "fastPLS_CMPB_supplement_cycle60_0.99.6_20260726.docx"
MAIN_OUT = DEST / "fastPLS_CMPB_main_cycle61_0.99.6_20260726.docx"
SUPP_OUT = DEST / "fastPLS_CMPB_supplement_cycle61_0.99.6_20260726.docx"


def replace_paragraph(doc, startswith, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(startswith):
            paragraph.text = text
            return paragraph
    raise RuntimeError(f"Paragraph not found: {startswith}")


def replace_figure_before_caption(doc, caption_prefix, image_path):
    caption = next(
        p for p in doc.paragraphs if p.text.startswith(caption_prefix)
    )
    node = caption._p.getprevious()
    while node is not None:
        drawings = node.xpath(".//w:drawing")
        if drawings:
            node.getparent().remove(node)
            break
        node = node.getprevious()
    else:
        raise RuntimeError(f"Figure preceding {caption_prefix} not found")

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(6.35))
    caption._p.addprevious(paragraph._p)


def revise_main():
    doc = Document(MAIN_SRC)
    replace_paragraph(
        doc,
        "We first evaluated whether",
        "We first evaluated whether the computational implementation of fastPLS "
        "improved the practical use of SIMPLS relative to independent R software "
        "under a controlled single-CPU setting. We then compared the compiled CPU "
        "route with the CUDA and Metal backends and examined when randomized SVD "
        "(rSVD) was preferable to the deterministic IRLBA route. The final analyses "
        "focus on two settings that extend the scale of PLS modelling: multivariate "
        "NMR prediction and million-sample ImageNet embeddings. Formal "
        "estimator-preservation tests, numerical audits, implementation ablations, "
        "component paths, and complete backend tables are reported in the "
        "Supplementary Material.",
    )
    replace_paragraph(
        doc,
        "Backend acceleration depended",
        "Backend acceleration depended on the amount and shape of the linear "
        "algebra rather than sample size alone (Figure 3A-B). Across 44 matched "
        "non-NMR CPU-CUDA comparisons, CUDA was faster in 13. The largest gains "
        "occurred on CIFAR-100: 8.90-fold for PLS-SVD, 14.50-fold for SIMPLS, "
        "4.39-fold for OPLS, and 14.55-fold for kernel PLS, with accuracy "
        "differences no larger than 0.11 percentage points. CUDA also accelerated "
        "several CBMC CITE-seq and GTEx v8 configurations, whereas compiled CPU "
        "execution remained preferable for small matrices because device setup and "
        "transfer costs were not amortized. Metal was faster than the matched Apple "
        "CPU route only for the two CIFAR-100 configurations tested, and its current "
        "metric differences require the numerical qualification reported in the "
        "Supplementary Material. These comparisons demonstrate compiled CPU and "
        "accelerator execution; they do not quantify multicore CPU scaling.",
    )
    replace_paragraph(
        doc,
        "The benefit of rSVD",
        "The benefit of rSVD over IRLBA was likewise workload-dependent "
        "(Figure 3C). For the nine classification datasets, matched CPU SIMPLS "
        "runs were 1.00- to 1.45-fold faster with rSVD (median, 1.04-fold). "
        "Eight datasets differed by no more than 0.21 percentage points in "
        "accuracy, whereas MetRef differed by 4.0 percentage points, illustrating "
        "that rSVD is an approximate solver rather than a deterministic "
        "replacement. The advantage was much larger for the NMR cross-covariance "
        "problem: at 100 components, rSVD reduced CPU SIMPLS time from 436.3 to "
        "19.6 s (22.3-fold) with unchanged displayed RMSD. We therefore use IRLBA "
        "for deterministic estimator-matched validation and rSVD for the principal "
        "performance benchmark, with fixed oversampling, power-iteration, and seed "
        "settings and with numerical audit criteria reported in the Supplementary "
        "Material.",
    )
    caption = replace_paragraph(
        doc,
        "Figure 3. Internal acceleration",
        "Figure 3. Internal acceleration and solver regimes in fastPLS. Values "
        "above one indicate faster execution by the accelerated route. (A) CUDA "
        "relative to matched CPU rSVD for the four PLS families; NMR is excluded "
        "because it is analysed separately. (B) Metal relative to matched Apple "
        "CPU rSVD; daggers identify absolute predictive-metric differences greater "
        "than 0.005. (C) CPU rSVD relative to CPU IRLBA for SIMPLS; the NMR point "
        "uses the fixed 100-component comparison. Complete accuracy, memory, "
        "uncertainty, and numerical-audit results are provided in the Supplementary "
        "Material. No multicore CPU speed-up is inferred.",
    )
    caption.style = "Caption"
    replace_figure_before_caption(doc, "Figure 3. Internal acceleration", FIGURE)
    doc.save(MAIN_OUT)


def revise_supplement():
    copy2(SUPP_SRC, SUPP_OUT)


def main():
    if not FIGURE.exists():
        raise FileNotFoundError(FIGURE)
    DEST.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
