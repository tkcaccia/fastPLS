#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle45"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle46"
EVIDENCE = ROOT / "benchmark_results" / "manuscript_revision_cycle46_20260726"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle45_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle45_0.99.6_20260726.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle46_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle46_0.99.6_20260726.docx"
FIGURE = EVIDENCE / "supp_cpu_cuda_metal_outer_test.png"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def replace_all(paragraph, replacements):
    text = paragraph.text
    for old, new in replacements:
        text = text.replace(old, new)
    if text != paragraph.text:
        style = paragraph.style
        paragraph.clear()
        paragraph.style = style
        paragraph.add_run(text)


def revise_main():
    document = Document(MAIN_SOURCE)
    caption = find_paragraph(
        document,
        "Figure 2. Matched CPU and CUDA outer-test performance",
    )
    previous = caption._p.getprevious()
    if previous is None or not previous.xpath(".//w:drawing"):
        raise RuntimeError("Expected Figure 2 picture immediately before caption")
    previous.getparent().remove(previous)
    remove_paragraph(caption)

    results = find_paragraph(document, "Table 1 and Figure 2 show")
    replace_all(
        results,
        [
            (
                (
                    "Table 1 and Figure 2 show both matched CPU and CUDA backends "
                    "for the complete twelve-task biomedical benchmark at the "
                    "training-set-selected component count within each prespecified "
                    "grid and PLS family."
                ),
                (
                    "Table 1 reports matched CPU and CUDA results for the complete "
                    "twelve-task biomedical benchmark at each family-specific "
                    "component count selected from the training data. Supplementary "
                    "Figure S25 visualizes these results and, in a separate row, "
                    "matched CPU/Metal validation for the four datasets evaluated "
                    "on Apple hardware."
                ),
            )
        ],
    )

    replacements = [
        ("Figure 5", "__FIGURE_4__"),
        ("Figure 4", "__FIGURE_3__"),
        ("Figure 3", "__FIGURE_2__"),
        ("Table 1 and Figure 2", "Table 1 and Supplementary Figure S25"),
        ("Figure 2", "Supplementary Figure S25"),
        ("__FIGURE_4__", "Figure 4"),
        ("__FIGURE_3__", "Figure 3"),
        ("__FIGURE_2__", "Figure 2"),
    ]
    for paragraph in document.paragraphs:
        replace_all(paragraph, replacements)

    document.core_properties.title = (
        "fastPLS CMPB manuscript - backend figure moved to supplement"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)
    document.add_heading(
        "S36. CPU, CUDA, and Metal backend performance",
        level=1,
    )
    document.add_paragraph(
        (
            "Supplementary Figure S25 separates two matched analyses. Rows 1-3 "
            "retain the complete twelve-task CPU/CUDA benchmark at each "
            "family-specific component count selected from the training data. "
            "The fourth row adds the independently executed Apple Metal "
            "validation for the four datasets available locally on that system. "
            "Within the Metal row, CPU and Metal use identical data, split, "
            "family, float64 precision, argmax decoder, rSVD controls, and "
            "component count. These Metal component counts were prespecified for "
            "backend validation and are not presented as the family-specific "
            "training-selected settings used in rows 1-3. A missing Metal "
            "combination is marked not evaluated rather than inferred."
        )
    )
    picture = document.add_paragraph()
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(FIGURE), width=Inches(6.75))
    caption = document.add_paragraph(
        (
            "Figure S25. Backend-specific outer-test performance. Rows 1-3 show "
            "the training-set-selected CPU/CUDA comparison for the twelve "
            "biomedical tasks. Row 4 shows the separate matched CPU/Metal "
            "validation for CIFAR-100 (A=50), MetRef (A=22), Retina (A=20), and "
            "Tabula Muris (A=20). Circles, triangles, and squares denote CPU, "
            "CUDA, and Metal; segments join only setting-matched backends. Error "
            "bars are conditional 95% intervals on the fixed held-out set; NE "
            "denotes not evaluated. The Metal row is a separate portability "
            "validation, not part of the training-selected CPU/CUDA analysis."
        ),
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = False
    document.core_properties.title = (
        "fastPLS CMPB supplement - CPU CUDA and Metal performance"
    )
    document.save(SUPP_OUT)


def main():
    if not FIGURE.exists():
        raise FileNotFoundError(FIGURE)
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
