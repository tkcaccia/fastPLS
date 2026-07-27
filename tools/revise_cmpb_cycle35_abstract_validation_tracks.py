#!/usr/bin/env python3

from pathlib import Path
import shutil

from docx import Document


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle34"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle35"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle34_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle34_0.99.6_20260726.docx"
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle34_20260726.docx"
)

MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle35_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle35_0.99.6_20260726.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle35_20260726.docx"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def revise_main():
    document = Document(MAIN_SOURCE)

    methods = find_paragraph(document, "Methods: fastPLS provides")
    methods.text = (
        "Methods: fastPLS provides PLS-SVD, SIMPLS, orthogonal PLS (OPLS), "
        "and kernel PLS through one R interface. Two numerical validation "
        "tracks were prespecified. Estimator preservation was assessed only "
        "for deterministic IRLBA SIMPLS against de Jong SIMPLS while reusing "
        "deflation, coefficient, and prediction state across component "
        "prefixes. Separately, stochastic rSVD was assessed as an approximate "
        "solver using prediction, subspace, label-agreement, and "
        "predictive-metric failure criteria and was used in workflow-level "
        "runtime and memory benchmarks; it was not used as "
        "estimator-equivalence evidence. The package combines implicit "
        "cross-covariance products, compact prediction, compiled validation, "
        "double-precision reference paths, conditional float32 paths, "
        "compiled CPU execution linked to system BLAS/LAPACK, NVIDIA CUDA, "
        "and Apple Metal."
    )

    results = find_paragraph(document, "Results: Deterministic IRLBA SIMPLS")
    results.text = (
        "Results: The deterministic IRLBA path passed all 117 "
        "component-level comparisons with de Jong SIMPLS, supporting "
        "estimator preservation for that optimized deterministic "
        "implementation. In the separate rSVD reliability study, the initial "
        "one-power setting failed 16 of 117 approximation checks and was "
        "rejected for confirmatory use. After removal of the one-vector "
        "warm-start shortcut, oversampling by 10 directions with two power "
        "iterations passed all prespecified approximation thresholds, but "
        "remained numerically approximate: maximum relative prediction error "
        "was 0.0332, minimum prediction correlation was 0.99939, maximum "
        "score-subspace angle was 4.93 degrees, and minimum label agreement "
        "was 0.99. Accordingly, rSVD runtime and memory results describe "
        "approximate workflows rather than estimator preservation. In the "
        "primary estimator-matched software comparison, float64 SIMPLS using "
        "deterministic CPU IRLBA was faster than pls::simpls.fit on seven of "
        "nine datasets, with identical median accuracy on all nine."
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - deterministic and approximate abstract tracks"
    )
    document.save(MAIN_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    heading = document.add_heading(
        "35. The abstract did not sufficiently separate deterministic "
        "preservation from approximate rSVD benchmarking",
        level=1,
    )
    heading.paragraph_format.keep_with_next = True
    comment = document.add_paragraph(
        "Reviewer comment: The abstract should distinguish deterministic "
        "estimator preservation from the approximate rSVD performance "
        "benchmark more explicitly."
    )
    comment.paragraph_format.keep_with_next = True
    response = document.add_paragraph(
        "Response: Corrected. The abstract now defines two prespecified and "
        "separate numerical tracks. Estimator preservation is attributed only "
        "to deterministic IRLBA SIMPLS compared with de Jong SIMPLS. The rSVD "
        "study is explicitly described as an approximation-reliability and "
        "workflow-performance assessment governed by numerical failure "
        "criteria, not as estimator-equivalence evidence. The Results portion "
        "first reports the 117 deterministic preservation comparisons and then "
        "reports the separate rSVD acceptance audit, including its nonzero "
        "prediction and subspace discrepancies. It also states directly that "
        "rSVD runtime and memory results characterize approximate workflows, "
        "whereas the primary estimator-matched external comparison uses "
        "deterministic float64 IRLBA."
    )
    response.paragraph_format.keep_together = True
    document.core_properties.title = (
        "fastPLS CMPB response - abstract validation-track distinction"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    shutil.copy2(SUPP_SOURCE, SUPP_OUT)
    revise_response()


if __name__ == "__main__":
    main()
