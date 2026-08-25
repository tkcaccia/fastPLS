from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle94"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle95"
OUTPUT.mkdir(parents=True, exist_ok=True)


def replace_paragraph(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def find_paragraph(document, phrase):
    matches = [p for p in document.paragraphs if phrase in p.text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph containing {phrase!r}; found {len(matches)}")
    return matches[0]


def insert_after(paragraph, text, style="Body Text"):
    new = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addnext(new._p)
    return new


main_source = SOURCE / "fastPLS_CMPB_main_cycle94_0.99.25_20260825.docx"
main_output = OUTPUT / "fastPLS_CMPB_main_cycle95_0.99.25_20260825.docx"
main = Document(main_source)

validation = find_paragraph(main, "Fold construction, fitting, prediction, and metric accumulation")
replace_paragraph(
    validation,
    "Fold construction, fitting, prediction, and metric accumulation remain compiled where supported; "
    "grouped observations can be constrained to one fold. Model-selection endpoints include accuracy, "
    "balanced accuracy, training R², cross-validated Q², and RMSD. Training R² uses fitted training "
    "responses and the complete-training response mean. Independent-test Q² uses test prediction error "
    "but centers the denominator on the training-response mean. Cross-validated Q² instead sums fold "
    "prediction errors and centers each held-out fold on the mean estimated from that fold's training "
    "observations. For PLS-DA, the same R² and Q² formulas are applied to dummy-coded responses; these "
    "quantities are distinct from decoded-label accuracy and balanced accuracy. Hybrid OPLS, nonlinear-"
    "kernel, and Metal paths are identified explicitly."
)

benchmark = find_paragraph(main, "Twelve tasks covered metabolomics")
replace_paragraph(
    benchmark,
    benchmark.text.replace(
        "multivariate regression used RMSD, Q2, and held-out bootstrap intervals",
        "multivariate regression used RMSD, independent-test Q² relative to the training-response mean, "
        "and held-out bootstrap intervals",
    )
)

nmr = find_paragraph(main, "At the family-selected settings")
replace_paragraph(
    nmr,
    nmr.text.replace(
        "At the family-selected settings,",
        "At the family-selected settings, independent-test Q² used the fixed training-response mean."
        " Accordingly,",
    )
)

nmr_caption = find_paragraph(main, "Figure 4. NMR predictive and computational analyses")
replace_paragraph(
    nmr_caption,
    "Figure 4. NMR predictive and computational analyses. Panels A-C separate family-selected "
    "held-out performance from the deposited 165-component historical context. Reported Q² is "
    "independent-test Q² relative to the training-response mean. Panels D-E overlay observed and "
    "predicted intensities for held-out sample AMI-00BP-8 (index 155), whose per-spectrum RMSD under "
    "50-component SIMPLS CUDA rSVD was closest to the median across the 321 held-out spectra, over "
    "the full response range and 1.7-0.5 ppm expansion. Panel F reports matched float64 solver/backend "
    "resources at fixed family-specific component counts. rSVD used oversampling 20, two power "
    "iterations, and seed 123.",
)

main.save(main_output)


supp_source = SOURCE / "fastPLS_CMPB_supplement_cycle94_0.99.25_20260825.docx"
supp_output = OUTPUT / "fastPLS_CMPB_supplement_cycle95_0.99.25_20260825.docx"
supp = Document(supp_source)

cv = find_paragraph(supp, "pls.single.cv() evaluates eligible combinations")
replace_paragraph(
    cv,
    cv.text.replace(
        "Selection can use accuracy, balanced accuracy, R2, Q2, or RMSD.",
        "Selection can use accuracy, balanced accuracy, observed-mean cross-validated R², "
        "fold-training-mean cross-validated Q², or RMSD.",
    )
)

candidates = [p for p in supp.paragraphs if "held-out fold" in p.text]
if len(candidates) != 1:
    raise RuntimeError("Could not identify the cross-validated Q2 definition paragraph")
metric_context = candidates[0]
replace_paragraph(
    metric_context,
    "For cross-validated Q², each held-out fold is evaluated against a response mean calculated only "
    "from the corresponding fold-training observations; fold-specific prediction sums of squares and "
    "denominators are then accumulated before forming the reported Q². Independent-test Q² instead uses "
    "the mean of the complete training responses. R² uses the mean of the responses being fitted or "
    "evaluated and is not substituted for Q² when a training reference is unavailable. Multivariate RMSD "
    "is one global error over all response entries. For PLS-DA, training R² and held-out Q² use centered "
    "dummy-coded responses and must not be interpreted as decoded-label accuracy. The public outputs record "
    "these conventions in metrics$definitions."
)

for table in supp.tables:
    for row in table.rows:
        for cell in row.cells:
            if "Q² 0." in cell.text:
                cell.text = cell.text.replace("Q² ", "independent-test Q² ")

supp.save(supp_output)

print(main_output)
print(supp_output)
