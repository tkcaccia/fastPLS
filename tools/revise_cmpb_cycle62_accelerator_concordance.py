#!/usr/bin/env python3
"""Qualify accelerator speed claims by paired numerical concordance."""

from pathlib import Path
from shutil import copy2

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle61"
DEST = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle62"
EVIDENCE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle62_20260726"
)
FIGURE = EVIDENCE / "accelerator_concordance_speedups.png"
AUDIT_CSV = EVIDENCE / "accelerator_paired_concordance_audit.csv"
SUMMARY_CSV = EVIDENCE / "accelerator_concordance_summary.csv"

MAIN_SRC = SRC / "fastPLS_CMPB_main_cycle61_0.99.6_20260726.docx"
SUPP_SRC = SRC / "fastPLS_CMPB_supplement_cycle61_0.99.6_20260726.docx"
MAIN_OUT = DEST / "fastPLS_CMPB_main_cycle62_0.99.6_20260726.docx"
SUPP_OUT = DEST / "fastPLS_CMPB_supplement_cycle62_0.99.6_20260726.docx"


def replace_paragraph(doc, startswith, text):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(startswith):
            paragraph.text = text
            return paragraph
    raise RuntimeError(f"Paragraph not found: {startswith}")


def replace_text_in_paragraph(doc, startswith, old, new):
    paragraph = next(
        p for p in doc.paragraphs if p.text.startswith(startswith)
    )
    if old not in paragraph.text:
        raise RuntimeError(f"Text not found in paragraph: {old}")
    paragraph.text = paragraph.text.replace(old, new)
    return paragraph


def replace_figure_before_caption(doc, caption_prefix, image_path):
    caption = next(
        p for p in doc.paragraphs if p.text.startswith(caption_prefix)
    )
    node = caption._p.getprevious()
    while node is not None:
        if node.xpath(".//w:drawing"):
            node.getparent().remove(node)
            break
        node = node.getprevious()
    else:
        raise RuntimeError(f"Figure preceding {caption_prefix} not found")

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image_path), width=Inches(6.35))
    caption._p.addprevious(paragraph._p)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_cell_margins(cell, value=55):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "B7B7B7")
    set_repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        for col_index, cell in enumerate(row.cells):
            cell.width = Inches(widths[col_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if col_index in (0, 1, len(widths) - 1)
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(6.2)
                    if row_index == 0:
                        run.bold = True


def family_label(value):
    return {
        "plssvd": "PLS-SVD",
        "simpls": "SIMPLS",
        "opls": "OPLS",
        "kernelpls": "kernel PLS",
    }.get(value, value)


def dataset_label(value):
    return {
        "cbmc_citeseq": "CBMC CITE-seq",
        "ccle": "CCLE",
        "cifar100": "CIFAR-100",
        "gtex_v8": "GTEx v8",
        "metref": "MetRef",
        "prism": "PRISM",
        "retina": "Retina",
        "tabula": "Tabula Muris",
        "tcga_brca": "TCGA-BRCA",
        "tcga_hnsc_methylation": "TCGA-HNSC",
        "tcga_pan_cancer": "TCGA Pan-Cancer",
    }.get(value, value)


def fmt_metric(value, metric_name):
    if pd.isna(value):
        return "NA"
    if metric_name == "accuracy":
        return f"{value:.4f}"
    return f"{value:.4g}"


def fmt_pair(left, right, digits=2):
    if pd.isna(left) or pd.isna(right):
        return "NA"
    return f"{left:.{digits}f}/{right:.{digits}f}"


def add_audit_table(doc, data, accelerator, caption):
    doc.add_paragraph(caption, style="Caption")
    headers = [
        "Dataset",
        "Family",
        "A",
        "Metric CPU/acc.",
        "Prediction agreement",
        "Time CPU/acc. (s)",
        "Host RSS CPU/acc. (MB)",
        "Acc. memory (MB)",
        "Numerical status",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text

    for _, row in data.iterrows():
        cells = table.add_row().cells
        values = [
            dataset_label(row["dataset"]),
            family_label(row["family"]),
            str(int(row["ncomp"])),
            (
                f"{fmt_metric(row['metric_cpu'], row['metric_name'])}/"
                f"{fmt_metric(row['metric_accelerator'], row['metric_name'])}"
            ),
            f"{row['prediction_agreement']:.4f}",
            fmt_pair(
                row["time_cpu_sec"],
                row["time_accelerator_sec"],
                3,
            ),
            fmt_pair(
                row["host_rss_cpu_mb"],
                row["host_rss_accelerator_mb"],
                0,
            ),
            (
                "NA"
                if pd.isna(row["accelerator_memory_mb"])
                else f"{row['accelerator_memory_mb']:.0f}"
            ),
            {
                "concordant": "Concordant; speed eligible",
                "discordant_metric": "Metric discordant; excluded",
                "discordant_prediction": "Prediction discordant; excluded",
                "prediction_not_archived": "Agreement not established; excluded",
            }[row["evidence_status"]],
        ]
        for cell, value in zip(cells, values):
            cell.text = value

    style_table(
        table,
        [1.05, 0.72, 0.35, 1.05, 0.85, 1.05, 1.20, 0.85, 1.35],
    )
    doc.add_paragraph(
        (
            f"For {accelerator}, speed eligibility required an absolute "
            "predictive-metric difference <=0.005 and sample-level prediction "
            "agreement >=0.995. Times and host RSS are medians from the original "
            "isolated benchmark. Accelerator memory is absolute process GPU "
            "memory for CUDA and incremental unified-process RSS for Metal; "
            "neither is isolated solver workspace."
        ),
        style="Body Text",
    )


def revise_main(audit, summary):
    doc = Document(MAIN_SRC)

    replace_text_in_paragraph(
        doc,
        "Methods: fastPLS provides",
        (
            "A separate Apple M3 campaign evaluated Metal portability, precision, "
            "shape-dependent performance, unified-memory use, and effective stage "
            "residency in isolated CPU-versus-Metal runs."
        ),
        (
            "Accelerator speed was considered interpretable only when the absolute "
            "paired predictive-metric difference was <=0.005 and sample-level "
            "prediction agreement was >=0.995; routes failing either criterion "
            "were retained as discordant but excluded from speed-up summaries. "
            "A separate Apple M3 campaign evaluated Metal portability, precision, "
            "shape-dependent performance, unified-memory use, and effective stage "
            "residency in isolated CPU-versus-Metal runs."
        ),
    )

    replace_text_in_paragraph(
        doc,
        "Results: The deterministic",
        (
            "In the separate Apple M3 study, Metal accelerated sufficiently large "
            "dense float64 workloads by up to 4.35-fold on CIFAR-100, but CPU "
            "float32 was faster on smaller omics tasks and NMR. A float64 Metal "
            "PLS-SVD accuracy discrepancy prevented a general backend-equivalence "
            "claim."
        ),
        (
            "Among 44 non-NMR CPU-CUDA pairs, 28 passed both numerical criteria "
            "and CUDA was faster in seven, with a maximum eligible speed-up of "
            "8.90-fold. Among 12 CPU-Metal pairs, six passed both criteria and "
            "none was faster with Metal. Larger nominal accelerator speed-ups in "
            "discordant routes were excluded rather than interpreted as "
            "acceleration."
        ),
    )

    replace_text_in_paragraph(
        doc,
        "Within each dataset, methods used",
        (
            "CPU and CUDA rSVD rows within a PLS family used the same preprocessing, "
            "response, component count, and argmax or regression prediction rule "
            "and are therefore estimator matched."
        ),
        (
            "CPU and CUDA rSVD rows within a PLS family used the same preprocessing, "
            "response, component count, and argmax or regression prediction rule. "
            "Because backend identity does not ensure numerical agreement for an "
            "approximate solver, accelerator speed-up eligibility additionally "
            "required an absolute predictive-metric difference <=0.005 and paired "
            "prediction agreement >=0.995."
        ),
    )

    replace_paragraph(
        doc,
        "We first evaluated whether",
        (
            "We first evaluated whether the computational implementation of "
            "fastPLS improved the practical use of SIMPLS relative to independent "
            "R software under a controlled single-CPU setting. We then compared "
            "the compiled CPU route with numerically concordant CUDA and Metal "
            "routes and examined when randomized SVD (rSVD) was preferable to the "
            "deterministic IRLBA route. The final analyses focus on multivariate "
            "NMR prediction and million-sample ImageNet embeddings. Formal "
            "estimator-preservation tests, numerical audits, quarantined backend "
            "routes, component paths, and complete paired resource tables are "
            "reported in the Supplementary Material."
        ),
    )

    replace_paragraph(
        doc,
        "Backend acceleration depended",
        (
            "Backend performance was interpreted only after a paired numerical "
            "audit (Figure 3A-B). Of 44 non-NMR CPU-CUDA comparisons, 28 met both "
            "the predictive-metric tolerance (absolute difference <=0.005) and "
            "prediction-agreement tolerance (>=0.995). CUDA was faster in seven of "
            "these 28 concordant routes: PLS-SVD on CIFAR-100 (8.90-fold), SIMPLS, "
            "OPLS, and kernel PLS on GTEx v8 (1.65-1.83-fold), and SIMPLS, OPLS, "
            "and kernel PLS on CBMC CITE-seq (1.78-2.09-fold). Sixteen CUDA routes "
            "were quarantined: three failed the metric tolerance and 13 additional "
            "routes failed prediction agreement. Consequently, the nominal "
            "14.50-fold CIFAR-100 SIMPLS speed-up was not counted because paired "
            "label agreement was 0.9922. Of 12 CPU-Metal comparisons, six were "
            "numerically concordant and none was faster with Metal. The nominal "
            "CIFAR-100 Metal speed-ups were excluded because PLS-SVD and SIMPLS "
            "label agreement was 0.9058 and 0.8904, respectively. The evidence "
            "therefore supports route-specific CUDA acceleration, not a general "
            "accelerator or Metal speed advantage."
        ),
    )

    caption = replace_paragraph(
        doc,
        "Figure 3. Internal acceleration",
        (
            "Figure 3. Numerically qualified backend and solver comparisons in "
            "fastPLS. Colored cells report speed-up only when the absolute paired "
            "predictive-metric difference was <=0.005 and sample-level prediction "
            "agreement was >=0.995. Gray D-metric and D-pred cells failed the "
            "metric and prediction criteria, respectively, and were excluded from "
            "speed summaries. (A) CUDA relative to matched CPU rSVD; NMR is "
            "analysed separately. (B) Metal relative to matched Apple CPU rSVD. "
            "(C) Approximate CPU rSVD relative to deterministic CPU IRLBA for "
            "SIMPLS; the NMR point uses the fixed 100-component comparison. "
            "Complete paired metrics, prediction agreement, runtime, host memory, "
            "and accelerator memory are reported in Supplementary Tables S53-S54."
        ),
    )
    caption.style = "Caption"
    replace_figure_before_caption(doc, "Figure 3. Numerically", FIGURE)

    replace_paragraph(
        doc,
        "The backend design is a second contribution",
        (
            "The backend design is a portability contribution, but accelerator "
            "availability must not be conflated with numerical equivalence or "
            "speed. Under the prespecified agreement criteria, CUDA was faster in "
            "seven of 28 concordant non-NMR routes; the remaining concordant routes "
            "favored CPU. No concordant route in the prespecified Apple campaign "
            "favored Metal. The nominal Metal CIFAR-100 gains were invalid as "
            "same-workflow speed comparisons because their predictions were "
            "discordant. CUDA and Metal therefore remain supported execution "
            "layers, but performance claims are restricted to the explicitly "
            "validated model, precision, shape, and prediction combinations in "
            "Supplementary Tables S53-S54. OPLS filtering, nonlinear-kernel "
            "construction, reduced decompositions, and public Metal "
            "cross-validation also retain host stages."
        ),
    )

    replace_paragraph(
        doc,
        "fastPLS combines a shape-aware",
        (
            "fastPLS combines a shape-aware accelerated sequential SIMPLS "
            "implementation with memory-aware PLS-SVD, compiled validation, "
            "compact prediction, conditional float32 execution, and latent-space "
            "LDA. Compiled CPU is the reference execution route. NVIDIA CUDA and "
            "Apple Metal provide additional hardware paths, but this study "
            "supports acceleration only for seven numerically concordant CUDA "
            "configurations and does not demonstrate a concordant Metal speed "
            "advantage. Float32 storage, rSVD, and accelerator execution are "
            "therefore reported as route-conditional capabilities rather than "
            "universal improvements. The GPL-3 R package calls reusable C++ "
            "components maintained with the MIT-licensed kodama-cpp codebase; "
            "future work will expand native accelerator residency and numerical "
            "validation across hardware architectures."
        ),
    )
    doc.save(MAIN_OUT)


def revise_supplement(audit, summary):
    doc = Document(SUPP_SRC)
    replace_paragraph(
        doc,
        "Supplementary Figure S25 separates",
        (
            "Supplementary Figure S25 displays all measured backend-specific "
            "outer-test metrics, including discordant routes, and must not be read "
            "as a speed-up figure. Rows 1-3 retain the twelve-task CPU/CUDA "
            "benchmark at each training-selected component count; row 4 is the "
            "separate Apple CPU/Metal campaign. Numerical concordance and speed "
            "eligibility are assessed independently in Section S42 and "
            "Supplementary Tables S52-S54."
        ),
    )
    replace_paragraph(
        doc,
        "Figure S25. Backend-specific outer-test performance",
        (
            "Figure S25. Backend-specific outer-test predictive metrics. Rows 1-3 "
            "show the training-selected CPU/CUDA comparison; row 4 shows the "
            "separate CPU/Metal campaign. Segments indicate matching settings, not "
            "numerical equivalence. The figure deliberately retains discordant "
            "routes; only Section S42 and Tables S52-S54 determine whether a route "
            "is eligible for speed interpretation."
        ),
    )
    replace_paragraph(
        doc,
        "Table S22. Repeated CPU-versus-Metal validation",
        (
            "Table S22. Repeated CPU-versus-Metal workflow measurements on Apple "
            "M3. Time and predictive metric are medians [minimum-maximum] across "
            "three seeded rSVD fits; NMR entries are single guarded feasibility "
            "runs. CPU/Metal time ratios are nominal measurements and are not "
            "speed-up evidence unless the corresponding route passes both "
            "concordance criteria in Section S42. Incremental RSS is "
            "baseline-corrected unified-process memory, not dedicated GPU VRAM."
        ),
    )
    replace_paragraph(
        doc,
        "All 1,064 isolated benchmark fits",
        (
            "All 1,064 isolated benchmark fits and decompositions completed. "
            "Metal had shorter nominal elapsed time for selected float64 "
            "CIFAR-100 and synthetic configurations, but these cells are "
            "quarantined from acceleration claims. CIFAR-100 PLS-SVD and SIMPLS "
            "failed the paired metric and prediction criteria, while complete "
            "sample-level prediction vectors were not archived for the synthetic "
            "shape campaign. In the prespecified real-data audit, six of 12 "
            "CPU/Metal routes were numerically concordant and all six were faster "
            "on CPU. Metal was also slower for MetRef, Retina, Tabula Muris, "
            "tall-thin synthetic matrices, and NMR. Float32 CPU/Metal median "
            "accuracy was close in the tested CIFAR routes, but CPU remained "
            "faster; these precision-specific results do not rescue the "
            "discordant float64 comparisons."
        ),
    )
    replace_paragraph(
        doc,
        "At each family-specific component count",
        (
            "At each family-specific component count selected using training data, "
            "CPU and CUDA rSVD paths used the same split, family, component count, "
            "prediction rule, and float64 inputs. This matching alone did not "
            "establish numerical concordance because rSVD and backend kernels can "
            "produce different approximate subspaces. The earlier unfiltered "
            "count of 13 nominal CUDA speed-ups is superseded by the paired audit "
            "in Section S42: seven speed-ups remained after requiring both metric "
            "and prediction agreement. NMR and ImageNet remain in their dedicated "
            "analyses."
        ),
    )

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = (
        section.page_height,
        section.page_width,
    )
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    doc.add_paragraph(
        "S42. Numerically qualified accelerator comparison",
        style="Heading 1",
    )
    doc.add_paragraph(
        (
            "Backend speed was separated from numerical agreement. A paired route "
            "was speed eligible only when (i) the absolute difference in held-out "
            "accuracy or regression metric was no greater than 0.005 and (ii) "
            "sample-level prediction agreement was at least 0.995. For "
            "classification, agreement is the fraction of identical labels. For "
            "multivariate regression, agreement is one minus the relative "
            "Frobenius prediction error. Predictions were regenerated with the "
            "same task objects, component count, rSVD controls, seed, precision, "
            "and classifier; timing and memory remain the medians from the original "
            "three isolated runs. Failing routes are retained below but excluded "
            "from all acceleration counts."
        ),
        style="First Paragraph",
    )

    doc.add_paragraph(
        (
            "Table S52. Summary of the paired numerical-concordance audit. "
            "Accelerator-faster counts include only concordant routes."
        ),
        style="Caption",
    )
    headers = [
        "Accelerator",
        "Paired",
        "Concordant",
        "Metric discordant",
        "Prediction discordant",
        "Faster among concordant",
        "Median speed-up",
        "Maximum speed-up",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for _, row in summary.iterrows():
        values = [
            row["accelerator"],
            str(int(row["paired_routes"])),
            str(int(row["concordant_routes"])),
            str(int(row["metric_discordant_routes"])),
            str(int(row["prediction_discordant_routes"])),
            str(int(row["accelerator_faster_concordant"])),
            f"{row['median_speedup_concordant']:.2f}x",
            f"{row['maximum_speedup_concordant']:.2f}x",
        ]
        for cell, value in zip(table.add_row().cells, values):
            cell.text = value
    style_table(table, [1.05, 0.60, 0.75, 1.10, 1.25, 1.35, 1.10, 1.10])

    add_audit_table(
        doc,
        audit[audit["accelerator"] == "CUDA"],
        "CUDA",
        (
            "Table S53. Complete paired CPU/CUDA numerical, runtime, and memory "
            "audit. CPU/acc. values are shown in that order."
        ),
    )
    add_audit_table(
        doc,
        audit[audit["accelerator"] == "Metal"],
        "Metal",
        (
            "Table S54. Complete paired Apple CPU/Metal numerical, runtime, and "
            "memory audit. CPU/acc. values are shown in that order."
        ),
    )
    doc.save(SUPP_OUT)


def main():
    required = [
        MAIN_SRC,
        SUPP_SRC,
        FIGURE,
        AUDIT_CSV,
        SUMMARY_CSV,
    ]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise FileNotFoundError("\n".join(missing))
    DEST.mkdir(parents=True, exist_ok=True)
    audit = pd.read_csv(AUDIT_CSV)
    summary = pd.read_csv(SUMMARY_CSV)
    revise_main(audit, summary)
    revise_supplement(audit, summary)
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
