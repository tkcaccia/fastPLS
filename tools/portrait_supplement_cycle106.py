from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = (
    ROOT
    / "artifacts"
    / "CMPB_rewrite_20260825_cycle105"
    / "fastPLS_CMPB_supplement_cycle105_0.99.25_20260825.docx"
)
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260826_cycle106"
OUTPUT = OUTPUT_DIR / "fastPLS_CMPB_supplement_cycle106_0.99.25_20260826.docx"

PAGE_WIDTH = Inches(8.5)
PAGE_HEIGHT = Inches(11)
MARGIN = Inches(0.45)
TABLE_WIDTH_DXA = 10944  # 7.6 inches, matching the portrait text width.


def set_cell_margins(table, value=45):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for edge in ("top", "start", "bottom", "end"):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, target_width):
    grid = table._tbl.tblGrid
    grid_columns = list(grid.gridCol_lst)
    original = [int(column.get(qn("w:w"))) for column in grid_columns]
    total = sum(original)
    if total <= 0:
        original = [1] * len(grid_columns)
        total = len(grid_columns)

    widths = [max(180, round(target_width * width / total)) for width in original]
    widths[-1] += target_width - sum(widths)

    for column, width in zip(grid_columns, widths):
        column.set(qn("w:w"), str(width))

    tbl_pr = table._tbl.tblPr
    tbl_width = tbl_pr.find(qn("w:tblW"))
    if tbl_width is None:
        tbl_width = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_width)
    tbl_width.set(qn("w:w"), str(target_width))
    tbl_width.set(qn("w:type"), "dxa")

    table.autofit = False
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell_width = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(cell_width)
            cell_width.set(qn("w:w"), str(widths[min(index, len(widths) - 1)]))
            cell_width.set(qn("w:type"), "dxa")


def set_table_font(table):
    columns = len(table.columns)
    size = 5.5 if columns >= 8 else 5.9 if columns >= 6 else 6.2
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                for run in paragraph.runs:
                    run.font.size = Pt(size)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document(SOURCE)

    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = PAGE_WIDTH
        section.page_height = PAGE_HEIGHT
        section.left_margin = MARGIN
        section.right_margin = MARGIN
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN

    for table in document.tables:
        set_table_width(table, TABLE_WIDTH_DXA)
        set_cell_margins(table)
        set_table_font(table)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
