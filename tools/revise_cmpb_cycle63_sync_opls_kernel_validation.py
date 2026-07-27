#!/usr/bin/env python3
"""Synchronize the expanded OPLS/kernel-PLS validation across the manuscript."""

import csv
import hashlib
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle62"
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle63"
EVIDENCE_DIR = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle63_20260726"
)

MAIN_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_main_cycle62_0.99.6_20260726.docx"
)
SUPP_SOURCE = (
    SOURCE_DIR
    / "fastPLS_CMPB_supplement_cycle62_0.99.6_20260726.docx"
)
MAIN_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_main_cycle63_0.99.6_20260726.docx"
)
SUPP_OUTPUT = (
    OUTPUT_DIR
    / "fastPLS_CMPB_supplement_cycle63_0.99.6_20260726.docx"
)

OLD_LEDGER = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle45_20260726"
    / "analysis_commit_provenance.csv"
)
NEW_LEDGER = EVIDENCE_DIR / "analysis_commit_provenance.csv"
EXPANDED_ARCHIVE = (
    ROOT
    / "benchmark_results"
    / "opls_kernel_setting_reliability_20260726"
)
EXPANDED_SCRIPT = ROOT / "benchmark" / "benchmark_opls_kernel_setting_reliability.R"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_paragraph(document, prefix, text):
    paragraph = find_paragraph(document, prefix)
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def insert_after(paragraph, text, style="Normal"):
    new_paragraph = paragraph._parent.add_paragraph(style=style)
    new_paragraph.add_run(text)
    paragraph._p.addnext(new_paragraph._p)
    return new_paragraph


def tree_digest(directory):
    digest = hashlib.sha256()
    files = sorted(path for path in directory.rglob("*") if path.is_file())
    for path in files:
        relative = path.relative_to(directory).as_posix()
        digest.update(relative.encode("utf-8"))
        digest.update(b"\0")
        with path.open("rb") as handle:
            for block in iter(lambda: handle.read(1024 * 1024), b""):
                digest.update(block)
        digest.update(b"\0")
    return digest.hexdigest()


def write_updated_ledger():
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    with OLD_LEDGER.open(newline="", encoding="utf-8") as handle:
        rows = list(csv.DictReader(handle))

    rows.append(
        {
            "analysis_id": "A15",
            "reported_output": (
                "Section S41; Tables S50-S51; Figure S41"
            ),
            "result_archive": (
                "benchmark_results/"
                "opls_kernel_setting_reliability_20260726"
            ),
            "package_version": "0.99.6",
            "package_commit": "not recorded in archived metadata",
            "package_commit_status": "not recoverable; not inferred",
            "benchmark_script": (
                "benchmark/benchmark_opls_kernel_setting_reliability.R"
            ),
            "benchmark_script_commit": (
                "not recorded in archived metadata"
            ),
            "benchmark_script_md5_current_copy": hashlib.md5(
                EXPANDED_SCRIPT.read_bytes()
            ).hexdigest(),
            "kodama_cpp_commit": "not recorded in archived metadata",
            "data_split_provenance": (
                "Six fixed synthetic/real tasks; seed 123; identical "
                "train/test partitions and five-fold assignments across 11 "
                "OPLS/kernel settings"
            ),
            "result_archive_sha256": tree_digest(EXPANDED_ARCHIVE),
        }
    )

    with NEW_LEDGER.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0]))
        writer.writeheader()
        writer.writerows(rows)
    return rows[-1]


def revise_main():
    document = Document(MAIN_SOURCE)

    replace_paragraph(
        document,
        "Methods: fastPLS provides",
        (
            "Methods: fastPLS provides PLS-SVD, SIMPLS, orthogonal PLS "
            "(OPLS), and kernel PLS through one R interface. Two numerical "
            "validation tracks were prespecified. Estimator preservation was "
            "assessed for deterministic IRLBA SIMPLS against de Jong SIMPLS. "
            "A separate deterministic float64 CPU reliability study compared "
            "OPLS with one, two, and three orthogonal components and kernel "
            "PLS under eight linear, RBF, and polynomial kernel settings "
            "against independent equation-level operators followed by "
            "pls::simpls.fit across six synthetic and real tasks. Separately, "
            "randomized singular value decomposition (rSVD) was assessed as "
            "an approximate solver using prediction, subspace, "
            "label-agreement, and predictive-metric failure criteria and was "
            "used in workflow-level runtime and memory benchmarks; it was not "
            "used as estimator-equivalence evidence. The package combines "
            "implicit cross-covariance products, compact prediction, compiled "
            "validation, double-precision reference paths, conditional "
            "float32 paths, compiled CPU execution linked to system "
            "BLAS/LAPACK, NVIDIA CUDA, and Apple Metal. Accelerator speed was "
            "considered interpretable only when the absolute paired "
            "predictive-metric difference was <=0.005 and sample-level "
            "prediction agreement was >=0.995; routes failing either "
            "criterion were retained as discordant but excluded from speed-up "
            "summaries. A separate Apple M3 campaign evaluated Metal "
            "portability, precision, shape-dependent performance, "
            "unified-memory use, and effective stage residency in isolated "
            "CPU-versus-Metal runs."
        ),
    )

    replace_paragraph(
        document,
        "Results: The deterministic",
        (
            "Results: The deterministic IRLBA SIMPLS path passed all 117 "
            "component-level comparisons with de Jong SIMPLS. In the expanded "
            "OPLS/kernel-PLS study, all 66 setting/task endpoint comparisons "
            "met every prespecified criterion, the selected component count "
            "agreed in all 66 comparisons, and all 1,540 fold-by-component "
            "fits completed without failure. This establishes implementation "
            "reliability for the evaluated deterministic float64 CPU settings, "
            "not predictive superiority of an OPLS or kernel configuration. "
            "In the separate rSVD reliability study, the initial one-power "
            "setting failed 16 of 117 approximation checks and was rejected "
            "for confirmatory use. After removal of the one-vector warm-start "
            "shortcut, oversampling by 10 directions with two power iterations "
            "passed all prespecified approximation thresholds, but remained "
            "numerically approximate: maximum relative prediction error was "
            "0.0332, minimum prediction correlation was 0.99939, maximum "
            "score-subspace angle was 4.93 degrees, and minimum label "
            "agreement was 0.99. Accordingly, rSVD runtime and memory results "
            "describe approximate workflows rather than estimator "
            "preservation. In a direct five-shape rSVD workflow comparison "
            "that changed only PLS family, the CUDA SIMPLS/PLS-SVD total-time "
            "ratio was 0.918-0.979; CPU and Metal ratios crossed unity, "
            "demonstrating matrix-shape-dependent rather than universal "
            "runtime parity. In the primary estimator-matched software "
            "comparison, float64 SIMPLS using deterministic CPU IRLBA was "
            "faster than pls::simpls.fit on seven of nine datasets, with "
            "identical median accuracy on all nine. Among 44 non-NMR CPU-CUDA "
            "pairs, 28 passed both numerical criteria and CUDA was faster in "
            "seven, with a maximum eligible speed-up of 8.90-fold. Among 12 "
            "CPU-Metal pairs, six passed both criteria and none was faster "
            "with Metal. Larger nominal accelerator speed-ups in discordant "
            "routes were excluded rather than interpreted as acceleration. "
            "Matched float32 inputs reduced stored matrix size by approximately "
            "48% on MetRef and PRISM, although speed and process memory were "
            "not uniformly improved. Within CPU SIMPLS, rSVD was only "
            "1.00-1.45-fold faster than IRLBA across nine classification tasks "
            "but 22.26-fold faster on the extreme-response NMR task at 100 "
            "components. Latent-space LDA improved accuracy in seven of eight "
            "deterministic SIMPLS comparisons (median gain 2.64 percentage "
            "points) with a median 1.13-fold total-time ratio relative to "
            "argmax. In the exploratory million-sample ImageNet stress test, "
            "SIMPLS-LDA outperformed argmax across all 100-1,000-component "
            "points, reaching top-1 accuracy 0.8093 at 1,000 components."
        ),
    )

    replace_paragraph(
        document,
        "A second prespecified estimator-validation study",
        (
            "A second prespecified estimator-validation study covered three "
            "OPLS settings (north=1, 2, and 3) and eight kernel-PLS settings: "
            "linear; RBF with gamma=0.25/p, 1/p, and 4/p; polynomial degrees "
            "2, 3, and 4 with offset 1; and a homogeneous degree-3 polynomial "
            "kernel with offset 0. An independent R reference implemented the "
            "Trygg-Wold orthogonal-weight and deflation equations without "
            "calling fastPLS filtering code; the filtered predictor matrix was "
            "then fitted with pls::simpls.fit. Kernel references independently "
            "constructed and training-centred Gram matrices, applied stored "
            "training centring to held-out kernels, and fitted "
            "pls::simpls.fit. Six tasks covered regression and classification, "
            "p<n and p>n, an ill-conditioned design, gasoline spectroscopy, "
            "and breast molecular classification, giving 66 setting/task "
            "comparisons. Prespecified endpoints were operator, coefficient, "
            "prediction, score-subspace, decoded-label, predictive-metric, "
            "failure, and fixed 5-fold selected-component agreement. "
            "Deterministic IRLBA and float64 CPU execution were used throughout; "
            "documented exact fallback was retained only when the smaller "
            "cross-covariance dimension was below six. The study was designed "
            "to test implementation reliability across settings, not to rank "
            "their predictive performance."
        ),
    )

    intro = find_paragraph(
        document,
        "We first evaluated whether the computational implementation",
    )
    insert_after(
        intro,
        (
            "Formal deterministic reliability testing was completed before "
            "interpreting workflow performance. Across three OPLS "
            "orthogonal-component settings and eight linear, RBF, and "
            "polynomial kernel-PLS settings, all 66 setting/task endpoint "
            "comparisons passed, all 66 fixed-fold component selections "
            "agreed with the independent reference, and all 1,540 "
            "fold-by-component fits completed. These results support the "
            "reliability of the evaluated implementations under deterministic "
            "float64 CPU execution; they do not show that any kernel or OPLS "
            "setting is predictively superior."
        ),
    )

    replace_paragraph(
        document,
        "The independent OPLS and nonlinear-kernel study",
        (
            "The expanded OPLS and kernel-PLS study reduces a previous "
            "validation gap across three orthogonal-component counts and eight "
            "linear, RBF, and polynomial kernel definitions. Agreement was "
            "demonstrated in all 66 setting/task comparisons at the orthogonal "
            "filter or centred Gram operator, coefficient, score-subspace, "
            "prediction, classification, and fixed-fold selection levels, and "
            "all 1,540 fold-by-component fits completed. This evidence "
            "supports implementation reliability across the evaluated settings; "
            "it does not establish predictive superiority of a kernel or OPLS "
            "configuration. Its scope is deliberately restricted to "
            "deterministic float64 CPU execution and the tested tasks and "
            "shapes, and does not extend equivalence to rSVD, float32, CUDA, "
            "Metal, or untested settings."
        ),
    )

    replace_paragraph(
        document,
        "The fastPLS R package, benchmark workflows",
        (
            "The fastPLS R package, benchmark workflows, analysis scripts, "
            "machine-readable result tables, synthetic generators, and "
            "aggregate benchmark outputs are available at "
            "https://github.com/tkcaccia/fastPLS. Low-level reusable C++ "
            "components are maintained at "
            "https://github.com/tkcaccia/kodama-cpp. Supplementary Table S41 "
            "and benchmark_results/manuscript_revision_cycle63_20260726/"
            "analysis_commit_provenance.csv map every reported analysis to "
            "its result archive, generating script, recorded package version, "
            "exact package commit when captured during execution, data/split "
            "provenance, and SHA-256 archive digest. Historical archives that "
            "recorded version 0.99.6 but not a Git SHA are marked explicitly "
            "as not recoverable; no later commit is assigned retrospectively. "
            "The package bundles only the GPL-compatible breast and colon "
            "examples; prepared real-data benchmark matrices are not "
            "redistributed. benchmark/DATA_ACQUISITION.md and "
            "benchmark/acquire_publication_datasets.R provide authoritative "
            "source links, exact release identifiers, executable public "
            "downloads, and checksum validation for user-authorized restricted "
            "files."
        ),
    )

    document.save(MAIN_OUTPUT)


def append_provenance_row(document, ledger_row):
    target = None
    for table in document.tables:
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header[:2] == ["ID", "Reported output"]:
            target = table
            break
    if target is None:
        raise RuntimeError("Provenance table not found")

    row = target.add_row()
    values = [
        ledger_row["analysis_id"],
        ledger_row["reported_output"],
        Path(ledger_row["result_archive"]).name,
        Path(ledger_row["benchmark_script"]).name,
        ledger_row["package_version"],
        "NR",
        ledger_row["result_archive_sha256"][:12],
    ]
    for index, (cell, value) in enumerate(zip(row.cells, values)):
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if index in (1, 2, 3)
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.font.size = Pt(5)


def revise_supplement(ledger_row):
    document = Document(SUPP_SOURCE)

    replace_paragraph(
        document,
        "S34. Independent OPLS",
        (
            "S34. Initial default-setting OPLS and nonlinear kernel-PLS "
            "estimator validation"
        ),
    )

    replace_paragraph(
        document,
        "Pass thresholds were operator relative error",
        (
            "Pass thresholds were operator relative error <=1e-10, prediction "
            "relative error <=1e-4, coefficient relative error <=1e-3, maximum "
            "predictive or orthogonal score-subspace angle <=0.1 degrees, "
            "classification label agreement >=0.995, and predictive-metric "
            "difference <=0.005. All 18 endpoints in this initial "
            "default-setting analysis passed. Maximum operator, prediction, "
            "and coefficient relative errors were 3.33e-15, 2.52e-12, and "
            "1.12e-10; the largest predictive and orthogonal angles were "
            "2.09e-6 and 1.21e-6 degrees; minimum label agreement was 1.000. "
            "All 420 fold-component fits completed, and the selected component "
            "count agreed in all 18 family-task comparisons. Section S41 "
            "extends this validation to 11 OPLS/kernel settings and 66 "
            "setting/task comparisons. Both studies validate deterministic "
            "float64 CPU implementation reliability under their tested "
            "conditions; neither establishes predictive superiority, nor "
            "equivalence for rSVD, float32, CUDA, or Metal."
        ),
    )

    replace_paragraph(
        document,
        "Table S41 is the authoritative mapping",
        (
            "Table S41 is the authoritative mapping between manuscript outputs "
            "and computational archives. A package commit identifies code "
            "loaded for computation; a benchmark-script commit identifies "
            "orchestration code; and the archive SHA-256 identifies the exact "
            "result tree. These identifiers are not interchangeable. Only A14 "
            "contains an exact package SHA written during execution. For "
            "A01-A13 and A15, the archived metadata identify fastPLS 0.99.6 "
            "but do not contain an immutable package or script SHA; these "
            "fields therefore remain unavailable. The current script MD5 "
            "values are retained in the machine-readable ledger for file "
            "identification but are not represented as historical run "
            "provenance. Future runs use benchmark/write_run_provenance.R to "
            "capture these fields before computation."
        ),
    )

    replace_paragraph(
        document,
        "Machine-readable ledger:",
        (
            "Machine-readable ledger: benchmark_results/"
            "manuscript_revision_cycle63_20260726/"
            "analysis_commit_provenance.csv. The result-archive digest permits "
            "verification of the archived outputs even where historical "
            "source commit capture was incomplete."
        ),
    )

    replace_paragraph(
        document,
        "The principal manuscript now emphasizes",
        (
            "The principal manuscript now emphasizes the comparative software "
            "benchmark, backend scaling, and the NMR and ImageNet applications. "
            "The technical evidence supporting those claims is retained in "
            "this Supplement: estimator-preservation and de Jong mapping "
            "(Sections S13-S14), precision validation (Sections S19 and S32), "
            "classifier agreement (Section S20), rSVD numerical reliability "
            "(Section S23), direct PLS-SVD versus SIMPLS shape experiments "
            "(Section S24), cross-validation comparisons (Section S29), "
            "optimization ablations (Section S30), initial and expanded OPLS "
            "and kernel-PLS reliability studies (Sections S34 and S41), and "
            "backend residency and Metal validation (Sections S36-S38). Figure "
            "S40 provides a compact visual index of the precision, solver, and "
            "classifier evidence; the underlying tables remain authoritative."
        ),
    )

    replace_paragraph(
        document,
        "The default-setting validation in Section S34 was extended",
        (
            "The default-setting validation in Section S34 was extended to "
            "test whether agreement depended on the number of OPLS orthogonal "
            "components or on the kernel definition. The same six fixed tasks "
            "were used: synthetic regression and classification with p<n and "
            "p>n, an ill-conditioned p>n regression design, gasoline "
            "spectroscopy regression, and breast molecular classification. "
            "All calculations used double precision, the CPU backend, "
            "deterministic IRLBA, seed 123, identical train/test partitions, "
            "and identical five-fold component-selection folds. This was an "
            "implementation-reliability study rather than a comparison of "
            "predictive superiority among settings."
        ),
    )

    replace_paragraph(
        document,
        "OPLS was evaluated with one, two, and three",
        (
            "OPLS was evaluated with one, two, and three orthogonal components. "
            "For each setting, the independent reference recomputed the "
            "Trygg-Wold orthogonal weight, score, loading, and deflation "
            "sequence before fitting pls::simpls.fit to the filtered "
            "predictors. Kernel PLS was evaluated under eight settings: a "
            "linear kernel; RBF kernels with gamma equal to 0.25/p, 1/p, and "
            "4/p; polynomial kernels of degree two, three, and four with "
            "offset one; and a homogeneous degree-three polynomial kernel "
            "with offset zero. Every Gram matrix and its train/test centring "
            "were constructed independently before the reference SIMPLS fit. "
            "The three OPLS and eight kernel settings applied to six tasks "
            "yielded 66 setting/task comparisons."
        ),
    )

    replace_paragraph(
        document,
        "All 1,540 fold-by-component fits completed",
        (
            "All 1,540 fold-by-component fits completed without failure, and "
            "fastPLS selected the same component count as the independent "
            "reference in all 66 of 66 setting/task comparisons (Table S51). "
            "These results establish implementation reliability for the tested "
            "OPLS and kernel-PLS settings under deterministic float64 CPU "
            "execution. They do not demonstrate predictive superiority of any "
            "setting and do not extend exact-equivalence claims to rSVD, "
            "float32, CUDA, or Metal, which remain governed by their separate "
            "numerical validations."
        ),
    )

    append_provenance_row(document, ledger_row)
    document.save(SUPP_OUTPUT)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    ledger_row = write_updated_ledger()
    revise_main()
    revise_supplement(ledger_row)
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)
    print(NEW_LEDGER)


if __name__ == "__main__":
    main()
