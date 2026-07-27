#!/usr/bin/env python3

import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle52"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle53"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle52_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle52_0.99.6_20260726.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle53_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle53_0.99.6_20260726.docx"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def set_cell_margins(cell, top=120, start=150, bottom=120, end=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_width(table, width_dxa=9072):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    grid_col = OxmlElement("w:gridCol")
    grid_col.set(qn("w:w"), str(width_dxa))
    grid.append(grid_col)
    cell = table.cell(0, 0)
    tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        cell._tc.get_or_add_tcPr().append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Inches(6.3)


def add_pseudocode(document):
    anchor = find_paragraph(
        document,
        "For tall matrices with moderate",
    )

    introduction = insert_paragraph_after(
        anchor,
        (
            "Algorithm 1 maps these execution optimizations to the standard "
            "SIMPLS component update."
        ),
        style="Body Text",
    )

    caption = document.add_paragraph(style="Caption")
    caption.add_run(
        "Algorithm 1. Accelerated SIMPLS component path in fastPLS. "
        "Direction extraction is either deterministic IRLBA or approximate "
        "rSVD; subsequent score construction, orthogonalization, and "
        "deflation follow de Jong [11]."
    )
    caption.paragraph_format.keep_with_next = True

    table = document.add_table(rows=1, cols=1)
    table.style = "Table"
    set_table_width(table)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    shade_cell(cell, "F3F6F8")

    lines = [
        ("Input:", " centred X; numeric, indicator, or label-aware response Y; maximum A; requested component-count set C; solver."),
        ("1.", " Define S0 = X'Y explicitly, or as products v -> X'(Yv) and u -> Y'(Xu); label-aware products use class sums."),
        ("2.", " If p <= n and storage permits, cache G = X'X. Initialize R, Q, and V as empty; set fitted values Yhat0 = 0."),
        ("3.", " For a = 1,...,A:"),
        ("3.1", " Extract the leading right direction ga of the current state S(a-1) using a fresh IRLBA solve or a fresh oversampled rSVD sketch."),
        ("3.2", " Set ra = S(a-1) ga and ta = X ra; divide ra and ta by ||ta||2."),
        ("3.3", " Compute predictor and response loadings pa = X' ta and qa = Y' ta."),
        ("3.4", " Orthogonalize va = (I - V V') pa and normalize va."),
        ("3.5", " Deflate S(a) = S(a-1) - va[va' S(a-1)], or update the equivalent implicit operator."),
        ("3.6", " Append ra, qa, and va; update Ba = R(1:a) Q(1:a)' and Yhata = Yhat(a-1) + ta qa'."),
        ("3.7", " If a is in C, store only the requested coefficient/prediction snapshot or compact latent representation."),
        ("Output:", " one maximal component path with requested prefixes, rather than independently refitted models."),
    ]

    cell.paragraphs[0].clear()
    for index, (label, text) in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(1.5)
        paragraph.paragraph_format.line_spacing = 1.0
        if label.startswith("3."):
            paragraph.paragraph_format.left_indent = Inches(0.18)
        label_run = paragraph.add_run(label)
        label_run.bold = True
        body_run = paragraph.add_run(text)
        for run in (label_run, body_run):
            run.font.name = "Courier New"
            run.font.size = Pt(7.8)

    introduction._p.addnext(caption._p)
    caption._p.addnext(table._tbl)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    document = Document(MAIN_SOURCE)
    add_pseudocode(document)
    document.save(MAIN_OUT)
    shutil.copy2(SUPP_SOURCE, SUPP_OUT)
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
