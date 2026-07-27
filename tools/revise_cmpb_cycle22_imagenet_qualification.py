#!/usr/bin/env python3

from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle21"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle22"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle21_0.99.6_20260725.docx"
SUPP_SOURCE = (
    SOURCE / "fastPLS_CMPB_supplement_cycle21_0.99.6_20260725.docx"
)
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle21_20260725.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle22_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle22_0.99.6_20260725.docx"
RESPONSE_OUT = (
    OUT / "fastPLS_CMPB_response_to_reviewers_cycle22_20260725.docx"
)

spec = spec_from_file_location(
    "cycle21_helpers",
    ROOT / "tools" / "revise_cmpb_cycle21_memory_baselines.py",
)
c21 = module_from_spec(spec)
spec.loader.exec_module(c21)
c16 = c21.c16


def replace_paragraph(document, prefix, text):
    paragraph = c16.find_paragraph(document, prefix)
    paragraph.text = text
    return paragraph


def revise_main():
    document = Document(MAIN_SOURCE)

    replace_paragraph(
        document,
        "The main biomedical benchmark included",
        "The main biomedical benchmark included twelve tasks spanning "
        "metabolomics, NMR spectroscopy, cancer and tissue omics, CITE-seq, "
        "Retina and Tabula Muris single-cell transcriptomics, and drug "
        "response. CIFAR-100 followed its documented 50,000/10,000 split [28]. "
        "A separate exploratory ImageNet/DINOv2 stress test used a pooled "
        "archive of 1,281,167 precomputed 1,024-dimensional embeddings from "
        "1,000 classes [26,27]. The archive did not retain an authoritative "
        "canonical train/validation flag. With seed 123, 1,000,000 row indices "
        "were sampled without replacement for development training and the "
        "complementary 281,167 rows formed a development holdout. This is not "
        "the standard ImageNet train/validation split. It assessed only the "
        "supervised representation stage after feature extraction and was not "
        "interpreted as biomedical or external predictive validation. Exact "
        "dimensions, checksums, split construction, and limitations are "
        "reported in the Supplementary Material.",
    )

    replace_paragraph(
        document,
        "The ImageNet representation experiment used",
        "The ImageNet representation experiment used the same fixed "
        "1,000,000/281,167 development split and float32 DINOv2 embeddings for "
        "every representation. Training and holdout indices were disjoint and "
        "exhausted all 1,281,167 rows. PCA and PLS were fitted using development-"
        "training rows only; the complementary rows were projected only after "
        "fitting, and FAISS indices contained training rows only. Raw embeddings "
        "were compared with PCA-rSVD and label-aware PLS-SVD/rSVD scores at 50, "
        "100, and 200 dimensions using exact CUDA cosine kNN with k=10. These "
        "dimensions and k were fixed before the matched raw/PCA/PLS run, but "
        "they were informed by earlier exploratory analyses on the same "
        "development holdout rather than selected by nested validation. The "
        "holdout is therefore not an untouched confirmatory test set. Two "
        "additional rSVD seeds repeated complete PLS and PCA fits; these repeats "
        "quantify randomized-algorithm and timing variation, not sampling or "
        "model-selection uncertainty. Timings separated fitting, train/holdout "
        "projection, query, inference, and end-to-end work. CUDA IVF recall@10 "
        "was evaluated against exact neighbours within each representation.",
    )

    replace_paragraph(
        document,
        "ImageNet/DINOv2 was used for a million-sample",
        "ImageNet/DINOv2 was used as a million-sample post-extraction stress "
        "test and exploratory supervised-reduction analysis, not as biomedical "
        "or external predictive validation. On the fixed development holdout, "
        "raw 1,024-dimensional embeddings gave top-1/top-5 accuracy "
        "0.6556/0.9392, whereas the 200-dimensional PLS endpoint gave median "
        "0.6521/0.9397 across three rSVD seeds. PLS therefore reduced dimension "
        "5.12-fold and held-out projection-plus-query time approximately "
        "3.9-fold, with a 0.35-percentage-point top-1 loss. The observed top-5 "
        "difference was only 0.0005 in proportion (0.05 percentage points); its "
        "conditional Wilson intervals overlapped, and no independent "
        "model-selection or external-validation uncertainty is available. It is "
        "therefore reported as descriptive and is not interpreted as an "
        "improvement. The 50-, 100-, and 200-dimensional path and k=10 were not "
        "nested-selected, and the same development holdout had informed earlier "
        "exploration. PCA remains the unsupervised dimension-matched control. "
        "Table 2 and Figure 4 present computational and compression trade-offs, "
        "not confirmatory accuracy rankings.",
    )

    replace_paragraph(
        document,
        "Table 2. Exploratory matched ImageNet/DINOv2 retrieval",
        "Table 2. Exploratory matched ImageNet/DINOv2 retrieval on a fixed "
        "random 1,000,000/281,167 development split drawn without replacement "
        "from a pooled 1,281,167-image embedding archive; this is not the "
        "canonical ImageNet train/validation split. Exact CUDA cosine kNN used "
        "k=10. The value of k and the 50/100/200-dimensional path were fixed for "
        "this matched run but informed by earlier exploration on the same "
        "development holdout, not nested validation. Brackets are Wilson 95% "
        "intervals conditional on this split. They do not account for split or "
        "hyperparameter-selection uncertainty, and the small top-5 difference "
        "is not interpreted as improvement. Transformation includes fitting and "
        "train/holdout projection; query is median (IQR) over three exact FAISS "
        "runs. H/G are peak host RSS and sampled GPU memory in MB.",
    )

    replace_paragraph(
        document,
        "Figure 4. Exploratory matched ImageNet/DINOv2 retrieval.",
        "Figure 4. Exploratory matched ImageNet/DINOv2 retrieval on the "
        "noncanonical random development split. (A) Descriptive top-1 and top-5 "
        "accuracy for raw embeddings and 50-, 100-, and 200-dimensional PCA and "
        "PLS representations. (B) Holdout projection plus exact FAISS query "
        "time. (C) End-to-end representation fitting, train/holdout projection, "
        "and query time. (D) Peak host RSS and sampled GPU memory. The displayed "
        "path uses seed 123; two additional seeded PLS/PCA fits are summarized "
        "in Supplementary Table S23. Repeated seeds do not replace independent "
        "split or hyperparameter-selection uncertainty.",
    )

    replace_paragraph(
        document,
        "The ImageNet control separates representation cost",
        "The ImageNet control separates representation cost from retrieval "
        "cost. It supports computational feasibility and a possible supervised-"
        "compression trade-off, not an accuracy advantage: PLS was slightly "
        "worse in top-1, and its 0.05-percentage-point top-5 difference from raw "
        "features is not interpreted as improvement. The PCA control separates "
        "supervised response information from dimensionality reduction alone. "
        "Because the split was noncanonical and its development holdout had "
        "informed prior exploration, an independently defined ImageNet "
        "validation set with training-only hyperparameter selection would be "
        "required for confirmatory predictive claims. These natural-image "
        "results do not establish biomedical predictive utility.",
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - qualified ImageNet development analysis"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)

    replace_paragraph(
        document,
        "CIFAR-100 and ImageNet are evaluated",
        "CIFAR-100 and ImageNet were evaluated as precomputed image embeddings "
        "rather than by training image encoders. CIFAR-100 used its standard "
        "50,000/10,000 partition [29]. The ImageNet archive contained 1,281,167 "
        "rows, 1,024 DINOv2 features, and 1,000 labels [27,28], but no "
        "authoritative canonical train/validation flag. The feature archive was "
        "reformatted without normalization or dimensionality reduction. Its "
        "exact DINOv2 checkpoint and pooling rule were not retained in "
        "independently auditable metadata, which limits representation-level "
        "reproducibility and is stated explicitly. Feature extraction was "
        "completed before the fastPLS benchmark and excluded from timing.",
    )

    replace_paragraph(
        document,
        "The external retrieval benchmark used",
        "The external retrieval benchmark used one fixed non-stratified split "
        "created by set.seed(123), sampling 1,000,000 of 1,281,167 pooled rows "
        "without replacement and assigning the complementary 281,167 rows to a "
        "development holdout. The sets had zero overlap and their union covered "
        "the archive. This is not the canonical ImageNet train/validation split. "
        "PLS-SVD and PCA were fitted only on development-training rows; holdout "
        "rows were projected after fitting, and FAISS indices contained training "
        "rows only. The maximum 200-dimensional randomized fit supplied 50- and "
        "100-dimensional prefixes. Exact CUDA cosine kNN used k=10 for all raw, "
        "PCA, and PLS representations. Although fixed before this matched run, "
        "k=10 and the component path were informed by earlier exploration on "
        "the same holdout rather than nested training-only selection. Complete "
        "PLS and PCA fits were repeated with three rSVD seeds; this measures "
        "algorithmic variation, not independent predictive uncertainty. IVF "
        "recall@10 compares approximate and exact neighbours within one "
        "representation and is not class recall. All accuracy results are "
        "descriptive.",
    )

    caption = c16.find_paragraph(
        document, "Table S18. Complete matched ImageNet retrieval results."
    )
    caption.text = (
        "Table S18. Complete exploratory ImageNet retrieval results on the "
        "noncanonical random development split. Query time is median (IQR); H/G "
        "are peak host/GPU memory in MB. k=10 and dimensions 50/100/200 were not "
        "nested-selected. Accuracy differences are descriptive."
    )

    uncertainty = c16.find_paragraph(
        document,
        "Table S31. Wilson 95% score intervals",
    )
    uncertainty.text = (
        "Table S31. Wilson 95% score intervals for the representative seed-123 "
        "ImageNet/DINOv2 retrieval paths. The intervals describe binomial "
        "uncertainty conditional on the fixed 281,167-image development holdout. "
        "They do not account for its reuse during prior exploration, "
        "hyperparameter selection, alternative random splits, or uncertainty in "
        "the pretrained feature extractor. The 0.0005 absolute top-5 difference "
        "between 200-dimensional PLS and raw features is not interpreted as an "
        "improvement."
    )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(
        "S28. ImageNet split provenance, leakage audit, and interpretation",
        level=1,
    )
    document.add_paragraph(
        "The source archive SHA-256 was "
        "b85fa6bdf7414216cf3f8f03c7d6862f881f1d8f61f61c85e2b6074f9bd5e2eb. "
        "The task metadata SHA-256 was "
        "f58ba6a2193c7f78a29d96a42b766d325852400a74a73a2b3be9251464655935. "
        "The task was generated on 22 July 2026 by "
        "benchmark/prepare_imagenet_float32_task.R. The archived split indices "
        "permit exact reconstruction. Training class counts ranged from 569 to "
        "1,058 and holdout counts from 157 to 333 because the split was random "
        "rather than stratified."
    )
    c16.add_table(
        document,
        ["Item", "Recorded value"],
        [
            ("Source archive", "imagenet_float32.RData"),
            ("Pooled observations", "1,281,167"),
            ("Features/classes", "1,024 / 1,000"),
            ("Split rule", "seed 123; 1,000,000 sampled without replacement"),
            ("Development holdout", "Complementary 281,167 rows"),
            ("Canonical ImageNet split", "No"),
            ("Train/holdout overlap", "0 rows"),
            ("Train/holdout union", "1,281,167 rows"),
            ("Retrieval setting", "Exact CUDA cosine kNN; k=10"),
            ("Dimensions", "Raw 1,024; PCA/PLS 50, 100, 200"),
            (
                "Selection status",
                "Informed by prior exploration; no nested validation",
            ),
            (
                "Interpretation",
                "Computational stress test and descriptive compression study",
            ),
        ],
        font_size=7.0,
    )
    document.add_paragraph(
        "Leakage audit. There was no row-level overlap: PLS labels and PCA "
        "fitting used training rows only, holdout rows were transformed after "
        "model fitting, and nearest-neighbour search queried a training-only "
        "index. However, the same holdout labels had been examined in earlier "
        "ImageNet development experiments. This constitutes model-development "
        "reuse at the workflow level even though no holdout row entered PLS/PCA "
        "fitting. Consequently, the term development holdout is used throughout "
        "and no unbiased external-generalization claim is made. Possible overlap "
        "between the pretrained DINOv2 corpus and ImageNet cannot be audited from "
        "the archived feature matrix."
    )
    document.add_paragraph(
        "Hyperparameter and uncertainty audit. k=10 and dimensions 50, 100, and "
        "200 were fixed for the final matched comparison, ensuring that raw, PCA, "
        "and PLS used identical retrieval settings. They were nevertheless "
        "informed by earlier experiments rather than prespecified independently "
        "or selected inside a nested training-only validation procedure. Wilson "
        "intervals quantify only finite-holdout binomial uncertainty. Repeated "
        "rSVD seeds quantify randomized representation variation, and repeated "
        "FAISS calls quantify runtime variation; neither resolves selection or "
        "split uncertainty. In particular, the observed PLS-minus-raw top-5 "
        "difference of 0.0005 is not evidence of improvement."
    )
    document.add_paragraph(
        "Machine-readable provenance is supplied as "
        "imagenet_split_provenance.csv, and the complete limitation statement is "
        "also stored in benchmark/IMAGENET_SPLIT_PROVENANCE.md."
    )

    document.core_properties.title = (
        "fastPLS CMPB supplement - qualified ImageNet development analysis"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    document.add_heading(
        "23. ImageNet split provenance, hyperparameters, and uncertainty",
        level=1,
    )
    document.add_paragraph(
        "Reviewer comment: The 1,000,000/281,167 split is not the standard "
        "ImageNet train/validation split and requires exact provenance and a "
        "leakage description. k=10 and the tested dimensionalities must be "
        "prespecified or selected without held-out labels. The small top-5 "
        "difference must not be interpreted as improvement without uncertainty."
    )
    document.add_paragraph(
        "Response: We agree and have corrected both the reporting and the "
        "interpretation. The archived source contains 1,281,167 pooled "
        "1,024-dimensional embeddings and does not retain an authoritative "
        "canonical train/validation flag. The benchmark task used seed 123 to "
        "sample 1,000,000 indices without replacement; the complementary "
        "281,167 rows formed a non-stratified development holdout. The sets have "
        "zero overlap and exhaust the archive. We now report the source and task "
        "SHA-256 checksums, preparation time, exact split algorithm, class-count "
        "ranges, and reconstruction script in Supplementary Section S28 and a "
        "machine-readable provenance file."
    )
    document.add_paragraph(
        "The leakage audit distinguishes row-level fitting separation from "
        "workflow-level reuse. No holdout row or label entered PCA/PLS fitting, "
        "and FAISS indices contained training rows only. However, the same "
        "holdout had informed earlier exploratory ImageNet analyses. We therefore "
        "renamed it a development holdout and no longer describe the analysis as "
        "external or confirmatory validation. We also disclose that the archived "
        "features do not permit independent auditing of the exact DINOv2 "
        "checkpoint, pooling rule, or pretraining-corpus overlap."
    )
    document.add_paragraph(
        "k=10 and dimensions 50, 100, and 200 were fixed before the final matched "
        "raw/PCA/PLS run, but they were informed by earlier exploration rather "
        "than independently prespecified or selected by nested training-only "
        "validation. This is now stated in Methods, Results, captions, and the "
        "Supplement. The 200-dimensional PLS endpoint differed from raw features "
        "by +0.0005 in top-5 proportion (0.05 percentage points), with overlapping "
        "conditional Wilson intervals. We no longer call this an improvement. "
        "Repeated rSVD seeds are described only as algorithmic variation, not "
        "sampling or model-selection uncertainty. The ImageNet evidence is now "
        "limited to computational feasibility and an exploratory compression-time "
        "trade-off."
    )
    document.core_properties.title = (
        "fastPLS CMPB response - qualified ImageNet development analysis"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    revise_response()
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
