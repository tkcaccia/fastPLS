#!/usr/bin/env python3

from pathlib import Path
from shutil import copy2

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle13"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle14"
POWER1 = (
    ROOT
    / "benchmark_results"
    / "simpls_estimator_preservation_reliable_20260725_1735"
    / "simpls_estimator_approximation_rsvd.csv"
)
POWER2 = (
    ROOT
    / "benchmark_results"
    / "simpls_estimator_preservation_reliable_power2_final_20260725"
    / "simpls_estimator_approximation_rsvd.csv"
)
CUDA = ROOT / "benchmark_results" / "rsvd_cuda_reliability_20260725.csv"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle13_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle13_0.99.6_20260725.docx"
RESPONSE_SOURCE = SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle13_20260725.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle14_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle14_0.99.6_20260725.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle14_20260725.docx"


def find_paragraph(document, start):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(start):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {start}")


def set_text(document, start, text):
    paragraph = find_paragraph(document, start)
    paragraph.text = text
    return paragraph


def style_table(table):
    table.style = "Table"
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7.5)


def add_table(document, columns, rows):
    table = document.add_table(rows=1, cols=len(columns))
    for index, label in enumerate(columns):
        table.rows[0].cells[index].text = label
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    style_table(table)
    return table


def summarize_cpu(path, setting):
    data = pd.read_csv(path)
    return (
        setting,
        f"{int(data['approximation_tolerance_pass'].sum())}/{len(data)}",
        f"{data['prediction_relative_error'].max():.4f}",
        f"{data['prediction_correlation'].min():.5f}",
        f"{data['score_subspace_max_angle_degrees'].max():.2f}",
        f"{data['classification_label_agreement'].min():.3f}",
        f"{data['metric_absolute_difference'].max():.3f}",
    )


def revise_main():
    document = Document(MAIN_SOURCE)
    set_text(
        document,
        "Methods:",
        "Methods: fastPLS provides PLS-SVD, SIMPLS, orthogonal PLS (OPLS), "
        "and kernel PLS through one R interface. Deterministic IRLBA SIMPLS "
        "preserves de Jong's sequential estimator while reusing deflation, "
        "coefficient, and prediction state across component prefixes. rSVD is "
        "treated as a stochastic approximation and is subject to prespecified "
        "prediction, subspace, label-agreement, and predictive-metric failure "
        "criteria. The package combines implicit cross-covariance products, "
        "compact prediction, compiled validation, double-precision reference "
        "paths, conditional float32 paths, multithread-capable CPU execution, "
        "NVIDIA CUDA, and Apple Metal."
    )
    set_text(
        document,
        "Results:",
        "Results: Deterministic IRLBA SIMPLS passed all 117 component-level "
        "comparisons with de Jong SIMPLS. An initial one-power rSVD setting "
        "failed 16 of 117 approximation comparisons and was rejected for "
        "confirmatory use. After removal of the one-vector warm-start shortcut, "
        "oversampling by 10 directions with two power iterations passed all "
        "117 prespecified CPU comparisons; maximum relative prediction error "
        "was 0.0332, minimum prediction correlation was 0.99939, maximum "
        "score-subspace angle was 4.93 degrees, and minimum label agreement "
        "was 0.99. Performance benchmarks therefore report rSVD as an "
        "approximate workflow, whereas estimator-preservation claims rely on "
        "deterministic IRLBA."
    )
    set_text(
        document,
        "The SIMPLS estimator follows",
        "The SIMPLS estimator follows de Jong's sequential construction [11]. "
        "At component k, the dominant direction of the current cross-covariance "
        "state is converted to a score and loading pair, orthogonalized against "
        "the preceding SIMPLS basis, and used in the standard rank-one "
        "deflation. When IRLBA is requested, each deflated update uses a fresh "
        "deterministic iterative solve. When rSVD is requested, each update uses "
        "a fresh oversampled randomized range finder on the current deflated "
        "operator; the rejected one-vector warm start and adaptive refresh are "
        "not used. In both cases, deflation terms, latent quantities, and "
        "prediction prefixes are retained incrementally, avoiding independent "
        "refits for each requested component count."
    )
    set_text(
        document,
        "Double-precision CPU fitting supports",
        "Double-precision CPU fitting supports bundled augmented implicitly "
        "restarted Lanczos bidiagonalization (IRLBA) [15] and randomized SVD "
        "(rSVD) [16]. IRLBA builds and restarts a bidiagonal Krylov subspace. "
        "rSVD draws a Gaussian range sketch, applies optional power iterations, "
        "orthonormalizes the range, and decomposes a reduced matrix. Exact SVD "
        "is used only when either matrix dimension is below six. IRLBA is the "
        "deterministic reference for confirmatory interpretation. The faster "
        "one-power rSVD setting is exploratory; difficult spectra may require "
        "additional power iterations or wider oversampling, and all important "
        "rSVD conclusions should be checked across seeds or against IRLBA. CUDA "
        "and Metal support rSVD, whereas bundled IRLBA is available on the CPU."
    )
    set_text(
        document,
        "Estimator preservation was evaluated",
        "Estimator preservation and randomized approximation were evaluated "
        "separately in a prespecified study. Eight synthetic regimes covered "
        "regression and classification, p<n and p>n, low- and high-rank "
        "responses, ill-conditioned predictors, and exact rank deficiency; four "
        "real tasks comprised Breast, Colon, MetRef, and an NMR spectral subset. "
        "Deterministic IRLBA was assessed against pls::simpls.fit. An rSVD row "
        "was classified as failed when relative prediction error exceeded 0.05, "
        "prediction correlation was below 0.99, any score/projection/loading "
        "subspace angle exceeded 10 degrees, label agreement was below 0.99, or "
        "the predictive metric differed by more than 0.01. Numerical completion "
        "alone was never counted as approximation success."
    )
    set_text(
        document,
        "rSVD was not used as estimator-equivalence evidence.",
        "rSVD was not used as estimator-equivalence evidence. Reanalysis exposed "
        "a reliability defect in the earlier warm-started one-vector direction "
        "shortcut: the original audit reached relative prediction error 1.06, "
        "prediction correlation 0.445, score-subspace angle 88.4 degrees, and "
        "label agreement 0.133. That shortcut was removed from the public CPU "
        "and CUDA SIMPLS routes. With a fresh oversampled range sketch at each "
        "deflation, oversampling 10 and one power iteration passed 101/117 "
        "comparisons; failures were confined to high-rank-response simulation "
        "and MetRef. Increasing to two power iterations passed 117/117, with "
        "worst relative prediction error 0.0332, minimum correlation 0.99939, "
        "maximum score-subspace angle 4.93 degrees, and minimum label agreement "
        "0.99. A focused CUDA audit required either four power iterations at "
        "oversampling 10 or oversampling 20 for all audited points to pass. "
        "Accordingly, IRLBA is recommended for confirmatory inference, "
        "ill-conditioned or rank-deficient matrices, slowly decaying spectra, "
        "or unstable rSVD results."
    )
    set_text(
        document,
        "Among the 46 completed",
        "The selected-point rSVD benchmark characterizes complete stochastic "
        "workflows rather than deterministic estimator equivalence. CPU was "
        "generally preferable for small matrices, where transfer and "
        "kernel-launch overhead exceeded accelerated work, whereas CUDA became "
        "advantageous for large dense products, high component counts, or large "
        "response dimensions. Every rSVD fit is now labelled as unaudited until "
        "it is compared across seeds or with IRLBA; rows failing the prespecified "
        "criteria are retained as failures rather than summarized as successful "
        "runs."
    )
    set_text(
        document,
        "Table 1. Twelve-task",
        "Table 1. Twelve-task biomedical workflow benchmark at the component "
        "count selected from training data. Each cell shows the completed "
        "CPU/CUDA rSVD workflow, effective component count, outer-test metric, "
        "median total fitting-plus-prediction time, peak host RSS, sampled GPU "
        "memory, precision, and completed runs. rSVD is stochastic and these "
        "rows are performance measurements, not evidence of equivalence with "
        "deterministic IRLBA."
    )
    document.core_properties.title = (
        "fastPLS CMPB manuscript - rSVD reliability revision cycle 14"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)
    set_text(
        document,
        "The next randomized direction",
        "Each randomized SIMPLS direction is generated from a fresh oversampled "
        "range sketch of the current deflated cross-covariance operator."
    )
    set_text(
        document,
        "These changes reorganize numerical work",
        "Incremental deflation, coefficient, and prediction updates reorganize "
        "numerical work while retaining sequential SIMPLS orthogonalization and "
        "deflation. The former one-vector warm-start and adaptive randomized "
        "refresh were rejected by formal validation and are not used by the "
        "public algorithm."
    )
    set_text(
        document,
        "For component ",
        "For component k, obtain one dominant left direction of the current "
        "deflated cross-covariance operator. IRLBA performs a fresh deterministic "
        "iterative solve. rSVD forms a fresh Gaussian sketch with target width "
        "one plus oversampling, applies the requested power iterations, performs "
        "QR orthonormalization, and extracts the leading vector from the reduced "
        "decomposition."
    )
    set_text(
        document,
        "For an operator ",
        "For an operator A and target rank k, rSVD draws a Gaussian matrix, forms "
        "a range sketch, optionally applies alternating products with A and its "
        "transpose, orthonormalizes the range, and decomposes the reduced matrix. "
        "The public one-power setting is intended for exploratory speed. It must "
        "not be interpreted as certified agreement; power, oversampling, and seed "
        "are recorded in every benchmark manifest."
    )
    set_text(
        document,
        "For deterministic validation",
        "For deterministic validation, every deflated component requests a fresh "
        "IRLBA direction. For randomized validation, every component requests a "
        "fresh oversampled rSVD range sketch. No preceding one-vector direction "
        "is reused."
    )
    set_text(
        document,
        "The rSVD results are deliberately",
        "The earlier rSVD audit exposed unacceptable divergence from the "
        "deterministic reference: maximum relative prediction error 1.06, minimum "
        "prediction correlation 0.445, maximum score-subspace angle 88.4 degrees, "
        "minimum label agreement 0.133, and maximum predictive-metric difference "
        "0.526. Investigation traced this behavior to a warm-started one-vector "
        "shortcut, which was removed. Corrected validation and explicit failure "
        "criteria are reported in Section S23."
    )

    document.add_heading("S23. rSVD numerical reliability and failure criteria", level=1)
    document.add_paragraph(
        "A completed randomized fit is not automatically a successful "
        "approximation. Relative prediction error, prediction correlation, "
        "score/projection/loading principal angles, classification-label "
        "agreement, and predictive-metric difference were compared with a "
        "deterministic CPU IRLBA fit on identical data, preprocessing, split, "
        "component count, and prediction head."
    )
    document.add_paragraph(
        "Failure was prespecified as prediction relative error >0.05, prediction "
        "correlation <0.99, any latent-subspace angle >10 degrees, label "
        "agreement <0.99, or absolute predictive-metric difference >0.01. "
        "Structural diagnostics in the fitted object check finite factors and "
        "effective rank, but report approximation_not_audited until an external "
        "agreement audit is performed."
    )
    document.add_paragraph(
        "Table S26. Corrected CPU SIMPLS-rSVD validation against deterministic "
        "IRLBA across 117 component-level comparisons."
    )
    add_table(
        document,
        [
            "Setting",
            "Passed",
            "Max rel. pred. error",
            "Min corr.",
            "Max score angle",
            "Min label agreement",
            "Max metric diff.",
        ],
        [
            summarize_cpu(POWER1, "oversample=10, power=1"),
            summarize_cpu(POWER2, "oversample=10, power=2"),
        ],
    )

    cuda = pd.read_csv(CUDA)
    cuda_rows = []
    for keys, group in cuda.groupby(["oversample", "power"], sort=True):
        cuda_rows.append(
            (
                f"oversample={keys[0]}, power={keys[1]}",
                f"{int(group['approximation_tolerance_pass'].sum())}/{len(group)}",
                f"{group['prediction_relative_error'].max():.4f}",
                f"{group['prediction_correlation'].min():.5f}",
                f"{group['label_agreement'].min():.3f}",
                f"{group['metric_absolute_difference'].max():.3f}",
            )
        )
    document.add_paragraph(
        "Table S27. Focused CUDA audit on high-rank-response regression and "
        "MetRef. The audit is smaller than the CPU suite and provides setting "
        "guidance rather than a universal guarantee."
    )
    add_table(
        document,
        [
            "Setting",
            "Passed",
            "Max rel. pred. error",
            "Min corr.",
            "Min label agreement",
            "Max metric diff.",
        ],
        cuda_rows,
    )
    document.add_paragraph(
        "IRLBA should be preferred for confirmatory inference, coefficient or "
        "loading interpretation, ill-conditioned or rank-deficient matrices, "
        "slow singular-value decay, instability across rSVD seeds, or any "
        "failed approximation criterion. rSVD remains appropriate when "
        "accelerator execution or speed is important and its approximation has "
        "been validated for the target task."
    )
    document.core_properties.title = (
        "fastPLS CMPB supplement - rSVD reliability revision cycle 14"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    document.add_heading("15. rSVD reliability and failure criteria", level=2)
    document.add_paragraph(
        "We agree that the earlier results were unacceptable. The benchmark had "
        "incorrectly labelled every numerically completed rSVD fit as successful, "
        "and the public SIMPLS path used a warm-started one-vector shortcut rather "
        "than an oversampled range finder. We removed that shortcut from CPU and "
        "CUDA SIMPLS and now generate a fresh oversampled sketch at every "
        "deflation. We added structural diagnostics to fitted models, singular-"
        "triplet residual diagnostics to fastsvd(), and prespecified approximation "
        "failure criteria. With oversampling 10 and one power iteration, the "
        "corrected CPU route passed 101/117 comparisons; those 16 failures remain "
        "reported. With two power iterations it passed 117/117, with maximum "
        "relative prediction error 0.0332, minimum correlation 0.99939, maximum "
        "score-subspace angle 4.93 degrees, and minimum label agreement 0.99. A "
        "focused CUDA audit showed that stronger power or oversampling was needed "
        "for all audited points to pass. The manuscript now states that rSVD is a "
        "stochastic workflow approximation and reserves estimator-preservation "
        "claims for deterministic IRLBA. It gives explicit guidance to prefer "
        "IRLBA for confirmatory inference, ill-conditioning, rank deficiency, "
        "slow spectral decay, seed instability, or any failed criterion."
    )
    document.core_properties.title = (
        "Response to reviewers - rSVD reliability revision cycle 14"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    revise_response()
    copy2(POWER1, OUT / "simpls_rsvd_cpu_power1.csv")
    copy2(POWER2, OUT / "simpls_rsvd_cpu_power2.csv")
    copy2(CUDA, OUT / "simpls_rsvd_cuda_setting_audit.csv")
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
