from pathlib import Path
from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "artifacts/CMPB_rewrite_20260826_cycle115"
OUT = ROOT / "artifacts/CMPB_rewrite_20260826_cycle116"
OUT.mkdir(parents=True, exist_ok=True)


def replace_once(document, old, new):
    for paragraph in document.paragraphs:
        if old in paragraph.text:
            for run in paragraph.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new, 1)
                    return
            # The target may span runs. Preserve paragraph-level style for this
            # narrowly scoped prose edit.
            paragraph.text = paragraph.text.replace(old, new, 1)
            return
    raise RuntimeError(f"Text not found: {old}")


main = Document(SRC / "fastPLS_CMPB_main_cycle115_0.99.25_20260826.docx")
replace_once(
    main,
    "The experiment used Breast, MetRef, and CIFAR-100 with 10, 22, and 50 components, respectively.",
    "The experiment used Breast, MetRef, and CIFAR-100 with 10, 22, and 50 components, respectively. IKPLS was not evaluated on the NMR or ImageNet case studies; therefore, no cross-language performance or scalability conclusion is drawn for either large-scale task.",
)
replace_once(
    main,
    "The archived CPU comparison completed all 36 planned runs.",
    "The archived CPU comparison, restricted to Breast, MetRef, and CIFAR-100 and excluding NMR and ImageNet, completed all 36 planned runs.",
)
main.save(OUT / "fastPLS_CMPB_main_cycle116_0.99.25_20260826.docx")


supp = Document(SRC / "fastPLS_CMPB_supplement_cycle115_0.99.25_20260826.docx")
replace_once(
    supp,
    "The common CPU contract comprised float64 input, externally applied training centring, identical centred one-hot responses, identical splits and component counts, final held-out prediction, three fresh-process repetitions, and one effective thread.",
    "The common CPU contract comprised float64 input, externally applied training centring, identical centred one-hot responses, identical splits and component counts, final held-out prediction, three fresh-process repetitions, and one effective thread. The comparison was restricted to Breast, MetRef, and CIFAR-100; IKPLS was not evaluated on NMR or ImageNet, and no cross-language conclusion is drawn for those case studies.",
)
supp.save(OUT / "fastPLS_CMPB_supplement_cycle116_0.99.25_20260826.docx")
