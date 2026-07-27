from pathlib import Path
import math

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import revise_cmpb_cycle9_simpls_validation as c9


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle11"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle12"
RESULTS = ROOT / "benchmark_results" / "imagenet_faiss_matched_1m_20260725"
PLOT = RESULTS / "plots" / "imagenet_matched_retrieval.png"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle11_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle11_0.99.6_20260725.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle12_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle12_0.99.6_20260725.docx"
RESPONSE_OUT = OUT / "response_imagenet_retrieval_20260725.docx"


def number(value, digits=3):
    try:
        value = float(value)
    except (TypeError, ValueError):
        return "NA"
    if not math.isfinite(value):
        return "NA"
    return f"{value:.{digits}f}"


def integer(value):
    try:
        value = float(value)
    except (TypeError, ValueError):
        return "NA"
    if not math.isfinite(value):
        return "NA"
    return f"{value:,.0f}"


def representation(row):
    return {
        "raw_dinov2": "Raw DINOv2",
        "pca_scores": "PCA-rSVD",
        "pls_scores": "PLS-SVD/rSVD",
    }[row["feature_space"]]


def move_after(anchor, element):
    anchor._p.addnext(element)


def add_main_table(document, anchor, table_data):
    rows = []
    for _, row in table_data.iterrows():
        rows.append((
            representation(row),
            integer(row["n_features"]),
            number(row["compression_ratio"], 2),
            number(row["top1_accuracy"], 4),
            number(row["top5_accuracy"], 4),
            number(row["transformation_time_sec"], 1),
            f"{number(row['query_time_median_sec'], 1)} "
            f"({number(row['query_time_iqr_sec'], 1)})",
            number(row["inference_time_median_sec"], 1),
            f"{integer(row['peak_host_rss_mb'])}/"
            f"{integer(row['peak_gpu_mem_mb'])}",
            number(row["ivf_neighbour_recall_at_10"], 4),
        ))
    caption = document.add_paragraph(
        "Table 2. Exploratory matched ImageNet/DINOv2 retrieval on the fixed "
        "1,000,000/281,167 train/test split. Exact CUDA cosine kNN (k=10) provides "
        "top-1 and top-5 accuracy. Transformation includes representation fitting "
        "and both train/test projection; query is median (IQR) over three exact "
        "FAISS runs. Inference comprises held-out projection plus exact query. "
        "H/G are peak host RSS and sampled GPU memory in MB. IVF recall@10 is "
        "measured against exact neighbours in the same representation from one "
        "exploratory IVF run.",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True
    table = c9.add_table(
        document,
        ["Representation", "Dim.", "1024/\ndim.", "Top-1", "Top-5",
         "Transform\n(s)", "Query\n(s)", "Inference\n(s)", "H/G\n(MB)",
         "IVF\nR@10"],
        rows,
        [1.02, 0.43, 0.48, 0.48, 0.48, 0.62, 0.72, 0.67, 0.78, 0.52],
        font_size=5.9,
    )
    move_after(anchor, caption._p)
    caption._p.addnext(table._tbl)
    return table


def add_main_figure(document, table, image):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(image), width=Inches(6.5))
    caption = document.add_paragraph(
        "Figure 3. Exploratory matched ImageNet/DINOv2 retrieval. (A) Top-1 and "
        "top-5 accuracy for raw embeddings and 50-, 100-, and 200-dimensional "
        "PCA and PLS representations. (B) Held-out transformation plus exact FAISS "
        "query time. (C) End-to-end representation fitting, train/test projection, "
        "and query time. (D) peak host RSS and sampled GPU memory. Search times "
        "are medians of three exact-index runs from one fixed representation fit; "
        "therefore the experiment is exploratory rather than an independent "
        "repeated-fit performance estimate.",
        style="Caption",
    )
    table._tbl.addnext(paragraph._p)
    paragraph._p.addnext(caption._p)


def image_net_sentences(table_data):
    raw = table_data[table_data["feature_space"] == "raw_dinov2"].iloc[0]
    pls = table_data[table_data["feature_space"] == "pls_scores"].sort_values(
        "top1_accuracy", ascending=False
    ).iloc[0]
    speedup = raw["inference_time_median_sec"] / pls["inference_time_median_sec"]
    return raw, pls, speedup


def revise_main(table_data):
    document = Document(MAIN_SOURCE)
    raw, pls, speedup = image_net_sentences(table_data)

    abstract = c9.find_paragraph(document, "Results:")
    c9.set_paragraph_text(
        abstract,
        abstract.text
        + " In an exploratory matched ImageNet/DINOv2 retrieval experiment with "
        "1,000,000 training and 281,167 held-out embeddings, the best tested "
        f"PLS-score representation ({int(pls['n_features'])} dimensions) achieved "
        f"top-1/top-5 accuracy of {pls['top1_accuracy']:.4f}/"
        f"{pls['top5_accuracy']:.4f}, compared with {raw['top1_accuracy']:.4f}/"
        f"{raw['top5_accuracy']:.4f} for 1,024-dimensional raw embeddings, while "
        f"reducing held-out transformation-plus-query time by {speedup:.1f}-fold."
    )

    design = c9.find_paragraph(document, "Estimator preservation was evaluated")
    protocol = c9.insert_after(
        design,
        "The ImageNet representation experiment used the fixed 1,000,000/281,167 "
        "split and float32 DINOv2 embeddings throughout. Raw embeddings were "
        "compared with unsupervised PCA-rSVD scores and supervised label-aware "
        "PLS-SVD/rSVD scores at 50, 100, and 200 dimensions. PCA and PLS used the "
        "same CUDA rSVD settings; PLS formed the centred predictor-response "
        "cross-covariance from class sums without constructing a one-million-by-"
        "1,000 dummy response. Exact FAISS cosine kNN (k=10) used identical labels, "
        "queries, distance weighting, and query blocks for all representations. "
        "CUDA IVF was evaluated against exact neighbours within the same "
        "representation, and neighbour recall@10 was the mean fraction of the ten "
        "exact neighbours recovered per query. Timings separated one-time fitting, "
        "training projection, held-out projection, index/query, inference, and "
        "end-to-end time. Exact index/query timing was repeated three times from "
        "one fixed representation fit; IVF and representation fitting were single "
        "runs, so all ImageNet predictive and end-to-end estimates are labelled "
        "exploratory.",
        style="Body Text",
    )

    result = c9.find_paragraph(document, "ImageNet/DINOv2 tested whether")
    c9.set_paragraph_text(
        result,
        "ImageNet/DINOv2 was used for two limited purposes: a million-sample "
        "computational stress test after foundation-model feature extraction and "
        "an exploratory test of PLS as supervised feature reduction. It was not "
        "used as biomedical validation. On the full fixed test set, exact CUDA "
        f"cosine kNN on raw 1,024-dimensional embeddings achieved top-1/top-5 "
        f"accuracy of {raw['top1_accuracy']:.4f}/{raw['top5_accuracy']:.4f}; "
        f"median query time was {raw['query_time_median_sec']:.1f} s. The best "
        f"tested PLS representation used {int(pls['n_features'])} dimensions "
        f"({pls['compression_ratio']:.2f}-fold compression) and achieved "
        f"{pls['top1_accuracy']:.4f}/{pls['top5_accuracy']:.4f}. Its one-time "
        f"fit and train/test projection required {pls['transformation_time_sec']:.1f} "
        f"s, whereas held-out projection plus exact query required "
        f"{pls['inference_time_median_sec']:.1f} s, a {speedup:.1f}-fold reduction "
        "relative to raw querying. PCA provided the unsupervised dimensionality-"
        "matched control. Approximate IVF recall was calculated only against exact "
        "neighbours from the corresponding raw, PCA, or PLS space. Table 2 and "
        "Figure 3 report the complete component path, memory, and timing endpoints. "
        "Because each representation was fitted once, these values are exploratory "
        "despite the three repeated exact-query timings."
    )
    table = add_main_table(document, result, table_data)
    add_main_figure(document, table, PLOT)

    discussion = c9.find_paragraph(document, "Float32 can reduce input")
    added = c9.insert_after(
        discussion,
        "The ImageNet control separates representation cost from retrieval cost. "
        "PLS did not automatically replace a strong raw-feature search on accuracy; "
        "its value is supported only when the observed accuracy is considered "
        "together with score dimension, held-out transformation time, query time, "
        "and memory. The PCA control is essential because it distinguishes a "
        "benefit of supervised response information from dimensionality reduction "
        "alone. These natural-image results establish computational feasibility "
        "and a supervised-compression hypothesis, not biomedical predictive utility.",
        style="Body Text",
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - completed ImageNet retrieval"
    )
    document.save(MAIN_OUT)


def revise_supplement(summary):
    document = Document(SUPP_SOURCE)
    document.add_page_break()
    document.add_heading(
        "S16. Exploratory matched ImageNet/DINOv2 retrieval", level=1
    )
    document.add_paragraph(
        "The external retrieval benchmark used the same fixed float32 training and "
        "held-out matrices for every representation. PLS-SVD used a label-aware "
        "class-sum cross-covariance and CUDA rSVD; PCA used the same CUDA rSVD "
        "settings without labels. The maximum 200-dimensional fit was reused for "
        "the 50- and 100-dimensional prefixes. Exact FAISS search defined the "
        "neighbour reference and classification endpoint. IVF recall@10 compares "
        "approximate and exact neighbour identifiers within the same representation; "
        "it is not class recall. Representation fits were single runs, exact "
        "index/query timings were repeated three times, and IVF was run once. The "
        "experiment is therefore explicitly exploratory."
    )
    rows = []
    for _, row in summary.sort_values(
        ["feature_space", "n_features", "faiss_method"]
    ).iterrows():
        rows.append((
            representation(row),
            integer(row["n_features"]),
            str(row["faiss_method"]).upper(),
            integer(row["n_repeats"]),
            number(row["top1_accuracy"], 4),
            number(row["top5_accuracy"], 4),
            number(row["balanced_accuracy"], 4),
            number(row["neighbour_recall_at_10"], 4),
            number(row["transformation_time_sec"], 1),
            f"{number(row['query_time_median_sec'], 1)} "
            f"({number(row['query_time_iqr_sec'], 1)})",
            f"{integer(row['peak_host_rss_mb'])}/"
            f"{integer(row['peak_gpu_mem_mb'])}",
        ))
    caption = document.add_paragraph(
        "Table S18. Complete matched ImageNet retrieval results. Query time is "
        "median (IQR); H/G are peak host/GPU memory in MB.",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True
    c9.add_table(
        document,
        ["Representation", "Dim.", "Index", "Runs", "Top-1", "Top-5",
         "Bal. acc.", "R@10", "Transform\n(s)", "Query\n(s)", "H/G\n(MB)"],
        rows,
        [1.0, 0.38, 0.42, 0.38, 0.45, 0.45, 0.52, 0.45, 0.62, 0.74, 0.75],
        font_size=5.8,
    )
    document.add_picture(str(PLOT), width=Inches(6.5))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        "Figure S20. Full exploratory ImageNet component path and computational "
        "endpoints. Exact search used three repeated index/query measurements from "
        "one fitted representation; representation fitting and IVF were single runs.",
        style="Caption",
    )
    document.add_paragraph(
        "Memory audit. Before the final operator-based run, two public-wrapper "
        "preparations were terminated by the operating system at approximately "
        "22 GiB peak RSS because they materialized large centring/response "
        "workspaces. The reported run instead formed the PLS cross-covariance from "
        "float32 class sums and evaluated centred projections as XR - μR; PCA "
        "centred only the benchmark-local in-memory copy and evaluated held-out "
        "scores as XV - μV. Source embeddings were never overwritten. These "
        "algebraic forms avoid dense one-hot responses and full centred test copies."
    )
    document.add_paragraph(
        "Machine-readable files include imagenet_faiss_matched_raw.csv, "
        "imagenet_faiss_matched_summary.csv, and "
        "imagenet_faiss_matched_main_table.csv. The archived exact-neighbour "
        "matrices permit independent recomputation of IVF recall."
    )
    document.core_properties.title = (
        "fastPLS CMPB supplement - completed ImageNet retrieval"
    )
    document.save(SUPP_OUT)


def write_response(table_data):
    raw, pls, speedup = image_net_sentences(table_data)
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    document.styles["Normal"].font.name = "Calibri"
    document.styles["Normal"].font.size = Pt(11)
    document.styles["Normal"].paragraph_format.space_after = Pt(7)
    document.styles["Heading 1"].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    document.add_heading("Response to reviewer: ImageNet analysis", level=1)
    reviewer = document.add_paragraph()
    reviewer.add_run("Reviewer comment. ").bold = True
    reviewer.add_run(
        "The ImageNet analysis is unfinished. Complete the matched raw-DINOv2 "
        "versus PCA-score versus PLS-score retrieval experiment, including "
        "transformation and query time, memory, top-1/top-5 accuracy and neighbour "
        "recall, or remove the performance claim. Identify single-run measurements "
        "as exploratory."
    )
    response = document.add_paragraph()
    response.add_run("Response. ").bold = True
    response.add_run(
        "We agree and completed the matched experiment on the fixed 1,000,000 "
        "training and 281,167 held-out embeddings. Raw DINOv2, PCA-rSVD, and "
        "label-aware PLS-SVD/rSVD used the same float32 matrices, cosine metric, "
        "k=10 voting, query blocks, and CUDA FAISS implementation. We report 50-, "
        "100-, and 200-dimensional component paths for PCA and PLS; top-1, top-5, "
        "balanced accuracy; fitting, train projection, held-out projection, query, "
        "inference, and end-to-end time; peak host/GPU memory; and IVF recall@10 "
        "against exact neighbours in the same representation."
    )
    document.add_paragraph(
        f"Raw exact retrieval achieved top-1/top-5 accuracy of "
        f"{raw['top1_accuracy']:.4f}/{raw['top5_accuracy']:.4f}. The best tested "
        f"PLS path used {int(pls['n_features'])} dimensions and achieved "
        f"{pls['top1_accuracy']:.4f}/{pls['top5_accuracy']:.4f}, with a "
        f"{speedup:.1f}-fold reduction in held-out projection-plus-query time. "
        "These results are now presented in main-text Table 2 and Figure 3, with "
        "the complete exact/IVF table in Supplementary Table S18."
    )
    document.add_paragraph(
        "We also corrected the uncertainty language. Exact index/query timing was "
        "repeated three times from one fixed representation fit, whereas each "
        "representation fit and IVF evaluation was performed once. Accordingly, "
        "the abstract, Methods, Results, table/figure captions, Discussion, and "
        "Supplementary Material explicitly identify the full ImageNet analysis as "
        "exploratory. No biomedical predictive-utility claim is made."
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    table_data = pd.read_csv(RESULTS / "imagenet_faiss_matched_main_table.csv")
    summary = pd.read_csv(RESULTS / "imagenet_faiss_matched_summary.csv")
    revise_main(table_data)
    revise_supplement(summary)
    write_response(table_data)
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
