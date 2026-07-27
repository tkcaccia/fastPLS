from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
OUT_DIR = ROOT / "artifacts" / "CMPB_independent_review_20260725"
OUT_FILE = OUT_DIR / "fastPLS_CMPB_independent_reviewer_comments_20260725.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def add_numbered_comment(doc, number, title, body):
    paragraph = doc.add_paragraph(style="List Number")
    title_run = paragraph.add_run(title + ". ")
    title_run.bold = True
    paragraph.add_run(body)


def add_minor_comment(doc, body):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(body)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)
    title.paragraph_format.space_after = Pt(6)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Calibri"
    subtitle.font.size = Pt(12)
    subtitle.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    subtitle.paragraph_format.space_after = Pt(14)

    for name, size, before, after in (
        ("Heading 1", 16, 16, 8),
        ("Heading 2", 13, 12, 6),
        ("Heading 3", 12, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Number", "List Bullet"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.25)
        style.paragraph_format.first_line_indent = Inches(-0.18)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.10


def build_report():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.text = "Independent reviewer report | Computer Methods and Programs in Biomedicine"
    header.style = doc.styles["Normal"]
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    add_page_number(section.footer.paragraphs[0])

    doc.add_paragraph("Independent reviewer report", style="Title")
    doc.add_paragraph(
        "Manuscript: “fastPLS: scalable partial least squares with compiled CPU and "
        "accelerator backends for high-dimensional biomedical data”",
        style="Subtitle",
    )

    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.65)
    table.columns[1].width = Inches(4.85)
    table.style = "Table Grid"
    metadata = (
        ("Journal", "Computer Methods and Programs in Biomedicine"),
        ("Recommendation", "Major revision"),
        ("Review basis", "Current main manuscript and supplementary material supplied for review"),
    )
    for row, (label, value) in zip(table.rows, metadata):
        row.cells[0].width = Inches(1.65)
        row.cells[1].width = Inches(4.85)
        set_cell_shading(row.cells[0], "E8EEF5")
        for cell in row.cells:
            set_cell_margins(cell)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)

    doc.add_heading("Overall assessment", level=1)
    doc.add_paragraph(
        "This manuscript presents an ambitious software framework for scalable partial least "
        "squares modelling, combining PLS-SVD, SIMPLS, OPLS, and kernel PLS with compiled CPU, "
        "CUDA, and Metal execution. The accelerated sequential SIMPLS implementation, compact "
        "prediction, matrix-free cross-covariance products, and explicit treatment of backend "
        "residency are potentially valuable contributions. The NMR application is scientifically "
        "relevant, and the authors are appropriately cautious that ImageNet is a computational "
        "stress test rather than biomedical validation."
    )
    doc.add_paragraph(
        "However, the current manuscript is not yet ready for acceptance. Several central claims "
        "are supported only partially, some speed comparisons combine changes in estimator, SVD "
        "solver, and hardware, and the ImageNet analysis is described more fully in the abstract "
        "and framing than in the quantitative results. The paper also needs a complete benchmark "
        "summary across all datasets, stronger uncertainty reporting, and a clearer distinction "
        "between production-ready and experimental precision/backend paths. These issues are "
        "addressable, and I would encourage resubmission after major revision."
    )

    doc.add_heading("Major comments", level=1)

    major_comments = [
        (
            "The principal methodological claim requires broader estimator-preservation evidence",
            "The manuscript states that accelerated SIMPLS changes execution rather than the "
            "statistical estimator. The synthetic and MetRef comparisons against "
            "pls::simpls.fit are encouraging, but they are too limited to support a general "
            "claim. Please provide a prespecified numerical validation spanning regression and "
            "classification, low- and high-rank responses, p<n and p>n, ill-conditioned data, "
            "and at least two real datasets. Report coefficient agreement where coefficients are "
            "identifiable, prediction agreement, principal angles between score/loading "
            "subspaces, selected-component agreement on fixed folds, convergence/failure counts, "
            "and tolerances. Deterministic IRLBA equivalence must remain separate from rSVD "
            "agreement, because rSVD is approximate. The supplement should also give executable "
            "pseudocode that maps each optimization to the original de Jong update.",
        ),
        (
            "The NMR speed comparison is confounded by the SVD solver",
            "The deposited reference uses PLS-SVD with IRLBA, whereas the fastest fastPLS rows "
            "use rSVD. The reported 26-fold and 387-fold speedups therefore combine software "
            "implementation, iterative solver, estimator/backend, and hardware effects. These "
            "numbers should not be interpreted as implementation-only speedups. Add matched "
            "comparisons such as reference IRLBA versus fastPLS IRLBA on CPU and fastPLS rSVD "
            "CPU versus CUDA with identical rSVD controls. If a fully matched comparison is not "
            "possible, label the current table as a complete-workflow comparison and decompose "
            "the sources of acceleration explicitly.",
        ),
        (
            "The NMR component-selection evidence is incomplete",
            "The selected value of 100 components is the upper boundary of the reported inner "
            "validation grid, and validation RMSD is still decreasing at that boundary. This "
            "does not establish that 100 is optimal. Extend the training-only component grid "
            "until a plateau or deterioration is observed, or state that 100 was the best value "
            "within a prespecified computational range. The same preprocessing, split, response "
            "target, component count, and water-region handling must be confirmed for the "
            "deposited reference and every fastPLS method. Retain the per-spectrum RMSD "
            "distribution, but also report response-wise errors and uncertainty over repeated "
            "splits or explain why only a fixed split is scientifically appropriate.",
        ),
        (
            "The complete multi-dataset benchmark is not visible in the main Results",
            "The Methods describe 12 tasks, but the main Results largely present SIMPLS "
            "agreement and NMR, while the quantitative all-dataset benchmark is not shown in a "
            "single interpretable figure or table. Add a main-text summary at the best "
            "training-selected component count for each dataset and method, displaying "
            "predictive metric, total time, host memory, GPU memory where applicable, precision, "
            "and execution status. Component-path plots for all datasets can remain "
            "supplementary. The paper should distinguish estimator-matched comparisons from "
            "workflow comparisons that use different classifiers or model families.",
        ),
        (
            "The ImageNet analysis is unfinished in the current manuscript",
            "The abstract states that fastPLS processed the million-sample ImageNet/DINOv2 task, "
            "yet the ImageNet Results section contains no accuracy, top-5 accuracy, runtime, "
            "memory, component-path, or comparator table. The final paragraph of the manuscript "
            "describes a planned FAISS analysis, which is not acceptable as a result. Either "
            "complete the matched raw-DINOv2 versus PCA-score versus PLS-score retrieval "
            "experiment, including PLS/PCA transformation and query time, memory, top-1/top-5 "
            "accuracy and neighbour recall, or remove the corresponding performance claim from "
            "the abstract. All ImageNet measurements appear to be single runs and must be "
            "identified as exploratory unless repeated.",
        ),
        (
            "Claims about float32 need to be narrowed and organized by supported path",
            "The supplement shows that float32 can reduce memory in selected cases, but on NMR "
            "it is hundreds of times slower for some paths and materially degrades SIMPLS, "
            "kernel-PLS, and OPLS RMSD. This is important negative evidence, but it conflicts "
            "with any broad implication that float32 is an end-to-end low-memory acceleration. "
            "Provide a model-by-backend precision table that distinguishes genuinely float32 "
            "operations from conversions or host fallbacks. Validate float32 versus float64 for "
            "all four PLS families, regression and classification, and every claimed backend. "
            "Until then, describe float32 as experimental or selectively supported rather than "
            "a general package advantage.",
        ),
        (
            "Backend residency and backend substitution require stricter presentation",
            "The stage-level residency table is useful, but the public interface can substitute "
            "label-aware PLS-SVD for requested CUDA SIMPLS in large-class settings, and the CUDA "
            "SIMPLS-LDA scores are reported to arise from PLS-SVD. A warning and metadata field "
            "are necessary but do not make these results SIMPLS. Every benchmark row, figure, "
            "and conclusion should use the executed estimator. Expand the residency table to "
            "identify fitting, reduced decomposition, prediction, LDA, and cross-validation "
            "separately for float64 and float32. Avoid the phrase GPU-native for any path that "
            "retains host QR, decomposition, filtering, kernel construction, or scoring.",
        ),
        (
            "External-package comparisons require a fully auditable fairness protocol",
            "The manuscript lists several R packages but does not yet present their exact "
            "versions, calls, preprocessing, response encodings, scaling conventions, component "
            "caps, prediction heads, thread counts, and timed regions in the main results. The "
            "primary comparison should be double-versus-double and estimator-matched. Sparse, "
            "NIPALS, OPLS, and discriminant workflows should not be treated as interchangeable "
            "with dense SIMPLS. Unsupported task types and memory/time failures should remain in "
            "the table, while adapter errors must be corrected and excluded from package-level "
            "conclusions. Please release the machine-readable manifest used to generate the "
            "tables.",
        ),
        (
            "Uncertainty and model-selection procedures need strengthening",
            "Three timing repetitions are a minimum, and several large tasks use only one run. "
            "Report medians with IQRs or ranges for every repeated benchmark and label single-run "
            "results as point estimates. Predictive uncertainty is more important than timing "
            "repeatability: where feasible, use repeated fixed splits or nested cross-validation "
            "and report variability. Clarify that test-set component curves are descriptive and "
            "are not used to select the reported best model. For all “best component” results, "
            "the selection rule and training-only folds should be explicit.",
        ),
        (
            "Cross-validation acceleration is asserted more strongly than it is demonstrated",
            "The manuscript explains why compiled fold handling should reduce overhead, but no "
            "direct quantitative comparison with an equivalent naive R-level K-fold workflow is "
            "shown. Add a matched experiment with identical folds, model, component grid, and "
            "predictions. Report total time, fold-level agreement, selected component, peak "
            "memory, and failures for representative small, large-low-p, and large-q datasets. "
            "The handling of grouped samples through constrain should be tested for leakage.",
        ),
        (
            "Reproducibility and software availability need completion before publication",
            "The Data and code availability section says that release identifiers, checksums, "
            "and manifests will be supplied later. These items must be present at submission, "
            "not promised. Provide an immutable package release, exact kodama-cpp revision, data "
            "checksums or acquisition scripts, split indices, compiler flags, BLAS/LAPACK and "
            "thread settings, CUDA/Metal versions, and commands that regenerate each table and "
            "figure. Continuous-integration evidence for Linux, macOS, Windows without CUDA, and "
            "a CUDA build would materially improve confidence in the package.",
        ),
        (
            "The biomedical relevance and scope should remain carefully bounded",
            "The manuscript appropriately states that ImageNet is not biomedical validation. "
            "This caveat should remain in the abstract and conclusion, and the biomedical value "
            "should rest primarily on the omics and NMR tasks. If the authors wish to motivate "
            "foundation-model embeddings, the claim should be computational portability to a "
            "similar matrix regime, not expected diagnostic performance. Likewise, kernel-PLS "
            "claims should be limited because nonlinear-kernel validation is described as "
            "incomplete.",
        ),
    ]
    for idx, (title, body) in enumerate(major_comments, start=1):
        add_numbered_comment(doc, idx, title, body)

    doc.add_heading("Minor comments", level=1)
    minor_comments = [
        "Remove the final “Package scope update” paragraph. It is editorial process text, not manuscript content, and it describes a planned analysis.",
        "Use “PLS-SVD” consistently. Define the abbreviation once and avoid alternating forms.",
        "Correct the Introduction sentence stating that predictors “may itself be multivariate”; the intended subject appears to be the response.",
        "Several mathematical expressions appear to have missing symbols in the prose surrounding OPLS, rank limits, and cross-covariance operations. Check the Word equations and the text version carefully after conversion.",
        "Define n, p, q, a, l, RSS, GPU memory, Q2, R2, RMSD, top-5 accuracy, prediction agreement, and principal angles at first use. State whether RMSD is global, per sample, or per response in every table.",
        "Explain why Q2 and R2 are very similar in some NMR tables and verify that training fitted values and held-out predictions are not being reused accidentally.",
        "The NMR figure caption says that the representative spectrum is selected by median RMSD, whereas the Results refer readers to supplementary representative spectra. Harmonize the description and identify which prediction method is plotted in panels A and B.",
        "The phrase “routine NMR preprocessing removed the water interval” should specify whether columns were removed or set to zero and whether this applies to X only, Y, or both.",
        "The supplement heading order is inconsistent: S10, S11, and S12 are followed by another S6. Renumber all sections, tables, and figure citations sequentially.",
        "Table S3 lists “SingleCell”, whereas the dataset text discusses Retina and Tabula Muris as distinct tasks. Resolve the naming and ensure the task count in the abstract, main text, table, and manifest agrees.",
        "The FlashSVD sentence is peripheral because FlashSVD is neither exposed nor benchmarked as a decomposition backend. Either provide a precise, justified implementation relationship or remove it to avoid confusing inspiration with implementation.",
        "Report CPU parallelism explicitly. “Compiled CPU” does not indicate whether BLAS, OpenMP, or other multithreading was enabled.",
        "Do not compare CUDA and Metal runtime as if they were hardware benchmarks unless the same machine and equivalent implementation are available. Prediction agreement and portability are appropriate cross-platform endpoints.",
        "Provide exact citations and version identifiers for all external packages, CUDA libraries, Apple Metal Performance Shaders, the float package, and the deposited NMR reference function.",
        "The manuscript should identify whether host RSS is an absolute process peak or an incremental allocation. The supplement does this correctly; table captions should do so consistently.",
        "Avoid phrases such as “complete statistical workflows” unless all documented combinations are tested. A capability matrix with tested, hybrid, experimental, and unsupported states would be clearer.",
        "Use consistent capitalization: CPU, CUDA, GPU, Metal, float32, and float64.",
        "Include readable panel labels and sufficiently large type in all benchmark figures. Large multi-dataset plots should be supplied as vector PDF or high-resolution supplementary figures.",
        "State the software license and source location once in the software-availability section, then avoid repeating licensing details throughout the Results.",
        "A language and typesetting edit is required for grammar, missing variables, punctuation around equations, and consistent hyphenation of training/test, cross-covariance, and held-out.",
    ]
    for body in minor_comments:
        add_minor_comment(doc, body)

    doc.add_heading("Recommendation", level=1)
    doc.add_paragraph(
        "Major revision. The package addresses a real computational need and the accelerated "
        "SIMPLS and NMR work are potentially publishable. Acceptance should depend on completing "
        "the all-dataset and ImageNet results, separating solver effects from implementation "
        "speedups, expanding estimator and precision validation, and releasing a fully "
        "reproducible benchmark manifest."
    )

    doc.core_properties.title = "Independent CMPB reviewer comments for fastPLS"
    doc.core_properties.subject = "Independent scientific review"
    doc.core_properties.author = "Independent reviewer"
    doc.core_properties.keywords = "fastPLS, PLS, SIMPLS, software validation, peer review"
    doc.save(OUT_FILE)
    print(OUT_FILE)


if __name__ == "__main__":
    build_report()
