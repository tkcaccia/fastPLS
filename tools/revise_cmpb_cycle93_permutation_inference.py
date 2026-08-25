from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle92"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle93"
OUTPUT.mkdir(parents=True, exist_ok=True)


def replace_paragraph(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def find_paragraph(document, phrase):
    matches = [p for p in document.paragraphs if phrase in p.text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph containing {phrase!r}; found {len(matches)}")
    return matches[0]


main_source = SOURCE / "fastPLS_CMPB_main_cycle92_0.99.25_20260825.docx"
main_output = OUTPUT / "fastPLS_CMPB_main_cycle93_0.99.25_20260825.docx"
main = Document(main_source)
paragraph = find_paragraph(main, "nested permutation testing uses the same selected endpoint")
replace_paragraph(
    paragraph,
    "Regression returns continuous predictions. Classification uses argmax PLS-DA or LDA fitted to PLS scores. "
    "LDA uses pooled within-class covariance, Cholesky solves, class priors, and deterministic trace-scaled "
    "regularization only when factorization fails [17,18]. For imbalanced classification, cross-validation can "
    "select by balanced accuracy, defined as the unweighted mean of class-specific recalls; nested permutation "
    "testing uses the same selected endpoint rather than substituting dummy-response Q2. Permutation inference "
    "uses the finite Monte Carlo correction (b+1)/(B+1), where b is the number of successful null statistics at "
    "least as extreme as observed and B is the number of successful null fits. Independent observations are "
    "permuted by row; repeated observations identified by the grouping constraint are exchanged as complete "
    "blocks within equal-size exchangeability strata. Group sizes, within-group responses, and class frequencies "
    "are thereby preserved, while fold assignments and randomized-solver seeds remain fixed across observed and "
    "null analyses. Failed null fits are excluded from B and reported explicitly."
)
main.save(main_output)


supp_source = SOURCE / "fastPLS_CMPB_supplement_cycle92_0.99.25_20260825.docx"
supp_output = OUTPUT / "fastPLS_CMPB_supplement_cycle93_0.99.25_20260825.docx"
supp = Document(supp_source)
paragraph = find_paragraph(supp, "When a nested permutation test is requested")
replace_paragraph(
    paragraph,
    "pls.single.cv() evaluates eligible combinations of component count and prediction settings within one K-fold "
    "cross-validation layer. pls.double.cv() adds an outer layer for unbiased performance estimation. Selection "
    "can use accuracy, balanced accuracy, R2, Q2, or RMSD. Balanced accuracy is the unweighted mean of class-specific "
    "recalls and is intended for unequal class frequencies. Nested permutation inference uses the same selection "
    "endpoint and its appropriate tail. If every row is an independent constraint group, predictor rows are "
    "permuted individually. If a constraint identifies repeated observations from one patient or subject, complete "
    "predictor blocks are exchanged only among constraint groups with equal row counts. Consequently, group sizes, "
    "within-group response patterns, and the complete response vector, including class frequencies, remain unchanged. "
    "The same outer and inner fold assignments and the same randomized-SVD seeds are reused for the observed and all "
    "null fits. For a statistic where larger is better, b counts successful null statistics greater than or equal to "
    "the observed value; for a loss, b counts values less than or equal to observed. With B successful permutations, "
    "the reported p-value is (b+1)/(B+1), so a finite Monte Carlo test cannot return zero. Failed null fits are stored "
    "with their messages, omitted from B, and summarized as requested, completed, and failed counts. A grouped test "
    "therefore requires at least two equal-size constraint groups that can be exchanged; otherwise the software stops "
    "with an explicit error rather than performing an invalid row-level permutation."
)
supp.save(supp_output)

print(main_output)
print(supp_output)
