#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
OUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260727_cycle80_review"
OUT = OUT_DIR / "fastPLS_CMPB_fresh_reviewer_report_cycle80_20260727.docx"

BLUE = RGBColor(31, 78, 121)
MUTED = RGBColor(89, 89, 89)
RED = RGBColor(156, 0, 6)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, text, separate, end])


def configure(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for name, size in (("Title", 21), ("Heading 1", 15), ("Heading 2", 12.5)):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = "Independent reviewer report | fastPLS"
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def labelled(document, label, text, color=None):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(label)
    run.bold = True
    if color:
        run.font.color.rgb = color
    paragraph.add_run(text)


def comment(document, number, title, assessment, revision):
    document.add_heading(f"{number}. {title}", level=2)
    document.add_paragraph(assessment)
    labelled(document, "Required revision: ", revision, RED)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure(document)

    document.add_paragraph("Reviewer Report", style="Title")
    document.add_paragraph(
        "fastPLS: scalable partial least squares with compiled CPU and accelerator backends for high-dimensional biomedical data"
    )
    labelled(document, "Journal: ", "Computer Methods and Programs in Biomedicine")
    labelled(document, "Recommendation: ", "Minor revision", RED)
    labelled(
        document,
        "Review basis: ",
        "Cycle80 main manuscript and Supplementary Material, read as a new submission.",
    )

    document.add_heading("Overall assessment", level=1)
    document.add_paragraph(
        "This is a substantially improved and technically credible software manuscript. "
        "The primary contribution is now focused on accelerated SIMPLS, with PLS-SVD, "
        "rSVD, float32, accelerator execution, and compiled cross-validation presented as "
        "supporting capabilities. The manuscript distinguishes deterministic estimator "
        "validation from approximate rSVD benchmarking, labels hybrid accelerator routes, "
        "reports failures rather than silently excluding them, and frames ImageNet as an "
        "exploratory computational and representation stress test rather than biomedical "
        "validation. The NMR application provides a relevant high-response-dimensional "
        "biomedical use case. The main text is concise and the detailed numerical evidence "
        "is appropriately delegated to the Supplement."
    )
    document.add_paragraph(
        "The remaining concerns are internal consistency and reproducibility details. They "
        "can be resolved without additional computation, but should be corrected before "
        "submission so that the definitive evidence is unambiguous."
    )

    document.add_heading("Major comments", level=1)
    comment(
        document,
        1,
        "The NMR representative-spectrum description is internally inconsistent",
        "The main Figure 4 displays a representative held-out observed and predicted "
        "spectrum selected near the median SIMPLS per-spectrum RMSD. Supplementary Section "
        "S17 still states that the main figure does not show such a spectrum and refers to "
        "an archived test sample 204. This conflicts with the current qualified analysis, "
        "which identifies sample AMI-00BP-8 (index 155).",
        "Synchronize Section S17, the Figure 4 legend, and the provenance record. State the "
        "selection rule once and identify the same sample everywhere.",
    )
    comment(
        document,
        2,
        "The NMR component-selection rule needs an executable description",
        "The text correctly calls the selected values the smallest values retained by a "
        "one-standard-error rule, rather than global optima. However, the complete candidate "
        "grids and eligibility decision are not stated together. This makes the five- and "
        "50-component choices difficult to reproduce from the manuscript alone.",
        "Report the PLS-SVD and SIMPLS grids, five paired training-only splits, the exact "
        "one-standard-error threshold, and the eligible component values. Preserve the "
        "wording 'selected within the evaluated grid.'",
    )
    comment(
        document,
        3,
        "The pseudocode conflicts with the declared notation",
        "The notation section reserves A for retained components, C for requested component "
        "counts, and lower-case k for retrieval/top-k quantities. Algorithm S1 nevertheless "
        "uses k for component prefixes.",
        "Replace this occurrence with an unambiguous component symbol, for example "
        "'for each requested a in C, update prefixes 1,...,a', and recheck all algorithm "
        "symbols against the notation table.",
    )

    document.add_heading("Minor comments", level=1)
    minor = [
        "Table S3 labels ImageNet only as a retrieval stress test, although the manuscript also reports exploratory SIMPLS-LDA classification. Rename the task accordingly.",
        "Clarify that the isolated CIFAR-100 SIMPLS-LDA evidence run used a 7,200-s limit, whereas the current package-comparison workflow defaults to 10,000 s.",
        "Name the exact reviewed source archive (fastPLS_0.99.6.tar.gz), its SHA-256 checksum, and the base commit in the availability/provenance statement.",
        "Retain the explicit qualification that ImageNet results are single-run exploratory measurements and that the 1,000-component point is a boundary stress point, not an optimum.",
        "After these edits, repeat the existing notation, citation, figure/table, and rendered-page audits.",
    ]
    for number, text in enumerate(minor, 1):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{number}. ")
        run.bold = True
        paragraph.add_run(text)

    document.add_heading("Recommendation", level=1)
    document.add_paragraph(
        "Minor revision. The methodological evidence is now sufficient for editorial "
        "assessment, subject to correction of the remaining inconsistencies above. I do not "
        "request additional experiments in this review round."
    )

    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
