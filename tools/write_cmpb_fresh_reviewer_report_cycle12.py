from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle12"
DOCX = OUT / "CMPB_fresh_reviewer_comments_cycle12_20260725.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(89, 89, 89)
RED = RGBColor(155, 28, 28)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Title", 22, BLUE, 0, 8),
        ("Subtitle", 11, GRAY, 0, 14),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Reviewer Lead" not in doc.styles:
        lead = doc.styles.add_style("Reviewer Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = doc.styles["Reviewer Lead"]
    lead.font.name = "Calibri"
    lead.font.size = Pt(11)
    lead.font.bold = True
    lead.font.color.rgb = RED
    lead.paragraph_format.space_before = Pt(4)
    lead.paragraph_format.space_after = Pt(8)
    lead.paragraph_format.line_spacing = 1.10

    header = section.header.paragraphs[0]
    header.text = "CONFIDENTIAL PEER-REVIEW REPORT"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.color.rgb = GRAY

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    add_field(footer, "PAGE")


def add_labeled_paragraph(doc, label, text):
    p = doc.add_paragraph()
    r = p.add_run(label)
    r.bold = True
    p.add_run(text)
    return p


def add_comment(doc, number, title, paragraphs):
    doc.add_heading(f"Major comment {number}. {title}", level=2)
    for text in paragraphs:
        doc.add_paragraph(text)


def add_minor(doc, number, text):
    p = doc.add_paragraph()
    r = p.add_run(f"{number}. ")
    r.bold = True
    p.add_run(text)


doc = Document()
configure_document(doc)

title = doc.add_paragraph(style="Title")
title.add_run("Reviewer report")
subtitle = doc.add_paragraph(style="Subtitle")
subtitle.add_run(
    "Manuscript: “fastPLS: scalable partial least squares with compiled CPU and "
    "accelerator backends for high-dimensional biomedical data”"
)

summary = doc.add_table(rows=4, cols=2)
summary.autofit = False
summary.columns[0].width = Inches(1.8)
summary.columns[1].width = Inches(4.7)
items = [
    ("Journal", "Computer Methods and Programs in Biomedicine"),
    ("Article type", "Computational methods and software"),
    ("Recommendation", "Major revision"),
    ("Review basis", "Main manuscript and Supplementary Material supplied for review"),
]
for i, (label, value) in enumerate(items):
    summary.cell(i, 0).width = Inches(1.8)
    summary.cell(i, 1).width = Inches(4.7)
    summary.cell(i, 0).text = label
    summary.cell(i, 1).text = value
    shade_cell(summary.cell(i, 0), "F2F4F7")
    for cell in summary.rows[i].cells:
        set_cell_margins(cell)
        cell.vertical_alignment = 1
    summary.cell(i, 0).paragraphs[0].runs[0].bold = True

doc.add_heading("Overall assessment", level=1)
doc.add_paragraph(
    "This manuscript addresses an important computational problem: extending partial least "
    "squares modelling and validation to matrices that are difficult to handle with conventional "
    "R implementations. The combination of compiled execution, memory-aware prediction, "
    "label-aware products, CUDA support, and a carefully optimized SIMPLS path is potentially "
    "valuable to the readership of Computer Methods and Programs in Biomedicine. The NMR case "
    "study is scientifically relevant, and the matched ImageNet experiment is now useful as a "
    "clearly delimited computational stress test and supervised-compression experiment."
)
doc.add_paragraph(
    "The manuscript is nevertheless not ready for publication. Its strongest evidence applies "
    "to deterministic IRLBA-based SIMPLS, whereas most headline performance results use rSVD, "
    "which the authors' own validation shows can depart substantially from the de Jong estimator. "
    "In addition, OPLS, nonlinear kernel PLS, Metal execution, cross-validation acceleration, and "
    "the external-software comparison are described more broadly than they are quantitatively "
    "validated. These are addressable issues, but they require changes to the experimental record "
    "and to the scope of the claims rather than editorial revision alone."
)
lead = doc.add_paragraph(style="Reviewer Lead")
lead.add_run(
    "Recommendation: major revision. The computational contribution is promising, but the "
    "estimator, approximation, backend, and workflow claims must be separated and supported "
    "consistently."
)

doc.add_heading("Major comments", level=1)

add_comment(
    doc,
    1,
    "Separate estimator-preserving SIMPLS from approximate SIMPLS-rSVD",
    [
        "The central methodological claim is currently split between two materially different "
        "objects. Deterministic IRLBA SIMPLS agrees closely with pls::simpls.fit and supports the "
        "statement that execution was reorganized without changing the de Jong estimator. By "
        "contrast, Supplementary Section S13 reports for rSVD a maximum relative prediction error "
        "of 1.06, a minimum prediction correlation of 0.445, a maximum score-subspace angle of "
        "88.4 degrees, and a minimum classification-label agreement of 0.133. These are not small "
        "numerical perturbations.",
        "Most headline timing results use rSVD. The paper should therefore define two separate "
        "claims: an estimator-preserving compiled SIMPLS implementation with deterministic IRLBA, "
        "and an approximate SIMPLS-rSVD variant with an empirical accuracy-speed trade-off. Tables, "
        "figures, abstract, and conclusions should identify the solver with every SIMPLS result. "
        "The authors should also state whether rSVD is the public default and, if so, explain how "
        "users are warned that it may change the fitted latent subspace and predictions."
    ],
)

add_comment(
    doc,
    2,
    "The claimed scope exceeds the validation of OPLS and kernel PLS",
    [
        "The package presents PLS-SVD, SIMPLS, OPLS, and kernel PLS as co-equal supported model "
        "families. However, the manuscript explicitly states that nonlinear kernel claims are "
        "limited to functionality until matched-reference validation is completed. The selected-"
        "point kernel PLS results appear to use the linear shortcut and are therefore not evidence "
        "for radial-basis or polynomial kernel implementations. OPLS validation is also limited, "
        "and the number and selection of orthogonal components are not evident in the main results.",
        "Before publication, the authors should provide matched-reference numerical tests for OPLS "
        "and for each nonlinear kernel, including prediction agreement, parameter selection using "
        "training data only, failures, and representative regression and classification tasks. "
        "A nonlinear synthetic problem would help demonstrate that kernel PLS is doing more than "
        "dispatching to linear SIMPLS. Otherwise, the manuscript should narrow its central claims "
        "to the model families that have been validated."
    ],
)

add_comment(
    doc,
    3,
    "The main benchmark obscures backend variability and uncertainty",
    [
        "Table 1 reports only the fastest completed CPU or CUDA row for each family and dataset. "
        "This selection is useful operationally but is not sufficient for scientific comparison: "
        "it hides the paired backend result, dispersion, the number of successful repetitions, and "
        "the magnitude of the speedup. The table also combines estimator-matched comparisons with "
        "workflow comparisons while using a dense within-cell format that is difficult to audit.",
        "Please report CPU and CUDA results side by side, with median and IQR (or another stated "
        "dispersion measure), completed/attempted runs, predictive metric, host RSS, GPU memory, "
        "and status. Fit and prediction times should be separated at least in the supplement. The "
        "selection rule for the displayed backend must be prespecified and should not depend on "
        "outer-test accuracy. The ImageNet row in Table 1 uses a 1,000-component SIMPLS/argmax "
        "workflow, whereas Section 3.3 uses 200-dimensional PLS-SVD scores and external kNN; these "
        "are different experiments and need explicit cross-referencing to prevent apparent "
        "contradiction."
    ],
)

add_comment(
    doc,
    4,
    "The external-package comparison is described but not presented",
    [
        "The Methods name nine external packages and carefully distinguish estimator-matched from "
        "workflow comparisons, but the Results provide no quantitative external-package table. A "
        "reader cannot verify the claimed software advantage, assess package restrictions, or "
        "determine whether timing scopes were matched.",
        "Add a precision-matched float64 table for compatible methods showing package and version, "
        "function, model family, response encoding, classification head, component count, timed "
        "operations, runtime dispersion, predictive metric, memory where available, and exact "
        "failure or restriction. Sparse PLS and other statistically different estimators should "
        "remain clearly separated from estimator-matched comparisons. If these data are not "
        "available, the external-comparison claims should be removed."
    ],
)

add_comment(
    doc,
    5,
    "Metal and multicore CPU claims require quantitative evidence",
    [
        "The manuscript title and architecture emphasize CPU and accelerator backends, yet Metal "
        "is supported mainly by a residency description and smoke tests. The Results even state "
        "that final reporting 'will compare' native platforms, which is future work and should not "
        "appear in a completed study. Likewise, the CPU thread configuration and any multicore "
        "parallelism are not sufficiently reported.",
        "Provide matched CPU-versus-Metal prediction agreement, runtime, host/unified-memory use, "
        "and failure status on at least two representative real datasets and more than one matrix "
        "shape. Report Apple hardware and software details and identify every host stage. State "
        "the number of CPU threads and BLAS implementation for all CPU benchmarks. If Metal is "
        "only a portability prototype, the title, abstract, and conclusions should say so."
    ],
)

add_comment(
    doc,
    6,
    "Float32 support is not yet validated as a general package capability",
    [
        "The manuscript appropriately reports negative float32 findings on NMR, including large "
        "RMSD degradation for SIMPLS, linear kernel PLS, and OPLS. Those results make a broad "
        "end-to-end float32 claim difficult to sustain. The fixed-score LDA agreement and the "
        "CIFAR-100 memory example do not validate all four PLS families across CPU, CUDA, and "
        "Metal.",
        "A model-by-backend precision table should report float32-versus-float64 prediction "
        "agreement, metric difference, selected-component agreement, runtime, host/GPU memory, "
        "and numerical failures across representative regression and classification tasks. Mark "
        "each route as validated, experimental, unsupported, or hybrid. Confirm that no hidden "
        "float64 conversion occurs in a route described as end-to-end float32. The Windows "
        "limitation should be presented as a platform limitation rather than left only in the "
        "supplement."
    ],
)

add_comment(
    doc,
    7,
    "The NMR comparison needs model-specific component-selection logic",
    [
        "The extended training-only selection is an important improvement and establishes an "
        "interior minimum for SIMPLS-rSVD. However, the selected 100 components are subsequently "
        "used for the deposited PLS-SVD/IRLBA reference and fastPLS PLS-SVD as well. It is not "
        "clear that 100 is training-optimal for those model/solver combinations, and this appears "
        "in tension with the Methods statement that components were selected separately for each "
        "PLS family.",
        "Either repeat training-only component selection for each compared NMR model or define 100 "
        "as a common prespecified computational comparison point and avoid calling it optimal for "
        "all methods. The approximately 12-20% RMSD increase of SIMPLS relative to the deposited "
        "reference also deserves direct interpretation rather than emphasis on similar median "
        "response-wise RMSD. Report uncertainty for pairwise RMSD differences and explain whether "
        "the tails of the response-wise error distribution account for the discrepancy."
    ],
)

add_comment(
    doc,
    8,
    "The ImageNet experiment is useful but remains a single-fit exploratory analysis",
    [
        "The matched raw-DINOv2, PCA, and PLS comparison is well motivated and substantially "
        "strengthens the paper. However, only query timing is repeated. Accuracy variability from "
        "the randomized PCA/PLS fit is not measured. The reported fourfold gain is an inference-"
        "time comparison; the corresponding end-to-end reduction is about 2.4-fold and should be "
        "reported alongside it.",
        "Repeat the representation fits with several fixed rSVD seeds if computationally feasible, "
        "or retain an explicit single-fit limitation in every claim. Define the cosine weighting, "
        "tie handling, class-score construction, top-5 calculation when k=10, and whether index "
        "construction is included in query time. Neighbour recall should remain clearly described "
        "as IVF-versus-exact retrieval recall within each representation, not predictive recall. "
        "The natural-image experiment must remain separate from biomedical validity claims."
    ],
)

add_comment(
    doc,
    9,
    "Cross-validation acceleration is asserted without a quantitative benchmark",
    [
        "Compiled cross-validation is presented as a major practical contribution, but the Results "
        "contain only a qualitative statement. No table shows the cost of a single fit, optimized "
        "10-fold validation, and a transparent reference implementation, nor whether the selected "
        "component and predictions agree.",
        "Add a quantitative cross-validation experiment across representative small, large-low-p, "
        "and extreme-response datasets. Report total time, speedup relative to repeated public "
        "fits, selected component, predictive metric, memory, and backend residency. CPU, CUDA, "
        "and Metal should be reported separately, and hybrid paths should not be described as "
        "fully native."
    ],
)

add_comment(
    doc,
    10,
    "The source of the SIMPLS speedup needs a broader ablation",
    [
        "The pseudocode and mapping to de Jong SIMPLS are valuable, but the performance evidence "
        "does not isolate the contribution of each optimization. The reported ablation is limited "
        "to MetRef and a 5,000-sample CIFAR-100 subset, and randomized workspace reuse itself "
        "changes accuracy. It is therefore difficult to determine which gains arise from compiled "
        "execution, cached X-transpose-X, incremental coefficients/predictions, implicit products, "
        "compact prediction, or the approximate solver.",
        "Provide an ablation over at least three contrasting matrix shapes, separating estimator-"
        "preserving optimizations from rSVD approximations. Include runtime, peak memory, prediction "
        "agreement, and failure status. Expand Table S2 with solver-dependent complexity terms and "
        "state the shape-dependent activation conditions sufficiently clearly for results to be "
        "reproduced."
    ],
)

add_comment(
    doc,
    11,
    "Dataset identity and reproducibility materials are not yet submission-ready",
    [
        "The manuscript alternates among 'SingleCell', Retina, and Tabula Muris, but the relationship "
        "between these objects is not clear. Each dataset in the main benchmark should have one "
        "unambiguous name, source, task definition, licence, prepared dimensions, preprocessing, "
        "and checksum. Similar detail is needed for the DINOv2 embeddings and the deposited NMR "
        "reference function.",
        "The Data and code availability statement says that the final submission will provide "
        "release identifiers, checksums, and manifests. These materials should already exist at "
        "review. Provide an immutable software release and archive DOI, exact package and compiler "
        "versions, BLAS/thread settings, optional-backend build instructions, benchmark commands, "
        "and machine-readable manifests. A clean CPU-only installation and test on Windows, Linux, "
        "and macOS should be documented."
    ],
)

add_comment(
    doc,
    12,
    "Correct the citation record and reconsider the FlashSVD attribution",
    [
        "Reference 5 contains an incorrect DOI. The inflammatory prostate-cancer article is indexed "
        "as doi:10.1186/s40170-021-00265-6, not 10.1186/s40170-021-00273-0. All references and "
        "dataset citations should be checked systematically.",
        "The statement that blocked latent prediction is inspired by FlashSVD is not currently "
        "convincing. FlashSVD concerns fused streaming inference for low-rank neural-network "
        "factors, whereas the implementation described here is blocked low-rank matrix prediction. "
        "If a concrete algorithmic element was transferred, identify it precisely and cite the "
        "peer-reviewed version if available. Otherwise, remove the attribution and describe the "
        "implementation directly."
    ],
)

doc.add_heading("Minor comments", level=1)
minor_comments = [
    "Use “PLS-SVD” consistently in prose, tables, figure labels, object names, and captions.",
    "Define every symbol immediately after its first equation, including all dimensions and the meaning of the randomized sketch width.",
    "Replace future-tense statements such as “final reporting will compare” with completed evidence or an explicit limitation.",
    "Clarify whether peak host RSS includes the loaded input matrices, package runtime, and index storage; report the GPU sampling interval and whether the value is process-specific.",
    "Table 1 is difficult to read at its current font size and information density. Consider one row per method/backend or move the full grid to the supplement and retain a simpler speedup summary in the main text.",
    "State the OPLS orthogonal-component setting and the kernel parameters used in every benchmark row.",
    "Explain whether the factor-response PLS-SVD cap is C minus 1 in every centred classification case and how requested values above the cap are shown to users.",
    "Report confidence intervals or paired bootstrap intervals for major predictive comparisons where feasible; large test sets make very small numerical differences appear visually important.",
    "Clarify whether the ImageNet transformation time includes index construction. The table caption currently separates transformation and query, but the implementation description should match exactly.",
    "The conclusion should distinguish capabilities that are production-ready from experimental or hybrid routes.",
    "Audit reference formatting, DOI accuracy, author lists, and capitalization. The reference list currently mixes abbreviated and full bibliographic styles.",
    "Verify that all equations, superscripts, multiplication symbols, and Greek letters survive the journal's Word-to-PDF conversion.",
]
for idx, text in enumerate(minor_comments, 1):
    add_minor(doc, idx, text)

doc.add_heading("Recommendation to the editor", level=1)
doc.add_paragraph(
    "The manuscript is within the journal's methodological and software scope and has the potential "
    "to become a useful contribution. I recommend major revision. A revised submission should be "
    "considered if it clearly separates deterministic estimator preservation from randomized "
    "approximation, validates or narrows the OPLS/kernel/Metal/float32 claims, presents the missing "
    "external and cross-validation benchmarks, and releases a complete reproducibility archive."
)

OUT.mkdir(parents=True, exist_ok=True)
doc.save(DOCX)
print(DOCX)
