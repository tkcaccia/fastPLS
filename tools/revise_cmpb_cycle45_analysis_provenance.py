#!/usr/bin/env python3

import csv
import hashlib
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle44"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle45"
EVIDENCE = ROOT / "benchmark_results" / "manuscript_revision_cycle45_20260726"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle44_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle44_0.99.6_20260726.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle45_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle45_0.99.6_20260726.docx"
LEDGER = EVIDENCE / "analysis_commit_provenance.csv"

UNKNOWN = "not recorded in archived metadata"
METAL_COMMIT = "6e50bd318f20289101f6b723953830aefa8b95d6"


ANALYSES = [
    {
        "id": "A01",
        "output": "Main Fig. 2; Tables S13-S16, S31, S35",
        "archive": "benchmark_results/manuscript_revision_cycle20_20260725",
        "script": "benchmark/summarize_paired_backend_selected.R",
        "data": "Prepared dataset-specific fixed outer splits; source rows retained under manuscript_multidataset_summary_20260725/source",
    },
    {
        "id": "A02",
        "output": "Main Table 3; Tables S17, S26",
        "archive": "benchmark_results/manuscript_revision_cycle15_20260725",
        "script": "benchmark/benchmark_pls_package_comparison.R",
        "data": "Same prepared training/test split and selected A within each dataset",
    },
    {
        "id": "A03",
        "output": "SIMPLS preservation; Tables S8-S11",
        "archive": "benchmark_results/simpls_estimator_preservation_reliable_power2_final_20260725",
        "script": "benchmark/benchmark_simpls_estimator_preservation.R",
        "data": "Fixed synthetic/real tasks, seed 123, stored five-fold assignments",
    },
    {
        "id": "A04",
        "output": "Main Fig. 3; Tables S5-S7, S12, S18-S19",
        "archive": "benchmark_results/review_nmr_extended_selection_20260725",
        "script": "benchmark/review_nmr_extended_component_selection.R",
        "data": "NMR protocol manifest; training-only repeated splits; water region handled identically",
    },
    {
        "id": "A05",
        "output": "NMR matched contrasts; Table S34",
        "archive": "benchmark_results/nmr_matched_contrasts_20260726",
        "script": "benchmark/plot_nmr_matched_contrasts.R",
        "data": "Matched NMR split/preprocessing with fixed family, A, precision, solver, and hardware per contrast",
    },
    {
        "id": "A06",
        "output": "Main Table 2; Tables S17b, S20-S21, S29, S32",
        "archive": "benchmark_results/imagenet_faiss_matched_1m_20260725",
        "script": "benchmark/benchmark_imagenet_faiss_matched_retrieval.R",
        "data": "Split provenance CSV; seed 123; disjoint 1,000,000/281,167 development split",
    },
    {
        "id": "A07",
        "output": "Table S33",
        "archive": "benchmark_results/cv_compiled_vs_r_loop_20260725",
        "script": "benchmark/benchmark_cv_compiled_vs_r_loop.R",
        "data": "Identical fixed folds and outputs for compiled and R-loop routes",
    },
    {
        "id": "A08",
        "output": "Table S34; Fig. S23",
        "archive": "benchmark_results/simpls_multidataset_ablation_20260725",
        "script": "benchmark/benchmark_simpls_multidataset_ablation.R",
        "data": "Stored train/test tasks; matched seeds and three isolated runs per ablation pair",
    },
    {
        "id": "A09",
        "output": "Tables S17c, S22, S36",
        "archive": "benchmark_results/float32_backend_agreement_cycle5",
        "script": "scripts/run_reviewer_precision_validation.sh",
        "data": "Matched float64/float32 prepared tasks and component counts",
    },
    {
        "id": "A10",
        "output": "LDA backend agreement",
        "archive": "benchmark_results/lda_backend_agreement_order_fixed",
        "script": "scripts/rerun_cuda_lda_package_comparison.sh",
        "data": "Fixed labels, folds, selected components, and deterministic regularization order",
    },
    {
        "id": "A11",
        "output": "Tables S23-S25",
        "archive": "benchmark_results/manuscript_revision_cycle43_20260726",
        "script": "scripts/run_kernel_sensitivity_after_suite.sh",
        "data": "Training-only kernel tuning; held-out outer test evaluation",
    },
    {
        "id": "A12",
        "output": "Tables S37; Fig. S24",
        "archive": "benchmark_results/simpls_vs_plssvd_shapes_20260726_cuda",
        "script": "benchmark/benchmark_simpls_vs_plssvd_shapes.R",
        "data": "Synthetic seed 777; matched matrices/splits; rSVD seeds 101-103",
    },
    {
        "id": "A13",
        "output": "Tables S39-S40",
        "archive": "benchmark_results/opls_kernel_estimator_validation_verified_20260726",
        "script": "benchmark/benchmark_opls_kernel_estimator_validation.R",
        "data": "Fixed synthetic/real tasks, seed 123, identical five-fold partitions",
    },
    {
        "id": "A14",
        "output": "Tables S18, S30; Metal validation figures",
        "archive": "benchmark_results/metal_validation_20260726",
        "script": "benchmark/metal_validation/run_metal_validation.R",
        "data": "Per-run task and seed fields retained in Metal campaign session files",
        "package_commit": METAL_COMMIT,
        "commit_status": "recorded by run",
    },
]


def tree_digest(directory):
    root = ROOT / directory
    if not root.exists():
        return "archive missing"
    digest = hashlib.sha256()
    files = sorted(path for path in root.rglob("*") if path.is_file())
    for path in files:
        relative = path.relative_to(root).as_posix()
        digest.update(relative.encode("utf-8"))
        digest.update(b"\0")
        with path.open("rb") as handle:
            for block in iter(lambda: handle.read(1024 * 1024), b""):
                digest.update(block)
        digest.update(b"\0")
    return digest.hexdigest()


def script_md5(script):
    path = ROOT / script
    if not path.exists():
        return "file unavailable"
    return hashlib.md5(path.read_bytes()).hexdigest()


def rows():
    output = []
    for item in ANALYSES:
        output.append(
            {
                "analysis_id": item["id"],
                "reported_output": item["output"],
                "result_archive": item["archive"],
                "package_version": "0.99.6",
                "package_commit": item.get("package_commit", UNKNOWN),
                "package_commit_status": item.get(
                    "commit_status", "not recoverable; not inferred"
                ),
                "benchmark_script": item["script"],
                "benchmark_script_commit": UNKNOWN,
                "benchmark_script_md5_current_copy": script_md5(item["script"]),
                "kodama_cpp_commit": UNKNOWN,
                "data_split_provenance": item["data"],
                "result_archive_sha256": tree_digest(item["archive"]),
            }
        )
    return output


def write_ledger(items):
    EVIDENCE.mkdir(parents=True, exist_ok=True)
    with LEDGER.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(items[0]))
        writer.writeheader()
        writer.writerows(items)


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_paragraph(document, prefix, text):
    paragraph = find_paragraph(document, prefix)
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def set_cell_width(cell, width_inches):
    width = int(width_inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def format_table(table, widths, font_size=5.0):
    table.autofit = False
    try:
        table.style = "Table Grid"
    except KeyError:
        table.style = "Table"
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    header_pr.append(OxmlElement("w:tblHeader"))
    for row_index, row in enumerate(table.rows):
        row_pr = row._tr.get_or_add_trPr()
        row_pr.append(OxmlElement("w:cantSplit"))
        for column_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
                    if row_index == 0:
                        run.bold = True


def add_table(document, headers, values, widths):
    table = document.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = str(value)
    for values_row in values:
        row = table.add_row()
        for cell, value in zip(row.cells, values_row):
            cell.text = str(value)
    format_table(table, widths)
    return table


def revise_main():
    document = Document(MAIN_SOURCE)
    paragraph = find_paragraph(document, "The main biomedical benchmark included")
    paragraph.text += (
        " Analysis-specific source provenance is reported in Supplementary "
        "Table S41 and in a machine-readable ledger. Exact commits are reported "
        "only when they were written by the run itself; an unrecorded historical "
        "SHA is not inferred from package version or result date."
    )
    replace_paragraph(
        document,
        "The fastPLS R package, benchmark workflows",
        (
            "The fastPLS R package, benchmark workflows, analysis scripts, "
            "machine-readable result tables, synthetic generators, and aggregate "
            "benchmark outputs are available at https://github.com/tkcaccia/fastPLS. "
            "Low-level reusable C++ components are maintained at "
            "https://github.com/tkcaccia/kodama-cpp. Supplementary Table S41 and "
            "benchmark_results/manuscript_revision_cycle45_20260726/"
            "analysis_commit_provenance.csv map every reported analysis to its "
            "result archive, generating script, recorded package version, exact "
            "package commit when captured during execution, data/split provenance, "
            "and SHA-256 archive digest. Historical archives that recorded version "
            "0.99.6 but not a Git SHA are marked explicitly as not recoverable; no "
            "later commit is assigned retrospectively. The package bundles only "
            "the GPL-compatible breast and colon examples; prepared real-data "
            "benchmark matrices are not redistributed. benchmark/"
            "DATA_ACQUISITION.md and benchmark/acquire_publication_datasets.R "
            "provide authoritative source links, exact release identifiers, "
            "executable public downloads, and checksum validation for "
            "user-authorized restricted files."
        ),
    )
    document.core_properties.title = (
        "fastPLS CMPB manuscript - analysis-specific source provenance"
    )
    document.save(MAIN_OUT)


def revise_supplement(items):
    document = Document(SUPP_SOURCE)
    replace_paragraph(
        document,
        "This supplement describes fastPLS version",
        (
            "This supplement describes fastPLS version 0.99.6. Source provenance "
            "is analysis specific: Table S41 maps each quantitative analysis to "
            "its result archive and records an exact package commit only when the "
            "run metadata captured it. Package version, result date, or a later "
            "manuscript commit is not treated as evidence of the historical "
            "computational SHA."
        ),
    )
    replace_paragraph(
        document,
        "The benchmark record includes package commit",
        (
            "Current benchmark workflows record repository commit and dirty state, "
            "benchmark-script checksum, package version, reusable-core commit when "
            "available, compiler, optimization flags, BLAS/LAPACK library, thread "
            "settings, R version, operating system, accelerator libraries, seeds, "
            "rSVD controls, and data/split identifiers. Earlier archives were less "
            "uniform: most CPU/CUDA campaigns retained version 0.99.6 and numerical "
            "settings but not an exact Git SHA, whereas the Metal campaign recorded "
            f"{METAL_COMMIT}. Table S41 distinguishes these cases and supplies a "
            "SHA-256 digest for every result archive. Missing historical SHAs are "
            "reported as unavailable rather than reconstructed."
        ),
    )

    env_table = document.tables[4]
    env_table.cell(2, 1).text = (
        "R 4.6.0 and R 4.5.1 across campaigns; fastPLS 0.99.6. "
        "See Table S41 for analysis-specific source provenance."
    )
    env_table.cell(2, 2).text = (
        "R 4.6.0; fastPLS 0.99.6; package commit recorded by each Metal run. "
        "See Table S41."
    )
    format_table(env_table, (1.10, 2.62, 2.62), font_size=5.2)

    document.add_heading("S35. Analysis-specific source provenance", level=1)
    document.add_paragraph(
        (
            "Table S41 is the authoritative mapping between manuscript outputs "
            "and computational archives. A package commit identifies code loaded "
            "for computation; a benchmark-script commit identifies orchestration "
            "code; and the archive SHA-256 identifies the exact result tree. These "
            "identifiers are not interchangeable. Only A14 contains an exact "
            "package SHA written during execution. For A01-A13, the archived "
            "metadata identify fastPLS 0.99.6 but do not contain an immutable "
            "package or script SHA; these fields therefore remain unavailable. "
            "The current script MD5 values are retained in the machine-readable "
            "ledger for file identification but are not represented as historical "
            "run provenance. Future runs use benchmark/write_run_provenance.R to "
            "capture these fields before computation."
        )
    )
    caption = document.add_paragraph(
        (
            "Table S41. Per-analysis provenance. SHA is the package commit recorded "
            "by the run; NR means not recorded and not inferred. Digest is the "
            "first 12 characters of the SHA-256 over sorted relative paths and "
            "file contents in the result archive. The complete values and "
            "data/split descriptions are in analysis_commit_provenance.csv."
        ),
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True
    compact = []
    for item in items:
        sha = (
            item["package_commit"][:12]
            if item["package_commit"] != UNKNOWN
            else "NR"
        )
        compact.append(
            (
                item["analysis_id"],
                item["reported_output"],
                item["result_archive"].replace("benchmark_results/", ""),
                Path(item["benchmark_script"]).name,
                item["package_version"],
                sha,
                item["result_archive_sha256"][:12],
            )
        )
    add_table(
        document,
        ("ID", "Reported output", "Result archive", "Generating script", "Ver.", "SHA", "Digest"),
        compact,
        (0.35, 1.25, 1.72, 1.45, 0.42, 0.72, 0.78),
    )
    document.add_paragraph(
        (
            "Machine-readable ledger: benchmark_results/"
            "manuscript_revision_cycle45_20260726/"
            "analysis_commit_provenance.csv. The result-archive digest permits "
            "verification of the archived outputs even where historical source "
            "commit capture was incomplete."
        )
    )
    document.core_properties.title = (
        "fastPLS CMPB supplement - analysis-specific source provenance"
    )
    document.save(SUPP_OUT)


def main():
    items = rows()
    write_ledger(items)
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement(items)
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(LEDGER)


if __name__ == "__main__":
    main()
