#!/usr/bin/env python3
"""Audit reader-facing notation in CMPB DOCX tables and figure PDFs."""

import re
import subprocess
import sys
from pathlib import Path

from docx import Document


NONCANONICAL = (
    ("plain Q2 metric", re.compile(r"(?<![A-Za-z0-9])Q2(?![A-Za-z0-9])")),
    ("PLSSVD", re.compile(r"\bPLSSVD\b", re.IGNORECASE)),
    ("PLS SVD", re.compile(r"\bPLS[ _/]SVD\b", re.IGNORECASE)),
    ("uppercase RSVD", re.compile(r"\bRSVD\b")),
    ("capitalized float precision", re.compile(r"\bFloat(?:32|64)\b")),
    ("noncanonical CPU", re.compile(r"\b(?:Cpu|cpu)\b")),
    ("noncanonical CUDA", re.compile(r"\b(?:Cuda|cuda)\b")),
    ("noncanonical Metal", re.compile(r"\b(?:METAL|metal)\b")),
)


def audit_text(source, text):
    findings = []
    for label, pattern in NONCANONICAL:
        for match in pattern.finditer(text):
            findings.append((source, label, match.group(0)))
    return findings


def audit_docx(path):
    document = Document(path)
    findings = []
    for index, paragraph in enumerate(document.paragraphs):
        findings.extend(
            audit_text(f"{path}:paragraph:{index}", paragraph.text)
        )
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                text = " ".join(paragraph.text for paragraph in cell.paragraphs)
                # Exact provenance filenames are identifiers, not display notation.
                if "simpls_vs_plssvd" in text:
                    continue
                findings.extend(
                    audit_text(
                        f"{path}:table:{table_index}:{row_index}:{cell_index}",
                        text,
                    )
                )
    return findings


def audit_pdf(path):
    result = subprocess.run(
        ["pdftotext", "-layout", str(path), "-"],
        capture_output=True,
        text=True,
        check=False,
    )
    if result.returncode != 0:
        return [(str(path), "PDF extraction failed", result.stderr.strip())]
    # Poppler flattens a visually superscripted Word equation, Q² = ..., to Q2.
    extracted = re.sub(r"(?<![A-Za-z0-9])Q2(?=\s*=)", "Q²", result.stdout)
    return audit_text(str(path), extracted)


def main(arguments):
    findings = []
    for argument in arguments:
        path = Path(argument)
        if path.suffix.lower() == ".docx":
            findings.extend(audit_docx(path))
        elif path.suffix.lower() == ".pdf":
            findings.extend(audit_pdf(path))
        else:
            raise SystemExit(f"Unsupported audit target: {path}")

    if findings:
        for source, label, token in findings:
            print(f"FAIL {source}: {label}: {token!r}")
        raise SystemExit(1)

    print(f"PASS: canonical notation in {len(arguments)} files")


if __name__ == "__main__":
    main(sys.argv[1:])
