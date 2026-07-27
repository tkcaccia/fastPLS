#!/usr/bin/env python3

from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path

from docx import Document


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle30"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle31"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle30_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle30_0.99.6_20260725.docx"
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle30_20260725.docx"
)

MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle31_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle31_0.99.6_20260725.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle31_20260725.docx"

spec = spec_from_file_location(
    "cycle9_table_helpers",
    ROOT / "tools" / "revise_cmpb_cycle9_simpls_validation.py",
)
c9 = module_from_spec(spec)
spec.loader.exec_module(c9)


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def revise_main():
    document = Document(MAIN_SOURCE)
    paragraph = find_paragraph(document, "The R package includes bundled IRLBA code")
    anchor = (
        "Because a controlled thread-scaling experiment was not performed, "
        "no CPU multithreading speedup is claimed. "
    )
    addition = (
        "Exact package versions, compiler and accelerator flags, BLAS/LAPACK "
        "implementations, thread settings, and GPU libraries are reported in "
        "Supplementary Table S4a. "
    )
    if anchor not in paragraph.text:
        raise RuntimeError("Reproducibility-table insertion point not found")
    paragraph.text = paragraph.text.replace(anchor, anchor + addition, 1)
    document.core_properties.title = (
        "fastPLS CMPB manuscript - exact reproducibility environment"
    )
    document.save(MAIN_OUT)


def reproducibility_rows():
    return [
        (
            "Platform",
            "Ubuntu 22.04.5 LTS; Linux 6.8.0-124; Intel Core i7-13700 "
            "(16 cores, 24 logical CPUs); 32,526,480 kB RAM; NVIDIA GeForce "
            "RTX 5060 Ti, 16,311 MiB",
            "macOS 14.5 (23F79); Apple M3 (8 CPU cores, 10 GPU cores); "
            "8,589,934,592 bytes unified memory; Metal 3",
        ),
        (
            "R and fastPLS",
            "R 4.6.0; fastPLS 0.99.6; source 6e50bd3 for the precision and "
            "external-comparison runs",
            "R 4.6.0; fastPLS 0.99.6",
        ),
        (
            "Core R packages",
            "Rcpp 1.1.1-1.1; RcppArmadillo 15.2.7-1; RcppEigen 0.3.4.0.2; "
            "Matrix 1.7-5; float 0.3-3; pls 2.9-0",
            "Rcpp 1.1.2; RcppArmadillo 15.4.0-1; RcppEigen 0.3.4.0.2; "
            "Matrix 1.7-5; float 0.3-3; pls 2.9-0",
        ),
        (
            "External comparison packages",
            "mdatools 0.15.0; chemometrics 1.4.4; pcv 1.1.0; plsdepot "
            "0.3.1; plsgenomics 1.5-3; mixOmics 6.36.0; spls 2.3-2; "
            "ropls 1.44.0",
            "Not used for the external-package comparison",
        ),
        (
            "C/C++ compiler",
            "GCC/G++ 11.4.0; C++17",
            "Homebrew Clang/Clang++ 22.1.1; C++17; macOS SDK 15.2",
        ),
        (
            "Effective C/C++ flags",
            "-std=gnu++17 -fpic -g -O2 "
            "-ffile-prefix-map=/build/r-base-gN72Ro/r-base-4.6.0=. "
            "-fstack-protector-strong -Wformat -Werror=format-security "
            "-Wdate-time -D_FORTIFY_SOURCE=2; package definitions: "
            "-DRANN -DARMA_64BIT_WORD=1 -DFASTPLS_HAS_CUDA "
            "-DFASTPLS_HAS_CUDA_KERNELS",
            "-std=gnu++17 -fPIC -falign-functions=64 -Wall -g -O2; "
            "Objective-C++: -x objective-c++ -fobjc-arc; package definitions: "
            "-DRANN -DARMA_64BIT_WORD=1 -DFASTPLS_HAS_METAL",
        ),
        (
            "Accelerator build/link",
            "NVCC 13.0.88: -std=c++17 --extended-lambda "
            "--expt-relaxed-constexpr -DFASTPLS_HAS_CUDA "
            "-DFASTPLS_HAS_CUDA_KERNELS -Xcompiler -fPIC; linked with "
            "-lcudart -lcublas -lcusolver -lcurand",
            "Linked frameworks: Foundation, Metal, and "
            "MetalPerformanceShaders",
        ),
        (
            "BLAS/LAPACK",
            "Reference BLAS 3.10.0 "
            "(/usr/lib/x86_64-linux-gnu/blas/libblas.so.3.10.0); LAPACK "
            "3.10.0 "
            "(/usr/lib/x86_64-linux-gnu/lapack/liblapack.so.3.10.0)",
            "R BLAS "
            "(/Library/Frameworks/R.framework/Versions/4.6/Resources/lib/"
            "libRblas.0.dylib); R LAPACK 3.12.1 "
            "(libRlapack.dylib)",
        ),
        (
            "CPU thread setting",
            "1 effective BLAS thread; reference BLAS and no OpenMP; "
            "OMP_NUM_THREADS, OPENBLAS_NUM_THREADS, and MKL_NUM_THREADS unset",
            "1 effective R-BLAS thread; no OpenMP; OMP_NUM_THREADS and "
            "VECLIB_MAXIMUM_THREADS unset",
        ),
        (
            "GPU runtime/libraries",
            "NVIDIA driver 595.71.05; CUDA SDK 13.0.3; CUDA runtime "
            "13.0.96; cuBLAS 13.1.1.3; cuSOLVER 12.0.4.66; cuRAND "
            "10.4.0.35",
            "Metal 3; Metal.framework and MetalPerformanceShaders.framework "
            "from macOS SDK 15.2",
        ),
    ]


def revise_supplement():
    document = Document(SUPP_SOURCE)
    anchor = find_paragraph(
        document,
        "The benchmark record includes package commit",
    )
    old_caption = find_paragraph(
        document,
        "Table S4. External implementation scope.",
    )
    old_caption.text = old_caption.text.replace("Table S4.", "Table S4b.", 1)

    caption = document.add_paragraph(
        "Table S4a. Exact reproducibility environment for the primary CUDA "
        "workstation and the Apple Metal validation system. Values were read "
        "from benchmark manifests, sessionInfo(), R configuration, package "
        "installation logs, linked runtime libraries, and operating-system "
        "queries. The CUDA and Metal columns are not direct hardware "
        "comparisons.",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True
    table = c9.add_table(
        document,
        ["Item", "CUDA workstation", "Metal validation system"],
        reproducibility_rows(),
        [1.10, 2.62, 2.62],
        font_size=5.2,
    )
    anchor._p.addnext(caption._p)
    caption._p.addnext(table._tbl)

    document.core_properties.title = (
        "fastPLS CMPB supplement - exact reproducibility environment"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    heading = document.add_heading(
        "31. Exact computational environment was not tabulated",
        level=1,
    )
    heading.paragraph_format.keep_with_next = True
    comment = document.add_paragraph(
        "Reviewer comment: Provide exact package versions, compiler flags, "
        "BLAS implementation, thread counts, and GPU libraries in a concise "
        "reproducibility table."
    )
    comment.paragraph_format.keep_with_next = True
    response = document.add_paragraph(
        "Response: Corrected. Supplementary Table S4a now reports the CUDA "
        "workstation and Metal validation system side by side, including R and "
        "package versions, external-comparator versions, C/C++ and accelerator "
        "compilers, effective compile definitions and flags, linked BLAS/LAPACK "
        "implementations, effective CPU thread counts, GPU hardware and driver, "
        "and exact CUDA or Metal libraries. The values were recovered from the "
        "benchmark manifests, session records, installation logs, linked "
        "libraries, and operating-system queries. The main Methods now points "
        "readers directly to this table."
    )
    response.paragraph_format.keep_together = True
    document.core_properties.title = (
        "fastPLS CMPB response to reviewers - reproducibility table"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    revise_response()


if __name__ == "__main__":
    main()
