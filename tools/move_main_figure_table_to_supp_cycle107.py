from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.table import Table

from portrait_supplement_cycle106 import set_cell_margins, set_table_font, set_table_width


ROOT = Path(__file__).resolve().parents[1]
MAIN_SOURCE = (
    ROOT
    / "artifacts"
    / "CMPB_rewrite_20260825_cycle105"
    / "fastPLS_CMPB_main_cycle105_0.99.25_20260825.docx"
)
SUPP_SOURCE = (
    ROOT
    / "artifacts"
    / "CMPB_rewrite_20260826_cycle106"
    / "fastPLS_CMPB_supplement_cycle106_0.99.25_20260826.docx"
)
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260826_cycle107"
MAIN_OUTPUT = OUTPUT_DIR / "fastPLS_CMPB_main_cycle107_0.99.25_20260826.docx"
SUPP_OUTPUT = OUTPUT_DIR / "fastPLS_CMPB_supplement_cycle107_0.99.25_20260826.docx"
FIGURE = (
    ROOT
    / "benchmark_results"
    / "frozen_release_0.99.25"
    / "figures"
    / "Figure_2_frozen_external_simpls.png"
)


def paragraph_by_prefix(document, prefix):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}; found {len(matches)}")
    return matches[0]


def replace_text(document, replacements):
    for paragraph in document.paragraphs:
        value = paragraph.text
        for old, new in replacements:
            value = value.replace(old, new)
        if value != paragraph.text:
            paragraph.text = value


def remove_figure_and_caption(caption):
    parent = caption._p.getparent()
    previous = caption._p.getprevious()
    while previous is not None:
        text = "".join(previous.itertext()).strip()
        has_drawing = bool(previous.xpath(".//w:drawing"))
        if text and not has_drawing:
            break
        candidate = previous
        previous = previous.getprevious()
        parent.remove(candidate)
    parent.remove(caption._p)


def insert_figure_before(document, anchor, image_path, caption_text):
    figure = document.add_paragraph()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.keep_with_next = True
    figure.add_run().add_picture(str(image_path), width=Inches(7.15))
    anchor._p.addprevious(figure._p)

    caption = document.add_paragraph(style=anchor.style)
    caption.text = caption_text
    caption.paragraph_format.keep_with_next = False
    anchor._p.addprevious(caption._p)


def insert_table_before(document, anchor, table_xml, caption_text, caption_style):
    caption = document.add_paragraph(style=caption_style)
    caption.text = caption_text
    caption.paragraph_format.keep_with_next = True
    anchor._p.addprevious(caption._p)

    copied_xml = deepcopy(table_xml)
    anchor._p.addprevious(copied_xml)
    copied = Table(copied_xml, document._body)
    set_table_width(copied, 10944)
    set_cell_margins(copied)
    set_table_font(copied)


def main():
    if not FIGURE.is_file():
        raise FileNotFoundError(FIGURE)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    main_doc = Document(MAIN_SOURCE)
    supp_doc = Document(SUPP_SOURCE)

    # Preserve the recommendation table before removing it from the main text.
    recommendation_table = main_doc.tables[1]._tbl

    remove_figure_and_caption(paragraph_by_prefix(main_doc, "Figure 2."))
    table_caption = paragraph_by_prefix(main_doc, "Table 1.")
    table_parent = recommendation_table.getparent()
    table_parent.remove(recommendation_table)
    table_caption._p.getparent().remove(table_caption._p)

    replace_text(main_doc, [
        (
            "The two timing profiles answer different questions and are not pooled "
            "(Figure 2; Supplementary Tables S10a-S10d).",
            "The two timing profiles answer different questions and are not pooled "
            "(Supplementary Figure S3 and Tables S10a-S10d).",
        ),
        ("(Figure 3; Supplementary Table S11)", "(Figure 2; Supplementary Table S11)"),
        ("Figure 4 displays held-out sample", "Figure 3 displays held-out sample"),
        ("Figure 3. Selected archived-release", "Figure 2. Selected archived-release"),
        ("Figure 4. Archived-release NMR", "Figure 3. Archived-release NMR"),
        ("Table 1 summarizes these decisions.", "Supplementary Table S6a summarizes these decisions."),
        ("Table S13, and Figure S3", "Table S13, and Figure S4"),
    ])

    # The recommendation table follows the evidence-status glossary.
    s12 = paragraph_by_prefix(supp_doc, "S12. Deterministic estimator validation")
    table_caption_style = paragraph_by_prefix(supp_doc, "Table S6.").style
    insert_table_before(
        supp_doc,
        s12,
        recommendation_table,
        "Table S6a. Practical interpretation of the computational evidence and recommended fastPLS route.",
        table_caption_style,
    )

    # The external-workflow figure follows the detailed repeated-comparison tables.
    s10e = paragraph_by_prefix(supp_doc, "Table S10e.")
    moved_caption = (
        "Figure S3. Repeated deterministic float64 single-CPU SIMPLS public workflows with fastPLS "
        "0.99.25 and pls 2.9.0. Panels show median total fitting-plus-prediction time, baseline-corrected "
        "complete-process peak RSS, and held-out argmax accuracy. Error bars are IQRs from three fresh "
        "processes per method-dataset pair. Splits and component counts were identical, and accuracy was "
        "identical for every pair. Full values and the separate minimum-output comparison are reported in "
        "Tables S10a-S10d."
    )
    insert_figure_before(supp_doc, s10e, FIGURE, moved_caption)

    replace_text(supp_doc, [
        ("Figure S3. Historical, partially reproducible", "Figure S4. Historical, partially reproducible"),
        ("Figure 4 displays AMI-0030-9", "Figure 3 displays AMI-0030-9"),
    ])

    main_doc.save(MAIN_OUTPUT)
    supp_doc.save(SUPP_OUTPUT)
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)


if __name__ == "__main__":
    main()
