#!/usr/bin/env python3
"""Standardize reader-facing numerical and backend notation."""

from pathlib import Path

from docx import Document

from revise_cmpb_cycle67_consolidate_evidence import (
    normalize_submission_terminology,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle73"
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle74"
MAIN_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_main_cycle73_0.99.6_20260726.docx"
)
SUPP_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_supplement_cycle73_0.99.6_20260726.docx"
)
MAIN_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_main_cycle74_0.99.6_20260726.docx"
)
SUPP_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_supplement_cycle74_0.99.6_20260726.docx"
)


REPLACEMENTS = (
    ("RMSD, Q2, and", "RMSD, Q², and"),
    ("; Q2 ", "; Q² "),
)


def replace_in_paragraph(paragraph):
    changed = False
    for run in paragraph.runs:
        revised = run.text
        for old, new in REPLACEMENTS:
            revised = revised.replace(old, new)
        if revised != run.text:
            run.text = revised
            changed = True
    return changed


def revise_document(document):
    changes = 0
    for paragraph in document.paragraphs:
        changes += int(replace_in_paragraph(paragraph))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    changes += int(replace_in_paragraph(paragraph))
    normalize_submission_terminology(document)
    return changes


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    main_document = Document(MAIN_SOURCE)
    main_changes = revise_document(main_document)
    main_document.save(MAIN_OUTPUT)

    supplement_document = Document(SUPP_SOURCE)
    supplement_changes = revise_document(supplement_document)
    supplement_document.save(SUPP_OUTPUT)

    print(f"main_changes={main_changes}")
    print(f"supplement_changes={supplement_changes}")
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)


if __name__ == "__main__":
    main()
