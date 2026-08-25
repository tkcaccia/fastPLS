#!/usr/bin/env python3

from pathlib import Path
from shutil import copy2

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle91"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle92"
RESULTS = ROOT / "benchmark_results" / "external_simpls_timing_publication_20260825"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle91_0.99.25_20260825.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle91_0.99.25_20260825.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle92_0.99.25_20260825.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle92_0.99.25_20260825.docx"
FIGURE = RESULTS / "external_simpls_timing_profiles.png"


def paragraph_start(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_paragraph(document, prefix, text):
    paragraph = paragraph_start(document, prefix)
    paragraph.text = text
    return paragraph


def set_cell_margins(cell, top=45, start=55, bottom=45, end=55):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table, font_size=4.8):
    table.style = "Table"
    table.autofit = True
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_index == 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
                    run.bold = row_index == 0


def insert_table_before(document, anchor, caption_text, columns, rows, font_size=4.8):
    caption = anchor.insert_paragraph_before(caption_text, style="Caption")
    table = document.add_table(rows=1, cols=len(columns))
    for index, label in enumerate(columns):
        table.rows[0].cells[index].text = str(label)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    style_table(table, font_size)
    anchor._p.addprevious(table._tbl)
    return caption, table


def replace_figure_before_caption(document, caption_prefix, image_path, width=5.8):
    caption = paragraph_start(document, caption_prefix)
    previous = caption._p.getprevious()
    while previous is not None and not previous.xpath(".//w:drawing"):
        previous = previous.getprevious()
    if previous is None:
        raise RuntimeError(f"Drawing before {caption_prefix} not found")
    for child in list(previous):
        previous.remove(child)
    paragraph = next(p for p in document.paragraphs if p._p is previous)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))


def fmt_time(median, iqr):
    if median < 0.1:
        return f"{median:.3f} ({iqr:.3f})"
    return f"{median:.2f} ({iqr:.2f})"


def build_pair_rows():
    pairs = pd.read_csv(RESULTS / "external_simpls_timing_pairs.csv")
    labels = {
        "ccle": "CCLE", "cifar100": "CIFAR-100", "gtex_v8": "GTEx v8",
        "metref": "MetRef", "retina": "Retina", "tabula": "Tabula Muris",
        "tcga_brca": "TCGA-BRCA", "tcga_hnsc_methylation": "TCGA-HNSC methyl.",
        "tcga_pan_cancer": "TCGA Pan-Cancer",
    }
    profile_labels = {
        "estimator_kernel": "Minimum outputs",
        "complete_workflow": "Public workflow",
    }
    rows = []
    for _, row in pairs.sort_values(["dataset", "comparison_profile"]).iterrows():
        rows.append([
            labels[row.dataset], profile_labels[row.comparison_profile],
            fmt_time(row.median_total_sec_fastpls, row.iqr_total_sec_fastpls),
            fmt_time(row.median_total_sec_pls, row.iqr_total_sec_pls),
            f"{row.speedup_pls_over_fastpls:.2f}",
            f"{row.median_accuracy_fastpls:.4f} / {row.median_accuracy_pls:.4f}",
            f"{int(row.repetitions_completed_fastpls)} / {int(row.repetitions_completed_pls)}",
        ])
    return rows


def revise_main():
    document = Document(MAIN_SOURCE)
    replace_paragraph(
        document,
        "Results: Deterministic fastPLS SIMPLS",
        "Results: Deterministic fastPLS SIMPLS met the prespecified tolerances in all 117 component-level comparisons. In 108 repeated isolated-process comparisons with pls::simpls.fit, accuracy was identical throughout. fastPLS was faster on four of nine datasets when both implementations retained the full coefficient path and suppressed dense scores, loadings, and fitted values, and on five of nine under ordinary public workflows; the largest workflow speed-up was 4.81-fold. In 486 controlled CPU/CUDA runs, qualified CUDA crossovers appeared at 5,000 observations and 2,000 predictors. Qualified CPU rSVD reduced 50-component NMR SIMPLS time from 350.7 to 9.8 s. The noncanonical single-run ImageNet analysis remains exploratory.",
    )
    results_heading = paragraph_start(document, "3. Results")
    methods_heading = paragraph_start(document, "3. Results")
    methods_heading.insert_paragraph_before(
        "External timing used two prespecified profiles. The minimum-output profile compared deterministic float64 SIMPLS while requiring the complete coefficient path, centering quantities, and final held-out predictions; scores, loadings, fitted-value arrays, and variance summaries were suppressed, with pls::simpls.fit called using stripped = TRUE. The public-workflow profile retained each implementation's ordinary model object and final predictions. Package and data loading preceded timing, no numerical warm-up was used, and every method-dataset repetition ran in a fresh process with one effective BLAS thread and a common 10,000-s timeout. Three repetitions were attempted for every pair; medians, IQRs, completed repetitions, failures, object sizes, and process RSS were retained.",
        style="Body Text",
    )
    replace_paragraph(
        document,
        "3.1 Shape-aware SIMPLS compared with independent R implementations",
        "3.1 Repeated comparison with an independent SIMPLS implementation",
    )
    replace_paragraph(
        document,
        "The float64 single-CPU comparison attempted 126 external-package runs",
        "The strict comparison completed all 108 planned runs: nine datasets, two output profiles, two implementations, and three fresh-process repetitions. Accuracy was identical for every pair. With minimum common prediction outputs, fastPLS was faster on four datasets and pls::simpls.fit on five; the largest fastPLS advantage was 1.54-fold on GTEx v8. Under ordinary public workflows, fastPLS was faster on five datasets, including 2.50-fold on CIFAR-100, 3.29-fold on Retina, and 4.81-fold on Tabula Muris. Corresponding accuracies were 0.8739 (8,739/10,000; Wilson 95% CI 0.8672-0.8803), 0.9678 (21,684/22,406; 0.9654-0.9700), and 0.8006 (40,077/50,059; 0.7971-0.8041). These timing profiles answer different questions and are not pooled (Figure 2; Supplementary Tables S10a-S10b).",
    )
    replace_paragraph(
        document,
        "Absolute process RSS was not uniformly reduced",
        "Workflow gains partly reflected output policy. On CIFAR-100, the compact fastPLS fit object was 1.38 MB versus 7,778 MB for ordinary pls::simpls.fit, and process RSS was 2.20 versus 13.42 GB. With complete coefficient paths required from both, object sizes were 59.30 and 58.60 MB and the speed-up narrowed to 1.17-fold. The broader package panel remains a workflow comparison with implementation-specific outputs (Supplementary Table S10c).",
    )
    replace_figure_before_caption(document, "Figure 2.", FIGURE)
    caption = replace_paragraph(
        document,
        "Figure 2.",
        "Figure 2. Repeated float64 single-CPU SIMPLS timing against pls::simpls.fit. Points are median fit-plus-prediction times; bars are IQRs from three fresh processes. Left: complete coefficient paths and final predictions with dense scores, loadings, and fitted arrays suppressed. Right: ordinary public objects. Accuracy was identical in every pair.",
    )
    caption.paragraph_format.keep_with_next = False
    caption.paragraph_format.keep_together = True
    replace_paragraph(
        document,
        "The principal contribution of fastPLS is a shape-dependent SIMPLS execution",
        "The principal contribution of fastPLS is a shape-dependent SIMPLS execution and storage layer, not a universally faster numerical kernel. The repeated external comparison showed identical predictions but route-dependent timing: minimum-output kernel differences were modest and changed sign across matrix shapes, whereas larger public-workflow gains on dense tasks also reflected compact output storage. Approximation, precision, and accelerator results remain conditional on numerical qualification.",
    )
    OUT.mkdir(parents=True, exist_ok=True)
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)
    replace_paragraph(
        document,
        "The primary software comparison used float64",
        "The definitive repeated comparison used float64, deterministic CPU SIMPLS, identical splits and component counts, one effective BLAS thread, and a 10,000-s timeout. Package and data loading occurred before timing; no numerical warm-up was performed. Every method-dataset pair ran in three fresh R processes. Two profiles were kept separate. In the minimum-output profile, fastPLS retained the full coefficient path, means, and compact internal prediction/audit factors, while pls::simpls.fit(stripped = TRUE) retained the full coefficient path and means; neither returned score, loading, fitted-value, residual, or variance arrays. In the public-workflow profile, fastPLS retained its ordinary compact object and variance summary, whereas pls::simpls.fit retained coefficients, X/Y scores and loadings, projection, fitted values, residuals, and X-variance quantities. All 108 runs completed and all paired accuracies were identical. The older 126-run multi-package panel is retained as a workflow comparison only: 110 completed, 12 were package limitations, two timed out, and two errored. It is not labelled an estimator-kernel comparison because outputs and model families differ.",
    )
    anchor = paragraph_start(document, "Table S10.")
    anchor.text = "Table S10c. Broad external-package workflow comparison. Cells report accuracy; total fitting-plus-prediction time; and peak process RSS. Implementations retain package-specific outputs, and unsupported runs, timeouts, and errors remain explicit."
    insert_table_before(
        document,
        anchor,
        "Table S10a. Repeated deterministic SIMPLS comparison. Times are median (IQR) seconds from three fresh processes. Speed-up is pls::simpls.fit divided by fastPLS; values above one favor fastPLS.",
        ["Dataset", "Profile", "fastPLS s", "pls s", "Speed-up", "Accuracy fastPLS / pls", "Completed reps"],
        build_pair_rows(),
        4.7,
    )
    insert_table_before(
        document,
        anchor,
        "Table S10b. Output-materialization contract used for the repeated comparison.",
        ["Profile", "fastPLS retained outputs", "pls::simpls.fit retained outputs", "Interpretation"],
        [
            [
                "Minimum outputs",
                "Full coefficient path, means, compact prediction/audit factors; no scores, loadings, fitted arrays, or variance summary",
                "stripped=TRUE: full coefficient path and means only",
                "Closest available common prediction endpoint; reports numerical-kernel plus unavoidable implementation state",
            ],
            [
                "Public workflow",
                "Ordinary compact fit object, means, latent prediction factors, variance summary; final predictions",
                "Ordinary object with coefficient, score, loading, projection, fitted, residual, and variance paths; final predictions",
                "Complete user workflow; time and memory include each package's normal output policy",
            ],
        ],
        5.2,
    )
    anchor.paragraph_format.page_break_before = True
    # Update the evidence-map authority and add a provenance record.
    evidence = document.tables[7]
    for row in evidence.rows[1:]:
        if row.cells[0].text.strip() == "fastPLS versus independent R implementations":
            row.cells[1].text = "Tables S10a-S10c"
            row.cells[2].text = "Repeated matched-output and ordinary-workflow timing; broad package workflow scope"
    provenance = document.tables[-1]
    cells = provenance.add_row().cells
    values = [
        "A22", "Tables S10a-S10b; Figure 2",
        "benchmark_results/external_simpls_timing_publication_20260825",
        "0.99.25", "ba80b65f0c66; archive SHA-256 74e134ef22d5",
        "scripts/run_external_simpls_timing.sh", "generated 2026-08-25",
    ]
    for cell, value in zip(cells, values):
        cell.text = value
    style_table(provenance, 4.7)
    OUT.mkdir(parents=True, exist_ok=True)
    document.save(SUPP_OUT)


if __name__ == "__main__":
    revise_main()
    revise_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)
