"""Create cycle-6 CMPB drafts with measured precision limitations and provenance."""

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.shared import Pt


ROOT = Path("/Users/stefano/Desktop/fastPLS/CMPB_rewrite_20260724_cycle5")
OUT = Path("/Users/stefano/Desktop/fastPLS/CMPB_rewrite_20260724_cycle6")
OUT.mkdir(parents=True, exist_ok=True)
MAIN_SRC = ROOT / "fastPLS_CMPB_main_cycle5_0.99.6_20260724.docx"
SUPP_SRC = ROOT / "fastPLS_CMPB_supplement_cycle5_0.99.6_20260724.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle6_0.99.6_20260724.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle6_0.99.6_20260724.docx"
REVIEW_OUT = OUT / "fastPLS_CMPB_independent_reviewer_report_cycle6_20260724.docx"


def replace_paragraph(doc, startswith, replacement):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(startswith):
            paragraph.text = replacement
            return True
    return False


def set_row(table, row, values):
    for cell, value in zip(table.rows[row].cells, values):
        cell.text = value


def compact_table(table, size=7):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = 0.9
                for run in paragraph.runs:
                    run.font.size = Pt(size)


def trim_rows(table, n_rows):
    while len(table.rows) > n_rows:
        table._tbl.remove(table.rows[-1]._tr)


def revise_main():
    copy2(MAIN_SRC, MAIN_OUT)
    doc = Document(MAIN_OUT)
    replace_paragraph(
        doc,
        "Double precision is the numerical reference.",
        "Double precision is the numerical reference. Input objects from the float package automatically select a float32 route for PLS-SVD, SIMPLS, OPLS, and kernel PLS on supported CPU, CUDA, and Metal builds. A controlled smoke screen covering all four families on a fixed classification task and a fixed univariate-regression task showed identical decoded labels between float32 and float64 on CPU, CUDA, and Metal; regression relative prediction differences were at most 2.4×10−6. These checks establish numerical compatibility, not a general speed claim. CUDA and Metal still contain host-side reduced-factorization or model-assembly stages, which are listed explicitly in Supplementary Table S1."
    )
    replace_paragraph(
        doc,
        "External comparisons used independent R implementations",
        "External comparisons use independent R implementations where model and response type can be matched, including functions from pls, mdatools, plsdepot, pcv, plsgenomics, chemometrics, mixOmics, spls, and ropls. The primary estimator-matched software comparison is performed with float64 inputs for both fastPLS and external packages; float32 is evaluated separately as a fastPLS precision and storage capability. Comparisons match preprocessing, response representation, component count, prediction head, thread setting, and timed work. Complete-workflow comparisons that use different model families or classification heads are reported separately and are not interpreted as estimator-only speed tests."
    )
    replace_paragraph(
        doc,
        "In a fixed-score validation of the revised LDA path",
        "In a fixed-score validation of the revised LDA path, float32 CPU and CUDA predictions agreed exactly across MetRef, CIFAR-100, and SingleCell at 2, 5, 10, and 20 components, with no factorization failures. The complementary four-family float32 smoke screen confirmed exact decoded-label agreement on CPU, CUDA, and Metal for PLS-SVD, SIMPLS, OPLS, and kernel PLS, and numerical agreement for univariate regression. A preliminary NMR precision control (ntrain=1200, ntest=321, p=13000, q=5000, ten components) preserved RMSD between float32 and float64 SIMPLS (6.4850×10−5 versus 6.4847×10−5), but float32 was slower in the current implementation (60.8 s CPU and 51.8 s CUDA versus 7.54 s float64 CPU). Float32 is therefore reported as a compatible reduced-precision path, not as a universal acceleration."
    )
    replace_paragraph(
        doc,
        "Float32 can reduce input and workspace storage",
        "Float32 can reduce input and workspace storage for supported routes, but it is not yet a universal performance benefit. The controlled backend checks establish agreement across the four public PLS families, whereas broader real-data memory and performance comparisons remain necessary before claiming a general memory advantage. Reproducibility should therefore be judged by prediction agreement, predictive metrics, selected components, numerical failures, and latent-subspace agreement rather than by assuming identical stochastic low-rank factors."
    )
    replace_paragraph(
        doc,
        "fastPLS combines an accelerated sequential SIMPLS implementation",
        "fastPLS combines an accelerated sequential SIMPLS implementation with memory-aware PLS-SVD, compiled validation, compact prediction, float32 input support, and CPU, CUDA, and Metal backends. These capabilities make established PLS workflows accessible for biomedical matrices that were previously limited by runtime or memory. The R package is distributed under GPL-3 and calls reusable C++ components maintained with the kodama-cpp codebase; CUDA and Metal backend code is presently maintained in the R-package layer. Future work will extend real-data precision validation, accelerator residency, and reproducibility testing across hardware architectures."
    )
    replace_paragraph(
        doc,
        "Double precision is the numerical reference.",
        "Double precision is the numerical reference. Inputs from the float package can select float32 arithmetic on supported CPU, CUDA, and Metal routes, but the matched real-data benchmark shows that this is not yet a universal speed or memory advantage. At 100 components on CIFAR-100, float32 CUDA reduced sampled GPU memory from 530 MB to 130–288 MB while preserving accuracy within 0.42 percentage points, but was slower than float64 because the current route is hybrid. On the full NMR task, float32 PLS-SVD retained similar RMSD (6.65×10−4 versus 6.45×10−4) but was markedly slower; float32 SIMPLS and linear kernel PLS had substantially worse RMSD. Float32 is therefore presented as a conditional reduced-precision capability, not as a generally validated replacement for float64. CUDA and Metal host-side stages are listed in Supplementary Table S1."
    )
    replace_paragraph(
        doc,
        "In a fixed-score validation of the revised LDA path",
        "In a fixed-score validation of the revised LDA path, float32 CPU and CUDA predictions agreed exactly across MetRef, CIFAR-100, and SingleCell at 2, 5, 10, and 20 components, with no factorization failures. This classifier-level check does not establish end-to-end float32 equivalence for every model shape. The matched full-resolution NMR results show that end-to-end float32 SIMPLS currently needs further numerical and performance work, whereas float32 PLS-SVD retained a similar held-out RMSD. Accordingly, the manuscript reports float32 results by model, backend, and dataset rather than pooling them into a single claim."
    )
    replace_paragraph(
        doc,
        "Float32 can reduce input and workspace storage",
        "Float32 can reduce input and workspace storage for selected routes, as observed for CUDA CIFAR-100, but it is not yet a universal performance or numerical-accuracy benefit. The full NMR comparison identifies an adverse SIMPLS/linear-kernel regime and is retained in the supplementary results. Reproducibility should therefore be judged by measured prediction agreement, predictive metrics, selected components, numerical failures, and latent-subspace agreement rather than by precision alone."
    )
    doc.save(MAIN_OUT)


def revise_supplement():
    copy2(SUPP_SRC, SUPP_OUT)
    doc = Document(SUPP_OUT)
    replace_paragraph(
        doc,
        "The double-precision CPU backend is the broadest reference implementation.",
        "The double-precision CPU backend is the broadest reference implementation. Native float32 is selected automatically when input matrices are float::float32 objects and is available for PLS-SVD, SIMPLS, OPLS, and kernel PLS on supported CPU, CUDA, and Metal builds. Native Windows float32 remains unavailable because the Windows R BLAS/LAPACK toolchain used by the package does not expose the required single-precision symbols. Unsupported combinations stop instead of silently converting to double."
    )
    replace_paragraph(
        doc,
        "This supplement describes fastPLS version",
        "This supplement describes fastPLS version 0.99.6. The released benchmark manifest records the exact package and reusable C++-component commits used to generate each quantitative table."
    )
    replace_paragraph(
        doc,
        "Table S1. Supported model, precision, and backend combinations",
        "Table S1. Stage-level CPU, CUDA, and Metal residency. ‘Hybrid’ denotes an accelerated path with one or more host-side operations."
    )
    replace_paragraph(
        doc,
        "CUDA and Metal offload the large matrix products.",
        "CUDA and Metal offload the large matrix products, but neither backend should be described as universally device-resident. In the current float32 CUDA implementation, QR and the reduced decomposition are finalized in host float32 after GPU range products. Metal likewise retains selected host-side reduced operations and SIMPLS direction updates. Supplementary Table S1 records residency by stage rather than assigning a single label to an entire model family."
    )
    replace_paragraph(
        doc,
        "The float32 CPU and CUDA routes retain score, covariance, factorization, and prediction buffers in single precision.",
        "The float32 CPU, CUDA, and Metal routes preserve the input precision through their supported PLS arithmetic. CUDA LDA uses single-precision score, covariance, factorization, and prediction buffers; Metal presently uses accelerated score projection with host-side LDA factorization. A controlled CPU/CUDA/Metal screen verified identical decoded labels for all four PLS families on a fixed classification task and relative prediction differences no larger than 2.4×10−6 on a fixed univariate-regression task."
    )
    replace_paragraph(
        doc,
        "S5.4 CBMC CITE-seq and SingleCell",
        "S5.4 CBMC CITE-seq, Retina, and Tabula Muris"
    )
    replace_paragraph(
        doc,
        "CBMC CITE-seq uses one measured modality as predictors",
        "CBMC CITE-seq uses one measured modality as predictors and the matched multivariate modality as response after cell-level filtering and modality-specific normalization. Retina and Tabula Muris are treated as distinct single-cell classification benchmarks and are named separately throughout the released manifest. The manifest distinguishes preprocessing embedded in each source object from preprocessing performed by the benchmark loader."
    )
    replace_paragraph(
        doc,
        "Float32 NMR control.",
        "Float32 NMR control. On the fixed NMR partition, a ten-component SIMPLS screening control with q=5000 responses gave RMSD 6.4850×10−5 for float32 CPU, 6.4850×10−5 for float32 CUDA, and 6.4847×10−5 for float64 CPU. The current float32 CPU and CUDA paths were slower (60.8 and 51.8 s) than float64 CPU (7.54 s). The full q=28355 float32 path did not satisfy the predeclared runtime screen. Together with the small controlled all-family agreement screen, these results support numerical compatibility but not a universal float32 performance claim."
    )
    replace_paragraph(
        doc,
        "For SIMPLS, a prediction at component count  can be evaluated",
        "For SIMPLS, a prediction at component count k is evaluated through compact latent factors rather than a dense coefficient matrix. Large test matrices are processed in row blocks and intermediate score matrices are released after use. This limits retained prediction state to the current block and low-rank factors; it does not introduce a separate SVD solver."
    )
    replace_paragraph(
        doc,
        "Float32 NMR control.",
        "Float32 NMR control. The matched full-resolution benchmark used 1,200 training spectra, 321 test spectra, 13,000 predictors, 28,355 responses, 100 components, and three isolated repetitions. Float32 CUDA PLS-SVD retained similar RMSD to float64 (median 6.65×10−4 versus 6.45×10−4) but required 248 s versus 0.843 s and sampled 1,578 MB versus 664 MB of GPU memory. Float32 CUDA SIMPLS and linear kernel PLS were slower (about 544–546 s versus 2.89–2.90 s) and had materially worse RMSD (about 4.34×10−3 versus 7.21×10−4). These results are reported as a negative precision result: current float32 execution is not recommended for this extreme multivariate-response regime."
    )

    residency = doc.tables[0]
    rows = [
        ["Stage", "CPU", "CUDA", "Metal", "Residency / precision note"],
        ["Core PLS fit", "Compiled C++", "GPU products/range; host small solve", "MPS products; host direction update", "CUDA/Metal hybrid; float32 smoke-tested"],
        ["OPLS filter", "Compiled C++", "Host filter + CUDA core", "Host filter + MPS core", "Hybrid; float32 smoke-tested"],
        ["Kernel PLS", "Host kernel + C++ core", "Host nonlinear Gram + CUDA core", "Host nonlinear Gram + MPS core", "Nonlinear K is O(n²); hybrid"],
        ["Prediction", "Compiled projection", "CUDA projection; host model/result I/O", "MPS projection; host post-processing", "Transfers remain part of execution"],
        ["LDA", "Compiled covariance/Cholesky", "GPU moments/solve/score", "MPS projection + host solve", "float32 CPU/CUDA agreement verified"],
        ["10-fold CV", "Compiled folds/fits/scoring", "Host scheduler + sequential GPU folds", "Compiled/hybrid fold fits", "No R-level refit loop; no concurrent folds"],
    ]
    trim_rows(residency, len(rows))
    while len(residency.rows) < len(rows): residency.add_row()
    for i, values in enumerate(rows):
        set_row(residency, i, values)
    compact_table(residency, size=7)

    replace_paragraph(
        doc,
        "Table S2. Dominant computational and storage terms.",
        "Table S2. Dominant computational and storage terms. Here, a is the requested number of PLS components and l is the randomized-SVD sketch width (requested rank plus oversampling). Exact constants depend on matrix shape, solver convergence, and backend."
    )
    complexity = doc.tables[1]
    rows = [
        ["Path", "Dominant work", "Additional large storage", "Memory-saving mechanism"],
        ["Explicit PLS-SVD", "Form S = XᵀY: O(npq), then one truncated decomposition", "S: O(pq); latent factors O((p+q)a)", "One decomposition supplies every requested component prefix"],
        ["Implicit PLS-SVD", "For sketch width l, SΩ = Xᵀ(YΩ) and SᵀQ = Yᵀ(XQ)", "Sketches/factors O((p+q)l); no S", "Matrix-free products; label-aware class sums avoid dense indicator Y"],
        ["SIMPLS", "a sequential leading-direction solves plus score/loading products", "Explicit S: O(pq); bases O((p+q)a)", "Cached deflation products and incremental prediction path; rSVD reuse is explicitly approximate"],
        ["OPLS", "Explicit S = XᵀY and one-vector filter SVD, then SIMPLS predictive core", "Filter S: O(pq); bases O(p·north) plus inner-model storage", "No dense deflated X copy; the current orthogonal-filter stage remains explicit"],
        ["Nonlinear kernel PLS", "Kernel construction plus SIMPLS in sample space", "Centred Gram matrix O(n²)", "Blocked test prediction only; nonlinear kernel storage remains quadratic"],
        ["Compact prediction", "XtestR followed by latent response mapping", "Test-score block and latent factors, not all p×q coefficient prefixes", "Row-block streaming and low-rank latent mapping"],
    ]
    trim_rows(complexity, len(rows))
    while len(complexity.rows) < len(rows): complexity.add_row()
    for i, values in enumerate(rows):
        set_row(complexity, i, values)
    compact_table(complexity, size=7)

    doc.save(SUPP_OUT)


def create_review():
    doc = Document()
    doc.add_heading("Independent reviewer report: cycle 6 update", level=0)
    doc.add_paragraph("Manuscript: fastPLS: scalable partial least squares with compiled CPU and accelerator backends for high-dimensional biomedical data")
    doc.add_heading("Recommendation", level=1)
    doc.add_paragraph("Major revision")
    doc.add_heading("Assessment after cycle 6", level=1)
    doc.add_paragraph(
        "The revised material is more credible. It distinguishes GPU-accelerated hybrid paths from fully device-resident algorithms and reports the adverse full-NMR float32 SIMPLS result rather than presenting reduced precision as a universal benefit. The complexity table now defines its notation."
    )
    doc.add_heading("Remaining major comments", level=1)
    for text in [
        "The revised float64 external-package benchmark and the corrected complete real-data benchmark must be executed, not only specified in scripts. Supply raw repetitions, dispersion, peak memory, failures, requested versus executed estimator/backend, and a release manifest.",
        "The float32 discussion is substantially improved, but the manuscript must complete the prespecified OPLS precision measurements and state clearly in the Abstract, Methods, package documentation, and Conclusions that float32 performance is model- and shape-dependent. A reader must not infer that float32 CUDA is generally faster or fully resident from the current text.",
        "The paper must identify an immutable release/archive from which all final tables and figures were generated. The current local working-tree fixes are not a sufficient provenance record.",
        "The estimator-preservation evidence is strong for deterministic IRLBA SIMPLS. Clearly separate it in all text and figures from rSVD, which is stochastic and approximate."
    ]:
        doc.add_paragraph(text, style="List Number")
    doc.add_heading("Minor comments", level=1)
    for text in [
        "Use Retina and Tabula Muris consistently instead of the ambiguous label ‘SingleCell’ in every main-text and supplementary table.",
        "Retain the explicit statement that ImageNet is a computational stress test rather than biomedical validation.",
        "For the NMR precision table, report uncertainty measures and the number of successful repetitions alongside every float32-versus-float64 estimate."
    ]:
        doc.add_paragraph(text, style="List Bullet")
    doc.save(REVIEW_OUT)


revise_main()
revise_supplement()
create_review()
print(MAIN_OUT)
print(SUPP_OUT)
print(REVIEW_OUT)
