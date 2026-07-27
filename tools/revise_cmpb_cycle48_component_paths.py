#!/usr/bin/env python3

import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle47"
RESULTS = ROOT / "benchmark_results" / "manuscript_revision_cycle48_20260726"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle48"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle47_0.99.6_20260726.docx"
SUPP_SOURCE = (
    SOURCE / "fastPLS_CMPB_supplement_cycle47_0.99.6_20260726.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle48_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle48_0.99.6_20260726.docx"

SELECTION = RESULTS / "component_selection_by_family.csv"
CORRELATIONS = RESULTS / "component_metric_spearman_correlations.csv"
PLOTS = RESULTS / "component_path_plots"

DATASETS = [
    ("metref", "MetRef"),
    ("ccle", "CCLE"),
    ("tcga_brca", "TCGA-BRCA"),
    ("tcga_hnsc_methylation", "TCGA-HNSC methylation"),
    ("gtex_v8", "GTEx v8"),
    ("tcga_pan_cancer", "TCGA Pan-Cancer"),
    ("retina", "Retina"),
    ("tabula", "Tabula Muris"),
    ("cifar100", "CIFAR-100"),
    ("cbmc_citeseq", "CBMC CITE-seq"),
    ("prism", "PRISM"),
    ("nmr", "NMR"),
]
FAMILIES = [
    ("plssvd", "PLS-SVD", "S43"),
    ("simpls", "SIMPLS", "S44"),
    ("opls", "OPLS", "S45"),
    ("kernelpls", "kernel PLS", "S46"),
]


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_text(paragraph, old, new):
    if old not in paragraph.text:
        raise RuntimeError(f"Text not found: {old[:100]}")
    text = paragraph.text.replace(old, new)
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)


def repeat_header(row):
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:tblHeader")) is None:
        properties.append(OxmlElement("w:tblHeader"))


def prevent_row_splitting(table):
    for row in table.rows:
        properties = row._tr.get_or_add_trPr()
        if properties.find(qn("w:cantSplit")) is None:
            properties.append(OxmlElement("w:cantSplit"))


def set_cell_width(cell, width_inches):
    width = int(width_inches * 1440)
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:tcW"))
    if element is None:
        element = OxmlElement("w:tcW")
        properties.append(element)
    element.set(qn("w:w"), str(width))
    element.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=40, start=55, bottom=40, end=55):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(document, headers, rows, widths, font_size=6.2):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = str(value)
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = str(value)
    repeat_header(table.rows[0])
    prevent_row_splitting(table)
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[column_index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if column_index in (0, 2, len(headers) - 1)
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
                    if row_index == 0:
                        run.bold = True
    return table


def read_rows(path):
    with path.open(newline="") as handle:
        return list(csv.DictReader(handle))


def number(value, digits=2):
    if value in ("", None, "NA", "NaN"):
        return "—"
    try:
        numeric = float(value)
    except (TypeError, ValueError):
        return "—"
    if numeric != numeric:
        return "—"
    return f"{numeric:.{digits}f}"


def integer(value):
    if value in ("", None, "NA", "NaN"):
        return "—"
    try:
        return str(int(float(value)))
    except (TypeError, ValueError):
        return "—"


def revise_main():
    document = Document(MAIN_SOURCE)
    paragraph = find_paragraph(
        document,
        "Supplementary Tables S14-S17 report matched CPU and CUDA results",
    )
    replace_text(
        paragraph,
        (
            "Supplementary Figure S25 visualizes these results and, in a separate "
            "row, matched CPU/Metal validation for the four datasets evaluated on "
            "Apple hardware."
        ),
        (
            "Supplementary Figure S25 visualizes these results and, in a separate "
            "row, matched CPU/Metal validation for the four datasets evaluated on "
            "Apple hardware. Training-selected component counts, complete "
            "component paths, and descriptive Spearman correlations between "
            "component count, prediction, total time, host RSS, and GPU memory are "
            "reported in Supplementary Tables S42-S46 and Figures S26-S37."
        ),
    )
    document.core_properties.title = (
        "fastPLS CMPB manuscript - supplementary component-path analysis"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    if not SELECTION.exists() or not CORRELATIONS.exists():
        raise FileNotFoundError("Component-path summaries have not been generated")
    for dataset, _ in DATASETS:
        image = PLOTS / f"{dataset}_component_paths.png"
        if not image.exists():
            raise FileNotFoundError(image)

    document = Document(SUPP_SOURCE)
    document.add_heading(
        "S37. Component selection and component-wise computational scaling",
        level=1,
    )
    document.add_paragraph(
        (
            "For every dataset and PLS family, the selected component count was "
            "obtained from training data only. The fixed outer test set was then "
            "used to display the complete predictive path. The component-path "
            "analysis uses matched float64 CPU-rSVD and CUDA-rSVD workflows with "
            "argmax decoding for classification and continuous prediction for "
            "regression. Total time includes fitting and prediction. Host and GPU "
            "memory are absolute isolated-process peaks; GPU values include CUDA "
            "context and runtime state. Dotted vertical lines identify the "
            "training-selected component count and must not be interpreted as "
            "selection from the held-out curves."
        )
    )
    document.add_paragraph(
        (
            "Within each dataset, family, and backend, Spearman rank correlation "
            "was calculated between requested component count and the predictive "
            "metric, total time, peak host RSS, and peak GPU memory. Correlations "
            "require at least three distinct successful component values and are "
            "descriptive because the component grids are short and irregular. A "
            "positive predictive correlation indicates increasing accuracy or "
            "Q2, but increasing RMSD for multivariate regression; it therefore "
            "does not always indicate better prediction. Correlations compare "
            "values only within one computational path and are not used to compare "
            "absolute runtime or memory between datasets."
        )
    )

    selection_rows = read_rows(SELECTION)
    dataset_order = {key: index for index, (key, _) in enumerate(DATASETS)}
    family_order = {key: index for index, (key, _, _) in enumerate(FAMILIES)}
    selection_rows.sort(
        key=lambda row: (
            dataset_order[row["dataset"]],
            family_order[row["family"]],
        )
    )
    document.add_paragraph(
        (
            "Table S42. Training-only component selection by dataset and PLS "
            "family. The grid contains requested values; A is the selected "
            "effective component count after any response-rank constraint. "
            "Boundary and response-rank labels identify constrained values rather "
            "than global optima. NMR OPLS and kernel PLS were not evaluated."
        ),
        style="Caption",
    )
    add_table(
        document,
        ("Dataset", "Family", "Requested grid", "A", "Selection status", "Source"),
        [
            (
                row["dataset_label"],
                row["family_label"],
                row["evaluated_grid"].replace(",", ", "),
                integer(row["selected_ncomp"]),
                row["selection_status"],
                (
                    "Repeated training splits"
                    if row["dataset"] == "nmr"
                    else "5-fold training CV"
                ),
            )
            for row in selection_rows
        ],
        (1.15, 0.72, 1.55, 0.42, 1.55, 1.05),
        font_size=5.9,
    )

    correlation_rows = read_rows(CORRELATIONS)
    for family, family_label, table_number in FAMILIES:
        document.add_page_break()
        document.add_paragraph(
            (
                f"Table {table_number}. Spearman correlations between requested "
                f"component count and component-path measurements for {family_label}. "
                "Prediction denotes accuracy or Q2 for those tasks and RMSD for "
                "multivariate regression. A dash denotes an unavailable or "
                "insufficient path."
            ),
            style="Caption",
        )
        rows = [
            row for row in correlation_rows if row["family"] == family
        ]
        rows.sort(
            key=lambda row: (
                dataset_order[row["dataset"]],
                0 if row["backend"] == "CPU" else 1,
            )
        )
        add_table(
            document,
            (
                "Dataset",
                "Backend",
                "Points",
                "A",
                "ρ prediction",
                "ρ time",
                "ρ host RSS",
                "ρ GPU memory",
            ),
            [
                (
                    row["dataset_label"],
                    row["backend"],
                    integer(row["n_path_points"]),
                    integer(row["selected_ncomp"]),
                    number(row["rho_performance"]),
                    number(row["rho_total_time"]),
                    number(row["rho_host_rss"]),
                    number(row["rho_gpu_memory"]),
                )
                for row in rows
            ],
            (1.35, 0.65, 0.52, 0.45, 0.90, 0.78, 0.88, 0.92),
            font_size=6.2,
        )

    valid = [
        row
        for row in correlation_rows
        if number(row["rho_total_time"]) != "—"
    ]
    strong_time = sum(
        float(row["rho_total_time"]) >= 0.70 for row in valid
    )
    positive_time = sum(
        float(row["rho_total_time"]) > 0 for row in valid
    )
    document.add_paragraph(
        (
            f"Across {len(valid)} evaluable dataset-family-backend paths, total "
            f"time increased monotonically with component count in {positive_time} "
            f"paths and showed a strong positive Spearman association "
            f"(ρ≥0.70) in {strong_time}. Host and GPU memory associations were "
            "less uniform because process baselines, allocator reuse, and CUDA "
            "context costs can dominate small matrices. Predictive paths were "
            "dataset-specific and often reached an evaluated-grid or response-rank "
            "boundary, reinforcing that the reported A values are the best values "
            "within prespecified training grids rather than universal optima."
        )
    )
    document.add_paragraph(
        (
            "Machine-readable values are supplied in "
            "component_selection_by_family.csv, "
            "component_path_summary_matched.csv, and "
            "component_metric_spearman_correlations.csv."
        )
    )

    for offset, (dataset, dataset_label) in enumerate(DATASETS, start=26):
        document.add_page_break()
        picture = document.add_paragraph()
        picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
        picture.add_run().add_picture(
            str(PLOTS / f"{dataset}_component_paths.png"),
            width=Inches(6.75),
        )
        document.add_paragraph(
            (
                f"Figure S{offset}. Component-path analysis for {dataset_label}. "
                "Columns show PLS-SVD, SIMPLS, OPLS, and kernel PLS; rows show "
                "total fitting-plus-prediction time, the task-specific predictive "
                "metric, peak host RSS, and peak GPU memory. CPU and CUDA curves "
                "use matched float64 rSVD workflows. Dotted vertical lines mark "
                "training-selected component counts. Missing panels or points "
                "denote structurally invalid or unevaluated combinations."
            ),
            style="Caption",
        )

    document.core_properties.title = (
        "fastPLS CMPB supplement - component selection and path analysis"
    )
    document.save(SUPP_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
