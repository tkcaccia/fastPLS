#!/usr/bin/env python3
"""Keep the synchronized NMR provenance explicit but compact."""

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260727_cycle84"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260727_cycle85"
MAIN_IN = SOURCE / "fastPLS_CMPB_main_cycle84_0.99.6_20260727.docx"
SUPP_IN = SOURCE / "fastPLS_CMPB_supplement_cycle84_0.99.6_20260727.docx"
MAIN_OUT = OUTPUT / "fastPLS_CMPB_main_cycle85_0.99.6_20260727.docx"
SUPP_OUT = OUTPUT / "fastPLS_CMPB_supplement_cycle85_0.99.6_20260727.docx"


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    copy2(MAIN_IN, MAIN_OUT)

    document = Document(SUPP_IN)
    provenance = document.tables[18]
    for row in provenance.rows[1:]:
        if row.cells[0].text == "A17":
            row.cells[1].text = "Figure 4/Table S12; AMI-00BP-8 #155"
            row.cells[4].text = "archive and selection CSV recorded"
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.font.size = Pt(6)
            break
    else:
        raise RuntimeError("A17 provenance row not found")
    document.save(SUPP_OUT)
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
