#!/usr/bin/env python3

from pathlib import Path
import shutil

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle35"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle36"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle35_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle35_0.99.6_20260726.docx"
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle35_20260726.docx"
)

MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle36_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle36_0.99.6_20260726.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle36_20260726.docx"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def insert_after(paragraph, text, style=None):
    new_xml = paragraph._parent.add_paragraph()._p
    paragraph._p.addnext(new_xml)
    new_paragraph = Paragraph(new_xml, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    new_paragraph.add_run(text)
    return new_paragraph


def replace_nodes(container, replacements):
    """Replace text without rebuilding paragraphs, tables, or math objects."""
    changed = 0
    for node in container._element.xpath(".//w:t"):
        value = node.text or ""
        updated = value
        for old, new in replacements:
            updated = updated.replace(old, new)
        if updated != value:
            node.text = updated
            changed += 1
    return changed


def replace_paragraph(document, prefix, replacements):
    paragraph = find_paragraph(document, prefix)
    before = paragraph.text
    changed = replace_nodes(paragraph, replacements)
    if not changed:
        raise RuntimeError(f"No replacement made in paragraph: {prefix}")
    if paragraph.text == before:
        raise RuntimeError(f"Paragraph did not change: {prefix}")
    return paragraph


def replace_cell_exact(cell, expected, replacement):
    if cell.text.strip() != expected:
        raise RuntimeError(
            f"Expected table cell {expected!r}, found {cell.text.strip()!r}"
        )
    nodes = cell._tc.xpath(".//w:t")
    if not nodes:
        raise RuntimeError(f"No text node in table cell: {expected!r}")
    nodes[0].text = replacement
    for node in nodes[1:]:
        node.text = ""


def revise_main():
    document = Document(MAIN_SOURCE)

    replace_paragraph(
        document,
        "Methods: fastPLS provides",
        [
            (
                "Separately, stochastic rSVD",
                "Separately, randomized singular value decomposition (rSVD)",
            )
        ],
    )
    architecture = replace_paragraph(
        document,
        "The public pls() function",
        [("performs nested validation", "performs nested cross-validation")],
    )
    notation = insert_after(
        architecture,
        (
            "Notation is fixed throughout: a indexes the current latent "
            "component, A denotes the retained component count, and K denotes "
            "the number of cross-validation folds. Lowercase k is reserved for "
            "retrieval or evaluation cutoffs, such as the neighbourhood size "
            "in k-nearest-neighbour analysis and top-k accuracy; it is never "
            "used as a PLS component index. The public R argument ncomp "
            "supplies A. "
            "Cross-validation denotes training-set resampling for model "
            "selection or assessment; nested cross-validation denotes an inner "
            "selection loop and a distinct outer assessment loop."
        ),
        style="Body Text",
    )
    notation.paragraph_format.keep_together = True

    replace_paragraph(
        document,
        "The SIMPLS estimator follows",
        [("At component k,", "At component a,")],
    )
    replace_paragraph(
        document,
        "Within each dataset, methods used",
        [("five-fold cross-validation", "5-fold cross-validation")],
    )
    replace_paragraph(
        document,
        "OPLS and kernel settings",
        [
            ("The reported k is", "The reported A is"),
            ("k-1 predictive components", "A-1 predictive components"),
            ("five-fold validation", "5-fold cross-validation"),
        ],
    )
    replace_paragraph(
        document,
        "The ImageNet representation experiment",
        [
            (
                "PCA-rSVD and label-aware PLS-SVD/rSVD scores",
                "PCA scores computed by rSVD and label-aware PLS-SVD scores "
                "computed by rSVD",
            ),
            ("nested validation", "nested cross-validation"),
        ],
    )
    replace_paragraph(
        document,
        "Table 1 and Figure 2 show",
        [
            (
                "best training-validation component value",
                "training-set-selected component count",
            )
        ],
    )
    replace_paragraph(
        document,
        "A formal estimator-preservation study",
        [
            (
                "fixed five-fold component-selection runs",
                "fixed 5-fold cross-validation component-selection runs",
            ),
            ("fixed-fold tasks", "fixed 5-fold cross-validation tasks"),
        ],
    )
    replace_paragraph(
        document,
        "A controlled implementation ablation",
        [("SIMPLS/IRLBA", "SIMPLS using IRLBA")],
    )
    replace_paragraph(
        document,
        "Table 1. Paired CPU/CUDA",
        [
            (
                "best training-validation value",
                "training-set-selected component count",
            ),
            ("k denotes the total budget", "A denotes the total budget"),
            ("k-1 predictive", "A-1 predictive"),
        ],
    )
    replace_paragraph(
        document,
        "Figure 2. Matched CPU",
        [
            (
                "best training-validation value",
                "training-set-selected component count",
            )
        ],
    )
    replace_paragraph(
        document,
        "Figure 3. Fixed-complexity NMR",
        [
            (
                "SIMPLS-rSVD-predicted",
                "rSVD-based SIMPLS prediction of the",
            ),
            (
                "fastPLS PLS-SVD and SIMPLS CPU/CUDA rSVD implementations",
                "fastPLS PLS-SVD and SIMPLS implementations using rSVD on "
                "CPU and CUDA",
            ),
            (
                "PLS-SVD/IRLBA reference",
                "PLS-SVD reference using IRLBA",
            ),
        ],
    )
    replace_paragraph(
        document,
        "Table 2. Exploratory matched ImageNet",
        [("not nested validation", "not nested cross-validation")],
    )
    replace_paragraph(
        document,
        "Cross-validation acceleration was tested",
        [
            (
                "each of ten prespecified folds",
                "each of 10 prespecified cross-validation folds",
            ),
            ("ten-fold validation", "10-fold cross-validation"),
            ("SIMPLS/IRLBA", "SIMPLS using IRLBA"),
        ],
    )

    replace_nodes(document.tables[0], [("k=", "A=")])
    replace_cell_exact(document.tables[3].rows[0].cells[1], "k", "A")
    replace_nodes(document, [
        ("PLS-SVD/rSVD", "PLS-SVD (rSVD)"),
        ("PCA-rSVD", "PCA (rSVD)"),
    ])

    document.core_properties.title = (
        "fastPLS CMPB manuscript - final language and notation edit"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)

    numerical_heading = find_paragraph(document, "S2. Numerical algorithms")
    notation = insert_after(
        numerical_heading,
        (
            "Notation follows the main text: a indexes the current latent "
            "component, A is the retained component count, K is the number of "
            "cross-validation folds, and r is the target rank of an SVD "
            "calculation. Lowercase k is used only for retrieval or evaluation "
            "cutoffs, such as nearest-neighbour neighbourhood size and top-k "
            "accuracy, and never for PLS component indexing. The public R "
            "argument ncomp supplies A."
        ),
        style="Body Text",
    )
    notation.paragraph_format.keep_together = True

    replace_paragraph(
        document,
        "For component k,",
        [("For component k,", "For component a,")],
    )
    replace_paragraph(
        document,
        "Table S2. Dominant computational",
        [
            (
                "Here, a is the requested number of PLS components",
                "Here, A is the retained number of PLS components",
            )
        ],
    )
    replace_paragraph(
        document,
        "For an operator A and target rank k,",
        [
            (
                "For an operator A and target rank k,",
                "For a linear operator M and target rank r,",
            ),
            ("products with A and its transpose", "products with M and its transpose"),
        ],
    )
    replace_paragraph(
        document,
        "For SIMPLS, a prediction at component count k",
        [
            (
                "prediction at component count k",
                "prediction at retained component count A",
            )
        ],
    )
    replace_paragraph(
        document,
        "pls.single.cv() evaluates",
        [
            (
                "within one cross-validation layer",
                "within one K-fold cross-validation layer",
            ),
            (
                "pls.double.cv() adds an outer layer",
                "pls.double.cv() performs nested cross-validation by adding "
                "an outer assessment layer",
            ),
        ],
    )
    replace_paragraph(
        document,
        "The corresponding deterministic results",
        [
            (
                "fixed five-fold validation",
                "fixed 5-fold cross-validation",
            )
        ],
    )
    replace_paragraph(
        document,
        "Figure S14.",
        [
            (
                "SIMPLS-rSVD fits",
                "SIMPLS fits using rSVD",
            )
        ],
    )
    replace_paragraph(
        document,
        "v <- p; if (k > 1)",
        [
            ("if (k > 1)", "if (a > 1)"),
            ("V[,1:(k-1)]", "V[,1:(a-1)]"),
        ],
    )
    replace_paragraph(
        document,
        "The full study comprised",
        [
            (
                "fixed five-fold component-selection runs",
                "fixed 5-fold cross-validation component-selection runs",
            ),
            (
                "fixed-fold tasks",
                "fixed 5-fold cross-validation tasks",
            ),
        ],
    )
    replace_paragraph(
        document,
        "Table S11. Fixed five-fold",
        [
            (
                "Fixed five-fold component selection",
                "Fixed 5-fold cross-validation component selection",
            )
        ],
    )
    replace_paragraph(
        document,
        "One maximal SIMPLS-rSVD model",
        [
            (
                "One maximal SIMPLS-rSVD model",
                "One maximal SIMPLS model using rSVD",
            )
        ],
    )
    replace_paragraph(
        document,
        "The archived benchmark compared",
        [
            ("ten-fold validation", "10-fold cross-validation"),
        ],
    )
    replace_paragraph(
        document,
        "Table S20. Optimized 10-fold validation",
        [
            (
                "Optimized 10-fold validation",
                "Optimized 10-fold cross-validation",
            )
        ],
    )
    replace_paragraph(
        document,
        "Kernel and component settings",
        [("five-fold validation", "5-fold cross-validation")],
    )
    replace_paragraph(
        document,
        "Table S26. Corrected CPU SIMPLS-rSVD",
        [
            (
                "CPU SIMPLS-rSVD validation",
                "CPU SIMPLS with rSVD validation",
            )
        ],
    )
    replace_paragraph(
        document,
        "Table S34. Controlled deterministic",
        [("n/p/q/k", "n/p/q/A")],
    )
    replace_paragraph(
        document,
        "for (k in seq_len(ncomp))",
        [("for (k in seq_len(ncomp))", "for (a in seq_len(ncomp))")],
    )
    replace_paragraph(
        document,
        "V[,k] <- v;",
        [("V[,k]", "V[,a]")],
    )
    replace_paragraph(
        document,
        "Ten fold groups were generated",
        [
            (
                "Ten fold groups were generated once",
                "For 10-fold cross-validation, fold-assignment groups were "
                "generated once",
            )
        ],
    )
    replace_paragraph(
        document,
        "Table S33. Matched ten-fold",
        [
            (
                "Matched ten-fold SIMPLS cross-validation",
                "Matched 10-fold SIMPLS cross-validation",
            )
        ],
    )
    replace_paragraph(
        document,
        "The principal selected-point benchmark",
        [
            ("remaining k-1 slots", "remaining A-1 slots"),
        ],
    )
    replace_paragraph(
        document,
        "The ablation used deterministic CPU SIMPLS/IRLBA",
        [
            ("SIMPLS/IRLBA", "SIMPLS using IRLBA"),
            ("p x q x k", "p x q x A"),
            ("component 1 to k", "component 1 to A"),
        ],
    )
    replace_paragraph(
        document,
        "The cached X'X production condition",
        [("requires k >= 20", "requires A >= 20")],
    )
    replace_paragraph(
        document,
        "Nonlinear kernel selection",
        [("Five-fold validation", "5-fold cross-validation")],
    )
    replace_paragraph(
        document,
        "Table S35. OPLS",
        [
            ("k is the component value", "A is the component value"),
            ("OPLS k equals", "OPLS A equals"),
        ],
    )
    replace_paragraph(
        document,
        "Warnings are evaluated",
        [
            (
                "q >= 10,000 and ncomp >= 50",
                "q >= 10,000 and A >= 50 (the public argument ncomp)",
            )
        ],
    )

    # Table headers and component-count labels.
    table_headers = {
        8: [(1, "Selected k", "Selected A")],
        14: [(2, "fastPLS k", "fastPLS A"), (3, "Reference k", "Reference A")],
        17: [(2, "k", "A")],
        18: [(2, "k", "A")],
        19: [(2, "k", "A")],
        20: [(2, "k", "A")],
        22: [(1, "k", "A")],
        25: [(3, "k", "A")],
        27: [(2, "k", "A")],
        28: [(3, "k", "A")],
        31: [(1, "k", "A")],
        33: [(2, "k", "A")],
        35: [(3, "k", "A")],
        37: [(2, "k", "A")],
        38: [(1, "n/p/q/k", "n/p/q/A")],
        39: [(2, "k shown", "A shown"), (3, "Predictive k", "Predictive A")],
    }
    for table_index, replacements in table_headers.items():
        for cell_index, expected, replacement in replacements:
            replace_cell_exact(
                document.tables[table_index].rows[0].cells[cell_index],
                expected,
                replacement,
            )

    replace_nodes(document.tables[12], [
        ("components 1,...,k", "components 1,...,A"),
        ("B_k=R_k Q_k'", "B_a=R_a Q_a'"),
        ("B_k=B_(k-1)+r_k q_k'", "B_a=B_(a-1)+r_a q_a'"),
        ("Yhat_k=Yhat_(k-1)+t_k q_k'", "Yhat_a=Yhat_(a-1)+t_a q_a'"),
        ("Yhat_k=T_k Q_k'", "Yhat_a=T_a Q_a'"),
    ])
    replace_nodes(document.tables[25], [
        ("simpls", "SIMPLS"),
        ("opls", "OPLS"),
    ])
    replace_nodes(document.tables[32], [("k=", "A=")])
    replace_nodes(document.tables[40], [
        ("ncomp >= 50", "A >= 50"),
    ])
    replace_nodes(document.tables[36], [
        ("no nested validation", "no nested cross-validation"),
    ])
    replace_nodes(document, [
        ("PLS-SVD/rSVD", "PLS-SVD (rSVD)"),
        ("PCA-rSVD", "PCA (rSVD)"),
        ("PLS-SVD/IRLBA", "PLS-SVD (IRLBA)"),
    ])

    document.core_properties.title = (
        "fastPLS CMPB supplement - final language and notation edit"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)

    replace_paragraph(
        document,
        "We agree that the previous table",
        [
            (
                "stochastic rSVD",
                "randomized SVD (rSVD)",
            )
        ],
    )
    replace_paragraph(
        document,
        "Response: Corrected. We removed the fastest-row",
        [
            ("PLS-SVD k=5", "PLS-SVD A=5"),
            ("SIMPLS k=50", "SIMPLS A=50"),
        ],
    )
    replace_paragraph(
        document,
        "Response: Corrected. Main OPLS",
        [
            ("displayed k comprises k-1", "displayed A comprises A-1"),
        ],
    )
    replace_paragraph(
        document,
        "Response: Corrected. We no longer characterize float32",
        [
            (
                "q >= 10,000 with ncomp >= 50",
                "q >= 10,000 with A >= 50 (the public argument ncomp)",
            )
        ],
    )
    replace_paragraph(
        document,
        "Response: Agreed. The previous ratio compared",
        [
            ("ten-fold validation", "10-fold cross-validation"),
            ("SIMPLS/IRLBA", "SIMPLS using IRLBA"),
        ],
    )
    replace_paragraph(
        document,
        "Response: Agreed. We added a controlled ablation",
        [("SIMPLS/IRLBA", "SIMPLS using IRLBA")],
    )

    heading = document.add_heading(
        "36. Language, notation, and cross-validation terminology required a "
        "final consistency edit",
        level=1,
    )
    heading.paragraph_format.keep_with_next = True
    comment = document.add_paragraph(
        "Reviewer comment: A final language and notation edit is needed for "
        "consistent use of PLS-SVD, rSVD, component index symbols, and "
        "cross-validation terminology."
    )
    comment.paragraph_format.keep_with_next = True
    response = document.add_paragraph(
        "Response: Corrected throughout the manuscript, supplement, tables, "
        "captions, and response letter. PLS-SVD and randomized SVD (rSVD) are "
        "now used consistently. The current latent component is indexed by a, "
        "the retained component count by A, and the number of "
        "cross-validation folds by K; lowercase k is reserved for retrieval "
        "or evaluation cutoffs such as kNN and top-k accuracy and is never used "
        "for PLS component indexing. The public R argument ncomp is "
        "identified explicitly as the input supplying A. We also standardized "
        "the terms cross-validation, nested cross-validation, 5-fold "
        "cross-validation, and 10-fold cross-validation, and removed the former "
        "ambiguous wording for training-set selection and nested assessment."
    )
    response.paragraph_format.keep_together = True

    document.core_properties.title = (
        "fastPLS CMPB response - final language and notation edit"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    revise_response()


if __name__ == "__main__":
    main()
