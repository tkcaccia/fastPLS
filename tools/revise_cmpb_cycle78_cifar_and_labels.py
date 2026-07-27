#!/usr/bin/env python3
"""Synchronize the corrected CIFAR-100 result and readable figure labels."""

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle77"
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260727_cycle78"
MAIN_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_main_cycle77_0.99.6_20260726.docx"
)
SUPP_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_supplement_cycle77_0.99.6_20260726.docx"
)
MAIN_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_main_cycle78_0.99.6_20260727.docx"
)
SUPP_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_supplement_cycle78_0.99.6_20260727.docx"
)
FIGURE_2 = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle57_20260726"
    / "external_single_cpu_accuracy_time_memory.png"
)
FIGURE_3 = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle62_20260726"
    / "accelerator_concordance_speedups.png"
)


def replace_paragraph_prefix(document, prefix, replacement):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            paragraph.text = replacement
            return
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_media(source, destination, replacements):
    temporary = destination.with_suffix(".media.docx")
    with ZipFile(source, "r") as zin, ZipFile(
        temporary, "w", compression=ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = replacements.get(item.filename, zin.read(item.filename))
            zout.writestr(item, data)
    temporary.replace(destination)


def set_cell_text(cell, value, size=7.2):
    cell.text = value
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(size)


def find_table(document, header_values):
    for table in document.tables:
        headers = [cell.text for cell in table.rows[0].cells] if table.rows else []
        if headers == header_values:
            return table
    raise RuntimeError(f"Table not found: {header_values}")


def update_external_table(document):
    table = find_table(
        document,
        [
            "Dataset",
            "A",
            "fastPLS argmax",
            "fastPLS LDA",
            "Fastest external",
            "Best-accuracy external",
        ],
    )
    for row in table.rows[1:]:
        if row.cells[0].text == "CIFAR-100":
            values = [cell.text for cell in row.cells]
            values[3] = "0.8710; 10.118s; 1687MB"
            for cell, value in zip(row.cells, values):
                set_cell_text(cell, value)
            return
    raise RuntimeError("CIFAR-100 row not found in Table S10")


def revise_main():
    document = Document(MAIN_SOURCE)
    replace_paragraph_prefix(
        document,
        "The float64 single-CPU comparison attempted 126 external-package runs:",
        (
            "The float64 single-CPU comparison attempted 126 external-package runs: "
            "110 completed, 12 were package limitations, two timed out, and two "
            "errored. In the estimator-matched argmax subset, fastPLS and "
            "pls::simpls.fit had identical accuracy on nine datasets; fastPLS was "
            "faster on seven, including 4.23-fold on CIFAR-100, 8.65-fold on Retina, "
            "and 8.90-fold on Tabula Muris. Matched accuracies were 0.8739 "
            "(8,739/10,000; Wilson 95% CI 0.8672-0.8803), 0.9678 "
            "(21,684/22,406; 0.9654-0.9700), and 0.8006 "
            "(40,077/50,059; 0.7971-0.8041), respectively. The separate fastPLS "
            "SIMPLS-LDA workflow reached CIFAR-100 accuracy 0.8710 in 10.118 s; "
            "because its prediction head differs, it is reported as a workflow "
            "comparison rather than estimator-matched evidence (Figure 2; "
            "Supplementary Table S10)."
        ),
    )
    replace_paragraph_prefix(
        document,
        "Figure 3. Supporting backend and solver workflow comparisons.",
        (
            "Figure 3. Supporting backend and solver workflow comparisons. Colored "
            "CPU/accelerator cells report speed-up only when the absolute predictive-"
            "metric difference was at most 0.005 and paired-prediction agreement was "
            "at least 0.995. Gray cells state whether the metric or predictions were "
            "discordant, or whether paired predictions were not retained. Panel C "
            "shows approximate rSVD speed relative to deterministic IRLBA; numerical "
            "audit status is reported separately in Supplementary Tables S8 and S11."
        ),
    )
    document.save(MAIN_OUTPUT)
    replace_media(
        MAIN_OUTPUT,
        MAIN_OUTPUT,
        {
            "word/media/image20.png": FIGURE_2.read_bytes(),
            "word/media/image24.png": FIGURE_3.read_bytes(),
        },
    )


def revise_supplement():
    document = Document(SUPP_SOURCE)
    update_external_table(document)
    replace_paragraph_prefix(
        document,
        "The primary software comparison used float64, deterministic CPU SIMPLS,",
        (
            "The primary software comparison used float64, deterministic CPU "
            "SIMPLS, the same split and component count, and one effective BLAS "
            "thread. fastPLS argmax is estimator matched to pls::simpls.fit; LDA is "
            "a workflow comparison because the prediction head differs. Memory is "
            "absolute process RSS and is reported for feasibility, not isolated "
            "algorithmic allocation. Across 126 attempted external-package "
            "dataset/method runs, 110 completed and 16 did not: 12 were documented "
            "package limitations, two were killed at the timeout, and two produced "
            "execution errors. The previously incomplete CIFAR-100 fastPLS "
            "SIMPLS-LDA row was rerun independently with a 7,200-s limit; all three "
            "replicates completed, with median accuracy 0.8710, median total time "
            "10.118 s, and median peak process RSS 1,687 MB. The isolated rerun is "
            "stored under benchmark_results/manuscript_revision_cycle78_20260726/"
            "cifar100_fastpls_simpls_lda."
        ),
    )
    document.save(SUPP_OUTPUT)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)


if __name__ == "__main__":
    main()
