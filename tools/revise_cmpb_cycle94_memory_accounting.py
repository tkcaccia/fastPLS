from pathlib import Path
import sys

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle93"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle94"
RESULTS = ROOT / "benchmark_results" / "external_simpls_memory_publication_20260825"
OUTPUT.mkdir(parents=True, exist_ok=True)

sys.path.insert(
    0,
    "/Users/stefano/.codex/plugins/cache/openai-primary-runtime/documents/26.819.11345/skills/documents/scripts",
)
from table_geometry import apply_table_geometry, column_widths_from_weights, section_content_width_dxa


LABELS = {
    "ccle": "CCLE",
    "cifar100": "CIFAR-100",
    "gtex_v8": "GTEx v8",
    "metref": "MetRef",
    "retina": "Retina",
    "tabula": "Tabula Muris",
    "tcga_brca": "TCGA-BRCA",
    "tcga_hnsc_methylation": "TCGA-HNSC methyl.",
    "tcga_pan_cancer": "TCGA Pan-Cancer",
}


def replace_paragraph(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def find_paragraph(document, phrase):
    matches = [p for p in document.paragraphs if phrase in p.text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph containing {phrase!r}; found {len(matches)}")
    return matches[0]


def replace_image_in_paragraph(paragraph, image_path):
    blips = paragraph._p.xpath(".//a:blip")
    if len(blips) != 1:
        raise RuntimeError(f"Expected one image in paragraph; found {len(blips)}")
    rel_id = blips[0].get(qn("r:embed"))
    image_part = paragraph.part.related_parts[rel_id]
    image_part._blob = Path(image_path).read_bytes()


def format_median_iqr(median, iqr):
    if median >= 1000 or iqr >= 1000:
        return f"{median:,.1f} ({iqr:,.1f})"
    if median >= 100:
        return f"{median:.1f} ({iqr:.1f})"
    return f"{median:.2f} ({iqr:.2f})"


def format_pair(left, right, formatter):
    return f"{formatter(left)} / {formatter(right)}"


def add_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def style_table(table, size=5.25):
    table.style = "Table"
    add_repeat_header(table.rows[0])
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(size)
                    run.bold = row_idx == 0


def set_table_rows(table, rows):
    if len(table.rows) != len(rows) + 1:
        raise RuntimeError(f"Expected {len(rows) + 1} rows; found {len(table.rows)}")
    for table_row, values in zip(table.rows[1:], rows):
        for cell, value in zip(table_row.cells, values):
            cell.text = str(value)
    style_table(table, size=4.5)


def insert_elements_before(anchor, elements):
    for element in elements:
        anchor._p.addprevious(element)


def build_memory_table(document, profile, memory):
    profile_data = memory[memory["comparison_profile"] == profile]
    rows = []
    for dataset in LABELS:
        d = profile_data[profile_data["dataset"] == dataset].set_index("implementation")
        fast = d.loc["fastpls"]
        ref = d.loc["pls"]
        rows.append(
            [
                LABELS[dataset],
                format_pair(
                    (fast.median_process_peak_rss_mb, fast.iqr_process_peak_rss_mb),
                    (ref.median_process_peak_rss_mb, ref.iqr_process_peak_rss_mb),
                    lambda x: format_median_iqr(*x),
                ),
                format_pair(
                    fast.median_prefit_process_rss_mb,
                    ref.median_prefit_process_rss_mb,
                    lambda x: f"{x:,.1f}",
                ),
                format_pair(
                    (
                        fast.median_baseline_corrected_peak_increment_mb,
                        fast.iqr_baseline_corrected_peak_increment_mb,
                    ),
                    (
                        ref.median_baseline_corrected_peak_increment_mb,
                        ref.iqr_baseline_corrected_peak_increment_mb,
                    ),
                    lambda x: format_median_iqr(*x),
                ),
                format_pair(
                    fast.theoretical_largest_retained_mb,
                    ref.theoretical_largest_retained_mb,
                    lambda x: f"{x:,.3f}" if x < 100 else f"{x:,.1f}",
                ),
            ]
        )

    table = document.add_table(rows=1, cols=5)
    headers = [
        "Dataset",
        "Absolute peak F / P",
        "Pre-fit baseline F / P",
        "Peak increment F / P",
        "Largest retained dense object F / P",
    ]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
    for values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, values):
            cell.text = value
    style_table(table)
    width = section_content_width_dxa(document.sections[0])
    widths = column_widths_from_weights([1.05, 1.45, 1.25, 1.45, 1.45], width)
    apply_table_geometry(table, widths, table_width_dxa=width)
    return table


summary = pd.read_csv(RESULTS / "external_simpls_timing_summary.csv")
pairs = pd.read_csv(RESULTS / "external_simpls_timing_pairs.csv")
memory = pd.read_csv(RESULTS / "external_simpls_memory_summary.csv")


main_source = SOURCE / "fastPLS_CMPB_main_cycle93_0.99.25_20260825.docx"
main_output = OUTPUT / "fastPLS_CMPB_main_cycle94_0.99.25_20260825.docx"
main = Document(main_source)

replace_paragraph(
    find_paragraph(main, "Isolated-process baseline and peak host RSS"),
    "Twelve tasks covered metabolomics, NMR, CITE-seq, tissue and cancer omics, single-cell "
    "transcriptomics, drug response, and CIFAR-100 [7,20-27,30]. Methods used identical stored splits "
    "and training-selected component grids. CIFAR-100 followed its documented 50,000/10,000 split [30]. "
    "Classification used accuracy and Wilson intervals; multivariate regression used RMSD, Q2, and held-out "
    "bootstrap intervals. Runtime included fitting and prediction. Absolute host RSS was the isolated R "
    "process high-water mark. A pre-fit baseline was recorded after package/data preparation and garbage "
    "collection; the baseline-corrected process increment was peak minus baseline. This increment still "
    "contains returned outputs, temporary BLAS allocations, allocator effects, and benchmark/runtime state, "
    "and is not labelled algorithmic workspace. Formula-based dense-object sizes were reported separately. "
    "Process-specific GPU increments similarly include runtime context rather than workspace alone."
)

replace_paragraph(
    find_paragraph(main, "The strict comparison completed all 108 planned runs"),
    "The strict comparison completed all 108 planned runs: nine datasets, two output profiles, two "
    "implementations, and three fresh-process repetitions. Accuracy was identical for every pair. With "
    "minimum common prediction outputs, fastPLS was faster on four datasets and pls::simpls.fit on five; "
    "the largest fastPLS advantage was 1.53-fold on GTEx v8. Under ordinary public workflows, fastPLS was "
    "faster on five datasets, including 2.41-fold on CIFAR-100, 3.35-fold on Retina, and 4.84-fold on Tabula "
    "Muris. Corresponding accuracies were 0.8739 (8,739/10,000; Wilson 95% CI 0.8672-0.8803), 0.9678 "
    "(21,684/22,406; 0.9654-0.9700), and 0.8006 (40,077/50,059; 0.7971-0.8041). These timing profiles answer "
    "different questions and are not pooled (Figure 2; Supplementary Tables S10a-S10b)."
)

replace_paragraph(
    find_paragraph(main, "Workflow gains partly reflected output policy"),
    "Workflow gains partly reflected output policy. On CIFAR-100, the compact fastPLS fit object was 1.38 MB "
    "versus 7,778 MB for ordinary pls::simpls.fit. Median complete-process peak RSS was 2.20 versus 13.42 GB; "
    "from nearly identical 1.91-GB pre-fit baselines, the corrected peak increments were 0.294 versus 11.51 GB. "
    "The theoretical largest retained dense object was the 0.586-MB final coefficient matrix for fastPLS and "
    "a 3,814.7-MB fitted or residual response path for pls::simpls.fit. These values describe the specified "
    "complete workflow, not isolated algorithmic workspace. When both methods retained complete coefficient "
    "paths, fit objects were 59.30 and 58.60 MB, corrected increments were 69.5 and 127.0 MB, and the speed-up "
    "was 1.18-fold. The broader package panel remains a workflow comparison with implementation-specific "
    "outputs (Supplementary Tables S10c-S10e)."
)

replace_image_in_paragraph(
    main.paragraphs[67], RESULTS / "external_simpls_timing_profiles.png"
)
main.save(main_output)


supp_source = SOURCE / "fastPLS_CMPB_supplement_cycle93_0.99.25_20260825.docx"
supp_output = OUTPUT / "fastPLS_CMPB_supplement_cycle94_0.99.25_20260825.docx"
supp = Document(supp_source)

measurement = find_paragraph(supp, "Each fit is run in an isolated R process")
replace_paragraph(
    measurement,
    measurement.text
    + " For the repeated external SIMPLS comparison, /usr/bin/time supplied the complete-worker high-water "
    "RSS and the worker sampled its current RSS after loading and garbage collection, immediately before the "
    "fit. Their difference is reported as a baseline-corrected process peak increment. Because the absolute "
    "high-water mark spans the worker lifetime, the increment can retain earlier loading high-water effects; "
    "it is therefore a conservative complete-process measure rather than fit-window workspace. Formula-based "
    "sizes of the cross-covariance, coefficient, fitted/residual, score, and prediction objects are reported "
    "alongside measured RSS."
)

external = find_paragraph(supp, "The definitive repeated comparison used float64")
replace_paragraph(
    external,
    "The definitive repeated comparison used float64, deterministic CPU SIMPLS, identical splits and "
    "component counts, one effective BLAS thread, and a 10,000-s timeout. Package and data loading occurred "
    "before timing; no numerical warm-up was performed. Every method-dataset pair ran in three fresh R "
    "processes. Two profiles were kept separate. In the minimum-output profile, fastPLS retained the full "
    "coefficient path, means, and compact internal prediction/audit factors, while pls::simpls.fit(stripped = "
    "TRUE) retained the full coefficient path and means; neither returned score, loading, fitted-value, "
    "residual, or variance arrays. In the public-workflow profile, fastPLS retained its ordinary compact object "
    "and variance summary, whereas pls::simpls.fit retained coefficients, X/Y scores and loadings, projection, "
    "fitted values, residuals, and X-variance quantities. Absolute lifetime peak RSS, the immediately pre-fit "
    "process baseline, and their difference were retained separately. The difference is a baseline-corrected "
    "complete-process increment, not isolated algorithmic allocation. The theoretical sizes of the largest "
    "retained dense objects were calculated from matrix dimensions and float64 element width. All 108 runs "
    "completed and all paired accuracies were identical. The older 126-run multi-package panel is retained as "
    "a workflow comparison only: 110 completed, 12 were package limitations, two timed out, and two errored. "
    "It is not labelled an estimator-kernel comparison because outputs and model families differ."
)

# Synchronize the repeated timing table with the memory-audit rerun.
timing_table = next(
    table for table in supp.tables
    if table.rows[0].cells[0].text == "Dataset" and table.rows[0].cells[1].text == "Profile"
)
timing_rows = []
for dataset in LABELS:
    for profile, profile_label in [
        ("complete_workflow", "Public workflow"),
        ("estimator_kernel", "Minimum outputs"),
    ]:
        row = pairs[(pairs.dataset == dataset) & (pairs.comparison_profile == profile)].iloc[0]
        timing_rows.append([
            LABELS[dataset],
            profile_label,
            format_median_iqr(row.median_total_sec_fastpls, row.iqr_total_sec_fastpls),
            format_median_iqr(row.median_total_sec_pls, row.iqr_total_sec_pls),
            f"{row.speedup_pls_over_fastpls:.2f}",
            f"{row.median_accuracy_fastpls:.4f} / {row.median_accuracy_pls:.4f}",
            f"{int(row.repetitions_completed_fastpls)} / {int(row.repetitions_completed_pls)}",
        ])
set_table_rows(timing_table, timing_rows)

old_broad_caption = find_paragraph(supp, "Table S10c. Broad external-package workflow comparison")
replace_paragraph(
    old_broad_caption,
    "Table S10e. Broad external-package workflow comparison. Cells report accuracy; total "
    "fitting-plus-prediction time; and absolute peak process RSS. Implementations retain package-specific "
    "outputs, and unsupported runs, timeouts, and errors remain explicit."
)

authority_table = next(
    table for table in supp.tables
    if table.rows[0].cells[0].text == "Main-text claim"
)
for row in authority_table.rows[1:]:
    if row.cells[1].text == "Tables S10a-S10c":
        row.cells[1].text = "Tables S10a-S10e"

provenance_table = next(
    table for table in supp.tables
    if table.rows[0].cells[0].text == "ID" and table.rows[0].cells[1].text == "Authoritative output"
)
for row in provenance_table.rows[1:]:
    if row.cells[0].text == "A22":
        row.cells[1].text = "Tables S10a-S10d; Figure 2"
        row.cells[2].text = "benchmark_results/external_simpls_memory_publication_20260825"
        row.cells[4].text = (
            "ba80b65f0c66; archive SHA-256 74e134ef22d5; "
            "worker SHA-256 88d9083a2328"
        )
        row.cells[5].text = (
            "scripts/run_external_simpls_timing.sh; "
            "benchmark/external_simpls_timing/worker.R"
        )
style_table(provenance_table, size=4.5)

minimum_table = build_memory_table(supp, "estimator_kernel", memory)
public_table = build_memory_table(supp, "complete_workflow", memory)
caption_min = supp.add_paragraph(
    "Table S10c. Host-memory accounting for the minimum-output profile. Values are MB; F/P denotes "
    "fastPLS / pls::simpls.fit. Absolute peak and corrected increment are median (IQR) across three fresh "
    "processes; pre-fit baseline and theoretical storage are medians or dimension-derived values. Both "
    "methods retain a p x q x A coefficient path in this profile."
)
caption_pub = supp.add_paragraph(
    "Table S10d. Host-memory accounting for ordinary public workflows. Values are MB; F/P denotes fastPLS / "
    "pls::simpls.fit. Absolute peak and corrected increment are median (IQR) across three fresh processes. "
    "The largest retained dense object is the p x q final coefficient matrix for fastPLS and one n_train x q "
    "x A fitted or residual path for pls::simpls.fit; the reference object ordinarily retains both fitted and "
    "residual paths. Measured RSS includes runtime, allocator, temporary, and output storage."
)
caption_pub.paragraph_format.page_break_before = True

insert_elements_before(
    old_broad_caption,
    [caption_min._p, minimum_table._tbl, caption_pub._p, public_table._tbl],
)

supp.save(supp_output)

print(main_output)
print(supp_output)
