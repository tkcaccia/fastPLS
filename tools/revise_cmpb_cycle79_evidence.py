#!/usr/bin/env python3
"""Strengthen the CMPB evidence hierarchy and separate exploratory analyses."""

from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260727_cycle78"
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260727_cycle79"
MAIN_SOURCE = SOURCE_DIR / "fastPLS_CMPB_main_cycle78_0.99.6_20260727.docx"
SUPP_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_supplement_cycle78_0.99.6_20260727.docx"
)
MAIN_OUTPUT = OUTPUT_DIR / "fastPLS_CMPB_main_cycle79_0.99.6_20260727.docx"
SUPP_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_supplement_cycle79_0.99.6_20260727.docx"
)
FIGURE_3 = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle62_20260726"
    / "accelerator_concordance_speedups.png"
)
FIGURE_1 = ROOT / "artifacts" / "figures" / "fastpls_architecture_current.png"
RSVD_SUPP_FIGURE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle62_20260726"
    / "rsvd_workflow_speed_supp.png"
)


def replace_paragraph_prefix(document, prefix, replacement):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            paragraph.text = replacement
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def paragraph_by_prefix(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_media(source, destination, replacements):
    temporary = destination.with_suffix(".media.docx")
    with ZipFile(source, "r") as zin, ZipFile(
        temporary, "w", compression=ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = replacements.get(item.filename, zin.read(item.filename))
            zout.writestr(item, data)
    temporary.replace(destination)


def set_cell_text(cell, value, size=7.1, bold=False):
    cell.text = value
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.bold = bold


def find_table(document, header_values):
    for table in document.tables:
        headers = [cell.text for cell in table.rows[0].cells] if table.rows else []
        if headers == header_values:
            return table
    raise RuntimeError(f"Table not found: {header_values}")


def append_table_row(table, values, size=7.1):
    row = table.add_row()
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value, size=size)
    return row


def insert_paragraph_before(document, target, text, style=None):
    paragraph = document.add_paragraph(text, style=style)
    target._p.addprevious(paragraph._p)
    return paragraph


def insert_table_before(document, target, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table"
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_text(cell, value, size=7.0, bold=True)
    for values in rows:
        append_table_row(table, values, size=6.9)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Inches(width)
    target._p.addprevious(table._tbl)
    return table


def insert_picture_before(document, target, path, width):
    paragraph = document.add_paragraph()
    paragraph.alignment = 1
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    target._p.addprevious(paragraph._p)
    return paragraph


def replace_math_symbol(paragraph, old, new):
    replacements = 0
    for element in paragraph._p.iter():
        if element.tag in {qn("m:t"), qn("w:t")} and element.text == old:
            element.text = new
            replacements += 1
    return replacements


def resize_embedded_image(document, target_name, width_inches, aspect_ratio):
    for shape in document.inline_shapes:
        blip = shape._inline.graphic.graphicData.pic.blipFill.blip
        relation = document.part.rels.get(blip.embed)
        if relation and relation.target_ref.endswith(target_name):
            shape.width = Inches(width_inches)
            shape.height = Inches(width_inches / aspect_ratio)
            return
    raise RuntimeError(f"Embedded image not found: {target_name}")


def revise_main():
    document = Document(MAIN_SOURCE)
    replace_paragraph_prefix(
        document,
        "Results: fastPLS SIMPLS met the prespecified numerical tolerances",
        (
            "Results: Deterministic fastPLS SIMPLS met the prespecified numerical "
            "tolerances in all 117 component-level comparisons. In matched "
            "single-CPU comparisons, it was faster than pls::simpls.fit on seven "
            "of nine datasets, with identical argmax accuracy and speed-up up to "
            "8.90-fold. Same-code ablations preserved predictions and showed that "
            "memory and runtime gains were optimization- and matrix-shape "
            "dependent. NMR and ImageNet were retained as exploratory large-scale "
            "feasibility studies rather than estimator-equivalence evidence."
        ),
    )
    replace_paragraph_prefix(
        document,
        "Deterministic fastPLS SIMPLS met the prespecified numerical tolerances",
        (
            "Deterministic fastPLS SIMPLS met the prespecified numerical "
            "tolerances in all 117 component-level comparisons with de Jong "
            "SIMPLS. Approximate rSVD was evaluated separately, and only fully "
            "audited settings support numerical-agreement claims; faster settings "
            "are labelled workflow-only. OPLS and kernel PLS met the prespecified "
            "tolerances in a separate deterministic reliability study. Same-code "
            "execution ablations and direct PLS-SVD/SIMPLS shape comparisons are "
            "reported in Supplementary Tables S7a-S7b."
        ),
    )
    replace_paragraph_prefix(
        document,
        "The float64 single-CPU comparison attempted 126 external-package runs:",
        (
            "The float64 single-CPU comparison attempted 126 external-package "
            "runs: 110 completed, 12 were package limitations, two timed out, and "
            "two errored. In the estimator-matched argmax subset, fastPLS and "
            "pls::simpls.fit [19] had identical accuracy on nine datasets; fastPLS "
            "was faster on seven, including 4.23-fold on CIFAR-100, 8.65-fold on "
            "Retina, and 8.90-fold on Tabula Muris. Matched accuracies were 0.8739 "
            "(8,739/10,000; Wilson 95% CI 0.8672-0.8803), 0.9678 "
            "(21,684/22,406; 0.9654-0.9700), and 0.8006 "
            "(40,077/50,059; 0.7971-0.8041), respectively. The separate fastPLS "
            "SIMPLS-LDA workflow reached CIFAR-100 accuracy 0.8710 in 10.118 s; "
            "because its prediction head differs, it is a workflow comparison "
            "rather than estimator-matched evidence (Figure 2; Supplementary "
            "Table S10)."
        ),
    )
    replace_paragraph_prefix(
        document,
        "Figure 2. Matched float64 single-CPU SIMPLS workflows.",
        (
            "Figure 2. float64 single-CPU SIMPLS classification workflows. "
            "Argmax rows compare matched SIMPLS estimators; LDA rows compare "
            "complete workflows with a different prediction head. Panels report "
            "accuracy, fitting-plus-prediction time, and absolute process RSS. NE "
            "denotes unavailable or incomplete runs."
        ),
    )
    replace_paragraph_prefix(
        document,
        "Hardware acceleration was a supporting workflow analysis.",
        (
            "Same-code ablations on CIFAR-100, MetRef, PRISM, and Retina produced "
            "zero endpoint-metric differences and identical classifications. "
            "Compact prediction reduced incremental RSS by up to 77.7% and time "
            "by up to 1.24-fold; implicit cross-covariance products reduced RSS by "
            "up to 70.6% but were faster only in the high-response PRISM regime. "
            "This shape dependence motivates adaptive internal routing rather than "
            "universal activation (Supplementary Table S7a)."
        ),
    )
    replace_paragraph_prefix(
        document,
        "Approximate rSVD reduced runtime in several large workflows,",
        (
            "A direct matched-shape timing study further separated estimator "
            "choice from implementation cost. On one CPU, the SIMPLS/PLS-SVD time "
            "ratio ranged from 1.00 to 3.84; on CUDA it ranged from 0.92 to 0.98 "
            "across five synthetic matrix regimes. Thus, the accelerated SIMPLS "
            "path approached one-shot PLS-SVD runtime on the tested CUDA shapes, "
            "without implying that the two PLS estimators are statistically "
            "identical (Supplementary Table S7b)."
        ),
    )
    replace_paragraph_prefix(
        document,
        "float32 approximately halved stored inputs",
        (
            "Hardware acceleration remained route and shape dependent. CPU, CUDA, "
            "and Metal speed-up was summarized only for paired predictions meeting "
            "the stated concordance criteria (Figure 3; Supplementary Table S11). "
            "The exploratory one-power rSVD setting met only 101/117 audit checks "
            "and its speed results are quarantined in Supplementary Figure S1; "
            "deterministic IRLBA remains the reference. float32 approximately "
            "halved stored inputs on MetRef and PRISM but did not uniformly improve "
            "runtime, incremental memory, or agreement (Supplementary Table S9)."
        ),
    )
    replace_paragraph_prefix(
        document,
        "Figure 3. Supporting backend and solver workflow comparisons.",
        (
            "Figure 3. Numerically concordant accelerator workflows. Values are "
            "CPU/accelerator total-time ratios and are colored only when the "
            "absolute predictive-metric difference was at most 0.005 and paired-"
            "prediction agreement was at least 0.995; ratios above one favor the "
            "accelerator. Gray cells identify metric or prediction discordance, or "
            "missing paired predictions. Full paired values are in Supplementary "
            "Table S11."
        ),
    )
    replace_paragraph_prefix(
        document,
        "fastPLS makes sequential SIMPLS and its validation feasible",
        (
            "fastPLS reduces avoidable computation and storage along the sequential "
            "SIMPLS component path while preserving the deterministic de Jong "
            "estimator within the prespecified numerical tolerances. OPLS, kernel "
            "PLS, approximate low-rank solvers, and accelerator backends extend "
            "this core contribution under route-specific validation; deterministic "
            "float64 CPU SIMPLS remains the confirmatory reference."
        ),
    )
    replace_paragraph_prefix(
        document,
        "Code and benchmark outputs are available at",
        (
            "Code and benchmark outputs are available at "
            "https://github.com/tkcaccia/fastPLS; reusable components are at "
            "https://github.com/tkcaccia/kodama-cpp. The reviewed software snapshot "
            "is fastPLS 0.99.6 at commit "
            "6e50bd318f20289101f6b723953830aefa8b95d6. Analysis-specific source "
            "status, scripts, and archive digests are reported in Supplementary "
            "Table S15."
        ),
    )
    document.save(MAIN_OUTPUT)
    replace_media(
        MAIN_OUTPUT,
        MAIN_OUTPUT,
        {
            "word/media/rId23.png": FIGURE_1.read_bytes(),
            "word/media/image24.png": FIGURE_3.read_bytes(),
        },
    )
    document = Document(MAIN_OUTPUT)
    resize_embedded_image(document, "rId23.png", 6.55, 1400 / 620)
    resize_embedded_image(document, "image24.png", 6.8, 10.5 / 5.4)
    document.save(MAIN_OUTPUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)
    replace_paragraph_prefix(
        document,
        "Route-level evidence is consolidated here",
        (
            "Route-level evidence is consolidated here rather than repeated in the "
            "main text. Table S1 defines stage residency; Tables S7-S7b separate "
            "deterministic estimator validation, same-code execution ablation, and "
            "the direct PLS-SVD/SIMPLS shape comparison; Table S8 and Figure S1 "
            "give rSVD controls, qualification, and quarantined workflow speed; "
            "Table S9 is the authoritative float32 capability matrix; Table S11 "
            "reports paired CPU/CUDA/Metal performance and concordance; and Table "
            "S15 maps each analysis to its source state and archive."
        ),
    )

    algorithm = paragraph_by_prefix(
        document, "Compute the leading  singular triplets once"
    )
    if replace_math_symbol(algorithm, "K", "A") < 2:
        raise RuntimeError("Algorithm S1 component symbol was not updated twice")

    evidence_map = find_table(
        document,
        ["Main-text claim", "Authority", "Evidence scope"],
    )
    append_table_row(
        evidence_map,
        [
            "SIMPLS speed and storage gains arise from specific execution changes",
            "Table S7a",
            "Five same-code ablations on four matrix regimes",
        ],
    )
    append_table_row(
        evidence_map,
        [
            "SIMPLS runtime relative to one-shot PLS-SVD is shape dependent",
            "Table S7b",
            "Five matched synthetic shapes on CPU and CUDA",
        ],
    )

    target_s13 = paragraph_by_prefix(document, "S13. Approximate rSVD reliability")
    insert_paragraph_before(
        document,
        target_s13,
        "S12.1 Execution ablation and direct PLS-family timing",
        "Heading 2",
    )
    insert_paragraph_before(
        document,
        target_s13,
        (
            "Each ablation changed one internal execution feature while keeping the "
            "data, split, SIMPLS estimator, deterministic IRLBA solver, component "
            "count, seed, and prediction head fixed. Three isolated runs were used "
            "per configuration. Speed-up is reference time divided by optimized "
            "time; positive RSS reduction denotes a lower fit-window incremental "
            "peak. Results show that memory reduction and speed are separate and "
            "matrix-shape dependent."
        ),
    )
    insert_paragraph_before(
        document,
        target_s13,
        (
            "Table S7a. Same-code SIMPLS execution ablation. All endpoint metric "
            "differences were zero and minimum classification agreement was 1.000."
        ),
    )
    insert_table_before(
        document,
        target_s13,
        [
            "Optimization",
            "Datasets",
            "Time speed-up range",
            "Incremental RSS reduction",
            "Interpretation",
        ],
        [
            [
                "Cached deflation products",
                "4",
                "0.94-1.02x",
                "0.16-70.60%",
                "Memory benefit; no uniform time gain",
            ],
            [
                "Cached X-transpose-X",
                "4 (active in 1)",
                "0.99-1.00x",
                "-0.06-0.57%",
                "Shape-gated; negligible in this panel",
            ],
            [
                "Compact prediction",
                "4",
                "0.97-1.24x",
                "-0.02-77.71%",
                "Largest benefit for high-response PRISM",
            ],
            [
                "Incremental coefficients",
                "4",
                "0.99-1.07x",
                "-3.66-2.68%",
                "Small isolated effect",
            ],
            [
                "Implicit cross-covariance",
                "4",
                "0.065-6.24x",
                "0.12-70.64%",
                "Faster only in the high-response regime",
            ],
        ],
        widths=[1.25, 0.65, 1.0, 1.1, 2.2],
    )
    insert_paragraph_before(
        document,
        target_s13,
        (
            "The matched PLS-family timing study used the same synthetic matrices, "
            "requested components, rSVD controls (oversampling 10, one power "
            "iteration, seeds 101-103), split, and backend for PLS-SVD and SIMPLS. "
            "The two estimators can differ predictively; this experiment compares "
            "execution time only."
        ),
    )
    insert_paragraph_before(
        document,
        target_s13,
        (
            "Table S7b. Direct matched-shape runtime comparison. Ratios are SIMPLS "
            "time divided by PLS-SVD time; values below one favor SIMPLS."
        ),
    )
    insert_table_before(
        document,
        target_s13,
        [
            "Shape",
            "n / p / q / A",
            "CPU ratio",
            "CUDA ratio",
            "Interpretation",
        ],
        [
            ["Wide", "400 / 2000 / 20 / 10", "1.18", "0.93", "Near parity"],
            ["Tall-thin", "5000 / 50 / 20 / 10", "1.00", "0.94", "Near parity"],
            [
                "High response",
                "1000 / 300 / 500 / 50",
                "1.11",
                "0.94",
                "Near parity",
            ],
            [
                "Balanced",
                "5000 / 500 / 50 / 50",
                "3.84",
                "0.98",
                "CPU favors PLS-SVD",
            ],
            [
                "High components",
                "3000 / 768 / 200 / 100",
                "1.34",
                "0.92",
                "CUDA near parity",
            ],
        ],
        widths=[1.15, 1.45, 0.8, 0.8, 1.7],
    )

    target_s14 = paragraph_by_prefix(document, "S14. float32 capability")
    insert_paragraph_before(
        document,
        target_s14,
        (
            "Figure S1. Exploratory one-power rSVD workflow speed relative to "
            "deterministic IRLBA. The setting used oversampling 10, one power "
            "iteration, and seed 123; it met only 101/117 audit checks and is "
            "excluded from estimator-preservation claims."
        ),
    )
    insert_picture_before(document, target_s14, RSVD_SUPP_FIGURE, 6.5)

    imagenet_table = find_table(
        document,
        [
            "Experiment",
            "Head / representation",
            "A/dim.",
            "Top-1",
            "Top-5",
            "End-to-end s",
            "Host RSS MB",
            "GPU MB",
            "Qualification",
        ],
    )
    for row in imagenet_table.rows[1:]:
        if row.cells[0].text == "Classification path":
            set_cell_text(row.cells[4], "not recorded", size=6.9)

    provenance = find_table(
        document,
        [
            "ID",
            "Authoritative output",
            "Result archive",
            "Version",
            "Source status",
            "Generating script",
            "SHA-256 prefix",
        ],
    )
    for row in provenance.rows[1:]:
        if row.cells[0].text == "A08":
            set_cell_text(row.cells[1], "Table S7a")
        elif row.cells[0].text == "A12":
            set_cell_text(row.cells[1], "Table S7b")

    replace_paragraph_prefix(
        document,
        "Full component paths, ablations, per-run rows,",
        (
            "Full component paths, per-run rows, sensitivity analyses, review-cycle "
            "figures, and the former cycle-66 expanded supplement are indexed in "
            "benchmark/MANUSCRIPT_EVIDENCE_ARCHIVE.md. Their underlying CSV, RDS, "
            "PDF, PNG, log, and session-information files remain available for "
            "audit. Compact definitive ablation results are retained in Table S7a."
        ),
    )
    document.save(SUPP_OUTPUT)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)


if __name__ == "__main__":
    main()
