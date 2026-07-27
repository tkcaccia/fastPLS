#!/usr/bin/env python3
"""Move route-level qualification detail from the main text to the Supplement."""

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle75"
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle76"
MAIN_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_main_cycle75_0.99.6_20260726.docx"
)
SUPP_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_supplement_cycle75_0.99.6_20260726.docx"
)
MAIN_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_main_cycle76_0.99.6_20260726.docx"
)
SUPP_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_supplement_cycle76_0.99.6_20260726.docx"
)


MAIN_REPLACEMENTS = {
    "Methods: The implementation reuses deflation products": (
        "Methods: The implementation reuses deflation products, latent quantities, "
        "coefficients, and predictions along one maximal component path, with compact "
        "prediction and optional implicit cross-covariance products. Deterministic "
        "float64 CPU IRLBA supported estimator-matched validation against de Jong "
        "SIMPLS and independent R software. Approximate rSVD and accelerator routes "
        "were audited separately under prespecified criteria detailed in the Supplement."
    ),
    "Results: fastPLS SIMPLS met the prespecified numerical tolerances": (
        "Results: fastPLS SIMPLS met the prespecified numerical tolerances throughout "
        "deterministic validation. In matched single-CPU comparisons, it was faster "
        "than pls::simpls.fit on seven of nine datasets, with identical accuracy and "
        "speed-up up to 8.90-fold. Exploratory accelerator stress tests gave CUDA "
        "SIMPLS an NMR RMSD of 0.000759 with a 5.94-fold CUDA/CPU speed-up; CUDA "
        "SIMPLS-LDA processed 1,000,000 training and 281,167 held-out ImageNet/DINOv2 "
        "embeddings, reaching top-1 accuracy 0.8093 at 1,000 components. These stress "
        "tests demonstrate workflow feasibility rather than confirmatory equivalence."
    ),
    "Direction extraction is modular rather than the principal estimator contribution.": (
        "Direction extraction is modular rather than the principal estimator "
        "contribution. float64 CPU fitting supports deterministic IRLBA [15] and "
        "approximate rSVD [16], which uses a Gaussian range sketch, orthonormalization, "
        "power iterations, and a reduced decomposition. Exact controls, seeds, audit "
        "thresholds, and route-specific qualification are consolidated in "
        "Supplementary Sections S13-S14 and Tables S8-S9."
    ),
    "Standard R matrices use eight-byte float64 values;": (
        "Standard R matrices use eight-byte float64 values; float-package inputs use "
        "four-byte float32 values. float64 is the reference. Precision support and "
        "backend residency are route specific; the authoritative capability "
        "classifications are in Supplementary Tables S1 and S9."
    ),
    "Five-fold training-only selection was performed separately by PLS family;": (
        "Five-fold training-only selection was performed separately by PLS family; "
        "boundary selections were described as best within the grid. Accelerator "
        "speed-up was interpreted only for numerically concordant paired predictions. "
        "CUDA timing covered data transfer, execution, synchronization, and returned "
        "model or predictions; exact concordance thresholds and timing boundaries are "
        "reported in the Supplement."
    ),
    "The primary evidence tested whether deterministic IRLBA SIMPLS": (
        "The primary evidence tested whether deterministic IRLBA SIMPLS met the "
        "prespecified numerical tolerances against de Jong SIMPLS and improved runtime. "
        "rSVD was assessed separately using computational screens covering prediction "
        "error and correlation, latent-subspace agreement, decoded labels, and endpoint "
        "metrics. These screens are engineering guardrails rather than clinical-"
        "equivalence margins. Their exact definitions, tolerances, discordant counts, "
        "and audit results are reported in Supplementary Sections S12-S14 and Tables "
        "S7-S9. Deterministic IRLBA remains the confirmatory route when individual "
        "prediction changes are consequential."
    ),
    "Deterministic fastPLS SIMPLS met the prespecified numerical tolerances in 117": (
        "Deterministic fastPLS SIMPLS met the prespecified numerical tolerances in all "
        "117 component-level comparisons with de Jong SIMPLS. Approximate rSVD was "
        "evaluated separately, and only fully audited settings support numerical-"
        "agreement claims; faster settings are labelled workflow-only. OPLS and kernel "
        "PLS also met the prespecified tolerances in a separate deterministic "
        "reliability study. Full errors, subspace angles, label agreement, audit counts, "
        "and setting-level results are in Supplementary Tables S7-S8."
    ),
    "Results are organized around the accelerated SIMPLS contribution.": (
        "Results are organized around the accelerated SIMPLS contribution. Primary "
        "evidence comprises deterministic estimator comparison and the matched IRLBA "
        "comparison with independent implementations. rSVD, precision, and hardware "
        "results are reported separately with their numerical-audit status; the NMR "
        "and ImageNet large-scale analyses are exploratory feasibility studies rather "
        "than confirmatory evidence."
    ),
    "Hardware acceleration was a supporting workflow analysis.": (
        "Hardware acceleration was a supporting workflow analysis. CPU, CUDA, and Metal "
        "comparisons contributed to speed-up summaries only when paired predictions met "
        "the prespecified concordance criteria; discordant routes were quarantined. "
        "CUDA benefits depended on matrix shape, whereas the tested Metal routes mainly "
        "established portability (Figure 3; Supplementary Table S11)."
    ),
    "The one-power speed comparison used oversampling 10": (
        "Approximate rSVD reduced runtime in several large workflows, including NMR, "
        "but the faster configuration did not meet the complete numerical audit and is "
        "therefore reported only as workflow evidence. Deterministic IRLBA remains the "
        "reference; qualified rSVD controls, seeds, thresholds, and audit counts are "
        "reported in Supplementary Table S8."
    ),
    "Figure 3. Supporting backend and solver workflow comparisons.": (
        "Figure 3. Supporting backend and solver workflow comparisons. Only numerically "
        "concordant routes contribute to speed-up summaries; solver controls and audit "
        "status are reported in Supplementary Tables S8 and S11."
    ),
    "NMR comprised 1,200 training and 321 held-out spectra": (
        "NMR comprised 1,200 training and 321 held-out spectra, with 13,000 predictors "
        "and 28,355 responses. Predictor columns between 4.6 and 4.8 ppm were zeroed in "
        "training and test data; responses were unmasked. Training-only one-standard-"
        "error selection retained five PLS-SVD and 50 SIMPLS components. The selected "
        "rSVD setting did not meet the complete audit, so the NMR rSVD results are "
        "exploratory; solver controls and audit status are in Supplementary Tables "
        "S8 and S12."
    ),
    "Under that workflow-only one-power configuration": (
        "In this exploratory workflow, selected CUDA PLS-SVD and SIMPLS achieved RMSD "
        "0.001043 (95% bootstrap interval 0.001000-0.001085) and 0.000759 "
        "(0.000665-0.000884), respectively. SIMPLS also had lower median per-spectrum, "
        "response-wise, and high-intensity errors (Figure 4). These results demonstrate "
        "feasibility but do not validate rSVD estimator agreement."
    ),
    "In the matched backend benchmark, family, split, float64 precision": (
        "In the matched backend benchmark, family, split, float64 precision, component "
        "count, and solver controls were fixed. CUDA reduced PLS-SVD time from 2.301 to "
        "0.648 s and SIMPLS time from 10.525 to 1.773 s, with CPU-CUDA prediction "
        "correlations 1.000000 and 0.999981. This establishes backend agreement within "
        "the evaluated workflow, not agreement with the deterministic estimator. "
        "CUDA device increments were 590 and 3,414 MB."
    ),
    "Figure 4. Exploratory NMR rSVD workflow analyses:": (
        "Figure 4. Exploratory NMR rSVD workflow analyses: training-only component "
        "selection, held-out errors, intensity-stratified RMSD, and matched CPU/CUDA "
        "resources. The evaluated rSVD route is workflow-only and excluded from primary "
        "numerical claims; full controls and audit status are reported in the Supplement. "
        "All 321 held-out spectra and response coordinates contribute."
    ),
    "Exploratory ImageNet experiment 1 used CUDA SIMPLS rSVD": (
        "Exploratory ImageNet experiment 1 used CUDA SIMPLS rSVD on 1,000,000 training "
        "and 281,167 held-out embeddings. The exact large-scale route was not "
        "independently numerically qualified and is therefore a workflow-feasibility "
        "analysis. LDA top-1 accuracy rose from 0.7793 at 100 components to 0.8093 at "
        "1,000 components; at 1,000 components, 227,535/281,167 observations were "
        "correct (Wilson 95% CI 0.8078-0.8107), and CUDA total time was 316.1 s "
        "(Figure 5)."
    ),
    "Exploratory FAISS retrieval used the same split and seed 123": (
        "Exploratory FAISS retrieval used the same split. Raw DINOv2 top-1/top-5 "
        "accuracy was 0.6556/0.9392; 200-component PLS gave 0.6516/0.9397 with "
        "5.12-fold compression and approximately fourfold lower projection-plus-query "
        "time. These single-run results show a possible compression trade-off, not an "
        "accuracy or numerical-equivalence claim (Supplementary Table S13)."
    ),
    "Figure 5. Exploratory ImageNet SIMPLS classification": (
        "Figure 5. Exploratory ImageNet SIMPLS classification across 100-1,000 "
        "components. The exact large-scale route was not independently numerically "
        "qualified; full solver controls and audit context are reported in the "
        "Supplement."
    ),
    "rSVD, implicit products, float32, CUDA, and Metal are optional": (
        "rSVD, implicit products, float32, CUDA, and Metal are optional implementation "
        "mechanisms around accelerated SIMPLS. Numerically qualified settings met the "
        "prespecified audit criteria, whereas the faster NMR and ImageNet configurations "
        "did not and remain exploratory. Detailed controls, tolerances, capability "
        "classifications, and route-specific outcomes are consolidated in the "
        "Supplement. CPU IRLBA remains the deterministic reference."
    ),
    "OPLS and kernel PLS demonstrate deterministic reuse": (
        "OPLS and kernel PLS demonstrate deterministic reuse of the accelerated core. "
        "NMR and ImageNet illustrate potential scale, but their rSVD results are "
        "exploratory. ImageNet does not establish biomedical validity. Limitations "
        "include finite component grids, conditional uncertainty, quadratic nonlinear "
        "kernels, approximate rSVD, and route-specific precision or accelerator "
        "disagreement."
    ),
}


SUPPLEMENT_NAVIGATION = (
    "Route-level evidence is consolidated here rather than repeated in the main text. "
    "Table S1 defines stage residency; Table S8 gives rSVD controls, thresholds, seeds, "
    "and audit counts; Table S9 is the authoritative float32 capability matrix; Table "
    "S11 reports paired CPU/CUDA/Metal performance and concordance; and Table S15 maps "
    "each analysis to its source state and archive."
)


def replace_main_paragraphs(document):
    replaced = set()
    for paragraph in document.paragraphs:
        for prefix, replacement in MAIN_REPLACEMENTS.items():
            if paragraph.text.startswith(prefix):
                paragraph.text = replacement
                replaced.add(prefix)
                break
    missing = set(MAIN_REPLACEMENTS) - replaced
    if missing:
        raise RuntimeError(f"Main-text paragraphs not found: {sorted(missing)}")


def insert_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    inserted.style = paragraph.style
    inserted.add_run(text)
    return inserted


def revise_supplement(document):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith("This supplement describes fastPLS version"):
            insert_after(paragraph, SUPPLEMENT_NAVIGATION)
            return
    raise RuntimeError("Supplement introduction paragraph not found")


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    main_document = Document(MAIN_SOURCE)
    replace_main_paragraphs(main_document)
    main_document.save(MAIN_OUTPUT)

    supplement_document = Document(SUPP_SOURCE)
    revise_supplement(supplement_document)
    supplement_document.save(SUPP_OUTPUT)

    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)


if __name__ == "__main__":
    main()
