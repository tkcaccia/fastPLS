from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle99"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle100"
OUTPUT.mkdir(parents=True, exist_ok=True)
MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle99_0.99.25_20260825.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle99_0.99.25_20260825.docx"
MAIN_OUTPUT = OUTPUT / "fastPLS_CMPB_main_cycle100_0.99.25_20260825.docx"
SUPP_OUTPUT = OUTPUT / "fastPLS_CMPB_supplement_cycle100_0.99.25_20260825.docx"


def paragraph_by_prefix(document, prefix):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}; found {len(matches)}")
    return matches[0]


def paragraph_containing(document, phrase):
    matches = [p for p in document.paragraphs if phrase in p.text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph containing {phrase!r}; found {len(matches)}")
    return matches[0]


def replace_paragraph(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    (paragraph.runs[0] if paragraph.runs else paragraph.add_run()).text = text


def insert_before(paragraph, text, style="Body Text"):
    new = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addprevious(new._p)
    return new


def insert_after(paragraph, text, style="Body Text"):
    new = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addnext(new._p)
    return new


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def format_table(table, widths):
    table.style = "Table"
    table.autofit = False
    table.alignment = 1
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.width = Inches(widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, "1F4E78")
            elif row_index % 2 == 0:
                set_cell_shading(cell, "EAF2F8")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if column_index < 2 else WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7.0)
                    if row_index == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_header(table.rows[0])


def insert_table_before(document, target, caption, headers, rows, widths):
    cap = document.add_paragraph(caption, style="Caption")
    cap.paragraph_format.keep_with_next = True
    target._p.addprevious(cap._p)
    table = document.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = value
    format_table(table, widths)
    target._p.addprevious(table._tbl)


main = Document(MAIN_SOURCE)

# Add a dedicated related-work subsection without conflating estimator identity.
contribution = paragraph_containing(main, "We present fastPLS, whose principal methodological contribution")
related_2 = insert_before(
    contribution,
    "A direct software comparator is IKPLS, which implements the two Improved Kernel PLS algorithms of Dayal and MacGregor using NumPy for CPU execution and JAX for CPU, GPU, and TPU execution [33,34]. IKPLS also combines fold-wise cross-product updates with vectorized JAX execution for fast validation [35]. It is not de Jong SIMPLS and therefore cannot establish estimator preservation, but it is an appropriate high-performance end-to-end comparator. Other accelerated PLS work includes iterative or randomized low-rank direction extraction, compressed or implicit products, parallel dense linear algebra, and GPU execution; these approaches trade different estimators, storage policies, precision, or compilation overhead against runtime [15,16,33-35].",
)
insert_before(
    related_2,
    "1.1 Related high-performance PLS software",
    style="Heading 2",
)
insert_after(
    related_2,
    "fastPLS addresses a different software regime: an R interface spanning PLS-SVD, de Jong SIMPLS, OPLS, and kernel PLS; compact latent prediction; optional implicit predictor-response products for large multivariate responses; compiled single and nested validation; and explicit solver/backend diagnostics. We therefore separate numerical-kernel validation against pls::simpls.fit from end-to-end software comparisons against IKPLS and other independent implementations.",
)

# Add the cross-language benchmark protocol as a distinct Methods subsection.
results_heading = paragraph_by_prefix(main, "3. Results")
methods_heading = insert_before(results_heading, "2.7 High-performance cross-language comparison", style="Heading 2")
insert_after(
    methods_heading,
    "Two comparisons answered different questions. Numerical-kernel validation compared deterministic float64 fastPLS SIMPLS with de Jong SIMPLS in pls::simpls.fit. The separate software comparison used fastPLS 0.99.25 and IKPLS 6.1.2 (NumPy Algorithms 1 and 2; JAX/CUDA Algorithms 1 and 2). Breast, MetRef, and CIFAR-100 used identical stored splits, externally training-centred float64 predictors, centred one-hot responses, 10, 22, and 50 components, respectively, and final held-out predictions as the common requested output. Each route ran three times in a fresh process. CPU runs used one effective thread. JAX results report cold host-to-device transfer, compilation, execution, prediction, and result transfer separately from warm execution; fastPLS CUDA totals include public-call transfer, synchronization, prediction, and result return. Because IKPLS and SIMPLS are different estimators and retain different internal state, these results are labelled end-to-end software comparisons rather than estimator-matched benchmarks.",
)

# Add concise results before the existing internal-acceleration section.
internal_heading = paragraph_by_prefix(main, "3.2 Internal acceleration")
replace_paragraph(internal_heading, "3.3 Internal acceleration and low-rank solver choice")
hp_heading = insert_before(internal_heading, "3.2 Cross-language high-performance PLS comparison", style="Heading 2")
hp_result_1 = insert_after(
    hp_heading,
    "The CPU comparison completed all 36 planned runs. IKPLS NumPy Algorithm 2 was fastest for all three tasks: median total times were 0.00068 s on Breast, 0.00388 s on MetRef, and 0.354 s on CIFAR-100, compared with 0.005, 0.037, and 5.155 s for fastPLS rSVD. Breast accuracy was identical (94.29%). IKPLS reached 75.0% on MetRef and 70.95% on CIFAR-100, whereas fastPLS rSVD reached 77.0% and 72.13%; deterministic fastPLS IRLBA reached 77.0% and 70.77%. Complete-process peak RSS on CIFAR-100 was 574 MB for IKPLS Algorithm 2 and 1,015 MB for fastPLS rSVD, but these cross-language RSS values include different runtimes and allocators.",
)
insert_after(
    hp_result_1,
    "On the NVIDIA RTX 5060 Ti, the cold end-to-end CUDA comparison also completed all 27 runs. Median fastPLS rSVD totals were 0.279 s on Breast, 0.309 s on MetRef, and 1.158 s on CIFAR-100; corresponding IKPLS JAX cold totals were 0.914-0.918, 0.981-1.058, and 1.859-1.861 s. Warm JAX CIFAR-100 totals remained 1.480-1.552 s. Accuracy was 94.29% for all Breast routes, 77.0% versus 75.0% on MetRef, and 71.57% versus 70.95% on CIFAR-100 for fastPLS and IKPLS, respectively. These findings do not establish universal superiority: IKPLS was substantially faster on CPU, whereas the tested fastPLS CUDA workflow was faster on this GPU. They define complementary software regimes and reinforce the need to report estimator, precision, hardware, compilation, transfer, and retained-output policies together (Supplementary Section S15).",
)

replace_paragraph(paragraph_by_prefix(main, "3.3 Multivariate NMR"), "3.4 Multivariate NMR prediction")
replace_paragraph(paragraph_by_prefix(main, "3.4 ImageNet"), "3.5 ImageNet-scale supervised representation")

# Expand the practical distinction in the Discussion without adding a superiority claim.
discussion = paragraph_by_prefix(main, "4. Discussion")
discussion_first = discussion._p.getnext()
discussion_paragraph = next(p for p in main.paragraphs if p._p is discussion_first)
insert_after(
    discussion_paragraph,
    "The IKPLS comparison places these claims in the current high-performance PLS landscape. Improved Kernel PLS was faster for the tested single-thread CPU workflows, especially on CIFAR-100, while fastPLS CUDA was faster than the tested JAX routes when transfer and cold compilation were included. The distinctive contribution of fastPLS is therefore not a claim to dominate every PLS kernel. It is the combination of de Jong SIMPLS validation, approximate rSVD with qualification metadata, compact and implicit multivariate-response workflows, multiple PLS families, R-native model selection, and route-aware CPU/CUDA/Metal execution. Users primarily requiring repeated classical IKPLS calibration in Python should consider IKPLS; users requiring the documented fastPLS workflow features should select routes according to matrix shape and the qualification tables.",
)

# Add primary references for Improved Kernel PLS and IKPLS.
main.add_paragraph(
    "[33] Dayal BS, MacGregor JF. Improved PLS algorithms. J Chemom. 1997;11:73-85. "
    "https://doi.org/10.1002/(SICI)1099-128X(199701)11:1<73::AID-CEM435>3.0.CO;2-#.",
    style="Body Text",
)
main.add_paragraph(
    "[34] Engstrom OCG, Dreier ES, Jespersen BM, Pedersen KS. IKPLS: Improved Kernel Partial Least Squares and fast cross-validation algorithms for Python with CPU and GPU implementations using NumPy and JAX. J Open Source Softw. 2024;9:6533. https://doi.org/10.21105/joss.06533.",
    style="Body Text",
)
main.add_paragraph(
    "[35] Engstrom OCG. Shortcutting cross-validation: efficiently deriving column-wise centred and scaled training-set X-transpose-X and X-transpose-Y without full recomputation. arXiv. 2024. https://doi.org/10.48550/arXiv.2401.13185.",
    style="Body Text",
)
main.save(MAIN_OUTPUT)


supp = Document(SUPP_SOURCE)
s16 = paragraph_by_prefix(supp, "S16. Selected-point")
s15_heading = insert_before(s16, "S15.1 Cross-language IKPLS software comparison", style="Heading 2")
intro = insert_after(
    s15_heading,
    "This comparison is not estimator matched. IKPLS 6.1.2 implements Dayal-MacGregor Improved Kernel PLS, whereas fastPLS implements de Jong SIMPLS with deterministic IRLBA or approximate rSVD direction extraction. The common contract was float64 input, externally applied training centring, identical centred one-hot responses, identical splits and component counts, and final held-out prediction. Each cell summarizes three fresh-process repetitions. CPU runs used one effective thread. Absolute and baseline-corrected RSS include language-runtime allocation and are interpreted as workflow feasibility rather than isolated algorithmic workspace.",
)

cpu_rows = [
    ("Breast", "fastPLS SIMPLS / IRLBA", "94.29", "0.0060", "0.0025", "186", "1.5"),
    ("Breast", "fastPLS SIMPLS / rSVD", "94.29", "0.0050", "0.0010", "186", "1.1"),
    ("Breast", "IKPLS NumPy Algorithm 1", "94.29", "0.00135", "0.00260", "131", "0.5"),
    ("Breast", "IKPLS NumPy Algorithm 2", "94.29", "0.00068", "0.00026", "131", "0.4"),
    ("MetRef", "fastPLS SIMPLS / IRLBA", "77.00", "0.023", "0.0015", "190", "1.8"),
    ("MetRef", "fastPLS SIMPLS / rSVD", "77.00", "0.037", "0.0020", "192", "4.3"),
    ("MetRef", "IKPLS NumPy Algorithm 1", "75.00", "0.00482", "0.00132", "135", "2.0"),
    ("MetRef", "IKPLS NumPy Algorithm 2", "75.00", "0.00388", "0.00110", "136", "3.4"),
    ("CIFAR-100", "fastPLS SIMPLS / IRLBA", "70.77", "4.408", "1.364", "1,083", "621"),
    ("CIFAR-100", "fastPLS SIMPLS / rSVD", "72.13", "5.155", "0.131", "1,015", "505"),
    ("CIFAR-100", "IKPLS NumPy Algorithm 1", "70.95", "1.084", "0.108", "560", "40"),
    ("CIFAR-100", "IKPLS NumPy Algorithm 2", "70.95", "0.354", "0.044", "574", "206"),
]
insert_table_before(
    supp, s16,
        "Table S10f. Single-thread CPU end-to-end comparison with IKPLS. Time and IQR are fit plus prediction seconds; RSS values are MB. rSVD used release-qualified controls: oversampling 20, two power iterations, and seed 123.",
    ("Dataset", "Workflow", "Accuracy (%)", "Median time", "Time IQR", "Peak RSS", "Incremental RSS"),
    cpu_rows, (0.70, 1.75, 0.72, 0.72, 0.66, 0.66, 0.76),
)

cuda_rows = [
    ("Breast", "fastPLS CUDA rSVD", "94.29", "0.279", "not JIT"),
    ("Breast", "IKPLS JAX Algorithm 1", "94.29", "0.918", "0.698"),
    ("Breast", "IKPLS JAX Algorithm 2", "94.29", "0.914", "0.724"),
    ("MetRef", "fastPLS CUDA rSVD", "77.00", "0.309", "not JIT"),
    ("MetRef", "IKPLS JAX Algorithm 1", "75.00", "1.058", "0.779"),
    ("MetRef", "IKPLS JAX Algorithm 2", "75.00", "0.981", "0.769"),
    ("CIFAR-100", "fastPLS CUDA rSVD", "71.57", "1.158", "not JIT"),
    ("CIFAR-100", "IKPLS JAX Algorithm 1", "70.95", "1.859", "1.480"),
    ("CIFAR-100", "IKPLS JAX Algorithm 2", "70.95", "1.861", "1.552"),
]
insert_table_before(
    supp, s16,
        "Table S10g. NVIDIA RTX 5060 Ti end-to-end comparison. Cold JAX time includes transfer, JIT compilation, execution, prediction, and result transfer; warm JAX time reuses compiled executables. fastPLS public-call time includes transfer, execution, synchronization, prediction, and result return. Times are medians of three fresh processes.",
    ("Dataset", "Workflow", "Accuracy (%)", "Cold total (s)", "Warm total (s)"),
    cuda_rows, (0.85, 2.20, 0.90, 1.05, 1.05),
)
insert_before(
    s16,
    "Reproducibility. CPU results are in benchmark_results/ikpls_cross_language_20260825/ikpls_cross_language_summary.csv. CUDA rows are in the cuda_rows directory and summarized by ikpls_cross_language_cuda_summary.csv. The benchmark used fastPLS 0.99.25, IKPLS 6.1.2, JAX 0.6.2, float64, three fresh processes, and fixed component counts of 10 (Breast), 22 (MetRef), and 50 (CIFAR-100). The deterministic de Jong comparison remains in Section S12 and must not be replaced by these end-to-end results.",
)

mapping = paragraph_containing(supp, "References are numbered as in the main manuscript")
replace_paragraph(
    mapping,
    "References are numbered as in the main manuscript. The final mapping is Retina [25], Tabula Muris [26], PRISM [27], ImageNet [28], DINOv2 [29], CIFAR-100 [30], UNI [31], Prov-GigaPath [32], Improved Kernel PLS [33], IKPLS [34], and fast cross-product validation [35].",
)
supp.save(SUPP_OUTPUT)

(OUTPUT / "README.md").write_text(
    "# Cycle 100\n\nAdds a dedicated high-performance PLS related-work section and a matched "
    "cross-language fastPLS 0.99.25 versus IKPLS 6.1.2 benchmark. Numerical-kernel "
    "validation against de Jong SIMPLS remains separate from end-to-end software timing.\n",
    encoding="utf-8",
)

print(MAIN_OUTPUT)
print(SUPP_OUTPUT)
