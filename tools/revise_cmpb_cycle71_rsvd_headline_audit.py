#!/usr/bin/env python3
"""Make rSVD controls and audit status explicit beside every headline result."""

from pathlib import Path
import re

from docx import Document

from revise_cmpb_cycle67_consolidate_evidence import (
    find_paragraph,
    normalize_submission_terminology,
    replace_paragraph,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE = (
    ROOT
    / "artifacts"
    / "CMPB_rewrite_20260726_cycle70"
    / "fastPLS_CMPB_main_cycle70_0.99.6_20260726.docx"
)
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle71"
OUTPUT = OUTPUT_DIR / "fastPLS_CMPB_main_cycle71_0.99.6_20260726.docx"


def word_count(text):
    return len(re.findall(r"\b[\w'-]+\b", text))


def revise_abstract(document):
    replace_paragraph(
        document,
        "Methods:",
        (
            "Methods: The accelerated SIMPLS path reuses deflation products, latent "
            "quantities, coefficients, and predictions across components and "
            "requested model prefixes. Compact prediction and optional implicit "
            "cross-covariance products reduce avoidable storage. Deterministic IRLBA "
            "provides the confirmatory solver. Approximate randomized SVD (rSVD) was "
            "audited separately: oversampling by 10 directions with one power "
            "iteration passed 101/117 checks, whereas two power iterations passed "
            "117/117 across synthetic seeds 101, 202, and 303 and real-data seed 123. "
            "Compiled CPU is the reference implementation, with qualified NVIDIA "
            "CUDA and Apple Metal routes. Runtime, memory, prediction agreement, and "
            "predictive performance were evaluated against de Jong SIMPLS and "
            "independent R software."
        ),
    )
    replace_paragraph(
        document,
        "Results:",
        (
            "Results: Deterministic fastPLS SIMPLS met all prespecified tolerances "
            "in 117 component-level comparisons with de Jong SIMPLS. In matched "
            "single-CPU IRLBA comparisons, it was faster than pls::simpls.fit on "
            "seven of nine datasets with identical accuracy. The two-power rSVD "
            "configuration passed all numerical checks; the faster one-power NMR "
            "and ImageNet runs did not inherit that status and are reported only as "
            "exploratory workflow-feasibility analyses. OPLS and kernel PLS "
            "extensions passed their separate deterministic reliability study."
        ),
    )
    replace_paragraph(
        document,
        "Conclusions:",
        (
            "Conclusions: fastPLS makes sequential SIMPLS substantially more "
            "practical while preserving the deterministic estimator. Primary claims "
            "are supported by deterministic IRLBA or the fully audited two-power "
            "rSVD configuration; unaudited or failed-audit randomized workflows are "
            "not used as confirmatory evidence."
        ),
    )


def revise_methods(document):
    replace_paragraph(
        document,
        "Direction extraction is modular",
        (
            "Direction extraction is modular rather than the principal estimator "
            "contribution. Float64 CPU fitting supports deterministic IRLBA [15] and "
            "approximate rSVD [16]. rSVD uses a Gaussian range sketch, "
            "orthonormalization, power iterations, and a reduced decomposition. The "
            "prespecified SIMPLS audit used oversampling 10: one power iteration "
            "passed 101/117 component-level checks, whereas two power iterations "
            "passed 117/117 across synthetic seeds 101, 202, and 303 and real-data "
            "seed 123. Only the two-power setting is numerically qualified by this "
            "audit. CUDA and Metal rSVD require route-specific qualification "
            "(Supplementary Table S8)."
        ),
    )
    replace_paragraph(
        document,
        "The primary validation tested whether accelerated deterministic SIMPLS",
        (
            "The primary evidence tested whether deterministic IRLBA SIMPLS "
            "preserved de Jong's estimator and improved runtime. The rSVD audit was "
            "a separate approximation study with controls and seeds recorded for "
            "each result. Backend portability, float32, OPLS/kernel extensions, "
            "cross-validation, and repeated outer partitions were also supporting "
            "analyses. A completed run was not considered numerically qualified "
            "unless it passed the prespecified audit criteria. Detailed designs, "
            "thresholds, provenance, and complete results are in Supplementary "
            "Tables S3-S15."
        ),
    )


def revise_results(document):
    replace_paragraph(
        document,
        "Results are organized around the accelerated SIMPLS contribution:",
        (
            "Results are organized around the accelerated SIMPLS contribution. "
            "Primary evidence comprises deterministic estimator preservation and "
            "the matched IRLBA comparison with independent implementations. rSVD, "
            "precision, and hardware results are reported separately with their "
            "controls and numerical-audit status; NMR and ImageNet one-power runs "
            "are exploratory feasibility studies rather than confirmatory evidence."
        ),
    )
    replace_paragraph(
        document,
        "The primary estimator-preservation study found",
        (
            "Deterministic fastPLS SIMPLS met all prespecified tolerances in 117 "
            "component-level comparisons with de Jong SIMPLS. In the separate rSVD "
            "audit, oversampling 10 with one power iteration passed 101/117 checks; "
            "oversampling 10 with two power iterations passed 117/117 across seeds "
            "101, 202, 303, and 123 for the real datasets. Separately, OPLS and "
            "kernel PLS extensions passed all 66 setting/task comparisons and all "
            "1,540 fold-component fits under deterministic float64 CPU settings "
            "(Supplementary Table S7)."
        ),
    )
    replace_paragraph(
        document,
        "Hardware acceleration was a supporting",
        (
            "Hardware acceleration was a supporting workflow analysis. The reported "
            "non-NMR rSVD backend rows used oversampling 10, one power iteration, and "
            "replicate seeds 124-126. Because the corresponding SIMPLS setting passed "
            "only 101/117 approximation checks, these rows are workflow-only and do "
            "not support estimator-preservation claims. Of 44 CPU-CUDA pairs, 28 met "
            "the paired predictive criteria and CUDA was faster in seven; six of 12 "
            "CPU-Metal pairs were concordant and none was faster with Metal "
            "(Figure 3; Supplementary Table S11)."
        ),
    )
    replace_paragraph(
        document,
        "The choice of direction solver modified",
        (
            "The one-power speed comparison used oversampling 10 and seeds 124-126; "
            "its numerical status was 101/117 checks passed and therefore "
            "workflow-only. CPU rSVD was 1.00-1.45-fold faster than IRLBA across nine "
            "classification tasks, but MetRef accuracy differed by 4.0 percentage "
            "points. On NMR at 100 components, it reduced CPU time from 436.3 to "
            "19.6 s. These values describe the speed-risk trade-off and do not "
            "support primary claims. IRLBA remains the deterministic reference; the "
            "qualified rSVD setting is oversampling 10, two power iterations, with "
            "117/117 audit checks passed (Supplementary Table S8)."
        ),
    )
    replace_paragraph(
        document,
        "Figure 3. Numerically qualified backend",
        (
            "Figure 3. Supporting backend and solver workflow comparisons. rSVD "
            "rows used oversampling 10, one power iteration, and seeds 124-126; this "
            "setting passed 101/117 approximation checks and is labelled "
            "workflow-only. Discordant routes are excluded from speed-up summaries."
        ),
    )
    replace_paragraph(
        document,
        "NMR comprised 1,200 training",
        (
            "NMR comprised 1,200 training and 321 held-out spectra, with 13,000 "
            "predictors and 28,355 responses. Predictor columns between 4.6 and "
            "4.8 ppm were zeroed in training and test data; responses were unmasked. "
            "Training selection and selected-point rSVD runs used oversampling 10, "
            "one power iteration, and fit/replicate seeds 123 and 124-126. Because "
            "this setting passed 101/117 checks, all NMR rSVD values below are "
            "exploratory workflow results, not primary numerical evidence. "
            "Training-only one-standard-error selection retained five PLS-SVD and "
            "50 SIMPLS components."
        ),
    )
    replace_paragraph(
        document,
        "Selected CUDA PLS-SVD and SIMPLS achieved",
        (
            "Under that workflow-only one-power configuration, selected CUDA "
            "PLS-SVD and SIMPLS achieved RMSD 0.001043 (95% bootstrap interval "
            "0.001000-0.001085) and 0.000759 (0.000665-0.000884), respectively. "
            "SIMPLS also had lower median per-spectrum, response-wise, and "
            "high-intensity errors (Figure 4). These results demonstrate feasibility "
            "but do not validate rSVD estimator agreement."
        ),
    )
    replace_paragraph(
        document,
        "Holding family, split, rSVD",
        (
            "In the matched backend benchmark, family, split, float64 precision, "
            "component count, oversampling 10, one power iteration, and seeds 124-126 "
            "were fixed. CUDA reduced PLS-SVD time from 2.301 to 0.648 s and SIMPLS "
            "time from 10.525 to 1.773 s, with CPU-CUDA prediction correlations "
            "1.000000 and 0.999981. This comparison establishes backend agreement "
            "within the one-power workflow, not agreement with the deterministic "
            "estimator. CUDA device increments were 590 and 3,414 MB."
        ),
    )
    replace_paragraph(
        document,
        "Figure 4. NMR predictive and implementation analyses:",
        (
            "Figure 4. Exploratory NMR rSVD workflow analyses: training-only "
            "component selection, held-out errors, intensity-stratified RMSD, and "
            "matched CPU/CUDA resources. All rSVD rows used oversampling 10, one "
            "power iteration, and fit/replicate seeds 123 and 124-126. Audit status: "
            "101/117 checks passed; workflow-only and excluded from primary claims. "
            "All 321 held-out spectra and response coordinates contribute."
        ),
    )
    replace_paragraph(
        document,
        "In exploratory ImageNet experiment 1",
        (
            "Exploratory ImageNet experiment 1 used CUDA SIMPLS rSVD with "
            "oversampling 10, one power iteration, and seed 123 on 1,000,000 "
            "training and 281,167 held-out embeddings. This exact large-scale route "
            "was not separately shown to pass the numerical audit; the related "
            "one-power SIMPLS configuration passed 101/117 checks. The results are "
            "therefore workflow-feasibility estimates only. LDA top-1 accuracy rose "
            "from 0.7793 at 100 components to 0.8093 at 1,000 components; at 1,000 "
            "components, 227,535/281,167 observations were correct (Wilson 95% CI "
            "0.8078-0.8107), and CUDA total time was 316.1 s (Figure 5)."
        ),
    )
    replace_paragraph(
        document,
        "Separate FAISS retrieval compared",
        (
            "Exploratory FAISS retrieval used the same split and seed 123; PLS "
            "representations used rSVD with oversampling 10 and one power iteration "
            "and were not numerically qualified. Raw DINOv2 top-1/top-5 accuracy was "
            "0.6556/0.9392; 200-component PLS gave 0.6516/0.9397 with 5.12-fold "
            "compression and approximately fourfold lower projection-plus-query "
            "time. These single-run results show a possible compression trade-off, "
            "not an accuracy or estimator-preservation claim (Supplementary Table "
            "S13)."
        ),
    )
    replace_paragraph(
        document,
        "Figure 5. Exploratory ImageNet SIMPLS",
        (
            "Figure 5. Exploratory ImageNet SIMPLS classification across 100-1,000 "
            "components. rSVD used oversampling 10, one power iteration, and seed "
            "123. Audit status: not separately qualified; the related setting passed "
            "101/117 checks. Results are workflow-feasibility estimates only."
        ),
    )


def revise_discussion(document):
    replace_paragraph(
        document,
        "rSVD, implicit products, float32",
        (
            "rSVD, implicit products, float32, CUDA, and Metal are optional "
            "implementation mechanisms around accelerated SIMPLS. Only oversampling "
            "10 with two power iterations passed all 117 rSVD checks in the broad "
            "audit. The faster one-power NMR and ImageNet workflows passed 101/117 "
            "at setting level or were not separately qualified; they demonstrate "
            "feasibility but cannot support confirmatory claims. CPU IRLBA remains "
            "the deterministic reference."
        ),
    )
    replace_paragraph(
        document,
        "OPLS and kernel PLS demonstrate reuse",
        (
            "OPLS and kernel PLS demonstrate deterministic reuse of the accelerated "
            "core. NMR and ImageNet illustrate potential scale but their one-power "
            "rSVD results are exploratory. ImageNet does not establish biomedical "
            "validity. Limitations include finite component grids, conditional "
            "uncertainty, quadratic nonlinear kernels, approximate rSVD, and "
            "route-specific precision or accelerator disagreement."
        ),
    )


def audit(document):
    abstract = " ".join(
        find_paragraph(document, prefix).text
        for prefix in (
            "Background and objective:",
            "Methods:",
            "Results:",
            "Conclusions:",
        )
    )
    introduction = next(
        i
        for i, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == "1. Introduction"
    )
    references = next(
        i
        for i, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == "References"
    )
    main_text = " ".join(
        paragraph.text for paragraph in document.paragraphs[introduction:references]
    )
    if word_count(abstract) > 350:
        raise RuntimeError("Abstract exceeds 350 words")
    if word_count(main_text) > 3500:
        raise RuntimeError("Main text exceeds 3,500 words")

    joined = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for required in (
        "one power iteration passed 101/117",
        "two power iterations passed 117/117",
        "seeds 124-126",
        "seed 123",
        "workflow-feasibility estimates only",
        "excluded from primary claims",
    ):
        if required not in joined:
            raise RuntimeError(f"Missing rSVD audit statement: {required}")
    print(f"Abstract words: {word_count(abstract)}")
    print(f"Main-text words: {word_count(main_text)}")


def main():
    document = Document(SOURCE)
    revise_abstract(document)
    revise_methods(document)
    revise_results(document)
    revise_discussion(document)
    normalize_submission_terminology(document)
    find_paragraph(document, "Algorithm 1.").paragraph_format.page_break_before = True
    audit(document)
    document.core_properties.title = (
        "fastPLS CMPB manuscript - rSVD headline audit controls"
    )
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
