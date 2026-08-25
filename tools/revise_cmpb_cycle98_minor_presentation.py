from copy import deepcopy
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MAIN_SOURCE = (
    ROOT
    / "artifacts"
    / "CMPB_rewrite_20260825_cycle97"
    / "fastPLS_CMPB_main_cycle97_0.99.25_20260825.docx"
)
SUPP_SOURCE = (
    ROOT
    / "artifacts"
    / "CMPB_rewrite_20260825_cycle96"
    / "fastPLS_CMPB_supplement_cycle96_0.99.25_20260825.docx"
)
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle98"
OUTPUT.mkdir(parents=True, exist_ok=True)
MAIN_OUTPUT = OUTPUT / "fastPLS_CMPB_main_cycle98_0.99.25_20260825.docx"
SUPP_OUTPUT = OUTPUT / "fastPLS_CMPB_supplement_cycle98_0.99.25_20260825.docx"
FIGURE_3 = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle62_20260726"
    / "accelerator_concordance_speedups.png"
)


def paragraph_by_prefix(document, prefix):
    matches = [p for p in document.paragraphs if p.text.startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}; found {len(matches)}")
    return matches[0]


def paragraph_containing(document, phrase):
    matches = [p for p in document.paragraphs if phrase in p.text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph containing {phrase!r}; found {len(matches)}")
    return matches[0]


def replace_paragraph(paragraph, text):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.text = text


def insert_before(paragraph, text, style="Body Text"):
    new = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addprevious(new._p)
    return new


def insert_after(paragraph, text, style="Body Text"):
    new = paragraph._parent.add_paragraph(text, style=style)
    paragraph._p.addnext(new._p)
    return new


def replace_media(source, destination, member, data):
    temporary = destination.with_suffix(".media.docx")
    with ZipFile(source, "r") as zin, ZipFile(
        temporary, "w", compression=ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            zout.writestr(item, data if item.filename == member else zin.read(item.filename))
    temporary.replace(destination)


def normalize_terms(document):
    replacements = (
        ("PLSSVD", "PLS-SVD"),
        ("RSVD", "rSVD"),
        ("randomized-SVD", "rSVD"),
        ("latent variable", "component"),
        ("latent-variable", "component"),
    )
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            for old, new in replacements:
                run.text = run.text.replace(old, new)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for old, new in replacements:
                            run.text = run.text.replace(old, new)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def format_table(table, widths):
    table.style = "Table"
    table.autofit = False
    table.alignment = 1
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            cell.width = Inches(widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                set_cell_shading(cell, "1F4E78")
            elif row_index % 2 == 0:
                set_cell_shading(cell, "EAF2F8")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if column_index == 0
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7.2)
                    if row_index == 0:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
    set_repeat_header(table.rows[0])


def insert_table_before(document, target, caption_text, headers, rows, widths):
    caption = document.add_paragraph(caption_text, style="Caption")
    caption.paragraph_format.keep_with_next = True
    target._p.addprevious(caption._p)
    table = document.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = value
    format_table(table, widths)
    target._p.addprevious(table._tbl)
    return table


main = Document(MAIN_SOURCE)
replace_paragraph(
    main.paragraphs[0],
    "fastPLS: a memory-aware accelerated SIMPLS implementation for high-dimensional "
    "biomedical data with optional CUDA and Metal routes",
)
replace_paragraph(
    main.paragraphs[9],
    "Background and objective: High-dimensional biomedical PLS workflows, including "
    "multivariate NMR prediction and repeated validation, can be limited by the sequential "
    "cost and storage of SIMPLS. We developed fastPLS to accelerate SIMPLS while retaining "
    "de Jong's component equations.",
)
replace_paragraph(
    main.paragraphs[10],
    "Methods: The reviewed public interface is fastPLS 0.99.25; analysis-specific executing "
    "archives and commits are mapped in the Supplement. The implementation uses compiled, "
    "shape-dependent execution, incremental coefficient and fitted-value updates, compact "
    "latent prediction, and optional implicit cross-covariance products. Fixed-control CPU "
    "IRLBA, an iterative truncated solver rather than an exact dense decomposition, was the "
    "deterministic numerical reference. Approximate rSVD and accelerator routes were audited "
    "separately under prespecified criteria.",
)
replace_paragraph(
    main.paragraphs[11],
    "Results: fastPLS SIMPLS met the prespecified deterministic numerical tolerances in all "
    "117 component-level comparisons. In 108 repeated single-CPU argmax comparisons with "
    "pls::simpls.fit, accuracy was identical throughout; fastPLS was faster on five of nine "
    "datasets under ordinary public workflows, with a largest speed-up of 4.84-fold. LDA "
    "results were evaluated separately as complete-workflow comparisons. For the 13,000 by "
    "28,355 NMR predictor-response problem at 50 SIMPLS components, qualified CPU and CUDA "
    "rSVD reduced the fixed-control IRLBA runtime from 350.7 s to 9.80 and 1.51 s, "
    "respectively, while meeting the prespecified numerical tolerances.",
)

intro = paragraph_by_prefix(main, "Among PLS formulations")
replace_paragraph(
    intro,
    "Among PLS formulations, SIMPLS constructs sequential components without explicitly "
    "deflating the predictor matrix [11]. PLS-SVD provides a one-shot cross-covariance "
    "comparator [10], while OPLS and kernel PLS extend the framework to orthogonal filtering "
    "and nonlinear relations [12-14]. Reference software such as pls::simpls.fit already "
    "returns coefficient and fitted-value arrays for every component prefix; fastPLS does not "
    "claim otherwise. Its distinct execution features are compact latent prediction, "
    "shape-dependent intermediate storage, optional implicit cross-covariance products, "
    "compiled low-rank solvers, and route-qualified CUDA and Metal support. Supplementary "
    "Table S1a defines this feature comparison explicitly.",
)

solver = paragraph_by_prefix(main, "Direction extraction is modular")
replace_paragraph(
    solver,
    "Direction extraction is modular rather than the principal estimator contribution. "
    "float64 CPU fitting supports fixed-control IRLBA [15] as the deterministic numerical "
    "reference and approximate rSVD [16], which uses a Gaussian range sketch, "
    "orthonormalization, power iterations, and a reduced decomposition. IRLBA is an iterative "
    "truncated solver and is not described as an exact dense SVD. In release 0.99.23, the rSVD "
    "default was changed to oversampling 20 and two power iterations, the CPU and CUDA setting "
    "that met every prespecified check across five seeds. Explicit unqualified overrides warn "
    "and are recorded in model diagnostics. Exact controls, seeds, audit thresholds, and "
    "route-specific qualification are consolidated in Supplementary Sections S13-S14 and "
    "Tables S8-S9.",
)

algorithm = main.tables[0].cell(0, 0)
replace_paragraph(
    algorithm.paragraphs[0],
    "Input: centred X; numeric, indicator, or label-aware response Y; maximum component count A; requested component-count set C; solver.\n"
    "1. Define S₀ = XᵀY explicitly, or define operator products S(v) = Xᵀ(Yv) and Sᵀ(u) = Yᵀ(Xu); label-aware products use class sums.\n"
    "2. If p ≤ n and storage permits, cache G = XᵀX. Initialize R, Q, and V as empty; set fitted values Ŷ₀ = 0.\n"
    "3. For component a = 1, …, A:\n"
    "3.1 Extract the leading right direction gₐ of state Sₐ₋₁ using a fresh IRLBA solve or a fresh oversampled rSVD sketch.\n"
    "3.2 Set rₐ = Sₐ₋₁gₐ and tₐ = Xrₐ; divide rₐ and tₐ by ‖tₐ‖₂.\n"
    "3.3 Compute predictor and response loadings pₐ = Xᵀtₐ and qₐ = Yᵀtₐ.\n"
    "3.4 Orthogonalize vₐ = (I − VVᵀ)pₐ and normalize vₐ.\n"
    "3.5 Deflate Sₐ = Sₐ₋₁ − vₐ(vₐᵀSₐ₋₁), or update the equivalent implicit operator.\n"
    "3.6 Append rₐ, qₐ, and vₐ; update Bₐ = Bₐ₋₁ + rₐqₐᵀ and Ŷₐ = Ŷₐ₋₁ + tₐqₐᵀ.\n"
    "3.7 If a ∈ C, store only the requested coefficient/prediction snapshot or compact latent representation.\n"
    "Output: the standard sequential component path; requested outputs are retained as dense snapshots or compact latent factors.",
)

boundary = paragraph_by_prefix(main, "Repeated outer partitions showed")
replace_paragraph(
    boundary,
    "Repeated outer partitions showed that predictive variation exceeded timing variation. "
    "Among 46 evaluated family-dataset selections, 24 (52.2%) occurred at a tested-grid "
    "boundary and another nine were response-rank limited; thus 33/46 (71.7%) were constrained "
    "rather than unconstrained optima. They are reported as best within the evaluated grid "
    "(Supplementary Table S14).",
)

figure2 = paragraph_by_prefix(main, "Figure 2.")
replace_paragraph(
    figure2,
    "Figure 2. Repeated float64 single-CPU SIMPLS timing against pls::simpls.fit. Points "
    "are median fit-plus-prediction times; bars are IQRs from n = 3 fresh-process "
    "repetitions per method-dataset pair. Left: complete coefficient paths and final "
    "predictions with dense scores, loadings, and fitted arrays suppressed. Right: ordinary "
    "public objects. Accuracy was identical in every argmax pair; full package-by-dataset "
    "values are in Supplementary Tables S10a-S10e.",
)

accelerator = paragraph_by_prefix(main, "Hardware acceleration remained route")
replace_paragraph(
    accelerator,
    "Hardware acceleration remained route and shape dependent. CPU, CUDA, and Metal runtime "
    "ratios were summarized only for paired predictions meeting the strict display criteria "
    "of absolute endpoint-metric difference ≤0.005 and prediction agreement ≥0.995 (Figure 3; "
    "Supplementary Table S11). These thresholds are deliberately stricter than the broader "
    "0.01/0.99 rSVD qualification limits, because Figure 3 labels selected workflows as "
    "near-concordant for direct runtime interpretation, whereas the qualification audit "
    "screens approximation across complete component paths and several seeds. The former "
    "oversampling-10, one-power CPU setting and rejected CUDA settings remain quarantined. "
    "Release 0.99.23 defaults to oversampling 20 and two power iterations, which met 585/585 "
    "CPU and 40/40 CUDA checks across five seeds; Metal rSVD remains unqualified. Fixed-control "
    "CPU IRLBA remains the deterministic numerical reference. float32 approximately halved "
    "stored inputs on MetRef and PRISM but did not uniformly improve runtime, incremental "
    "memory, or agreement (Supplementary Table S9).",
)

figure3 = paragraph_by_prefix(main, "Figure 3.")
replace_paragraph(
    figure3,
    "Figure 3. CPU/accelerator runtime ratios for numerically concordant workflows. Cells "
    "are colored only when the absolute selected-endpoint metric difference was at most 0.005 "
    "and paired-prediction agreement was at least 0.995; ratios above one favor the "
    "accelerator and ratios below one favor CPU. These stricter display thresholds identify "
    "near-concordant selected workflows; the broader multi-seed rSVD qualification uses "
    "0.01/0.99 across component paths. Gray cells indicate discordance or missing paired "
    "predictions. Full paired values are in Supplementary Table S11.",
)

figure4 = paragraph_by_prefix(main, "Figure 4.")
replace_paragraph(
    figure4,
    "Figure 4. NMR predictive and computational analyses. Panels A-C separate family-selected "
    "held-out performance from the historical deposited 165-component workflow, which differs "
    "in component count, implementation, and hardware and is not a matched reference. Reported "
    "Q² is independent-test Q² relative to the training-response mean. Panels D-E overlay "
    "observed and predicted intensities for held-out sample AMI-00BP-8 (index 155), selected "
    "as closest to the median per-spectrum RMSD under 50-component SIMPLS CUDA rSVD, over the "
    "full response range and 1.7-0.5 ppm expansion. Panel F reports matched float64 "
    "solver/backend resources at fixed family-specific component counts. rSVD used "
    "oversampling 20, two power iterations, and seed 123.",
)

figure5 = paragraph_by_prefix(main, "Figure 5.")
replace_paragraph(
    figure5,
    "Figure 5. Exploratory ImageNet matrix-processing feasibility with float32 SIMPLS-LDA. "
    "Seed 123 assigned 1,000,000 rows from the pooled 1,281,167-row DINOv2 embedding archive "
    "to training and the complementary 281,167 rows to holdout; this was not the canonical "
    "ImageNet split. Accuracy and blocked prediction time are shown for requested 100-1,000-"
    "component prefixes from one shared fit; 1,000 is a boundary stress point, not an optimum. "
    "The route is hybrid: SIMPLS deflation and score projection are host-resident, whereas "
    "rSVD range products and LDA use CUDA. Controls were oversampling 20, two power "
    "iterations, and seed 123. Values are single-run exploratory feasibility measurements.",
)

declaration = paragraph_by_prefix(main, "Declaration of competing interest")
credit_heading = insert_before(
    declaration, "CRediT authorship contribution statement", style="Heading 1"
)
insert_after(
    credit_heading,
    "Dupe Ojo: Investigation, validation, data curation, visualization, writing - review and "
    "editing. Alessia Vignoli: Investigation, validation, data curation, visualization, writing "
    "- review and editing. Stefano Cacciatore: Conceptualization, methodology, software, "
    "supervision, project administration, writing - original draft, writing - review and "
    "editing. Leonardo Tenori: Conceptualization, methodology, supervision, writing - review "
    "and editing. This draft allocation must be confirmed by all authors before submission.",
)
ack_heading = insert_before(declaration, "Acknowledgements", style="Heading 1")
insert_after(
    ack_heading,
    "The authors thank the University of Cape Town's ICTS High Performance Computing team for "
    "providing a high-performance computing facility for this study "
    "(https://ucthpc.uct.ac.za/).",
)

availability = paragraph_by_prefix(main, "Code and benchmark outputs are available")
replace_paragraph(
    availability,
    availability.text
    + " The accepted source, vignette, reference manual, benchmark scripts, and definitive key "
    "tables will be deposited as one immutable tagged release with a persistent archive "
    "identifier; until that deposit is minted, the SHA-256-identified source archive is the "
    "immutable review object.",
)

normalize_terms(main)
main.save(MAIN_OUTPUT)
replace_media(
    MAIN_OUTPUT,
    MAIN_OUTPUT,
    "word/media/image24.png",
    FIGURE_3.read_bytes(),
)


supp = Document(SUPP_SOURCE)
replace_paragraph(
    supp.paragraphs[1],
    "fastPLS: a memory-aware accelerated SIMPLS implementation for high-dimensional "
    "biomedical data with optional CUDA and Metal routes",
)

s2 = paragraph_by_prefix(supp, "S2. Numerical algorithms")
insert_table_before(
    supp,
    s2,
    "Table S1a. Feature comparison between the primary deterministic SIMPLS reference and fastPLS. A check indicates that the public route provides the feature; qualifications are stated explicitly.",
    ["Feature", "pls::simpls.fit", "fastPLS 0.99.25"],
    [
        ("One fit returns component prefixes", "Yes", "Yes; not claimed as novel"),
        ("Minimum-output fit", "stripped mode retains coefficients", "Optional suppression of scores, loadings, fitted values, and variance summaries"),
        ("Compact latent prediction", "No documented compact latent model", "Yes; blocked prediction without a dense coefficient/prediction path"),
        ("Implicit XᵀY operator", "No", "Yes; shape-dependent and primarily memory-oriented"),
        ("Truncated solvers", "Reference SIMPLS direction calculation", "Fixed-control CPU IRLBA or approximate rSVD"),
        ("Accelerator routes", "No native CUDA/Metal route", "Route-specific CUDA and Metal; hybrid/experimental status remains explicit"),
    ],
    [1.65, 2.15, 2.70],
)

irlba = paragraph_by_prefix(supp, "The double CPU route bundles")
replace_paragraph(
    irlba,
    "The float64 CPU route bundles augmented implicitly restarted Lanczos bidiagonalization. "
    "Under fixed inputs, starting state, and controls it is the deterministic numerical "
    "reference used here, but it remains an iterative truncated solver rather than an exact "
    "dense decomposition. It expands and restarts a bidiagonal Krylov subspace to approximate "
    "the requested dominant singular triplets. The default public controls are maxit=1000, "
    "tol=1e-5, eps=1e-9, and svtol=1e-5 unless overridden through .... Exact fallback is used "
    "only when either input dimension is below six. IRLBA is otherwise retained even when the "
    "requested rank approaches the smaller matrix dimension. The native float32 CPU and Metal "
    "routes use separate single-precision IRLBA-style implementations; CUDA float32 currently "
    "supports rSVD only.",
)

lda = paragraph_by_prefix(supp, "For each class, the system")
insert_after(
    lda,
    "Fallback-use audit. In the fixed-score numerical study, all 144 candidate float32 LDA "
    "fits (72 CPU and 72 CUDA) succeeded at the first nominal relative regularization value, "
    "rho = 1e-8; no escalation occurred and CPU/CUDA selected values did not differ beyond "
    "float32 representation. The archived public end-to-end workflow rows did not retain the "
    "selected fallback value, so invocation frequencies are not claimed for those older "
    "benchmarks. Future publication reruns must record rho for every fitted prefix and backend.",
)

repo_only = paragraph_by_prefix(supp, "S21. Repository-only detailed material")
replace_paragraph(repo_only, "S22. Repository-only detailed material")
synthetic_heading = insert_before(
    repo_only, "S21. Unrestricted synthetic end-to-end reproduction", style="Heading 1"
)
insert_after(
    synthetic_heading,
    "benchmark/synthetic_end_to_end_example.R reproduces every public analysis stage without "
    "restricted data: seeded simulation, a fixed train/holdout split, training-only component "
    "and classifier selection, final refitting, held-out prediction, classification and "
    "regression evaluation, timing, result serialization, and session capture. It writes its "
    "outputs to benchmark_results/synthetic_end_to_end_example/ and is executed with "
    "Rscript benchmark/synthetic_end_to_end_example.R from the repository root.",
)

references = paragraph_by_prefix(supp, "Supplementary references")
ci_heading = insert_before(
    references, "S23. Continuous-integration and platform-test status", style="Heading 1"
)
insert_after(
    ci_heading,
    "The latest public BiocStaging matrix available during this audit tested fastPLS 0.99.22 "
    "at commit ba80b65f0c660478af0573428e0957b4e0a7e382. Four Linux CPU, four macOS build, and "
    "five Windows CPU configurations completed with package-check status OK; source and WebAssembly "
    "jobs also completed. These hosted macOS jobs establish compilation and CPU fallback, not "
    "Metal runtime execution. Linux CUDA and macOS Metal runtime qualification used dedicated "
    "hardware benchmarks rather than continuous integration, and this remains a limitation. "
    "The local 0.99.25 R CMD check --as-cran audit completed with zero errors and warnings and "
    "one expected version-jump note. Unit tests request unavailable backend/solver combinations "
    "and verify informative failure before model allocation. Public CI metadata: "
    "https://biocstaging.r-universe.dev/fastPLS.",
)

normalize_terms(supp)
supp.save(SUPP_OUTPUT)

print(MAIN_OUTPUT)
print(SUPP_OUTPUT)
