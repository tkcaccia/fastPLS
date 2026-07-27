#!/usr/bin/env python3
"""Clarify that Table S7 excludes float conversion from timed execution."""

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle59"
DEST = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle60"

MAIN_SRC = SRC / "fastPLS_CMPB_main_cycle59_0.99.6_20260726.docx"
SUPP_SRC = SRC / "fastPLS_CMPB_supplement_cycle59_0.99.6_20260726.docx"
MAIN_OUT = DEST / "fastPLS_CMPB_main_cycle60_0.99.6_20260726.docx"
SUPP_OUT = DEST / "fastPLS_CMPB_supplement_cycle60_0.99.6_20260726.docx"


def replace_text(doc, old, new):
    for paragraph in doc.paragraphs:
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)
            return True
    return False


def revise_main():
    doc = Document(MAIN_SRC)
    old = (
        "Single-precision inputs provide a separate memory-oriented capability. "
        "Matched float32 storage reduced the input-matrix footprint by approximately "
        "one half in the validated routes, but did not provide a universal runtime or "
        "peak-RSS advantage because conversions, numerical fallbacks, and backend "
        "residency differ by model family. Precision-specific agreement and the "
        "validated CPU, CUDA, and Metal combinations are therefore reported in the "
        "Supplementary Material rather than pooled with the double-precision software "
        "comparison."
    )
    new = (
        "Single-precision inputs provide a separate memory-oriented capability. "
        "Matched float32 storage reduced the input-matrix footprint by approximately "
        "one half in the validated routes, but did not provide a universal runtime or "
        "peak-RSS advantage because numerical kernels, fallbacks, and backend residency "
        "differ by model family. Input conversion was completed before timing and was "
        "not included in fitting-plus-prediction time. Precision-specific agreement and "
        "the validated CPU, CUDA, and Metal combinations are therefore reported in the "
        "Supplementary Material rather than pooled with the double-precision software "
        "comparison."
    )
    if not replace_text(doc, old, new):
        raise RuntimeError("Could not locate the main-text float32 paragraph")
    doc.save(MAIN_OUT)


def revise_supplement():
    doc = Document(SUPP_SRC)
    old_control = (
        "Float32 NMR control. The matched full-resolution benchmark used 1,200 "
        "training spectra, 321 test spectra, 13,000 predictors, 28,355 responses, "
        "100 components, and three isolated repetitions unless a first run exceeded "
        "the predeclared 1,200-s timeout."
    )
    new_control = (
        "Float32 NMR control. The full-resolution benchmark used 1,200 training "
        "spectra, 321 test spectra, 13,000 predictors, 28,355 responses, 100 "
        "components, and three isolated repetitions unless a first run exceeded the "
        "predeclared 1,200-s timeout. Float32 conversion and task serialization were "
        "completed before each isolated benchmark process; reported time contains only "
        "model fitting and prediction."
    )
    if not replace_text(doc, old_control, new_control):
        raise RuntimeError("Could not locate the supplement float32 NMR paragraph")

    old_intro = (
        "Table S7 reports medians and interquartile ranges (IQR) from three isolated "
        "repetitions for the completed NMR precision measurements. A timeout means "
        "that the first attempted repetition exceeded 1,200 s; subsequent repetitions "
        "were intentionally not started. RSS denotes peak host resident set size and "
        "GPU memory is the sampled process-level peak, including CUDA context and "
        "runtime state."
    )
    new_intro = (
        "Table S7 reports medians and interquartile ranges (IQR) from three isolated "
        "repetitions for the completed NMR precision measurements. Input conversion, "
        "dataset loading, and task serialization occurred before timing; total time is "
        "strictly fitting plus prediction. A timeout means that the first attempted "
        "repetition exceeded 1,200 s; subsequent repetitions were intentionally not "
        "started. RSS denotes peak host resident set size and GPU memory is the sampled "
        "process-level peak, including CUDA context and runtime state. The CUDA rows "
        "compare the package's current execution routes rather than precision alone: "
        "the float64 route uses the optimized native matrix-free implementation, "
        "whereas the float32 route explicitly forms the full cross-covariance on the "
        "host, offloads randomized range sampling, and completes QR and the reduced SVD "
        "on the host. The float32 slowdown therefore cannot be attributed to input "
        "conversion and should not be interpreted as an intrinsic disadvantage of "
        "single-precision GPU arithmetic."
    )
    if not replace_text(doc, old_intro, new_intro):
        raise RuntimeError("Could not locate the Table S7 introduction")

    old_caption = (
        "Table S7. Full NMR matched float64/float32 comparison at 100 components. "
        "Values are median (IQR); the two entries in the final column give float64 "
        "then float32 completed/attempted repetitions."
    )
    new_caption = (
        "Table S7. Full NMR comparison of the current float64 and float32 execution "
        "routes at 100 components. Times are fitting plus prediction and exclude input "
        "conversion, dataset loading, and task serialization. Values are median (IQR); "
        "the two entries in the final column give float64 then float32 "
        "completed/attempted repetitions."
    )
    if not replace_text(doc, old_caption, new_caption):
        raise RuntimeError("Could not locate the Table S7 caption")

    for table in doc.tables:
        if not table.rows or not table.rows[0].cells:
            continue
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if headers and headers[0] == "Method / backend" and "float64 time (s)" in headers:
            for cell in table.rows[0].cells:
                if cell.text.strip() == "float64 time (s)":
                    cell.text = "float64 fit + predict (s)"
                elif cell.text.strip() == "float32 time (s)":
                    cell.text = "float32 fit + predict (s)"
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(6.5)
            break
    else:
        raise RuntimeError("Could not locate the Table S7 data table")

    doc.save(SUPP_OUT)


def main():
    DEST.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    for pdf in SRC.glob("*.pdf"):
        copy2(pdf, DEST / pdf.name.replace("cycle59", "cycle60"))
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
