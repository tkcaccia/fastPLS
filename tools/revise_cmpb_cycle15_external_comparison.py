#!/usr/bin/env python3

from pathlib import Path
from shutil import copy2

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle14"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle15"
EVIDENCE_OUT = (
    ROOT / "benchmark_results" / "manuscript_revision_cycle15_20260725"
)
SOURCE_DATA = (
    ROOT
    / "benchmark_results"
    / "manuscript_multidataset_summary_20260725"
    / "source"
    / "external_float64_summary.csv"
)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle14_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle14_0.99.6_20260725.docx"
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle14_20260725.docx"
)

MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle15_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle15_0.99.6_20260725.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle15_20260725.docx"

DATASET_LABELS = {
    "ccle": "CCLE",
    "cifar100": "CIFAR-100",
    "gtex_v8": "GTEx v8",
    "metref": "MetRef",
    "retina": "Retina",
    "tabula": "Tabula Muris",
    "tcga_brca": "TCGA-BRCA",
    "tcga_hnsc_methylation": "TCGA-HNSC methyl.",
    "tcga_pan_cancer": "TCGA Pan-Cancer",
}


def set_cell_margins(cell, top=55, start=70, bottom=55, end=70):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table, font_size=5.8):
    table.style = "Table"
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER
                    if row_index == 0
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
                    run.bold = row_index == 0


def find_paragraph(document, start):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(start):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {start}")


def set_paragraph(document, start, text):
    paragraph = find_paragraph(document, start)
    paragraph.text = text
    return paragraph


def replace_sentence(paragraph, old, new):
    if old not in paragraph.text:
        raise RuntimeError(f"Text not found in paragraph: {old}")
    paragraph.text = paragraph.text.replace(old, new)


def find_comparison_table(document):
    expected = [
        "Dataset",
        "k",
        "fastPLS",
        "fastPLS ms",
        "pls ms",
        "Speedup",
        "fastPLS metric",
        "pls metric",
    ]
    for table in document.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if headers == expected:
            return table
    raise RuntimeError("External SIMPLS comparison table not found")


def fill_table(table, columns, rows, font_size=5.8):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for index, label in enumerate(columns):
        table.rows[0].cells[index].text = label
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    style_table(table, font_size=font_size)


def add_table(document, columns, rows, font_size=5.8):
    table = document.add_table(rows=1, cols=len(columns))
    for index, label in enumerate(columns):
        table.rows[0].cells[index].text = label
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = str(value)
    style_table(table, font_size=font_size)
    return table


def build_comparisons():
    data = pd.read_csv(SOURCE_DATA)
    common = (
        data["input_precision"].eq("float64")
        & data["execution_precision"].eq("float64")
        & data["classifier"].eq("argmax")
        & data["algorithm"].eq("SIMPLS")
    )

    reference = data.loc[
        common & data["method_id"].eq("pls_simpls_fit")
    ].copy()
    irlba = data.loc[
        common & data["method_id"].eq("fastPLS_simpls_cpu_irlba")
    ].copy()

    keys = ["dataset", "ncomp_requested"]
    deterministic = irlba.merge(
        reference,
        on=keys,
        suffixes=("_fastpls", "_reference"),
        validate="one_to_one",
    )
    deterministic["fastpls_method"] = "fastPLS CPU IRLBA"
    deterministic["reference_method"] = "pls::simpls.fit"
    deterministic["speedup_vs_pls"] = (
        deterministic["median_time_ms_reference"]
        / deterministic["median_time_ms_fastpls"]
    )
    deterministic["metric_difference"] = (
        deterministic["median_metric_fastpls"]
        - deterministic["median_metric_reference"]
    )
    keep = [
        "dataset",
        "ncomp_requested",
        "fastpls_method",
        "reference_method",
        "reps_ok_fastpls",
        "reps_ok_reference",
        "median_time_ms_fastpls",
        "iqr_time_ms_fastpls",
        "median_time_ms_reference",
        "iqr_time_ms_reference",
        "speedup_vs_pls",
        "median_metric_fastpls",
        "iqr_metric_fastpls",
        "median_metric_reference",
        "iqr_metric_reference",
        "metric_difference",
    ]
    deterministic = deterministic[keep].sort_values("dataset")

    rsvd = data.loc[
        common
        & data["method_id"].str.match(
            r"fastPLS_simpls_(cpu|cuda)_rsvd$", na=False
        )
    ].copy()
    rsvd = (
        rsvd.sort_values(["dataset", "median_time_ms"])
        .groupby(keys, as_index=False)
        .first()
    )
    approximate = rsvd.merge(
        reference,
        on=keys,
        suffixes=("_fastpls", "_reference"),
        validate="one_to_one",
    )
    approximate["fastpls_method"] = approximate["method_id_fastpls"].map(
        {
            "fastPLS_simpls_cpu_rsvd": "fastPLS CPU rSVD",
            "fastPLS_simpls_cuda_rsvd": "fastPLS CUDA rSVD",
        }
    )
    approximate["reference_method"] = "pls::simpls.fit"
    approximate["speedup_vs_pls"] = (
        approximate["median_time_ms_reference"]
        / approximate["median_time_ms_fastpls"]
    )
    approximate["metric_difference"] = (
        approximate["median_metric_fastpls"]
        - approximate["median_metric_reference"]
    )
    approximate = approximate[keep].sort_values("dataset")

    EVIDENCE_OUT.mkdir(parents=True, exist_ok=True)
    deterministic.to_csv(
        EVIDENCE_OUT / "external_simpls_irlba_estimator_matched.csv",
        index=False,
    )
    approximate.to_csv(
        EVIDENCE_OUT / "external_simpls_rsvd_approximate_workflow.csv",
        index=False,
    )
    return deterministic, approximate


def table_rows(data):
    rows = []
    for _, row in data.iterrows():
        rows.append(
            (
                DATASET_LABELS[row["dataset"]],
                int(row["ncomp_requested"]),
                row["fastpls_method"],
                (
                    f'{row["median_time_ms_fastpls"]:.0f} '
                    f'({row["iqr_time_ms_fastpls"]:.1f})'
                ),
                (
                    f'{row["median_time_ms_reference"]:.0f} '
                    f'({row["iqr_time_ms_reference"]:.1f})'
                ),
                f'{row["speedup_vs_pls"]:.2f}',
                f'{row["median_metric_fastpls"]:.4f}',
                f'{row["median_metric_reference"]:.4f}',
            )
        )
    return rows


TABLE_COLUMNS = [
    "Dataset",
    "k",
    "fastPLS",
    "fastPLS ms",
    "pls ms",
    "Speedup",
    "fastPLS metric",
    "pls metric",
]


def revise_main(deterministic):
    document = Document(MAIN_SOURCE)

    set_paragraph(
        document,
        "External comparisons use",
        "External comparisons use independent R implementations where model and "
        "response type can be matched, including functions from pls, mdatools, "
        "plsdepot, pcv, plsgenomics, chemometrics, mixOmics, spls, and ropls. "
        "The primary estimator-matched SIMPLS comparison uses deterministic "
        "fastPLS CPU IRLBA and pls::simpls.fit with float64 inputs, identical "
        "preprocessing, response representation, component count, argmax "
        "decoding, thread setting, and timed work. Stochastic rSVD comparisons "
        "are reported separately as approximate workflow comparisons and are "
        "not interpreted as estimator-only speed tests. Package restrictions, "
        "timeouts, memory kills, and unsupported responses remain explicit.",
    )
    set_paragraph(
        document,
        "The precision-matched float64 SIMPLS comparison",
        "The primary float64 estimator-matched SIMPLS comparison is summarized "
        "in Table 3. Deterministic fastPLS CPU IRLBA and pls::simpls.fit had "
        "identical median test accuracy on all nine datasets. fastPLS was faster "
        "on seven datasets, including 8.90-fold on Tabula Muris, 8.65-fold on "
        "Retina, and 4.23-fold on CIFAR-100. pls::simpls.fit was faster on the "
        "two smallest absolute-time comparisons, TCGA-BRCA and TCGA-HNSC "
        "methylation, by 6 and 10 ms, respectively. The earlier rSVD table has "
        "been moved to the Supplement and is explicitly labelled as an "
        "approximate workflow comparison; its four-percentage-point MetRef "
        "difference is not presented as estimator agreement.",
    )
    set_paragraph(
        document,
        "Table 3.",
        "Table 3. Estimator-matched float64 SIMPLS comparison. fastPLS uses "
        "deterministic CPU IRLBA and the reference uses pls::simpls.fit, with "
        "identical split, preprocessing, component count, dummy response, and "
        "argmax decoder. Runtime is median total fitting plus prediction in "
        "milliseconds (IQR) over three isolated runs. rSVD is excluded.",
    )
    fill_table(
        find_comparison_table(document),
        TABLE_COLUMNS,
        table_rows(deterministic),
    )

    results = find_paragraph(document, "Results:")
    results.text = (
        results.text
        + " In the external estimator-matched comparison, float64 SIMPLS using "
        "deterministic CPU IRLBA was faster than pls::simpls.fit on seven of "
        "nine datasets, with identical median accuracy on all nine."
    )
    document.core_properties.title = (
        "fastPLS CMPB manuscript - deterministic external comparison cycle 15"
    )
    document.save(MAIN_OUT)


def revise_supplement(deterministic, approximate):
    document = Document(SUPP_SOURCE)
    set_paragraph(
        document,
        "The primary software comparison used",
        "The primary software comparison uses float64 inputs and a deterministic "
        "SIMPLS direction solver in fastPLS. CPU IRLBA and pls::simpls.fit use "
        "the same preprocessing, response representation, component count, "
        "argmax decoder, thread setting, and timed fitting-plus-prediction work. "
        "Approximate rSVD workflows are reported separately in Section S24 and "
        "are not treated as estimator-matched evidence. The complete archive "
        "retains every package restriction, failure, and memory measurement.",
    )
    set_paragraph(
        document,
        "Table S19.",
        "Table S19. Estimator-matched deterministic fastPLS CPU IRLBA versus "
        "pls::simpls.fit. Times are median (IQR) milliseconds from three isolated "
        "float64 runs; predictive values are median outer-test accuracy.",
    )
    fill_table(
        find_comparison_table(document),
        TABLE_COLUMNS,
        table_rows(deterministic),
    )

    document.add_page_break()
    document.add_heading(
        "S24. Approximate rSVD workflow comparison", level=1
    )
    document.add_paragraph(
        "This secondary analysis compares the fastest completed fastPLS CPU or "
        "CUDA rSVD SIMPLS workflow with deterministic pls::simpls.fit. It is not "
        "an estimator-matched comparison because rSVD is stochastic and "
        "approximate. Metric differences therefore combine implementation, "
        "randomized approximation, and backend effects. In particular, the "
        "MetRef rSVD median accuracy was 0.810 versus 0.770 for the deterministic "
        "reference and must not be interpreted as numerical equivalence."
    )
    document.add_paragraph(
        "Table S28. Approximate float64 rSVD workflow comparison. The fastest "
        "completed CPU/CUDA rSVD row is shown for each dataset; times are median "
        "(IQR) milliseconds over three isolated runs.",
        style="Caption",
    )
    add_table(
        document,
        TABLE_COLUMNS,
        table_rows(approximate),
    )
    document.core_properties.title = (
        "fastPLS CMPB supplement - deterministic external comparison cycle 15"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    set_paragraph(
        document,
        "A float64, estimator-matched SIMPLS table",
        "We agree that the previous table was incorrectly labelled because most "
        "fastPLS rows used stochastic rSVD, whereas pls::simpls.fit is "
        "deterministic. Table 3 and Supplementary Table S19 now use fastPLS CPU "
        "IRLBA as the deterministic estimator-matched route. Both methods use "
        "float64 inputs, identical preprocessing, split, component count, dummy "
        "response, argmax decoder, and three isolated runs. Median accuracy is "
        "identical on all nine datasets; fastPLS is faster on seven of nine. The "
        "former rSVD table is retained only as Supplementary Table S28, titled "
        "and discussed as an approximate workflow comparison. The four-point "
        "MetRef difference is no longer presented as estimator agreement.",
    )
    document.add_heading(
        "16. External comparison was not estimator matched", level=1
    )
    document.add_paragraph(
        "Reviewer comment: Table 3 compared approximate fastPLS rSVD-SIMPLS "
        "with deterministic pls::simpls.fit. The primary comparison should use "
        "deterministic IRLBA, and rSVD should be labelled approximate."
    )
    document.add_paragraph(
        "Response: Corrected. The primary table now uses deterministic CPU IRLBA "
        "and shows identical median accuracy on all nine datasets. The rSVD "
        "results were moved to a separate supplementary workflow table and are "
        "not used for estimator-equivalence or implementation-only claims."
    )
    document.core_properties.title = (
        "fastPLS CMPB response - deterministic external comparison cycle 15"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    deterministic, approximate = build_comparisons()
    revise_main(deterministic)
    revise_supplement(deterministic, approximate)
    revise_response()
    copy2(
        EVIDENCE_OUT / "external_simpls_irlba_estimator_matched.csv",
        OUT / "external_simpls_irlba_estimator_matched.csv",
    )
    copy2(
        EVIDENCE_OUT / "external_simpls_rsvd_approximate_workflow.csv",
        OUT / "external_simpls_rsvd_approximate_workflow.csv",
    )
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
