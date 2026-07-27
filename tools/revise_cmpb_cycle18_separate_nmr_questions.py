#!/usr/bin/env python3

from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle17"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle18"
EVIDENCE = (
    ROOT / "benchmark_results" / "manuscript_revision_cycle18_20260725"
)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle17_0.99.6_20260725.docx"
SUPP_SOURCE = (
    SOURCE / "fastPLS_CMPB_supplement_cycle17_0.99.6_20260725.docx"
)
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle17_20260725.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle18_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle18_0.99.6_20260725.docx"
RESPONSE_OUT = (
    OUT / "fastPLS_CMPB_response_to_reviewers_cycle18_20260725.docx"
)

PLSSVD_ROWS = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle16_20260725"
    / "nmr_plssvd_extended_lower_grid"
    / "heldout"
    / "rows"
)
SIMPLS_ROWS = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle17_20260725"
    / "nmr_simpls_one_se"
    / "heldout"
    / "rows"
)

spec = spec_from_file_location(
    "cycle17_helpers", ROOT / "tools" / "revise_cmpb_cycle17_nmr_one_se.py"
)
c17 = module_from_spec(spec)
spec.loader.exec_module(c17)
c16 = c17.c16


def read_rows(directory):
    files = sorted(directory.glob("*.csv"))
    if not files:
        raise RuntimeError(f"No result rows found in {directory}")
    return pd.concat(
        [pd.read_csv(path) for path in files],
        ignore_index=True,
    )


def summarize_selected(rows, family, selection_basis):
    rows = rows.loc[rows["status"].eq("ok")].copy()
    if rows.empty:
        raise RuntimeError(f"No completed selected-setting rows for {family}")
    return {
        "family": family,
        "backend": "CUDA rSVD",
        "precision": str(rows["precision"].iloc[0]),
        "selected_ncomp": int(rows["ncomp"].iloc[0]),
        "selection_basis": selection_basis,
        "RMSD": float(rows["RMSD"].iloc[0]),
        "Q2": float(rows["Q2"].iloc[0]),
        "total_time_median_sec": float(rows["total_time_sec"].median()),
        "total_time_iqr_sec": float(
            rows["total_time_sec"].quantile(0.75)
            - rows["total_time_sec"].quantile(0.25)
        ),
        "peak_host_rss_median_mb": float(rows["host_rss_mb"].median()),
        "peak_gpu_memory_median_mb": float(rows["gpu_peak_mb"].median()),
        "replicates": int(len(rows)),
        "status": "ok",
    }


def build_evidence():
    EVIDENCE.mkdir(parents=True, exist_ok=True)
    selected = pd.DataFrame(
        [
            summarize_selected(
                read_rows(PLSSVD_ROWS),
                "PLS-SVD",
                "unique one-SE candidate",
            ),
            summarize_selected(
                read_rows(SIMPLS_ROWS),
                "SIMPLS",
                "smallest one-SE candidate; eligible 50, 75, 100",
            ),
        ]
    )
    selected.to_csv(
        EVIDENCE / "nmr_family_selected_predictive_benchmark.csv",
        index=False,
    )
    return selected


def insert_heading_before(paragraph, text, style="Heading 3"):
    heading = paragraph.insert_paragraph_before(text)
    heading.style = style
    return heading


def selected_rows_for_table(selected):
    rows = []
    for _, row in selected.iterrows():
        rows.append(
            (
                row["family"],
                int(row["selected_ncomp"]),
                row["selection_basis"],
                f'{row["RMSD"]:.7f}',
                f'{row["Q2"]:.5f}',
                (
                    f'{row["total_time_median_sec"]:.3f} '
                    f'({row["total_time_iqr_sec"]:.3f})'
                ),
                (
                    f'{row["peak_host_rss_median_mb"]:.0f} / '
                    f'{row["peak_gpu_memory_median_mb"]:.0f}'
                ),
                int(row["replicates"]),
            )
        )
    return rows


def revise_main(selected):
    document = Document(MAIN_SOURCE)

    table_caption = c16.find_paragraph(document, "Table 1.")
    if "different component counts" not in table_caption.text:
        table_caption.text += (
            " Because component counts are family-specific, these runtimes "
            "describe selected predictive workflows and are not used as "
            "fixed-complexity implementation speed comparisons."
        )

    figure_caption = c16.find_paragraph(document, "Figure 2.")
    if "family-specific predictive settings" not in figure_caption.text:
        figure_caption.text += (
            " This is the family-specific predictive-setting analysis; the "
            "separate NMR fixed-complexity implementation analysis is shown in "
            "Figure 3."
        )

    selection = c16.find_paragraph(document, "NMR represented the extreme")
    insert_heading_before(
        selection,
        "3.2.1 Family-specific predictive model selection",
    )
    selection.text = (
        "NMR represented the extreme multivariate-response setting (1,200 "
        "training and 321 held-out spectra; p=13,000; q=28,355). Every method "
        "used the predefined outer split and full response spectrum; predictors "
        "were centred without variance scaling, and the routine 4.6-4.8 ppm "
        "residual-water interval was set to zero in both training and held-out "
        "predictor matrices. Five repeated training-only 80/20 splits were used "
        "for component selection. PLS-SVD selected 5 components as its unique "
        "one-standard-error candidate. SIMPLS split-specific minima ranged from "
        "25 to 100 components; 50, 75, and 100 were one-standard-error eligible, "
        "so the parsimonious rule selected 50. At these family-specific settings, "
        "CUDA PLS-SVD gave held-out RMSD 0.0010426, Q2 0.98916, and median total "
        "time 0.898 s, whereas CUDA SIMPLS gave RMSD 0.0007591, Q2 0.99425, and "
        "median total time 1.971 s. These values answer the predictive "
        "model-selection question; they do not isolate implementation speed "
        "because model complexity differs. Table 1 and Figure 2 report the same "
        "selected-setting principle across datasets, and Supplementary Table "
        "S12c and Figure S18 provide the NMR details."
    )

    implementation = c16.find_paragraph(
        document, "The predefined outer split was retained"
    )
    insert_heading_before(
        implementation,
        "3.2.2 Fixed-complexity implementation benchmark",
    )
    implementation.text = (
        "A separate benchmark fixed every implementation at 100 components to "
        "isolate execution time, memory use, and numerical agreement at equal "
        "model size. This operating point was not selected independently for "
        "PLS-SVD or SIMPLS and is not interpreted as the preferred predictive "
        "model. The predefined outer split was retained to reproduce the "
        "deposited analysis, and all implementations received identical "
        "preprocessed matrices and prediction targets. Figure 3 therefore "
        "addresses an implementation question only: panels A-C verify "
        "prediction behavior at the common component count, and panel D compares "
        "computational resources. Predictive conclusions remain based on the "
        "family-specific settings reported above."
    )
    implementation.add_run().add_break(WD_BREAK.PAGE)

    c16.set_paragraph(
        document,
        "Figure 3.",
        "Figure 3. Fixed-complexity NMR implementation benchmark at 100 "
        "components. This common component count was imposed to compare "
        "implementations and was not the family-specific predictive selection. "
        "(A) Observed and SIMPLS-rSVD-predicted held-out spectrum across the "
        "full chemical-shift range. (B) Enlarged 0.5-1.7 ppm region. (C) "
        "Distribution of per-spectrum RMSD across the 321 held-out spectra for "
        "the deposited PLS-SVD/IRLBA reference and fastPLS PLS-SVD and SIMPLS "
        "CPU/CUDA rSVD implementations. (D) Median total "
        "fitting-plus-prediction time, peak host RSS, and sampled peak GPU "
        "memory across three isolated runs. The representative spectrum was "
        "selected mechanically as the test spectrum whose RMSD was closest to "
        "the held-out median. All comparisons used float64 data and identical "
        "routinely preprocessed training and held-out spectra.",
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - separated NMR analyses cycle 18"
    )
    document.save(MAIN_OUT)


def revise_supplement(selected):
    document = Document(SUPP_SOURCE)

    c16.set_paragraph(
        document,
        "Routine NMR spectral preprocessing and component selection.",
        "Routine NMR spectral preprocessing and component selection. The task "
        "comprised 1,200 training and 321 held-out spectra, 13,000 predictors, "
        "and 28,355 numeric responses. The 4.6-4.8 ppm residual-water interval "
        "was set to zero in Xtrain and Xtest; Ytrain and Ytest were unchanged. "
        "Predictors were centred without variance scaling. Five repeated 80/20 "
        "training-only splits used seeds 123, 456, 789, 1011, and 2027. Mean "
        "RMSD and SE were calculated across paired splits, and the smallest "
        "component count within one SE of the minimum mean was selected. Table "
        "S12b reports the selection curves; Table S12c reports outer-test "
        "performance only at the family-specific selected settings.",
    )

    fixed_benchmark = c16.find_paragraph(document, "Held-out NMR results.")
    caption = fixed_benchmark.insert_paragraph_before(
        "Table S12c. NMR predictive benchmark at the family-specific "
        "training-selected component count. Values were obtained on the "
        "untouched outer test set using float64 CUDA rSVD; time and memory are "
        "medians across three isolated repetitions, with time IQR in "
        "parentheses. These rows answer predictive model selection and are not "
        "a fixed-complexity implementation comparison."
    )
    caption.style = "Caption"
    caption.paragraph_format.page_break_before = True
    table = c16.add_table(
        document,
        [
            "Family",
            "Selected k",
            "Selection basis",
            "RMSD",
            "Q2",
            "Total s median (IQR)",
            "Host / GPU MB",
            "n",
        ],
        selected_rows_for_table(selected),
        font_size=5.1,
    )
    c16.prevent_row_splitting(table)
    fixed_benchmark._p.addprevious(table._tbl)

    fixed_benchmark.text = (
        "Fixed-complexity NMR implementation benchmark. A distinct analysis "
        "held the component count at 100 for the deposited fastsimpls "
        "PLS-SVD/IRLBA function and fastPLS PLS-SVD and SIMPLS CPU/CUDA rSVD "
        "implementations. The same routinely preprocessed training and held-out "
        "spectra were used, with three isolated runs per implementation. This "
        "analysis evaluates runtime, host and GPU memory, and numerical "
        "agreement at equal model size; it does not identify the preferred "
        "predictive component count. The deposited reference required median "
        "total time 431.23 s (IQR 2.71 s) and peak host RSS 6,101 MB (IQR "
        "826 MB), with RMSD 0.0007194 and Q2 0.99484. Full implementation "
        "results are reported in Table S6 and Figures S12-S14."
    )
    table_s6_caption = c16.set_paragraph(
        document,
        "Table S6.",
        "Table S6. Fixed-complexity float64 NMR implementation comparison with "
        "the deposited fastsimpls PLS-SVD/IRLBA reference at the imposed "
        "100-component operating point. Time and memory are medians from three "
        "isolated runs; prediction agreement is the Pearson correlation between "
        "vectorized held-out predictions and the deposited reference "
        "prediction. This table is not a component-selection analysis.",
    )
    table_s6_caption.paragraph_format.page_break_before = True
    c16.set_paragraph(
        document,
        "Figure S12.",
        "Figure S12. Observed and predicted full held-out NMR spectrum at the "
        "common 100-component fixed-complexity implementation point. The "
        "spectrum was selected mechanically by median per-spectrum RMSD; the "
        "component count was imposed for implementation comparison rather than "
        "selected as the preferred predictive setting.",
    )
    c16.set_paragraph(
        document,
        "Figure S14.",
        "Figure S14. Distribution of held-out per-spectrum RMSD for the matched "
        "CPU and CUDA SIMPLS-rSVD fits at the common 100-component "
        "fixed-complexity implementation point.",
    )
    c16.set_paragraph(
        document,
        "The outer split was intentionally fixed",
        "The outer split was intentionally fixed to reproduce the deposited "
        "reference implementation comparison. For all five methods, the input "
        "checksum, preprocessing rule, centred predictor matrices, full "
        "multivariate target, outer training/test allocation, and imposed "
        "100-component model size were identical. This fixed-complexity "
        "analysis is separate from the family-specific predictive benchmark in "
        "Table S12c. Error distributions were retained at two levels: 321 "
        "per-spectrum RMSD values and 28,355 response-wise RMSD values.",
    )

    document.core_properties.title = (
        "fastPLS CMPB supplement - separated NMR analyses cycle 18"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    document.add_heading(
        "19. The NMR comparisons mixed two scientific questions",
        level=1,
    )
    document.add_paragraph(
        "Reviewer comment: Table 1 uses family-specific selected component "
        "counts, whereas Figure 3 compares every implementation at 100 "
        "components. The implementation benchmark and predictive "
        "model-selection benchmark should be presented separately."
    )
    document.add_paragraph(
        "Response: Corrected. Section 3.2 is now divided into a family-specific "
        "predictive model-selection analysis and a fixed-complexity "
        "implementation analysis. The predictive analysis reports PLS-SVD at "
        "its selected 5 components and SIMPLS at its one-SE-selected 50 "
        "components, including outer-test RMSD, Q2, and total time; these "
        "results are shown in Table 1, Figure 2, and new Supplementary Table "
        "S12c. Figure 3, Table S6, and Figures S12-S14 are now labelled "
        "exclusively as fixed-100-component implementation comparisons of "
        "runtime, memory, and numerical agreement. Their captions state that "
        "100 components was imposed for equal model size and is not a "
        "family-specific predictive selection."
    )
    document.core_properties.title = (
        "fastPLS CMPB response - separated NMR analyses cycle 18"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    selected = build_evidence()
    revise_main(selected)
    revise_supplement(selected)
    revise_response()
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
