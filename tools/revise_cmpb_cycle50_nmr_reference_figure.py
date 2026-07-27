#!/usr/bin/env python3

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle49"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle50"
EVIDENCE = ROOT / "benchmark_results" / "nmr_reference_metrics_20260726"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle49_0.99.6_20260726.docx"
SUPP_SOURCE = (
    SOURCE / "fastPLS_CMPB_supplement_cycle49_0.99.6_20260726.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle50_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle50_0.99.6_20260726.docx"

PANEL_C = EVIDENCE / "nmr_reference_predictive.png"
PANEL_D = EVIDENCE / "nmr_reference_resources.png"
PANEL_A = (
    ROOT
    / "benchmark_results"
    / "review_nmr_20260724"
    / "plots"
    / "nmr_spectrum_full.png"
)
PANEL_B = (
    ROOT
    / "benchmark_results"
    / "review_nmr_20260724"
    / "plots"
    / "nmr_spectrum_zoom.png"
)


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_paragraph(document, prefix, text):
    paragraph = find_paragraph(document, prefix)
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def set_panel(cell, label, image, width=3.05):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = paragraph.add_run(f"{label}\n")
    label_run.bold = True
    label_run.font.size = Pt(10)
    paragraph.add_run().add_picture(str(image), width=Inches(width))


def revise_main():
    document = Document(MAIN_SOURCE)

    replace_paragraph(
        document,
        "A separate analysis fixed fastPLS PLS-SVD and SIMPLS at 100 components",
        (
            "A separate analysis fixed the deposited Nature Communications "
            "PLS-SVD workflow and the four fastPLS PLS-SVD and SIMPLS workflows "
            "at 100 components, float64 precision, the predefined split, identical "
            "preprocessing, and the same multivariate prediction target (Figure 3). "
            "The deposited fastsimpls PLS-SVD/IRLBA CPU workflow required 431.234 s "
            "and 6,101 MB peak host RSS, with RMSD 0.0007194 and Q2 0.99484. "
            "fastPLS PLS-SVD/rSVD required 16.323 s on CPU and 1.115 s on CUDA, "
            "with RMSD 0.0007292 and 0.0007183, respectively; prediction "
            "correlations with the deposited result were 0.999942 and 0.999974. "
            "fastPLS SIMPLS/rSVD required 20.140 s on CPU and 3.055 s on CUDA, "
            "with RMSD 0.0008606 and 0.0008047 and prediction correlations 0.998220 "
            "and 0.998667. Host RSS was 2,964-3,469 MB for the fastPLS workflows; "
            "sampled process-level GPU peaks were 664 MB for CUDA PLS-SVD and "
            "3,432 MB for CUDA SIMPLS. Panels C-D show these five workflows "
            "directly. Because the deposited workflow uses IRLBA whereas the "
            "displayed fastPLS workflows use rSVD, its speed difference is a "
            "composite historical workflow comparison rather than an "
            "implementation-only effect. The matched backend and family contrasts "
            "remain in Supplementary Table S6a and Figures S14-S15."
        ),
    )

    figure_table = document.tables[0]
    if len(figure_table.rows) != 2 or len(figure_table.columns) != 2:
        raise RuntimeError("The expected 2 x 2 NMR figure table was not found")
    set_panel(figure_table.cell(0, 0), "A", PANEL_A)
    set_panel(figure_table.cell(0, 1), "B", PANEL_B)
    set_panel(figure_table.cell(1, 0), "C", PANEL_C)
    set_panel(figure_table.cell(1, 1), "D", PANEL_D)
    page_break = document.add_page_break()
    figure_table._tbl.addprevious(page_break._p)

    replace_paragraph(
        document,
        "Figure 3. Fixed-complexity NMR analysis at 100 components",
        (
            "Figure 3. NMR benchmark at 100 components. (A, B) Observed and CUDA "
            "SIMPLS/rSVD-predicted spectrum, globally and at 0.5-1.7 ppm. "
            "(C) RMSD, Q2, and prediction correlation relative to the deposited "
            "Nature Communications PLS-SVD/IRLBA CPU reference. (D) Total time, "
            "host RSS, and GPU memory. All workflows used identical data "
            "processing and float64 inputs; points and bars are medians and "
            "interquartile ranges from three runs."
        ),
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - NMR reference metric comparison"
    )
    document.save(MAIN_OUT)


def main():
    for path in (
        MAIN_SOURCE,
        SUPP_SOURCE,
        PANEL_A,
        PANEL_B,
        PANEL_C,
        PANEL_D,
    ):
        if not path.exists():
            raise FileNotFoundError(path)
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    shutil.copy2(SUPP_SOURCE, SUPP_OUT)
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
