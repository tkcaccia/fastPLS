#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
OUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle37_review"
OUT_FILE = (
    OUT_DIR / "fastPLS_CMPB_fresh_reviewer_comments_cycle37_20260726.docx"
)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(0x20, 0x29, 0x31)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    title = document.styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(21)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x12, 0x36, 0x55)
    title.paragraph_format.space_after = Pt(5)

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle.font.size = Pt(11.5)
    subtitle.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    subtitle.paragraph_format.space_after = Pt(14)

    for name, size in (("Heading 1", 15), ("Heading 2", 12.5)):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1F, 0x5E, 0x8C)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    for name in ("List Number", "List Bullet"):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.28)
        style.paragraph_format.first_line_indent = Inches(-0.20)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.08


def add_numbered_comment(document, title, body):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.keep_together = True
    title_run = paragraph.add_run(f"{title}. ")
    title_run.bold = True
    paragraph.add_run(body)


def add_minor_comment(document, body):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.add_run(body)


def build_report():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_styles(document)

    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    header = section.header.paragraphs[0]
    header.text = (
        "Fresh reviewer report | Computer Methods and Programs in Biomedicine"
    )
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    add_page_number(section.footer.paragraphs[0])

    document.add_paragraph("Reviewer report", style="Title")
    document.add_paragraph(
        "Manuscript: fastPLS: scalable partial least squares with compiled CPU "
        "and accelerator backends for high-dimensional biomedical data",
        style="Subtitle",
    )

    metadata = document.add_table(rows=3, cols=2)
    metadata.style = "Table Grid"
    metadata.autofit = False
    values = (
        ("Journal", "Computer Methods and Programs in Biomedicine"),
        ("Recommendation", "Major revision"),
        (
            "Review basis",
            "Fresh reading of the current main manuscript and Supplementary Material",
        ),
    )
    for row, (label, value) in zip(metadata.rows, values):
        row.cells[0].width = Inches(1.55)
        row.cells[1].width = Inches(5.05)
        shade_cell(row.cells[0], "E7EEF5")
        for cell in row.cells:
            set_cell_margins(cell)
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].paragraphs[0].add_run(value)

    document.add_heading("Overall assessment", level=1)
    document.add_paragraph(
        "This is an ambitious and potentially useful software paper. The authors "
        "provide a broad PLS interface, formal numerical checks for deterministic "
        "SIMPLS, explicit CPU/CUDA/Metal capability boundaries, a scientifically "
        "relevant NMR case study, and unusually transparent negative results for "
        "float32 and approximate rSVD routes. The distinction between deterministic "
        "estimator preservation and randomized approximation is a particular "
        "strength."
    )
    document.add_paragraph(
        "Nevertheless, several central claims remain insufficiently isolated. The "
        "headline randomized benchmark does not expose the rSVD controls or audit "
        "status, the NMR speed comparison changes both implementation and solver, "
        "and the claimed SIMPLS computational advance is not tested directly "
        "against PLS-SVD across matrix shapes. GPU memory accounting also includes "
        "context initialization while being interpreted as workspace memory. These "
        "issues affect the interpretation of the principal results and require "
        "major revision rather than editorial correction."
    )

    document.add_heading("Major comments", level=1)
    major_comments = [
        (
            "The central SIMPLS contribution needs a direct quantitative test",
            "The deterministic validation convincingly shows that the optimized "
            "IRLBA route can reproduce de Jong SIMPLS within stated tolerances. It "
            "does not, however, establish the central performance proposition: "
            "when and why the reorganized SIMPLS path approaches or exceeds the "
            "speed of PLS-SVD. Add a matched experiment using identical float64 "
            "data, split, preprocessing, component count, prediction rule, solver "
            "controls, and hardware. It should span representative n, p, q, and A "
            "regimes and report fitting time, prediction time, incremental memory, "
            "and prediction agreement. A shape-regime summary would make the "
            "methodological contribution much clearer than the current collection "
            "of family-specific selected points.",
        ),
        (
            "The rSVD settings behind the headline benchmark are not auditable",
            "The manuscript shows that oversampling 10 with one power iteration "
            "failed 16 of 117 CPU checks, while the focused CUDA audit required "
            "four power iterations or oversampling 20 to pass all eight difficult "
            "points. Table 1 nevertheless reports CPU and CUDA rSVD rows simply as "
            "OK and does not state oversampling, power iterations, seed, or "
            "approximation-audit status. The exact settings used for every headline "
            "row must be shown. Rows produced with a rejected or unaudited setting "
            "should be rerun with an accepted configuration or explicitly marked "
            "exploratory/unaudited. The abstract must also distinguish CPU from "
            "CUDA reliability evidence; the present 117-of-117 statement can be "
            "read as applying to both.",
        ),
        (
            "The NMR implementation comparison confounds solver and implementation",
            "The deposited reference uses PLS-SVD with IRLBA, whereas Figure 3 and "
            "Table S6 compare it with fastPLS PLS-SVD and SIMPLS using rSVD. The "
            "large speed difference therefore combines implementation, solver, PLS "
            "family, and hardware effects. Add fastPLS PLS-SVD with deterministic "
            "IRLBA at the same 100-component operating point on CPU. If available, "
            "also report SIMPLS with IRLBA. Use these matched rows to quantify the "
            "implementation gain, and retain rSVD/CUDA rows as separate workflow "
            "acceleration results. The family-specific selected-component analysis "
            "should remain separate, as it is now.",
        ),
        (
            "CPU/CUDA rSVD rows should not be called estimator matched without an audit",
            "The same PLS family, split, and component count define a matched "
            "workflow, but randomized range finding and backend-specific numerical "
            "paths need not produce the same fitted estimator. This is visible in "
            "several non-identical CPU/CUDA metrics. Replace estimator-matched with "
            "workflow-matched unless predictions, subspaces, and metrics satisfy the "
            "prespecified agreement criteria for that pair. Include the executed "
            "solver configuration and agreement status in the machine-readable and "
            "human-readable summaries.",
        ),
        (
            "GPU memory is not baseline corrected in the claimed sense",
            "All reported pre-fit GPU baselines are zero because the CUDA context is "
            "created during fitting. Consequently, the reported increment includes "
            "context creation, libraries, allocator pools, and algorithm workspaces. "
            "It cannot be interpreted as isolated algorithmic or prediction "
            "workspace memory. Initialize and synchronize the CUDA context before "
            "recording the baseline, or relabel the current quantity as "
            "context-inclusive process GPU growth. Because 20 ms nvidia-smi polling "
            "may miss brief peaks, validate representative measurements with CUDA "
            "memory APIs or discuss the resulting sampling limitation.",
        ),
        (
            "Computational uncertainty is incomplete",
            "Several CIFAR-100 and PRISM cells in Table 1 have n=1 and an IQR of "
            "zero, despite runtimes that appear feasible for replication. These "
            "rows should be repeated at least three times under isolated-process "
            "conditions, preferably five times for the principal speed claims. "
            "ImageNet remains appropriately labelled exploratory, but the paper "
            "should distinguish repeated query timing, repeated randomized fitting, "
            "and repeated full end-to-end workflow timing. A single memory trace is "
            "not a dispersion estimate.",
        ),
        (
            "Component selection remains boundary dominated",
            "Many SIMPLS, OPLS, and linear kernel-PLS selections occur at the upper "
            "tested boundary, and several PLS-SVD values are response-rank limited. "
            "The dagger notation is helpful, but the text still uses "
            "training-set-selected values as if the grids had resolved the model "
            "complexity. Report the proportion of boundary selections, extend grids "
            "for the principal predictive comparisons where feasible, and use the "
            "phrase best value within the evaluated grid consistently. Predictive "
            "claims on the small test sets also need repeated outer splits or nested "
            "cross-validation; Wilson or bootstrap intervals conditional on one "
            "fixed split do not capture training-set uncertainty.",
        ),
        (
            "OPLS and nonlinear kernel PLS need independent numerical validation",
            "Formal estimator-preservation evidence is supplied only for SIMPLS. "
            "However, the package and conclusions also emphasize OPLS and kernel "
            "PLS. The main kernel rows use the linear shortcut and therefore largely "
            "duplicate SIMPLS, while nonlinear kernels are evaluated on only three "
            "datasets without an independent implementation check. On small fixed "
            "tasks, compare OPLS predictions, predictive/orthogonal scores, and "
            "filtering with an established OPLS implementation. Likewise validate "
            "linear and nonlinear kernel centering and predictions against an "
            "independent kernel-PLS implementation. Otherwise, narrow the general "
            "validation claim to SIMPLS.",
        ),
        (
            "Float32 and Metal should be positioned as conditional capabilities",
            "The negative float32 evidence is valuable, but it is too important to "
            "remain mainly in the Supplement. Extreme-response NMR shows severe "
            "runtime and accuracy failures, and Metal has only limited, mostly "
            "single-run validation on two classification shapes. Add a compact "
            "main-text capability table distinguishing validated, conditional, "
            "experimental, hybrid, and unsupported routes. The abstract should "
            "state that float32 and Metal support is conditional rather than listing "
            "them alongside the deterministic CPU reference without qualification.",
        ),
        (
            "The cross-validation acceleration claim is currently narrow",
            "The matched comparison against an R-level loop is appropriate, but it "
            "covers only MetRef, Retina, and PRISM and gives modest speedups of "
            "1.07-1.35 for deterministic IRLBA. If compiled validation is a central "
            "contribution, add at least one larger feasible classification task, "
            "report memory as well as time, and compare complete returned outputs. "
            "If this is not feasible, present compiled cross-validation as an "
            "implementation convenience with measured modest gains rather than a "
            "general acceleration result.",
        ),
        (
            "The main benchmark presentation is too dense and is mislabelled",
            "Table 1 is difficult to read at publication size because each cell "
            "contains two backends, uncertainty, time, run count, status, and two "
            "memory triplets. Move the detailed B/P/delta memory audit to the "
            "Supplement and retain a concise main table or paired figure with "
            "predictive metric, total time, and incremental memory. Also replace "
            "twelve-task biomedical benchmark with heterogeneous benchmark: "
            "CIFAR-100 is a computational image benchmark, not a biomedical task.",
        ),
        (
            "The reproducibility record contains multiple source identifiers",
            "The availability statement cites one review-cycle package commit, "
            "whereas the Supplementary reproducibility table identifies a different "
            "fastPLS source for precision and external-comparison runs. Per-analysis "
            "commits are acceptable, but the relationship must be explicit and the "
            "exact commit/configuration must be attached to every result table. "
            "Archive the submission code and machine-readable results with an "
            "immutable DOI, include the rSVD controls and backend build flags in the "
            "manifest, and provide one end-to-end public or synthetic workflow that "
            "reproduces a principal table without restricted data.",
        ),
    ]
    for title, body in major_comments:
        add_numbered_comment(document, title, body)

    document.add_heading("Minor comments", level=1)
    minor_comments = [
        "The abstract is dominated by numerical-validation details but does not "
        "summarize the principal scaling result, NMR result, or the conditional "
        "status of float32 and Metal. Rebalance it around objective, implementation, "
        "key computational evidence, and limitations.",
        "Explain how the 117 component-level comparisons are distributed across "
        "datasets, seeds, and component values. Passing a tolerance should be "
        "described as agreement within prespecified tolerances, not equivalence.",
        "Justify the rSVD failure thresholds scientifically. A 3.32% relative "
        "prediction error, 4.93-degree subspace angle, or 1% label disagreement may "
        "not be negligible in every biomedical application.",
        "Cite the external R packages used in the comparison or provide formal "
        "software references in the Supplement, not only package names and versions.",
        "For every regression dataset, state the response scale and units as clearly "
        "as done for CBMC CITE-seq. Cross-dataset RMSD values are not directly "
        "comparable.",
        "Figure 2 combines accuracy and RMSD tasks. Ensure that the panels, scales, "
        "and axis labels make these fundamentally different metrics impossible to "
        "confuse.",
        "In Figure 3, identify the selected spectrum in a reproducible table and "
        "show the observed and predicted curves with a legend that remains legible "
        "in grayscale.",
        "The ImageNet section is appropriately cautious. Keep it explicitly "
        "exploratory and avoid using the small top-5 difference as evidence for "
        "supervised improvement.",
        "Use consistent typographic treatment for mathematical symbols n, p, q, A, "
        "a, K, and r. Literal R arguments such as ncomp and north should be formatted "
        "as code.",
        "Clarify whether the public package version used for the paper is a stable "
        "release, a Bioconductor development version, or a review snapshot.",
    ]
    for comment in minor_comments:
        add_minor_comment(document, comment)

    document.add_heading("Recommendation", level=1)
    document.add_paragraph(
        "Major revision. The software framework is promising and the manuscript "
        "shows commendable transparency, but the main speed, memory, and rSVD "
        "claims require better isolation and more auditable reporting before the "
        "paper is suitable for publication."
    )

    document.core_properties.title = (
        "Fresh CMPB reviewer report for the fastPLS manuscript"
    )
    document.core_properties.subject = (
        "Independent methodological, computational, and presentation review"
    )
    document.core_properties.author = "Independent reviewer"
    document.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    print(build_report())
