#!/usr/bin/env python3

from pathlib import Path

from docx import Document


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle29"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle30"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle29_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle29_0.99.6_20260725.docx"
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle29_20260725.docx"
)

MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle30_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle30_0.99.6_20260725.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle30_20260725.docx"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def revise_main():
    document = Document(MAIN_SOURCE)
    abstract = find_paragraph(document, "Methods: fastPLS provides")
    old = "multithread-capable CPU execution"
    new = "compiled CPU execution linked to the system BLAS/LAPACK libraries"
    if old not in abstract.text:
        raise RuntimeError("Multithread claim not found in abstract")
    abstract.text = abstract.text.replace(old, new, 1)

    backend = find_paragraph(document, "The R package includes bundled IRLBA code")
    anchor = (
        "CUDA uses NVIDIA CUDA libraries and cuBLAS; Metal uses Apple's Metal "
        "Performance Shaders. "
    )
    addition = (
        "The CPU backend uses compiled C/C++ code and the BLAS/LAPACK libraries "
        "linked to R, which may be OpenBLAS on suitable installations. This is an "
        "implementation capability, not evidence of multicore acceleration. All "
        "primary software comparisons used one effective BLAS thread. Because a "
        "controlled thread-scaling experiment was not performed, no multicore "
        "speed-up is claimed. "
    )
    if anchor not in backend.text:
        raise RuntimeError("Backend insertion point not found")
    backend.text = backend.text.replace(anchor, anchor + addition, 1)
    document.core_properties.title = (
        "fastPLS CMPB manuscript - qualified CPU threading statement"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)
    paragraph = find_paragraph(
        document,
        "Reproducibility experiments use identical data",
    )
    old = (
        "Reproducibility experiments use identical data, preprocessing, folds, "
        "requested components, and seeds for single-thread CPU, multithread CPU, "
        "CUDA, and Metal. "
    )
    new = (
        "Reproducibility experiments use identical data, preprocessing, folds, "
        "requested components, and seeds for the configured CPU, CUDA, and Metal "
        "routes. The CPU route uses compiled code and the BLAS/LAPACK libraries "
        "linked to R and can use OpenBLAS where it is installed. That build "
        "capability is recorded separately from measured performance. The primary "
        "software-comparison runs used one effective BLAS thread. No controlled "
        "one-thread versus multiple-thread scaling experiment was conducted, and "
        "no multicore speed-up is inferred. "
    )
    if old not in paragraph.text:
        raise RuntimeError("Threading sentence not found in supplement")
    paragraph.text = paragraph.text.replace(old, new, 1)
    document.core_properties.title = (
        "fastPLS CMPB supplement - qualified CPU threading statement"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    document.add_heading(
        "30. CPU multithreading claim required scaling evidence",
        level=1,
    )
    document.add_paragraph(
        "Reviewer comment: The claim of “multithread-capable CPU execution” "
        "requires a thread-scaling experiment or should be softened."
    )
    document.add_paragraph(
        "Response: Corrected by softening the claim. The abstract now describes "
        "compiled CPU execution linked to the system BLAS/LAPACK libraries. The "
        "Methods and Supplement explain that these libraries may use one or more "
        "threads according to the platform build and runtime configuration and "
        "that active settings are recorded. Because we did not conduct a "
        "controlled one-thread versus multiple-thread scaling experiment, the "
        "manuscript no longer claims a CPU multithreading speedup."
    )
    document.core_properties.title = (
        "fastPLS CMPB response to reviewers - CPU threading qualification"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    revise_response()


if __name__ == "__main__":
    main()
