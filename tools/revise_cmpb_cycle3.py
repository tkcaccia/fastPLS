"""Create cycle-3 CMPB drafts after three isolated NMR repetitions."""

from pathlib import Path
from shutil import copy2

from docx import Document


ROOT = Path("/Users/stefano/Desktop/fastPLS/CMPB_rewrite_20260724_cycle2")
OUT = Path("/Users/stefano/Desktop/fastPLS/CMPB_rewrite_20260724_cycle3")
OUT.mkdir(parents=True, exist_ok=True)

MAIN_SRC = ROOT / "fastPLS_CMPB_main_cycle2_0.99.7_20260724.docx"
SUPP_SRC = ROOT / "fastPLS_CMPB_supplement_cycle2_0.99.7_20260724.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle3_0.99.8_20260724.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle3_0.99.8_20260724.docx"
REVIEW_OUT = OUT / "fastPLS_CMPB_independent_reviewer_report_cycle3_20260724.docx"


def replace_paragraph(doc, startswith, replacement):
    for p in doc.paragraphs:
        if p.text.startswith(startswith):
            p.clear()
            p.add_run(replacement)
            return
    raise RuntimeError(f"No paragraph starts with: {startswith}")


def revise_main():
    copy2(MAIN_SRC, MAIN_OUT)
    doc = Document(MAIN_OUT)
    replace_paragraph(
        doc,
        "Results: With deterministic IRLBA",
        "Results: With deterministic IRLBA, fastPLS SIMPLS agreed with pls::simpls.fit on controlled synthetic matrices to numerical precision and on preprocessed MetRef data across 2-18 components (prediction correlation 1.000; maximum predictive-subspace angle 0.0027 degrees). "
        "rSVD was evaluated separately as an approximate alternative. In multivariate NMR, training-only validation selected 100 components. Across three isolated runs, held-out SIMPLS-rSVD median total time was 20.14 s (IQR 0.40 s) on CPU and 3.06 s (IQR 0.03 s) on CUDA, with stable RMSD of 0.000861 and 0.000805, respectively. "
        "fastPLS also processed a million-sample ImageNet/DINOv2 post-feature-extraction stress test. Under the documented interfaces and resource limits, the evaluated external R workflows did not complete that task."
    )
    replace_paragraph(
        doc,
        "NMR represented the extreme multivariate-response setting",
        "NMR represented the extreme multivariate-response setting (1,200 training and 321 held-out spectra; p=13,000; q=28,355). "
        "The predictor water region between 4.6 and 4.8 ppm was set to zero in both training and test predictors before any fit. A fixed 20% inner split of the training spectra selected components from 10, 25, 50, 75, and 100 using validation RMSD, leaving the held-out test spectra untouched. "
        "Validation RMSD decreased from 0.001137 at 10 to 0.000894 at 100 components, which was selected. Three isolated matched runs at 100 components gave CPU median R2/Q2 0.9926, RMSD 0.000861, and fit-plus-prediction time 20.14 s (IQR 0.40 s); CUDA gave R2/Q2 0.9935, RMSD 0.000805, and 3.06 s (IQR 0.03 s). "
        "Median peak host RSS was 3,143 MB on CPU and 3,469 MB on CUDA; sampled CUDA compute-applications memory was 3,432 MB in every repetition. Per-spectrum and per-response error distributions, as well as global and zoomed observed-versus-predicted spectra, are provided in the Supplementary Material."
    )
    replace_paragraph(
        doc,
        "The matched CPU/CUDA analysis should be interpreted",
        "The matched CPU/CUDA analysis is a three-run fixed-split computational comparison, not an uncertainty estimate of biological generalization. The representative spectra shown in the Supplementary Material were selected mechanically as the held-out spectrum whose RMSD was closest to the held-out median; they were not selected by visual inspection."
    )
    doc.save(MAIN_OUT)


def revise_supplement():
    copy2(SUPP_SRC, SUPP_OUT)
    doc = Document(SUPP_OUT)
    replace_paragraph(
        doc,
        "Held-out NMR results. At 100 components",
        "Held-out NMR results. At 100 components, three isolated CPU SIMPLS-rSVD runs gave median R2/Q2=0.9926, RMSD=0.000861, MAE=0.000132, and median per-spectrum RMSD=0.000517. CUDA gave median R2/Q2=0.9935, RMSD=0.000805, MAE=0.000133, and median per-spectrum RMSD=0.000518. "
        "Median fit-plus-prediction time was 20.14 s (IQR 0.40 s) on CPU and 3.06 s (IQR 0.03 s) on CUDA. Median maximum process RSS was 3,143 MB and 3,469 MB; CUDA compute-applications memory sampled at 0.2-second intervals was 3,432 MB in all runs."
    )
    doc.save(SUPP_OUT)


def create_review():
    doc = Document()
    doc.add_heading("Independent reviewer report: cycle 3 update", level=0)
    doc.add_paragraph("Manuscript: fastPLS: scalable partial least squares with compiled CPU and accelerator backends for high-dimensional biomedical data")
    doc.add_heading("Recommendation", level=1)
    doc.add_paragraph("Major revision")
    doc.add_heading("Assessment after cycle 3", level=1)
    doc.add_paragraph(
        "The NMR computational comparison now has three isolated repetitions with central tendency and dispersion, and the manuscript clearly distinguishes computational reproducibility from biological generalizability. "
        "This resolves the previous concern about reporting a single timing point. The estimator-preservation wording remains substantially improved by distinguishing deterministic IRLBA from approximate rSVD."
    )
    doc.add_heading("Remaining major comments", level=1)
    for text in [
        "Regenerate the complete real-data benchmark after the corrected IRLBA SIMPLS dispatch and provide raw repetitions, dispersion, peak memory, failures, and requested versus executed estimator/backend in the final tables.",
        "Add a compact real-data ablation that separates incremental component-prefix prediction, cached deflation quantities, and rSVD workspace reuse. It should state which mechanisms are used in deterministic IRLBA versus approximate rSVD.",
        "Present the external-package software comparison in precision-matched float64 form. Float32 remains a compatibility/footprint capability until a full family-by-backend accuracy and memory study demonstrates a practical benefit.",
        "Provide a concise complexity and residency table for all four PLS families and CPU/CUDA/Metal, identifying device-resident and host stages, including prediction, LDA, and cross-validation.",
        "Keep ImageNet clearly restricted to a non-biomedical post-feature-extraction scalability experiment and do not use it as evidence of biomedical utility."
    ]:
        doc.add_paragraph(text, style="List Number")
    doc.add_heading("Minor comments", level=1)
    for text in [
        "Use a tagged software release and archived benchmark manifest before submission.",
        "State explicitly in the NMR figure legend that the water-region mask applies only to the predictor spectrum.",
        "Ensure all speed claims use the same definition of total time and declare whether memory is process RSS, incremental RSS, or device allocation."
    ]:
        doc.add_paragraph(text, style="List Bullet")
    doc.save(REVIEW_OUT)


revise_main()
revise_supplement()
create_review()
print(MAIN_OUT)
print(SUPP_OUT)
print(REVIEW_OUT)
