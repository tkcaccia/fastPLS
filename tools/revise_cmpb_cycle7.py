"""Create Cycle 7 CMPB drafts using the completed matched precision benchmark."""

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.shared import Pt


SRC = Path("/Users/stefano/Desktop/fastPLS/CMPB_rewrite_20260724_cycle6")
OUT = Path("/Users/stefano/Desktop/fastPLS/CMPB_rewrite_20260724_cycle7")
OUT.mkdir(parents=True, exist_ok=True)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle7_0.99.6_20260724.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle7_0.99.6_20260724.docx"
REVIEW_OUT = OUT / "fastPLS_CMPB_independent_reviewer_report_cycle7_20260724.docx"


def replace_paragraph(doc, startswith, replacement):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(startswith):
            paragraph.text = replacement
            return True
    raise RuntimeError(f"Paragraph not found: {startswith}")


def compact_table(table, size=7):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.line_spacing = 0.9
                for run in paragraph.runs:
                    run.font.size = Pt(size)


def add_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for i, values in enumerate(rows):
        for j, value in enumerate(values):
            table.cell(i, j).text = value
    compact_table(table)
    return table


def revise_main():
    copy2(SRC / "fastPLS_CMPB_main_cycle6_0.99.6_20260724.docx", MAIN_OUT)
    doc = Document(MAIN_OUT)
    replace_paragraph(
        doc,
        "Double precision is the numerical reference.",
        "Double precision is the numerical reference. Inputs from the float package can select float32 arithmetic on supported CPU, CUDA, and Metal routes, but the matched real-data benchmark establishes that this is not a universal speed, memory, or numerical-accuracy benefit. At 100 components on CIFAR-100, float32 CUDA reduced sampled GPU memory from 530 MB to 130–288 MB while preserving accuracy within 0.42 percentage points, but was slower than float64 because the current route is hybrid. On the full NMR task (1,200 training spectra, 321 test spectra, 13,000 predictors and 28,355 responses), float32 PLS-SVD retained a similar RMSD (median 6.648×10−4 versus 6.458×10−4), but took 248 s versus 0.843 s on CUDA. Float32 SIMPLS and linear kernel PLS had materially worse RMSD (4.340×10−3 versus 7.211×10−4); CPU OPLS timed out, whereas CUDA OPLS completed but had median RMSD 1.189×10−2. Float32 is therefore a conditional reduced-storage capability, not a generally validated replacement for float64. Replicate IQRs, success counts, and timeout status are reported in Supplementary Table S3."
    )
    replace_paragraph(
        doc,
        "In a fixed-score validation of the revised LDA path",
        "In a fixed-score validation of the revised LDA path, float32 CPU and CUDA predictions agreed exactly across MetRef, CIFAR-100, and SingleCell at 2, 5, 10, and 20 components, with no factorization failures. This classifier-level check does not establish end-to-end float32 equivalence for every model shape. In the matched three-repetition full-NMR study, float32 PLS-SVD remained close to float64 in RMSD but was much slower, and float32 SIMPLS, kernel PLS, and CUDA OPLS had materially larger RMSD; CPU OPLS exceeded the 1,200-s limit in its first run and remaining repetitions were not started. Accordingly, float32 results are reported by model, backend, dataset, and execution status rather than pooled into a single capability claim."
    )
    replace_paragraph(
        doc,
        "Float32 can reduce input and workspace storage",
        "Float32 can reduce input and workspace storage for selected routes, as observed for CUDA CIFAR-100, but it is not yet a universal performance or numerical-accuracy benefit. The completed full-NMR comparison identifies adverse SIMPLS, linear-kernel, and OPLS regimes and is retained with uncertainty and failure status in the supplementary results. Reproducibility should therefore be judged by measured prediction agreement, predictive metrics, selected components, numerical failures, and latent-subspace agreement rather than by precision alone."
    )
    doc.save(MAIN_OUT)


def revise_supplement():
    copy2(SRC / "fastPLS_CMPB_supplement_cycle6_0.99.6_20260724.docx", SUPP_OUT)
    doc = Document(SUPP_OUT)
    replace_paragraph(
        doc,
        "Float32 NMR control.",
        "Float32 NMR control. The matched full-resolution benchmark used 1,200 training spectra, 321 test spectra, 13,000 predictors, 28,355 responses, 100 components, and three isolated repetitions unless a first run exceeded the predeclared 1,200-s timeout. Float32 CUDA PLS-SVD retained similar RMSD to float64 (median 6.648×10−4 [IQR 5.41×10−6] versus 6.450×10−4 [IQR 1.90×10−7]) but required 248.0 s [IQR 7.51] versus 0.843 s [IQR 0.0065] and sampled 1,578 MB versus 664 MB of GPU memory. Float32 CUDA SIMPLS and linear kernel PLS were slower (544–546 s versus 2.89–2.90 s) and had materially worse RMSD (4.340×10−3 versus 7.211×10−4). CPU OPLS timed out in both precision modes; CUDA float64 OPLS also timed out, whereas CUDA float32 OPLS completed in 542.5 s [IQR 0.036] with RMSD 1.189×10−2 [IQR 1.88×10−3]. These results are retained as negative precision evidence: current float32 execution is not recommended for this extreme multivariate-response regime."
    )
    doc.add_page_break()
    doc.add_heading("S6. Matched precision benchmark", level=1)
    doc.add_paragraph(
        "Table S3 reports medians and interquartile ranges (IQR) from three isolated repetitions for the completed NMR measurements. A timeout means that the first attempted repetition exceeded 1,200 s; subsequent repetitions were intentionally not started. RSS denotes peak host resident set size and GPU memory is the sampled process peak."
    )
    rows = [
        ["Method / backend", "float64 time (s)", "float32 time (s)", "float64 RMSD", "float32 RMSD", "f64 GPU MB", "f32 GPU MB", "Completed runs"],
        ["PLS-SVD / CPU", "15.191 (0.073)", "295.533 (0.433)", "6.458e-4 (1.95e-6)", "6.478e-4 (3.89e-6)", "–", "–", "3 / 3; 3 / 3"],
        ["PLS-SVD / CUDA", "0.843 (0.0065)", "247.961 (7.506)", "6.450e-4 (1.90e-7)", "6.646e-4 (5.41e-6)", "664", "1578", "3 / 3; 3 / 3"],
        ["SIMPLS / CPU", "19.753 (0.070)", "1046.774 (10.951)", "7.295e-4 (1.52e-5)", "4.359e-3 (1.88e-4)", "–", "–", "3 / 3; 3 / 3"],
        ["SIMPLS / CUDA", "2.892 (0.007)", "544.227 (0.401)", "7.211e-4 (1.39e-5)", "4.340e-3 (1.68e-4)", "3432", "1550", "3 / 3; 3 / 3"],
        ["Linear kernel PLS / CPU", "19.718 (0.023)", "1046.313 (0.275)", "7.295e-4 (1.52e-5)", "4.359e-3 (1.88e-4)", "–", "–", "3 / 3; 3 / 3"],
        ["Linear kernel PLS / CUDA", "2.896 (0.004)", "545.794 (1.496)", "7.211e-4 (1.39e-5)", "4.340e-3 (1.68e-4)", "3432", "1550", "3 / 3; 3 / 3"],
        ["OPLS / CPU", "timeout", "timeout", "–", "–", "–", "–", "0 / 1; 0 / 1"],
        ["OPLS / CUDA", "timeout", "542.470 (0.036)", "–", "1.189e-2 (1.88e-3)", "–", "1738", "0 / 1; 3 / 3"],
    ]
    add_table(doc, rows)
    doc.add_paragraph("Table S3. Full NMR matched float64/float32 comparison at 100 components. Values are median (IQR); the two entries in the final column give float64 then float32 completed/attempted repetitions.")
    doc.save(SUPP_OUT)


def create_review():
    doc = Document()
    doc.add_heading("Independent reviewer report: cycle 7 update", level=0)
    doc.add_paragraph("Manuscript: fastPLS: scalable partial least squares with compiled CPU and accelerator backends for high-dimensional biomedical data")
    doc.add_heading("Recommendation", level=1)
    doc.add_paragraph("Major revision")
    doc.add_heading("Assessment after cycle 7", level=1)
    doc.add_paragraph("The float32 section is now substantially more convincing. It reports a complete, matched NMR table with uncertainty, successful-run counts, and transparent CPU/CUDA OPLS outcomes. The manuscript no longer implies that reduced precision is universally faster, less memory intensive, or device resident.")
    doc.add_heading("Remaining major comments", level=1)
    for text in [
        "The primary float64 external-package comparison still needs to complete. Provide the raw repetitions, dispersion, peak memory, execution status, requested versus executed estimator/backend, and the final wide package-comparison tables before claiming a broad software advantage.",
        "Archive the exact clean source tree and benchmark artifacts under an immutable release identifier (for example, a tagged GitHub release plus DOI). The new manifest fields are useful, but a manuscript needs the actual release tag and archive location.",
        "The complete real-data benchmark must use the final tagged source tree. Do not combine results from an earlier dirty working tree with a later manuscript claim."
    ]:
        doc.add_paragraph(text, style="List Number")
    doc.add_heading("Minor comments", level=1)
    for text in [
        "Keep the wording that ImageNet is a computational stress test rather than biomedical validation.",
        "Use Retina and Tabula Muris consistently in all final figures, legends, and data manifests.",
        "State in the main text that an OPLS timeout suppresses later repetitions by design, and report completed/attempted counts in all final performance tables."
    ]:
        doc.add_paragraph(text, style="List Bullet")
    doc.save(REVIEW_OUT)


revise_main()
revise_supplement()
create_review()
print(MAIN_OUT)
print(SUPP_OUT)
print(REVIEW_OUT)
