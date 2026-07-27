#!/usr/bin/env python3
"""Synchronize the NMR representative spectrum across text and provenance."""

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260727_cycle83"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260727_cycle84"
MAIN_IN = SOURCE / "fastPLS_CMPB_main_cycle83_0.99.6_20260727.docx"
SUPP_IN = SOURCE / "fastPLS_CMPB_supplement_cycle83_0.99.6_20260727.docx"
MAIN_OUT = OUTPUT / "fastPLS_CMPB_main_cycle84_0.99.6_20260727.docx"
SUPP_OUT = OUTPUT / "fastPLS_CMPB_supplement_cycle84_0.99.6_20260727.docx"

SELECTION = (
    "held-out sample AMI-00BP-8 (index 155), whose per-spectrum RMSD "
    "under 50-component SIMPLS CUDA rSVD was closest to the median across "
    "the 321 held-out spectra"
)


def replace_prefix(document, prefix, replacement):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            paragraph.text = replacement
            return
    raise RuntimeError(f"Paragraph not found: {prefix}")


def revise_main():
    document = Document(MAIN_IN)
    replace_prefix(
        document,
        "At the family-selected settings, CUDA PLS-SVD",
        (
            "At the family-selected settings, CUDA PLS-SVD and SIMPLS achieved "
            "RMSD 0.001043 (Q² 0.98916) and 0.000759 (Q² 0.99425), "
            "respectively. SIMPLS had lower median per-spectrum and "
            f"response-wise error. Figure 4 displays {SELECTION}. This "
            "prespecified descriptive rule did not select the best-predicted "
            "or most visually concordant spectrum."
        ),
    )
    replace_prefix(
        document,
        "Figure 4. NMR predictive and computational analyses.",
        (
            "Figure 4. NMR predictive and computational analyses. Panels A-C "
            "separate family-selected held-out performance from the deposited "
            "165-component historical context. Panels D-E overlay observed and "
            f"predicted intensities for {SELECTION}, over the full response "
            "range and 1.7-0.5 ppm expansion. Panel F reports matched float64 "
            "solver/backend resources at fixed family-specific component "
            "counts. rSVD used oversampling 20, two power iterations, and seed "
            "123."
        ),
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_IN)
    replace_prefix(
        document,
        "The paired backend analysis changes only CPU versus CUDA",
        (
            "The paired backend analysis changes only CPU versus CUDA within "
            "family. The deposited 165-component workflow uses the original "
            "centring-only protocol and is historical context. Predictor "
            "columns with chemical shifts strictly between 4.6 and 4.8 ppm "
            "were set to zero in both training and test predictor matrices "
            "before inner splitting or fitting. No response column was zeroed, "
            "masked, or excluded; all response metrics use all 28,355 "
            f"coordinates. Main-text Figure 4 displays {SELECTION}. This "
            "prespecified descriptive rule did not select the best-predicted "
            "or most visually concordant spectrum."
        ),
    )

    provenance = document.tables[18]
    for row in provenance.rows[1:]:
        if row.cells[0].text == "A17":
            row.cells[1].text = (
                "Figure 4 and Table S12; overlay AMI-00BP-8 (index 155)"
            )
            row.cells[4].text = (
                "source archive recorded; median-RMSD selection in "
                "nmr_representative_spectrum_selection.csv"
            )
            break
    else:
        raise RuntimeError("A17 provenance row not found")

    document.save(SUPP_OUT)


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
