#!/usr/bin/env python3

from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle24"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle25"
SETTINGS = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle13_20260725"
    / "main_benchmark_opls_kernel_settings.csv"
)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle24_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle24_0.99.6_20260725.docx"
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle24_20260725.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle25_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle25_0.99.6_20260725.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle25_20260725.docx"

spec = spec_from_file_location(
    "cycle24_helpers",
    ROOT / "tools" / "revise_cmpb_cycle24_simpls_ablation.py",
)
c24 = module_from_spec(spec)
spec.loader.exec_module(c24)
c16 = c24.c16


def settings_rows(data):
    dataset_labels = {
        "cbmc_citeseq": "CBMC CITE-seq",
        "ccle": "CCLE",
        "cifar100": "CIFAR-100",
        "gtex_v8": "GTEx v8",
        "metref": "MetRef",
        "nmr": "NMR",
        "prism": "PRISM",
        "retina": "Retina",
        "tabula": "Tabula Muris",
        "tcga_brca": "TCGA-BRCA",
        "tcga_hnsc_methylation": "TCGA-HNSC methylation",
        "tcga_pan_cancer": "TCGA Pan-Cancer",
    }
    rows = []
    for _, x in data.iterrows():
        evaluated = x["status"] == "evaluated"
        rows.append(
            (
                dataset_labels[x["dataset"]],
                "OPLS" if x["method"] == "opls" else "kernel PLS",
                str(int(x["reported_total_ncomp"])) if evaluated else "not evaluated",
                str(int(x["predictive_ncomp"])) if evaluated else "-",
                str(int(x["orthogonal_ncomp"])) if evaluated else "-",
                str(x["kernel_type"]) if evaluated and pd.notna(x["kernel_type"]) else "-",
                "-" if not evaluated or pd.isna(x["gamma"]) else f'{x["gamma"]:.6g}',
                "-" if not evaluated or pd.isna(x["degree"]) else str(int(x["degree"])),
                "-" if not evaluated or pd.isna(x["coef0"]) else f'{x["coef0"]:.6g}',
            )
        )
    return rows


def revise_main():
    document = Document(MAIN_SOURCE)

    design = c16.find_paragraph(
        document, "Within each dataset, methods used identical fixed outer splits"
    )
    c24.insert_after(
        design,
        "OPLS and kernel settings were prespecified explicitly for the principal "
        "all-dataset benchmark. OPLS used one orthogonal component (north=1) "
        "whenever at least two total components were evaluated. The reported k "
        "is the total component budget: k-1 predictive components plus one "
        "orthogonal component; north was not tuned. Kernel PLS used the linear "
        "kernel, for which gamma, polynomial degree, and intercept do not "
        "apply. Its rows are therefore linear-kernel implementation controls "
        "and may agree with SIMPLS. Nonlinear RBF and polynomial kernel PLS were "
        "evaluated only in a separate sensitivity analysis, with kernel "
        "parameters and component count selected by five-fold validation on "
        "the training partition (Supplementary Tables S24-S25). All main "
        "benchmark settings are listed in Supplementary Table S35.",
        style="Body Text",
    )

    result = c16.find_paragraph(document, "Kernel sensitivity confirmed")
    result.text = (
        "The kernel-PLS rows in Table 1 use a linear kernel; agreement with "
        "SIMPLS is therefore expected and is not evidence about nonlinear "
        "kernel performance. In the separate training-only sensitivity study, "
        "kernel choice changed predictive behaviour. Polynomial kernel PLS "
        "obtained the highest outer-test accuracy on MetRef (0.990) and CCLE "
        "(0.789), whereas RBF produced the lowest PRISM RMSD (0.5450 on CUDA). "
        "The selected gamma, polynomial degree, intercept, and component count "
        "are reported in Supplementary Tables S24-S25."
    )

    caption = c16.find_paragraph(
        document, "Table 1. Paired CPU/CUDA biomedical workflow benchmark"
    )
    caption.text += (
        " OPLS uses one prespecified orthogonal component; k denotes the total "
        "budget (k-1 predictive plus one orthogonal). Kernel PLS uses the "
        "linear kernel, so gamma, degree, and intercept are not applicable."
    )

    for table in document.tables:
        header = [cell.text.strip() for cell in table.rows[0].cells]
        if header[:3] == ["Dataset", "PLS-SVD", "SIMPLS"] and len(header) >= 5:
            table.rows[0].cells[3].text = "OPLS (north=1)"
            table.rows[0].cells[4].text = "linear kernel PLS"
            break

    document.core_properties.title = (
        "fastPLS CMPB manuscript - explicit OPLS and kernel settings"
    )
    document.save(MAIN_OUT)


def revise_supplement(data):
    document = Document(SUPP_SOURCE)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(
        "S31. OPLS and kernel-PLS settings in the principal benchmark", level=1
    )
    document.add_paragraph(
        "The principal selected-point benchmark did not tune the number of "
        "orthogonal OPLS components or compare nonlinear kernels. OPLS used "
        "north=1 whenever the total component budget was at least two. The "
        "reported component count is total: one slot is orthogonal and the "
        "remaining k-1 slots are predictive. Kernel PLS used the linear kernel; "
        "gamma, degree, and coef0 are consequently not defined. This linear "
        "kernel route is an implementation and backend control and can produce "
        "the same predictions as SIMPLS. It must not be interpreted as an RBF "
        "or polynomial kernel result."
    )
    document.add_paragraph(
        "Nonlinear kernel selection was evaluated separately on MetRef, CCLE, "
        "and PRISM. Five-fold validation on the training partition selected "
        "component count and kernel parameters before one evaluation on the "
        "unchanged test partition. The exact selected gamma, polynomial degree, "
        "intercept, and component count are given in Table S24, with outer-test "
        "runtime, memory, and prediction in Table S25. The external "
        "pls::kernelpls.fit comparator is also a linear predictor-space "
        "algorithm and is not a nonlinear RBF or polynomial kernel model."
    )
    caption = document.add_paragraph(
        "Table S35. OPLS and kernel-PLS settings for the principal paired "
        "CPU/CUDA benchmark. k is the component value displayed in the main "
        "table. OPLS k equals predictive plus orthogonal components. A dash "
        "denotes a parameter that is not applicable or a model that was not "
        "evaluated.",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True
    table = c16.add_table(
        document,
        [
            "Dataset",
            "Model",
            "k shown",
            "Predictive k",
            "north",
            "Kernel",
            "gamma",
            "degree",
            "coef0",
        ],
        settings_rows(data),
        font_size=5.0,
    )
    c24.repeat_header(table.rows[0])
    c16.prevent_row_splitting(table)
    document.core_properties.title = (
        "fastPLS CMPB supplement - explicit OPLS and kernel settings"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    document.add_heading(
        "26. OPLS and kernel-PLS settings were underspecified", level=1
    )
    document.add_paragraph(
        "Reviewer comment: The main benchmark did not report the selected "
        "number of orthogonal components, kernel type, or kernel parameters; "
        "identical SIMPLS and kernel-PLS results suggested frequent use of the "
        "linear shortcut."
    )
    document.add_paragraph(
        "Response: Corrected. Main OPLS fixed north=1; displayed k comprises "
        "k-1 predictive plus one orthogonal component. Main kernel-PLS was "
        "linear, so gamma, degree, and intercept do not apply and agreement "
        "with SIMPLS is expected. We relabelled Table 1 and added Methods and "
        "Results clarifications, Section S31, Table S35, and machine-readable "
        "fields. Separate nonlinear selections remain in Tables S24-S25."
    )
    document.core_properties.title = (
        "fastPLS CMPB response - explicit OPLS and kernel settings"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    data = pd.read_csv(SETTINGS)
    revise_main()
    revise_supplement(data)
    revise_response()
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
