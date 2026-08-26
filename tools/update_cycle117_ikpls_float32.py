from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "artifacts" / "CMPB_rewrite_20260826_cycle116"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260826_cycle117"
OUT.mkdir(parents=True, exist_ok=True)

MAIN_IN = SRC / "fastPLS_CMPB_main_cycle116_0.99.25_20260826.docx"
SUPP_IN = SRC / "fastPLS_CMPB_supplement_cycle116_0.99.25_20260826.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle117_0.99.25_20260826.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle117_0.99.25_20260826.docx"


def replace_paragraph(doc, startswith, text):
    matches = [p for p in doc.paragraphs if p.text.startswith(startswith)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph starting {startswith!r}, found {len(matches)}")
    p = matches[0]
    for run in p.runs:
        run.text = ""
    p.add_run(text)
    return p


def insert_paragraph_before(target, text, style=None):
    p = OxmlElement("w:p")
    target._p.addprevious(p)
    new_p = target._parent.add_paragraph()
    new_p._p.getparent().remove(new_p._p)
    p.addnext(new_p._p)
    p.getparent().remove(p)
    if style:
        new_p.style = style
    new_p.add_run(text)
    return new_p


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table):
    table.style = "Table"
    table.autofit = True
    for r_idx, row in enumerate(table.rows):
        if r_idx == 0:
            tr_pr = row._tr.get_or_add_trPr()
            tbl_header = OxmlElement("w:tblHeader")
            tbl_header.set(qn("w:val"), "true")
            tr_pr.append(tbl_header)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(6.8)
                    if r_idx == 0:
                        run.bold = True
        if r_idx == 0:
            for cell in row.cells:
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "D9EAF7")
                cell._tc.get_or_add_tcPr().append(shd)


def update_main():
    doc = Document(MAIN_IN)
    replace_paragraph(
        doc,
        "Results: An independent dense-LAPACK de Jong panel",
        "Results: An independent dense-LAPACK de Jong panel completed all 82 component-prefix comparisons without numerical failure. Outside an intentionally near-tied singular-value case, maximum held-out prediction error was 4.53 x 10^-15; the tied case retained a 0.015-degree maximum subspace angle and 3.36 x 10^-4 relative prediction error. Against pls::simpls.fit, fastPLS was faster on five of nine datasets in repeated ordinary public-workflow comparisons, with identical accuracy and a maximum 4.85-fold speed-up. Across the broader R-package panel, fastPLS was the fastest tested workflow on seven of nine classification datasets. fastPLS also completed the 13,000 by 28,355-response NMR benchmark and the 1,000,000-row ImageNet/DINOv2 stress test; no tested external R workflow completed ImageNet, and only limited external routes were feasible for NMR. At 50 NMR SIMPLS components, CPU IRLBA, CPU rSVD, and CUDA rSVD required 692.21, 152.09, and 3.91 s, respectively, with RMSD 0.0007561, 0.0007560, and 0.0007561. In a separate float32 comparison, IKPLS matched the ImageNet argmax endpoint but could not fit the selected NMR model because its retained 50-component coefficient tensor alone required 68.66 GiB."
    )
    replace_paragraph(
        doc,
        "Numerical equivalence and software-level performance were assessed separately.",
        "Numerical equivalence and software-level performance were assessed separately. Deterministic float64 fastPLS SIMPLS was first compared with the de Jong SIMPLS implementation in pls::simpls.fit to evaluate the numerical kernel. A separate archived-release CPU experiment compared fastPLS 0.99.25 with the two NumPy formulations provided by IKPLS 6.1.2. The score-explicit formulation computes and retains training scores, whereas the cross-product formulation fits from the predictor and predictor-response cross-products without retaining training scores. The repeated float64 comparison used Breast, MetRef, and CIFAR-100 with 10, 22, and 50 components, respectively. Both IKPLS formulations, deterministic fastPLS IRLBA, and approximate fastPLS rSVD were executed three times for each dataset, giving 36 planned runs. All routes used identical stored splits, externally training-centred predictors, centred one-hot responses, component counts, final held-out prediction, fresh processes, and one effective thread. A separate float32 feasibility extension evaluated public IKPLS 6.1.2 on NMR and ImageNet. NMR was attempted at one and five components under a 10-GiB virtual-memory guard; the coefficient-path storage requirement at 50 components was calculated analytically. ImageNet was evaluated at 100, 200, 500, and 1,000 components using pre-centred float32 arrays, blocked held-out prediction, and one CPU thread; float32 centring and one-hot-response construction were timed separately from model fitting. Because IKPLS and SIMPLS implement different estimators and retain different internal state, both experiments were interpreted as end-to-end software or feasibility comparisons rather than estimator-matched benchmarks."
    )
    replace_paragraph(
        doc,
        "The archived CPU comparison, restricted to Breast, MetRef, and CIFAR-100",
        "The repeated float64 CPU comparison on Breast, MetRef, and CIFAR-100 completed all 36 planned runs. The IKPLS cross-product formulation was fastest, with median total times of 0.00059 s on Breast, 0.00298 s on MetRef, and 0.218 s on CIFAR-100, compared with 0.005, 0.033, and 3.585 s for fastPLS rSVD. Breast accuracy was identical (94.29%). IKPLS reached 75.0% on MetRef and 70.95% on CIFAR-100; fastPLS rSVD reached 77.0% and 72.13%, while deterministic fastPLS IRLBA reached 77.0% and 70.77%. On CIFAR-100, complete-process peak resident-set size was 584 MiB for the IKPLS cross-product formulation and 1,178 MiB for fastPLS rSVD, including their different language runtimes and memory allocators. In the separate single-run float32 extension, IKPLS completed all four ImageNet settings. At 1,000 components it reached 79.99% top-1 and 95.61% top-5 accuracy, required 197.06 s for fitting plus blocked prediction, and reached 11.79 GiB peak process RSS; the separately measured float32 centring and one-hot preparation required 26.36 s. Accuracy closely matched fastPLS argmax at the same prefix (79.98% and 95.61%), but runtime is not estimator matched because fastPLS used one maximal CUDA-rSVD SIMPLS fit to supply every prefix."
    )
    nmr_anchor = replace_paragraph(
        doc,
        "The deposited 165-component PLS-SVD/IRLBA workflow required",
        "The deposited 165-component PLS-SVD/IRLBA workflow required 447.6 s and achieved RMSD 0.000710. It is shown as scientific context for the earlier Nature Communications analysis, not as an implementation-only contrast, because component count, workflow, and hardware differ. Public IKPLS 6.1.2 was also tested in float32. Its one-component fit required 56.58 s and 4.18 GiB incremental RSS but emitted a near-zero-weight warning and returned the training-response mean (RMSD 0.002113), so it did not provide a valid predictive component. The five-component run exceeded the 10-GiB guard, and the retained coefficient path alone would require 68.66 GiB at the selected 50-component SIMPLS setting. IKPLS was therefore infeasible for the scientifically relevant NMR model on the evaluated hardware."
    )
    replace_paragraph(
        doc,
        "Archived fastPLS 0.99.25 completed the million-row SIMPLS fit",
        "Archived fastPLS 0.99.25 completed the million-row SIMPLS fit and blocked prediction of all 281,167 held-out embeddings (Figure 5). Public IKPLS 6.1.2 also completed float32 cross-product fits at 100, 200, 500, and 1,000 components on one CPU thread. Its top-1/top-5 accuracy increased from 62.91%/86.40% at 100 components to 79.99%/95.61% at 1,000 components, while model fitting plus blocked prediction increased from 52.67 to 197.06 s and peak process RSS from 8.24 to 11.79 GiB. These IKPLS accuracies closely followed the corresponding fastPLS argmax trajectory, but the timing comparison is not estimator matched: IKPLS refitted each requested component count, whereas one maximal fastPLS CUDA-rSVD SIMPLS fit supplied all prefixes. The results demonstrate downstream matrix-processing feasibility; they do not establish representation-level reproducibility, biomedical utility, or an optimized ImageNet classifier, and the 1,000-component endpoint remains a boundary stress point. Full values and provenance limitations are reported in Supplementary Sections S15.2 and S18."
    )
    replace_paragraph(
        doc,
        "The principal gain arises from reusing sequential SIMPLS quantities",
        "The principal gain arises from reusing sequential SIMPLS quantities, updating coefficients and fitted values incrementally, and predicting from compact latent factors rather than from retained dense paths. The benefit increases when the response dimension, test set, or number of requested component prefixes is large: compact prediction reduced incremental RSS by up to 77.7%, but offered little benefit when those outputs were intrinsically small. PLS-family and hardware choices remain workload dependent. On one CPU, PLS-SVD was as fast as or faster than SIMPLS across the five matched synthetic matrices (SIMPLS/PLS-SVD time ratio 1.00-3.84), as expected for a one-shot decomposition. On the qualified CUDA cases, SIMPLS approached or marginally exceeded PLS-SVD speed (ratio 0.92-0.98), while retaining sequential orthogonalization and support for component counts not restricted by response rank. Family selection should nevertheless use training-only predictive validation because PLS-SVD and SIMPLS are different estimators. In the broader R-package panel, fastPLS had the lowest observed total time on seven of nine classification datasets, although the matched minimal-output comparison with pls::simpls.fit showed smaller, dataset-dependent differences. IKPLS was faster in the single-thread cross-language experiments and reproduced the ImageNet argmax trajectory closely, but its retained coefficient path made the selected NMR model infeasible even in float32. This contrast emphasizes that performance depends on matrix shape and storage policy: fastPLS contributes an R-native de Jong SIMPLS workflow, compact multivariate-response prediction, nested validation, multiple PLS families, and route diagnostics rather than universal superiority over every PLS implementation."
    )
    replace_paragraph(
        doc,
        "Code and benchmark outputs are available at https://github.com/tkcaccia/fastPLS",
        "Code and benchmark outputs are available at https://github.com/tkcaccia/fastPLS; reusable components are at https://github.com/tkcaccia/kodama-cpp. Central current-version evidence was generated with fastPLS 0.99.25, Git commit 7887401b09e25f54a546a253c255741cb1ab48e5, from source archive fastPLS_0.99.25.tar.gz (SHA-256 604f74d72f5e9540f7378efd37f2ca6ed87ccb9e73b912211331a8cd97233481). The IKPLS large-case float32 scripts and archived summary are in benchmark/ikpls_cross_language/. The ImageNet PLS fitting and prediction stages were rerun from that archive, but the older DINOv2 feature-extraction checkpoint, pooling rule, script, and image-to-row mapping are not fully recoverable; those results are therefore historical and partially reproducible. The deposited 165-component NMR workflow remains a separately identified historical context row."
    )
    doc.save(MAIN_OUT)


def update_supplement():
    doc = Document(SUPP_IN)
    replace_paragraph(
        doc,
        "This archived-release comparison is not estimator matched.",
        "This archived-release comparison is not estimator matched. IKPLS 6.1.2 implements Dayal-MacGregor Improved Kernel PLS in two formulations: the score-explicit variant computes and retains the training-score matrix, whereas the cross-product variant fits from the predictor cross-product and predictor-response cross-product without retaining training scores. fastPLS instead implements de Jong SIMPLS with deterministic IRLBA or approximate rSVD. The repeated CPU contract comprised float64 input, externally applied training centring, identical centred one-hot responses, identical splits and component counts, final held-out prediction, three fresh-process repetitions, and one effective thread. This repeated comparison used Breast, MetRef, and CIFAR-100. A separate single-run float32 feasibility extension evaluated the public IKPLS cross-product formulation on NMR and ImageNet, as detailed in Section S15.2. Resident-set size includes language-runtime allocation and is therefore interpreted as a workflow-feasibility measurement."
    )
    target = next(p for p in doc.paragraphs if p.text == "S16. Selected-point CPU and CUDA benchmark")
    insert_paragraph_before(target, "S15.2 Large-case float32 feasibility extension", "Heading 2")
    insert_paragraph_before(
        target,
        "The extension used IKPLS 6.1.2, its public NumPy cross-product formulation, float32 input and arithmetic, one effective CPU thread, and no conversion time inside model timing. NMR had 1,200 training and 321 held-out rows, 13,000 predictors, and 28,355 responses. ImageNet had 1,000,000 training and 281,167 held-out rows, 1,024 predictors, and 1,000 classes. ImageNet predictors and centred one-hot responses were prepared as float32 arrays before fitting; held-out prediction was blocked. The float32 centring and one-hot preparation required 26.36 s and is reported separately. NMR runs used a 10-GiB virtual-memory guard so that structurally infeasible coefficient paths failed safely rather than exhausting the workstation. Each configuration was measured once and is therefore exploratory.",
    )
    caption = insert_paragraph_before(
        target,
        "Table S10g. IKPLS 6.1.2 float32 feasibility on the NMR and ImageNet case studies. Time excludes external conversion and centring; ImageNet preprocessing required an additional 26.36 s. Peak RSS is absolute complete-process resident-set size. NMR at 50 components was not run because the retained coefficient tensor alone would require 68.66 GiB.",
        "Caption",
    )
    caption.paragraph_format.page_break_before = True
    caption.paragraph_format.keep_with_next = True

    rows = [
        ["Dataset", "Components", "Status", "Fit (s)", "Predict (s)", "Metric", "Peak RSS (MiB)"],
        ["NMR", "1", "Degenerate", "53.13", "3.45", "RMSD 0.002113", "4,281"],
        ["NMR", "5", "Memory guard", "-", "-", "Not estimated", "-"],
        ["NMR", "50", "Not run", "-", "-", "B tensor 68.66 GiB", "-"],
        ["ImageNet", "100", "Success", "43.98", "8.69", "Top-1 62.91%; top-5 86.40%", "8,239"],
        ["ImageNet", "200", "Success", "62.21", "6.19", "Top-1 70.22%; top-5 91.36%", "8,634"],
        ["ImageNet", "500", "Success", "109.93", "7.16", "Top-1 76.61%; top-5 94.64%", "9,783"],
        ["ImageNet", "1,000", "Success", "190.84", "6.22", "Top-1 79.99%; top-5 95.61%", "11,792"],
    ]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r_idx, values in enumerate(rows):
        for c_idx, value in enumerate(values):
            table.cell(r_idx, c_idx).text = value
    style_table(table)
    target._p.addprevious(table._tbl)
    insert_paragraph_before(
        target,
        "The NMR one-component run emitted the IKPLS near-zero-weight warning and produced predictions identical to the training-response mean; its RMSD is therefore a baseline value rather than evidence of a predictive latent component. The five-component attempt failed safely when a 13,000 x 13,000 float32 allocation could not be made under the guard. Public IKPLS retains a coefficient tensor with dimensions components x predictors x responses, corresponding to 1.373 GiB at one component, 6.866 GiB at five components, and 68.66 GiB at 50 components. On ImageNet, all four runs succeeded and closely tracked the archived fastPLS argmax accuracy trajectory. Runtime remains a software-workflow comparison rather than an estimator-matched comparison because IKPLS and SIMPLS differ mathematically and the archived fastPLS timing represents one maximal 1,000-component CUDA fit shared across prefixes.",
    )
    insert_paragraph_before(
        target,
        "Reproducibility. Scripts are in benchmark/ikpls_cross_language/. Raw logs, per-configuration CSV files, preprocessing time, and the joined comparison are in benchmark_results/ikpls_large_float32_20260826/. The experiment used IKPLS 6.1.2, NumPy float32 arrays, one CPU thread, blocked ImageNet prediction, and conversion outside timing.",
    )
    doc.save(SUPP_OUT)


if __name__ == "__main__":
    update_main()
    update_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)
