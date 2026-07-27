#!/usr/bin/env python3

import argparse
import re
from pathlib import Path

from docx import Document


REFERENCE_RE = re.compile(r"^\[(\d+)\]\s+")
CITATION_RE = re.compile(r"\[((?:\d+\s*(?:[-,]\s*)?)+)\]")


def document_text(document, include_tables=False):
    for paragraph in document.paragraphs:
        yield paragraph.text
    if include_tables:
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield cell.text


def expand_citation(text):
    numbers = []
    for part in text.replace(" ", "").split(","):
        if "-" in part:
            start, end = (int(x) for x in part.split("-", 1))
            numbers.extend(range(start, end + 1))
        elif part:
            numbers.append(int(part))
    return numbers


def audit(main_path, supplement_path):
    main = Document(main_path)
    supplement = Document(supplement_path)

    references = []
    for paragraph in main.paragraphs:
        match = REFERENCE_RE.match(paragraph.text.strip())
        if match:
            references.append(int(match.group(1)))

    errors = []
    expected = list(range(1, len(references) + 1))
    if references != expected:
        errors.append(
            f"reference sequence is {references}; expected {expected}"
        )
    if len(references) != len(set(references)):
        errors.append("duplicate reference numbers detected")

    valid = set(references)
    citations = []
    for label, document in (("main", main), ("supplement", supplement)):
        for text in document_text(document):
            if REFERENCE_RE.match(text.strip()):
                continue
            for match in CITATION_RE.finditer(text):
                for number in expand_citation(match.group(1)):
                    citations.append((label, number, text.strip()))
                    if number not in valid:
                        errors.append(
                            f"{label} cites [{number}], which is absent from references"
                        )

    joined_main = "\n".join(document_text(main))
    joined_supplement = "\n".join(document_text(supplement))
    required = {
        "main CIFAR-100": (
            joined_main,
            r"CIFAR-100 followed its documented 50,000/10,000 split \[30\]",
        ),
        "main ImageNet/DINOv2": (
            joined_main,
            r"ImageNet/DINOv2 stress test.*?\[28,29\]",
        ),
        "main pathology": (
            joined_main,
            r"UNI and Prov-GigaPath.*?\[31,32\]",
        ),
        "supplement CIFAR-100": (
            joined_supplement,
            r"CIFAR-100 used its standard 50,000/10,000 partition \[30\]",
        ),
        "supplement ImageNet/DINOv2": (
            joined_supplement,
            r"1,024 DINOv2 features.*?\[28,29\]",
        ),
        "supplement pathology": (
            joined_supplement,
            r"UNI and Prov-GigaPath.*?\[31,32\]",
        ),
    }
    for label, (text, pattern) in required.items():
        if not re.search(pattern, text, flags=re.DOTALL):
            errors.append(f"expected citation mapping not found: {label}")

    if errors:
        raise SystemExit("Reference audit failed:\n- " + "\n- ".join(errors))

    counts = {number: 0 for number in references}
    for _, number, _ in citations:
        counts[number] += 1
    print(
        f"Reference audit passed: {len(references)} sequential references, "
        f"{len(citations)} resolved citation links."
    )
    print("Citation counts:", counts)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("main_docx", type=Path)
    parser.add_argument("supplement_docx", type=Path)
    args = parser.parse_args()
    audit(args.main_docx, args.supplement_docx)


if __name__ == "__main__":
    main()
