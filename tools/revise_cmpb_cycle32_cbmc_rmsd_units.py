#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle31"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle32"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle31_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle31_0.99.6_20260725.docx"
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle31_20260725.docx"
)

MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle32_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle32_0.99.6_20260725.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle32_20260725.docx"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def insert_after(paragraph, text, style=None):
    new_xml = paragraph._parent.add_paragraph()._p
    paragraph._p.addnext(new_xml)
    new_paragraph = Paragraph(new_xml, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    new_paragraph.add_run(text)
    return new_paragraph


def revise_main():
    document = Document(MAIN_SOURCE)
    methods = find_paragraph(document, "Within each dataset, methods used identical")
    cbmc_methods = (
        "For CBMC CITE-seq specifically, the prepared object contained "
        "untransformed integer RNA assay counts for 1,000 preselected genes "
        "as predictors and untransformed integer antibody-derived tag (ADT) "
        "assay counts for 10 surface-protein markers as the multivariate "
        "response. The model subtracted training-column means from both "
        "matrices without variance scaling, centered held-out predictors with "
        "the training predictor means, and restored the training ADT means to "
        "the predictions before evaluation. Its RMSD is therefore in original "
        "ADT count units per cell, pooled over 862 held-out cells and 10 "
        "markers, rather than a unitless normalized error."
    )
    insert_after(methods, cbmc_methods, style="Body Text")

    caption = find_paragraph(
        document,
        "Table 1. Paired CPU/CUDA biomedical workflow benchmark",
    )
    caption.text += (
        " For CBMC CITE-seq, RMSD is in original ADT assay-count units per "
        "cell after response-mean restoration and is pooled across 862 test "
        "cells and 10 markers."
    )

    figure_caption = find_paragraph(
        document,
        "Figure 2. Matched CPU and CUDA outer-test performance",
    )
    figure_caption.text += (
        " Regression RMSD retains each dataset's response units; CBMC "
        "CITE-seq values are pooled ADT assay-count errors, not normalized or "
        "unitless errors."
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - CBMC CITE-seq RMSD units"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)
    paragraph = find_paragraph(
        document,
        "CBMC CITE-seq uses one measured modality as predictors",
    )
    paragraph.text = (
        "CBMC CITE-seq [21] used the prepared SeuratData CBMC object "
        "(dataset version 3.1.4; 8,617 matched cells). X was the stored "
        "untransformed integer RNA assay-count matrix restricted before "
        "benchmarking to 1,000 preselected genes. Y was the stored "
        "untransformed integer ADT assay-count matrix for CD3, CD4, CD8, "
        "CD45RA, CD56, CD16, CD11c, CD14, CD19, and CD34. The stored split "
        "contained 7,755 training and 862 held-out cells and was loaded "
        "unchanged. The benchmark loader applied no library-size "
        "normalization, log transformation, centred-log-ratio transformation, "
        "or additional feature selection. The PLS call used mean centering "
        "without variance scaling: training means were subtracted from X and "
        "Y, held-out X was centered with the training X means, and the "
        "training Y means were restored to predictions. Consequently, the "
        "reported global RMSD is the square root of the mean squared error "
        "across all 8,620 held-out response entries (862 cells x 10 markers), "
        "in original ADT assay-count units per cell. The marker scales are "
        "heterogeneous, so "
        "this pooled RMSD is weighted toward high-count markers and is not "
        "numerically comparable with RMSD measured on normalized responses in "
        "another dataset. For the selected CUDA SIMPLS result, the global "
        "RMSD of 1,056.461 combined marker-wise RMSE values from 56.4 for "
        "CD34 to 3,191.3 for CD45RA. Retina and Tabula Muris are separate "
        "single-cell classification benchmarks and are named separately "
        "throughout the released manifest."
    )

    dimensions_caption = find_paragraph(
        document,
        "Table S3. Dataset dimensions represented",
    )
    dimensions_caption.text += (
        " CBMC CITE-seq q=10 denotes the ten raw ADT count responses; its "
        "RMSD is reported in those assay-count units."
    )

    uncertainty_caption = find_paragraph(
        document,
        "Table S30. Predictive estimates and 95% intervals",
    )
    uncertainty_caption.text += (
        " CBMC CITE-seq RMSD is pooled across its 10 raw ADT count responses "
        "and therefore retains ADT assay-count units per cell."
    )

    document.core_properties.title = (
        "fastPLS CMPB supplement - CBMC CITE-seq RMSD units"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    heading = document.add_heading(
        "32. CBMC CITE-seq RMSD units and preprocessing were unclear",
        level=1,
    )
    heading.paragraph_format.keep_with_next = True
    comment = document.add_paragraph(
        "Reviewer comment: Explain the units and preprocessing underlying "
        "CBMC CITE-seq RMSD; its magnitude cannot otherwise be interpreted."
    )
    comment.paragraph_format.keep_with_next = True
    response = document.add_paragraph(
        "Response: Corrected. We inspected the exact prepared object used by "
        "the benchmark and its loader. The benchmark predicts ten "
        "untransformed integer ADT assay-count responses from 1,000 "
        "untransformed integer RNA assay-count predictors in a fixed "
        "7,755/862-cell split. The loader performs no log, library-size, CLR, "
        "or other modality normalization. The PLS fit mean-centers X and Y "
        "without variance scaling and restores the training ADT means before "
        "held-out scoring. We now define the reported RMSD as the pooled root "
        "mean squared error over 862 x 10 held-out response entries, in "
        "original ADT assay-count units per cell. We also state that the "
        "heterogeneous marker scales make this value scale-weighted and not "
        "directly comparable with normalized-unit RMSD from other datasets. "
        "The Methods, Table 1 caption, Figure 2 caption, Supplementary dataset "
        "protocol, Table S3 caption, and uncertainty-table caption now carry "
        "this clarification."
    )
    response.paragraph_format.keep_together = True
    document.core_properties.title = (
        "fastPLS CMPB response to reviewers - CBMC CITE-seq RMSD units"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    revise_response()


if __name__ == "__main__":
    main()
