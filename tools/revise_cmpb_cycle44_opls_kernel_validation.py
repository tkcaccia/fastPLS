#!/usr/bin/env python3

import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle43"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle44"
RESULTS = (
    ROOT
    / "benchmark_results"
    / "opls_kernel_estimator_validation_verified_20260726"
)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle43_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle43_0.99.6_20260726.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle44_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle44_0.99.6_20260726.docx"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_text(paragraph, old, new):
    if old not in paragraph.text:
        raise RuntimeError(f"Text not found: {old[:120]}")
    text = paragraph.text.replace(old, new)
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    if style is not None:
        inserted.style = style
    if text:
        inserted.add_run(text)
    return inserted


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def prevent_row_splitting(table):
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_width(cell, width_inches):
    width = int(width_inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def format_table(table, widths, font_size=5.7):
    table.autofit = False
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    repeat_header(table.rows[0])
    prevent_row_splitting(table)
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if column_index in (0, 1)
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
                    if row_index == 0:
                        run.bold = True


def add_table(document, headers, rows, widths, font_size=5.7):
    table = document.add_table(rows=1, cols=len(headers))
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = str(value)
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = str(value)
    format_table(table, widths, font_size)
    return table


def fmt_scientific(value):
    value = float(value)
    if value == 0:
        return "0"
    return f"{value:.2e}"


def endpoint_rows():
    path = RESULTS / "opls_kernel_estimator_validation_raw.csv"
    rows = []
    family_label = {
        "OPLS": "OPLS",
        "kernelPLS_rbf": "Kernel PLS/RBF",
        "kernelPLS_poly": "Kernel PLS/poly",
    }
    with path.open(newline="") as handle:
        for row in csv.DictReader(handle):
            shape = (
                f'{row["n_train"]}/{row["p"]}/{row["q"]}/{row["ncomp"]}'
            )
            rows.append(
                (
                    family_label[row["family"]],
                    row["case"].replace("_", " "),
                    row["task"][:5],
                    shape,
                    fmt_scientific(row["operator_relative_error"]),
                    fmt_scientific(row["prediction_relative_error"]),
                    fmt_scientific(row["coefficient_relative_error"]),
                    f'{float(row["max_predictive_score_angle_deg"]):.2e}',
                    "NA"
                    if row["label_agreement"] in ("", "NA", "NaN", "nan")
                    else f'{float(row["label_agreement"]):.3f}',
                    f'{float(row["metric_absolute_difference"]):.2e}',
                    "Pass" if row["passes_all"] == "TRUE" else "Fail",
                )
            )
    return rows


def selection_rows():
    path = RESULTS / "opls_kernel_component_selection_summary.csv"
    rows = []
    family_label = {
        "OPLS": "OPLS",
        "kernelPLS_rbf": "Kernel PLS/RBF",
        "kernelPLS_poly": "Kernel PLS/poly",
    }
    with path.open(newline="") as handle:
        for row in csv.DictReader(handle):
            rows.append(
                (
                    family_label[row["family"]],
                    row["case"].replace("_", " "),
                    row["grid"].replace(";", ","),
                    row["fast_selected_ncomp"],
                    row["reference_selected_ncomp"],
                    "Yes"
                    if row["selected_component_agreement"] == "TRUE"
                    else "No",
                    row["failed_folds"],
                )
            )
    return rows


def revise_main():
    document = Document(MAIN_SOURCE)

    abstract_methods = find_paragraph(document, "Methods: fastPLS provides")
    replace_text(
        abstract_methods,
        (
            "Estimator preservation was assessed only for deterministic IRLBA "
            "SIMPLS against de Jong SIMPLS while reusing deflation, coefficient, "
            "and prediction state across component prefixes."
        ),
        (
            "Estimator preservation was assessed for deterministic IRLBA SIMPLS "
            "against de Jong SIMPLS and, separately, for OPLS and nonlinear RBF "
            "and polynomial kernel PLS against independent equation-level "
            "operators followed by pls::simpls.fit."
        ),
    )

    abstract_results = find_paragraph(document, "Results: The deterministic IRLBA")
    replace_text(
        abstract_results,
        (
            "The deterministic IRLBA path passed all 117 component-level "
            "comparisons with de Jong SIMPLS, supporting estimator preservation "
            "for that optimized deterministic implementation."
        ),
        (
            "The deterministic IRLBA SIMPLS path passed all 117 component-level "
            "comparisons with de Jong SIMPLS. Independent OPLS, RBF-kernel, and "
            "polynomial-kernel validation passed all 18 endpoint and all 18 "
            "fixed-fold component-selection comparisons."
        ),
    )

    algorithms = find_paragraph(document, "OPLS first estimates orthogonal scores")
    replace_text(
        algorithms,
        (
            "Linear, polynomial, and RBF kernels were therefore tuned and "
            "evaluated separately; this sensitivity analysis tests implemented "
            "workflows but is not presented as estimator-equivalence evidence "
            "against an independent nonlinear kernel implementation."
        ),
        (
            "Linear, polynomial, and RBF kernels were therefore tuned and "
            "evaluated separately. Deterministic OPLS and nonlinear kernel PLS "
            "were additionally validated against independent equation-level "
            "filter/kernel construction followed by de Jong SIMPLS; randomized "
            "workflow rows remained outside the equivalence claim."
        ),
    )

    validation = find_paragraph(
        document, "Estimator preservation and randomized approximation were evaluated"
    )
    insert_after(
        validation,
        (
            "A second prespecified estimator-validation study covered OPLS and "
            "nonlinear kernel PLS. An independent R reference implemented the "
            "Trygg-Wold orthogonal-weight and deflation equations without calling "
            "fastPLS filtering code; the filtered predictor matrix was then fitted "
            "with pls::simpls.fit. RBF and polynomial references independently "
            "constructed and training-centred Gram matrices, applied the stored "
            "training centring to held-out kernels, and fitted pls::simpls.fit. "
            "Six tasks covered regression and classification, p<n and p>n, an "
            "ill-conditioned design, gasoline spectroscopy, and breast molecular "
            "classification. Prespecified endpoints were operator, coefficient, "
            "prediction, score-subspace, decoded-label, predictive-metric, failure, "
            "and fixed 5-fold selected-component agreement. Deterministic IRLBA "
            "was requested; documented exact fallback was retained only when the "
            "smaller cross-covariance dimension was below six."
        ),
        validation.style,
    )

    benchmark = find_paragraph(
        document, "Table 1 and Figure 2 show both matched CPU and CUDA backends"
    )
    replace_text(
        benchmark,
        (
            "while PLS-SVD, OPLS, and kernel PLS were not formally audited for "
            "estimator equivalence."
        ),
        (
            "while the PLS-SVD and selected-point randomized OPLS/kernel PLS rows "
            "were not audited for deterministic estimator equivalence. Separate "
            "deterministic OPLS and nonlinear-kernel validation is reported below."
        ),
    )

    simpls_results = find_paragraph(
        document, "A formal estimator-preservation study compared accelerated fastPLS"
    )
    insert_after(
        simpls_results,
        (
            "Independent validation extended this evidence to deterministic OPLS "
            "and nonlinear kernel PLS. Across six synthetic and real tasks, all "
            "18 endpoint comparisons and all 420 fold-component fits completed. "
            "All endpoint rows met the prespecified tolerances: maximum operator, "
            "prediction, and coefficient relative errors were 3.33×10−15, "
            "2.52×10−12, and 1.12×10−10; maximum predictive and OPLS orthogonal "
            "score-subspace angles were 2.09×10−6° and 1.21×10−6°; classification "
            "label agreement was 1.000; and the maximum predictive-metric "
            "difference was 3.19×10−12. The selected component count agreed in "
            "all 18 family-task comparisons, with no failed fold. This evidence "
            "supports the deterministic OPLS filter and explicit nonlinear-kernel "
            "estimators, but not the approximate rSVD rows in Table 1 "
            "(Supplementary Section S34, Tables S39-S40)."
        ),
        simpls_results.style,
    )

    kernel_results = find_paragraph(
        document, "The kernel-PLS rows in Table 1 use a linear kernel"
    )
    replace_text(
        kernel_results,
        (
            "The kernel-PLS rows in Table 1 use a linear kernel; agreement with "
            "SIMPLS is therefore expected and is not evidence about nonlinear "
            "kernel performance."
        ),
        (
            "The kernel-PLS rows in Table 1 use a linear kernel; agreement with "
            "SIMPLS is therefore expected. Nonlinear estimator preservation is "
            "instead supported by the independent deterministic RBF and polynomial "
            "validation above."
        ),
    )

    discussion = find_paragraph(
        document, "fastPLS extends established PLS algorithms through implementation"
    )
    insert_after(
        discussion,
        (
            "The independent OPLS and nonlinear-kernel study reduces a previous "
            "validation gap: agreement was demonstrated at the orthogonal filter "
            "or centred Gram operator, coefficient, score-subspace, prediction, "
            "classification, and fixed-fold selection levels. The evidence is "
            "deliberately limited to deterministic CPU execution and the evaluated "
            "kernels and shapes. It does not establish equivalence for randomized "
            "solvers, reduced precision, or every accelerator stage."
        ),
        discussion.style,
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - independent OPLS and nonlinear kernel validation"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)

    table16 = find_paragraph(document, "Table S16. OPLS selected-point benchmark.")
    replace_text(
        table16,
        "Numerical audit: not formally assessed; workflow-only.",
        (
            "Numerical audit: selected-point rSVD rows are workflow-only; "
            "deterministic OPLS estimator validation is reported in Section S34."
        ),
    )
    table17 = find_paragraph(
        document, "Table S17. kernel PLS selected-point benchmark."
    )
    replace_text(
        table17,
        "Numerical audit: not formally assessed; workflow-only.",
        (
            "Numerical audit: selected-point rSVD rows are workflow-only; "
            "deterministic nonlinear-kernel estimator validation is reported in "
            "Section S34."
        ),
    )

    document.add_heading(
        "S34. Independent OPLS and nonlinear kernel-PLS estimator validation",
        level=1,
    )
    document.add_paragraph(
        (
            "This validation was independent of the fastPLS OPLS filtering and "
            "kernel-construction helpers. For OPLS, the reference centred and "
            "scaled X with training statistics, centred Y, extracted the leading "
            "left singular direction w from X'Y, computed p=X't/(t't), formed "
            "w_o=p-w(w'p)/(w'w), and removed t_o p_o' sequentially. The filtered "
            "training matrix was fitted with pls::simpls.fit, and the stored "
            "orthogonal filter was applied to test data. For nonlinear kernel PLS, "
            "the reference independently formed RBF or degree-three polynomial "
            "Gram matrices, double-centred the training matrix, centred the test "
            "cross-kernel with training means, and fitted pls::simpls.fit. The "
            "fastPLS implementation was not called to construct either reference "
            "operator."
        )
    )
    document.add_paragraph(
        (
            "Six fixed tasks covered synthetic multivariate regression and "
            "classification with p<n and p>n, one ill-conditioned p>n regression "
            "design, the gasoline spectroscopy regression data, and breast "
            "molecular classification. OPLS used one or two orthogonal components "
            "as prespecified by task. RBF gamma was 1/p; the polynomial kernel used "
            "gamma=1/p, degree=3, and coef0=1. Component grids were 1-5 except "
            "Breast (1-3). Seed 123 and identical five-fold partitions were used. "
            "Deterministic IRLBA was requested; exact SVD was used only by the "
            "documented fallback when the smaller cross-covariance dimension was "
            "below six."
        )
    )
    document.add_paragraph(
        (
            "Pass thresholds were operator relative error <=1e-10, prediction "
            "relative error <=1e-4, coefficient relative error <=1e-3, maximum "
            "predictive or orthogonal score-subspace angle <=0.1 degrees, "
            "classification label agreement >=0.995, and predictive-metric "
            "difference <=0.005. All 18 endpoints passed. Maximum operator, "
            "prediction, and coefficient relative errors were 3.33e-15, 2.52e-12, "
            "and 1.12e-10; the largest predictive and orthogonal angles were "
            "2.09e-6 and 1.21e-6 degrees; minimum label agreement was 1.000. "
            "All 420 fold-component fits completed, and the selected component "
            "count agreed in all 18 family-task comparisons. These results validate "
            "the deterministic CPU estimators under the tested conditions. They do "
            "not establish equivalence for rSVD, float32, or accelerator-resident "
            "variants."
        )
    )

    caption_style = find_paragraph(document, "Table S37.").style
    caption39 = document.add_paragraph(
        (
            "Table S39. Independent deterministic endpoint validation. Shape is "
            "ntrain/p/q/A. Op, pred, and coef are relative errors for the "
            "orthogonal-filter or centred-kernel operator, held-out prediction, "
            "and kernel/filtered-space coefficient. Angle is the maximum predictive "
            "score-subspace angle in degrees; labels is decoded-label agreement; "
            "metric is the absolute accuracy or RMSD difference."
        ),
        style=caption_style,
    )
    add_table(
        document,
        (
            "Family",
            "Case",
            "Task",
            "Shape",
            "Op",
            "Pred",
            "Coef",
            "Angle",
            "Labels",
            "Metric",
            "Status",
        ),
        endpoint_rows(),
        (0.78, 1.35, 0.45, 0.72, 0.57, 0.57, 0.57, 0.57, 0.48, 0.57, 0.48),
        5.2,
    )
    caption39.paragraph_format.keep_with_next = True

    caption40 = document.add_paragraph(
        (
            "Table S40. Fixed five-fold component-selection agreement. The same "
            "folds and component grid were supplied to fastPLS and the independent "
            "reference. Failed folds count numerical or execution failures."
        ),
        style=caption_style,
    )
    add_table(
        document,
        (
            "Family",
            "Case",
            "Grid",
            "fastPLS A",
            "Reference A",
            "Agree",
            "Failed folds",
        ),
        selection_rows(),
        (0.95, 2.05, 0.72, 0.72, 0.82, 0.62, 0.82),
        5.8,
    )
    caption40.paragraph_format.keep_with_next = True

    document.add_paragraph(
        (
            "Reproducibility files: "
            "benchmark/benchmark_opls_kernel_estimator_validation.R and "
            "benchmark_results/opls_kernel_estimator_validation_verified_20260726/. "
            "The machine-readable output retains every endpoint, fold, tolerance, "
            "and failure field."
        )
    )

    document.core_properties.title = (
        "fastPLS CMPB supplement - independent OPLS and nonlinear kernel validation"
    )
    document.save(SUPP_OUT)


def main():
    required = [
        RESULTS / "opls_kernel_estimator_validation_raw.csv",
        RESULTS / "opls_kernel_component_selection_summary.csv",
    ]
    for path in required:
        if not path.exists():
            raise FileNotFoundError(path)
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
