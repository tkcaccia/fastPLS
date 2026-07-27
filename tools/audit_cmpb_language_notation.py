#!/usr/bin/env python3

import argparse
from pathlib import Path
import re
import sys

from docx import Document


FORBIDDEN_LITERAL = (
    "PLSSVD",
    "RSVD",
    "nested validation",
    "five-fold validation",
    "five-fold cross-validation",
    "ten-fold validation",
    "ten-fold cross-validation",
    "training-validation",
    "component k",
    "component count k",
    "Selected k",
    "fastPLS k",
    "Reference k",
    "k shown",
    "Predictive k",
    "n/p/q/k",
    "SIMPLS-rSVD",
    "PLS-SVD/rSVD",
)


def text_blocks(document):
    for index, paragraph in enumerate(document.paragraphs):
        yield f"P{index}", paragraph.text
    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                yield f"T{table_index}R{row_index}C{cell_index}", cell.text


def audit(path):
    document = Document(path)
    findings = []
    all_text = "\n".join(text for _, text in text_blocks(document))

    for location, text in text_blocks(document):
        for term in FORBIDDEN_LITERAL:
            if term in text:
                findings.append(f"{location}: forbidden term {term!r}: {text}")
        if re.search(r"\bk-\s*1\b", text) and "k-nearest" not in text:
            findings.append(f"{location}: component expression uses k-1: {text}")

    if "main" in path.name or "supplement" in path.name:
        if "PLS-SVD" not in all_text:
            findings.append("document does not contain PLS-SVD")
        if "rSVD" not in all_text:
            findings.append("document does not contain rSVD")
        if "A denotes" not in all_text and "A is the retained" not in all_text:
            findings.append("document does not define retained component count A")
        if "K denotes" not in all_text and "K is the number" not in all_text:
            findings.append("document does not define fold count K")

    return findings


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("documents", nargs="+", type=Path)
    args = parser.parse_args()

    failed = False
    for path in args.documents:
        findings = audit(path)
        if findings:
            failed = True
            print(f"FAIL {path}")
            for finding in findings:
                print(f"  {finding}")
        else:
            print(f"PASS {path}")
    sys.exit(1 if failed else 0)


if __name__ == "__main__":
    main()
