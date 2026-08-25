from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
IN_DIR = ROOT / "artifacts/CMPB_rewrite_20260808_cycle87"
OUT_DIR = ROOT / "artifacts/CMPB_rewrite_20260824_cycle88"
MAIN_IN = IN_DIR / "fastPLS_CMPB_main_cycle87_0.99.10_20260808.docx"
SUPP_IN = IN_DIR / "fastPLS_CMPB_supplement_cycle87_0.99.10_20260808.docx"
MAIN_OUT = OUT_DIR / "fastPLS_CMPB_main_cycle88_0.99.23_20260824.docx"
SUPP_OUT = OUT_DIR / "fastPLS_CMPB_supplement_cycle88_0.99.23_20260824.docx"


def replace_exact(document, old, new):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == old]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph, found {len(matches)}: {old[:100]}")
    paragraph = matches[0]
    style = paragraph.style
    paragraph.text = new
    paragraph.style = style


def replace_starting(document, start, new):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text.startswith(start)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph, found {len(matches)}: {start[:100]}")
    paragraph = matches[0]
    style = paragraph.style
    paragraph.text = new
    paragraph.style = style


def set_cell(cell, text, size=5):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(size)


OUT_DIR.mkdir(parents=True, exist_ok=True)

main = Document(MAIN_IN)
replace_exact(
    main,
    "The public pls() interface selects PLS family, component count, solver, backend, and, for classification, argmax or latent-space linear discriminant analysis (LDA). pls.single.cv() selects settings within one cross-validation layer; pls.double.cv() uses nested cross-validation. The audited 0.99.10 interface exposes PLS modelling, prediction, evaluation, cross-validation, score plotting, and standalone truncated SVD; PCA and nearest-neighbour classifiers are not package APIs. Requested estimators are never silently substituted.",
    "The public pls() interface selects PLS family, component count, solver, backend, and, for classification, argmax or latent-space linear discriminant analysis (LDA). pls.single.cv() selects settings within one cross-validation layer; pls.double.cv() uses nested cross-validation. The audited 0.99.23 interface exposes PLS modelling, prediction, evaluation, cross-validation, score plotting, and standalone truncated SVD; PCA and nearest-neighbour classifiers are not package APIs. Requested estimators are never silently substituted.",
)
replace_exact(
    main,
    "Direction extraction is modular rather than the principal estimator contribution. float64 CPU fitting supports deterministic IRLBA [15] and approximate rSVD [16], which uses a Gaussian range sketch, orthonormalization, power iterations, and a reduced decomposition. Exact controls, seeds, audit thresholds, and route-specific qualification are consolidated in Supplementary Sections S13-S14 and Tables S8-S9.",
    "Direction extraction is modular rather than the principal estimator contribution. float64 CPU fitting supports deterministic IRLBA [15] and approximate rSVD [16], which uses a Gaussian range sketch, orthonormalization, power iterations, and a reduced decomposition. In release 0.99.23, rSVD remains the default solver but its effective default was changed to oversampling 20 and two power iterations, the CPU and CUDA configuration that met every prespecified check across five seeds. Explicit unqualified overrides remain possible but emit a warning and are recorded in model diagnostics. Exact controls, seeds, audit thresholds, and route-specific qualification are consolidated in Supplementary Sections S13-S14 and Tables S8-S9.",
)
replace_exact(
    main,
    "The primary evidence tested whether deterministic IRLBA SIMPLS met the prespecified numerical tolerances against de Jong SIMPLS and improved runtime. rSVD was assessed separately using computational screens covering prediction error and correlation, latent-subspace agreement, decoded labels, and endpoint metrics. These screens are engineering guardrails rather than clinical-equivalence margins. Their exact definitions, tolerances, discordant counts, and audit results are reported in Supplementary Sections S12-S14 and Tables S7-S9. Deterministic IRLBA remains the confirmatory route when individual prediction changes are consequential.",
    "The primary evidence tested whether deterministic IRLBA SIMPLS met the prespecified numerical tolerances against de Jong SIMPLS and improved runtime. rSVD was assessed separately using computational screens covering prediction error and correlation, latent-subspace agreement, decoded labels, and endpoint metrics. The definitive default-control audit used seeds 1, 7, 19, 43, and 123. These screens are engineering guardrails rather than clinical-equivalence margins, and qualifying a control configuration does not certify every future dataset-specific fit. Their exact definitions, tolerances, discordant counts, and audit results are reported in Supplementary Sections S12-S14 and Tables S7-S9. Deterministic IRLBA remains the confirmatory route when individual prediction changes are consequential.",
)
replace_exact(
    main,
    "Hardware acceleration remained route and shape dependent. CPU, CUDA, and Metal speed-up was summarized only for paired predictions meeting the stated concordance criteria (Figure 3; Supplementary Table S11). The exploratory one-power rSVD setting met only 101/117 audit checks and its speed results are quarantined in Supplementary Figure S1; deterministic IRLBA remains the reference. float32 approximately halved stored inputs on MetRef and PRISM but did not uniformly improve runtime, incremental memory, or agreement (Supplementary Table S9).",
    "Hardware acceleration remained route and shape dependent. CPU, CUDA, and Metal speed-up was summarized only for paired predictions meeting the stated concordance criteria (Figure 3; Supplementary Table S11). The former oversampling-10, one-power CPU setting met only 101/117 checks in the historical single-seed audit; in the expanded five-seed screen, CPU 10/2 failed 5/255 checks and CUDA 10/1 and 10/2 failed 40/40 and 13/40 checks, respectively. These settings and their speed results are quarantined. Release 0.99.23 instead defaults to oversampling 20 and two power iterations, which met 585/585 CPU and 40/40 CUDA checks across five seeds; Metal rSVD remains unqualified. Deterministic IRLBA remains the reference. float32 approximately halved stored inputs on MetRef and PRISM but did not uniformly improve runtime, incremental memory, or agreement (Supplementary Table S9).",
)
replace_exact(
    main,
    "Code and benchmark outputs are available at https://github.com/tkcaccia/fastPLS; reusable components are at https://github.com/tkcaccia/kodama-cpp. Quantitative results remain tied to fastPLS 0.99.6, base commit 6e50bd318f20289101f6b723953830aefa8b95d6, and source-archive SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85. The current audited interface is fastPLS 0.99.10; its source archive has SHA-256 163ac7bd5c0c241f3817fac989e219f71b3956b388f6fcefa2f3420c45051b25. Analysis-specific scripts and archive digests are reported in Supplementary Table S15.",
    "Code and benchmark outputs are available at https://github.com/tkcaccia/fastPLS; reusable components are at https://github.com/tkcaccia/kodama-cpp. Historical quantitative results remain tied to fastPLS 0.99.6, base commit 6e50bd318f20289101f6b723953830aefa8b95d6, and source-archive SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85; they are not relabelled as reruns. The reviewed interface and definitive multi-seed rSVD qualification use fastPLS 0.99.23. Its exact source archive fastPLS_0.99.23.tar.gz has SHA-256 bd8bd6b0dd85219cd9dbe25b429cea02d0a6633df490618d3a3ad73f177fd5e8. Analysis-specific scripts and archive digests are reported in Supplementary Table S15.",
)
main.save(MAIN_OUT)

supp = Document(SUPP_IN)
replace_exact(
    supp,
    "This supplement distinguishes the current audited fastPLS 0.99.10 interface from the 0.99.6 source archive used for the quantitative benchmarks. Table S15 maps each analysis to its result archive and records an exact package commit only when run metadata captured it. A later package version, result date, or manuscript commit is never treated as evidence of a historical computational SHA.",
    "This supplement distinguishes the reviewed fastPLS 0.99.23 interface and multi-seed rSVD qualification from the 0.99.6 source archive used for historical quantitative benchmarks. Table S15 maps each analysis to its result archive and records an exact package commit only when run metadata captured it. A later package version, result date, or manuscript commit is never treated as evidence of a historical computational SHA, and historical values are not relabelled as 0.99.23 reruns.",
)
replace_exact(
    supp,
    "For a linear operator M and target rank r, rSVD draws a Gaussian matrix, forms a range sketch, optionally applies alternating products with M and its transpose, orthonormalizes the range, and decomposes the reduced matrix. The public one-power setting is intended for exploratory speed. It must not be interpreted as certified agreement; power, oversampling, and seed are recorded in every benchmark manifest.",
    "For a linear operator M and target rank r, rSVD draws a Gaussian matrix, forms a range sketch, applies alternating products with M and its transpose, orthonormalizes the range, and decomposes the reduced matrix. Release 0.99.23 uses oversampling 20 and two power iterations as the effective CPU and CUDA default. This setting met every prespecified check in the five-seed qualification panel; qualification applies to the controls and audited panel, not to every future fit. Explicit unqualified controls remain available for research use but produce an unavoidable warning and are recorded in fitted-model diagnostics. Power, oversampling, seed, backend, panel, and qualification status are retained in every fit.",
)
replace_exact(
    supp,
    "For approximate rSVD, oversampling 10 with one power iteration met the prespecified numerical tolerances in 101 of 117 component-level checks, whereas two power iterations met them in 117 of 117. The focused CUDA audit met the prespecified numerical tolerances at all evaluated points with either four power iterations at oversampling 10 or oversampling 20. The authoritative numerical audit is Table S8; full endpoints remain in the repository archive. These are solver-reliability results, not estimator-equivalence evidence.",
    "The historical single-seed CPU audit found 101/117 checks meeting tolerance at oversampling 10 with one power iteration and 117/117 at 10/2. The expanded audit showed that 10/2 was not stable across seeds: 250/255 checks met tolerance, with five MetRef failures. The definitive release audit therefore evaluated oversampling 20 with two power iterations over seeds 1, 7, 19, 43, and 123. All 585 CPU checks and all 40 CUDA checks met the prespecified tolerances. By contrast, CUDA 10/1 and 10/2 met 0/40 and 27/40 checks. The authoritative numerical audit is Table S8; full component-level endpoints and session information remain in publication_release_0.99.23/rsvd_qualification. These are solver-control reliability results, not estimator-equivalence evidence.",
)
replace_exact(
    supp,
    "Current benchmark workflows record repository state, benchmark-script checksum, package version, source-archive SHA-256, compiler, BLAS/LAPACK, thread settings, accelerator libraries, seeds, rSVD controls, and data/split identifiers. Table S15 maps each analysis to its exact evidence archive. NMR and ImageNet used fastPLS_0.99.6.tar.gz (SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85) from base commit 6e50bd318f20289101f6b723953830aefa8b95d6. The audited 0.99.10 archive (SHA-256 163ac7bd5c0c241f3817fac989e219f71b3956b388f6fcefa2f3420c45051b25) changes installation, API cleanup, documentation, and metric-selection behavior; benchmark values are not relabelled as 0.99.10 reruns.",
    "Current benchmark workflows record repository state, benchmark-script checksum, package version, source-archive SHA-256, compiler, BLAS/LAPACK, thread settings, accelerator libraries, seeds, rSVD controls, and data/split identifiers. Table S15 maps each analysis to its exact evidence archive. NMR and ImageNet used fastPLS_0.99.6.tar.gz (SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85) from base commit 6e50bd318f20289101f6b723953830aefa8b95d6; these benchmark values are not relabelled as later-version reruns. The reviewed 0.99.23 source archive has SHA-256 bd8bd6b0dd85219cd9dbe25b429cea02d0a6633df490618d3a3ad73f177fd5e8 and generated the definitive multi-seed rSVD qualification.",
)
replace_starting(
    supp,
    "rSVD uses a fixed seed, Gaussian range sketch, oversampling, and power iterations.",
    "rSVD uses a Gaussian range sketch, oversampling, power iterations, and a recorded seed. Qualification required all of the following: relative Frobenius prediction error <= 0.05, prediction correlation >= 0.99, each score/projection/loading subspace angle <= 10 degrees, decoded-label agreement >= 0.99, and absolute predictive-metric difference <= 0.01. These prespecified values were computational non-inferiority screens rather than clinical or estimator-equivalence margins. The relative error is scale normalized but can conceal localized errors, so it was never interpreted without the endpoint-specific metric and, for classification, exact discordant counts. An agreement threshold of 0.99 permits no changed prediction for test sets smaller than 100 and at most one for n = 100. Release qualification used seeds 1, 7, 19, 43, and 123 and required every check to meet tolerance. The 20/2 default met 585/585 CPU and 40/40 CUDA checks; maximum relative prediction error was 1.24e-09 on CPU and 6.33e-06 on CUDA, minimum label agreement was 1.000, and maximum metric difference was 5.32e-10 and 3.78e-07, respectively. Metal remains unqualified.",
)
replace_exact(
    supp,
    "Table S8. Definitive rSVD numerical-audit summary. CPU contains 117 component-level tests; CUDA contains eight task/endpoint tests per setting.",
    "Table S8. Definitive rSVD numerical-audit summary. The release-default rows use five seeds (1, 7, 19, 43, and 123): CPU contains 585 component-level checks and CUDA contains 40 task/endpoint checks. Rejected CUDA controls were evaluated over the same five seeds.",
)
replace_exact(
    supp,
    "Figure S1. Exploratory one-power rSVD workflow speed relative to deterministic IRLBA. The setting used oversampling 10, one power iteration, and seed 123; it met only 101/117 audit checks and is excluded from estimator-preservation claims.",
    "Figure S1. Historical exploratory one-power rSVD workflow speed relative to deterministic IRLBA. The setting used oversampling 10, one power iteration, and seed 123; it met only 101/117 checks and is excluded from estimator-preservation and release-default claims.",
)
replace_exact(
    supp,
    "The ledger never infers a historical commit from a package version or result date. Where a run did not record its Git SHA, the source status is explicitly not recoverable. SHA-256 prefixes identify immutable result archives; full digests are retained in the machine-readable ledger. The 0.99.10 code audit is documented separately in benchmark/CODE_AUDIT_0.99.10.md and does not alter historical result provenance.",
    "The ledger never infers a historical commit from a package version or result date. Where a run did not record its Git SHA, the source status is explicitly not recoverable. SHA-256 prefixes identify immutable result archives; full digests are retained in the machine-readable ledger. The 0.99.23 release qualification is documented in publication_release_0.99.23/rsvd_qualification and does not alter historical result provenance.",
)

# Replace Table S8 with the definitive five-seed release audit. Preserve the
# existing table style and compact typography.
audit = supp.tables[11]
while len(audit.rows) > 1:
    audit._tbl.remove(audit.rows[-1]._tr)
rows = [
    ("CPU", "20", "2", "585/585", "1.24e-09", "1.00000", "<0.001", "1.000", "5.32e-10", "Qualified release default"),
    ("CUDA", "10", "1", "0/40", "0.2486", "0.96884", "NR", "0.830", "0.0300", "Rejected"),
    ("CUDA", "10", "2", "27/40", "0.0822", "0.99602", "NR", "0.940", "0.0200", "Rejected"),
    ("CUDA", "10", "4", "40/40", "0.0078", "0.99996", "NR", "0.990", "6.75e-05", "Qualified alternative"),
    ("CUDA", "20", "1", "40/40", "0.0027", "1.00000", "NR", "1.000", "0.0002", "Qualified alternative"),
    ("CUDA", "20", "2", "40/40", "6.33e-06", "1.00000", "NR", "1.000", "3.78e-07", "Qualified release default"),
    ("CUDA", "20", "4", "40/40", "6.20e-06", "1.00000", "NR", "1.000", "1.05e-07", "Qualified alternative"),
    ("Metal", "20", "2", "NR", "NR", "NR", "NR", "NR", "NR", "Unqualified; warning"),
]
for values in rows:
    cells = audit.add_row().cells
    for cell, value in zip(cells, values):
        set_cell(cell, value)

# Append the frozen release qualification to the provenance ledger.
ledger = supp.tables[18]
cells = ledger.add_row().cells
values = (
    "A19",
    "Table S8; definitive multi-seed release-default qualification",
    "publication_release_0.99.23/rsvd_qualification",
    "0.99.23",
    "exact source archive and session metadata recorded",
    "benchmark/benchmark_simpls_estimator_preservation.R; benchmark/benchmark_rsvd_cuda_reliability.R",
    "bd8bd6b0dd85",
)
for cell, value in zip(cells, values):
    set_cell(cell, value)

supp.save(SUPP_OUT)
print(MAIN_OUT)
print(SUPP_OUT)
