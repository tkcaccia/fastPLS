from pathlib import Path
import csv
import shutil

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle8"
OUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle9"
RESULTS = ROOT / "benchmark_results" / "simpls_estimator_preservation_20260725"

MAIN_SOURCE = SOURCE_DIR / "fastPLS_CMPB_main_cycle8_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE_DIR / "fastPLS_CMPB_supplement_cycle8_0.99.6_20260725.docx"
MAIN_OUT = OUT_DIR / "fastPLS_CMPB_main_cycle9_0.99.6_20260725.docx"
SUPP_OUT = OUT_DIR / "fastPLS_CMPB_supplement_cycle9_0.99.6_20260725.docx"
RESPONSE_OUT = OUT_DIR / "response_SIMPLS_estimator_preservation_20260725.docx"


def read_csv(path):
    with path.open(newline="", encoding="utf-8-sig") as handle:
        return list(csv.DictReader(handle))


def find_paragraph(document, starts_with):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(starts_with):
            return paragraph
    raise ValueError(f"Paragraph not found: {starts_with}")


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p.getparent().remove(new_paragraph._p)
    new_p.getparent().replace(new_p, new_paragraph._p)
    if style is not None:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def set_paragraph_text(paragraph, text):
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def set_cell_margins(cell, top=55, start=70, bottom=55, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_widths(table, widths_inches):
    table.autofit = False
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_inches:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)
    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths_inches)):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def format_table(table, header_fill="E8EEF5", font_size=7.2):
    table.style = "Table"
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "B7C2CE")
        borders.append(node)
    for row_index, row in enumerate(table.rows):
        row_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        row_pr.append(cant_split)
        if row_index == 0:
            table_header = OxmlElement("w:tblHeader")
            table_header.set(qn("w:val"), "true")
            row_pr.append(table_header)
        for cell in row.cells:
            if row_index == 0:
                shade_cell(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
                    if row_index == 0:
                        run.bold = True


def add_table(document, headers, rows, widths, font_size=7.2):
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    set_table_widths(table, widths)
    format_table(table, font_size=font_size)
    return table


def add_code_block(document, lines):
    for index, line in enumerate(lines):
        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.left_indent = Inches(0.15)
        paragraph.paragraph_format.right_indent = Inches(0.15)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_together = True
        run = paragraph.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        if index == 0:
            paragraph.paragraph_format.space_before = Pt(4)
        if index == len(lines) - 1:
            paragraph.paragraph_format.space_after = Pt(4)
        p_pr = paragraph._p.get_or_add_pPr()
        shade = OxmlElement("w:shd")
        shade.set(qn("w:fill"), "F4F6F9")
        p_pr.append(shade)


def unique_tasks():
    manifest = read_csv(RESULTS / "simpls_estimator_preservation_task_manifest.csv")
    grouped = {}
    for row in manifest:
        key = row["dataset"]
        if key not in grouped:
            grouped[key] = dict(row)
            grouped[key]["seeds"] = []
        grouped[key]["seeds"].append(row["seed"])
    rows = []
    for row in grouped.values():
        seeds = ",".join(sorted(set(row["seeds"]), key=int))
        condition = row["condition"].replace("_", " ")
        rows.append(
            (
                row["dataset"].replace("_", " "),
                row["source"],
                row["task_type"],
                condition,
                seeds,
                row["n_train"],
                row["p"],
                row["q"],
                row["response_rank"],
                "yes" if row["coefficient_identifiable"] == "TRUE" else "no",
                row["ncomp_grid"].replace(";", ","),
            )
        )
    return rows


def deterministic_summary_rows():
    rows = read_csv(RESULTS / "simpls_estimator_preservation_summary.csv")
    output = []
    for row in rows:
        output.append(
            (
                row["source"],
                row["task_type"],
                f'{float(row["prediction_relative_error"]):.3g}',
                f'{float(row["coefficient_relative_error"]):.3g}',
                f'{float(row["score_subspace_max_angle_degrees"]):.3g}',
                f'{float(row["projection_subspace_max_angle_degrees"]):.3g}',
                f'{float(row["loading_subspace_max_angle_degrees"]):.3g}',
                f'{float(row["metric_absolute_difference"]):.3g}',
            )
        )
    return output


def deterministic_cv_rows():
    rows = read_csv(RESULTS / "simpls_estimator_preservation_cv_selection.csv")
    output = []
    for row in rows:
        if row["solver"] != "irlba":
            continue
        output.append(
            (
                row["dataset"].replace("_", " "),
                row["task_type"],
                row["fastpls_selected_ncomp"],
                row["reference_selected_ncomp"],
                row["selected_component_agreement"],
                f'{float(row["maximum_cv_curve_absolute_difference"]):.3g}',
            )
        )
    return output


def revise_main():
    document = Document(MAIN_SOURCE)

    abstract_results = find_paragraph(document, "Results:")
    set_paragraph_text(
        abstract_results,
        "Results: In a prespecified validation spanning eight synthetic regimes "
        "across three seeds and four real datasets, deterministic IRLBA SIMPLS "
        "completed 117 component-level comparisons against de Jong SIMPLS without "
        "failure. All met the stated tolerances; the maximum relative prediction "
        "error was 1.09×10−5, the maximum latent-subspace angle was 0.00145°, and "
        "the component selected by fixed five-fold validation agreed in all 12 "
        "tested tasks. rSVD was evaluated separately as an approximate solver. In "
        "multivariate NMR, training-only validation selected 100 components. Across "
        "three isolated runs, held-out SIMPLS-rSVD median total time was 20.14 s "
        "(IQR 0.40 s) on CPU and 3.06 s (IQR 0.03 s) on CUDA, with RMSD of "
        "0.000861 and 0.000805, respectively."
    )

    benchmark_methods = find_paragraph(document, "External comparisons use independent")
    inserted = insert_after(
        benchmark_methods,
        "Estimator preservation was evaluated in a separate prespecified study. "
        "Eight synthetic regimes covered regression and classification, p<n and "
        "p>n, low- and high-rank responses, ill-conditioned predictors, and exact "
        "rank deficiency; each synthetic regime was repeated with seeds 101, 202, "
        "and 303. Four real-data tasks comprised Breast, Colon, MetRef, and a "
        "deterministically sampled NMR spectral subset. fastPLS deterministic IRLBA "
        "SIMPLS was compared with pls::simpls.fit using identical centred matrices, "
        "component grids, splits, and five-fold assignments. Prespecified endpoints "
        "were relative prediction error, coefficient error when X had full column "
        "rank, score/projection/loading principal angles, decoded-label agreement, "
        "predictive-metric difference, selected-component agreement, and failures. "
        "rSVD used the same design but was analysed separately as an approximate "
        "direction solver.",
        style="Body Text",
    )
    inserted.paragraph_format.space_after = Pt(6)

    existing_result = find_paragraph(document, "A controlled numerical check compared")
    set_paragraph_text(
        existing_result,
        "A formal estimator-preservation study compared accelerated fastPLS SIMPLS "
        "using deterministic IRLBA with pls::simpls.fit. The study comprised eight "
        "synthetic regimes repeated with three seeds and four real datasets, yielding "
        "117 component-level comparisons. All 56 endpoint runs and 24 fixed five-fold "
        "component-selection runs completed without failure. Every deterministic "
        "comparison met the prespecified tolerances: the maximum relative prediction "
        "and coefficient errors were 1.09×10−5 and 1.13×10−5, respectively; the "
        "maximum score, projection, and loading subspace angles were 0.00143°, "
        "0.00145°, and 0.00141°; and decoded-label agreement met or exceeded 0.995. "
        "The selected component count agreed with the reference in all 12 fixed-fold "
        "tasks. These results support preservation of the sequential de Jong estimator "
        "for the deterministic IRLBA path."
    )

    rsvd_result = find_paragraph(document, "A targeted rSVD execution ablation")
    set_paragraph_text(
        rsvd_result,
        "rSVD was not used as estimator-equivalence evidence. In the same formal "
        "suite, its behavior ranged from close agreement to substantial divergence: "
        "the maximum relative prediction error was 1.06, the minimum prediction "
        "correlation was 0.445, and the largest score-subspace angle was 88.4°. "
        "Although the selected component happened to agree in the 12 fixed-fold "
        "tasks, the maximum difference between complete validation curves was 0.105. "
        "A separate execution ablation on MetRef and a fixed CIFAR-100 subset "
        "therefore interprets rSVD workspace reuse only as an accuracy-speed "
        "approximation. Complete task definitions, tolerances, pseudocode, "
        "component-selection results, and failures are reported in the Supplementary "
        "Material."
    )

    package_scope = find_paragraph(document, "Package scope update.")
    delete_paragraph(package_scope)

    document.core_properties.title = (
        "fastPLS CMPB manuscript - formal SIMPLS estimator-preservation revision"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)
    document.add_page_break()
    document.add_heading("S13. Formal SIMPLS estimator-preservation validation", level=1)
    document.add_paragraph(
        "The validation plan was fixed before the full run. Deterministic IRLBA and "
        "approximate rSVD were analysed separately. The reference was "
        "pls::simpls.fit, which implements the sequential de Jong SIMPLS update. "
        "Every comparison used the same centred matrices, split, component grid, "
        "response representation, and seed."
    )

    document.add_heading("S13.1 Prespecified regimes and tolerances", level=2)
    document.add_paragraph(
        "Eight synthetic regimes covered multivariate regression and dummy-response "
        "classification, p<n and p>n, low- and high-rank responses, ill-conditioned "
        "predictors, and exact rank deficiency. Each synthetic regime used seeds 101, "
        "202, and 303. Four real-data tasks were Breast, Colon, MetRef, and a "
        "deterministically sampled NMR spectral subset. The NMR subset used 600 "
        "training and 150 held-out spectra, 300 evenly spaced predictor bins after "
        "excluding 4.6-4.8 ppm, and 120 evenly spaced response bins. It was used only "
        "for numerical agreement, not as a replacement for the full NMR benchmark."
    )
    document.add_paragraph(
        "The deterministic tolerances were fixed as follows: relative prediction "
        "error ≤10−4; relative coefficient error ≤10−3 when the centred predictor "
        "matrix had full column rank; maximum score, projection, and loading "
        "subspace angle ≤0.1°; classification label agreement ≥0.995; classification "
        "accuracy difference ≤0.005; and regression RMSD difference ≤10−4. "
        "Selected-component agreement was assessed exactly. All errors, non-finite "
        "outputs, and failed fits were retained."
    )

    caption = document.add_paragraph(
        "Table S8. Prespecified estimator-preservation tasks. Coefficient "
        "identifiability denotes full column rank of the centred training predictor "
        "matrix.",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True
    add_table(
        document,
        [
            "Task",
            "Source",
            "Type",
            "Condition",
            "Seeds",
            "n",
            "p",
            "q",
            "rank(Y)",
            "Coeff.",
            "Components",
        ],
        unique_tasks(),
        [1.18, 0.42, 0.56, 1.12, 0.52, 0.38, 0.38, 0.40, 0.46, 0.42, 0.66],
    )

    document.add_heading("S13.2 Executable mapping to de Jong SIMPLS", level=2)
    document.add_paragraph(
        "Algorithm S3 gives executable-style R pseudocode. The direction solver is "
        "called once for the current deflated cross-covariance. The remaining "
        "normalization, orthogonalization, deflation, coefficient, and fitted-value "
        "updates are the sequential de Jong operations."
    )
    add_code_block(
        document,
        [
            "X <- sweep(X, 2, colMeans(X), \"-\"); Y <- sweep(Y, 2, colMeans(Y), \"-\")",
            "S <- crossprod(X, Y); B <- 0; Yhat <- 0; V <- matrix(0, ncol(X), ncomp)",
            "for (k in seq_len(ncomp)) {",
            "  r <- dominant_left(S)                 # IRLBA or rSVD direction",
            "  t <- X %*% r; z <- sqrt(crossprod(t)); t <- t/z; r <- r/z",
            "  p <- crossprod(X, t); q <- crossprod(Y, t)",
            "  v <- p; if (k > 1) v <- v - V[,1:(k-1)] %*% crossprod(V[,1:(k-1)], v)",
            "  v <- v/sqrt(crossprod(v)); S <- S - v %*% crossprod(v, S)",
            "  V[,k] <- v; B <- B + r %*% t(q); Yhat <- Yhat + t %*% t(q)",
            "}",
        ],
    )

    document.add_paragraph(
        "Table S9. Mapping of execution optimizations to preserved de Jong "
        "quantities.",
        style="Caption",
    ).paragraph_format.keep_with_next = True
    add_table(
        document,
        ["Optimization", "Preserved quantity", "Execution change"],
        [
            (
                "One maximal path",
                "Sequential components 1,...,k",
                "Requested prefixes are snapshots rather than independent refits",
            ),
            (
                "Shape-aware X'X cache",
                "Score norm and p=X't",
                "Reuses algebraically equivalent products on eligible tall matrices",
            ),
            (
                "Deflation-row cache",
                "S <- S-v(v'S)",
                "Evaluates v'S once before the rank-one update",
            ),
            (
                "Incremental coefficients",
                "B_k=R_k Q_k'",
                "Updates B_k=B_(k-1)+r_k q_k'",
            ),
            (
                "Incremental fitted values",
                "Yhat_k=T_k Q_k'",
                "Updates Yhat_k=Yhat_(k-1)+t_k q_k'",
            ),
            (
                "Implicit cross-covariance",
                "Global operator S=X'Y",
                "Applies Sz=X'(Yz) and S'u=Y'(Xu) without storing S",
            ),
        ],
        [1.35, 1.65, 3.50],
    )
    document.add_paragraph(
        "For deterministic validation, every deflated component requests a fresh "
        "IRLBA direction from the current operator. Randomized workspace reuse is "
        "confined to rSVD and is not part of the estimator-preservation claim."
    )

    document.add_heading("S13.3 Numerical agreement results", level=2)
    document.add_paragraph(
        "The full study comprised 56 endpoint runs and 24 fixed five-fold "
        "component-selection runs. No run failed. Deterministic IRLBA contributed "
        "117 component-level comparisons; all passed every applicable tolerance. "
        "The selected component agreed in all 12 deterministic fixed-fold tasks."
    )

    document.add_paragraph(
        "Table S10. Worst deterministic IRLBA discrepancy within each source and "
        "task type across all component-level comparisons.",
        style="Caption",
    ).paragraph_format.keep_with_next = True
    add_table(
        document,
        [
            "Source",
            "Task",
            "Pred. rel. error",
            "Coeff. rel. error",
            "Score angle (°)",
            "Projection angle (°)",
            "Loading angle (°)",
            "Metric difference",
        ],
        deterministic_summary_rows(),
        [0.58, 0.72, 0.78, 0.78, 0.78, 0.86, 0.82, 0.88],
    )

    document.add_paragraph(
        "Table S11. Fixed five-fold component selection for deterministic IRLBA "
        "SIMPLS and pls::simpls.fit.",
        style="Caption",
    ).paragraph_format.keep_with_next = True
    add_table(
        document,
        [
            "Task",
            "Type",
            "fastPLS k",
            "Reference k",
            "Agreement",
            "Maximum curve difference",
        ],
        deterministic_cv_rows(),
        [1.80, 1.00, 0.72, 0.78, 0.76, 1.44],
        font_size=6.5,
    )

    document.add_paragraph(
        "The rSVD results are deliberately not pooled with deterministic IRLBA. "
        "Across the same component-level comparisons, the maximum relative "
        "prediction error was 1.06, minimum prediction correlation was 0.445, "
        "maximum score-subspace angle was 88.4°, minimum classification label "
        "agreement was 0.133, and maximum predictive-metric difference was 0.526. "
        "The selected component happened to agree in all 12 fixed-fold tasks, but "
        "the maximum difference between complete validation curves was 0.105. These "
        "results confirm that rSVD is a potentially useful approximation rather than "
        "evidence of exact estimator preservation."
    )

    prediction_plot = RESULTS / "plots" / "simpls_prediction_agreement.png"
    angle_plot = RESULTS / "plots" / "simpls_subspace_agreement.png"
    document.add_picture(str(prediction_plot), width=Inches(6.45))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        "Figure S16. Relative held-out prediction error against "
        "pls::simpls.fit. Each point is one component-level comparison. The dashed "
        "line is the prespecified 10−4 deterministic tolerance.",
        style="Caption",
    )
    document.add_picture(str(angle_plot), width=Inches(6.45))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        "Figure S17. Maximum principal angle between training-score subspaces. "
        "Each point is one component-level comparison. The dashed line is the "
        "prespecified 0.1° deterministic tolerance.",
        style="Caption",
    )

    old_preliminary = find_paragraph(
        document,
        "Controlled SIMPLS equivalence."
    )
    set_paragraph_text(
        old_preliminary,
        "Controlled SIMPLS equivalence. The preliminary three-scenario screen has "
        "been superseded by the prespecified formal validation in Section S13."
    )
    old_agreement = find_paragraph(document, "Estimator agreement.")
    set_paragraph_text(
        old_agreement,
        "Estimator agreement. The final prespecified deterministic IRLBA and "
        "separate rSVD results are reported in Section S13 and Tables S8-S11."
    )
    supplement_scope = find_paragraph(document, "Supplementary scope update.")
    delete_paragraph(supplement_scope)

    document.core_properties.title = (
        "fastPLS CMPB supplementary material - formal SIMPLS validation"
    )
    document.save(SUPP_OUT)


def write_response():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.space_after = Pt(7)
    styles["Heading 1"].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    document.add_heading("Response to reviewer: SIMPLS estimator preservation", level=1)
    reviewer = document.add_paragraph()
    reviewer.add_run("Reviewer comment. ").bold = True
    reviewer.add_run(
        "The principal methodological claim requires broader estimator-preservation "
        "evidence, including regression and classification, low- and high-rank "
        "responses, p<n and p>n, ill-conditioned data, real datasets, coefficient "
        "and subspace agreement, selected components, failures, tolerances, and "
        "separate treatment of deterministic IRLBA and approximate rSVD."
    )
    response = document.add_paragraph()
    response.add_run("Response. ").bold = True
    response.add_run(
        "We agree and added a prespecified formal validation comprising eight "
        "synthetic regimes repeated with three seeds and four real datasets "
        "(Breast, Colon, MetRef, and a deterministic NMR spectral subset). The "
        "study included regression and classification, p<n and p>n, low- and "
        "high-rank responses, ill-conditioning, and exact rank deficiency. We "
        "reported relative prediction and coefficient errors, score/projection/"
        "loading principal angles, decoded-label agreement, predictive-metric "
        "differences, fixed-fold selected-component agreement, and all failures "
        "under prespecified tolerances."
    )
    document.add_paragraph(
        "All 56 endpoint runs and 24 five-fold component-selection runs completed "
        "without failure. Across 117 deterministic IRLBA component-level "
        "comparisons, all applicable tolerances were met. The maximum relative "
        "prediction error was 1.09×10−5, maximum coefficient error was 1.13×10−5, "
        "and maximum score, projection, and loading subspace angles were 0.00143°, "
        "0.00145°, and 0.00141°. The selected component agreed with "
        "pls::simpls.fit in all 12 fixed-fold tasks."
    )
    document.add_paragraph(
        "rSVD is now reported separately and is not used as estimator-equivalence "
        "evidence. Its maximum relative prediction error was 1.06 and maximum "
        "score-subspace angle was 88.4° in difficult synthetic regimes, confirming "
        "that it is an explicitly approximate solver. We added executable-style "
        "pseudocode and a mapping table showing how each acceleration preserves the "
        "corresponding de Jong update."
    )
    document.add_paragraph(
        "Changes appear in the main Methods and Results and in Supplementary "
        "Section S13, Tables S8-S11, and Figures S16-S17. The complete scripts and "
        "machine-readable outputs are stored under benchmark/ and "
        "benchmark_results/simpls_estimator_preservation_20260725/."
    )
    document.save(RESPONSE_OUT)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    write_response()
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
