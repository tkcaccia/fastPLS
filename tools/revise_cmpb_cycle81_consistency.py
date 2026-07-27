#!/usr/bin/env python3
"""Resolve the consistency and provenance comments from the cycle80 review."""

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260727_cycle80"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260727_cycle81"
MAIN_IN = SOURCE / "fastPLS_CMPB_main_cycle80_0.99.6_20260727.docx"
SUPP_IN = SOURCE / "fastPLS_CMPB_supplement_cycle80_0.99.6_20260727.docx"
MAIN_OUT = OUTPUT / "fastPLS_CMPB_main_cycle81_0.99.6_20260727.docx"
SUPP_OUT = OUTPUT / "fastPLS_CMPB_supplement_cycle81_0.99.6_20260727.docx"

ARCHIVE = "fastPLS_0.99.6.tar.gz"
ARCHIVE_SHA = "c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85"
BASE_COMMIT = "6e50bd318f20289101f6b723953830aefa8b95d6"


def replace_prefix(document, prefix, replacement):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            paragraph.text = replacement
            return
    raise RuntimeError(f"Paragraph not found: {prefix}")


def revise_main():
    document = Document(MAIN_IN)
    replace_prefix(
        document,
        "Code and benchmark outputs are available at",
        (
            "Code and benchmark outputs are available at "
            "https://github.com/tkcaccia/fastPLS; reusable components are at "
            "https://github.com/tkcaccia/kodama-cpp. The reviewed software "
            f"snapshot is fastPLS 0.99.6. Commit {BASE_COMMIT} identifies the "
            f"base source, and the exact experimental source archive {ARCHIVE} "
            f"has SHA-256 {ARCHIVE_SHA}. Analysis-specific scripts and archive "
            "digests are reported in Supplementary Table S15."
        ),
    )
    OUTPUT.mkdir(parents=True, exist_ok=True)
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_IN)

    # Algorithm S1: use the component notation declared in S1.
    target = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.startswith(
            "Form the compact latent prediction factors from prefixes"
        )
    )
    target.text = (
        "Form the compact latent prediction factors from prefixes 1,...,a "
        "for each requested a in C."
    )

    replace_prefix(
        document,
        "Current benchmark workflows record repository state,",
        (
            "Current benchmark workflows record repository state, benchmark-"
            "script checksum, package version, source-archive SHA-256, compiler, "
            "BLAS/LAPACK, thread settings, accelerator libraries, seeds, rSVD "
            "controls, and data/split identifiers. Table S15 maps each analysis "
            "to its exact evidence archive. The current NMR and ImageNet runs "
            f"used {ARCHIVE} (SHA-256 {ARCHIVE_SHA}) from base commit "
            f"{BASE_COMMIT}."
        ),
    )

    replace_prefix(
        document,
        "The primary software comparison used float64,",
        (
            "The primary software comparison used float64, deterministic CPU "
            "SIMPLS, the same split and component count, and one effective BLAS "
            "thread. fastPLS argmax is estimator matched to pls::simpls.fit; "
            "LDA is a workflow comparison because the prediction head differs. "
            "Memory is absolute process RSS and is reported for feasibility, "
            "not isolated algorithmic allocation. Across 126 attempted "
            "external-package dataset/method runs, 110 completed and 16 did "
            "not: 12 were documented package limitations, two were killed at "
            "the timeout, and two produced execution errors. The previously "
            "incomplete CIFAR-100 fastPLS SIMPLS-LDA row was rerun "
            "independently with a 7,200-s limit; all three replicates completed, "
            "with median accuracy 0.8710, median total time 10.118 s, and median "
            "peak process RSS 1,687 MB. The current package-comparison workflow "
            "uses a 10,000-s default timeout. The isolated evidence is stored "
            "under benchmark_results/manuscript_revision_cycle78_20260726/"
            "cifar100_fastpls_simpls_lda."
        ),
    )

    replace_prefix(
        document,
        "The family-selected predictive analysis and paired backend analysis",
        (
            "The family-selected predictive analysis and paired backend "
            "analysis answer different questions. Component selection used five "
            "paired training-only splits. The PLS-SVD candidate grid was "
            "1, 2, 3, 5, 7, 10, 25, 50, 75, 100, 125, 150, 165, 175, 200, "
            "250, and 300; the SIMPLS grid was 10, 25, 50, 75, 100, 125, 150, "
            "165, 175, 200, 250, and 300. For each family, the one-standard-error "
            "threshold was the minimum mean validation RMSD plus the standard "
            "error at its minimizing component count, and the smallest eligible "
            "count was retained. Only five components were eligible for "
            "PLS-SVD; 50, 75, and 100 were eligible for SIMPLS, so 50 was "
            "retained. These are the selected values within the evaluated grids, "
            "not global optima. The paired backend analysis changes only CPU "
            "versus CUDA within family. The deposited 165-component workflow "
            "uses the original centring-only protocol and is historical context. "
            "Predictor columns with chemical shifts strictly between 4.6 and "
            "4.8 ppm were set to zero in both training and test predictor "
            "matrices before inner splitting or fitting. No response column was "
            "zeroed, masked, or excluded; all response metrics use all 28,355 "
            "coordinates. Main-text Figure 4 displays held-out sample "
            "AMI-00BP-8 (index 155), selected by the prespecified descriptive "
            "rule of SIMPLS RMSD closest to the held-out median; it was not the "
            "best-predicted spectrum."
        ),
    )

    # Table S3: the ImageNet evidence now includes classification and retrieval.
    data_table = document.tables[2]
    for row in data_table.rows[1:]:
        if row.cells[0].text == "ImageNet/DINOv2":
            row.cells[1].text = (
                "Exploratory classification/retrieval stress test"
            )
            break
    else:
        raise RuntimeError("ImageNet/DINOv2 row not found in Table S3")

    OUTPUT.mkdir(parents=True, exist_ok=True)
    document.save(SUPP_OUT)


def main():
    revise_main()
    revise_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
