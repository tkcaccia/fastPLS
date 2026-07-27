#!/usr/bin/env python3

from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle25"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle26"
CAPABILITY = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle13_20260725"
    / "float32_capability_table.csv"
)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle25_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle25_0.99.6_20260725.docx"
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle25_20260725.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle26_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle26_0.99.6_20260725.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle26_20260725.docx"

spec = spec_from_file_location(
    "cycle25_helpers",
    ROOT / "tools" / "revise_cmpb_cycle25_opls_kernel_settings.py",
)
c25 = module_from_spec(spec)
spec.loader.exec_module(c25)
c24 = c25.c24
c16 = c25.c16


def capability_rows(data):
    rows = []
    for _, x in data.iterrows():
        status = "conditional"
        if "small controlled validation only" in x["validation_status"]:
            status = "limited"
        if "avoid float32" in x["recommendation"]:
            status = "unsafe for extreme q"
        rows.append(
            (
                x["method"],
                x["backend"],
                x["supported_svd"],
                status,
                x["validation_status"],
                x["observed_limitation"],
                x["automatic_warning"],
            )
        )
    return rows


def revise_main():
    document = Document(MAIN_SOURCE)

    methods = c16.find_paragraph(
        document, "Double precision is the numerical reference."
    )
    methods.text = (
        "Double precision is the numerical reference. Inputs from the float "
        "package select float32 arithmetic on supported CPU, CUDA, and Metal "
        "routes, but float32 is treated as a conditional reduced-storage mode, "
        "not a general acceleration or equivalence claim. Before allocation, "
        "the public interface applies method- and shape-based checks. Float32 "
        "SIMPLS and kernel-PLS classification warn about the accuracy "
        "sensitivity observed in matched real-data tests. For multivariate "
        "regression with q >= 10,000 and at least 50 requested components, "
        "PLS-SVD warns about severe runtime and device-memory penalties, while "
        "SIMPLS, OPLS, and kernel PLS warn about numerical and runtime risk. "
        "Nonlinear float32 kernel PLS warns about limited validation and the "
        "retained n-by-n Gram matrix. The checks depend on matrix shape and "
        "requested estimator, not dataset identity. The complete capability "
        "matrix and warning policy are reported in Supplementary Table S36."
    )

    results = c16.find_paragraph(
        document, "In a fixed-score validation of the revised LDA path"
    )
    results.text = (
        "Float32 support was not uniformly beneficial. Small controlled "
        "CPU/CUDA/Metal tests showed exact decoded-label agreement for all four "
        "families and negligible univariate-regression prediction differences. "
        "On MetRef, PLS-SVD preserved accuracy, whereas float32 SIMPLS and "
        "linear kernel PLS differed from matched float64 results by two to five "
        "percentage points depending on backend. On PRISM, input storage was "
        "approximately halved, but float32 PLS-SVD, SIMPLS, and linear kernel "
        "PLS were 2.4-6.0 times slower despite similar RMSD; OPLS was faster in "
        "float32 and retained similar RMSD. The full-response NMR experiment "
        "identified an unsafe regime: float32 PLS-SVD retained similar RMSD but "
        "was approximately 294 times slower on CUDA and used more sampled "
        "device memory, while float32 SIMPLS and linear kernel PLS were "
        "approximately 188 times slower and had about six-fold larger RMSD. "
        "OPLS either timed out or produced poor RMSD. These findings motivate "
        "the capability-specific warnings and recommendations in Supplementary "
        "Tables S7, S21, and S36."
    )

    discussion = c16.find_paragraph(
        document, "Float32 can reduce input and workspace storage"
    )
    discussion.text = (
        "Float32 approximately halves raw input storage and can reduce selected "
        "workspaces, but it is not established as a general package advantage. "
        "Its utility is route- and matrix-dependent: moderate PLS-SVD and some "
        "OPLS cases retained predictions, whereas SIMPLS/kernel classification "
        "showed measurable accuracy sensitivity and extreme-response NMR "
        "exposed severe speed, memory, and numerical failures. The package "
        "therefore retains float64 as the reference, emits automatic warnings "
        "for the observed risk regimes, and requires held-out agreement checks "
        "before float32 results are used for confirmatory analysis."
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - float32 capability boundaries"
    )
    document.save(MAIN_OUT)


def revise_supplement(data):
    document = Document(SUPP_SOURCE)
    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(
        "S32. Float32 capability boundaries and automatic warnings", level=1
    )
    document.add_paragraph(
        "Implementation support and numerical validation are distinct. A route "
        "is listed as supported when the package executes it without silently "
        "upcasting the input; this does not establish speed or float64 "
        "equivalence. 'Conditional' denotes moderate real-data evidence with "
        "route-specific limitations. 'Limited' denotes controlled smoke or "
        "portability evidence only. 'Unsafe for extreme q' denotes a measured "
        "full-response regime in which runtime, device memory, or RMSD made the "
        "float32 route unsuitable."
    )
    document.add_paragraph(
        "Warnings are evaluated before model workspaces are allocated. They "
        "use the requested estimator, task type, response dimension, component "
        "count, and kernel family; no dataset name is inspected. Extreme "
        "multivariate regression is guarded at q >= 10,000 and ncomp >= 50, "
        "matching the scale at which the full NMR validation exposed failures. "
        "Classification warnings for SIMPLS and kernel PLS reflect the two- to "
        "five-percentage-point float32/float64 differences observed on MetRef. "
        "The warning thresholds are empirical safety boundaries, not proofs "
        "that smaller problems are equivalent."
    )
    caption = document.add_paragraph(
        "Table S36. Float32 implementation and validation capability matrix. "
        "The solver column lists accepted float32 low-rank routes. Warning "
        "conditions are automatic and shape based. CPU and Metal accept rSVD "
        "and IRLBA-style float32 solvers; CUDA accepts rSVD.",
        style="Caption",
    )
    caption.paragraph_format.keep_with_next = True
    table = c16.add_table(
        document,
        [
            "Model",
            "Backend",
            "Solver",
            "Status",
            "Evidence",
            "Observed limitation",
            "Automatic warning",
        ],
        capability_rows(data),
        font_size=4.8,
    )
    c24.repeat_header(table.rows[0])
    c16.prevent_row_splitting(table)
    document.add_paragraph(
        "Measured anchors. Float32 approximately halved MetRef/PRISM input "
        "storage. MetRef PLS-SVD accuracy remained 0.80, while SIMPLS and "
        "linear kernel PLS differed by 0.02 on CPU and 0.05 on CUDA. PRISM RMSD "
        "remained close, but float32 PLS-SVD, SIMPLS, and linear kernel PLS were "
        "slower. On full-response NMR, CUDA SIMPLS float32/float64 time was "
        "544/2.89 s and RMSD was 4.340e-3/7.211e-4. These failures preclude a "
        "general float32 validation claim."
    )
    document.core_properties.title = (
        "fastPLS CMPB supplement - float32 capability boundaries"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    document.add_heading(
        "27. Float32 was not established as a general package advantage",
        level=1,
    )
    document.add_paragraph(
        "Reviewer comment: Several float32 routes were slower or materially "
        "less accurate, particularly on NMR. Supported and validated "
        "combinations should be separated in a capability table, with automatic "
        "warnings for unsafe routes."
    )
    document.add_paragraph(
        "Response: Corrected. We no longer characterize float32 as a general "
        "advantage. Supplementary Table S36 separates implementation support, "
        "validation strength, observed limitations, and recommendations for "
        "all four model families on CPU, CUDA, and Metal. The public pls() "
        "interface now emits shape-based warnings before allocation: SIMPLS "
        "and kernel-PLS classification warn about observed accuracy "
        "sensitivity; q >= 10,000 with ncomp >= 50 triggers an extreme-response "
        "performance warning for PLS-SVD and a numerical-risk warning for "
        "SIMPLS, OPLS, and kernel PLS; nonlinear kernel PLS warns about limited "
        "validation and retained Gram-matrix storage. Unit tests verify these "
        "rules, and the manuscript now reports MetRef, PRISM, and NMR positive "
        "and negative evidence explicitly."
    )
    document.core_properties.title = (
        "fastPLS CMPB response - float32 capability boundaries"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    data = pd.read_csv(CAPABILITY)
    revise_main()
    revise_supplement(data)
    revise_response()
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
