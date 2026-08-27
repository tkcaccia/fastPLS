from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260826_cycle107"
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260826_cycle108"
FIGURE_DIR = OUTPUT_DIR / "figures"

MAIN_SOURCE = SOURCE_DIR / "fastPLS_CMPB_main_cycle107_0.99.25_20260826.docx"
SUPP_SOURCE = SOURCE_DIR / "fastPLS_CMPB_supplement_cycle107_0.99.25_20260826.docx"
MAIN_OUTPUT = OUTPUT_DIR / "fastPLS_CMPB_main_cycle108_0.99.25_20260826.docx"
SUPP_OUTPUT = OUTPUT_DIR / "fastPLS_CMPB_supplement_cycle108_0.99.25_20260826.docx"

FIGURE_2 = FIGURE_DIR / "Figure_2_external_packages_updated.png"
FIGURE_3 = FIGURE_DIR / "Figure_3_accelerator_concordance_updated.png"
FIGURE_5 = FIGURE_DIR / "Figure_5_imagenet_updated.png"
FIGURE_S4 = FIGURE_DIR / "Figure_S4_selected_cpu_cuda.png"


def paragraph_by_prefix(document, prefix):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph beginning {prefix!r}; found {len(matches)}"
        )
    return matches[0]


def replace_text(document, replacements):
    for paragraph in document.paragraphs:
        value = paragraph.text
        for old, new in replacements:
            value = value.replace(old, new)
        if value != paragraph.text:
            paragraph.text = value


def remove_figure_and_caption(caption):
    parent = caption._p.getparent()
    previous = caption._p.getprevious()
    while previous is not None:
        text = "".join(previous.itertext()).strip()
        has_drawing = bool(previous.xpath(".//w:drawing"))
        if text and not has_drawing:
            break
        candidate = previous
        previous = previous.getprevious()
        parent.remove(candidate)
    parent.remove(caption._p)


def insert_page_break_before(document, anchor):
    paragraph = document.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    anchor._p.addprevious(paragraph._p)


def insert_figure_before(
    document,
    anchor,
    image_path,
    caption_text,
    width,
    caption_style,
    page_break=False,
):
    if page_break:
        insert_page_break_before(document, anchor)
    figure = document.add_paragraph()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.keep_with_next = True
    figure.add_run().add_picture(str(image_path), width=Inches(width))
    anchor._p.addprevious(figure._p)

    caption = document.add_paragraph(style=caption_style)
    caption.text = caption_text
    caption.paragraph_format.keep_with_next = False
    anchor._p.addprevious(caption._p)
    return caption


def rewrite_main():
    document = Document(MAIN_SOURCE)
    caption_style = paragraph_by_prefix(document, "Figure 2.").style

    # Replace the restricted three-dataset backend figure with the complete
    # concordance map requested for the main text.
    remove_figure_and_caption(paragraph_by_prefix(document, "Figure 2."))

    intro = paragraph_by_prefix(document, "Results are organized around")
    intro.text = (
        "Results are organized around the shape-dependent SIMPLS execution "
        "contribution. We first compare complete single-CPU classification "
        "workflows across independent R implementations, then separate "
        "numerically concordant CPU, CUDA, and Metal execution. NMR provides "
        "the principal biomedical high-response case study, while ImageNet is "
        "retained as a qualified foundation-model-scale feasibility analysis. "
        "Exact estimator validation, approximate-solver qualification, route "
        "diagnostics, and complete result tables remain in the Supplementary "
        "Material."
    )

    comparison = paragraph_by_prefix(document, "The strict comparison completed")
    comparison.text = comparison.text.replace(
        "The two timing profiles answer different questions and are not pooled "
        "(Supplementary Figure S3 and Tables S10a-S10d).",
        "The two timing profiles answer different questions and are not pooled "
        "(Supplementary Figure S3 and Tables S10a-S10d). The broader archived "
        "workflow panel compares fastPLS argmax and LDA with independent R "
        "implementations across nine classification datasets (Figure 2; "
        "Supplementary Table S10e).",
    )

    section_33 = paragraph_by_prefix(document, "3.3 Internal acceleration")
    insert_figure_before(
        document,
        section_33,
        FIGURE_2,
        (
            "Figure 2. Archived float64 single-CPU SIMPLS classification "
            "workflows. Panels report fixed-split accuracy, total fitting-plus-"
            "prediction time, and absolute complete-process peak RSS from three "
            "isolated runs; NE denotes not evaluated. Argmax rows provide the "
            "closest estimator comparison; LDA and other PLS-DA rows compare "
            "complete workflows with different heads or retained outputs."
        ),
        width=5.4,
        caption_style=caption_style,
        page_break=True,
    )

    hardware = paragraph_by_prefix(document, "Hardware acceleration remained")
    hardware.text = (
        "Hardware acceleration remained route and shape dependent. In the "
        "archived real-dataset analysis, paired CPU/CUDA and CPU/Metal timing was "
        "interpreted only when the absolute predictive-metric difference was at "
        "most 0.005 and paired-prediction agreement was at least 0.995 (Figure 3). "
        "Among 44 CPU/CUDA pairs, 28 met both criteria and CUDA was faster in "
        "seven, including an 8.90-fold PLS-SVD ratio on CIFAR-100. Six of 12 "
        "CPU/Metal pairs met both criteria, but none favored Metal. Gray cells "
        "retain discordant routes without converting them into speed claims. A "
        "separate frozen-release SIMPLS-rSVD comparison is reported in "
        "Supplementary Figure S4 and Table S11. CPU IRLBA remains the "
        "deterministic numerical reference, and float32 does not uniformly "
        "improve runtime, process memory, or agreement (Supplementary Tables "
        "S8-S9)."
    )

    insert_figure_before(
        document,
        hardware,
        FIGURE_3,
        (
            "Figure 3. Archived CPU/accelerator runtime ratios for numerically "
            "concordant workflows. Ratios are CPU time divided by accelerator "
            "time, so values above one favor CUDA or Metal and values below one "
            "indicate accelerator slowdown. Cells are colored only when the "
            "absolute paired predictive-metric difference is at most 0.005 and "
            "prediction agreement is at least 0.995. Gray cells are retained as "
            "metric- or prediction-discordant and excluded from acceleration "
            "claims."
        ),
        width=6.8,
        caption_style=caption_style,
    )

    replace_text(
        document,
        [
            ("Figure 3 displays held-out sample", "Figure 4 displays held-out sample"),
            ("Figure 3. Archived-release NMR", "Figure 4. Archived-release NMR"),
        ],
    )

    imagenet = paragraph_by_prefix(document, "Foundation-model embeddings are")
    imagenet.text = (
        "Foundation-model embeddings are increasingly relevant to medical-image "
        "pipelines because downstream analysis may involve millions of dense tile "
        "representations. As a historical engineering example, archived fastPLS "
        "0.99.25 completed label-aware float32 SIMPLS fitting on 1,000,000 stored "
        "ImageNet/DINOv2 rows and blocked prediction of 281,167 held-out rows "
        "(Figure 5). Argmax and LDA were evaluated across 100-1,000 requested "
        "components; one shared fit supplied all prefixes. These single-run, "
        "noncanonical-split results demonstrate matrix-processing feasibility, "
        "not reproducible feature extraction, biomedical utility, or an optimized "
        "ImageNet classifier. The 1,000-component point is a boundary stress point. "
        "Full values and provenance limitations are reported in Supplementary "
        "Section S18 and Table S13."
    )

    discussion = paragraph_by_prefix(document, "4. Discussion")
    insert_figure_before(
        document,
        discussion,
        FIGURE_5,
        (
            "Figure 5. Historical, partially reproducible ImageNet/DINOv2 "
            "downstream SIMPLS feasibility experiment. fastPLS 0.99.25 reran "
            "label-aware float32 CUDA-rSVD fitting and blocked prediction on "
            "1,000,000 training and 281,167 held-out stored embeddings. Panels "
            "report single-run top-1 and top-5 accuracy for argmax and LDA, "
            "shared-path fitting and prediction time, and complete-process host "
            "and device-memory measurements across 100-1,000 component prefixes. "
            "rSVD used oversampling 20, two power iterations, and seed 123. The "
            "split was noncanonical; checkpoint, pooling, extraction, and complete "
            "image-to-row provenance were unavailable. The figure is an "
            "engineering stress test, not biomedical or representation-quality "
            "validation, and 1,000 components is a boundary stress point rather "
            "than an optimum."
        ),
        width=6.8,
        caption_style=caption_style,
        page_break=True,
    )

    document.save(MAIN_OUTPUT)


def rewrite_supplement():
    document = Document(SUPP_SOURCE)
    caption_style = paragraph_by_prefix(document, "Figure S4.").style

    # ImageNet is promoted to the main text. Preserve the more focused current-
    # release CPU/CUDA SIMPLS panel here as the new Figure S4.
    remove_figure_and_caption(paragraph_by_prefix(document, "Figure S4."))
    section_17 = paragraph_by_prefix(document, "S17. NMR case study")
    insert_figure_before(
        document,
        section_17,
        FIGURE_S4,
        (
            "Figure S4. Frozen-release paired CPU/CUDA SIMPLS-rSVD workflows "
            "for MetRef, Retina, and CIFAR-100. Panels show total runtime, the "
            "post-prediction minus pre-fit host-RSS snapshot, and held-out "
            "accuracy. Points and IQRs summarize three runs. Every route used "
            "fastPLS 0.99.25, oversampling 20, two power iterations, seed 123, "
            "and identical splits and component counts. Device/context allocation "
            "is included; host-memory differences are not isolated solver "
            "workspace."
        ),
        width=6.8,
        caption_style=caption_style,
    )

    replace_text(
        document,
        [
            ("The repeated repeated comparison", "The repeated comparison"),
            ("Figure 3 displays AMI-0030-9", "Figure 4 displays AMI-0030-9"),
            (
                "Figure S4. Historical, partially reproducible ImageNet/DINOv2 downstream SIMPLS feasibility experiment.",
                "Figure 5 in the main text presents the historical, partially reproducible ImageNet/DINOv2 downstream SIMPLS feasibility experiment.",
            ),
        ],
    )

    image_section = paragraph_by_prefix(document, "The historical archive contained")
    image_section.text = image_section.text + (
        " The accuracy, runtime, and memory visualization is presented as main-text Figure 5."
    )

    document.save(SUPP_OUTPUT)


def main():
    required = [MAIN_SOURCE, SUPP_SOURCE, FIGURE_2, FIGURE_3, FIGURE_5, FIGURE_S4]
    missing = [str(path) for path in required if not path.is_file()]
    if missing:
        raise FileNotFoundError("Missing required files:\n" + "\n".join(missing))
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    rewrite_main()
    rewrite_supplement()
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)


if __name__ == "__main__":
    main()
