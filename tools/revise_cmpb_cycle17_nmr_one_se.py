#!/usr/bin/env python3

from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path
from shutil import copy2
import subprocess

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Inches


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle16"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle17"
EVIDENCE = (
    ROOT / "benchmark_results" / "manuscript_revision_cycle17_20260725"
)
PLOT_SCRIPT = ROOT / "tools" / "plot_cmpb_cycle17_nmr_one_se.R"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle16_0.99.6_20260725.docx"
SUPP_SOURCE = (
    SOURCE / "fastPLS_CMPB_supplement_cycle16_0.99.6_20260725.docx"
)
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle16_20260725.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle17_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle17_0.99.6_20260725.docx"
RESPONSE_OUT = (
    OUT / "fastPLS_CMPB_response_to_reviewers_cycle17_20260725.docx"
)

SIMPLS_RAW = (
    ROOT
    / "benchmark_results"
    / "review_nmr_extended_selection_20260725"
    / "nmr_component_selection_raw.csv"
)
PLSSVD_RAW = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle16_20260725"
    / "nmr_plssvd_extended_lower_grid"
    / "results"
    / "nmr_component_selection_raw.csv"
)
HELDOUT_DIR = EVIDENCE / "nmr_simpls_one_se" / "heldout" / "rows"
SELECTED_SOURCE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle16_20260725"
    / "selected_backend_cycle16_chosen.csv"
)
STATUS_SOURCE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle16_20260725"
    / "component_selection_status.csv"
)
ALL_DATASET_FIGURE = EVIDENCE / "plots" / "selected_performance_all_datasets.png"
NMR_SELECTION_FIGURE = EVIDENCE / "plots" / "nmr_component_selection_one_se.png"

spec = spec_from_file_location(
    "cycle16_helpers", ROOT / "tools" / "revise_cmpb_cycle16_component_boundaries.py"
)
c16 = module_from_spec(spec)
spec.loader.exec_module(c16)


def summarize(raw):
    ok = raw.loc[raw["status"].eq("ok")].copy()
    summary = (
        ok.groupby("ncomp", as_index=False)["RMSD"]
        .agg(["mean", "std", "median", "count"])
        .reset_index()
    )
    summary["se"] = summary["std"] / np.sqrt(summary["count"])
    quantiles = (
        ok.groupby("ncomp")["RMSD"]
        .quantile([0.25, 0.75])
        .unstack()
        .rename(columns={0.25: "q25", 0.75: "q75"})
        .reset_index()
    )
    summary = summary.merge(quantiles, on="ncomp", how="left")
    minimizing = summary.loc[summary["mean"].idxmin()]
    threshold = minimizing["mean"] + minimizing["se"]
    summary["within_one_se"] = summary["mean"].le(threshold)
    selected = int(summary.loc[summary["within_one_se"], "ncomp"].min())
    return ok, summary, int(minimizing["ncomp"]), float(threshold), selected


def paired_differences(raw, reference_k=100):
    ok = raw.loc[raw["status"].eq("ok")].copy()
    reference = (
        ok.loc[ok["ncomp"].eq(reference_k), ["split", "RMSD"]]
        .rename(columns={"RMSD": "reference_RMSD"})
    )
    rows = []
    for component in sorted(ok["ncomp"].unique()):
        values = (
            ok.loc[ok["ncomp"].eq(component), ["split", "RMSD"]]
            .merge(reference, on="split", how="inner")
        )
        difference = values["RMSD"] - values["reference_RMSD"]
        n = len(difference)
        standard_error = difference.std(ddof=1) / np.sqrt(n)
        t_critical = 2.7764451051977987 if n == 5 else np.nan
        rows.append(
            {
                "ncomp": int(component),
                "reference_ncomp": reference_k,
                "mean_paired_difference": difference.mean(),
                "paired_se": standard_error,
                "ci95_lower": difference.mean() - t_critical * standard_error,
                "ci95_upper": difference.mean() + t_critical * standard_error,
                "splits_better_than_reference": int((difference < 0).sum()),
                "n_paired_splits": n,
            }
        )
    return pd.DataFrame(rows)


def update_selected(heldout):
    selected = pd.read_csv(SELECTED_SOURCE)
    mask = selected["dataset"].eq("nmr") & selected["method_panel"].eq("simpls")
    selected.loc[mask, "variant_name"] = "fastpls_simpls_cuda_rsvd"
    selected.loc[mask, "engine"] = "CUDA"
    selected.loc[mask, "backend"] = "cuda_rsvd"
    selected.loc[mask, "requested_ncomp"] = 50
    selected.loc[mask, "effective_ncomp"] = 50
    selected.loc[mask, "n_runs"] = len(heldout)
    fields = {
        "metric_median": ("RMSD", "median"),
        "metric_q25": ("RMSD", "q25"),
        "metric_q75": ("RMSD", "q75"),
        "total_time_sec_median": ("total_time_sec", "median"),
        "total_time_sec_q25": ("total_time_sec", "q25"),
        "total_time_sec_q75": ("total_time_sec", "q75"),
        "host_rss_mb_median": ("host_rss_mb", "median"),
        "host_rss_mb_q25": ("host_rss_mb", "q25"),
        "host_rss_mb_q75": ("host_rss_mb", "q75"),
        "gpu_mem_mb_median": ("gpu_peak_mb", "median"),
        "gpu_mem_mb_q25": ("gpu_peak_mb", "q25"),
        "gpu_mem_mb_q75": ("gpu_peak_mb", "q75"),
    }
    for destination, (source, statistic) in fields.items():
        value = (
            heldout[source].median()
            if statistic == "median"
            else heldout[source].quantile(0.25 if statistic == "q25" else 0.75)
        )
        selected.loc[mask, destination] = value
    selected.loc[mask, "selection_status"] = "interior tested value"
    selected.to_csv(EVIDENCE / "selected_backend_cycle17_chosen.csv", index=False)

    status = pd.read_csv(STATUS_SOURCE)
    status.loc[status["dataset"].eq("nmr"), "simpls"] = (
        "k=50; interior tested value; one-SE rule"
    )
    status.to_csv(EVIDENCE / "component_selection_status.csv", index=False)
    return selected, status


def build_evidence():
    EVIDENCE.mkdir(parents=True, exist_ok=True)
    simpls = pd.read_csv(SIMPLS_RAW)
    plssvd = pd.read_csv(PLSSVD_RAW)
    heldout = pd.concat(
        [pd.read_csv(path) for path in sorted(HELDOUT_DIR.glob("*.csv"))],
        ignore_index=True,
    )
    simpls_ok, simpls_summary, simpls_min, simpls_threshold, simpls_selected = (
        summarize(simpls)
    )
    plssvd_ok, plssvd_summary, plssvd_min, plssvd_threshold, plssvd_selected = (
        summarize(plssvd)
    )
    assert simpls_selected == 50
    assert plssvd_selected == 5

    simpls_summary.to_csv(EVIDENCE / "nmr_simpls_one_se_summary.csv", index=False)
    plssvd_summary.to_csv(EVIDENCE / "nmr_plssvd_one_se_summary.csv", index=False)
    paired = paired_differences(simpls)
    paired.to_csv(EVIDENCE / "nmr_simpls_paired_differences.csv", index=False)
    split_minima = (
        simpls_ok.loc[simpls_ok.groupby("split")["RMSD"].idxmin()]
        [["split", "split_seed", "ncomp", "RMSD"]]
        .rename(columns={"ncomp": "split_minimizing_ncomp"})
    )
    split_minima.to_csv(EVIDENCE / "nmr_simpls_split_minima.csv", index=False)
    decision = pd.DataFrame(
        [
            {
                "family": "PLS-SVD",
                "mean_minimizing_ncomp": plssvd_min,
                "one_se_threshold": plssvd_threshold,
                "one_se_eligible_ncomp": ",".join(
                    str(int(x))
                    for x in plssvd_summary.loc[
                        plssvd_summary["within_one_se"], "ncomp"
                    ]
                ),
                "selected_ncomp": plssvd_selected,
            },
            {
                "family": "SIMPLS",
                "mean_minimizing_ncomp": simpls_min,
                "one_se_threshold": simpls_threshold,
                "one_se_eligible_ncomp": ",".join(
                    str(int(x))
                    for x in simpls_summary.loc[
                        simpls_summary["within_one_se"], "ncomp"
                    ]
                ),
                "selected_ncomp": simpls_selected,
            },
        ]
    )
    decision.to_csv(EVIDENCE / "nmr_one_se_decision.csv", index=False)
    selected, status = update_selected(heldout)
    return {
        "simpls_summary": simpls_summary,
        "plssvd_summary": plssvd_summary,
        "paired": paired,
        "split_minima": split_minima,
        "decision": decision,
        "heldout": heldout,
        "selected": selected,
        "status": status,
    }


def nmr_table_rows(evidence):
    simpls = evidence["simpls_summary"].set_index("ncomp")
    plssvd = evidence["plssvd_summary"].set_index("ncomp")
    components = sorted(set(simpls.index) | set(plssvd.index))

    def format_cells(frame, component):
        if component not in frame.index:
            return "\u2014", "\u2014"
        row = frame.loc[component]
        return (
            f'{row["mean"]:.7f} ({row["se"]:.7f})',
            f'{row["median"]:.7f} ({row["q25"]:.7f}-{row["q75"]:.7f})',
        )

    rows = []
    for component in components:
        simpls_mean, simpls_median = format_cells(simpls, component)
        plssvd_mean, plssvd_median = format_cells(plssvd, component)
        rows.append(
            (
                component,
                simpls_mean,
                simpls_median,
                plssvd_mean,
                plssvd_median,
            )
        )
    return rows


def revise_main(evidence):
    document = Document(MAIN_SOURCE)
    nmr_row = next(
        row for row in document.tables[0].rows if row.cells[0].text.strip() == "NMR"
    )
    nmr_row.cells[2].text = (
        "CUDA rSVD | k=50\n"
        "RMSD=7.59e-04 | 1.97 s (IQR 0.01)\n"
        "H=3250; G=3414 MB | f64 | n=3"
    )
    c16.style_table(document.tables[0], font_size=5.3)
    c16.replace_picture_before_caption(
        document, "Figure 2.", ALL_DATASET_FIGURE, Inches(6.65)
    )

    methods = c16.find_paragraph(document, "Within each dataset")
    methods.text = (
        methods.text
        + " For NMR SIMPLS, mean RMSD and its standard error were calculated "
        "across the same five paired training-only splits. The smallest "
        "component count whose mean was no more than one standard error above "
        "the minimum mean was selected."
    )
    c16.set_paragraph(
        document,
        "NMR represented the extreme",
        "NMR represented the extreme multivariate-response setting (1,200 "
        "training and 321 held-out spectra; p=13,000; q=28,355). Every method "
        "used the predefined outer split and full response spectrum; predictors "
        "were centred without variance scaling, and the routine 4.6-4.8 ppm "
        "residual-water interval was set to zero in both training and held-out "
        "predictor matrices. Five repeated training-only 80/20 splits were used "
        "for component selection. PLS-SVD retained a unique one-standard-error "
        "selection at 5 components after extending its grid below 10. For "
        "SIMPLS, split-specific minima occurred at 25, 50, 75, 75, and 100 "
        "components. Mean validation RMSD was lowest at 50 components "
        "(0.0009195; SE 0.0000370); 50, 75, and 100 all lay within one standard "
        "error of that minimum. The parsimonious one-standard-error rule "
        "therefore selected 50 components. The paired mean RMSD difference "
        "between 50 and 100 components was -2.02e-05 (95% CI -6.27e-05 to "
        "2.22e-05), so 100 is not interpreted as a unique optimum. At the "
        "selected 50-component value, CUDA SIMPLS achieved held-out RMSD "
        "0.0007591 with median total time 1.971 s across three repetitions. "
        "Figure 3 retains the common prespecified 100-component operating point "
        "only to compare implementations at equal model size.",
    )
    c16.find_paragraph(document, "NMR represented the extreme").add_run().add_break(
        WD_BREAK.PAGE
    )
    document.core_properties.title = (
        "fastPLS CMPB manuscript - NMR one-standard-error revision cycle 17"
    )
    document.save(MAIN_OUT)


def revise_supplement(evidence):
    document = Document(SUPP_SOURCE)
    c16.set_paragraph(
        document,
        "Routine NMR spectral preprocessing and component selection.",
        "Routine NMR spectral preprocessing and component selection. The task "
        "comprised 1,200 training and 321 held-out spectra, 13,000 predictors, "
        "and 28,355 numeric responses. The 4.6-4.8 ppm residual-water interval "
        "was set to zero in Xtrain and Xtest; Ytrain and Ytest were unchanged. "
        "Predictors were centred without variance scaling. Five repeated 80/20 "
        "training-only splits used seeds 123, 456, 789, 1011, and 2027. At each "
        "component count, mean RMSD and SE were calculated across the paired "
        "splits. The one-standard-error threshold was the minimum mean RMSD plus "
        "its SE, and the smallest eligible component count was selected. This "
        "selected 5 components for PLS-SVD and 50 for SIMPLS.",
    )
    c16.set_paragraph(
        document,
        "One maximal SIMPLS-rSVD model",
        "One maximal SIMPLS-rSVD model was fitted per inner split, and requested "
        "component prefixes were evaluated in response blocks. Split-specific "
        "SIMPLS minima were 100, 75, 25, 75, and 50 components. The mean curve "
        "was minimized at 50 components (0.0009195; SE 0.0000370), and the "
        "one-standard-error set comprised 50, 75, and 100. The paired 50-minus-"
        "100 RMSD difference was -2.02e-05 (95% t interval -6.27e-05 to "
        "2.22e-05; 50 was lower in four of five splits). The corresponding "
        "75-minus-100 difference was -1.60e-05 (95% interval -5.35e-05 to "
        "2.16e-05). These data support a stable 50-100 component region, not a "
        "unique optimum at 100; the smallest one-SE candidate, 50, was selected.",
    )
    c16.set_paragraph(
        document,
        "Table S12b.",
        "Table S12b. Family-specific repeated training-only NMR component "
        "selection. Each cell gives mean RMSD (SE), followed by median RMSD "
        "(interquartile range), over the same five paired splits.",
    )
    c16.fill_table(
        document.tables[5],
        [
            "Components",
            "SIMPLS mean (SE)",
            "SIMPLS median (IQR)",
            "PLS-SVD mean (SE)",
            "PLS-SVD median (IQR)",
        ],
        nmr_table_rows(evidence),
        font_size=5.3,
    )
    c16.replace_picture_before_caption(
        document, "Figure S18.", NMR_SELECTION_FIGURE, Inches(6.65)
    )
    c16.set_paragraph(
        document,
        "Figure S18.",
        "Figure S18. Repeated training-only NMR component selection. Grey lines "
        "show the five paired splits; coloured lines and ribbons show mean RMSD "
        "and one standard error. The dotted horizontal line is the one-SE "
        "threshold, yellow squares mark eligible values, and the dashed "
        "vertical line marks the smallest eligible component count. (A) "
        "PLS-SVD selects 5 components. (B) SIMPLS has an eligible 50-100 range "
        "and selects 50 components; the figure does not support a unique optimum "
        "at 100.",
    )
    boundary = c16.find_paragraph(
        document, "Component counts are reported as the best"
    )
    boundary.text = (
        "Component counts are reported as training-selected values within "
        "prespecified grids. Lower and upper tested-grid boundaries and "
        "PLS-SVD response-rank limits are labelled explicitly. NMR PLS-SVD "
        "selected 5 components; NMR SIMPLS selected the smallest value in its "
        "one-standard-error set, 50. With only five paired splits, the error "
        "differences among 50, 75, and 100 components were not resolved."
    )
    nmr_status_row = next(
        row for row in document.tables[29].rows
        if row.cells[0].text.strip() == "NMR"
    )
    nmr_status_row.cells[3].text = (
        "k=50; interior tested value; one-SE rule (eligible 50,75,100)"
    )
    c16.style_table(document.tables[29], font_size=5.2)
    document.core_properties.title = (
        "fastPLS CMPB supplement - NMR one-standard-error revision cycle 17"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    c16.set_paragraph(
        document,
        "SIMPLS was evaluated over 10-300",
        "SIMPLS component selection has been reanalysed using the paired "
        "one-standard-error rule. Split-specific minima ranged from 25 to 100 "
        "components. Mean RMSD was minimized at 50, and 50, 75, and 100 were "
        "within one standard error; 50 was therefore selected as the "
        "parsimonious value. PLS-SVD retains its interior 5-component "
        "selection after extension below 10. All claims now refer to selected "
        "values within evaluated grids rather than global optima.",
    )
    document.add_heading("18. The NMR optimum was overstated", level=1)
    document.add_paragraph(
        "Reviewer comment: SIMPLS split-specific minima range from 25 to 100 "
        "components, and the aggregate differences around 75-100 are small "
        "relative to dispersion. A one-standard-error rule, paired uncertainty "
        "analysis, or more cautious interpretation is required."
    )
    document.add_paragraph(
        "Response: Corrected. We calculated mean RMSD and SE across the same "
        "five paired training-only splits. The minimum mean occurred at 50 "
        "components, and the one-SE set was 50, 75, and 100; the smallest value, "
        "50, was selected. The paired 50-minus-100 difference was -2.02e-05 "
        "(95% CI -6.27e-05 to 2.22e-05), which does not establish a unique "
        "100-component optimum. The manuscript, Table 1, Figure 2, Table S12b, "
        "Figure S18, and Table S29 now report the parsimonious 50-component "
        "selection and describe 50-100 as a stable low-error region. A new "
        "three-repetition held-out CUDA SIMPLS evaluation at 50 components "
        "gave RMSD 0.0007591 and median total time 1.971 s."
    )
    document.core_properties.title = (
        "fastPLS CMPB response - NMR one-standard-error revision cycle 17"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    evidence = build_evidence()
    subprocess.run(["Rscript", str(PLOT_SCRIPT)], check=True)
    revise_main(evidence)
    revise_supplement(evidence)
    revise_response()
    for source in (
        EVIDENCE / "selected_backend_cycle17_chosen.csv",
        EVIDENCE / "component_selection_status.csv",
        EVIDENCE / "nmr_simpls_one_se_summary.csv",
        EVIDENCE / "nmr_simpls_paired_differences.csv",
        EVIDENCE / "nmr_simpls_split_minima.csv",
        EVIDENCE / "nmr_one_se_decision.csv",
        ALL_DATASET_FIGURE,
        NMR_SELECTION_FIGURE,
    ):
        copy2(source, OUT / source.name)
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
