#!/usr/bin/env python3
"""Provide an executable NMR one-standard-error component-selection rule."""

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260727_cycle82"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260727_cycle83"
MAIN_IN = SOURCE / "fastPLS_CMPB_main_cycle82_0.99.6_20260727.docx"
SUPP_IN = SOURCE / "fastPLS_CMPB_supplement_cycle82_0.99.6_20260727.docx"
MAIN_OUT = OUTPUT / "fastPLS_CMPB_main_cycle83_0.99.6_20260727.docx"
SUPP_OUT = OUTPUT / "fastPLS_CMPB_supplement_cycle83_0.99.6_20260727.docx"


def insert_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    if paragraph.style is not None:
        result.style = paragraph.style
    result.add_run(text)
    return result


def replace_prefix(document, prefix, replacement):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            paragraph.text = replacement
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def revise_main():
    document = Document(MAIN_IN)
    replace_prefix(
        document,
        "NMR comprised 1,200 training and 321 held-out spectra,",
        (
            "NMR comprised 1,200 training and 321 held-out spectra, with "
            "13,000 predictors and 28,355 responses. Predictor columns between "
            "4.6 and 4.8 ppm were zeroed in training and test data as standard "
            "water-region preprocessing; responses were unmasked. Training-only "
            "one-standard-error selection retained five PLS-SVD and 50 SIMPLS "
            "components. These are family-specific predictive settings rather "
            "than a matched implementation comparison. The complete candidate "
            "grids and executable selection rule are reported in Supplementary "
            "Section S17."
        ),
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_IN)
    anchor = replace_prefix(
        document,
        "The family-selected predictive analysis and paired backend analysis",
        (
            "The family-selected predictive analysis and paired backend "
            "analysis answer different questions. Component selection used five "
            "paired training-only splits. The PLS-SVD candidate grid was "
            "1, 2, 3, 5, 7, 10, 25, 50, 75, 100, 125, 150, 165, 175, 200, "
            "250, and 300 components. The SIMPLS candidate grid was 10, 25, "
            "50, 75, 100, 125, 150, 165, 175, 200, 250, and 300 components."
        ),
    )
    rule = insert_after(
        anchor,
        (
            "For family f, component count a, and split s = 1,...,5, let "
            "e_f,s(a) denote validation RMSD. The implementation computes "
            "m_f(a) = (1/5) sum_s e_f,s(a), chooses the smallest "
            "a_min in argmin_a m_f(a), and defines the one-standard-error "
            "threshold h_f = m_f(a_min) + sd_s[e_f,s(a_min)]/sqrt(5), where "
            "sd is the sample standard deviation across the five splits. The "
            "eligible set is E_f = {a: m_f(a) <= h_f}, and the returned count "
            "is a_selected = min(E_f)."
        ),
    )
    decision = insert_after(
        rule,
        (
            "For PLS-SVD, a_min = 5, h_f = 0.001528979123, and E_f = {5}; "
            "therefore a_selected = 5. For SIMPLS, a_min = 50, "
            "h_f = 0.0009565399754, and E_f = {50, 75, 100}; therefore "
            "a_selected = 50. Each value is selected within the evaluated grid, "
            "not asserted to be a global optimum."
        ),
    )
    insert_after(
        decision,
        (
            "The paired backend analysis changes only CPU versus CUDA within "
            "family. The deposited 165-component workflow uses the original "
            "centring-only protocol and is historical context. Predictor "
            "columns with chemical shifts strictly between 4.6 and 4.8 ppm "
            "were set to zero in both training and test predictor matrices "
            "before inner splitting or fitting. No response column was zeroed, "
            "masked, or excluded; all response metrics use all 28,355 "
            "coordinates. Main-text Figure 4 displays held-out sample "
            "AMI-00BP-8 (index 155), selected by the prespecified descriptive "
            "rule of SIMPLS RMSD closest to the held-out median; it was not the "
            "best-predicted spectrum."
        ),
    )
    document.save(SUPP_OUT)


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    copy2(MAIN_IN, MAIN_OUT)
    copy2(SUPP_IN, SUPP_OUT)
    revise_main()
    revise_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
