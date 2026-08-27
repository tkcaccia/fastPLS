from pathlib import Path

from docx import Document


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts/CMPB_rewrite_20260826_cycle108/fastPLS_CMPB_main_cycle108_0.99.25_20260826.docx"
OUTDIR = ROOT / "artifacts/CMPB_rewrite_20260826_cycle109"
OUTPUT = OUTDIR / "fastPLS_CMPB_main_cycle109_0.99.25_20260826.docx"


ABSTRACT_RESULTS = (
    "Results: An independent dense-LAPACK de Jong panel completed all 82 "
    "component-prefix comparisons without numerical failure. Outside an intentionally "
    "near-tied singular-value case, maximum held-out prediction error was 4.53 x "
    "10^-15; the tied case retained a 0.015-degree maximum subspace angle and 3.36 x "
    "10^-4 relative prediction error. Against pls::simpls.fit, fastPLS was faster on "
    "five of nine datasets in repeated ordinary public-workflow comparisons, with "
    "identical accuracy and a maximum 4.85-fold speed-up. Across the broader R-package "
    "panel, fastPLS was the fastest tested workflow on seven of nine classification "
    "datasets. fastPLS also completed the 13,000 by 28,355-response NMR benchmark and "
    "the 1,000,000-row ImageNet/DINOv2 stress test; no tested external R workflow "
    "completed ImageNet, and only limited external routes were feasible for NMR. At 50 "
    "NMR SIMPLS components, CPU IRLBA, CPU rSVD, and CUDA rSVD required 692.21, 152.09, "
    "and 3.91 s, respectively, with RMSD 0.0007561, 0.0007560, and 0.0007561."
)


STRICT_RESULTS_OLD = (
    "The strict comparison completed all 108 planned runs: nine datasets, two output "
    "profiles, two implementations, and three fresh-process repetitions. Accuracy was "
    "identical for every pair. With minimum common prediction outputs, fastPLS was faster "
    "on four datasets and pls::simpls.fit on five; the largest fastPLS advantage was "
    "1.48-fold on GTEx v8. Under ordinary public workflows, fastPLS was faster on five "
    "datasets, including 2.39-fold on CIFAR-100, 3.35-fold on Retina, and 4.85-fold on "
    "Tabula Muris. Corresponding exact held-out counts were 8,739/10,000, 21,684/22,406, "
    "and 40,077/50,059. These are conditional row-level computational endpoints. In "
    "particular, cell-level Retina and Tabula Muris rows are not independent biological "
    "replicates, so binomial intervals are not interpreted as biological generalization "
    "intervals. The two timing profiles answer different questions and are not pooled "
    "(Supplementary Figure S3 and Tables S10a-S10d). The broader archived workflow panel "
    "compares fastPLS argmax and LDA with independent R implementations across nine "
    "classification datasets (Figure 2; Supplementary Table S10e)."
)


STRICT_RESULTS_NEW = (
    "The strict comparison completed all 108 planned runs: nine datasets, two output "
    "profiles, two implementations, and three fresh-process repetitions. Accuracy was "
    "identical for every pair. With minimum common prediction outputs, fastPLS was faster "
    "on four datasets and pls::simpls.fit on five; the largest fastPLS advantage was "
    "1.48-fold on GTEx v8. Under ordinary public workflows, fastPLS was faster than "
    "pls::simpls.fit on five datasets, including 2.39-fold on CIFAR-100, 3.35-fold on "
    "Retina, and 4.85-fold on Tabula Muris. Corresponding exact held-out counts were "
    "8,739/10,000, 21,684/22,406, and 40,077/50,059. These are conditional row-level "
    "computational endpoints. In particular, cell-level Retina and Tabula Muris rows are "
    "not independent biological replicates, so binomial intervals are not interpreted as "
    "biological generalization intervals. The two timing profiles answer different "
    "questions and are not pooled (Supplementary Figure S3 and Tables S10a-S10d). In the "
    "broader archived workflow panel, which compared fastPLS argmax and LDA with multiple "
    "independent R implementations, fastPLS had the lowest observed total time on seven "
    "of nine classification datasets. The exceptions were TCGA-BRCA and TCGA-HNSC, where "
    "plsgenomics was faster (Figure 2; Supplementary Table S10e)."
)


def replace_exact(document: Document, old: str, new: str) -> None:
    matches = [p for p in document.paragraphs if p.text == old]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one exact paragraph match, found {len(matches)}")
    matches[0].text = new


def main() -> None:
    OUTDIR.mkdir(parents=True, exist_ok=True)
    document = Document(SOURCE)
    abstract_matches = [p for p in document.paragraphs if p.text.startswith("Results: An independent dense-LAPACK")]
    if len(abstract_matches) != 1:
        raise RuntimeError(f"Expected one abstract Results paragraph, found {len(abstract_matches)}")
    abstract_matches[0].text = ABSTRACT_RESULTS
    replace_exact(document, STRICT_RESULTS_OLD, STRICT_RESULTS_NEW)
    document.save(OUTPUT)


if __name__ == "__main__":
    main()
