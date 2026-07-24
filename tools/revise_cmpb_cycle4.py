"""Create cycle-4 CMPB drafts after the SIMPLS execution ablation."""

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Desktop/fastPLS/CMPB_rewrite_20260724_cycle3")
OUT = Path("/Users/stefano/Desktop/fastPLS/CMPB_rewrite_20260724_cycle4")
OUT.mkdir(parents=True, exist_ok=True)
MAIN_SRC = ROOT / "fastPLS_CMPB_main_cycle3_0.99.8_20260724.docx"
SUPP_SRC = ROOT / "fastPLS_CMPB_supplement_cycle3_0.99.8_20260724.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle4_0.99.9_20260724.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle4_0.99.9_20260724.docx"
REVIEW_OUT = OUT / "fastPLS_CMPB_independent_reviewer_report_cycle4_20260724.docx"


def paragraph_after(paragraph, text, style=None):
    node = OxmlElement("w:p")
    paragraph._p.addnext(node)
    inserted = Paragraph(node, paragraph._parent)
    if style:
        inserted.style = style
    inserted.add_run(text)
    return inserted


def insert_after(doc, startswith, text):
    for p in doc.paragraphs:
        if p.text.startswith(startswith):
            return paragraph_after(p, text)
    raise RuntimeError(f"No paragraph starts with: {startswith}")


def revise_main():
    copy2(MAIN_SRC, MAIN_OUT)
    doc = Document(MAIN_OUT)
    insert_after(
        doc,
        "A controlled numerical check compared compiled deterministic IRLBA SIMPLS",
        "A targeted rSVD execution ablation was performed separately on preprocessed MetRef (22 components) and a fixed stratified CIFAR-100 subset (5,000 training and 1,000 test observations; 100 components). "
        "On CIFAR-100, the default rSVD workspace-reuse route had median total time 0.926 s (IQR 0.057 s) and accuracy 0.880, compared with 1.265 s (IQR 0.101 s) and accuracy 0.876 when workspace reuse was disabled. "
        "Prediction agreement between these two approximate rSVD executions was 0.984. Disabling cached deflation preserved predictions in this screen and did not provide a reliable speed advantage. Thus, workspace reuse is reported as an approximate rSVD execution option, whereas deterministic IRLBA is the estimator-preserving implementation."
    )
    doc.save(MAIN_OUT)


def revise_supplement():
    copy2(SUPP_SRC, SUPP_OUT)
    doc = Document(SUPP_OUT)
    anchor = None
    for p in doc.paragraphs:
        if p.text.startswith("Estimator agreement. On fixed synthetic data"):
            anchor = p
            break
    if anchor is None:
        raise RuntimeError("Estimator-agreement anchor not found")
    paragraph_after(
        anchor,
        "Execution ablation. The rSVD workspace-reuse and cached-deflation controls were examined on preprocessed MetRef (22 components) and a fixed stratified CIFAR-100 subset (5,000/1,000 train/test observations; 100 components), with three repetitions and seed 123. On CIFAR-100, default rSVD workspace reuse had median total time 0.926 s (IQR 0.057 s), accuracy 0.880, and agreement 1.000 with the default reference run. Disabling workspace reuse gave 1.265 s (IQR 0.101 s), accuracy 0.876, and prediction agreement 0.984. Disabling cached deflation gave 0.997 s (IQR 0.218 s), accuracy 0.880, and agreement 1.000. On MetRef, execution times were too short and variable for meaningful speed ranking; workspace-reuse and no-reuse accuracies were 0.920 and 0.900, respectively. These results show that rSVD workspace reuse is an accuracy-speed approximation, not a deterministic-equivalence mechanism."
    )
    doc.save(SUPP_OUT)


def create_review():
    doc = Document()
    doc.add_heading("Independent reviewer report: cycle 4 update", level=0)
    doc.add_paragraph("Manuscript: fastPLS: scalable partial least squares with compiled CPU and accelerator backends for high-dimensional biomedical data")
    doc.add_heading("Recommendation", level=1)
    doc.add_paragraph("Major revision")
    doc.add_heading("Assessment after cycle 4", level=1)
    doc.add_paragraph(
        "The authors now provide an informative execution ablation. It correctly establishes that rSVD workspace reuse is a speed-oriented approximate solver strategy rather than proof of exact SIMPLS equivalence. The deterministic IRLBA route remains the appropriate evidence for estimator preservation."
    )
    doc.add_heading("Remaining major comments", level=1)
    for text in [
        "Regenerate the complete real-data benchmark after the corrected IRLBA SIMPLS dispatch and provide raw repetitions, dispersion, peak memory, failures, and requested versus executed estimator/backend in the final tables.",
        "Present the external-package comparison in precision-matched float64 form. Float32 should remain a secondary compatibility/footprint analysis until broader numerical and memory evidence is available.",
        "Provide a concise complexity and residency table for PLS-SVD, SIMPLS, OPLS, and kernel PLS on CPU, CUDA, and Metal, including fitting, decomposition, prediction, LDA, and cross-validation stages.",
        "The paper should identify a final immutable software release and archived data/benchmark manifest, and ensure that all results were generated from that revision.",
        "Maintain the explicit separation between biomedical validation and the ImageNet/DINOv2 computational stress test."
    ]:
        doc.add_paragraph(text, style="List Number")
    doc.add_heading("Minor comments", level=1)
    doc.add_paragraph("Retain the NMR water-mask clarification in all figure legends and state the GPU-memory sampling method next to the relevant table or figure.")
    doc.save(REVIEW_OUT)


revise_main()
revise_supplement()
create_review()
print(MAIN_OUT)
print(SUPP_OUT)
print(REVIEW_OUT)
