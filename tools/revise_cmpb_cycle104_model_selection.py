from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle103"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle104"
OUTPUT.mkdir(parents=True, exist_ok=True)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle103_0.99.25_20260825.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle103_0.99.25_20260825.docx"
MAIN_OUTPUT = OUTPUT / "fastPLS_CMPB_main_cycle104_0.99.25_20260825.docx"
SUPP_OUTPUT = OUTPUT / "fastPLS_CMPB_supplement_cycle104_0.99.25_20260825.docx"


def paragraph_by_prefix(document, prefix):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}; found {len(matches)}")
    return matches[0]


def insert_after(document, paragraph, text):
    new = document.add_paragraph()
    new.style = paragraph.style
    new.text = text
    paragraph._p.addnext(new._p)
    return new


def replace_everywhere(document, old, new):
    for paragraph in document.paragraphs:
        if old in paragraph.text:
            paragraph.text = paragraph.text.replace(old, new)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old in paragraph.text:
                        paragraph.text = paragraph.text.replace(old, new)


def update_main():
    document = Document(MAIN_SOURCE)

    metrics = paragraph_by_prefix(document, "Fold construction, fitting, prediction")
    metrics.text = (
        "Fold construction, fitting, prediction, and metric accumulation remain compiled where "
        "supported; grouped observations can be constrained to one fold. Classification tuning "
        "uses held-out accuracy by default or balanced accuracy when equal weighting of class "
        "recalls is required. Regression tuning uses held-out RMSD by default and can instead use "
        "observed-mean cross-validated R² or fold-training-mean cross-validated Q². Full-data "
        "training R² is returned only as a descriptive fit statistic and is never a component-"
        "selection endpoint. Independent-test Q² uses the complete training-response mean. For "
        "PLS-DA, training R² and cross-validated Q² are calculated on dummy-coded responses and "
        "are distinct from decoded-label accuracy and balanced accuracy. Exact formulas and "
        "denominator conventions are given in Supplementary Section S4."
    )

    benchmark = paragraph_by_prefix(document, "The broader benchmark design covered")
    benchmark.text = (
        "The broader benchmark design covered biomedical and computational tasks spanning "
        "metabolomics, NMR, CITE-seq, tissue and cancer omics, single-cell transcriptomics, drug "
        "response, and image embeddings [7,20-30]. Frozen central evidence comprises deterministic "
        "SIMPLS validation, external comparison, controlled scaling and solver qualification, "
        "selected CPU/CUDA routes, and the NMR case study. Methods used identical stored splits "
        "and training-only component grids within each analysis. Cross-dataset selected points are "
        "treated as fixed computational workloads when their grids were boundary or rank constrained. "
        "Runtime included fitting and prediction; memory definitions, split units, endpoint "
        "definitions, and uncertainty limitations are detailed in the Supplement."
    )

    selection = paragraph_by_prefix(document, "Five-fold training-only selection")
    selection.text = (
        "Five-fold training-only selection was performed separately by PLS family. Of 46 evaluated "
        "family-dataset choices, 24 were at a tested-grid boundary and nine were limited by response "
        "rank; these 33 choices are reported as constrained, not optimal. They define reproducible "
        "benchmark workloads and do not support unconstrained family-level predictive claims. The "
        "central NMR analysis used an expanded 1-300-component training-only grid and a one-standard-"
        "error rule, yielding interior family-specific settings. Accelerator speed-up was interpreted "
        "only for numerically concordant paired predictions."
    )

    components = paragraph_by_prefix(document, "Component counts are described")
    components.text = (
        "Component counts are described as selected within the evaluated training grids, not as "
        "unconstrained optima. The 24 boundary and nine rank-limited choices are used only as fixed "
        "computational workloads in the all-dataset benchmark. NMR grid extension and one-standard-"
        "error selection provide the principal component-sensitivity analysis; historical repeated-"
        "partition analyses for representative datasets are labelled supporting development evidence."
    )

    external = paragraph_by_prefix(document, "The strict comparison completed")
    external.text = (
        "The strict comparison completed all 108 planned runs: nine datasets, two output profiles, "
        "two implementations, and three fresh-process repetitions. Accuracy was identical for every "
        "pair. With minimum common prediction outputs, fastPLS was faster on four datasets and "
        "pls::simpls.fit on five; the largest fastPLS advantage was 1.48-fold on GTEx v8. Under "
        "ordinary public workflows, fastPLS was faster on five datasets, including 2.39-fold on "
        "CIFAR-100, 3.35-fold on Retina, and 4.85-fold on Tabula Muris. Corresponding exact held-out "
        "counts were 8,739/10,000, 21,684/22,406, and 40,077/50,059. These are conditional row-level "
        "computational endpoints. In particular, cell-level Retina and Tabula Muris rows are not "
        "independent biological replicates, so binomial intervals are not interpreted as biological "
        "generalization intervals. The two timing profiles answer different questions and are not "
        "pooled (Figure 2; Supplementary Tables S10a-S10d)."
    )

    nmr = paragraph_by_prefix(document, "NMR comprised 1,200 training")
    nmr.text = (
        "NMR comprised 1,200 training and 321 held-out spectra, with 13,000 predictors and 28,355 "
        "responses. Predictor columns between 4.6 and 4.8 ppm were zeroed in training and test data "
        "as standard water-region preprocessing; responses were unmasked. Training-only selection "
        "used five paired splits, PLS-SVD components 1, 2, 3, 5, 7, 10, 25, 50, 75, 100, 125, 150, "
        "165, 175, 200, 250, and 300, and SIMPLS components 10, 25, 50, 75, 100, 125, 150, 165, 175, "
        "200, 250, and 300. The one-standard-error rule retained five PLS-SVD and 50 SIMPLS components, "
        "both interior to the evaluated grids. These are family-specific predictive settings rather "
        "than a matched implementation comparison."
    )

    replace_everywhere(document, "prespecified", "predefined")
    document.save(MAIN_OUTPUT)


def update_supplement():
    document = Document(SUPP_SOURCE)

    cv = paragraph_by_prefix(document, "pls.single.cv() evaluates eligible")
    cv.text = (
        "pls.single.cv() evaluates eligible combinations of component count and prediction settings "
        "within one K-fold layer; pls.double.cv() adds an outer layer for performance estimation after "
        "tuning. For classification, auto selection uses accuracy; balanced accuracy, the unweighted "
        "mean of class-specific recalls, must be requested when class imbalance makes overall accuracy "
        "inappropriate. Dummy-response Q² is also available but is not decoded-label performance. For "
        "regression, auto selection uses RMSD. The 'r2' selection metric denotes an observed-mean "
        "held-out R², whereas 'q2' uses fold-training response means. Full-data training R² is "
        "descriptive and is never optimized. Nested permutation inference uses the same selected endpoint "
        "and its appropriate tail. Grouped permutation and fold rules preserve complete equal-size "
        "constraint blocks; failed null fits are recorded and excluded from the successful-null count."
    )

    grids = paragraph_by_prefix(document, "Two uses of component grids are distinguished")
    grids.text = (
        "Two uses of component grids are distinguished. Benchmark trajectory figures evaluate fixed "
        "component counts on a fixed test set and are descriptive. Training-only model selection reports "
        "the best value within the evaluated grid. Among 46 evaluated family-dataset choices, 20 reached "
        "the upper boundary, four reached the lower boundary, nine were response-rank limited, and 13 "
        "were interior. Thus 33/46 choices were constrained and are treated as fixed computational "
        "workloads rather than unconstrained predictive optima. The complete status is stored in "
        "benchmark_results/manuscript_revision_cycle48_20260726/component_selection_by_family.csv. "
        "NMR was the central sensitivity analysis: its grid was extended to 300 components and the "
        "one-standard-error rule selected interior settings. Table S14 reports supporting repeated-"
        "partition sensitivity, explicitly separated from frozen-release central claims."
    )

    formulas = paragraph_by_prefix(document, "For cross-validated Q²")
    formulas.text = (
        "In the displayed R² equation, full-data training R² uses fitted responses and the mean "
        "of each complete training-response column. The selectable observed-mean cross-validated R² "
        "uses the same numerator-denominator form, but substitutes out-of-fold predictions and the "
        "mean of each complete observed-response column. In the displayed Q² equation, every held-out "
        "fold contributes its prediction sum of squares, while its denominator is centred on the mean "
        "estimated from that fold's training observations; fold numerators and denominators are summed "
        "before forming Q². Independent-test Q² instead uses the complete training-response mean. "
        "RMSD is the square root of the out-of-fold sum of squared errors divided by the number of "
        "predicted response entries. For PLS-DA, R² and Q² use dummy-coded responses and are not "
        "decoded-label accuracy. The public metrics$definitions element records the convention used."
    )

    datasets = paragraph_by_prefix(document, "The benchmark writes a machine-generated dataset manifest")
    insert_after(
        document,
        datasets,
        "Split-unit qualification. The archived cross-dataset benchmark does not contain an auditable "
        "subject or donor grouping variable for every prepared object. Its held-out accuracies and any "
        "row-level binomial intervals therefore quantify conditional computational performance only. "
        "Cell-level rows in Retina, Tabula Muris, and CBMC CITE-seq are not treated as independent "
        "biological replicates. Biomedical generalization claims require donor- or subject-grouped "
        "splits through constrain; where such identifiers are unavailable, the analysis is explicitly "
        "labelled a row-level computational benchmark."
    )

    nmr_grid = paragraph_by_prefix(document, "The family-selected predictive analysis")
    nmr_grid.text = (
        "The family-selected predictive analysis and paired backend analysis answer different questions. "
        "Component selection used five paired training-only splits. The PLS-SVD candidate grid was 1, 2, "
        "3, 5, 7, 10, 25, 50, 75, 100, 125, 150, 165, 175, 200, 250, and 300 components. The SIMPLS "
        "candidate grid was 10, 25, 50, 75, 100, 125, 150, 165, 175, 200, 250, and 300 components. "
        "Both selected values were interior. This extended grid is the principal biomedical sensitivity "
        "analysis; the cross-dataset boundary rows remain fixed computational workloads."
    )

    table_s14 = paragraph_by_prefix(document, "Table S14.")
    table_s14.text = (
        "Table S14. Supporting repeated-partition predictive dispersion and component-selection "
        "sensitivity. These development results show training-partition variability but were not "
        "regenerated from the frozen 0.99.25 archive and are not used as central release claims. "
        "Selection frequencies at a boundary are reported explicitly."
    )

    replace_everywhere(document, "prespecified", "predefined")
    document.save(SUPP_OUTPUT)


if __name__ == "__main__":
    update_main()
    update_supplement()
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)
