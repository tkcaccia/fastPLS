from copy import deepcopy
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.shared import Pt


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle55"
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle56"

MAIN_SOURCE = SOURCE_DIR / "fastPLS_CMPB_main_cycle55_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE_DIR / "fastPLS_CMPB_supplement_cycle55_0.99.6_20260726.docx"
MAIN_OUTPUT = OUTPUT_DIR / "fastPLS_CMPB_main_cycle56_0.99.6_20260726.docx"
SUPP_OUTPUT = OUTPUT_DIR / "fastPLS_CMPB_supplement_cycle56_0.99.6_20260726.docx"


def clear_paragraph(paragraph):
    for child in list(paragraph._p):
        if child.tag.endswith("}pPr"):
            continue
        paragraph._p.remove(child)


def add_text(paragraph, text, superscript=False):
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.font.superscript = superscript
    return run


def set_main_byline(paragraph):
    clear_paragraph(paragraph)
    author_parts = [
        ("Dupe Ojo", "a,\u2020"),
        (", Alessia Vignoli", "b,c,\u2020"),
        (", Stefano Cacciatore", "a,d,*"),
        (", Leonardo Tenori", "b,c,*"),
    ]
    for name, marks in author_parts:
        add_text(paragraph, name)
        add_text(paragraph, marks, superscript=True)


def set_supplement_byline(paragraph):
    clear_paragraph(paragraph)
    author_parts = [
        ("Dupe Ojo", "\u2020"),
        (", Alessia Vignoli", "\u2020"),
        (", Stefano Cacciatore", "*"),
        (", Leonardo Tenori", "*"),
    ]
    for name, marks in author_parts:
        add_text(paragraph, name)
        add_text(paragraph, marks, superscript=True)


def set_note(paragraph, text):
    clear_paragraph(paragraph)
    add_text(paragraph, text)


def add_note_after(paragraph, text):
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        if not child.tag.endswith("}pPr"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph

    inserted = Paragraph(new_p, paragraph._parent)
    set_note(inserted, text)
    return inserted


def update_main():
    document = Document(MAIN_OUTPUT)
    set_main_byline(document.paragraphs[1])
    set_note(
        document.paragraphs[6],
        "\u2020These authors contributed equally and share first authorship.",
    )
    add_note_after(
        document.paragraphs[6],
        "*Co-corresponding authors: Stefano Cacciatore, "
        "stefano.cacciatore@icgeb.org; Leonardo Tenori, tenori@cerm.unifi.it.",
    )
    document.save(MAIN_OUTPUT)


def update_supplement():
    document = Document(SUPP_OUTPUT)
    set_supplement_byline(document.paragraphs[2])
    equal_note = add_note_after(
        document.paragraphs[2],
        "\u2020These authors contributed equally and share first authorship.",
    )
    add_note_after(
        equal_note,
        "*Co-corresponding authors: Stefano Cacciatore, "
        "stefano.cacciatore@icgeb.org; Leonardo Tenori, tenori@cerm.unifi.it.",
    )
    document.save(SUPP_OUTPUT)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    copy2(MAIN_SOURCE, MAIN_OUTPUT)
    copy2(SUPP_SOURCE, SUPP_OUTPUT)
    update_main()
    update_supplement()
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)


if __name__ == "__main__":
    main()
