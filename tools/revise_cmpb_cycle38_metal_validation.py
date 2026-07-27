#!/usr/bin/env python3

import csv
import math
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle36"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle38"
SUMMARY = (
    ROOT
    / "benchmark_results"
    / "metal_validation_20260726"
    / "summary"
    / "metal_validation_summary.csv"
)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle36_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle36_0.99.6_20260726.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle38_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle38_0.99.6_20260726.docx"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def set_paragraph(paragraph, text):
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def insert_after(paragraph, text, style="Body Text"):
    new_xml = paragraph._parent.add_paragraph()._p
    paragraph._p.addnext(new_xml)
    new_paragraph = Paragraph(new_xml, paragraph._parent)
    new_paragraph.style = style
    new_paragraph.add_run(text)
    return new_paragraph


def read_summary():
    with SUMMARY.open(newline="", encoding="utf-8") as stream:
        return list(csv.DictReader(stream))


def number(row, name):
    value = row.get(name, "")
    try:
        return float(value)
    except (TypeError, ValueError):
        return math.nan


def lookup(rows, **conditions):
    matches = [
        row
        for row in rows
        if all(str(row.get(name, "")) == str(value) for name, value in conditions.items())
    ]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one result for {conditions}, found {len(matches)}")
    return matches[0]


def fmt_time(row):
    median = number(row, "median_total_sec")
    low = number(row, "min_total_sec")
    high = number(row, "max_total_sec")
    if math.isfinite(low) and math.isfinite(high) and abs(high - low) > 1e-12:
        return f"{median:.3f} [{low:.3f}-{high:.3f}]"
    return f"{median:.3f}"


def fmt_metric(row):
    median = number(row, "median_metric")
    low = number(row, "min_metric")
    high = number(row, "max_metric")
    if median < 0.01:
        value = f"{median:.3e}"
        bounds = f"{low:.3e}-{high:.3e}"
    else:
        value = f"{median:.4f}"
        bounds = f"{low:.4f}-{high:.4f}"
    if math.isfinite(low) and math.isfinite(high) and abs(high - low) > 1e-12:
        return f"{value} [{bounds}]"
    return value


def fmt_rss_pair(cpu, metal):
    cpu_rss = number(cpu, "median_incremental_peak_rss_mb")
    metal_rss = number(metal, "median_incremental_peak_rss_mb")
    return f"{cpu_rss:.1f}/{metal_rss:.1f}"


def set_cell_margins(cell, top=35, start=45, bottom=35, end=45):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell(cell, text, bold=False, size=6.4, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def set_widths(table, widths):
    table.autofit = False
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(int(width * 1440)))
        grid.append(node)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def prevent_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def build_metal_rows(rows):
    default = [
        row
        for row in rows
        if row["experiment"] == "real_dataset"
        and row["svd_method"] == "rsvd"
        and row["oversample"] == "10"
        and row["power"] == "1"
    ]
    specifications = []
    for dataset, methods, classifiers in (
        ("metref", ("plssvd", "simpls", "opls", "kernelpls"), ("lda",)),
        ("retina", ("plssvd", "simpls", "opls", "kernelpls"), ("lda",)),
        ("tabula", ("plssvd", "simpls"), ("lda",)),
        ("cifar100", ("plssvd", "simpls"), ("argmax", "lda")),
    ):
        for method in methods:
            for classifier in classifiers:
                for precision in ("float64", "float32"):
                    specifications.append((dataset, method, classifier, precision))

    output = []
    labels = {
        "metref": "MetRef",
        "retina": "Retina",
        "tabula": "Tabula Muris",
        "cifar100": "CIFAR-100",
        "plssvd": "PLS-SVD",
        "simpls": "SIMPLS",
        "opls": "OPLS",
        "kernelpls": "kernel PLS",
    }
    for dataset, method, classifier, precision in specifications:
        cpu = lookup(
            default,
            dataset=dataset,
            method=method,
            classifier=classifier,
            precision=precision,
            backend_requested="cpu",
        )
        metal = lookup(
            default,
            dataset=dataset,
            method=method,
            classifier=classifier,
            precision=precision,
            backend_requested="metal",
        )
        speedup = number(cpu, "median_total_sec") / number(metal, "median_total_sec")
        output.append(
            [
                labels[dataset],
                f"{labels[method]} / {'LDA' if classifier == 'lda' else classifier}",
                precision.replace("float", "f"),
                cpu["ncomp"],
                fmt_time(cpu),
                fmt_time(metal),
                f"{speedup:.2f}",
                f"{fmt_metric(cpu)} / {fmt_metric(metal)}",
                fmt_rss_pair(cpu, metal),
            ]
        )

    for dataset, method, label in (
        ("nmr_q28355", "plssvd", "NMR full q=28,355"),
        ("nmr_q5000", "simpls", "NMR guarded q=5,000"),
    ):
        cpu = lookup(
            rows,
            dataset=dataset,
            method=method,
            precision="float32",
            backend_requested="cpu",
        )
        metal = lookup(
            rows,
            dataset=dataset,
            method=method,
            precision="float32",
            backend_requested="metal",
        )
        speedup = number(cpu, "median_total_sec") / number(metal, "median_total_sec")
        output.append(
            [
                label,
                "PLS-SVD / regression" if method == "plssvd" else "SIMPLS / regression",
                "f32",
                cpu["ncomp"],
                fmt_time(cpu),
                fmt_time(metal),
                f"{speedup:.2f}",
                f"{fmt_metric(cpu)} / {fmt_metric(metal)}",
                fmt_rss_pair(cpu, metal),
            ]
        )
    return output


def replace_metal_table(document, rows):
    old_table = document.tables[25]
    expected = ["Dataset", "Method", "Head", "A", "CPU metric", "Metal metric", "Agreement"]
    if [cell.text.strip() for cell in old_table.rows[0].cells] != expected:
        raise RuntimeError("The expected cycle-36 Metal table was not found.")

    headers = [
        "Dataset",
        "Model / head",
        "Prec.",
        "A",
        "CPU time s",
        "Metal time s",
        "CPU/Metal speed ratio",
        "CPU / Metal metric",
        "CPU / Metal incremental RSS MB",
    ]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = old_table.style
    old_table._tbl.addnext(table._tbl)
    old_table._element.getparent().remove(old_table._element)

    for cell, header in zip(table.rows[0].cells, headers):
        set_cell(cell, header, bold=True, size=6.1)
        shade(cell, "D9E7F5")
    repeat_header(table.rows[0])
    prevent_split(table.rows[0])

    for values in rows:
        row = table.add_row()
        for index, (cell, value) in enumerate(zip(row.cells, values)):
            align = WD_ALIGN_PARAGRAPH.LEFT if index in (0, 1) else WD_ALIGN_PARAGRAPH.CENTER
            set_cell(cell, value, size=6.0, align=align)
        prevent_split(row)

    set_widths(table, [0.78, 1.03, 0.32, 0.28, 0.72, 0.72, 0.57, 1.00, 1.03])


def revise_main(rows):
    document = Document(MAIN_SOURCE)

    methods = find_paragraph(document, "Methods: fastPLS provides")
    set_paragraph(
        methods,
        methods.text
        + " A separate Apple M3 campaign evaluated Metal portability, precision, "
        "shape-dependent performance, unified-memory use, and effective stage "
        "residency in isolated CPU-versus-Metal runs.",
    )

    results = find_paragraph(document, "Results: The deterministic IRLBA path")
    set_paragraph(
        results,
        results.text
        + " In the separate Apple M3 study, Metal accelerated sufficiently large "
        "dense float64 workloads by up to 4.35-fold on CIFAR-100, but CPU float32 "
        "was faster on smaller omics tasks and NMR. A float64 Metal PLS-SVD "
        "accuracy discrepancy prevented a general backend-equivalence claim.",
    )

    precision_methods = find_paragraph(
        document, "Double precision is the numerical reference."
    )
    set_paragraph(
        precision_methods,
        (
            "Double precision is the numerical reference. Inputs from the float "
            "package select float32 arithmetic on supported CPU, CUDA, and Metal "
            "routes, but float32 is treated as a conditional reduced-storage mode, "
            "not a general acceleration or equivalence claim. The Apple Metal "
            "campaign used a MacBook Pro Mac15,3 with an Apple M3, eight CPU cores, "
            "ten GPU cores, and 8 GB unified memory. It included isolated CPU and "
            "Metal fits for all four PLS families on MetRef and Retina, PLS-SVD and "
            "SIMPLS on Tabula Muris and CIFAR-100, synthetic matrix-shape regimes, "
            "guarded and full-response NMR, OPLS orthogonal-component and kernel "
            "settings, standalone rSVD reliability, and 10-fold cross-validation. "
            "Real-data results used three independently seeded rSVD fits; timings "
            "covered fitting and prediction. Because Apple silicon shares physical "
            "memory between CPU and GPU, memory was measured as pre-fit process RSS "
            "and baseline-corrected peak process RSS rather than as separate host "
            "RAM and GPU VRAM. Requested and reported stage backends were retained "
            "to distinguish native Metal work from host-assisted execution. "
            "Method- and shape-based warnings remained active for precision-sensitive "
            "classification, extreme multivariate responses, and nonlinear kernels "
            "(Supplementary Sections S19-S20)."
        ),
    )

    metal_results = find_paragraph(
        document, "Metal established numerical portability on two classification shapes."
    )
    set_paragraph(
        metal_results,
        (
            "The expanded Apple M3 campaign completed 1,064 isolated fits and "
            "decompositions without process failure, but did not support a universal "
            "Metal advantage. CPU was faster for MetRef, Retina, Tabula Muris, "
            "tall-thin synthetic matrices, and the extreme-response NMR task. On "
            "CIFAR-100 at 50 components, float64 Metal accelerated PLS-SVD by "
            "4.35-fold with argmax and 4.03-fold with LDA, and accelerated SIMPLS "
            "by 1.94- and 2.20-fold, respectively. These speedups were not uniformly "
            "estimator matched: median PLS-SVD/argmax accuracy was 0.7418 on CPU "
            "and 0.7082 on Metal, whereas the LDA difference was 0.8323 versus "
            "0.8274. For float32 CIFAR-100, CPU was 1.25-1.62 times faster than "
            "Metal, while median CPU/Metal accuracy was identical for both PLS-SVD "
            "heads and SIMPLS-LDA and differed by 0.17 percentage points for "
            "SIMPLS/argmax."
        ),
    )
    inserted = insert_after(
        metal_results,
        (
            "The NMR result reinforced the shape dependence. Full-response float32 "
            "PLS-SVD (13,000 predictors and 28,355 responses) required 2.428 s and "
            "912 MB incremental RSS on CPU versus 10.664 s and 2,210 MB on Metal; "
            "RMSD was 7.14e-4 and 7.27e-4, respectively. Metal SIMPLS "
            "was not extended beyond q=5,000 because that guarded run was 7.29-fold "
            "slower than CPU without predictive benefit. Increasing rSVD power "
            "iterations reduced the CIFAR-100 CPU/Metal SIMPLS accuracy difference "
            "from 2.71 to 0.30 percentage points, whereas the float64 Metal PLS-SVD "
            "result remained insensitive to the exposed power and oversampling "
            "controls. Finally, public Metal cross-validation completed all 192 "
            "runs, but returned metadata identified CPU rSVD or IRLBA decomposition "
            "and compiled CPU LDA scoring for important stages; these runs establish "
            "cross-platform usability, not fully Metal-resident cross-validation "
            "(Supplementary Table S22)."
        ),
    )
    inserted.paragraph_format.keep_together = True

    backend_discussion = find_paragraph(
        document, "The backend design is a second contribution."
    )
    set_paragraph(
        backend_discussion,
        (
            "The backend design is a second contribution, but the Apple results "
            "show why hardware selection must be shape and precision aware. Metal "
            "Performance Shaders accelerated the large dense float64 CIFAR-100 and "
            "balanced synthetic products once useful work amortized dispatch and "
            "synchronization. The same backend was slower for small omics, tall-thin, "
            "and extreme-response NMR matrices on an 8 GB M3, and usually consumed "
            "more incremental unified-process RSS. CUDA remains the more complete "
            "accelerator route for the primary benchmarks. OPLS filtering, nonlinear "
            "kernel construction, reduced decompositions, and public Metal "
            "cross-validation retain host stages; requested and effective stage "
            "backends must therefore be reported rather than inferred from the "
            "top-level backend argument."
        ),
    )

    float_discussion = find_paragraph(
        document, "Float32 approximately halves raw input storage"
    )
    set_paragraph(
        float_discussion,
        (
            "Float32 approximately halves raw input storage and can reduce selected "
            "workspaces, but the combined CUDA and Metal evidence does not establish "
            "a general speed or memory advantage. On the M3, float32 CPU execution "
            "was the fastest evaluated route for CIFAR-100 and NMR and closely "
            "matched float32 Metal predictions, whereas Metal added dispatch and "
            "unified-memory overhead. Float64 Metal randomized paths were more "
            "sensitive: SIMPLS agreement improved with an additional power "
            "iteration, but PLS-SVD retained an unresolved CIFAR-100 discrepancy. "
            "The package therefore retains float64 CPU IRLBA as the confirmatory "
            "reference, treats rSVD and float32 as validated conditionally by route "
            "and matrix shape, and requires held-out agreement checks before reduced "
            "precision or accelerator results are used for scientific inference."
        ),
    )

    conclusion = find_paragraph(
        document, "fastPLS combines an accelerated sequential SIMPLS"
    )
    set_paragraph(
        conclusion,
        (
            "fastPLS combines an accelerated sequential SIMPLS implementation with "
            "memory-aware PLS-SVD, compiled validation, compact prediction, "
            "conditional float32 execution, and CPU, CUDA, and Metal backends. "
            "The expanded Apple validation demonstrates portability but also defines "
            "its present boundary: Metal is advantageous for selected large dense "
            "float64 workloads, whereas CPU float32 remains preferable on the tested "
            "M3 for smaller omics and extreme-response NMR matrices. These "
            "capabilities make established PLS workflows accessible at scales that "
            "were previously limited by runtime or memory while preserving explicit "
            "numerical and residency diagnostics. The GPL-3 R package calls reusable "
            "C++ components maintained with the MIT-licensed kodama-cpp codebase; "
            "future work will increase native accelerator residency and validate "
            "agreement across larger Apple GPU configurations."
        ),
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - Apple Metal validation revision"
    )
    document.save(MAIN_OUT)


def revise_supplement(rows):
    document = Document(SUPP_SOURCE)
    heading = find_paragraph(document, "S20. CPU and Metal reproducibility")
    set_paragraph(
        heading,
        "S20. Apple Metal performance, precision, and residency validation",
    )

    methods = find_paragraph(document, "MetRef used 773/100 train/test observations")
    set_paragraph(
        methods,
        (
            "The local Apple campaign used a MacBook Pro Mac15,3 with an Apple M3 "
            "(eight CPU cores, ten GPU cores) and 8 GB unified memory. Each model "
            "ran in an isolated R process. The real-data suite used the stored "
            "MetRef, Retina, Tabula Muris, and CIFAR-100 splits, rSVD oversampling "
            "10, one power iteration, and three independent seeds. It evaluated "
            "float64 and float32 inputs, argmax and LDA where applicable, and all "
            "four PLS families where matrix shape permitted. Separate suites covered "
            "five synthetic shape regimes, guarded/full NMR, OPLS orthogonal "
            "components, linear/RBF/polynomial kernels, standalone rSVD reliability, "
            "and 10-fold cross-validation. Time includes fitting plus prediction. "
            "Memory is the isolated-process peak minus the RSS recorded immediately "
            "before fitting. It represents incremental unified-process memory, not "
            "dedicated GPU VRAM."
        ),
    )

    caption = find_paragraph(document, "Table S22. Matched CPU/Metal predictive agreement.")
    set_paragraph(
        caption,
        (
            "Table S22. Repeated CPU-versus-Metal validation on Apple M3. Time and "
            "predictive metric are medians [minimum-maximum] across three independently "
            "seeded rSVD fits for real classification tasks; NMR entries are single "
            "guarded feasibility runs. CPU/Metal speed ratio above one favors Metal. "
            "Metrics are accuracy for classification and RMSD for regression. "
            "Incremental RSS is baseline-corrected unified-process memory. A is the "
            "retained component count."
        ),
    )
    replace_metal_table(document, build_metal_rows(rows))

    findings = find_paragraph(document, "On CIFAR-100, Metal PLS-SVD/argmax required")
    set_paragraph(
        findings,
        (
            "All 1,064 isolated benchmark fits and decompositions completed. Metal "
            "was consistently slower for MetRef, Retina, Tabula Muris, tall-thin "
            "synthetic matrices, and NMR. In contrast, float64 Metal accelerated "
            "CIFAR-100 PLS-SVD by 4.03-4.35-fold and SIMPLS by 1.94-2.20-fold. "
            "Metal also accelerated float64 PLS-SVD 2.10-fold in the balanced "
            "synthetic regime and 4.08-fold in the high-component regime. These "
            "runtime gains cannot be interpreted alone: float64 CIFAR-100 "
            "PLS-SVD/argmax differed by 3.36 percentage points between CPU and "
            "Metal, and SIMPLS/argmax differed by 2.71 percentage points under the "
            "one-power setting. Float32 CPU/Metal median accuracy was identical for "
            "PLS-SVD argmax and LDA and SIMPLS-LDA, and differed by only 0.17 "
            "percentage points for SIMPLS/argmax, but CPU was faster for all four "
            "float32 CIFAR routes."
        ),
    )
    sensitivity = insert_after(
        findings,
        (
            "A sensitivity analysis used oversampling 10 or 20 and one or two power "
            "iterations on CIFAR-100. Two power iterations reduced the median "
            "CPU/Metal SIMPLS argmax difference from 2.71 to 0.30 percentage points. "
            "The float64 Metal PLS-SVD prediction remained unchanged across the "
            "tested settings while CPU predictions changed with the randomized "
            "sketch, indicating that the exposed settings do not currently govern "
            "the effective Metal PLS-SVD route. This route is therefore excluded "
            "from numerical-equivalence claims pending correction. In the standalone "
            "rSVD suite, larger oversampling and two power iterations reduced "
            "singular-value and subspace errors. Large principal angles on Gaussian "
            "matrices with clustered singular values were interpreted together with "
            "singular-value and prediction errors because the individual truncated "
            "subspace is weakly identifiable in that setting."
        ),
    )
    residency = insert_after(
        sensitivity,
        (
            "Full-response float32 NMR PLS-SVD completed in 2.428 s on CPU and "
            "10.664 s on Metal, with RMSD 7.1419e-4 and 7.2741e-4 and "
            "incremental RSS 912 and 2,210 MB, respectively. At q=5,000, float32 "
            "SIMPLS produced the same RMSD on CPU and Metal but Metal required "
            "21.582 s versus 2.959 s; full q=28,355 Metal SIMPLS was not attempted "
            "on the 8 GB machine. All 192 public cross-validation requests completed, "
            "but effective metadata reported CPU rSVD or IRLBA decomposition and "
            "compiled CPU LDA scoring for important stages. Consequently, the "
            "current public Metal cross-validation route is hybrid and is not "
            "described as fully device resident."
        ),
    )
    residency.paragraph_format.keep_together = True

    document.core_properties.title = (
        "fastPLS CMPB supplementary material - Apple Metal validation revision"
    )
    document.save(SUPP_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    rows = read_summary()
    revise_main(rows)
    revise_supplement(rows)
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
