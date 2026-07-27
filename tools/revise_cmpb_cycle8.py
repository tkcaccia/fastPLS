"""Create cKNN-free Cycle 8 CMPB drafts from the verified Cycle 7 documents."""

import csv
from pathlib import Path
from shutil import copy2

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


SRC = Path("/Users/stefano/Desktop/fastPLS/CMPB_rewrite_20260724_cycle7")
OUT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github/artifacts/CMPB_rewrite_20260725_cycle8")
OUT.mkdir(parents=True, exist_ok=True)

FILES = {
    "main": (
        SRC / "fastPLS_CMPB_main_cycle7_0.99.6_20260724.docx",
        OUT / "fastPLS_CMPB_main_cycle8_0.99.6_20260725.docx",
    ),
    "supplement": (
        SRC / "fastPLS_CMPB_supplement_cycle7_0.99.6_20260724.docx",
        OUT / "fastPLS_CMPB_supplement_cycle8_0.99.6_20260725.docx",
    ),
    "review": (
        SRC / "fastPLS_CMPB_independent_reviewer_report_cycle7_20260724.docx",
        OUT / "fastPLS_CMPB_independent_reviewer_report_cycle8_20260725.docx",
    ),
}

NMR_PLOTS = Path(
    "/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github/"
    "benchmark_results/review_nmr_20260724/plots"
)
NMR_METHOD_SUMMARY = NMR_PLOTS / "nmr_final_method_summary.csv"


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def delete_cknn_content(doc):
    terms = ("cknn", "candidate-knn", "candidate knn")
    for paragraph in list(doc.paragraphs):
        if any(term in paragraph.text.lower() for term in terms):
            delete_paragraph(paragraph)
    for table in doc.tables:
        for row in table.rows:
            if any(any(term in cell.text.lower() for term in terms) for cell in row.cells):
                for cell in row.cells:
                    cell.text = "Removed: obsolete package-native nearest-neighbour analysis."


def replace_startswith(doc, prefix, replacement):
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(prefix):
            paragraph.text = replacement
            return
    raise RuntimeError(f"Paragraph beginning '{prefix}' was not found")


def remove_table_borders(table):
    table_properties = table._tbl.tblPr
    borders = table_properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def add_nmr_main_figure(doc):
    figure_paths = [
        NMR_PLOTS / "nmr_spectrum_full.png",
        NMR_PLOTS / "nmr_spectrum_zoom.png",
        NMR_PLOTS / "nmr_per_spectrum_rmsd.png",
        NMR_PLOTS / "nmr_speed_memory.png",
    ]
    missing = [str(path) for path in figure_paths if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing NMR figure panels: " + ", ".join(missing))

    anchor = next(
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.startswith("The matched CPU/CUDA analysis is a three-run")
    )
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    remove_table_borders(table)

    for label, path, cell in zip("ABCD", figure_paths, [c for row in table.rows for c in row.cells]):
        cell.width = Inches(3.18)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        label_run = paragraph.add_run(label + "\n")
        label_run.bold = True
        label_run.font.size = Pt(9)
        paragraph.add_run().add_picture(str(path), width=Inches(3.05))

    caption = doc.add_paragraph(style="Image Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    caption.add_run(
        "Figure 2. NMR prediction and computational performance at the selected "
        "100-component model. (A) Observed and SIMPLS-rSVD-predicted held-out "
        "spectrum across the full chemical-shift range. (B) Enlarged 0.5-1.7 ppm "
        "region. (C) Distribution of per-spectrum RMSD across the 321 held-out "
        "spectra for the deposited PLS-SVD/IRLBA reference and fastPLS PLS-SVD "
        "and SIMPLS CPU/CUDA rSVD implementations. (D) Median total "
        "fitting-plus-prediction time, peak host RSS, and sampled peak GPU memory "
        "across three isolated runs. The representative spectrum was selected "
        "mechanically as the test spectrum whose RMSD was closest to the held-out "
        "median. All comparisons used float64 data and the same routinely "
        "preprocessed training and held-out spectra."
    )

    anchor._p.addnext(caption._p)
    anchor._p.addnext(table._tbl)


def read_nmr_method_summary():
    if not NMR_METHOD_SUMMARY.exists():
        raise FileNotFoundError(f"Missing NMR method summary: {NMR_METHOD_SUMMARY}")
    with NMR_METHOD_SUMMARY.open(newline="") as handle:
        return list(csv.DictReader(handle))


def add_nmr_reference_table(doc):
    rows = read_nmr_method_summary()
    anchor = next(
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.startswith("Held-out NMR results.")
    )
    caption = doc.add_paragraph()
    caption.add_run(
        "Table S6. Precision-matched float64 comparison with the deposited "
        "fastsimpls PLS-SVD/IRLBA reference at 100 components. Time and memory "
        "are medians from three isolated runs; prediction agreement is the Pearson "
        "correlation between vectorized held-out predictions and the deposited "
        "reference prediction."
    ).bold = True
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = (
        "Implementation",
        "Total time (s)",
        "Host RSS (MB)",
        "GPU memory (MB)",
        "RMSD",
        "Q2",
        "Agreement",
    )
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(7.5)
    for result in rows:
        cells = table.add_row().cells
        gpu = result["gpu_peak_mb"]
        values = (
            result["method"],
            f"{float(result['total_time_sec']):.3f}",
            f"{float(result['host_rss_mb']):.0f}",
            "NA" if not gpu or gpu == "NA" else f"{float(gpu):.0f}",
            f"{float(result['RMSD']):.7f}",
            f"{float(result['Q2']):.5f}",
            f"{float(result['prediction_agreement']):.6f}",
        )
        for cell, value in zip(cells, values):
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)
    anchor._p.addnext(table._tbl)
    anchor._p.addnext(caption._p)


def revise_main():
    src, dest = FILES["main"]
    copy2(src, dest)
    doc = Document(dest)
    replace_startswith(
        doc,
        "Regression returns continuous predictions",
        "Regression returns continuous predictions along the requested component path. "
        "For factor responses, fastPLS provides standard argmax PLS-DA decoding "
        "or latent-space linear discriminant analysis (LDA).",
    )
    replace_startswith(
        doc,
        "NMR represented the extreme multivariate-response setting",
        "NMR represented the extreme multivariate-response setting (1,200 training "
        "and 321 held-out spectra; p=13,000; q=28,355). Routine NMR spectral "
        "preprocessing removed the residual water-resonance interval between 4.6 and "
        "4.8 ppm consistently from the training and held-out data before modelling "
        "and evaluation. A fixed 20% inner split of the training spectra "
        "selected components from 10, 25, 50, 75, and 100 using validation RMSD, "
        "leaving the held-out test spectra untouched. "
        "Validation RMSD decreased from 0.001137 at 10 to 0.000894 at 100 components, "
        "which was selected. The primary comparison used float64 data and the "
        "deposited fastsimpls implementation from the earlier NMR study, run as "
        "PLS-SVD with IRLBA (fast=TRUE, iter=FALSE). Across three isolated runs, "
        "the reference required a median 431.23 s and 6,101 MB peak host RSS and "
        "achieved RMSD 0.0007194. fastPLS PLS-SVD/rSVD required 16.32 s and "
        "2,964 MB on CPU, a 26.4-fold speedup and 51% lower host RSS, with RMSD "
        "0.0007292 and prediction correlation 0.999942 against the reference. "
        "The CUDA implementation required 1.115 s, 3,338 MB host RSS, and 664 MB "
        "GPU memory, a 386.8-fold speedup, with RMSD 0.0007183 and prediction "
        "correlation 0.999974. fastPLS SIMPLS/rSVD required 20.14 s on CPU and "
        "3.06 s on CUDA, with RMSD 0.000861 and 0.000805, respectively. Figure 2 "
        "shows the representative observed-versus-predicted spectrum, held-out "
        "error distributions, and matched computational resources; detailed "
        "response-wise errors and the full comparison table are provided in the "
        "Supplementary Material.",
    )
    add_nmr_main_figure(doc)
    delete_cknn_content(doc)
    doc.add_paragraph(
        "Package scope update. Candidate-nearest-neighbour classification was removed "
        "from the public fastPLS interface and is not evaluated as a package method. "
        "The planned ImageNet representation analysis instead uses an independent FAISS "
        "retrieval workflow to compare raw DINOv2, PCA, and PLS scores under matched "
        "train/test splits and query accounting."
    )
    doc.save(dest)


def revise_supplement():
    src, dest = FILES["supplement"]
    copy2(src, dest)
    doc = Document(dest)
    replace_startswith(
        doc,
        "NMR predictor preprocessing and selection.",
        "Routine NMR spectral preprocessing and component selection. The NMR task "
        "comprised 1,200 training and 321 held-out spectra, 13,000 predictors, and "
        "28,355 numeric responses. The residual water-resonance interval from 4.6 to "
        "4.8 ppm was removed consistently from the training and held-out spectra as "
        "part of routine preprocessing before modelling and evaluation. A fixed seed "
        "(123) selected 20% of training spectra as an inner validation set. "
        "SIMPLS-rSVD models with 10, 25, 50, 75, and 100 components were fitted "
        "independently. Validation RMSD was 0.001137, 0.000998, 0.000929, 0.000913, "
        "and 0.000894, respectively; 100 components were selected.",
    )
    replace_startswith(
        doc,
        "Held-out NMR results.",
        "Held-out NMR results. The precision-matched float64 comparison used the "
        "deposited fastsimpls PLS-SVD/IRLBA function and fastPLS PLS-SVD and "
        "SIMPLS CPU/CUDA rSVD implementations on the same routinely preprocessed "
        "training and held-out spectra at 100 components. Three isolated runs were "
        "performed for every implementation. The deposited reference required "
        "median total time 431.23 s (IQR 2.71 s) and peak host RSS 6,101 MB "
        "(IQR 826 MB), with RMSD 0.0007194 and Q2 0.99484. CPU and CUDA fastPLS "
        "PLS-SVD reduced total time to 16.32 and 1.115 s and host RSS to 2,964 "
        "and 3,338 MB, respectively. Their prediction correlations with the "
        "reference were 0.999942 and 0.999974. CPU and CUDA SIMPLS required "
        "20.14 and 3.06 s, with RMSD 0.000861 and 0.000805. CUDA PLS-SVD and "
        "SIMPLS sampled 664 and 3,432 MB peak GPU memory, respectively.",
    )
    replace_startswith(
        doc,
        "Table S3 reports medians",
        "Table S7 reports medians and interquartile ranges (IQR) from three isolated "
        "repetitions for the completed NMR precision measurements. A timeout means "
        "that the first attempted repetition exceeded 1,200 s; subsequent repetitions "
        "were intentionally not started. RSS denotes peak host resident set size and "
        "GPU memory is the sampled process peak.",
    )
    replace_startswith(
        doc,
        "Table S3. Full NMR matched float64/float32 comparison",
        "Table S7. Full NMR matched float64/float32 comparison at 100 components. "
        "Values are median (IQR); the two entries in the final column give float64 "
        "then float32 completed/attempted repetitions.",
    )
    add_nmr_reference_table(doc)
    delete_cknn_content(doc)
    doc.add_paragraph(
        "Supplementary scope update. Candidate-nearest-neighbour classification has been "
        "removed from fastPLS and from this supplement. ImageNet retrieval is reserved for "
        "an external FAISS analysis with matched raw-feature, PCA-score, and PLS-score "
        "query experiments; it is not evidence for a package-native nearest-neighbour head."
    )
    doc.save(dest)


def revise_review():
    src, dest = FILES["review"]
    copy2(src, dest)
    doc = Document(dest)
    doc.add_paragraph(
        "Cycle 8 scope change: the package-native cKNN head and all related manuscript "
        "and supplement claims were removed. Any future ImageNet retrieval result must use "
        "the independent FAISS workflow and matched PCA/raw-feature controls."
    )
    doc.save(dest)


revise_main()
revise_supplement()
revise_review()
for _, (_, dest) in FILES.items():
    print(dest)
