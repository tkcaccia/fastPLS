from pathlib import Path
from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
MAIN_IN = ROOT / "artifacts/CMPB_rewrite_20260727_cycle86/fastPLS_CMPB_main_cycle86_0.99.7_20260801.docx"
SUPP_IN = ROOT / "artifacts/CMPB_rewrite_20260727_cycle85/fastPLS_CMPB_supplement_cycle85_0.99.6_20260727.docx"
OUT = ROOT / "artifacts/CMPB_rewrite_20260808_cycle87"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle87_0.99.10_20260808.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle87_0.99.10_20260808.docx"


def replace_paragraph(document, old, new):
    matches = [p for p in document.paragraphs if p.text == old]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph match, found {len(matches)}: {old[:100]}")
    matches[0].text = new


def replace_cell_text(document, row_key, column_name, new_text):
    for table in document.tables:
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        if column_name not in headers:
            continue
        key_col = 0
        value_col = headers.index(column_name)
        for row in table.rows[1:]:
            if row.cells[key_col].text.strip() == row_key:
                set_compact_table_cell(row.cells[value_col], new_text)
                return
    raise RuntimeError(f"Could not find table row {row_key!r} and column {column_name!r}")


def set_compact_table_cell(cell, text):
    """Replace a dense supplementary-table cell without losing its 5 pt type."""
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(5)


OUT.mkdir(parents=True, exist_ok=True)

main = Document(MAIN_IN)
replace_paragraph(
    main,
    "The public pls() interface selects PLS family, component count, solver, backend, and, for classification, argmax or latent-space linear discriminant analysis (LDA). pls.single.cv() selects settings within one cross-validation layer; pls.double.cv() uses nested cross-validation. Requested estimators are never silently substituted.",
    "The public pls() interface selects PLS family, component count, solver, backend, and, for classification, argmax or latent-space linear discriminant analysis (LDA). pls.single.cv() selects settings within one cross-validation layer; pls.double.cv() uses nested cross-validation. The audited 0.99.10 interface exposes PLS modelling, prediction, evaluation, cross-validation, score plotting, and standalone truncated SVD; PCA and nearest-neighbour classifiers are not package APIs. Requested estimators are never silently substituted.",
)
replace_paragraph(
    main,
    "Standard R matrices use eight-byte float64 values; float-package inputs use four-byte float32 values. float64 is the reference. Precision support and backend residency are route specific; the authoritative capability classifications are in Supplementary Tables S1 and S9.",
    "Standard R matrices use eight-byte float64 values; float-package inputs use four-byte float32 values. float64 is the reference. Precision support and backend residency are route specific; the authoritative capability classifications are in Supplementary Tables S1 and S9. On Windows, a portable float-package CPU fallback supports rSVD PLS-SVD, SIMPLS, and linear-kernel SIMPLS with argmax, whereas native compiled float32 OPLS, nonlinear kernel PLS, and LDA remain unavailable.",
)
replace_paragraph(
    main,
    "Regression returns continuous predictions. Classification uses argmax PLS-DA or LDA fitted to PLS scores. LDA uses pooled within-class covariance, Cholesky solves, class priors, and deterministic trace-scaled regularization only when factorization fails [17,18].",
    "Regression returns continuous predictions. Classification uses argmax PLS-DA or LDA fitted to PLS scores. LDA uses pooled within-class covariance, Cholesky solves, class priors, and deterministic trace-scaled regularization only when factorization fails [17,18]. For imbalanced classification, cross-validation can select by balanced accuracy, defined as the unweighted mean of class-specific recalls; nested permutation testing uses the same selected endpoint rather than substituting dummy-response Q².",
)
replace_paragraph(
    main,
    "Fold construction, fitting, prediction, and metric accumulation remain compiled where supported; grouped observations can be constrained to one fold. Hybrid OPLS, nonlinear-kernel, and Metal paths are identified explicitly.",
    "Fold construction, fitting, prediction, and metric accumulation remain compiled where supported; grouped observations can be constrained to one fold. Model-selection endpoints include accuracy, balanced accuracy, R², Q², and RMSD. Hybrid OPLS, nonlinear-kernel, and Metal paths are identified explicitly.",
)
replace_paragraph(
    main,
    "Code and benchmark outputs are available at https://github.com/tkcaccia/fastPLS; reusable components are at https://github.com/tkcaccia/kodama-cpp. The reviewed software snapshot is fastPLS 0.99.6. Commit 6e50bd318f20289101f6b723953830aefa8b95d6 identifies the base source, and the exact experimental source archive fastPLS_0.99.6.tar.gz has SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85. Analysis-specific scripts and archive digests are reported in Supplementary Table S15.",
    "Code and benchmark outputs are available at https://github.com/tkcaccia/fastPLS; reusable components are at https://github.com/tkcaccia/kodama-cpp. Quantitative results remain tied to fastPLS 0.99.6, base commit 6e50bd318f20289101f6b723953830aefa8b95d6, and source-archive SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85. The current audited interface is fastPLS 0.99.10; its source archive has SHA-256 163ac7bd5c0c241f3817fac989e219f71b3956b388f6fcefa2f3420c45051b25. Analysis-specific scripts and archive digests are reported in Supplementary Table S15.",
)
main.save(MAIN_OUT)

supp = Document(SUPP_IN)
replace_paragraph(
    supp,
    "fastPLS: scalable partial least squares with compiled CPU and accelerator backends for high-dimensional biomedical data",
    "fastPLS: accelerated SIMPLS for high-dimensional biomedical data with compiled CPU and accelerator backends",
)
replace_paragraph(
    supp,
    "This supplement describes fastPLS version 0.99.6. Source provenance is analysis specific: Table S15 maps each quantitative analysis to its result archive and records an exact package commit only when the run metadata captured it. Package version, result date, or a later manuscript commit is not treated as evidence of the historical computational SHA.",
    "This supplement distinguishes the current audited fastPLS 0.99.10 interface from the 0.99.6 source archive used for the quantitative benchmarks. Table S15 maps each analysis to its result archive and records an exact package commit only when run metadata captured it. A later package version, result date, or manuscript commit is never treated as evidence of a historical computational SHA.",
)
replace_paragraph(
    supp,
    "The double-precision CPU backend is the broadest reference implementation. Native float32 is selected automatically when input matrices are float::float32 objects and is available for PLS-SVD, SIMPLS, OPLS, and kernel PLS on supported CPU, CUDA, and Metal builds. Native Windows float32 remains unavailable because the Windows R BLAS/LAPACK toolchain used by the package does not expose the required single-precision symbols. Unsupported combinations stop instead of silently converting to double.",
    "The double-precision CPU backend is the broadest reference implementation. Native float32 is selected automatically when input matrices are float::float32 objects and is available for route-specific PLS-SVD, SIMPLS, OPLS, and kernel PLS combinations on supported CPU, CUDA, and Metal builds. Windows now has a portable float-package CPU fallback for rSVD PLS-SVD, SIMPLS, and linear-kernel SIMPLS with argmax; it is not the native compiled single-precision path. Windows float32 OPLS, nonlinear kernel PLS, LDA, CUDA, and Metal combinations stop before fitting instead of silently converting to double.",
)
replace_paragraph(
    supp,
    "pls.single.cv() evaluates eligible combinations of component count and prediction settings within one K-fold cross-validation layer. pls.double.cv() adds an outer layer for unbiased performance estimation. The constrain input assigns related samples, such as repeated measurements from one patient, to the same fold. Leave-one-group-out behaviour therefore respects the constraint rather than separating related observations.",
    "pls.single.cv() evaluates eligible combinations of component count and prediction settings within one K-fold cross-validation layer. pls.double.cv() adds an outer layer for unbiased performance estimation. Selection can use accuracy, balanced accuracy, R², Q², or RMSD. Balanced accuracy is the unweighted mean of class-specific recalls and is intended for unequal class frequencies. When a nested permutation test is requested, the sampled null distribution uses the same selection endpoint and the appropriate upper or lower tail. The constrain input assigns related samples, such as repeated measurements from one patient, to the same fold; leave-one-group-out behaviour therefore never separates related observations.",
)
replace_paragraph(
    supp,
    "Current benchmark workflows record repository state, benchmark-script checksum, package version, source-archive SHA-256, compiler, BLAS/LAPACK, thread settings, accelerator libraries, seeds, rSVD controls, and data/split identifiers. Table S15 maps each analysis to its exact evidence archive. The current NMR and ImageNet runs used fastPLS_0.99.6.tar.gz (SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85) from base commit 6e50bd318f20289101f6b723953830aefa8b95d6.",
    "Current benchmark workflows record repository state, benchmark-script checksum, package version, source-archive SHA-256, compiler, BLAS/LAPACK, thread settings, accelerator libraries, seeds, rSVD controls, and data/split identifiers. Table S15 maps each analysis to its exact evidence archive. NMR and ImageNet used fastPLS_0.99.6.tar.gz (SHA-256 c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85) from base commit 6e50bd318f20289101f6b723953830aefa8b95d6. The audited 0.99.10 archive (SHA-256 163ac7bd5c0c241f3817fac989e219f71b3956b388f6fcefa2f3420c45051b25) changes installation, API cleanup, documentation, and metric-selection behavior; benchmark values are not relabelled as 0.99.10 reruns.",
)
replace_paragraph(
    supp,
    "The ledger never infers a historical commit from a package version or result date. Where a run did not record its Git SHA, the source status is explicitly not recoverable. SHA-256 prefixes identify immutable result archives; full digests are retained in the machine-readable ledger.",
    "The ledger never infers a historical commit from a package version or result date. Where a run did not record its Git SHA, the source status is explicitly not recoverable. SHA-256 prefixes identify immutable result archives; full digests are retained in the machine-readable ledger. The 0.99.10 code audit is documented separately in benchmark/CODE_AUDIT_0.99.10.md and does not alter historical result provenance.",
)

replace_cell_text(
    supp,
    "PLS-SVD",
    "Windows",
    "portable CPU rSVD regression/argmax; native LDA unavailable",
)

# The capability table contains several repeated family labels. Update the
# Windows cells by family, kernel, and backend rather than relying on row keys.
for table in supp.tables:
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    if headers[:4] != ["Family", "Kernel", "Backend", "Endpoint status"] or "Windows" not in headers:
        continue
    win_col = headers.index("Windows")
    for row in table.rows[1:]:
        family = row.cells[0].text.strip()
        kernel = row.cells[1].text.strip()
        backend = row.cells[2].text.strip()
        if backend != "CPU":
            continue
        if family in {"PLS-SVD", "SIMPLS"}:
            set_compact_table_cell(
                row.cells[win_col],
                "portable CPU rSVD regression/argmax; native LDA unavailable",
            )
        elif family == "kernel PLS" and kernel == "linear":
            set_compact_table_cell(
                row.cells[win_col],
                "portable CPU rSVD regression/argmax; native LDA unavailable",
            )
        else:
            set_compact_table_cell(row.cells[win_col], "unavailable")
    break

supp.save(SUPP_OUT)
print(MAIN_OUT)
print(SUPP_OUT)
