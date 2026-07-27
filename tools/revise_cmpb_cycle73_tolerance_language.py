#!/usr/bin/env python3
"""Replace empirical equivalence wording with explicit tolerance language."""

from pathlib import Path

from docx import Document

from revise_cmpb_cycle67_consolidate_evidence import (
    normalize_submission_terminology,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle72"
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle73"
MAIN_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_main_cycle72_0.99.6_20260726.docx"
)
SUPP_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_supplement_cycle72_0.99.6_20260726.docx"
)
MAIN_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_main_cycle73_0.99.6_20260726.docx"
)
SUPP_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_supplement_cycle73_0.99.6_20260726.docx"
)


REPLACEMENTS = (
    (
        "met all prespecified tolerances",
        "met the prespecified numerical tolerances",
    ),
    (
        "one power iteration passed 101/117 checks, whereas two power "
        "iterations passed 117/117",
        "one power iteration met the prespecified numerical tolerances in "
        "101/117 checks, whereas two power iterations met them in 117/117",
    ),
    (
        "one power iteration passed 101/117 checks; oversampling 10 with two "
        "power iterations passed 117/117",
        "one power iteration met the prespecified numerical tolerances in "
        "101/117 checks; oversampling 10 with two power iterations met them "
        "in 117/117",
    ),
    (
        "The two-power rSVD configuration passed all numerical checks",
        "The two-power rSVD configuration met the prespecified numerical "
        "tolerances in all checks",
    ),
    (
        "OPLS and kernel PLS extensions passed their separate deterministic "
        "reliability study",
        "OPLS and kernel PLS extensions met the prespecified numerical "
        "tolerances in their separate deterministic reliability study",
    ),
    (
        "while preserving the deterministic estimator",
        "while meeting the prespecified deterministic numerical tolerances "
        "against de Jong SIMPLS",
    ),
    (
        "preserving the deterministic estimator when IRLBA is used",
        "and deterministic IRLBA comparisons are evaluated against "
        "prespecified numerical tolerances",
    ),
    (
        "one power iteration passed 101/117 component-level checks, whereas "
        "two power iterations passed 117/117",
        "one power iteration met the prespecified numerical tolerances in "
        "101/117 component-level checks, whereas two power iterations met "
        "them in 117/117",
    ),
    (
        "deterministic IRLBA SIMPLS preserved de Jong's estimator",
        "deterministic IRLBA SIMPLS met the prespecified deterministic "
        "numerical tolerances against de Jong SIMPLS",
    ),
    (
        "alongside pass/fail status",
        "alongside the tolerance-status classification",
    ),
    (
        "deterministic estimator preservation",
        "deterministic estimator comparison",
    ),
    (
        "OPLS and kernel PLS extensions passed all 66 setting/task "
        "comparisons and all 1,540 fold-component fits",
        "OPLS and kernel PLS extensions met the prespecified numerical "
        "tolerances in all 66 setting/task comparisons and all 1,540 "
        "fold-component fits",
    ),
    (
        "the corresponding SIMPLS setting passed only 101/117 approximation "
        "checks",
        "the corresponding SIMPLS setting met the prespecified numerical "
        "tolerances in only 101/117 approximation checks",
    ),
    (
        "estimator-preservation claims",
        "claims of numerical equivalence",
    ),
    (
        "its numerical status was 101/117 checks passed",
        "it met the prespecified numerical tolerances in 101/117 checks",
    ),
    (
        "with 117/117 audit checks passed",
        "and met the prespecified numerical tolerances in 117/117 audit checks",
    ),
    (
        "this setting passed 101/117 approximation checks",
        "this setting met the prespecified numerical tolerances in 101/117 "
        "approximation checks",
    ),
    (
        "Because this setting passed 101/117 checks",
        "Because this setting met the prespecified numerical tolerances in "
        "only 101/117 checks",
    ),
    (
        "Audit status: 101/117 checks passed",
        "Audit status: met the prespecified numerical tolerances in 101/117 "
        "checks",
    ),
    (
        "the related one-power SIMPLS configuration passed 101/117 checks",
        "the related one-power SIMPLS configuration met the prespecified "
        "numerical tolerances in 101/117 checks",
    ),
    (
        "was not separately shown to pass the numerical audit",
        "was not separately shown to meet the prespecified numerical "
        "tolerances",
    ),
    (
        "the related setting passed 101/117 checks",
        "the related setting met the prespecified numerical tolerances in "
        "101/117 checks",
    ),
    (
        "not an accuracy or estimator-preservation claim",
        "not an accuracy or numerical-equivalence claim",
    ),
    (
        "reuse of sequential quantities preserved de Jong SIMPLS",
        "reuse of sequential quantities met the prespecified deterministic "
        "numerical tolerances against de Jong SIMPLS",
    ),
    (
        "Only oversampling 10 with two power iterations passed all 117 "
        "broad-audit checks",
        "Only oversampling 10 with two power iterations met the prespecified "
        "numerical tolerances in all 117 broad-audit checks",
    ),
    (
        "Passing denotes numerical non-inferiority within the stated",
        "Meeting these tolerances denotes numerical non-inferiority within "
        "the stated",
    ),
    (
        "The faster one-power NMR and ImageNet workflows passed 101/117 at "
        "setting level",
        "The faster one-power NMR and ImageNet workflows met the prespecified "
        "numerical tolerances in 101/117 checks at setting level",
    ),
    (
        "preserve the input precision through their supported PLS arithmetic",
        "keep their supported PLS arithmetic in the input precision",
    ),
    (
        "SIMPLS estimator-preservation and rSVD reliability analyses",
        "SIMPLS estimator-comparison and rSVD reliability analyses",
    ),
    (
        "All passed the prespecified tolerances",
        "All met the prespecified numerical tolerances",
    ),
    (
        "one power iteration passed 101 of 117 component-level checks, whereas "
        "two power iterations passed 117 of 117",
        "one power iteration met the prespecified numerical tolerances in 101 "
        "of 117 component-level checks, whereas two power iterations met them "
        "in 117 of 117",
    ),
    (
        "The focused CUDA audit passed all evaluated points",
        "The focused CUDA audit met the prespecified numerical tolerances at "
        "all evaluated points",
    ),
    (
        "Deterministic estimator preservation and approximate-solver "
        "agreement are separate questions",
        "Deterministic estimator comparison and approximate-solver agreement "
        "are separate questions",
    ),
    (
        "Primary check that backend changes preserve model output",
        "Primary check of model-output agreement under backend changes",
    ),
    (
        "Deterministic SIMPLS preserves de Jong SIMPLS within stated "
        "tolerances",
        "Deterministic SIMPLS meets the prespecified numerical tolerances "
        "against de Jong SIMPLS",
    ),
    (
        "Preserved",
        "Met deterministic tolerances",
    ),
    (
        "Endpoint pass",
        "Endpoint tolerance status",
    ),
    (
        "Passes",
        "Tolerance checks met",
    ),
    (
        "route passes the SIMPLS method identifier",
        "route supplies the SIMPLS method identifier",
    ),
)


def replace_in_paragraph(paragraph):
    original = paragraph.text
    revised = original
    for old, new in REPLACEMENTS:
        revised = revised.replace(old, new)
    if revised == original:
        return False
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(revised)
    return True


def revise_document(document):
    changes = 0
    for paragraph in document.paragraphs:
        changes += int(replace_in_paragraph(paragraph))
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    changes += int(replace_in_paragraph(paragraph))
    normalize_submission_terminology(document)
    return changes


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    main_document = Document(MAIN_SOURCE)
    main_changes = revise_document(main_document)
    main_document.save(MAIN_OUTPUT)

    supplement_document = Document(SUPP_SOURCE)
    supplement_changes = revise_document(supplement_document)
    supplement_document.save(SUPP_OUTPUT)

    print(f"main_changes={main_changes}")
    print(f"supplement_changes={supplement_changes}")
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)


if __name__ == "__main__":
    main()
