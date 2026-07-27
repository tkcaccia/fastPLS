#!/usr/bin/env python3

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle54"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle55"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle54_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle54_0.99.6_20260726.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle55_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle55_0.99.6_20260726.docx"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def revise_main():
    document = Document(MAIN_SOURCE)
    opening = find_paragraph(
        document,
        "Supplementary Tables S14-S17 report matched CPU and CUDA results",
    )

    opening.text = (
        "We evaluated fastPLS in three stages. First, we tested whether the "
        "accelerated implementations preserved the predictive behaviour of the "
        "underlying PLS estimators across heterogeneous biomedical and molecular "
        "datasets. Second, we examined how matrix shape, component count, "
        "numerical precision, and execution backend affected total runtime and "
        "memory. Third, NMR and ImageNet were used as targeted stress tests for "
        "extreme response dimensionality and million-sample feature matrices, "
        "respectively. The benchmark-wide analysis comprised twelve supervised "
        "tasks and four PLS families. For every family and dataset, component "
        "selection used only the training data before matched evaluation on the "
        "same held-out observations."
    )

    benchmark_summary = insert_paragraph_after(
        opening,
        (
            "Predictive performance was generally stable across matched CPU and "
            "CUDA workflows. All 46 evaluated backend pairs completed; NMR OPLS "
            "and kernel PLS were explicitly marked as not evaluated. Predictive "
            "intervals overlapped for every pair, and 26 of 36 classification "
            "pairs had identical observed accuracy; the largest difference was "
            "2.27 percentage points for TCGA-BRCA OPLS. Backend choice affected "
            "computation more strongly. CPU was faster in 31 of 46 pairs, mainly "
            "for small or moderate matrices, whereas CUDA was faster in 15. The "
            "largest CUDA advantage occurred on CIFAR-100 (up to 14.55-fold); "
            "for NMR, CUDA accelerated the selected PLS-SVD and SIMPLS workflows "
            "by 2.72- and 5.39-fold, respectively."
        ),
        style="Body Text",
    )

    details = insert_paragraph_after(
        benchmark_summary,
        (
            "Complete family-specific CPU/CUDA results are reported in "
            "Supplementary Tables S14-S17, with CPU/Metal validation in "
            "Supplementary Figure S25 and selected-point computational summaries "
            "in Figure S38. Component-selection paths and their associations "
            "with prediction, runtime, host RSS, and GPU memory are provided in "
            "Supplementary Tables S42-S46 and Figures S26-S37. Of the 46 "
            "family-specific selections, 33 were at a tested-grid boundary or "
            "the PLS-SVD response-rank ceiling and are therefore described as "
            "best within the evaluated grid rather than global optima. Headline "
            "rSVD workflows used oversampling 10, one power iteration, and seeds "
            "124-126. Numerical-audit status is reported separately from "
            "execution status because the one-power SIMPLS setting passed 101 "
            "of 117 prespecified approximation checks. Incremental host RSS "
            "ranged from 2 to 1,037 MB and represented a median 41.3% of the "
            "absolute process peak. Sampled GPU increments ranged from 192 to "
            "3,414 MB and included CUDA context, library state, allocator pools, "
            "data, models, and workspaces; they are process-level footprints, "
            "not isolated algorithm workspaces."
        ),
        style="Body Text",
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - revised Results introduction"
    )
    document.save(MAIN_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    shutil.copy2(SUPP_SOURCE, SUPP_OUT)
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
