#!/usr/bin/env python3

from pathlib import Path

from docx import Document


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle46"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle47"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle46_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle46_0.99.6_20260726.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle47_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle47_0.99.6_20260726.docx"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def remove_following_table(paragraph):
    element = paragraph._p.getnext()
    while element is not None and not element.tag.endswith("}tbl"):
        element = element.getnext()
    if element is None:
        raise RuntimeError("Expected a table after the caption")
    element.getparent().remove(element)


def replace_text(paragraph, replacements):
    text = paragraph.text
    for old, new in replacements:
        text = text.replace(old, new)
    if text != paragraph.text:
        style = paragraph.style
        paragraph.clear()
        paragraph.style = style
        paragraph.add_run(text)


def revise_main():
    document = Document(MAIN_SOURCE)

    caption = find_paragraph(
        document,
        "Table 1. Paired CPU/CUDA biomedical workflow benchmark",
    )
    remove_following_table(caption)
    remove_paragraph(caption)

    targeted_replacements = [
        (
            "Table 1 reports matched CPU and CUDA results for the complete "
            "twelve-task biomedical benchmark",
            "Supplementary Tables S14-S17 report matched CPU and CUDA results, "
            "split by PLS family, for the complete twelve-task biomedical benchmark",
        ),
        (
            "The table reports numerical-audit status separately from execution status",
            "The four tables report numerical-audit status separately from execution status",
        ),
        (
            "the approximate rSVD rows in Table 1",
            "the approximate rSVD rows in Supplementary Tables S16-S17",
        ),
        (
            "The kernel-PLS rows in Table 1",
            "The kernel-PLS rows in Supplementary Table S17",
        ),
        (
            "Table 1 and Supplementary Figure S25 report the same selected-setting "
            "principle across datasets",
            "Supplementary Tables S14-S17 and Figure S25 report the same "
            "selected-setting principle across datasets",
        ),
        (
            "Accordingly, Table 1 compares reproducible workflows",
            "Accordingly, Supplementary Tables S14-S17 compare reproducible workflows",
        ),
        (
            "Supplementary Tables S14-S17 compare reproducible workflows at "
            "prespecified training-selected operating points; it does not establish",
            "Supplementary Tables S14-S17 compare reproducible workflows at "
            "prespecified training-selected operating points; they do not establish",
        ),
    ]
    for paragraph in document.paragraphs:
        replace_text(paragraph, targeted_replacements)

    renumber = [
        ("Table 3", "__MAIN_TABLE_2__"),
        ("Table 2", "__MAIN_TABLE_1__"),
        ("__MAIN_TABLE_2__", "Table 2"),
        ("__MAIN_TABLE_1__", "Table 1"),
    ]
    for paragraph in document.paragraphs:
        replace_text(paragraph, renumber)

    document.core_properties.title = (
        "fastPLS CMPB manuscript - detailed benchmark tables in supplement"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)

    heading = find_paragraph(
        document,
        "S15. Complete selected-point multi-dataset benchmark",
    )
    heading.text = (
        "S15. Complete selected-point benchmark split by PLS family"
    )

    intro = heading._p.getnext()
    if intro is None or not intro.tag.endswith("}p"):
        raise RuntimeError("Expected explanatory paragraph after S15 heading")
    intro_paragraph = next(
        paragraph
        for paragraph in document.paragraphs
        if paragraph._p is intro
    )
    replace_text(
        intro_paragraph,
        [
            (
                "The tables below retain",
                "Supplementary Tables S14-S17 replace the combined main-text "
                "benchmark table and retain",
            )
        ],
    )

    for paragraph in document.paragraphs:
        replace_text(
            paragraph,
            [
                (
                    "shown in Table 1",
                    "reported in Supplementary Tables S14-S17",
                ),
                (
                    "workflow in Table 1",
                    "workflow in Supplementary Tables S14-S17",
                ),
            ],
        )

    document.core_properties.title = (
        "fastPLS CMPB supplement - family-specific benchmark tables"
    )
    document.save(SUPP_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
