#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle38"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle39"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle38_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle38_0.99.6_20260726.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle39_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle39_0.99.6_20260726.docx"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_paragraph(document, prefix, text):
    paragraph = find_paragraph(document, prefix)
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, text, size=4.0):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.name = "Arial"
            run.font.size = Pt(size)


def insert_solver_control_row(table):
    row = table.add_row()
    table._tbl.remove(row._tr)
    table.rows[0]._tr.addnext(row._tr)
    values = [
        "rSVD controls and numerical audit",
        (
            "o=10; power=1; seeds=124/125/126. "
            "Audit: not formally assessed for PLS-SVD; workflow-only."
        ),
        (
            "o=10; power=1; seeds=124/125/126. "
            "Setting-level audit: 101/117 passed (16 failed); workflow-only."
        ),
        (
            "o=10; power=1; seeds=124/125/126. "
            "Audit: not formally assessed for OPLS; workflow-only."
        ),
        (
            "o=10; power=1; seeds=124/125/126. "
            "Audit: not formally assessed for kernel PLS; workflow-only."
        ),
    ]
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value)
        shade_cell(cell, "E7EEF8")
    row.cells[0].paragraphs[0].runs[0].bold = True


def revise_main():
    document = Document(MAIN_SOURCE)

    replace_paragraph(
        document,
        "Table 1 and Figure 2 show",
        (
            "Table 1 and Figure 2 show both matched CPU and CUDA backends for the "
            "complete twelve-task biomedical benchmark at the training-set-selected "
            "component count within each prespecified grid and PLS family. All "
            "headline rSVD workflows used oversampling 10, one power iteration, "
            "and replicate seeds 124, 125, and 126. The table reports numerical-audit "
            "status separately from execution status: the one-power SIMPLS setting "
            "passed 101 of 117 prespecified approximation checks and is therefore "
            "workflow-only evidence, while PLS-SVD, OPLS, and kernel PLS were not "
            "formally audited for estimator equivalence. A dagger identifies a lower "
            "or upper tested-grid boundary or a PLS-SVD response-rank limit; these "
            "values are not claimed as global optima. Retina and Tabula Muris, the "
            "two single-cell resources, are reported separately under their specific "
            "names. All 46 evaluated CPU/CUDA pairs completed, while NMR OPLS and "
            "kernel PLS are marked as not evaluated rather than omitted. The "
            "predictive intervals overlapped for every pair; 26 of 36 classification "
            "pairs had identical observed accuracy, and the largest difference was "
            "2.27 percentage points for TCGA-BRCA OPLS. Backend choice affected "
            "computation more strongly. CPU was faster in 31 of 46 pairs, primarily "
            "on small or modest matrix problems, whereas CUDA was faster in 15 pairs. "
            "The largest CUDA advantages occurred on CIFAR-100 (up to 14.55-fold); "
            "on NMR, CUDA accelerated selected PLS-SVD and SIMPLS by 2.72- and "
            "5.39-fold, respectively. The memory audit separates data and runtime "
            "overhead from fitting allocations. Incremental host RSS ranged from 2 "
            "to 1037 MB and represented a median 41.3% of the absolute process peak. "
            "All measured CUDA pre-fit process baselines were 0 MB because CUDA "
            "contexts were initialized during fitting; incremental device memory "
            "ranged from 192 to 3414 MB. Thus absolute host RSS documents whether a "
            "workflow fits on the machine, whereas the baseline-corrected values "
            "expose algorithmic and prediction workspace costs."
        ),
    )

    replace_paragraph(
        document,
        "Table 1. Paired CPU/CUDA",
        (
            "Table 1. Paired CPU/CUDA biomedical workflow benchmark at the "
            "training-set-selected component count within each evaluated component "
            "grid. Each cell retains both matched backends at the same family-specific "
            "component count and reports the outer-test metric with 95% interval, "
            "median total fitting-plus-prediction time with run IQR, completed runs, "
            "and memory in MB. The solver-control row gives oversampling (o), power "
            "iterations, replicate seeds, and numerical-audit status for every "
            "headline rSVD row. Execution status OK means that a run completed and "
            "passed structural and finiteness checks; it does not mean that the "
            "randomized result passed an estimator-equivalence audit. H and G denote "
            "host and process-specific GPU memory; B/P/Δ denotes the immediately "
            "pre-fit baseline, absolute isolated-process peak, and incremental "
            "fit-window peak above baseline, respectively. †Lower or upper tested-grid "
            "boundary, or response-rank limit for PLS-SVD; these entries are not "
            "global optima. Accuracy uses Wilson intervals and RMSD uses "
            "10,000-resample held-out-sample bootstrap intervals, conditional on the "
            "fixed outer split. NMR OPLS and kernel PLS are explicitly labelled not "
            "evaluated. OPLS uses one prespecified orthogonal component; A denotes "
            "the total budget (A-1 predictive plus one orthogonal). Kernel PLS uses "
            "the linear kernel, so gamma, degree, and intercept are not applicable. "
            "For CBMC CITE-seq, RMSD is in original ADT assay-count units per cell "
            "after response-mean restoration and is pooled across 862 test cells and "
            "10 markers."
        ),
    )
    insert_solver_control_row(document.tables[0])

    replace_paragraph(
        document,
        "Figure 3. Fixed-complexity NMR",
        (
            "Figure 3. Fixed-complexity NMR implementation benchmark at 100 "
            "components. This common component count was imposed to compare "
            "implementations and was not the family-specific predictive selection. "
            "(A) Observed and rSVD-based SIMPLS prediction of the held-out spectrum "
            "across the full chemical-shift range. (B) Enlarged 0.5-1.7 ppm region. "
            "(C) Distribution of per-spectrum RMSD across the 321 held-out spectra "
            "for the deposited PLS-SVD reference using IRLBA and fastPLS PLS-SVD and "
            "SIMPLS implementations using rSVD on CPU and CUDA. (D) Median total "
            "fitting-plus-prediction time, peak host RSS, and sampled peak GPU memory "
            "across three isolated runs. The rSVD runs used oversampling 10, one "
            "power iteration, and seeds 124, 125, and 126; their numerical-audit "
            "status is workflow-only, not deterministic estimator equivalence. The "
            "representative spectrum was selected mechanically as the test spectrum "
            "whose RMSD was closest to the held-out median. All comparisons used "
            "float64 data and identical routinely preprocessed training and held-out "
            "spectra."
        ),
    )

    replace_paragraph(
        document,
        "Table 2. Exploratory matched ImageNet/DINOv2",
        (
            "Table 2. Exploratory matched ImageNet/DINOv2 retrieval on a fixed random "
            "1,000,000/281,167 development split drawn without replacement from a "
            "pooled 1,281,167-image embedding archive; this is not the canonical "
            "ImageNet train/validation split. PCA and PLS-SVD rSVD used oversampling "
            "10, one power iteration, and seed 123 for the displayed path; seeds 456 "
            "and 789 were complete repeated representation fits reported in "
            "Supplementary Table S23. Numerical-audit status is not formally assessed "
            "against a deterministic decomposition, so all rSVD representation rows "
            "are exploratory workflow results. Exact CUDA cosine kNN used k=10. The "
            "value of k and the 50/100/200-dimensional path were fixed for this "
            "matched run but informed by earlier exploration on the same development "
            "holdout, not nested cross-validation. Brackets are Wilson 95% intervals "
            "conditional on this split. They do not account for split or "
            "hyperparameter-selection uncertainty, and the small top-5 difference is "
            "not interpreted as improvement. Transformation includes fitting and "
            "train/holdout projection; query is median (IQR) over three exact FAISS "
            "runs. H/G are peak host RSS and sampled GPU memory in MB."
        ),
    )

    replace_paragraph(
        document,
        "Figure 4. Exploratory matched ImageNet/DINOv2",
        (
            "Figure 4. Exploratory matched ImageNet/DINOv2 retrieval on the "
            "noncanonical random development split. (A) Descriptive top-1 and top-5 "
            "accuracy for raw embeddings and 50-, 100-, and 200-dimensional PCA and "
            "PLS representations. (B) Holdout projection plus exact FAISS query time. "
            "(C) End-to-end representation fitting, train/holdout projection, and "
            "query time. (D) Peak host RSS and sampled GPU memory. The displayed rSVD "
            "path used oversampling 10, one power iteration, and seed 123; its audit "
            "status is exploratory/not formally assessed. Two additional seeded "
            "PLS/PCA fits (456 and 789) are summarized in Supplementary Table S23. "
            "Repeated seeds do not replace independent split or hyperparameter-"
            "selection uncertainty."
        ),
    )

    document.core_properties.title = (
        "fastPLS CMPB manuscript - explicit rSVD controls and audit status"
    )
    document.save(MAIN_OUT)


def append_audit_text(document, prefix, addition):
    paragraph = find_paragraph(document, prefix)
    text = paragraph.text.rstrip()
    if not text.endswith("."):
        text += "."
    replace_paragraph(document, prefix, f"{text} {addition}")


def revise_supplement():
    document = Document(SUPP_SOURCE)

    replace_paragraph(
        document,
        "The tables below retain both CPU and CUDA rSVD rows",
        (
            "The tables below retain both CPU and CUDA rSVD rows, all failed or "
            "unavailable executions, precision, and resource measurements. Every "
            "selected-point row used oversampling 10, one power iteration, and "
            "replicate seeds 124, 125, and 126. Execution status and numerical-audit "
            "status are distinct: OK records completion plus structural/finiteness "
            "checks, whereas the one-power SIMPLS setting passed only 101 of 117 "
            "prespecified approximation checks. It is therefore labelled workflow-"
            "only. PLS-SVD, OPLS, and kernel PLS rows were not formally audited for "
            "estimator equivalence. Component counts were selected using training "
            "data only, except for the explicitly labelled exploratory ImageNet "
            "stress point. Argmax decoding was used for classification and the "
            "prespecified regression metric for numeric responses. Results involving "
            "LDA, external packages, or different PLS families are complete-workflow "
            "comparisons and are interpreted separately."
        ),
    )

    common = (
        " rSVD controls: oversampling 10, one power iteration, seeds "
        "124/125/126. OK is execution status, not numerical equivalence."
    )
    append_audit_text(
        document,
        "Table S14. PLS-SVD selected-point benchmark.",
        common.strip() + " Numerical audit: not formally assessed; workflow-only.",
    )
    append_audit_text(
        document,
        "Table S15. SIMPLS selected-point benchmark.",
        common.strip()
        + " Setting-level numerical audit: 101/117 passed and 16 failed; "
        "workflow-only.",
    )
    append_audit_text(
        document,
        "Table S16. OPLS selected-point benchmark.",
        common.strip() + " Numerical audit: not formally assessed; workflow-only.",
    )
    append_audit_text(
        document,
        "Table S17. kernel PLS selected-point benchmark.",
        common.strip() + " Numerical audit: not formally assessed; workflow-only.",
    )

    append_audit_text(
        document,
        "Table S18. Complete exploratory ImageNet retrieval results",
        (
            "rSVD controls: oversampling 10, one power iteration, displayed seed "
            "123; numerical audit not formally assessed."
        ),
    )
    append_audit_text(
        document,
        "Table S22. Repeated CPU-versus-Metal validation",
        (
            "Real-data rSVD controls were oversampling 10, one power iteration, "
            "and seeds 201/202/203; NMR used seed 123. Audit status: exploratory "
            "workflow-only; CIFAR-100 one-power backend discrepancies failed the "
            "agreement criterion, and float64 Metal PLS-SVD is excluded from "
            "equivalence claims."
        ),
    )
    append_audit_text(
        document,
        "Table S23. Repeated full ImageNet/DINOv2 representation fits.",
        (
            "rSVD controls: oversampling 10, one power iteration, seeds 123/456/789. "
            "Numerical audit: not formally assessed against a deterministic "
            "decomposition; exploratory workflow-only."
        ),
    )
    append_audit_text(
        document,
        "Table S28. Approximate float64 rSVD workflow comparison.",
        (
            "rSVD controls: oversampling 10, one power iteration, seeds "
            "124/125/126. Setting-level SIMPLS audit: 101/117 passed and 16 failed; "
            "these rows are not estimator-equivalence evidence."
        ),
    )

    document.core_properties.title = (
        "fastPLS CMPB supplementary material - explicit rSVD controls and audit status"
    )
    document.save(SUPP_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
