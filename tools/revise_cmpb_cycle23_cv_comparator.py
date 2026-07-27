#!/usr/bin/env python3

from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_BREAK


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle22"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle23"
EVIDENCE = ROOT / "benchmark_results" / "cv_compiled_vs_r_loop_20260725"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle22_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle22_0.99.6_20260725.docx"
RESPONSE_SOURCE = SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle22_20260725.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle23_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle23_0.99.6_20260725.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle23_20260725.docx"

spec = spec_from_file_location(
    "cycle22_helpers",
    ROOT / "tools" / "revise_cmpb_cycle22_imagenet_qualification.py",
)
c22 = module_from_spec(spec)
spec.loader.exec_module(c22)
c16 = c22.c16


def replace(document, prefix, text):
    paragraph = c16.find_paragraph(document, prefix)
    paragraph.text = text


def evidence_rows():
    data = pd.read_csv(EVIDENCE / "cv_compiled_vs_r_loop_summary.csv")
    rows = []
    for _, x in data.iterrows():
        backend = f'{x["backend"]}/{x["svd_method"]}'
        metric = (
            f'{x["compiled_metric"]:.4f}/{x["r_loop_metric"]:.4f}'
        )
        agreement = (
            f'{x["prediction_agreement_min"]:.4f}'
            if pd.notna(x["prediction_agreement_min"])
            else f'r={x["prediction_correlation_min"]:.4f}'
        )
        rows.append(
            (
                x["dataset"].replace("metref", "MetRef").replace(
                    "retina", "Retina"
                ).replace("prism", "PRISM"),
                backend,
                str(int(x["ncomp"])),
                f'{x["compiled_median_sec"]:.3f} '
                f'({x["compiled_iqr_sec"]:.3f})',
                f'{x["r_loop_median_sec"]:.3f} ({x["r_loop_iqr_sec"]:.3f})',
                f'{x["speedup_median"]:.2f}',
                metric,
                agreement,
            )
        )
    return rows


def revise_main():
    document = Document(MAIN_SOURCE)
    replace(
        document,
        "Across 150 completed argmax/LDA comparisons",
        "Cross-validation acceleration was tested against an equivalent "
        "R-level loop rather than inferred from the cost relative to one fit. "
        "The comparator called the same fastPLS estimator independently in "
        "each of ten prespecified folds and returned the same out-of-fold "
        "outputs. For deterministic SIMPLS/IRLBA, compiled fold management was "
        "1.35-fold faster on MetRef, 1.26-fold faster on Retina, and 1.07-fold "
        "faster on PRISM (three paired runs). Fold partitions, predictions, "
        "accuracy or RMSD, and effective component count were identical; "
        "regression predictions correlated 1.000. CPU rSVD gave 1.21-1.39-fold "
        "speedups with the same predictions in these cases. CUDA was not "
        "uniformly advantageous for the small MetRef workload, emphasizing "
        "that compiled orchestration and accelerator execution are separate "
        "claims. The previously reported 3.78-13.20-fold ratio is retained only "
        "as the cost of ten-fold validation relative to one fit, not as an "
        "acceleration estimate (Supplementary Tables S20 and S32).",
    )
    document.core_properties.title = (
        "fastPLS CMPB manuscript - matched cross-validation comparator"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)
    replace(
        document,
        "The archived 10-fold benchmark compared",
        "The archived benchmark compared one optimized ten-fold validation "
        "call with one fit-plus-prediction workflow. It therefore quantifies "
        "validation overhead, not acceleration. A new matched comparator "
        "freezes fold groups and compares the compiled engine with an explicit "
        "R-level loop that calls the same public fastPLS estimator once per "
        "fold; its results are reported in Section S29 and Table S32.",
    )

    document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    document.add_heading(
        "S29. Matched compiled versus R-level cross-validation", level=1
    )
    document.add_paragraph(
        "Ten fold groups were generated once with seed 123 and passed as "
        "constraints to both engines. The compiled route performed fold "
        "construction, fitting, prediction, and scoring inside the compiled "
        "backend. The reference route iterated over the same groups in R and "
        "called fastPLS::pls() with identical method, scaling, component count, "
        "SVD controls, classifier, and fold-specific seed. Runtime excludes "
        "dataset loading and fold generation. Three isolated paired processes "
        "were used per setting; order was alternated. Partition identity was "
        "verified from fold contingency, and out-of-fold predictions were "
        "compared row by row."
    )
    document.add_paragraph(
        "Table S32. Matched ten-fold SIMPLS cross-validation. Times are median "
        "(IQR) seconds over three runs. Speedup is R-loop time divided by "
        "compiled time. Metric cells show compiled / R-loop accuracy or RMSD. "
        "Agreement is label agreement for classification and Pearson "
        "correlation for regression.",
        style="Caption",
    )
    c16.add_table(
        document,
        [
            "Dataset",
            "Backend/SVD",
            "k",
            "Compiled s",
            "R loop s",
            "Speedup",
            "Metric C/R",
            "Prediction agreement",
        ],
        evidence_rows(),
        font_size=5.7,
    )
    document.add_paragraph(
        "The deterministic IRLBA rows are the primary estimator-matched "
        "evidence. CPU rSVD is reported separately because it is approximate, "
        "although predictions agreed exactly in these evaluated settings. The "
        "CUDA MetRef row had identical folds but 93.47% label agreement and a "
        "0.69-percentage-point accuracy difference between the compiled and "
        "fold-wise GPU routes; it is not treated as estimator-equivalence "
        "evidence. Its median speedup was below one because launch and workspace "
        "costs dominated this small problem. The explicit R-loop comparator for "
        "CIFAR-100 at 50 components exceeded the 31 GiB host-memory limit and "
        "was recorded as infeasible; no speedup was inferred from that failure."
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    document.add_heading(
        "24. Cross-validation lacked an equivalent comparator", level=1
    )
    document.add_paragraph(
        "Response: Agreed. The previous ratio compared ten-fold validation with "
        "one fit and measured overhead, not acceleration. We added a matched "
        "benchmark against an explicit R-level loop that calls the same "
        "fastPLS::pls() estimator on the same prespecified folds and returns "
        "the same outputs. Deterministic SIMPLS/IRLBA showed 1.35-fold, "
        "1.26-fold, and 1.07-fold compiled speedups on MetRef, Retina, and "
        "PRISM, respectively, with identical predictions and metrics. CPU rSVD "
        "showed 1.21-1.39-fold speedups and is labelled approximate. CUDA was "
        "not faster on small MetRef and is reported separately. The CIFAR-100 "
        "R-loop comparator exceeded available RAM and is retained as an "
        "explicit feasibility failure rather than converted into a speedup. "
        "Methods, Results, Supplementary Section S29, Table S32, benchmark "
        "scripts, and machine-readable tables were updated."
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    revise_response()
    print(MAIN_OUT)
    print(SUPP_OUT)
    print(RESPONSE_OUT)


if __name__ == "__main__":
    main()
