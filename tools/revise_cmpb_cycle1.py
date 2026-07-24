"""Create the first evidence-based CMPB manuscript-review cycle drafts.

The source manuscript is preserved.  This script makes only local prose
replacements and inserts clearly labelled preliminary validation text so that
claims do not outrun the available benchmark evidence.
"""

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Desktop/fastPLS/CMPB_rewrite_20260721")
OUT = Path("/Users/stefano/Desktop/fastPLS/CMPB_rewrite_20260724_cycle1")
OUT.mkdir(parents=True, exist_ok=True)

MAIN_SRC = ROOT / "fastPLS_CMPB_main_current_0.99.4_20260721.docx"
SUPP_SRC = ROOT / "fastPLS_CMPB_supplement_current_0.99.4_20260721.docx"
REVIEW_SRC = ROOT / "fastPLS_CMPB_independent_reviewer_report_20260721.docx"
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle1_0.99.6_20260724.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle1_0.99.6_20260724.docx"
REVIEW_OUT = OUT / "fastPLS_CMPB_independent_reviewer_report_cycle1_20260724.docx"


def paragraph_after(paragraph, text, style=None):
    node = OxmlElement("w:p")
    paragraph._p.addnext(node)
    inserted = Paragraph(node, paragraph._parent)
    if style:
        inserted.style = style
    inserted.add_run(text)
    return inserted


def replace_paragraph(doc, startswith, replacement):
    for p in doc.paragraphs:
        if p.text.startswith(startswith):
            p.clear()
            p.add_run(replacement)
            return p
    raise RuntimeError(f"No paragraph starts with: {startswith}")


def insert_after_paragraph(doc, startswith, text, style=None):
    for p in doc.paragraphs:
        if p.text.startswith(startswith):
            return paragraph_after(p, text, style=style)
    raise RuntimeError(f"No paragraph starts with: {startswith}")


def revise_main():
    copy2(MAIN_SRC, MAIN_OUT)
    doc = Document(MAIN_OUT)

    replace_paragraph(
        doc,
        "Methods: fastPLS provides",
        "Methods: fastPLS provides PLS-SVD, SIMPLS, orthogonal PLS (OPLS), "
        "and kernel PLS through one R interface. Its SIMPLS implementation "
        "preserves de Jong’s sequential estimator while reusing the preceding "
        "direction, caching rank-one deflation products, exploiting favourable "
        "cross-product shapes, and updating coefficient and prediction paths "
        "incrementally. The package combines randomized singular value "
        "decomposition (rSVD), bundled CPU IRLBA, implicit cross-covariance "
        "products, compact prediction, compiled validation, and CPU, NVIDIA CUDA, "
        "and Apple Metal backends. Selected float32 execution paths are available "
        "and are evaluated separately from the double-precision reference. "
        "Performance was evaluated using fixed data partitions, prediction "
        "metrics, total fitting-plus-prediction time, and peak memory."
    )
    replace_paragraph(
        doc,
        "Results: Accelerated backends retained",
        "Results: On controlled synthetic matrices, accelerated SIMPLS matched "
        "pls::simpls.fit closely: prediction correlations were at least 0.9999999, "
        "the largest relative prediction difference was 3.86×10−4, and the maximum "
        "principal angle between fitted predictive subspaces was 0.050°. "
        "Accelerated backends retained predictive behaviour while reducing "
        "computational cost in selected large or high-response-dimensional settings. "
        "The optimized SIMPLS path narrowed the runtime gap to one-shot PLS-SVD "
        "without changing sequential SIMPLS deflation. fastPLS processed multivariate "
        "NMR and a million-sample foundation-model embedding stress test based on "
        "ImageNet/DINOv2. Under the documented interfaces and resource limits, the "
        "evaluated external R workflows did not complete the latter task."
    )
    replace_paragraph(
        doc,
        "Here we present fastPLS, which accelerates",
        "Here we present fastPLS, which accelerates established PLS estimators "
        "rather than redefining them. The principal methodological contribution is "
        "a reorganized SIMPLS execution path that reuses information across sequential "
        "components and across requested component prefixes. The package also separates "
        "the outer PLS estimator from the truncated-SVD and hardware layers, supports "
        "compact and implicit products when large matrices would otherwise dominate "
        "memory, and provides a consistent prediction and validation interface. The "
        "GPL-3 R package is complemented by low-level C++ implementations maintained "
        "in the MIT-licensed kodama-cpp project; accelerator support described here "
        "belongs to the R-package backend layer."
    )
    replace_paragraph(
        doc,
        "The R package includes bundled IRLBA code",
        "The R package includes bundled IRLBA code and is distributed under GPL-3. "
        "The low-level C++ implementations are also maintained in the MIT-licensed "
        "kodama-cpp project. This manuscript evaluates the R-package CPU, CUDA, and "
        "Metal backend layer; it does not claim that the standalone core independently "
        "provides every accelerator path."
    )
    insert_after_paragraph(
        doc,
        "Across completed benchmarks, the numerical backend",
        "A controlled numerical check directly compared the retained non-optimized "
        "SIMPLS route and the state-reuse route with pls::simpls.fit on fixed "
        "well-conditioned, ill-conditioned, and rank-deficient multivariate regression "
        "matrices (ntrain=180, p=60, q=4, five components). Across these cases, the "
        "optimized route had prediction correlation ≥0.9999999 relative to the "
        "independent reference, relative prediction error ≤3.86×10−4, coefficient "
        "relative error ≤5.49×10−4, and maximum predictive-subspace angle ≤0.050°. "
        "The retained and optimized fastPLS routes were indistinguishable at the "
        "reported precision in this small controlled study. The full benchmark and "
        "real-data confirmation remain reported separately."
    )
    replace_paragraph(
        doc,
        "In a fixed-score validation of the revised LDA path",
        "In a fixed-score validation of the revised LDA path, float32 CPU and CUDA "
        "predictions agreed exactly across MetRef, CIFAR-100, and SingleCell at 2, 5, "
        "10, and 20 components, with no factorization failures. A preliminary NMR "
        "precision control (ntrain=1200, ntest=321, p=13000, q=5000, ten components) "
        "preserved RMSD between float32 and float64 SIMPLS (6.4850×10−5 versus "
        "6.4847×10−5), but float32 was slower in the current implementation (60.8 s "
        "CPU and 51.8 s CUDA versus 7.54 s float64 CPU). Float32 is therefore reported "
        "as a memory-format and numerical-capability evaluation, not as a general speed "
        "advantage, until the full high-response execution path is further optimized."
    )
    replace_paragraph(
        doc,
        "Float32 is important for memory-limited biomedical",
        "Float32 can reduce input and workspace storage for supported routes, but it "
        "is not yet a universal performance benefit. Current native support is "
        "restricted to selected PLS-SVD and SIMPLS paths, and the preliminary "
        "high-response NMR control showed preserved prediction but slower float32 "
        "execution. Reproducibility should therefore be judged by prediction agreement, "
        "predictive metrics, selected components, numerical failures, and latent-subspace "
        "agreement rather than by assuming identical stochastic singular vectors or a "
        "speedup from reduced precision."
    )
    replace_paragraph(
        doc,
        "The fastPLS R package, benchmark workflows",
        "The fastPLS R package, benchmark workflows, and analysis scripts are available "
        "at https://github.com/tkcaccia/fastPLS (review-cycle package commit "
        "72e178b9e3c9510dc86c4b287d68b9c717f9fdf5). Low-level C++ implementations "
        "are maintained at https://github.com/tkcaccia/kodama-cpp. Public data should "
        "be obtained from the sources cited in the Supplementary Material. Restricted "
        "data are not redistributed; the final submission will provide acquisition and "
        "preprocessing instructions, immutable release identifiers, data checksums, and "
        "benchmark manifests."
    )
    doc.save(MAIN_OUT)


def revise_supplement():
    copy2(SUPP_SRC, SUPP_OUT)
    doc = Document(SUPP_OUT)
    anchor = None
    for p in doc.paragraphs:
        if p.text.startswith("S10. Reporting conventions"):
            anchor = p
            break
    if anchor is None:
        raise RuntimeError("S10 anchor not found")
    heading = paragraph_after(anchor, "S11. Preliminary review-cycle validation", style="Heading 1")
    paragraph_after(
        heading,
        "Controlled SIMPLS equivalence. The retained non-optimized SIMPLS route, "
        "the state-reuse route, and pls::simpls.fit were compared on fixed "
        "well-conditioned, ill-conditioned, and rank-deficient synthetic multivariate "
        "regression tasks (ntrain=180, p=60, q=4, five components). For the state-reuse "
        "route relative to pls::simpls.fit, prediction correlation was at least "
        "0.9999999, relative prediction error was at most 3.86×10−4, coefficient "
        "relative error was at most 5.49×10−4, and the maximum principal angle between "
        "predictive subspaces was 0.050°. This establishes controlled numerical agreement, "
        "but does not replace the planned real-data and memory ablations."
    )
    paragraph_after(
        heading,
        "Float32 NMR control. On the fixed NMR partition, a ten-component SIMPLS "
        "screening control with q=5000 responses gave RMSD 6.4850×10−5 for float32 "
        "CPU, 6.4850×10−5 for float32 CUDA, and 6.4847×10−5 for float64 CPU. The "
        "current float32 CPU and CUDA paths were slower (60.8 and 51.8 s) than float64 "
        "CPU (7.54 s). The full q=28355 float32 path did not satisfy the predeclared "
        "runtime screen. These results are retained as a limitation and motivate further "
        "alignment of the float32 and optimized double-precision SIMPLS execution paths."
    )
    doc.save(SUPP_OUT)


def create_review():
    doc = Document()
    doc.add_heading("Independent reviewer report: cycle 1 update", level=0)
    doc.add_paragraph("Manuscript: fastPLS: scalable partial least squares with compiled CPU and accelerator backends for high-dimensional biomedical data")
    doc.add_heading("Recommendation", level=1)
    doc.add_paragraph("Major revision")
    doc.add_heading("Assessment after cycle 1", level=1)
    doc.add_paragraph(
        "The revised working manuscript now reports a controlled SIMPLS equivalence "
        "experiment and corrects the float32 interpretation. The numerical agreement "
        "study is encouraging: the optimized implementation closely matched "
        "pls::simpls.fit on three synthetic matrix conditions. The authors also avoid "
        "overstating float32, since the current high-response screening result preserves "
        "RMSD but is slower than float64. These changes improve scientific transparency."
    )
    doc.add_heading("Major comments", level=1)
    for text in [
        "The central benchmark evidence is still incomplete. The manuscript requires final tables and plots with raw repetitions, dispersion, peak host RSS, peak GPU memory, failures, requested versus executed estimator, and exact dataset provenance.",
        "The SIMPLS equivalence experiment is a useful start but remains limited to small synthetic data. Add a real-data comparison, report several component counts, include an explicit ablation of each stated reuse strategy, and report memory as well as time.",
        "The float32 result is currently a limitation rather than a contribution to performance. Either align the float32 core with the optimized double route and repeat the full precision study, or narrow the package and manuscript claims to verified numerical storage compatibility.",
        "The automatic large-class SIMPLS-to-label-aware-PLS-SVD substitution must be retained in all final records. Any resulting ImageNet row must be analysed as PLS-SVD, not sequential SIMPLS.",
        "NMR requires the planned detailed validation: preprocessing confirmation including the excluded water region, training-only component selection, per-spectrum and per-response error distributions, and matched CPU/CUDA float64 results at the selected component count.",
        "ImageNet should remain an explicitly non-biomedical post-feature-extraction stress test. The main biomedical evidence should be based on NMR, metabolomics, cancer-omics, and the selected single-cell datasets."
    ]:
        doc.add_paragraph(text, style="List Number")
    doc.add_heading("Minor comments", level=1)
    for text in [
        "Use PLS-SVD consistently and define all performance metrics where first reported.",
        "Replace future-tense data and code availability statements with immutable release identifiers before submission.",
        "Clarify the relationship between the GPL-3 R package, bundled code, and the MIT-licensed kodama-cpp project without implying unsupported accelerator features.",
        "Ensure that all backend tables distinguish native, hybrid, and host-side operations.",
        "State that the current high-response float32 screen was stopped by a predeclared runtime criterion rather than reporting it as a successful benchmark."
    ]:
        doc.add_paragraph(text, style="List Bullet")
    doc.add_heading("Required next evidence", level=1)
    doc.add_paragraph(
        "The next cycle should run the pinned full benchmark suite, the real-data SIMPLS "
        "equivalence/ablation experiment, the NMR-focused error and component-selection "
        "analysis, and a precision-matched float64/float32 study on configurations that "
        "are genuinely supported. Only then should the main conclusions and abstract be "
        "finalized."
    )
    doc.save(REVIEW_OUT)


revise_main()
revise_supplement()
create_review()
print(MAIN_OUT)
print(SUPP_OUT)
print(REVIEW_OUT)
