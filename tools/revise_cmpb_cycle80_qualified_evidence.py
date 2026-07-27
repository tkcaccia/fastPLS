#!/usr/bin/env python3
"""Replace stale NMR/ImageNet evidence with qualified current-package runs."""

import csv
import math
import re
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260727_cycle79"
OUTPUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260727_cycle80"
MAIN_SOURCE = SOURCE_DIR / "fastPLS_CMPB_main_cycle79_0.99.6_20260727.docx"
SUPP_SOURCE = (
    SOURCE_DIR / "fastPLS_CMPB_supplement_cycle79_0.99.6_20260727.docx"
)
MAIN_OUTPUT = OUTPUT_DIR / "fastPLS_CMPB_main_cycle80_0.99.6_20260727.docx"
SUPP_OUTPUT = (
    OUTPUT_DIR / "fastPLS_CMPB_supplement_cycle80_0.99.6_20260727.docx"
)
RESULT_ROOT = (
    ROOT / "benchmark_results" / "manuscript_revision_cycle80_20260727"
)
NMR_ROOT = RESULT_ROOT / "nmr_qualified"
IMAGENET_ROOT = RESULT_ROOT / "imagenet_float32_simpls_lda_path"
NMR_FIGURE = NMR_ROOT / "nmr_qualified_main_figure.png"
IMAGENET_FIGURE = (
    IMAGENET_ROOT / "imagenet_float32_simpls_lda_main_figure.png"
)
ARCHIVE_SHA = "c868770fee196af24bfb15b4da9fad1d7765deaf68c0ef1fd3e5a15670ecaa85"


def read_csv(path):
    with path.open(newline="") as handle:
        return list(csv.DictReader(handle))


def number(row, key):
    value = row.get(key, "")
    return float(value) if value not in {"", None} else math.nan


def replace_paragraph_prefix(document, prefix, replacement):
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            paragraph.text = replacement
            return
    raise RuntimeError(f"Paragraph not found: {prefix}")


def set_cell_text(cell, value, size=6.8, bold=False):
    cell.text = str(value)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.bold = bold


def replace_table_rows(table, rows, size=6.6):
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            set_cell_text(cell, value, size=size)


def find_table(document, headers):
    for table in document.tables:
        current = [cell.text for cell in table.rows[0].cells]
        if current == headers:
            return table
    raise RuntimeError(f"Table not found: {headers}")


def replace_media(source, destination, replacements):
    temporary = destination.with_suffix(".media.docx")
    with ZipFile(source, "r") as zin, ZipFile(
        temporary, "w", compression=ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = replacements.get(item.filename, zin.read(item.filename))
            zout.writestr(item, data)
    temporary.replace(destination)


def resize_embedded_image(document, target_name, width_inches, height_inches):
    for shape in document.inline_shapes:
        blip = shape._inline.graphic.graphicData.pic.blipFill.blip
        relation = document.part.rels.get(blip.embed)
        if relation and relation.target_ref.endswith(target_name):
            shape.width = Inches(width_inches)
            shape.height = Inches(height_inches)
            return
    raise RuntimeError(f"Embedded image not found: {target_name}")


def wilson(correct, total, z=1.959963984540054):
    p = correct / total
    denom = 1 + z * z / total
    centre = (p + z * z / (2 * total)) / denom
    spread = z * math.sqrt(
        p * (1 - p) / total + z * z / (4 * total * total)
    ) / denom
    return centre - spread, centre + spread


def parse_peak_rss_mb(path):
    text = path.read_text(errors="replace")
    match = re.search(
        r"Maximum resident set size \(kbytes\):\s*([0-9]+)",
        text,
    )
    return float(match.group(1)) / 1024 if match else math.nan


def load_evidence():
    nmr = read_csv(NMR_ROOT / "nmr_qualified_summary.csv")
    nmr_by_label = {row["label"]: row for row in nmr}
    agreement = read_csv(NMR_ROOT / "nmr_qualified_agreement.csv")
    agreement_by_candidate = {row["candidate"]: row for row in agreement}
    historical = read_csv(
        NMR_ROOT / "nmr_family_selected_and_historical_metrics.csv"
    )
    historical_row = next(
        row for row in historical
        if row["workflow"].startswith("Deposited") and row["metric"] == "RMSD"
    )
    gpu_rows = read_csv(NMR_ROOT / "nmr_cuda_gpu_memory_summary.csv")
    gpu_by_family = {row["family"]: row for row in gpu_rows}

    image = read_csv(
        IMAGENET_ROOT / "imagenet_float32_simpls_lda_path.csv"
    )
    image.sort(key=lambda row: int(row["ncomp_effective"]))
    expected_grid = list(range(100, 1001, 100))
    observed_grid = [int(row["ncomp_effective"]) for row in image]
    if observed_grid != expected_grid or any(
        row["status"] != "success" for row in image
    ):
        raise RuntimeError(
            f"Incomplete qualified ImageNet path: {observed_grid}"
        )
    image_peak_host = parse_peak_rss_mb(
        IMAGENET_ROOT / "imagenet_float32_simpls_lda_path.time"
    )
    gpu_trace = read_csv(
        IMAGENET_ROOT / "imagenet_float32_simpls_lda_path_gpu_trace.csv"
    )
    gpu_values = [float(row["memory_used_mb"]) for row in gpu_trace]
    image_gpu_baseline = gpu_values[0]
    image_gpu_peak = max(gpu_values)
    return {
        "nmr": nmr_by_label,
        "agreement": agreement_by_candidate,
        "historical": historical_row,
        "nmr_gpu": gpu_by_family,
        "image": image,
        "image_peak_host": image_peak_host,
        "image_gpu_baseline": image_gpu_baseline,
        "image_gpu_peak": image_gpu_peak,
    }


def revise_main(evidence):
    nmr = evidence["nmr"]
    agree = evidence["agreement"]
    image = evidence["image"]
    image_first = image[0]
    image_last = image[-1]
    correct = round(
        number(image_last, "top1_accuracy") *
        int(image_last["test_n"])
    )
    lower, upper = wilson(correct, int(image_last["test_n"]))

    document = Document(MAIN_SOURCE)
    replace_paragraph_prefix(
        document,
        "Results: Deterministic fastPLS SIMPLS met",
        (
            "Results: Deterministic fastPLS SIMPLS met the prespecified "
            "numerical tolerances in all 117 component-level comparisons. In "
            "matched single-CPU comparisons, it was faster than "
            "pls::simpls.fit on seven of nine datasets, with identical argmax "
            "accuracy and speed-up up to 8.90-fold. On NMR, qualified CPU rSVD "
            "(oversampling 20, two power iterations, seed 123) reduced "
            "50-component SIMPLS fitting-plus-prediction time from 350.7 to "
            "9.8 s while its predictions differed from deterministic IRLBA by "
            "relative Frobenius error 6.29e-11. A current-package hybrid "
            "float32 SIMPLS/LDA route processed one million ImageNet/DINOv2 "
            "training embeddings; at the requested 1,000-component boundary, "
            f"top-1/top-5 accuracy was "
            f"{number(image_last, 'top1_accuracy'):.4f}/"
            f"{number(image_last, 'top5_accuracy'):.4f}. This noncanonical "
            "single-run analysis is exploratory."
        ),
    )
    replace_paragraph_prefix(
        document,
        "NMR comprised 1,200 training and 321 held-out spectra,",
        (
            "NMR comprised 1,200 training and 321 held-out spectra, with "
            "13,000 predictors and 28,355 responses. Predictor columns between "
            "4.6 and 4.8 ppm were zeroed in training and test data as standard "
            "water-region preprocessing; responses were unmasked. Training-only "
            "one-standard-error selection retained five PLS-SVD and 50 SIMPLS "
            "components. These are family-specific predictive settings rather "
            "than a matched implementation comparison."
        ),
    )
    replace_paragraph_prefix(
        document,
        "Compact latent prediction processes large test matrices",
        (
            "Compact latent prediction processes large test matrices in row "
            "blocks and releases temporary scores after each block. For the "
            "ImageNet top-5 analysis, predictions were similarly blocked and "
            "classification counts were accumulated online, avoiding a full "
            "281,167-by-1,000 double score matrix."
        ),
    )
    replace_paragraph_prefix(
        document,
        "In this exploratory workflow, selected CUDA PLS-SVD",
        (
            "At the family-selected settings, CUDA PLS-SVD and SIMPLS achieved "
            "RMSD 0.001043 (Q² 0.98916) and 0.000759 (Q² 0.99425), "
            "respectively. SIMPLS had lower median per-spectrum and "
            "response-wise error. The representative spectrum in Figure 4 was "
            "chosen reproducibly as the held-out sample nearest the median "
            "SIMPLS RMSD, not selected for visual fit."
        ),
    )
    replace_paragraph_prefix(
        document,
        "In the matched backend benchmark, family, split, float64 precision,",
        (
            "A separate matched solver/backend analysis held family, split, "
            "float64 precision, and component count fixed. For PLS-SVD at five "
            "components, CPU IRLBA, CPU rSVD, and CUDA rSVD required 264.3, "
            "4.74, and 0.41 s; rSVD predictions had relative error at most "
            "3.75e-11 against IRLBA. For SIMPLS at 50 components, the times "
            "were 350.7, 9.80, and 1.51 s. CPU rSVD relative prediction error "
            "was 6.29e-11; CUDA rSVD error was 0.00566 with correlation "
            "0.999981. All rSVD rows used oversampling 20, two power iterations, "
            "and seed 123 and met the prespecified approximate-route tolerances."
        ),
    )
    replace_paragraph_prefix(
        document,
        "The deposited 165-component historical workflow required",
        (
            "The deposited 165-component PLS-SVD/IRLBA workflow required "
            "447.6 s and achieved RMSD 0.000710. It is shown as scientific "
            "context for the earlier Nature Communications analysis, not as an "
            "implementation-only contrast, because component count, workflow, "
            "and hardware differ."
        ),
    )
    replace_paragraph_prefix(
        document,
        "Figure 4. Exploratory NMR rSVD workflow analyses:",
        (
            "Figure 4. NMR predictive and computational analyses. Panels A-C "
            "separate family-selected held-out performance from the deposited "
            "165-component historical context. Panels D-E overlay observed and "
            "predicted intensities for the prespecified median-RMSD held-out "
            "spectrum over the full response range and 1.7-0.5 ppm expansion. "
            "Panel F reports matched float64 solver/backend resources at fixed "
            "family-specific component counts. rSVD used oversampling 20, two "
            "power iterations, and seed 123."
        ),
    )
    replace_paragraph_prefix(
        document,
        "Exploratory ImageNet experiment 1 used CUDA SIMPLS rSVD",
        (
            "Exploratory ImageNet experiment 1 used the current package on "
            "1,000,000 training and 281,167 held-out 1,024-dimensional "
            "embeddings. The actual route was label-aware float32 SIMPLS with "
            "GPU rSVD and native CUDA LDA, but sequential SIMPLS deflation and "
            "score projection remained host-resident; it is therefore hybrid, "
            "not fully GPU-resident. rSVD used oversampling 20, two power "
            "iterations, and seed 123."
        ),
    )
    replace_paragraph_prefix(
        document,
        "Exploratory FAISS retrieval used the same split.",
        (
            f"A single shared component-path fit evaluated requested prefixes "
            f"from 100 to 1,000 components. The 1,000-component row was the "
            f"prespecified boundary stress point, not a selected optimum. "
            f"Top-1/top-5 accuracy changed from "
            f"{number(image_first, 'top1_accuracy'):.4f}/"
            f"{number(image_first, 'top5_accuracy'):.4f} at 100 components to "
            f"{number(image_last, 'top1_accuracy'):.4f}/"
            f"{number(image_last, 'top5_accuracy'):.4f} at "
            f"{int(image_last['ncomp_effective']):,} components; the latter "
            f"corresponded to {correct:,}/{int(image_last['test_n']):,} correct "
            f"top-1 predictions (Wilson 95% CI {lower:.4f}-{upper:.4f}). The "
            f"shared fit required {number(image_last, 'shared_path_fit_time_sec'):.1f} "
            f"s. Peak process RSS was {evidence['image_peak_host']:.0f} MB and "
            f"peak device memory was {evidence['image_gpu_peak']:.0f} MB "
            f"({evidence['image_gpu_peak'] - evidence['image_gpu_baseline']:.0f} "
            f"MB above baseline). Because the split was noncanonical and the "
            f"measurements were single runs, these are exploratory feasibility "
            f"estimates, not confirmatory accuracy claims."
        ),
    )
    replace_paragraph_prefix(
        document,
        "Figure 5. Exploratory ImageNet SIMPLS classification",
        (
            "Figure 5. Exploratory ImageNet float32 SIMPLS-LDA component path. "
            "Accuracy and prefix-specific held-out prediction time are shown for "
            "requested 100-1,000-component prefixes from one shared fit; 1,000 "
            "is a boundary stress point, not a selected optimum. Resource values describe "
            "the complete shared path. The route is hybrid: SIMPLS deflation and "
            "score projection are host-resident, while rSVD range products and "
            "LDA use CUDA. Controls were oversampling 20, two power iterations, "
            "and seed 123."
        ),
    )
    replace_paragraph_prefix(
        document,
        "rSVD, implicit products, float32, CUDA, and Metal are optional",
        (
            "rSVD, implicit products, float32, CUDA, and Metal are optional "
            "implementation mechanisms around accelerated SIMPLS. Qualified "
            "NMR controls met the prespecified approximate-route tolerances, "
            "but rSVD remains stochastic and CPU IRLBA remains the deterministic "
            "reference. The million-sample ImageNet route demonstrated "
            "feasibility with current package code, while its hybrid residency, "
            "single run, noncanonical split, and lack of an estimator-matched "
            "large-scale control preclude a general accelerator or accuracy "
            "claim."
        ),
    )
    replace_paragraph_prefix(
        document,
        "Code and benchmark outputs are available at",
        (
            "Code and benchmark outputs are available at "
            "https://github.com/tkcaccia/fastPLS; reusable components are at "
            "https://github.com/tkcaccia/kodama-cpp. The reviewed software "
            "snapshot is fastPLS 0.99.6. Commit "
            "6e50bd318f20289101f6b723953830aefa8b95d6 identifies the base "
            f"source and source-archive SHA-256 {ARCHIVE_SHA} identifies the "
            "exact current experimental build. Analysis-specific scripts and "
            "archive digests are reported in Supplementary Table S15."
        ),
    )
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(MAIN_OUTPUT)
    replace_media(
        MAIN_OUTPUT,
        MAIN_OUTPUT,
        {
            "word/media/image25.png": NMR_FIGURE.read_bytes(),
            "word/media/image19.png": IMAGENET_FIGURE.read_bytes(),
        },
    )
    document = Document(MAIN_OUTPUT)
    resize_embedded_image(document, "image25.png", 6.00, 7.65)
    resize_embedded_image(document, "image19.png", 6.35, 2.54)
    document.save(MAIN_OUTPUT)


def revise_supplement(evidence):
    nmr = evidence["nmr"]
    agree = evidence["agreement"]
    historical = evidence["historical"]
    image = evidence["image"]
    nmr_gpu = evidence["nmr_gpu"]
    document = Document(SUPP_SOURCE)

    replace_paragraph_prefix(
        document,
        "Current benchmark workflows record repository state,",
        (
            "Current benchmark workflows record repository state, benchmark-"
            "script checksum, package version, source-archive SHA-256, compiler, "
            "BLAS/LAPACK, thread settings, accelerator libraries, seeds, rSVD "
            "controls, and data/split identifiers. Table S15 maps each analysis "
            "to its exact evidence archive. The current NMR and ImageNet runs "
            f"used source archive SHA-256 {ARCHIVE_SHA}."
        ),
    )
    replace_paragraph_prefix(
        document,
        "Table S12. Definitive NMR evidence.",
        (
            "Table S12. Qualified NMR evidence. Family-selected predictive "
            "comparisons, matched solver/backend comparisons, and the deposited "
            "historical workflow answer different questions and are separated. "
            "Host memory is baseline-corrected process RSS; CUDA memory is total "
            "device use above the pre-run baseline and includes context."
        ),
    )
    replace_paragraph_prefix(
        document,
        "The pooled archive contained 1,281,167 precomputed DINOv2 embeddings",
        (
            "The pooled archive contained 1,281,167 precomputed DINOv2 "
            "embeddings with 1,024 features and 1,000 classes. Seed 123 assigned "
            "1,000,000 rows to training and 281,167 to a complementary holdout; "
            "this was not the canonical ImageNet split. The current-package "
            "classification experiment used label-aware float32 SIMPLS, rSVD "
            "oversampling 20, two power iterations, seed 123, and CUDA LDA. "
            "SIMPLS deflation and score projection were host-resident, so the "
            "route is hybrid. A full-score top-5 attempt was killed at 27.1 GB "
            "process RSS; the reported run used 10,000-row prediction blocks "
            "and online metric accumulation without changing fitted scores or "
            "class decisions. The separate FAISS experiment used a different "
            "estimator and objective and remains a compression/retrieval study."
        ),
    )
    replace_paragraph_prefix(
        document,
        "Table S13. Definitive ImageNet classification",
        (
            "Table S13. Current exploratory ImageNet classification and "
            "representation results. Shared-path fit time is repeated only to "
            "identify the common fit and must not be interpreted as ten "
            "independent component-specific fits. The 1,000-component row is "
            "the requested boundary stress point, not a selected optimum. All "
            "classification rows are single-run feasibility estimates."
        ),
    )

    capability = find_table(
        document,
        [
            "Family", "Kernel", "Backend", "Endpoint status", "Solver",
            "Residency", "Windows", "Extreme q",
        ],
    )
    for row in capability.rows[1:]:
        values = [cell.text for cell in row.cells]
        if values[0] == "SIMPLS" and values[2] == "CUDA":
            set_cell_text(
                row.cells[5],
                "Hybrid in float32 label-aware classification: host sequential "
                "SIMPLS and score projection; CUDA rSVD range products and LDA",
                size=6.2,
            )
            set_cell_text(
                row.cells[3],
                "regression=validated; argmax=experimental; LDA=experimental",
                size=6.2,
            )

    nmr_table = find_table(
        document,
        [
            "Analysis", "Family", "Implementation", "A",
            "Predictive metric", "Time s", "Delta host MB", "Delta GPU MB",
            "Error detail / scope",
        ],
    )
    nmr_rows = []
    for label in [
        "PLS-SVD CPU IRLBA",
        "PLS-SVD CPU rSVD",
        "PLS-SVD CUDA rSVD",
        "SIMPLS CPU IRLBA",
        "SIMPLS CPU rSVD",
        "SIMPLS CUDA rSVD",
    ]:
        row = nmr[label]
        family = "PLS-SVD" if row["family"] == "plssvd" else "SIMPLS"
        implementation = f"{row['backend'].upper()} {row['solver']}"
        gpu_delta = "0.0"
        if row["backend"] == "cuda":
            gpu_delta = evidence["nmr_gpu"][row["family"]][
                "incremental_gpu_mb"
            ]
        detail = "deterministic reference"
        if row["solver"] == "rsvd":
            a = agree[label]
            detail = (
                f"vs IRLBA: rel. pred. error "
                f"{float(a['relative_frobenius_error']):.3g}; "
                f"corr {float(a['prediction_correlation']):.6f}; "
                "os=20, power=2, seed=123"
            )
        nmr_rows.append([
            "Matched solver/backend",
            family,
            implementation,
            row["ncomp"],
            f"RMSD {float(row['RMSD']):.6g}; Q² {float(row['Q2']):.6f}",
            f"{float(row['total_time_sec_median']):.3f}",
            f"{float(row['incremental_process_peak_rss_mb']):.1f}",
            gpu_delta,
            detail,
        ])
    nmr_rows.append([
        "Historical scientific context",
        "PLS-SVD",
        "deposited R/IRLBA",
        "165",
        f"RMSD {float(historical['value']):.6g}",
        "447.601",
        "3605.5",
        "n/a",
        "Different component count/workflow/hardware; not backend-only",
    ])
    replace_table_rows(nmr_table, nmr_rows, size=6.15)

    image_table = find_table(
        document,
        [
            "Experiment", "Head / representation", "A/dim.", "Top-1",
            "Top-5", "End-to-end s", "Host RSS MB", "GPU MB",
            "Qualification",
        ],
    )
    image_rows = []
    host_peak = evidence["image_peak_host"]
    gpu_peak = evidence["image_gpu_peak"]
    for row in image:
        image_rows.append([
            "Shared SIMPLS path",
            "LDA",
            row["ncomp_effective"],
            f"{float(row['top1_accuracy']):.4f}",
            f"{float(row['top5_accuracy']):.4f}",
            (
                f"shared fit {float(row['shared_path_fit_time_sec']):.1f} + "
                f"prediction {float(row['prediction_time_sec']):.1f}"
            ),
            f"{host_peak:.0f}",
            f"{gpu_peak:.0f}",
            "single exploratory run; hybrid; os=20, power=2, seed=123",
        ])
    old_rows = [
        [cell.text for cell in row.cells]
        for row in image_table.rows[1:]
        if row.cells[0].text == "FAISS exact retrieval"
    ]
    image_rows.extend(old_rows)
    replace_table_rows(image_table, image_rows, size=6.1)

    provenance = find_table(
        document,
        [
            "ID", "Authoritative output", "Result archive", "Version",
            "Source status", "Generating script", "SHA-256 prefix",
        ],
    )
    existing = [
        [cell.text for cell in row.cells]
        for row in provenance.rows[1:]
        if row.cells[0].text not in {"A17", "A18"}
    ]
    existing.extend([
        [
            "A17",
            "Figure 4 and Table S12",
            "benchmark_results/manuscript_revision_cycle80_20260727/nmr_qualified",
            "0.99.6",
            "source archive recorded",
            "benchmark/benchmark_nmr_qualified_solver.R",
            ARCHIVE_SHA[:12],
        ],
        [
            "A18",
            "Figure 5 and Table S13 classification path",
            "benchmark_results/manuscript_revision_cycle80_20260727/"
            "imagenet_float32_simpls_lda_path",
            "0.99.6",
            "source archive recorded",
            "benchmark/benchmark_imagenet_float32_simpls_lda_path.R",
            ARCHIVE_SHA[:12],
        ],
    ])
    replace_table_rows(provenance, existing, size=6.0)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    document.save(SUPP_OUTPUT)


def main():
    required = [
        NMR_FIGURE,
        NMR_ROOT / "nmr_cuda_gpu_memory_summary.csv",
        IMAGENET_FIGURE,
        IMAGENET_ROOT / "imagenet_float32_simpls_lda_path.csv",
    ]
    missing = [str(path) for path in required if not path.exists()]
    if missing:
        raise SystemExit("Missing evidence:\n" + "\n".join(missing))
    evidence = load_evidence()
    revise_main(evidence)
    revise_supplement(evidence)
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)


if __name__ == "__main__":
    main()
