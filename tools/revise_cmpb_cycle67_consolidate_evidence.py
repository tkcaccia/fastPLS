#!/usr/bin/env python3
"""Create a compact, authoritative CMPB supplement and de-duplicate the main text."""

from pathlib import Path
import shutil

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle66"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle68"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle66_0.99.6_20260726.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle66_0.99.6_20260726.docx"
MAIN_OUTPUT = OUTPUT / "fastPLS_CMPB_main_cycle68_0.99.6_20260726.docx"
SUPP_OUTPUT = OUTPUT / "fastPLS_CMPB_supplement_cycle68_0.99.6_20260726.docx"

SIMPLS = (
    ROOT
    / "benchmark_results"
    / "simpls_estimator_preservation_reliable_power2_final_20260725"
)
OPLS_KERNEL = ROOT / "benchmark_results" / "opls_kernel_setting_reliability_20260726"
CUDA_RSVD = ROOT / "benchmark_results" / "rsvd_cuda_reliability_20260725.csv"
CAPABILITY = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle13_20260725"
    / "float32_capability_matrix.csv"
)
EXTERNAL = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle57_20260726"
    / "external_single_cpu_accuracy_time_memory.csv"
)
SELECTED = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle20_20260725"
    / "paired_backend_selected_summary.csv"
)
MEMORY = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle21_20260725"
    / "selected_memory_baseline_summary.csv"
)
NMR = ROOT / "benchmark_results" / "manuscript_revision_cycle64_20260726"
IMAGENET = ROOT / "benchmark_results" / "manuscript_revision_cycle54_20260726"
OUTER = ROOT / "benchmark_results" / "manuscript_revision_cycle66_20260726"
PROVENANCE = (
    ROOT
    / "benchmark_results"
    / "manuscript_revision_cycle63_20260726"
    / "analysis_commit_provenance.csv"
)


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_paragraph(document, prefix, text):
    paragraph = find_paragraph(document, prefix)
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def replace_text_everywhere(document, old, new):
    """Replace text in paragraphs and table cells without rebuilding document parts."""
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace(old, new)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace(old, new)


def normalize_submission_terminology(document):
    """Apply the manuscript's canonical numerical and backend terminology."""
    replacements = (
        ("PLSSVD", "PLS-SVD"),
        ("PLS SVD", "PLS-SVD"),
        ("RSVD", "rSVD"),
        ("R-SVD", "rSVD"),
        ("Float32", "float32"),
        ("Float64", "float64"),
        ("Cuda", "CUDA"),
        ("METAL", "Metal"),
        ("prediction-head", "prediction head"),
        ("PLS-SVD/rSVD", "PLS-SVD using rSVD"),
        ("SIMPLS/rSVD", "SIMPLS using rSVD"),
        ("nested validation", "nested cross-validation"),
    )
    for old, new in replacements:
        replace_text_everywhere(document, old, new)


def insert_after(paragraph, text="", style=None):
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    if style is not None:
        inserted.style = style
    if text:
        inserted.add_run(text)
    return inserted


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def synchronize_algorithm_notation(document):
    """Reserve K for CV folds and C for requested component-count snapshots."""
    notation = find_paragraph(document, "Notation is fixed throughout:")
    notation_text = notation.text
    notation_text = notation_text.replace(
        "A denotes the retained component count, and K denotes the number of "
        "cross-validation folds.",
        "A denotes the maximum retained component count, C denotes the set of "
        "requested component counts, and K denotes the number of cross-validation "
        "folds.",
    )
    notation.clear()
    notation.add_run(notation_text)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace(
                            "requested prefixes K",
                            "requested component-count set C",
                        ).replace(
                            "If a is in K",
                            "If a is in C",
                        )


def truncate_from_heading(document, prefix):
    anchor = find_paragraph(document, prefix)
    body = document._element.body
    remove = False
    for child in list(body):
        if child is anchor._p:
            remove = True
        if remove and child.tag != qn("w:sectPr"):
            body.remove(child)


def format_table(table, font_size=6.0, first_col_left=True):
    table.style = "Table"
    table.autofit = True
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    for row_index, row in enumerate(table.rows):
        row_properties = row._tr.get_or_add_trPr()
        row_properties.append(OxmlElement("w:cantSplit"))
        for column_index, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if first_col_left and column_index == 0
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.bold = row_index == 0


def add_table(document, headers, rows, font_size=6.0):
    table = document.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = str(header)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    format_table(table, font_size)
    return table


def add_caption(document, text):
    paragraph = document.add_paragraph(style="Caption")
    paragraph.add_run(text)
    return paragraph


def fmt(value, digits=3):
    if value is None or pd.isna(value):
        return "NR"
    value = float(value)
    if value != 0 and (abs(value) < 10 ** (-digits) or abs(value) >= 10000):
        return f"{value:.2e}"
    return f"{value:.{digits}f}"


def metric_value(row):
    name = str(row.metric_name)
    if pd.isna(row.metric_median):
        return str(row.status)
    value = float(row.metric_median)
    if name == "accuracy":
        return f"Acc {value:.4f}"
    if name == "rmsd":
        return f"RMSD {value:.6g}"
    return f"{name} {value:.4g}"


def compact_route(value):
    value = str(value)
    replacements = {
        "compiled CPU": "CPU",
        "device-accelerated": "device",
        "Metal-accelerated": "Metal",
        "host-assisted": "hybrid",
        "classification: ": "",
        "not applicable": "n/a",
    }
    for old, new in replacements.items():
        value = value.replace(old, new)
    return value


def build_claim_map():
    return [
        (
            "Deterministic SIMPLS preserves de Jong SIMPLS within stated tolerances",
            "Table S7",
            "117 component-level endpoints and 12 fixed-fold selections",
        ),
        (
            "OPLS and kernel-PLS implementation reliability",
            "Table S7",
            "66 setting/task endpoints, 66 selections, 1,540 fold-component fits",
        ),
        (
            "rSVD is approximate and requires numerical qualification",
            "Table S8",
            "CPU and CUDA oversampling/power-iteration audits",
        ),
        (
            "Backend stages are native, hybrid, or host-resident",
            "Table S1",
            "Stage-level CPU, CUDA, and Metal residency",
        ),
        (
            "Float32 support is route conditional",
            "Table S9",
            "Family/backend/endpoint capability and warning policy",
        ),
        (
            "fastPLS versus independent R implementations",
            "Table S10",
            "Matched float64 one-CPU SIMPLS workflows",
        ),
        (
            "Selected-point predictive, time, and memory performance",
            "Table S11",
            "Paired CPU/CUDA dataset-family rows",
        ),
        (
            "NMR predictive selection and backend attribution are separate",
            "Table S12",
            "Family-selected errors, paired backends, historical workflow",
        ),
        (
            "ImageNet feasibility and supervised compression are exploratory",
            "Table S13",
            "Component path and matched FAISS representation study",
        ),
        (
            "Predictive dispersion includes training-sample variation",
            "Table S14",
            "Repeated outer partitions and selection frequencies",
        ),
        (
            "Each analysis maps to a reproducible archive",
            "Table S15",
            "Script, package metadata, split provenance, and archive digest",
        ),
    ]


def build_estimator_validation():
    sim = pd.read_csv(SIMPLS / "simpls_estimator_preservation_summary.csv")
    sim_cv = pd.read_csv(SIMPLS / "simpls_estimator_preservation_cv_selection.csv")
    op = pd.read_csv(OPLS_KERNEL / "opls_kernel_setting_reliability_summary.csv")
    op_cv = pd.read_csv(
        OPLS_KERNEL / "opls_kernel_setting_selection_setting_summary.csv"
    )
    sim_max = {
        "pred": sim.prediction_relative_error.max(),
        "coef": sim.coefficient_relative_error.max(),
        "angle": sim[
            [
                "score_subspace_max_angle_degrees",
                "projection_subspace_max_angle_degrees",
                "loading_subspace_max_angle_degrees",
            ]
        ]
        .max()
        .max(),
        "metric": sim.metric_absolute_difference.max(),
    }
    op_max = {
        "pred": op.max_prediction_relative_error.max(),
        "coef": op.max_coefficient_relative_error.max(),
        "angle": op[
            ["max_predictive_score_angle_deg", "max_orthogonal_score_angle_deg"]
        ]
        .max()
        .max(),
        "metric": op.max_metric_absolute_difference.max(),
    }
    return [
        (
            "SIMPLS / deterministic IRLBA",
            "pls::simpls.fit / de Jong",
            "117 / 117",
            f"{len(sim_cv)} / {len(sim_cv)}",
            fmt(sim_max["pred"], 3),
            fmt(sim_max["coef"], 3),
            fmt(sim_max["angle"], 4),
            fmt(sim_max["metric"], 3),
            "Preserved",
        ),
        (
            "OPLS (north 1-3) and kernel PLS (8 settings)",
            "Independent filter/kernel + pls::simpls.fit",
            f"{op.passes_all.sum()} / {op.runs.sum()}",
            f"{op_cv.selected_component_agreements.sum()} / {op_cv.comparisons.sum()}",
            fmt(op_max["pred"], 3),
            fmt(op_max["coef"], 3),
            fmt(op_max["angle"], 4),
            fmt(op_max["metric"], 3),
            "Reliable in deterministic float64 CPU scope",
        ),
    ]


def build_rsvd_reliability():
    cpu_p2 = pd.read_csv(SIMPLS / "simpls_estimator_approximation_rsvd.csv")
    cpu_p1 = pd.read_csv(
        ROOT
        / "benchmark_results"
        / "simpls_estimator_preservation_reliable_20260725_1735"
        / "simpls_estimator_approximation_rsvd.csv"
    )
    cuda = pd.read_csv(CUDA_RSVD)
    rows = []
    for label, frame, oversample, power in [
        ("CPU", cpu_p1, 10, 1),
        ("CPU", cpu_p2, 10, 2),
    ]:
        passed = int(frame.approximation_tolerance_pass.fillna(False).sum())
        rows.append(
            (
                label,
                oversample,
                power,
                f"{passed}/{len(frame)}",
                fmt(frame.prediction_relative_error.max(), 4),
                fmt(frame.prediction_correlation.min(), 5),
                fmt(frame.score_subspace_max_angle_degrees.max(), 2),
                fmt(frame.classification_label_agreement.min(), 3),
                fmt(frame.metric_absolute_difference.max(), 4),
                "Rejected" if passed < len(frame) else "Qualified approximate",
            )
        )
    for (oversample, power), frame in cuda.groupby(["oversample", "power"]):
        passed = int(frame.approximation_tolerance_pass.fillna(False).sum())
        rows.append(
            (
                "CUDA",
                int(oversample),
                int(power),
                f"{passed}/{len(frame)}",
                fmt(frame.prediction_relative_error.max(), 4),
                fmt(frame.prediction_correlation.min(), 5),
                "NR",
                fmt(frame.label_agreement.min(), 3),
                fmt(frame.metric_absolute_difference.max(), 4),
                "Qualified approximate" if passed == len(frame) else "Rejected",
            )
        )
    return rows


def build_float32_capability():
    data = pd.read_csv(CAPABILITY).fillna("n/a")
    rows = []
    keys = ["family", "kernel_scope", "backend"]
    for key, frame in data.groupby(keys, sort=False):
        family, scope, backend = key
        endpoints = "; ".join(
            f"{compact_route(row.endpoint)}={row.status}" for row in frame.itertuples()
        )
        residency = "/".join(sorted(set(compact_route(x) for x in frame.execution_residency)))
        solver = "/".join(sorted(set(str(x) for x in frame.supported_solver)))
        windows = "/".join(sorted(set(str(x) for x in frame.windows_status)))
        extreme = "/".join(sorted(set(str(x) for x in frame.extreme_response_status)))
        rows.append(
            (
                family,
                compact_route(scope),
                backend,
                endpoints,
                solver,
                residency,
                windows,
                extreme,
            )
        )
    return rows


def build_external_comparison():
    data = pd.read_csv(EXTERNAL)
    rows = []
    for dataset, frame in data.groupby("dataset", sort=False):
        fast = frame[frame.package.str.lower() == "fastpls"]
        arg = fast[fast.classifier == "argmax"].iloc[0]
        lda = fast[fast.classifier == "lda"]
        ext = frame[frame.package.str.lower() != "fastpls"].copy()
        fastest = ext.sort_values(["median_time_ms", "median_accuracy"], ascending=[True, False]).iloc[0]
        best = ext.sort_values(["median_accuracy", "median_time_ms"], ascending=[False, True]).iloc[0]
        lda_text = "NR"
        if len(lda):
            r = lda.iloc[0]
            lda_text = (
                f"{r.median_accuracy:.4f}; {r.median_time_ms/1000:.3f}s; "
                f"{r.median_peak_host_rss_mb:.0f}MB"
            )
        rows.append(
            (
                str(arg.dataset_label).replace("\n", " "),
                int(arg.ncomp_requested),
                f"{arg.median_accuracy:.4f}; {arg.median_time_ms/1000:.3f}s; {arg.median_peak_host_rss_mb:.0f}MB",
                lda_text,
                f"{fastest.package}/{fastest.algorithm}: {fastest.median_accuracy:.4f}; {fastest.median_time_ms/1000:.3f}s; {fastest.median_peak_host_rss_mb:.0f}MB",
                f"{best.package}/{best.algorithm}: {best.median_accuracy:.4f}; {best.median_time_ms/1000:.3f}s; {best.median_peak_host_rss_mb:.0f}MB",
            )
        )
    return rows


def build_selected_performance():
    selected = pd.read_csv(SELECTED)
    memory = pd.read_csv(MEMORY)
    rows = []
    order = {
        name: index
        for index, name in enumerate(
            [
                "metref",
                "ccle",
                "tcga_brca",
                "tcga_hnsc_methylation",
                "gtex_v8",
                "tcga_pan_cancer",
                "retina",
                "tabula",
                "cifar100",
                "cbmc_citeseq",
                "prism",
                "nmr",
            ]
        )
    }
    selected["_order"] = selected.dataset.map(order)
    for (dataset, family), frame in selected.sort_values(
        ["_order", "method_panel", "engine"]
    ).groupby(["dataset", "method_panel"], sort=False):
        cpu = frame[frame.engine == "CPU"]
        cuda = frame[frame.engine == "CUDA"]
        if not len(cpu) or not len(cuda):
            continue
        cpu, cuda = cpu.iloc[0], cuda.iloc[0]
        mem = memory[
            (memory.dataset == dataset)
            & (memory.method_panel == family)
        ]
        mcpu = mem[mem.engine == "CPU"]
        mcuda = mem[mem.engine == "CUDA"]
        host_cpu = (
            mcpu.iloc[0].incremental_host_rss_mb_median if len(mcpu) else np.nan
        )
        host_cuda = (
            mcuda.iloc[0].incremental_host_rss_mb_median if len(mcuda) else np.nan
        )
        gpu_cuda = (
            mcuda.iloc[0].incremental_gpu_mem_mb_median if len(mcuda) else np.nan
        )
        if pd.isna(cpu.metric_median) or pd.isna(cuda.metric_median):
            status = str(cpu.status)
        else:
            metric_diff = abs(float(cpu.metric_median) - float(cuda.metric_median))
            status = (
                "metric concordant" if metric_diff <= 0.005 else "metric discordant"
            )
        family_label = {
            "plssvd": "PLS-SVD",
            "simpls": "SIMPLS",
            "opls": "OPLS",
            "kernelpls": "kernel PLS",
        }[family]
        rows.append(
            (
                dataset.replace("_", " "),
                family_label,
                "n/a" if pd.isna(cpu.effective_ncomp) else int(cpu.effective_ncomp),
                str(cpu.selection_status).replace(" tested-grid", ""),
                f"{metric_value(cpu)} / {metric_value(cuda)}",
                (
                    f"{cpu.total_time_sec_median:.3f} / "
                    f"{cuda.total_time_sec_median:.3f}"
                    if pd.notna(cpu.total_time_sec_median)
                    and pd.notna(cuda.total_time_sec_median)
                    else "n/a"
                ),
                f"{fmt(host_cpu,1)} / {fmt(host_cuda,1)}",
                fmt(gpu_cuda, 1),
                status,
            )
        )
    return rows


def build_nmr():
    errors = pd.read_csv(NMR / "nmr_family_selected_error_summary.csv")
    paired = pd.read_csv(NMR / "nmr_paired_backend_only_summary.csv")
    historical = pd.read_csv(NMR / "nmr_historical_reference_165_summary.csv")
    rows = []
    for family in ["PLS-SVD", "SIMPLS"]:
        error = errors[errors.family == family].iloc[0]
        for engine in ["CPU", "CUDA"]:
            run = paired[(paired.family == family) & (paired.engine == engine)].iloc[0]
            rows.append(
                (
                    "Family-selected paired backend",
                    family,
                    engine,
                    int(run.effective_ncomp),
                    f"RMSD {error.global_RMSD:.6f}; Q2 {error.Q2_correlation_squared:.4f}",
                    f"{run.total_time_sec_median:.3f}",
                    fmt(run.incremental_host_rss_mb_median, 1),
                    fmt(run.incremental_gpu_mem_mb_median, 1),
                    f"sample median {error.sample_RMSD_median:.6f}; response median {error.response_RMSD_median:.6f}",
                )
            )
    ref = historical[historical.variant_name == "nature_fastsimpls_plssvd"].iloc[0]
    rows.append(
        (
            "Historical published workflow",
            "PLS-SVD",
            "deposited R",
            165,
            f"RMSD {ref.rmsd_median:.6f}",
            f"{ref.total_time_sec_median:.3f}",
            fmt(ref.incremental_peak_host_rss_mb_median, 1),
            "n/a",
            "Original centring-only protocol; contextual, not backend-only",
        )
    )
    return rows


def build_imagenet():
    path = pd.read_csv(IMAGENET / "imagenet_argmax_lda_component_path.csv")
    rows = []
    for ncomp in [100, 300, 500, 700, 1000]:
        for classifier in ["Argmax", "LDA"]:
            hit = path[(path.ncomp == ncomp) & (path.classifier == classifier)]
            if not len(hit):
                continue
            row = hit.iloc[0]
            rows.append(
                (
                    "Classification path",
                    classifier,
                    ncomp,
                    f"{row.accuracy:.4f}",
                    "reported in archive",
                    f"{row.total_fit_predict_sec:.1f}",
                    f"{row.peak_host_rss_mb:.0f}",
                    f"{row.peak_gpu_compute_apps_mb:.0f}",
                    "single exploratory run",
                )
            )
    retrieval = pd.read_csv(IMAGENET / "imagenet_retrieval_table.csv")
    for feature_space, ncomp in [
        ("raw_dinov2", np.nan),
        ("pca_scores", 200),
        ("pls_scores", 200),
    ]:
        hit = retrieval[retrieval.feature_space == feature_space]
        if pd.notna(ncomp):
            hit = hit[hit.ncomp == ncomp]
        row = hit.iloc[0]
        rows.append(
            (
                "FAISS exact retrieval",
                str(row.representation),
                "raw" if pd.isna(row.ncomp) else int(row.ncomp),
                f"{row.top1_accuracy:.4f}",
                f"{row.top5_accuracy:.4f}",
                f"{row.end_to_end_time_median_sec:.1f}",
                f"{row.peak_host_rss_mb:.0f}",
                f"{row.peak_gpu_mem_mb:.0f}",
                f"{row.compression_ratio:.2f}x compression; {int(row.n_repeats)} timing repeats",
            )
        )
    return rows


def build_outer_uncertainty():
    summary = pd.read_csv(OUTER / "repeated_outer_predictive_dispersion_summary.csv")
    rows = []
    for row in summary.itertuples():
        classifier = row.classifier
        rows.append(
            (
                row.dataset.replace("_", " "),
                str(row.method).replace("plssvd", "PLS-SVD").replace("kernelpls", "kernel PLS").upper()
                if row.method != "kernelpls"
                else "kernel PLS",
                classifier,
                row.n_outer_success,
                row.metric_name,
                f"{row.metric_mean:.6g} ({row.metric_sd:.3g})",
                f"{row.metric_q025:.6g}-{row.metric_q975:.6g}",
                f"{row.selected_ncomp_min}-{row.selected_ncomp_max}",
                f"{row.upper_boundary_frequency:.2f}",
                "rank constrained" if row.rank_constrained_grid else "",
            )
        )
    return rows


def build_provenance():
    data = pd.read_csv(PROVENANCE).fillna("NR")
    rows = []
    remap = {
        "A01": "Tables S11 and selected-point main figure",
        "A02": "Table S10",
        "A03": "Tables S7-S8",
        "A04": "Table S12",
        "A05": "Table S12",
        "A06": "Table S13",
        "A07": "Table S9",
        "A08": "Table S11",
        "A09": "Table S11",
        "A10": "Table S7",
        "A11": "Table S10",
        "A12": "Table S11",
        "A13": "Table S8",
        "A14": "Table S7",
        "A15": "Table S7",
    }
    for row in data.itertuples():
        rows.append(
            (
                row.analysis_id,
                remap.get(row.analysis_id, row.reported_output),
                row.result_archive,
                row.package_version,
                str(row.package_commit_status).replace("; ", ";"),
                row.benchmark_script,
                str(row.result_archive_sha256)[:12],
            )
        )
    rows.append(
        (
            "A16",
            "Table S14",
            "benchmark_results/manuscript_revision_cycle66_20260726",
            "0.99.6",
            "run metadata and failures retained",
            "benchmark/benchmark_repeated_outer_selection.R",
            "see archive README",
        )
    )
    return rows


def revise_main():
    document = Document(MAIN_SOURCE)
    synchronize_algorithm_notation(document)

    replace_paragraph(
        document,
        "OPLS first estimates orthogonal scores",
        (
            "OPLS first estimates orthogonal scores and loadings and then fits the "
            "predictive SIMPLS core [12]; the same orthogonal filter is applied at "
            "prediction. Kernel PLS constructs a centred linear, radial-basis, or "
            "polynomial kernel and uses the internal SIMPLS core [13,14]. The linear "
            "kernel dispatches to the ordinary linear route and avoids an n-by-n "
            "Gram matrix. Nonlinear kernels explicitly retain that matrix, requiring "
            "8n^2 bytes for one float64 Gram matrix before centring, copies, and "
            "solver workspaces (0.75 GiB at n=10,000; 2.98 GiB at n=20,000; "
            "6.71 GiB at n=30,000). The deterministic nonlinear-kernel validation "
            "reached n_train=180 and establishes numerical reliability only within "
            "that tested range. For memory planning on the evaluated hardware, we "
            "therefore recommend n<=20,000 on the 32-GB workstation and n<=10,000 "
            "on the 8-GB unified-memory Mac. These are conservative operational "
            "ceilings, not empirically validated performance limits; larger "
            "nonlinear-kernel jobs require explicit memory assessment. CUDA "
            "nonlinear kernel PLS remains host assisted, so host memory, rather than "
            "the nominal 16-GB device capacity alone, governs feasibility."
        ),
    )

    replace_paragraph(
        document,
        "The R package includes bundled IRLBA code",
        (
            "The R package includes bundled IRLBA code and is distributed under "
            "GPL-3. Reusable low-level CPU components are also maintained in the "
            "MIT-licensed kodama-cpp project. CUDA uses NVIDIA CUDA libraries and "
            "cuBLAS; Metal uses Apple's Metal Performance Shaders. On CPU, compiled "
            "C/C++ delegates dense linear algebra to the BLAS/LAPACK implementation "
            "linked to R, which may be OpenBLAS on suitable installations. This is "
            "an implementation capability, not evidence of multicore acceleration. "
            "All primary external-software comparisons used one effective BLAS "
            "thread. No controlled CPU thread-scaling experiment was conducted; "
            "therefore, no multicore speed-up is claimed. Exact package versions, "
            "compiler and accelerator flags, BLAS/LAPACK implementations, thread "
            "settings, and GPU libraries are reported in Supplementary Table S4a. "
            "The stage-level residency table identifies operations that remain on "
            "the host, so hybrid routes are not presented as fully device resident."
        ),
    )
    replace_paragraph(
        document,
        "The main biomedical benchmark included twelve tasks",
        (
            find_paragraph(
                document, "The main biomedical benchmark included twelve tasks"
            ).text
            .replace(
                "Exact dimensions, checksums, split construction, and limitations "
                "are reported in the Supplementary Material.",
                "Exact dimensions, checksums, split construction, and limitations "
                "are reported in Supplementary Tables S3 and S3a.",
            )
            .replace("Supplementary Table S41", "Supplementary Table S15")
        ),
    )
    replace_paragraph(
        document,
        "Standard R numeric matrices store each value",
        (
            "Standard R numeric matrices store each value in double precision "
            "(float64; eight bytes), whereas matrices supplied through the float "
            "package use four-byte float32 values. Float64 is the numerical "
            "reference. Float32 support is route conditional rather than "
            "package-wide: every PLS-family, backend, solver, and prediction route "
            "is classified as validated, experimental, hybrid, unavailable, or "
            "failed. The single authoritative capability and public-API policy "
            "matrix is Supplementary Table S9."
        ),
    )
    replace_paragraph(
        document,
        "Estimator preservation and randomized approximation",
        (
            "Estimator preservation and randomized approximation were evaluated "
            "separately in a prespecified study. Eight synthetic regimes covered "
            "regression and classification, p<n and p>n, low- and high-rank "
            "responses, ill-conditioned predictors, and exact rank deficiency; "
            "four real tasks comprised Breast, Colon, MetRef, and an NMR spectral "
            "subset. Deterministic IRLBA was assessed against pls::simpls.fit. An "
            "rSVD row was classified as failed when relative prediction error "
            "exceeded 0.05, prediction correlation was below 0.99, any "
            "score/projection/loading subspace angle exceeded 10 degrees, label "
            "agreement was below 0.99, or the predictive metric differed by more "
            "than 0.01. Numerical completion alone was never counted as "
            "approximation success. The deterministic estimator evidence is "
            "reported in Supplementary Table S7 and the separate approximate-solver "
            "audit in Supplementary Table S8."
        ),
    )
    replace_paragraph(
        document,
        "OPLS and kernel settings were prespecified explicitly",
        (
            "OPLS and kernel settings were prespecified explicitly. The principal "
            "all-dataset benchmark used one orthogonal component for OPLS and a "
            "linear kernel for kernel PLS; nonlinear RBF and polynomial settings "
            "were evaluated in the independent deterministic validation. The "
            "authoritative settings and reliability evidence are reported in "
            "Supplementary Table S7, with complete paths retained in the repository "
            "archive."
        ),
    )
    replace_paragraph(
        document,
        "Single precision was route-dependent",
        (
            "Single precision was route-dependent rather than uniformly "
            "beneficial. Float32 approximately halved stored input size on MetRef "
            "and PRISM, but runtime, incremental host memory, and predictive "
            "agreement varied by route. The authoritative capability matrix and "
            "paired precision/resource evidence are consolidated in Supplementary "
            "Table S9; unsupported or measured-risk paths remain visible and are "
            "not pooled into a general float32 speed claim."
        ),
    )
    replace_paragraph(
        document,
        "Formal deterministic reliability testing",
        (
            find_paragraph(
                document, "Formal deterministic reliability testing"
            ).text
            + " The definitive evidence is Supplementary Table S7."
        ),
    )
    replace_paragraph(
        document,
        "We first evaluated whether the computational implementation",
        (
            "We first evaluated whether the computational implementation of "
            "fastPLS improved the practical use of SIMPLS relative to independent "
            "R software under a controlled single-CPU setting. We then compared "
            "the compiled CPU route with numerically concordant CUDA and Metal "
            "routes and examined when randomized SVD (rSVD) was preferable to the "
            "deterministic IRLBA route. The final analyses focus on multivariate "
            "NMR prediction and million-sample ImageNet embeddings. The "
            "authoritative evidence map is Supplementary Table S6; estimator "
            "validation, numerical audits, route qualification, and provenance are "
            "reported in Supplementary Tables S7-S15."
        ),
    )
    replace_paragraph(
        document,
        "Repeated outer partitions quantified",
        (
            find_paragraph(document, "Repeated outer partitions quantified").text
            + " The definitive dispersion and selection summary is Supplementary "
            "Table S14."
        ),
    )
    replace_paragraph(
        document,
        "Peak host memory in Figure 2",
        (
            find_paragraph(document, "Peak host memory in Figure 2").text.replace(
                "reported in the Supplementary Material.",
                "reported in Supplementary Table S10.",
            )
        ),
    )
    replace_paragraph(
        document,
        "The primary software comparison used double-precision inputs",
        (
            "The primary software comparison used double-precision inputs, fixed "
            "outer splits, and one effective BLAS thread for fastPLS and all external "
            "packages (Figure 2). Across 126 attempted external-package "
            "dataset/method runs, 110 completed successfully and 16 did not: 12 "
            "documented package limitations, two timeout kills, and two execution "
            "errors. The figure displays the 65 successful external classification "
            "rows with compatible outputs together with 17 fastPLS rows. "
            "Deterministic fastPLS SIMPLS with argmax decoding produced the same "
            "test accuracy as pls::simpls.fit on all nine completed matched "
            "classification datasets and was faster on seven, including 4.23-fold "
            "on CIFAR-100, 8.65-fold on Retina, and 8.90-fold on Tabula Muris. "
            "Matched accuracies were 8,739/10,000 for CIFAR-100 (0.8739; Wilson 95% "
            "CI 0.8672-0.8803), 21,684/22,406 for Retina (0.9678; 0.9654-0.9700), "
            "and 40,077/50,059 for Tabula Muris (0.8006; 0.7971-0.8041). With the "
            "same latent-space LDA workflow, fastPLS and plsgenomics agreed on all "
            "eight completed matched datasets; fastPLS was faster on six, including "
            "6.44-fold on Retina and 6.77-fold on Tabula Muris. LDA classified "
            "21,714/22,406 Retina cells correctly (0.9691; 0.9668-0.9713) and "
            "43,813/50,059 Tabula Muris cells (0.8752; 0.8723-0.8781). Other "
            "independent implementations provide broader workflow context and are "
            "not interpreted as estimator-matched when their algorithms or "
            "prediction heads differ."
        ),
    )
    replace_paragraph(
        document,
        "The benefit of rSVD over IRLBA",
        (
            find_paragraph(document, "The benefit of rSVD over IRLBA").text.replace(
                "reported in the Supplementary Material.",
                "reported in Supplementary Table S8.",
            )
        ),
    )
    replace_paragraph(
        document,
        "Figure 3. Numerically qualified backend",
        (
            "Figure 3. Numerically qualified backend and solver comparisons in "
            "fastPLS. Speed-up is reported only for paired routes meeting the "
            "prespecified metric and sample-level prediction criteria. Discordant "
            "routes are quarantined from speed summaries. The authoritative paired "
            "metrics, runtime, incremental host memory, device memory, and status "
            "are reported in Supplementary Table S11."
        ),
    )
    replace_paragraph(
        document,
        "The deposited Nature Communications workflow",
        (
            find_paragraph(
                document, "The deposited Nature Communications workflow"
            ).text.replace(
                "in the Supplementary Material rather than backend-only effects.",
                "in Supplementary Table S12 rather than backend-only effects.",
            )
        ),
    )
    replace_paragraph(
        document,
        "The NMR task contained 1,200 training spectra",
        (
            "The NMR task contained 1,200 training spectra, 321 held-out spectra, "
            "13,000 NOESY predictor bins, and 28,355 diffusion-edited response "
            "intensities. As routine spectral preprocessing, predictor columns with "
            "chemical shifts strictly between 4.6 and 4.8 ppm were set to zero in "
            "both X training and X test matrices before inner splitting or fitting. "
            "No response column was zeroed, masked, or excluded: RMSD, Q2, "
            "response-wise errors, and intensity-stratified errors used all 28,355 "
            "response coordinates. We separated predictive model selection from "
            "implementation benchmarking (Figure 4). In the predictive analysis, "
            "five paired training-only splits and the one-standard-error rule "
            "selected five components for PLS-SVD and 50 for SIMPLS. These are "
            "family-specific predictive settings, not a common-complexity "
            "comparison."
        ),
    )
    replace_paragraph(
        document,
        "Figure 4. Separated NMR predictive",
        (
            "Figure 4. Separated NMR predictive and implementation analyses. "
            "(A) Training-only component paths and one-standard-error selections. "
            "(B) Held-out per-spectrum RMSD distributions at five PLS-SVD and 50 "
            "SIMPLS components. (C) Response-wise RMSD over all 28,355 unmasked "
            "response coordinates. (D) Aggregate RMSD across response coordinates "
            "stratified by mean absolute intensity in Y training. (E) Paired "
            "CPU/CUDA implementation benchmark within each family; split, "
            "preprocessing, target, rSVD, float64 precision, and component count are "
            "fixed. Host and GPU memory are increments above the pre-fit process "
            "baseline; GPU values include runtime context and are not workspace-only "
            "allocations. No representative spectrum was selected for this figure: "
            "panels B-D summarize all 321 held-out spectra. The separate archived "
            "illustrative spectrum used test sample 204, selected because its "
            "per-spectrum RMSD was closest to the held-out median."
        ),
    )
    replace_paragraph(
        document,
        "An independent FAISS benchmark",
        (
            "Experiment 2, representation retrieval, was independent of the LDA "
            "classification experiment and addressed a different objective. Exact "
            "CUDA FAISS cosine k-nearest-neighbour retrieval compared raw "
            "1,024-dimensional DINOv2 embeddings, unsupervised PCA scores, and "
            "supervised PLS scores using the same training index and query set. "
            "Raw-feature retrieval produced 184,328/281,167 correct top-1 queries "
            "(0.6556; Wilson 95% CI 0.6538-0.6573) and 264,078/281,167 correct "
            "top-5 queries (0.9392; 0.9383-0.9401). At 200 PLS components, the "
            "corresponding values were 183,213/281,167 (0.6516; 0.6499-0.6534) and "
            "264,221/281,167 (0.9397; 0.9388-0.9406), giving a 5.12-fold "
            "representation reduction, a 0.40-percentage-point top-1 loss, and "
            "approximately fourfold lower projection-plus-query time. PCA was the "
            "unsupervised dimensionality-reduction control. These single-run "
            "retrieval measurements are exploratory; Wilson intervals are "
            "conditional on this fixed development holdout and do not establish a "
            "top-5 improvement. The component-level comparison is Supplementary "
            "Table S13."
        ),
    )
    replace_paragraph(
        document,
        "ImageNet/DINOv2 was used as a computational stress test",
        (
            "ImageNet/DINOv2 was used as a computational stress test and as an "
            "exploratory evaluation of PLS as supervised feature extraction, not as "
            "evidence of biomedical predictive validity. Experiment 1 evaluated "
            "SIMPLS classification with argmax and latent-space LDA on the fixed "
            "1,000,000/281,167 development split (1,024 features; 1,000 classes). "
            "CUDA SIMPLS/rSVD was evaluated from 100 to 1,000 components. At 100 "
            "components, argmax classified 176,303/281,167 observations correctly "
            "(0.6270; Wilson 95% CI 0.6253-0.6288), whereas LDA classified "
            "219,113/281,167 (0.7793; 0.7778-0.7808). At 1,000 components, argmax "
            "classified 224,791/281,167 (0.7995; 0.7980-0.8010) and CUDA LDA "
            "227,535/281,167 (0.8093; 0.8078-0.8107). CPU and CUDA predictions "
            "agreed to displayed precision, while CUDA reduced total LDA time from "
            "218.8 to 14.5 s at 100 components and from 2,199.7 to 316.1 s at 1,000 "
            "components (Figure 5). These intervals quantify binomial uncertainty "
            "conditional on the fixed noncanonical holdout, not model-selection or "
            "training-sample uncertainty."
        ),
    )
    replace_paragraph(
        document,
        "Figure 5. ImageNet-scale SIMPLS classification",
        (
            "Figure 5. ImageNet experiment 1: SIMPLS classification. CUDA "
            "SIMPLS/rSVD was fitted to 1,000,000 DINOv2 embeddings and evaluated on "
            "281,167 held-out embeddings over 100-1,000 components. Panels report "
            "top-1 accuracy, total fitting plus prediction time, peak host RSS, and "
            "peak GPU allocation for argmax and latent-space LDA. Measurements are "
            "single exploratory runs on a fixed, nonstandard split. The separate "
            "raw/PCA/PLS FAISS representation-retrieval experiment is not part of "
            "this figure and is reported in Supplementary Table S13."
        ),
    )
    replace_paragraph(
        document,
        "The fastPLS R package, benchmark workflows",
        (
            "The fastPLS R package, benchmark workflows, analysis scripts, "
            "machine-readable result tables, synthetic generators, and aggregate "
            "outputs are available at https://github.com/tkcaccia/fastPLS. "
            "Low-level reusable C++ components are maintained at "
            "https://github.com/tkcaccia/kodama-cpp. The reviewed source snapshot "
            "is fastPLS version 0.99.6 at commit "
            "6e50bd318f20289101f6b723953830aefa8b95d6; no version-matched GitHub "
            "release tag existed when the analysis was frozen, so this commit, "
            "rather than the older v1.2.0 tag, is the exact reviewed identifier. "
            "A version-matched archival tag must be minted before final publication. "
            "Supplementary Table S15 maps "
            "every reported analysis to its result archive, generating script, "
            "recorded package metadata, split provenance, and archive digest. "
            "Superseded review-cycle material is indexed in "
            "benchmark/MANUSCRIPT_EVIDENCE_ARCHIVE.md and is not duplicated in the "
            "submission supplement."
        ),
    )

    evidence_anchor = find_paragraph(document, "Within each dataset, methods used identical")
    insert_after(
        evidence_anchor,
        (
            "Supplementary Table S6 provides a claim-to-evidence map. Each central "
            "quantitative statement in the main text points to one authoritative "
            "Supplementary table; expanded paths and review-cycle diagnostics are "
            "available only through the repository archive."
        ),
        style="Body Text",
    )

    discussion = find_paragraph(document, "fastPLS extends established PLS algorithms")
    conclusion = find_paragraph(document, "5. Conclusions")
    first = (
        "fastPLS accelerates established PLS estimators rather than introducing a "
        "new statistical objective. Deterministic IRLBA SIMPLS preserved de Jong's "
        "estimator in the prespecified validation, while OPLS and kernel PLS passed "
        "their independent deterministic setting study (Supplementary Table S7). "
        "rSVD remains an approximate workflow whose admissible settings and failure "
        "criteria are reported separately (Supplementary Table S8)."
    )
    second = (
        "Computational benefit was conditional on matrix shape and route. "
        "Deterministic float64 CPU SIMPLS improved on independent R implementations "
        "without changing median accuracy on the matched tasks (Supplementary Table "
        "S10). CUDA was advantageous only for a subset of numerically concordant "
        "large workloads, whereas Metal primarily established portability in the "
        "tested campaign. Stage residency, route qualification, selected-point "
        "runtime, incremental host memory, and device memory are separated in "
        "Supplementary Tables S1, S9, and S11."
    )
    third = (
        "The NMR study separates predictive model selection from backend "
        "attribution: selected SIMPLS was more accurate, while matched CPU/CUDA "
        "contrasts isolated accelerator effects within each family. The historical "
        "deposited workflow remains contextual rather than an implementation-only "
        "comparator (Supplementary Table S12). ImageNet demonstrates feasibility "
        "after foundation-model feature extraction and an exploratory supervised "
        "compression trade-off, not biomedical validity or a confirmed accuracy "
        "gain (Supplementary Table S13)."
    )
    fourth = (
        "The remaining limitations are explicit. Many component selections are "
        "best only within finite grids; repeated outer partitions quantify, but do "
        "not eliminate, training-sample uncertainty (Supplementary Table S14). "
        "Float32 halves representation size but is not uniformly faster or "
        "numerically interchangeable with float64. Confirmatory work should use "
        "deterministic float64 CPU IRLBA unless the chosen approximate, accelerator, "
        "or float32 route has passed the documented checks for the target shape."
    )
    discussion.clear()
    discussion.add_run(first)
    cursor = discussion
    for text in [second, third, fourth]:
        cursor = insert_after(cursor, text, style="Body Text")
    current = cursor._p.getnext()
    while current is not None and current is not conclusion._p:
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt

    replace_paragraph(
        document,
        "[29] Oquab M",
        (
            "[29] Oquab M, Darcet T, Moutakanni T, Vo HV, Szafraniec M, Khalidov V, "
            "et al. DINOv2: learning robust visual features without supervision. "
            "Trans Mach Learn Res. 2024. "
            "https://openreview.net/forum?id=a68SUt6zFt; arXiv:2304.07193."
        ),
    )
    normalize_submission_terminology(document)
    document.core_properties.title = (
        "fastPLS CMPB manuscript - compact authoritative evidence"
    )
    document.save(MAIN_OUTPUT)


def revise_supplement():
    shutil.copy2(SUPP_SOURCE, SUPP_OUTPUT)
    document = Document(SUPP_OUTPUT)
    truncate_from_heading(document, "S11. Preliminary review-cycle validation")
    replace_paragraph(
        document,
        "The linear kernel dispatches to the ordinary linear SIMPLS route",
        (
            "The linear kernel dispatches to the ordinary linear SIMPLS route, "
            "avoiding an unnecessary n-by-n Gram matrix. RBF and polynomial kernels "
            "construct and centre an explicit training Gram matrix and apply stored "
            "training means to held-out kernels. One float64 Gram matrix requires "
            "8n^2 bytes before copies and workspaces: 0.75 GiB at n=10,000, "
            "2.98 GiB at n=20,000, 6.71 GiB at n=30,000, and 18.63 GiB at "
            "n=50,000. Kernel centring, temporary products, and the R process require "
            "additional memory. The deterministic nonlinear-kernel validation "
            "included n_train=42-180; it establishes numerical agreement in that "
            "range, not large-sample feasibility. Conservative planning limits for "
            "the tested hardware are n<=20,000 on the 32-GB workstation and "
            "n<=10,000 on the 8-GB unified-memory Mac. These are operational "
            "ceilings based on reserving memory for multiple dense buffers, not "
            "measured performance limits. Larger nonlinear-kernel jobs were not "
            "validated and require an explicit memory calculation. The current "
            "CUDA nonlinear-kernel route is host assisted, so the 16-GB device does "
            "not remove the host-memory constraint."
        ),
    )
    replace_paragraph(
        document,
        "Reproducibility experiments use identical data",
        (
            "Reproducibility experiments use identical data, preprocessing, folds, "
            "requested components, and seeds for the configured CPU, CUDA, and "
            "Metal routes. CPU builds delegate dense linear algebra to the "
            "BLAS/LAPACK implementation linked to R and can use OpenBLAS where it "
            "is installed. That build capability is recorded separately from "
            "measured performance. The primary software-comparison runs used one "
            "effective BLAS thread. No controlled one-thread versus multiple-thread "
            "scaling experiment was conducted, and no multicore speed-up is "
            "inferred. Because rSVD is stochastic and nearly tied singular values "
            "can rotate an equivalent basis, raw component columns are sign-aligned "
            "only for descriptive comparisons. Primary endpoints are prediction "
            "agreement, absolute predictive-metric difference, relative Frobenius "
            "prediction error, selected component count, numerical failures, and "
            "run-to-run variability. Principal angles compare latent subspaces."
        ),
    )
    replace_paragraph(
        document,
        "This supplement describes fastPLS version",
        (
            "This supplement describes fastPLS version 0.99.6. Source provenance "
            "is analysis specific: Table S15 maps each quantitative analysis to "
            "its result archive and records an exact package commit only when the "
            "run metadata captured it. Package version, result date, or a later "
            "manuscript commit is not treated as evidence of the historical "
            "computational SHA."
        ),
    )
    replace_paragraph(
        document,
        "Current benchmark workflows record repository commit",
        (
            "Current benchmark workflows record repository state, benchmark-script "
            "checksum, package version, reusable-core commit when available, "
            "compiler, BLAS/LAPACK, thread settings, accelerator libraries, seeds, "
            "rSVD controls, and data/split identifiers. Earlier archives did not "
            "always capture an exact Git SHA. Table S15 distinguishes recorded and "
            "unrecoverable source states and supplies an archive digest rather than "
            "assigning later commits retrospectively."
        ),
    )
    replace_paragraph(
        document,
        "The simulated datasets were used only",
        (
            "The simulated datasets were used only for the formal SIMPLS "
            "estimator-preservation and rSVD reliability analyses; no separate "
            "n/p/q performance-scaling sweep is claimed. Five multivariate "
            "regression regimes and three dummy-response classification regimes "
            "were generated with seeds 101, 202, and 303. Regression data used "
            "independent standard-normal latent scores and Gaussian predictor and "
            "response loading matrices. Predictor and response matrices were "
            "formed by multiplying the latent-score matrix by the transposed "
            "loading matrices and adding independent Gaussian noise with standard "
            "deviation 0.05. The regimes covered p < n, p > n, low- and high-rank "
            "Y, near-collinear predictors, and exact rank deficiency. "
            "Near-collinearity was created by making two loading columns linear "
            "combinations of the first loading column plus noise with standard "
            "deviation 1 × 10⁻⁷; exact rank deficiency was created by duplicating "
            "ten predictor columns. Training and held-out rows were generated in "
            "one draw and separated before model fitting. The authoritative "
            "deterministic and approximate summaries are Tables S7 and S8; exact "
            "regime dimensions remain in the repository result archive."
        ),
    )
    replace_paragraph(
        document,
        "The corresponding deterministic results comprised",
        (
            "The corresponding deterministic results comprised 117 component-level "
            "comparisons. All passed the prespecified tolerances. Within synthetic "
            "classification, the worst relative prediction and coefficient errors "
            "were 2.55 × 10⁻⁸ and 4.91 × 10⁻⁸; within synthetic regression they "
            "were 1.09 × 10⁻⁵ and 1.13 × 10⁻⁵. Maximum score-subspace angles were "
            "2.09 × 10⁻⁶ degrees for classification and 0.00143 degrees for "
            "regression. Fixed 5-fold component selection agreed with "
            "pls::simpls.fit in all eight synthetic tasks. The authoritative "
            "deterministic summary is Table S7; complete task-level endpoints and "
            "curves are retained in the repository archive."
        ),
    )
    replace_paragraph(
        document,
        "For approximate rSVD, oversampling 10",
        (
            "For approximate rSVD, oversampling 10 with one power iteration passed "
            "101 of 117 component-level checks, whereas two power iterations passed "
            "117 of 117. The focused CUDA audit passed all evaluated points with "
            "either four power iterations at oversampling 10 or oversampling 20. "
            "The authoritative numerical audit is Table S8; full endpoints remain "
            "in the repository archive. These are solver-reliability results, not "
            "estimator-equivalence evidence."
        ),
    )

    reporting_heading = find_paragraph(document, "S10. Reporting conventions")
    insert_after(
        reporting_heading,
        (
            "Completed, failed, timed-out, memory-killed, and unavailable routes "
            "remain explicit. Component counts at a tested boundary are described "
            "as best within the evaluated grid. Deterministic estimator validation "
            "is not pooled with approximate rSVD agreement, and speed-up is "
            "summarized only for numerically concordant paired routes."
        ),
        style="Body Text",
    )

    document.add_heading("S11. Authoritative evidence map", level=1)
    document.add_paragraph(
        (
            "This compact supplement contains one authoritative table for each "
            "central evidence class. Expanded component paths, intermediate "
            "review-cycle summaries, sensitivity analyses, and diagnostic figures "
            "remain available through benchmark/MANUSCRIPT_EVIDENCE_ARCHIVE.md. "
            "They are reproducibility records, not parallel sources for manuscript "
            "claims."
        )
    )
    add_caption(
        document,
        "Table S6. Claim-to-evidence map. Each central main-text claim is assigned "
        "to one authoritative Supplementary source.",
    )
    add_table(
        document,
        ["Main-text claim", "Authority", "Evidence scope"],
        build_claim_map(),
        6.5,
    )

    document.add_heading("S12. Deterministic estimator validation", level=1)
    document.add_paragraph(
        (
            "Deterministic estimator preservation and approximate-solver agreement "
            "are separate questions. Table S7 contains only deterministic float64 "
            "CPU validation. rSVD is excluded and evaluated in Section S13."
        )
    )
    add_caption(
        document,
        "Table S7. Definitive deterministic estimator-validation summary. Errors "
        "are worst observed values; angles are in degrees.",
    )
    add_table(
        document,
        [
            "fastPLS scope",
            "Reference",
            "Endpoint pass",
            "Selection agreement",
            "Max pred. rel. error",
            "Max coef. rel. error",
            "Max angle",
            "Max metric diff.",
            "Conclusion",
        ],
        build_estimator_validation(),
        5.5,
    )

    document.add_heading("S13. Approximate rSVD reliability", level=1)
    document.add_paragraph(
        (
            "rSVD uses a fixed seed, Gaussian range sketch, oversampling, and power "
            "iterations. A route is qualified only when all prespecified prediction, "
            "subspace, label, and metric checks pass. Qualified rSVD remains "
            "approximate and is not estimator-equivalence evidence."
        )
    )
    add_caption(
        document,
        "Table S8. Definitive rSVD numerical-audit summary. CPU contains 117 "
        "component-level tests; CUDA contains eight task/endpoint tests per setting.",
    )
    add_table(
        document,
        [
            "Backend",
            "Oversample",
            "Power",
            "Passes",
            "Max pred. rel. error",
            "Min pred. corr.",
            "Max score angle",
            "Min label agree.",
            "Max metric diff.",
            "Status",
        ],
        build_rsvd_reliability(),
        5.5,
    )

    document.add_heading("S14. Float32 capability", level=1)
    document.add_paragraph(
        (
            "Float64 is the confirmatory reference. Float32 input halves raw matrix "
            "storage but does not guarantee lower peak process memory, faster "
            "execution, or numerical agreement. Public calls allow validated "
            "routes, warn for experimental, hybrid, or measured-risk routes, and "
            "stop before allocation for unavailable routes."
        )
    )
    add_caption(
        document,
        "Table S9. Definitive float32 capability matrix. Endpoint statuses are "
        "validated, experimental, hybrid, unavailable, or failed; extreme-response "
        "status refers to q >= 10,000 with at least 50 components.",
    )
    add_table(
        document,
        [
            "Family",
            "Kernel",
            "Backend",
            "Endpoint status",
            "Solver",
            "Residency",
            "Windows",
            "Extreme q",
        ],
        build_float32_capability(),
        5.1,
    )

    document.add_heading("S15. External software comparison", level=1)
    document.add_paragraph(
        (
            "The primary software comparison used float64, deterministic CPU "
            "SIMPLS, the same split and component count, and one CPU process. "
            "fastPLS argmax is estimator matched to pls::simpls.fit; LDA is a "
            "workflow comparison because the prediction head differs. Memory is "
            "absolute process RSS and is reported for feasibility, not isolated "
            "algorithmic allocation. Across 126 attempted external-package "
            "dataset/method runs, 110 completed and 16 did not: 12 were documented "
            "package limitations, two were killed at the timeout, and two returned "
            "execution errors. Table S10 summarizes the 65 successful external "
            "classification rows with compatible outputs and 17 fastPLS rows; "
            "non-successful attempts remain in the machine-readable run table."
        )
    )
    add_caption(
        document,
        "Table S10. Definitive one-CPU external comparison. Cells report accuracy; "
        "total fitting-plus-prediction time; peak process RSS.",
    )
    add_table(
        document,
        [
            "Dataset",
            "A",
            "fastPLS argmax",
            "fastPLS LDA",
            "Fastest external",
            "Best-accuracy external",
        ],
        build_external_comparison(),
        5.4,
    )

    document.add_heading("S16. Selected-point CPU and CUDA benchmark", level=1)
    document.add_paragraph(
        (
            "Every row uses the component count selected from the training data for "
            "that dataset and family. CPU and CUDA values share the split, model "
            "family, precision, solver family, and requested component count. "
            "Incremental host RSS is the fit-window peak minus the immediately "
            "pre-fit baseline. Incremental GPU memory includes runtime/context "
            "state and is not workspace-only allocation."
        )
    )
    add_caption(
        document,
        "Table S11. Definitive selected-point performance. Metric, time, and host "
        "memory are CPU / CUDA; GPU memory is the CUDA increment. Endpoint choices "
        "are best within the evaluated training grid.",
    )
    add_table(
        document,
        [
            "Dataset",
            "Family",
            "A",
            "Selection status",
            "Metric CPU / CUDA",
            "Time s CPU / CUDA",
            "Delta host MB CPU / CUDA",
            "Delta GPU MB",
            "Numerical status",
        ],
        build_selected_performance(),
        5.0,
    )

    document.add_heading("S17. NMR case study", level=1)
    document.add_paragraph(
        (
            "The family-selected predictive analysis and paired backend analysis "
            "answer different questions. The former compares the best values within "
            "the training grids (five PLS-SVD and 50 SIMPLS components); the latter "
            "changes only CPU versus CUDA within family. The deposited 165-component "
            "workflow uses the original centring-only protocol and is historical "
            "context. Predictor columns with chemical shifts strictly between 4.6 "
            "and 4.8 ppm were set to zero in both training and test predictor "
            "matrices before inner splitting or fitting. No response column was "
            "zeroed, masked, or excluded; all reported response metrics use all "
            "28,355 response coordinates. The current main-text Figure 4 summarizes "
            "all 321 held-out spectra and does not display one representative "
            "spectrum. In the separate archived spectrum-overlay diagnostic, test "
            "sample 204 was selected by the prespecified descriptive rule of RMSD "
            "closest to the held-out median; it was not the best-predicted spectrum."
        )
    )
    add_caption(
        document,
        "Table S12. Definitive NMR evidence. Time is fitting plus prediction; host "
        "and GPU memory are increments above the pre-fit process baseline.",
    )
    add_table(
        document,
        [
            "Analysis",
            "Family",
            "Implementation",
            "A",
            "Predictive metric",
            "Time s",
            "Delta host MB",
            "Delta GPU MB",
            "Error detail / scope",
        ],
        build_nmr(),
        5.4,
    )

    document.add_heading("S18. ImageNet exploratory stress test", level=1)
    document.add_paragraph(
        (
            "The pooled archive contained 1,281,167 precomputed DINOv2 embeddings "
            "with 1,024 features and 1,000 classes. Seed 123 assigned 1,000,000 rows "
            "to development training and 281,167 to a complementary holdout. This "
            "was not the canonical ImageNet split. Two separate experiments used "
            "that split. Experiment 1 fitted SIMPLS and evaluated argmax or "
            "latent-space LDA classification; its rows are single exploratory runs. "
            "Experiment 2 compared exact cosine-neighbour retrieval on raw DINOv2, "
            "PCA, and PLS representations with FAISS. It used a different estimator "
            "and objective from LDA classification. FAISS exact-query timing used "
            "three repeats after one representation fit; transformation, held-out "
            "projection, and query are included in end-to-end time."
        )
    )
    add_caption(
        document,
        "Table S13. Definitive ImageNet classification and representation results. "
        "The experiment establishes computational feasibility and compression "
        "trade-offs, not biomedical validity.",
    )
    add_table(
        document,
        [
            "Experiment",
            "Head / representation",
            "A/dim.",
            "Top-1",
            "Top-5",
            "End-to-end s",
            "Host RSS MB",
            "GPU MB",
            "Qualification",
        ],
        build_imagenet(),
        5.3,
    )

    document.add_heading("S19. Repeated outer-partition uncertainty", level=1)
    document.add_paragraph(
        (
            "MetRef, GTEx v8, and Retina used ten stratified 80/20 outer "
            "partitions with 5-fold training-only component selection. NMR used "
            "five random 80/20 outer partitions and 3-fold selection. "
            "Classification used deterministic float64 CPU IRLBA; NMR used fixed-"
            "seed float64 CUDA rSVD for feasibility. Empirical ranges are "
            "descriptive across partitions, not confidence intervals."
        )
    )
    add_caption(
        document,
        "Table S14. Definitive repeated-partition predictive dispersion and "
        "selection stability. Values are mean (SD), empirical 2.5th-97.5th "
        "percentile range, and selected-component range.",
    )
    add_table(
        document,
        [
            "Dataset",
            "Family",
            "Head",
            "Outer n",
            "Metric",
            "Mean (SD)",
            "Empirical range",
            "Selected A range",
            "Upper-bound freq.",
            "Constraint",
        ],
        build_outer_uncertainty(),
        5.2,
    )

    document.add_heading("S20. Analysis provenance", level=1)
    document.add_paragraph(
        (
            "The ledger never infers a historical commit from a package version or "
            "result date. Where a run did not record its Git SHA, the source status "
            "is explicitly not recoverable. SHA-256 prefixes identify immutable "
            "result archives; full digests are retained in the machine-readable "
            "ledger."
        )
    )
    add_caption(
        document,
        "Table S15. Definitive analysis provenance. Archive paths are relative to "
        "the repository root.",
    )
    add_table(
        document,
        [
            "ID",
            "Authoritative output",
            "Result archive",
            "Version",
            "Source status",
            "Generating script",
            "SHA-256 prefix",
        ],
        build_provenance(),
        4.9,
    )

    document.add_heading("S21. Repository-only detailed material", level=1)
    document.add_paragraph(
        (
            "Full component paths, ablations, per-run rows, sensitivity analyses, "
            "review-cycle figures, and the former cycle-66 expanded supplement are "
            "indexed in benchmark/MANUSCRIPT_EVIDENCE_ARCHIVE.md. Their underlying "
            "CSV, RDS, PDF, PNG, log, and session-information files remain "
            "available for audit. They are deliberately not reproduced here."
        )
    )
    document.add_heading("Supplementary references", level=1)
    document.add_paragraph(
        (
            "References are numbered as in the main manuscript. The final mapping "
            "is Retina [25], Tabula Muris [26], PRISM [27], ImageNet [28], DINOv2 "
            "[29], CIFAR-100 [30], UNI [31], and Prov-GigaPath [32]."
        )
    )

    replace_text_everywhere(document, "Supplementary Table S41", "Supplementary Table S15")
    replace_text_everywhere(document, "Table S41", "Table S15")
    normalize_submission_terminology(document)
    document.core_properties.title = (
        "fastPLS compact supplementary information - authoritative evidence"
    )
    document.save(SUPP_OUTPUT)


def audit_opls_kernel_counts():
    """Prevent the superseded 18-comparison study from re-entering submission files."""
    main_document = Document(MAIN_OUTPUT)
    supplement_document = Document(SUPP_OUTPUT)

    def document_text(document):
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        cells = [
            cell.text
            for table in document.tables
            for row in table.rows
            for cell in row.cells
        ]
        return "\n".join(paragraphs + cells)

    main_text = document_text(main_document)
    supplement_text = document_text(supplement_document)
    combined = f"{main_text}\n{supplement_text}"
    obsolete = (
        "18 endpoint",
        "18 fixed-fold",
        "18 setting/task",
        "all 18",
    )
    found = [phrase for phrase in obsolete if phrase in combined]
    if found:
        raise RuntimeError(
            "Superseded OPLS/kernel-PLS counts remain: " + ", ".join(found)
        )

    required_main = (
        "all 66 setting/task endpoint comparisons",
        "selected component count agreed in all 66 comparisons",
        "all 1,540 fold-by-component fits",
    )
    missing_main = [phrase for phrase in required_main if phrase not in main_text]
    if missing_main:
        raise RuntimeError(
            "Main manuscript is missing synchronized OPLS/kernel-PLS counts: "
            + ", ".join(missing_main)
        )

    required_supplement = ("66 / 66", "1,540")
    missing_supplement = [
        phrase for phrase in required_supplement if phrase not in supplement_text
    ]
    if missing_supplement:
        raise RuntimeError(
            "Supplement is missing synchronized OPLS/kernel-PLS counts: "
            + ", ".join(missing_supplement)
        )


def audit_algorithm_notation():
    """Keep component-prefix and cross-validation symbols unambiguous."""
    document = Document(MAIN_OUTPUT)
    paragraphs = "\n".join(paragraph.text for paragraph in document.paragraphs)
    cells = "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    combined = f"{paragraphs}\n{cells}"
    obsolete = ("requested prefixes K", "If a is in K")
    found = [phrase for phrase in obsolete if phrase in combined]
    if found:
        raise RuntimeError(
            "Ambiguous component-prefix notation remains: " + ", ".join(found)
        )
    required = (
        "C denotes the set of requested component counts",
        "requested component-count set C",
        "If a is in C",
        "K denotes the number of cross-validation folds",
    )
    missing = [phrase for phrase in required if phrase not in combined]
    if missing:
        raise RuntimeError(
            "Synchronized algorithm notation is incomplete: " + ", ".join(missing)
        )


def audit_discussion_deduplication():
    """Keep implementation-mechanism wording out of the Discussion recap."""
    document = Document(MAIN_OUTPUT)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    duplicated_openings = (
        "The computational gain derives from compiled execution",
        "The speed gain comes from compiled sequential execution",
    )
    found = [phrase for phrase in duplicated_openings if phrase in text]
    if found:
        raise RuntimeError(
            "Duplicated Discussion wording remains: " + ", ".join(found)
        )


def audit_cpu_parallel_claims():
    """Separate optional BLAS linkage from measured multicore performance."""
    main_document = Document(MAIN_OUTPUT)
    supplement_document = Document(SUPP_OUTPUT)
    main_text = "\n".join(p.text for p in main_document.paragraphs)
    supplement_text = "\n".join(p.text for p in supplement_document.paragraphs)
    combined = f"{main_text}\n{supplement_text}"

    forbidden = (
        "Those libraries may use one or more threads",
        "multithread-capable CPU execution",
        "CPU multithreading speedup",
    )
    found = [phrase for phrase in forbidden if phrase in combined]
    if found:
        raise RuntimeError(
            "Unqualified CPU parallel-performance wording remains: "
            + ", ".join(found)
        )

    required = (
        "implementation capability, not evidence of multicore acceleration",
        "All primary external-software comparisons used one effective BLAS thread",
        "no multicore speed-up is claimed",
        "build capability is recorded separately from measured performance",
        "no multicore speed-up is inferred",
    )
    missing = [phrase for phrase in required if phrase not in combined]
    if missing:
        raise RuntimeError(
            "CPU capability/performance separation is incomplete: "
            + ", ".join(missing)
        )


def audit_cycle68_reviewer_items():
    """Verify the requested kernel, run-accounting, NMR, and ImageNet clarifications."""
    main_document = Document(MAIN_OUTPUT)
    supplement_document = Document(SUPP_OUTPUT)

    def document_text(document):
        return "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ]
        )

    main_text = document_text(main_document)
    supplement_text = document_text(supplement_document)
    combined = f"{main_text}\n{supplement_text}"
    required = (
        "8n^2 bytes",
        "n<=20,000",
        "n<=10,000",
        "110 completed successfully and 16 did not",
        "12 documented package limitations",
        "8,739/10,000",
        "219,113/281,167",
        "No response column was zeroed, masked, or excluded",
        "No representative spectrum was selected for this figure",
        "test sample 204",
        "Experiment 1 evaluated SIMPLS classification",
        "Experiment 2, representation retrieval",
        "https://openreview.net/forum?id=a68SUt6zFt",
        "fastPLS version 0.99.6 at commit",
    )
    missing = [phrase for phrase in required if phrase not in combined]
    if missing:
        raise RuntimeError(
            "Cycle-68 reviewer clarifications are incomplete: " + ", ".join(missing)
        )


def audit_cross_references_and_terminology():
    """Reject stale cross-references and noncanonical terminology."""
    main_document = Document(MAIN_OUTPUT)
    supplement_document = Document(SUPP_OUTPUT)

    def document_text(document):
        return "\n".join(
            [paragraph.text for paragraph in document.paragraphs]
            + [
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ]
        )

    combined = f"{document_text(main_document)}\n{document_text(supplement_document)}"
    forbidden = (
        "Supplementary Table S41",
        "Table S41",
        "PLSSVD",
        "PLS SVD",
        "RSVD",
        "R-SVD",
        "Float32",
        "Float64",
        "Cuda",
        "METAL",
        "prediction-head",
    )
    found = [phrase for phrase in forbidden if phrase in combined]
    if found:
        raise RuntimeError(
            "Stale cross-reference or terminology remains: " + ", ".join(found)
        )
    required = (
        "PLS-SVD",
        "rSVD",
        "float32",
        "float64",
        "CPU",
        "CUDA",
        "Metal",
        "prediction head",
        "Supplementary Table S15",
    )
    missing = [phrase for phrase in required if phrase not in combined]
    if missing:
        raise RuntimeError(
            "Canonical submission terminology is incomplete: " + ", ".join(missing)
        )


def main():
    OUTPUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    audit_opls_kernel_counts()
    audit_algorithm_notation()
    audit_discussion_deduplication()
    audit_cpu_parallel_claims()
    audit_cycle68_reviewer_items()
    audit_cross_references_and_terminology()
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)


if __name__ == "__main__":
    main()
