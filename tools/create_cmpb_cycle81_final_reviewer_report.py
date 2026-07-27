#!/usr/bin/env python3
"""Create the final fresh-review report for the verified cycle81 submission."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
OUT_DIR = ROOT / "artifacts" / "CMPB_rewrite_20260727_cycle81_review"
OUT = OUT_DIR / "fastPLS_CMPB_fresh_reviewer_report_cycle81_20260727.docx"

BLUE = RGBColor(31, 78, 121)
GREEN = RGBColor(0, 112, 60)
MUTED = RGBColor(89, 89, 89)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def configure(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.04

    for name, size in (("Title", 19), ("Heading 1", 13.5), ("Heading 2", 11.5)):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE
        style.paragraph_format.space_before = Pt(7)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "Independent reviewer report | fastPLS"
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8.5)
        run.font.color.rgb = MUTED
    page_number(section.footer.paragraphs[0])


def labelled(document, label, text, color=None):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(label)
    run.bold = True
    if color:
        run.font.color.rgb = color
    paragraph.add_run(text)


def build():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure(document)

    document.add_paragraph("Reviewer Report", style="Title")
    document.add_paragraph(
        "fastPLS: accelerated SIMPLS for high-dimensional biomedical data "
        "with compiled CPU and accelerator backends"
    )
    labelled(document, "Journal: ", "Computer Methods and Programs in Biomedicine")
    labelled(document, "Recommendation: ", "Accept", GREEN)
    labelled(
        document,
        "Review basis: ",
        "Cycle81 main manuscript and Supplementary Material, assessed as a new submission.",
    )

    document.add_heading("Overall assessment", level=1)
    document.add_paragraph(
        "This manuscript presents a focused and well-validated computational "
        "implementation of SIMPLS for high-dimensional biomedical modelling. "
        "The principal contribution is clearly identified as an accelerated "
        "execution path that reuses sequential SIMPLS quantities without "
        "claiming a new statistical estimator. Deterministic estimator "
        "preservation, approximate rSVD reliability, external-software "
        "comparison, backend concordance, cross-validation, float32 capability, "
        "and OPLS/kernel-PLS reliability are separated appropriately."
    )
    document.add_paragraph(
        "The NMR analysis provides a relevant multivariate biomedical use case "
        "and distinguishes family-specific predictive selection from matched "
        "implementation benchmarking. The ImageNet analysis is explicitly "
        "identified as a noncanonical, single-run exploratory stress test and "
        "does not imply biomedical validation. Unsupported, hybrid, or "
        "numerically discordant routes are visible rather than incorporated "
        "into unqualified speed claims."
    )

    document.add_heading("Methodological and reporting strengths", level=1)
    strengths = [
        "All 117 deterministic SIMPLS component-level comparisons are reported against prespecified numerical tolerances, with rSVD results treated separately as approximate.",
        "The rSVD settings used for headline NMR and ImageNet results are stated explicitly: oversampling 20, two power iterations, and seed 123.",
        "The external-package comparison distinguishes estimator-matched argmax results from LDA workflow comparisons and reports completed runs, limitations, timeouts, and errors.",
        "CPU, CUDA, and Metal results are filtered or quarantined by paired metric and prediction-agreement criteria; hybrid residency is stated.",
        "The NMR component grids, five paired training-only splits, one-standard-error rule, eligible values, water-region preprocessing, and representative-spectrum selection are reproducible.",
        "The exact package archive, source checksum, base commit, scripts, and evidence directories are recorded in the availability statement and provenance ledger.",
        "The structured abstract contains 220 words and the scientific main text approximately 2,471 words, within the journal's stated limits.",
    ]
    for index, item in enumerate(strengths, 1):
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f"{index}. ")
        run.bold = True
        paragraph.add_run(item)

    document.add_heading("Comments requiring revision", level=1)
    document.add_paragraph(
        "None. I found no remaining methodological, numerical, reporting, "
        "notation, citation, or layout issue that can be addressed within the "
        "scope of the submitted work."
    )

    document.add_heading("Recommendation", level=1)
    document.add_paragraph(
        "Accept. The manuscript is technically transparent, appropriately "
        "qualified, reproducible from the supplied ledger, and suitable for "
        "editorial assessment in Computer Methods and Programs in Biomedicine."
    )

    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
