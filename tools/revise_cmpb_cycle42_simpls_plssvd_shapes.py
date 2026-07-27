#!/usr/bin/env python3

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle41"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260726_cycle42"
EVIDENCE = ROOT / "benchmark_results" / "simpls_vs_plssvd_shapes_20260726"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle41_0.99.6_20260726.docx"
SUPP_SOURCE = (
    SOURCE / "fastPLS_CMPB_supplement_cycle41_0.99.6_20260726.docx"
)
MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle42_0.99.6_20260726.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle42_0.99.6_20260726.docx"
FIGURE = EVIDENCE / "simpls_vs_plssvd_shapes_runtime_ratio.png"
PAIRED = EVIDENCE / "simpls_vs_plssvd_shapes_paired.csv"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_paragraph(document, prefix, text):
    paragraph = find_paragraph(document, prefix)
    style = paragraph.style
    paragraph.clear()
    paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def insert_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    inserted = Paragraph(new_p, paragraph._parent)
    if style is not None:
        inserted.style = style
    if text:
        inserted.add_run(text)
    return inserted


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))


def prevent_row_splitting(table):
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:cantSplit")) is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def set_cell_width(cell, width_inches):
    width = int(width_inches * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def style_table(table, widths, font_size=5.2):
    table.autofit = False
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    repeat_header(table.rows[0])
    prevent_row_splitting(table)
    for row_index, row in enumerate(table.rows):
        for column_index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[column_index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.alignment = (
                    WD_ALIGN_PARAGRAPH.LEFT
                    if column_index in (0, 1)
                    else WD_ALIGN_PARAGRAPH.CENTER
                )
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
                    if row_index == 0:
                        run.bold = True


def add_table_after(paragraph, headers, rows, widths):
    document = paragraph._parent
    table = document.add_table(
        rows=1, cols=len(headers), width=Inches(sum(widths))
    )
    for cell, value in zip(table.rows[0].cells, headers):
        cell.text = value
    for values in rows:
        row = table.add_row()
        for cell, value in zip(row.cells, values):
            cell.text = str(value)
    paragraph._p.addnext(table._tbl)
    style_table(table, widths)
    return table


def format_time(value, iqr):
    return f"{value:.3f} ({iqr:.3f})"


def table_rows():
    data = pd.read_csv(PAIRED)
    environment = {
        ("Chiamaka Linux", "CPU"): "Linux CPU",
        ("Chiamaka Linux", "CUDA"): "Linux CUDA",
        ("Apple M3", "CPU"): "M3 CPU",
        ("Apple M3", "METAL"): "M3 Metal",
    }
    shape_name = {
        "wide": "Wide",
        "tall_thin": "Tall-thin",
        "high_response": "High response",
        "balanced": "Balanced",
        "high_components": "High components",
    }
    rows = []
    for _, row in data.iterrows():
        rows.append(
            (
                environment[(row.platform, row.backend)],
                (
                    f"{shape_name[row['shape']]} "
                    f"{int(row.n_train)}/{int(row.p)}/{int(row.q)}/{int(row.ncomp)}"
                ),
                format_time(
                    row.median_total_sec_plssvd, row.iqr_total_sec_plssvd
                ),
                format_time(
                    row.median_total_sec_simpls, row.iqr_total_sec_simpls
                ),
                f"{row.simpls_over_plssvd_time:.3f}",
                f"{row.median_rmsd_plssvd:.3f}/{row.median_rmsd_simpls:.3f}",
                (
                    f"{int(row.completed_runs_plssvd)}/"
                    f"{int(row.completed_runs_simpls)}"
                ),
            )
        )
    return rows


def revise_main():
    document = Document(MAIN_SOURCE)

    results = find_paragraph(document, "Results: The deterministic IRLBA path")
    old = results.text
    marker = (
        "In the primary estimator-matched software comparison, float64 SIMPLS "
        "using deterministic CPU IRLBA"
    )
    addition = (
        "In a direct five-shape rSVD workflow comparison that changed only PLS "
        "family, the CUDA SIMPLS/PLS-SVD total-time ratio was 0.918-0.979; CPU "
        "and Metal ratios crossed unity, demonstrating matrix-shape-dependent "
        "rather than universal runtime parity. "
    )
    if marker not in old:
        raise RuntimeError("Abstract insertion marker not found")
    results.clear()
    results.add_run(old.replace(marker, addition + marker))

    methods_anchor = find_paragraph(
        document, "A controlled implementation ablation then isolated"
    )
    insert_after(
        methods_anchor,
        (
            "A separate prespecified family-speed experiment compared PLS-SVD "
            "and SIMPLS across five synthetic multivariate-regression shapes: "
            "wide (400/2,000/20, A=10), tall-thin (5,000/50/20, A=10), "
            "high-response (1,000/300/500, A=50), balanced "
            "(5,000/500/50, A=50), and high-component "
            "(3,000/768/200, A=100), where dimensions are n/p/q. Within each "
            "execution environment and shape, generated X and Y, train/test "
            "split, component count, float64 precision, centering, public "
            "fit-plus-predict path, and rSVD controls were identical; only the "
            "PLS family changed. Oversampling was 10, power iterations were one, "
            "and seeds were 101-103. Each fit ran in an isolated process. "
            "Ratios compare SIMPLS with PLS-SVD within the same machine and "
            "backend; absolute CPU times are not compared across machines."
        ),
        methods_anchor.style,
    )

    figure2_caption = find_paragraph(
        document, "Figure 2. Matched CPU and CUDA outer-test performance"
    )
    heading = insert_after(
        figure2_caption,
        "3.1.1 Direct PLS-SVD versus SIMPLS speed comparison",
        find_paragraph(document, "3.2 NMR multivariate regression").style,
    )
    result_text = insert_after(
        heading,
        (
            "The direct matched experiment showed a crossover rather than a "
            "universal winner (Figure 3; Supplementary Table S37). On Chiamaka "
            "CUDA, SIMPLS was 2.1-8.2% faster than PLS-SVD across "
            "all five shapes (SIMPLS/PLS-SVD ratio 0.918-0.979). On the same "
            "Linux CPU, PLS-SVD was equal or faster (ratio 1.000-3.843). On the "
            "Apple M3 CPU, SIMPLS was faster in four of five shapes "
            "(ratio 0.605-0.769) but 2.29-fold slower for the balanced shape. "
            "Metal SIMPLS was 4.69-fold faster for the tall-thin shape and near "
            "parity for the high-response shape, whereas PLS-SVD was "
            "1.42-2.66-fold faster for the remaining three shapes. These results "
            "support a conditional computational claim: optimized sequential "
            "SIMPLS can approach or exceed one-shot PLS-SVD, but the crossover "
            "depends on n, p, q, A, and backend. Because the two families define "
            "different estimators, their held-out RMSD values are reported for "
            "context but are not treated as estimator-agreement evidence."
        ),
        find_paragraph(document, "Table 1 and Figure 2 show").style,
    )
    picture = insert_after(result_text)
    picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture.add_run().add_picture(str(FIGURE), width=Inches(6.75))
    caption = insert_after(
        picture,
        (
            "Figure 3. Direct matched total-time comparison of SIMPLS with "
            "PLS-SVD across matrix shapes. The y-axis is the median "
            "SIMPLS/PLS-SVD fit-plus-prediction time ratio on a logarithmic "
            "scale; values below one favor SIMPLS and values above one favor "
            "PLS-SVD. Dashed lines delimit +/-5% around parity. Each point "
            "summarizes three isolated runs with identical X, Y, split, A, "
            "float64 precision, centering, prediction path, and rSVD controls "
            "(oversampling 10, one power iteration, seeds 101-103). Comparisons "
            "are within machine and backend."
        ),
        figure2_caption.style,
    )

    replace_paragraph(
        document,
        "Figure 2. Matched CPU and CUDA outer-test performance",
        figure2_caption.text.replace("Figure 4", "Figure 5"),
    )
    replace_paragraph(
        document,
        "A separate analysis fixed fastPLS PLS-SVD and SIMPLS at 100 components",
        find_paragraph(
            document,
            "A separate analysis fixed fastPLS PLS-SVD and SIMPLS at 100 components",
        ).text.replace("Figure 3C-D", "Figure 4C-D"),
    )
    replace_paragraph(
        document,
        "Figure 3. Fixed-complexity NMR analysis",
        find_paragraph(
            document, "Figure 3. Fixed-complexity NMR analysis"
        ).text.replace("Figure 3.", "Figure 4.", 1),
    )
    replace_paragraph(
        document,
        "Figure 4. Exploratory matched ImageNet/DINOv2 retrieval",
        find_paragraph(
            document, "Figure 4. Exploratory matched ImageNet/DINOv2 retrieval"
        ).text.replace("Figure 4.", "Figure 5.", 1),
    )
    replace_paragraph(
        document,
        "ImageNet/DINOv2 was used as a million-sample",
        find_paragraph(
            document, "ImageNet/DINOv2 was used as a million-sample"
        ).text.replace("Table 2 and Figure 4", "Table 2 and Figure 5"),
    )

    discussion = find_paragraph(
        document, "fastPLS extends established PLS algorithms"
    )
    insert_after(
        discussion,
        (
            "The direct family-speed experiment also bounds the SIMPLS claim. "
            "The optimized sequential implementation is not uniformly as fast "
            "as PLS-SVD. It reached near parity across all five CUDA shapes and "
            "outperformed PLS-SVD in several CPU and Metal regimes, but PLS-SVD "
            "remained faster for other shapes, especially the balanced Linux-CPU "
            "case and three Metal cases. Thus the contribution is removal of much "
            "of the sequential-execution penalty in favorable regimes, while "
            "retaining SIMPLS component construction; method choice should still "
            "consider matrix shape, component count, backend, and predictive "
            "behaviour."
        ),
        discussion.style,
    )

    conclusion = find_paragraph(
        document, "fastPLS combines an accelerated sequential SIMPLS"
    )
    conclusion_text = conclusion.text.replace(
        "fastPLS combines an accelerated sequential SIMPLS implementation",
        (
            "fastPLS combines a shape-aware accelerated sequential SIMPLS "
            "implementation that can approach PLS-SVD runtime in suitable regimes"
        ),
    )
    conclusion.clear()
    conclusion.add_run(conclusion_text)

    document.core_properties.title = (
        "fastPLS CMPB manuscript - matched SIMPLS versus PLS-SVD shapes"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)
    heading_style = find_paragraph(
        document, "S32. Float32 capability boundaries"
    ).style
    body_style = find_paragraph(
        document, "Measured anchors. Float32 approximately halved"
    ).style
    caption_style = find_paragraph(
        document, "Table S34. Controlled deterministic SIMPLS"
    ).style

    document.add_paragraph(
        "S33. Direct matched PLS-SVD versus SIMPLS speed comparison",
        style=heading_style,
    )
    document.add_paragraph(
        (
            "This experiment was designed to test the computational claim about "
            "accelerated SIMPLS directly. Five synthetic multivariate-regression "
            "shapes covered wide, tall-thin, high-response, balanced, and "
            "high-component regimes. Data were generated once per shape with "
            "data seed 777. Within every shape, machine, backend, and replicate, "
            "PLS-SVD and SIMPLS received identical matrices, training/test rows, "
            "component count, centering, float64 storage, rSVD oversampling "
            "(10), power iterations (one), random seed (101-103), and public "
            "fit-plus-predict workflow. Only method changed. Each run used an "
            "isolated process and all 120 family-specific executions completed. "
            "The table reports median (IQR) seconds across three runs. The ratio "
            "is SIMPLS/PLS-SVD; it is interpreted only within a machine/backend "
            "pair. Held-out RMSD is descriptive because PLS-SVD and SIMPLS are "
            "different estimators."
        ),
        style=body_style,
    )
    table_caption = document.add_paragraph(
        (
            "Table S37. Direct matched PLS-family speed comparison across matrix "
            "shapes. Shape dimensions are n/p/q/A. Times are median (IQR) total "
            "fit-plus-prediction seconds. Ratio is SIMPLS/PLS-SVD. RMSD is "
            "PLS-SVD/SIMPLS. Runs reports completed PLS-SVD/SIMPLS replicates."
        ),
        style=caption_style,
    )
    add_table_after(
        table_caption,
        (
            "Environment",
            "Shape n/p/q/A",
            "PLS-SVD s (IQR)",
            "SIMPLS s (IQR)",
            "Ratio",
            "RMSD P/S",
            "Runs",
        ),
        table_rows(),
        (0.75, 1.25, 1.05, 1.05, 0.55, 0.80, 0.45),
    )
    figure_paragraph = document.add_paragraph()
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.add_run().add_picture(str(FIGURE), width=Inches(6.75))
    document.add_paragraph(
        (
            "Figure S24. Matched SIMPLS/PLS-SVD total-time ratio across the five "
            "matrix-shape regimes. The ratio axis is logarithmic; one denotes "
            "parity and the dashed lines delimit +/-5%. Absolute CPU times should "
            "not be compared between Chiamaka Linux and Apple M3."
        ),
        style=caption_style,
    )
    document.core_properties.title = (
        "fastPLS CMPB supplement - matched SIMPLS versus PLS-SVD shapes"
    )
    document.save(SUPP_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    if not FIGURE.exists() or not PAIRED.exists():
        raise FileNotFoundError("Missing matched shape evidence")
    revise_main()
    revise_supplement()
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
