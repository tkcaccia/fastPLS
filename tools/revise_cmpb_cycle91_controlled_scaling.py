from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
IN_DIR = ROOT / "artifacts/CMPB_rewrite_20260824_cycle90"
OUT_DIR = ROOT / "artifacts/CMPB_rewrite_20260825_cycle91"
RESULTS = ROOT / "benchmark_results/controlled_scaling_publication_cuda_20260825"
MAIN_IN = IN_DIR / "fastPLS_CMPB_main_cycle90_0.99.25_20260824.docx"
SUPP_IN = IN_DIR / "fastPLS_CMPB_supplement_cycle90_0.99.25_20260824.docx"
MAIN_OUT = OUT_DIR / "fastPLS_CMPB_main_cycle91_0.99.25_20260825.docx"
SUPP_OUT = OUT_DIR / "fastPLS_CMPB_supplement_cycle91_0.99.25_20260825.docx"


def find_exact(document, text):
    matches = [paragraph for paragraph in document.paragraphs if paragraph.text == text]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph, found {len(matches)}: {text[:100]}")
    return matches[0]


def replace_exact(document, old, new):
    paragraph = find_exact(document, old)
    style = paragraph.style
    paragraph.text = new
    paragraph.style = style


def insert_paragraph_after(document, element, text="", style=None):
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    paragraph = Paragraph(new_p, document._body)
    if text:
        paragraph.add_run(text)
    if style is not None:
        paragraph.style = style
    return paragraph


def insert_table_after(document, element, rows, widths=None, font_size=5.0):
    table = document.add_table(rows=1, cols=len(rows[0]))
    if document.tables and document.tables[0].style is not None:
        table.style = document.tables[0].style
    table.autofit = False
    for j, value in enumerate(rows[0]):
        table.rows[0].cells[j].text = str(value)
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "true")
    header_pr.append(repeat)
    for values in rows[1:]:
        cells = table.add_row().cells
        for j, value in enumerate(values):
            cells[j].text = str(value)
    for row_i, row in enumerate(table.rows):
        for col_i, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cell.width = Inches(widths[col_i])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.space_before = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.bold = row_i == 0
    element.addnext(table._tbl)
    return table


def fmt_num(value, digits=3):
    if pd.isna(value):
        return "-"
    return f"{float(value):.{digits}g}"


OUT_DIR.mkdir(parents=True, exist_ok=True)

main = Document(MAIN_IN)

methods_anchor = find_exact(
    main,
    "Five-fold training-only selection was performed separately by PLS family; boundary selections were described as best within the grid. Accelerator speed-up was interpreted only for numerically concordant paired predictions. CUDA timing covered data transfer, execution, synchronization, and returned model or predictions; exact concordance thresholds and timing boundaries are reported in the Supplement.",
)
insert_paragraph_after(
    main,
    methods_anchor._p,
    "A controlled one-factor-at-a-time SIMPLS study varied training observations (500-10,000), predictors (50-2,000), responses (10-1,000), retained components (2-80), requested prefixes (1-20), exact cross-covariance rank (5-80), classes (5-200), and explicit cross-covariance storage (2-64 MB). Each point used three matched seeds, a deterministic float64 CPU-IRLBA reference, and rSVD with oversampling 20 and two power iterations. Fresh processes recorded fit and prediction time, baseline-corrected peak host RSS, CUDA memory, and numerical agreement; speed crossovers were inferred only when every replicate met tolerance.",
    style="Body Text",
)

results_anchor = find_exact(
    main,
    "A direct matched-shape timing study further separated estimator choice from implementation cost. On one CPU, the SIMPLS/PLS-SVD time ratio ranged from 1.00 to 3.84; on CUDA it ranged from 0.92 to 0.98 across five synthetic matrix regimes. Thus, shape-aware SIMPLS approached one-shot PLS-SVD runtime on the tested CUDA shapes, without implying that the two PLS estimators are statistically identical or that SIMPLS is universally faster (Supplementary Table S7b).",
)
insert_paragraph_after(
    main,
    results_anchor._p,
    "The controlled study completed all 486 CPU/CUDA runs. Automatic rSVD met tolerance at all sample-size, predictor-size, requested-prefix, and cross-covariance-size points, but not throughout the retained-component, class-count, rank, or CUDA response-dimension sweeps; 73 of 276 automatic-route runs were therefore excluded from speed claims. Among qualified points, CUDA first exceeded CPU rSVD at n = 5,000 and p = 2,000 and was 3.04-5.97-fold faster across the 2-64 MB cross-covariance sweep. CPU implicit products reduced incremental RSS by 29.7-47.5%, but were slower below 32 MB and approximately time-neutral at 32-64 MB. A 56-run Metal diagnostic completed without execution failure, but all 20 Metal rSVD routes were numerically discordant and were quarantined (Supplementary Tables S7c-S7d; Figure S1).",
    style="Body Text",
)

replace_exact(
    main,
    "rSVD, implicit products, float32, CUDA, and Metal are optional implementation mechanisms around shape-aware SIMPLS. Qualified NMR controls met the prespecified approximate-route tolerances, but rSVD remains stochastic and CPU IRLBA remains the deterministic reference. The million-sample ImageNet route demonstrated feasibility with current package code, while its hybrid residency, single run, noncanonical split, and lack of an estimator-matched large-scale control preclude a general accelerator or accuracy claim.",
    "rSVD, implicit products, float32, CUDA, and Metal are optional implementation mechanisms around shape-aware SIMPLS. Controlled scaling showed that neither approximation nor acceleration is uniformly preferable: route-specific numerical qualification must precede speed interpretation, and implicit CPU products primarily exchange runtime for lower memory. Qualified NMR controls met tolerance, but deterministic CPU IRLBA remains the confirmatory reference. The million-sample ImageNet route demonstrated feasibility, while its hybrid residency, single run, noncanonical split, and lack of an estimator-matched large-scale control preclude a general accelerator or accuracy claim.",
)

replace_exact(
    main,
    "Results: Deterministic fastPLS SIMPLS met the prespecified numerical tolerances in all 117 component-level comparisons. In matched single-CPU comparisons, it was faster than pls::simpls.fit on seven of nine datasets, with identical argmax accuracy and speed-up up to 8.90-fold. On NMR, qualified CPU rSVD (oversampling 20, two power iterations, seed 123) reduced 50-component SIMPLS fitting-plus-prediction time from 350.7 to 9.8 s while its predictions differed from deterministic IRLBA by relative Frobenius error 6.29e-11. A current-package hybrid float32 SIMPLS/LDA route processed one million ImageNet/DINOv2 training embeddings; at the requested 1,000-component boundary, top-1/top-5 accuracy was 0.8094/0.9393. This noncanonical single-run analysis is exploratory.",
    "Results: Deterministic fastPLS SIMPLS met the prespecified tolerances in all 117 component-level comparisons and was faster than pls::simpls.fit on seven of nine matched single-CPU datasets, with identical argmax accuracy and speed-up up to 8.90-fold. In 486 controlled CPU/CUDA runs, qualified CUDA crossovers appeared at 5,000 observations and 2,000 predictors, whereas implicit CPU products reduced incremental RSS by 29.7-47.5% but were not uniformly faster. Qualified CPU rSVD reduced 50-component NMR SIMPLS time from 350.7 to 9.8 s. A hybrid float32 SIMPLS/LDA route processed one million ImageNet/DINOv2 training embeddings; this noncanonical single-run analysis is exploratory.",
)
main.save(MAIN_OUT)

supp = Document(SUPP_IN)
replace_exact(
    supp,
    "The simulated datasets were used only for the formal SIMPLS estimator-comparison and rSVD reliability analyses; no separate n/p/q performance-scaling sweep is claimed. Five multivariate regression regimes and three dummy-response classification regimes were generated with seeds 101, 202, and 303. Regression data used independent standard-normal latent scores and Gaussian predictor and response loading matrices. Predictor and response matrices were formed by multiplying the latent-score matrix by the transposed loading matrices and adding independent Gaussian noise with standard deviation 0.05. The regimes covered p < n, p > n, low- and high-rank Y, near-collinear predictors, and exact rank deficiency. Near-collinearity was created by making two loading columns linear combinations of the first loading column plus noise with standard deviation 1 × 10⁻⁷; exact rank deficiency was created by duplicating ten predictor columns. Training and held-out rows were generated in one draw and separated before model fitting. The authoritative deterministic and approximate summaries are Tables S7 and S8; exact regime dimensions remain in the repository result archive.",
    "Two prespecified synthetic designs answered different questions. The formal estimator-comparison and rSVD reliability study used five multivariate regression and three dummy-response classification regimes with seeds 101, 202, and 303. Regression data used independent standard-normal latent scores and Gaussian predictor and response loading matrices. Predictor and response matrices were formed by multiplying the latent-score matrix by the transposed loading matrices and adding independent Gaussian noise with standard deviation 0.05. The regimes covered p < n, p > n, low- and high-rank Y, near-collinear predictors, and exact rank deficiency. Near-collinearity was created by making two loading columns linear combinations of the first loading column plus noise with standard deviation 1 × 10⁻⁷; exact rank deficiency was created by duplicating ten predictor columns. Training and held-out rows were generated in one draw and separated before model fitting. The separate controlled scaling design is reported in Section S12.2. The authoritative deterministic and approximate summaries are Tables S7 and S8; exact regime dimensions remain in the repository result archive.",
)

old_fig_caption = "Figure S1. Historical exploratory one-power rSVD workflow speed relative to deterministic IRLBA. The setting used oversampling 10, one power iteration, and seed 123; it met only 101/117 checks and is excluded from estimator-preservation and release-default claims."
replace_exact(supp, old_fig_caption, old_fig_caption.replace("Figure S1.", "Figure S2."))

anchor = find_exact(
    supp,
    "Table S7b. Direct matched-shape runtime comparison. Ratios are SIMPLS time divided by PLS-SVD time; values below one favor SIMPLS.",
)
heading = insert_paragraph_after(supp, anchor._p, "S12.2 Controlled scaling and route crossovers", style="Heading 2")
methods = insert_paragraph_after(
    supp,
    heading._p,
    "The baseline regression scenario used 2,000 training and 400 test observations, p = 400, q = 100, rank 30, and 20 retained components. One factor varied at a time: n = 500, 1,000, 2,000, 5,000, 10,000; p = 50, 100, 250, 500, 1,000, 2,000; q = 10, 25, 50, 100, 250, 500, 1,000; retained components = 2, 5, 10, 20, 40, 80; requested prefixes = 1, 2, 5, 10, 20; exact cross-covariance rank = 5, 10, 20, 40, 80; classes = 5, 10, 25, 50, 100, 200; and explicit cross-covariance storage = 2, 4, 8, 16, 32, 64 MB. The rank sweep used noise-free responses to preserve exact rank; other regression sweeps used fixed Gaussian response noise. Classification labels arose from class scores in a fixed latent space. Three matched data/solver seeds were used per point. Every scenario included deterministic float64 CPU IRLBA and automatic CPU/CUDA rSVD with oversampling 20 and two power iterations; explicit and implicit rSVD were both forced in the storage sweep.",
    style="Body Text",
)
results = insert_paragraph_after(
    supp,
    methods._p,
    "All 486 CPU/CUDA runs completed. Automatic rSVD produced 73 tolerance failures among 276 runs; failures were retained and excluded from crossover inference. All 72 forced explicit/implicit runs met tolerance. Qualified CUDA speed crossovers occurred at n = 5,000, p = 2,000, and throughout the 2-64 MB storage sweep. The CPU implicit route reduced incremental RSS by 29.7-47.5% but was slower below 32 MB and near parity thereafter. CUDA explicit/implicit times differed by less than 1% through 32 MB and by 0.1% at 64 MB. The separate Metal smoke study completed 56 runs; all 20 Metal rSVD routes were outside tolerance, so no Metal performance crossover is claimed.",
    style="Body Text",
)

overview = pd.read_csv(RESULTS / "controlled_scaling_factor_overview.csv")
table_rows = [["Factor", "Tested values", "CPU qualified", "CUDA qualified", "First qualified CUDA crossover", "Median total-time range, CPU / CUDA (s)"]]
label_map = {
    "n": "Training observations",
    "p": "Predictors",
    "q": "Responses",
    "ncomp": "Retained components",
    "prefix_count": "Requested prefixes",
    "rank": "Exact rank",
    "class_count": "Classes",
    "crosscov_mb": "Cross-covariance MB",
}
for _, row in overview.iterrows():
    crossover = fmt_num(row["first_cuda_value_qualified_and_faster"], 5)
    table_rows.append([
        label_map[row["factor_name"]],
        str(row["tested_values"]),
        f"{int(row['cpu_qualified_points'])}/{int(row['cpu_tested_points'])}",
        f"{int(row['cuda_qualified_points'])}/{int(row['cuda_tested_points'])}",
        crossover,
        f"{row['cpu_total_time_range_sec']} / {row['cuda_total_time_range_sec']}",
    ])
caption = insert_paragraph_after(supp, results._p, "Table S7c. Controlled one-factor-at-a-time scaling. A point is qualified only when every replicate met the numerical tolerances; crossover values are the first tested values, not extrapolated thresholds.", style="Caption")
table = insert_table_after(supp, caption._p, table_rows, widths=[0.95, 1.7, 0.65, 0.65, 0.85, 1.35], font_size=4.8)

cross = pd.read_csv(RESULTS / "explicit_implicit_crossover.csv")
cross_rows = [["Backend", "S size (MB)", "Explicit time (s)", "Implicit time (s)", "Implicit / explicit", "Incremental RSS explicit / implicit (MB)", "Qualified time / memory preference"]]
for _, row in cross.iterrows():
    cross_rows.append([
        row["backend"].upper(), fmt_num(row["crosscov_mb"], 4),
        fmt_num(row["explicit_total_sec"]), fmt_num(row["implicit_total_sec"]),
        fmt_num(row["implicit_over_explicit_time"]),
        f"{fmt_num(row['explicit_incremental_rss_mb'], 4)} / {fmt_num(row['implicit_incremental_rss_mb'], 4)}",
        f"{row['qualified_time_preference']} / {row['qualified_memory_preference']}",
    ])
cross_caption = insert_paragraph_after(supp, table._tbl, "Table S7d. Forced explicit versus implicit cross-covariance routes. Incremental RSS excludes the loaded synthetic input but includes fit and prediction workspaces.", style="Caption")
cross_table = insert_table_after(supp, cross_caption._p, cross_rows, widths=[0.55, 0.75, 0.75, 0.75, 0.75, 1.25, 1.6], font_size=4.7)

image_paragraph = insert_paragraph_after(supp, cross_table._tbl)
image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
image_paragraph.add_run().add_picture(str(RESULTS / "controlled_scaling_overview.png"), width=Inches(6.5))
insert_paragraph_after(
    supp,
    image_paragraph._p,
    "Figure S1. Controlled SIMPLS scaling. Median fit-plus-prediction time, baseline-corrected host RSS, numerical error relative to deterministic CPU IRLBA, and the forced implicit/explicit crossover are shown for the tested grid. Timing interpretation excludes routes outside the prespecified numerical tolerances.",
    style="Caption",
)

claim_table = supp.tables[7]
cells = claim_table.add_row().cells
for cell, value in zip(cells, [
    "Controlled shape-dependent routing",
    "Tables S7c-S7d; Figure S1",
    "46 scenarios, 486 CPU/CUDA runs, three seeds; separate 56-run Metal diagnostic; time, prediction, incremental RSS, GPU memory, and numerical qualification",
]):
    cell.text = value
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(5)

ledger = supp.tables[18]
cells = ledger.add_row().cells
for cell, value in zip(cells, [
    "A21",
    "Controlled one-factor SIMPLS scaling",
    "benchmark_results/controlled_scaling_publication_cuda_20260825",
    "0.99.25",
    "exact source archive; isolated CPU/CUDA processes; Metal diagnostic on reviewed Mac build",
    "scripts/run_controlled_scaling.sh; benchmark/controlled_scaling/*.R",
    "74e134ef22d5",
]):
    cell.text = value
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(5)

supp.save(SUPP_OUT)
print(MAIN_OUT)
print(SUPP_OUT)
