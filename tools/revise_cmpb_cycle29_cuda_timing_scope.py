#!/usr/bin/env python3

from pathlib import Path

from docx import Document


ROOT = Path("/Users/stefano/Documents/GPUPLS/fastPLS_fresh_github")
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle28"
OUT = ROOT / "artifacts" / "CMPB_rewrite_20260725_cycle29"

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle28_0.99.6_20260725.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle28_0.99.6_20260725.docx"
RESPONSE_SOURCE = (
    SOURCE / "fastPLS_CMPB_response_to_reviewers_cycle28_20260725.docx"
)

MAIN_OUT = OUT / "fastPLS_CMPB_main_cycle29_0.99.6_20260725.docx"
SUPP_OUT = OUT / "fastPLS_CMPB_supplement_cycle29_0.99.6_20260725.docx"
RESPONSE_OUT = OUT / "fastPLS_CMPB_response_to_reviewers_cycle29_20260725.docx"


def find_paragraph(document, prefix):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def revise_main():
    document = Document(MAIN_SOURCE)
    paragraph = find_paragraph(document, "Within each dataset, methods used")
    old = "Runtime comprised fitting and prediction. "
    new = (
        "Runtime comprised fitting and prediction. CUDA times were measured "
        "around the complete public fit and prediction calls rather than the "
        "device kernels alone. Fit time therefore included host-side argument "
        "marshalling, transfer of training data to the device, CUDA execution "
        "and synchronization, and transfer of fitted model components returned "
        "to R. Prediction time included transfer of the test matrix and any "
        "nonresident model arrays to the device, CUDA execution and "
        "synchronization, and transfer of predictions or class scores back to "
        "host memory. Resident workspaces were reused where supported. Dataset "
        "loading, package initialization, and metric calculation after "
        "prediction were excluded. "
    )
    if old not in paragraph.text:
        raise RuntimeError("Runtime sentence not found in main Methods")
    paragraph.text = paragraph.text.replace(old, new, 1)
    document.core_properties.title = (
        "fastPLS CMPB manuscript - explicit CUDA timing boundary"
    )
    document.save(MAIN_OUT)


def revise_supplement():
    document = Document(SUPP_SOURCE)
    paragraph = find_paragraph(document, "Each fit is run in an isolated R process.")
    paragraph.text = (
        "Each fit is run in an isolated R process. Total time is fitting plus "
        "prediction and excludes package installation, dataset acquisition and "
        "loading, package initialization, and metric calculation after "
        "prediction. The fit timer starts immediately before the public pls() "
        "call and stops after the fitted R model is returned. For CUDA routes, "
        "this interval includes host-side argument marshalling, host-to-device "
        "transfer of training matrices, device allocation, CUDA kernels and "
        "library calls, synchronization, and device-to-host transfer of model "
        "components exposed in the returned R object. The prediction timer "
        "starts immediately before predict() and stops after predictions or "
        "class scores are available in R. It includes host-to-device transfer "
        "of the test matrix and any model arrays not retained on the device, "
        "CUDA execution and synchronization, and device-to-host result "
        "transfer. Routes supporting resident workspaces reuse them; the "
        "benchmark does not add an artificial host round trip, but the public "
        "R model remains the interface between the separately timed fit and "
        "prediction calls. Therefore, reported CUDA time is end-to-end backend "
        "time rather than kernel-only time. After data and libraries are loaded, "
        "garbage collection is completed and host RSS and process-specific GPU "
        "memory are recorded immediately before fitting. The R process then "
        "signals a synchronized external sampler, which polls host RSS and "
        "PID-matched nvidia-smi device memory every 0.02 s throughout fitting "
        "and prediction. /usr/bin/time records the absolute maximum RSS over the "
        "complete isolated process. Incremental host memory is the maximum of "
        "fit-window samples and post-fit/post-prediction snapshots minus the "
        "pre-fit baseline; incremental GPU memory is the PID-specific sampled "
        "peak minus its pre-fit baseline. Negative differences are truncated at "
        "zero. Thus absolute process RSS measures whole-workflow feasibility, "
        "whereas the baseline-corrected increment estimates fitting and "
        "prediction workspace."
    )
    document.core_properties.title = (
        "fastPLS CMPB supplement - explicit CUDA timing boundary"
    )
    document.save(SUPP_OUT)


def revise_response():
    document = Document(RESPONSE_SOURCE)
    document.add_heading("29. CUDA timing boundary was not explicit", level=1)
    document.add_paragraph(
        "Reviewer comment: Define whether CUDA timing includes host-to-device "
        "transfer, model transfer, synchronization, and result transfer."
    )
    document.add_paragraph(
        "Response: Corrected. The Methods and Supplementary Methods now define "
        "the timers around the complete public fit and prediction calls. CUDA "
        "fit time includes argument marshalling, training-data transfer, device "
        "allocation and execution, synchronization, and transfer of returned "
        "model components to R. Prediction time includes test-data and required "
        "nonresident model transfer, synchronized device execution, and return "
        "of predictions or class scores to host memory. Resident workspaces are "
        "reused where supported. The reported value is therefore end-to-end "
        "backend time, not kernel-only time; loading, initialization, and "
        "post-prediction metric calculation are excluded."
    )
    document.core_properties.title = (
        "fastPLS CMPB response to reviewers - CUDA timing boundary"
    )
    document.save(RESPONSE_OUT)


def main():
    OUT.mkdir(parents=True, exist_ok=True)
    revise_main()
    revise_supplement()
    revise_response()


if __name__ == "__main__":
    main()
