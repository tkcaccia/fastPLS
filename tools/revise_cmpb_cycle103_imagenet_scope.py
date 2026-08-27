from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle102"
OUTPUT = ROOT / "artifacts" / "CMPB_rewrite_20260825_cycle103"
OUTPUT.mkdir(parents=True, exist_ok=True)

MAIN_SOURCE = SOURCE / "fastPLS_CMPB_main_cycle102_0.99.25_20260825.docx"
SUPP_SOURCE = SOURCE / "fastPLS_CMPB_supplement_cycle102_0.99.25_20260825.docx"
MAIN_OUTPUT = OUTPUT / "fastPLS_CMPB_main_cycle103_0.99.25_20260825.docx"
SUPP_OUTPUT = OUTPUT / "fastPLS_CMPB_supplement_cycle103_0.99.25_20260825.docx"
IMAGENET_FIGURE = (
    ROOT / "benchmark_results" / "frozen_release_0.99.25" / "figures" /
    "Figure_5_frozen_imagenet.png"
)


def paragraph_by_prefix(document, prefix):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}; found {len(matches)}")
    return matches[0]


def remove_paragraph(paragraph):
    paragraph._element.getparent().remove(paragraph._element)


def set_cell_text(cell, value, size=6.5):
    cell.text = str(value)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(size)


def remove_preceding_drawing(caption, limit=8):
    previous = caption._p.getprevious()
    checked = 0
    while previous is not None and checked < limit:
        prior = previous.getprevious()
        if previous.xpath(".//w:drawing") or previous.xpath(".//w:pict"):
            previous.getparent().remove(previous)
            return
        if "".join(previous.itertext()).strip():
            break
        previous = prior
        checked += 1
    raise RuntimeError("Could not locate the figure preceding the caption")


def insert_figure_before(document, target, image_path, caption_text, width=6.7):
    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_with_next = True
    image_paragraph.paragraph_format.space_after = Pt(3)
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    target._p.addprevious(image_paragraph._p)

    caption_paragraph = document.add_paragraph()
    caption_paragraph.style = target.style
    caption_paragraph.paragraph_format.keep_together = True
    caption_paragraph.paragraph_format.keep_with_next = True
    run = caption_paragraph.add_run(caption_text)
    run.italic = True
    target._p.addprevious(caption_paragraph._p)


def update_main():
    document = Document(MAIN_SOURCE)

    introduction = paragraph_by_prefix(document, "Increasing spectral resolution")
    introduction.text = (
        "Increasing spectral resolution, single-cell sample counts, foundation-model embeddings, "
        "and repeated validation make PLS computationally demanding. KODAMA, for example, "
        "repeatedly fits cross-validated PLS discriminant models [8,9]. In medical imaging, "
        "foundation models such as UNI and Prov-GigaPath generate large dense tile or slide "
        "representations for downstream supervised analysis [31,32]. Historical ImageNet/DINOv2 "
        "embeddings were therefore retained only as a partially reproducible engineering proxy for "
        "this post-extraction matrix regime, not as biomedical validation [28,29]."
    )

    benchmark = paragraph_by_prefix(document, "The broader benchmark design covered")
    benchmark.text = (
        "The broader benchmark design covered biomedical and computational tasks spanning "
        "metabolomics, NMR, CITE-seq, tissue and cancer omics, single-cell transcriptomics, drug "
        "response, and image embeddings [7,20-30]. Frozen central evidence comprises deterministic "
        "SIMPLS validation, external comparison, controlled scaling and solver qualification, "
        "selected CPU/CUDA routes, and the NMR case study. The historical ImageNet representation "
        "is reported separately as an engineering feasibility demonstration. Methods used identical "
        "stored splits and training-only component grids within each analysis. Runtime included "
        "fitting and prediction; memory definitions and uncertainty procedures are detailed in the "
        "Supplement."
    )

    exploratory = paragraph_by_prefix(document, "A separate exploratory ImageNet/DINOv2 analysis")
    exploratory.text = (
        "A separate historical ImageNet/DINOv2 analysis used 1,281,167 stored 1,024-dimensional "
        "embeddings. The exact DINOv2 checkpoint, pooling rule, extraction script, and auditable "
        "image-to-row mapping were not retained; the 1,000,000/281,167 split was noncanonical and "
        "had informed earlier component choices. Only the downstream PLS fitting and prediction "
        "stages were rerun with frozen fastPLS 0.99.25. Consequently, this analysis is partially "
        "reproducible and supports only matrix-processing feasibility, not representation-level, "
        "biomedical, or comparative predictive claims."
    )

    organization = paragraph_by_prefix(document, "Results are organized around")
    organization.text = (
        "Results are organized around the shape-dependent SIMPLS execution contribution. Primary "
        "evidence comprises exact and deterministic estimator comparison, independent software "
        "comparison, and the current-version biomedical NMR case study. rSVD, precision, and "
        "hardware results are reported separately with their numerical-audit status. The historical "
        "ImageNet analysis is retained as a supplementary feasibility demonstration because its "
        "representation provenance is incomplete."
    )

    heading = paragraph_by_prefix(document, "3.5 ImageNet-scale supervised representation")
    heading.text = "3.5 Foundation-model embedding feasibility"

    first = paragraph_by_prefix(document, "The exploratory ImageNet classification experiment")
    first.text = (
        "Foundation-model embeddings are increasingly relevant to medical-image pipelines because "
        "the downstream analysis may involve millions of dense tile representations. As a historical "
        "engineering example, frozen fastPLS 0.99.25 completed label-aware float32 SIMPLS fitting on "
        "1,000,000 stored ImageNet/DINOv2 rows and blocked prediction of 281,167 held-out rows. "
        "This demonstrates feasibility for a foundation-model-scale matrix, but not reproducible "
        "feature extraction or biomedical utility. Detailed single-run classification, timing, "
        "memory, and provenance limitations are reported in Supplementary Section S18, Table S13, "
        "and Figure S3."
    )

    remove_paragraph(paragraph_by_prefix(document, "One shared fit supplied 100-1,000-component prefixes"))
    figure_caption = paragraph_by_prefix(document, "Figure 5.")
    remove_preceding_drawing(figure_caption)
    remove_paragraph(figure_caption)

    availability = paragraph_by_prefix(document, "Code and benchmark outputs are available")
    availability.text = (
        "Code and benchmark outputs are available at https://github.com/tkcaccia/fastPLS; reusable "
        "components are at https://github.com/tkcaccia/kodama-cpp. Central current-version evidence "
        "was generated with fastPLS 0.99.25, Git commit "
        "7887401b09e25f54a546a253c255741cb1ab48e5, from source archive fastPLS_0.99.25.tar.gz "
        "(SHA-256 604f74d72f5e9540f7378efd37f2ca6ed87ccb9e73b912211331a8cd97233481). "
        "The ImageNet PLS fitting and prediction stages were rerun from that archive, but the older "
        "DINOv2 feature-extraction checkpoint, pooling rule, script, and image-to-row mapping are not "
        "fully recoverable; those results are therefore historical and partially reproducible. The "
        "deposited 165-component NMR workflow is separately labelled historical context. A persistent "
        "archive identifier will replace the checksum-only review object before acceptance."
    )

    document.save(MAIN_OUTPUT)


def update_supplement():
    document = Document(SUPP_SOURCE)

    source = paragraph_by_prefix(document, "CIFAR-100 and ImageNet were evaluated")
    source.text = (
        "CIFAR-100 and ImageNet were evaluated as precomputed image embeddings rather than by "
        "training image encoders. CIFAR-100 used its standard 50,000/10,000 partition [30]. The "
        "historical ImageNet archive contained 1,281,167 rows, 1,024 DINOv2 features, and 1,000 "
        "labels [28,29], but no authoritative canonical train/validation flag. Its exact DINOv2 "
        "checkpoint, image preprocessing, pooling rule, extraction script, and independently "
        "auditable image-to-row mapping were not retained. The stored feature matrix and labels can "
        "be checksummed and the downstream fastPLS analysis rerun, but the representation cannot be "
        "regenerated independently from the available metadata. Feature extraction was completed "
        "before the fastPLS benchmark and excluded from timing."
    )

    role = paragraph_by_prefix(document, "ImageNet is included as a large-scale foundation-model")
    role.text = (
        "ImageNet is retained as a historical, partially reproducible foundation-model embedding "
        "stress test, not as a biomedical validation set. Its relevance is the post-extraction matrix "
        "regime: computational pathology models such as UNI and Prov-GigaPath produce large dense "
        "tile or slide representations for supervised downstream analysis [31,32]. The experiment "
        "tests whether fastPLS can process a matrix of that scale. It does not establish natural-to-"
        "medical-image transfer, diagnostic performance, or representation-level reproducibility."
    )

    central = paragraph_by_prefix(document, "Every central benchmark used the frozen")
    central.text = (
        "Current-version central analyses used the frozen 0.99.25 execution archive and recorded its "
        "SHA-256, package version, input checksum, script identity, compiler, BLAS/LAPACK, thread "
        "settings, accelerator libraries, seed, rSVD controls, repetition, and status. NMR was rerun "
        "end to end from the prepared predictors and responses. For ImageNet, only the downstream "
        "PLS fitting and prediction stages were rerun; feature extraction remains historical and "
        "partially reproducible. Table S15 maps each claim to its result directory and provenance scope."
    )

    scientific_role = paragraph_by_prefix(document, "The scientific role of this experiment is computational")
    scientific_role.text = (
        "The role of this experiment is computational rather than biomedical. It supplies a large "
        "multiclass embedding matrix resembling the post-extraction data form produced by medical-"
        "image foundation models [31,32]. Because representation provenance is incomplete, it is "
        "interpreted only as a historical feasibility and memory-throughput demonstration."
    )

    section = paragraph_by_prefix(document, "S18. ImageNet exploratory stress test")
    section.text = "S18. Historical ImageNet foundation-model embedding stress test"

    details = paragraph_by_prefix(document, "The pooled archive contained 1,281,167")
    details.text = (
        "The historical archive contained 1,281,167 precomputed DINOv2 embeddings with 1,024 "
        "features and 1,000 class labels. Seed 123 assigned 1,000,000 rows to training and 281,167 "
        "to a noncanonical holdout that had informed earlier component choices. The exact checkpoint, "
        "pooling rule, extraction script, and auditable image-to-row mapping were unavailable. Frozen "
        "fastPLS 0.99.25 reran only downstream label-aware float32 SIMPLS fitting and blocked "
        "prediction using CUDA rSVD oversampling 20, two power iterations, and seed 123. All values "
        "are single-run historical feasibility estimates; they are not used for biomedical, "
        "representation-quality, or comparative predictive claims."
    )

    table_caption = paragraph_by_prefix(document, "Table S13.")
    insert_figure_before(
        document,
        table_caption,
        IMAGENET_FIGURE,
        "Figure S3. Historical, partially reproducible ImageNet/DINOv2 downstream SIMPLS "
        "feasibility experiment. The frozen 0.99.25 package reran PLS fitting and blocked prediction "
        "on the stored embedding matrix; DINOv2 checkpoint, pooling, extraction, and complete "
        "image-to-row provenance were unavailable. Panels report single-run exploratory accuracy, "
        "runtime, and memory across 100-1,000 component prefixes. The experiment is a proxy for the "
        "post-extraction scale encountered with medical-image foundation models, not biomedical "
        "validation.",
    )
    table_caption.text = (
        "Table S13. Historical frozen-release ImageNet downstream classification. One shared fit "
        "supplies all component prefixes for each head; fit time is repeated only to expose the common "
        "cost. The 1,000-component row is a boundary stress point, not a selected optimum. All rows "
        "are single-run, partially reproducible feasibility estimates."
    )

    evidence_map = document.tables[8]
    for row in evidence_map.rows[1:]:
        if row.cells[0].text.startswith("ImageNet feasibility"):
            set_cell_text(row.cells[0], "Historical ImageNet downstream matrix-processing feasibility")
            set_cell_text(row.cells[2], (
                "Frozen 0.99.25 PLS rerun on stored embeddings; feature-extraction provenance incomplete"
            ))
            break

    provenance = document.tables[26]
    for row in provenance.rows[1:]:
        if row.cells[0].text == "F09":
            set_cell_text(row.cells[1], "ImageNet downstream PLS fitting and prediction")
            set_cell_text(row.cells[4], (
                "7887401b09e2 downstream rerun; historical embedding extraction incomplete"
            ))
            break

    document.save(SUPP_OUTPUT)


if __name__ == "__main__":
    if not IMAGENET_FIGURE.exists():
        raise FileNotFoundError(IMAGENET_FIGURE)
    update_main()
    update_supplement()
    print(MAIN_OUTPUT)
    print(SUPP_OUTPUT)
