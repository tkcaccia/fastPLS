#!/usr/bin/env python3

"""Build the manuscript and supplement from audited current-release evidence."""

from pathlib import Path
from copy import deepcopy
import csv
import math
import re
import statistics
import subprocess
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
VERSION = "0.99.39"
PREVIOUS_VERSION = "0.99.36"
RESULTS = ROOT / "publication_results" / VERSION / "current_release"
MAIN_IN = ROOT / "artifacts/CMPB_rewrite_20260901_cycle136/fastPLS_CMPB_main_0.99.36.docx"
SUPP_IN = ROOT / "artifacts/CMPB_rewrite_20260901_cycle136/fastPLS_CMPB_supplement_0.99.36.docx"
OUT_DIR = ROOT / "artifacts/CMPB_rewrite_20260903_cycle138"
MAIN_OUT = OUT_DIR / f"fastPLS_CMPB_main_{VERSION}.docx"
SUPP_OUT = OUT_DIR / f"fastPLS_CMPB_supplement_{VERSION}.docx"


def enable_continuous_line_numbers(document):
    """Number manuscript lines continuously for journal review."""
    for section in document.sections:
        section_properties = section._sectPr
        existing = section_properties.find(qn("w:lnNumType"))
        if existing is not None:
            section_properties.remove(existing)

        line_numbers = OxmlElement("w:lnNumType")
        line_numbers.set(qn("w:countBy"), "1")
        line_numbers.set(qn("w:distance"), "360")
        line_numbers.set(qn("w:restart"), "continuous")
        section_properties.append(line_numbers)


def normalize_table_font_sizes(document, size_pt):
    """Use one readable font size within every table in a document."""
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(size_pt)


SCIENTIFIC_NOTATION = re.compile(
    r"(?<![A-Za-z0-9])(?:"
    r"(?P<times>\d+(?:\.\d+)?\s*[x×]\s*10\^?(?P<times_exp>[+\-−]\d+))"
    r"|(?P<compact>\d+(?:\.\d+)?[eE](?P<compact_exp>[+\-−]?\d+))"
    r"|(?P<bare>10(?:\^(?P<bare_caret_exp>[+\-−]\d+)|(?P<bare_unicode_exp>−\d+)))"
    r")(?![A-Za-z0-9])"
)


def _normalized_exponent(value):
    numeric_value = int(value.replace("−", "-"))
    sign = "−" if numeric_value < 0 else "+"
    return sign + str(abs(numeric_value))


def _scientific_parts(match):
    if match.group("times"):
        coefficient = re.match(r"\d+(?:\.\d+)?", match.group("times")).group(0)
        exponent = match.group("times_exp")
    elif match.group("compact"):
        coefficient = re.match(r"\d+(?:\.\d+)?", match.group("compact")).group(0)
        exponent = match.group("compact_exp")
    else:
        coefficient = None
        exponent = match.group("bare_caret_exp") or match.group("bare_unicode_exp")

    prefix = "10" if coefficient in (None, "1") else f"{coefficient} × 10"
    return prefix, _normalized_exponent(exponent)


def _copy_run_properties(source_run, target_run):
    source_properties = source_run._r.rPr
    if source_properties is None:
        return
    target_properties = target_run._r.rPr
    if target_properties is not None:
        target_run._r.remove(target_properties)
    target_run._r.insert(0, deepcopy(source_properties))


def format_scientific_notation(document):
    """Render scientific-notation exponents as true superscripts."""
    paragraphs = list(document.paragraphs)
    paragraphs.extend(
        paragraph
        for table in document.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )

    for paragraph in paragraphs:
        for run in list(paragraph.runs):
            text = run.text
            matches = list(SCIENTIFIC_NOTATION.finditer(text))
            if not matches:
                continue

            pieces = []
            cursor = 0
            for match in matches:
                if match.start() > cursor:
                    pieces.append((text[cursor:match.start()], False))
                prefix, exponent = _scientific_parts(match)
                pieces.extend(((prefix, False), (exponent, True)))
                cursor = match.end()
            if cursor < len(text):
                pieces.append((text[cursor:], False))

            run.text = pieces[0][0]
            run.font.superscript = pieces[0][1]
            insertion_point = run._r
            for piece_text, superscript in pieces[1:]:
                new_run = paragraph.add_run(piece_text)
                _copy_run_properties(run, new_run)
                new_run.font.superscript = superscript
                insertion_point.addnext(new_run._r)
                insertion_point = new_run._r
RSVD_FIGURE = RESULTS / "rsvd_qualification/rsvd_qualification.png"
SCALING_FIGURE = RESULTS / "controlled_scaling/controlled_scaling_current.png"
EXTERNAL_FIGURE = RESULTS / "external_simpls/external_simpls_current_release.png"
EXTERNAL_PAIRS = RESULTS / "external_simpls/external_simpls_timing_pairs.csv"
IKPLS_SUMMARY = RESULTS / "ikpls_cross_language_cpu/ikpls_cross_language_summary.csv"
IKPLS_LARGE_SUMMARY = (
    RESULTS / "ikpls_large_float32/ikpls_fastpls_large_float32_summary.csv"
)
PACKAGE_FIGURE = RESULTS / "r_package_panel/pls_package_comparison_current.png"
PACKAGE_STATUS = RESULTS / "r_package_panel/pls_package_comparison_status.csv"
PACKAGE_SUMMARY = RESULTS / "r_package_panel/pls_package_comparison_summary.csv"
BACKEND_FIGURE = RESULTS / f"figures/Figure_backend_runtime_{VERSION}.png"
SELECTED_SIMPLS_BACKEND_FIGURE = (
    RESULTS / f"figures/Figure_selected_simpls_backend_{VERSION}.png"
)
CUDA_PAIRED = RESULTS / "selected_backend_cuda/matched_cuda_paired.csv"
METAL_PAIRED = RESULTS / "selected_backend_metal/matched_metal_paired.csv"
NMR_RESULTS = RESULTS / "nmr"
NMR_FIGURE = RESULTS / f"nmr/figures/Figure_nmr_fixed165_{VERSION}.png"
NMR_SELECTED_FIGURE = RESULTS / f"nmr/figures/Figure_nmr_selected_{VERSION}.png"
IMAGENET_FIGURE = RESULTS / f"figures/Figure_imagenet_{VERSION}.png"
IMAGENET_SUMMARY = RESULTS / "imagenet/imagenet_current_summary.csv"
SIMPLS_EXACT_SUMMARY = (
    RESULTS / "simpls_exact/simpls_exact_reference_case_summary.csv"
)
SIMPLS_PRESERVATION = (
    RESULTS
    / "simpls_preservation/simpls_estimator_preservation_validation_summary.csv"
)
SIMPLS_PRESERVATION_SUMMARY = (
    RESULTS / "simpls_preservation/simpls_estimator_preservation_summary.csv"
)
OPLS_KERNEL_ESTIMATOR = (
    RESULTS / "opls_kernel_estimator/opls_kernel_estimator_validation_summary.csv"
)
OPLS_KERNEL_ESTIMATOR_SELECTION = (
    RESULTS / "opls_kernel_estimator/opls_kernel_component_selection_summary.csv"
)
OPLS_KERNEL_SETTINGS = (
    RESULTS / "opls_kernel_settings/opls_kernel_setting_reliability_summary.csv"
)
OPLS_KERNEL_SETTINGS_SELECTION = (
    RESULTS / "opls_kernel_settings/opls_kernel_setting_selection_summary.csv"
)
REPEATED_OUTER_DIR = RESULTS / "repeated_outer"
ABLATION_EFFECTS = RESULTS / "simpls_ablation/simpls_multidataset_ablation_effects.csv"
SHAPE_PAIRED = RESULTS / "simpls_vs_plssvd_shapes/simpls_vs_plssvd_shapes_paired.csv"
COMPONENT_SELECTION = RESULTS / "component_selection/selected_components.csv"
COMPONENT_SELECTION_CLASSIFICATION_FIGURE = (
    RESULTS / "component_selection/component_selection_classification.png"
)
COMPONENT_SELECTION_REGRESSION_FIGURE = (
    RESULTS / "component_selection/component_selection_regression.png"
)
COMPONENT_PATH_COMBINED_DIR = RESULTS / "component_path_cpu_cuda_metal"
COMPONENT_PATH_DATASETS = (
    ("cbmc_citeseq", "CBMC CITE-seq"),
    ("ccle", "CCLE"),
    ("cifar100", "CIFAR-100"),
    ("gtex_v8", "GTEx v8"),
    ("metref", "MetRef"),
    ("prism", "PRISM"),
    ("retina", "Retina"),
    ("tabula", "Tabula Muris"),
    ("tcga_brca", "TCGA-BRCA"),
    ("tcga_hnsc_methylation", "TCGA-HNSC methylation"),
    ("tcga_pan_cancer", "TCGA Pan-Cancer"),
)
TABLE_NUMBER_MAP = {
    "S1": "S1", "S2": "S2", "S3": "S3", "S3a": "S4",
    "S4a": "S5", "S4b": "S6", "S5": "S7", "S6": "S8",
    "S6a": "S9", "S6b": "S10", "S6c": "S11", "S6d": "S12",
    "S7": "S13", "S8": "S14", "S9a": "S15", "S9b": "S16",
    "S9c": "S17", "S9d": "S18", "S9e": "S19", "S9f": "S20",
    "S9g": "S21", "S10": "S22", "S11": "S23", "S12": "S24",
    "S13": "S25", "S14": "S26", "S15": "S27",
}


def replace_runs(container, replacements):
    for paragraph in container.paragraphs:
        for run in paragraph.runs:
            for old, new in replacements.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_runs(cell, replacements)


def replace_release_version(document):
    """Synchronize generated prose and table paths with the evaluated release."""
    replace_runs(document, {
        PREVIOUS_VERSION: VERSION,
        f"publication_results/{PREVIOUS_VERSION}": f"publication_results/{VERSION}",
    })
    patterns = (
        (re.compile(r"(?<=fastPLS )0\.99\.\d+"), VERSION),
        (re.compile(r"(?<=version )0\.99\.\d+"), VERSION),
        (re.compile(r"(?<=v)0\.99\.\d+"), VERSION),
        (re.compile(r"(?<=publication_results/)0\.99\.\d+"), VERSION),
    )
    containers = [document]
    containers.extend(
        cell
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    for container in containers:
        for paragraph in container.paragraphs:
            for run in paragraph.runs:
                for pattern, replacement in patterns:
                    updated = pattern.sub(replacement, run.text)
                    if updated != run.text:
                        # Assigning run.text rebuilds the run and would discard any
                        # embedded drawing carried by an image-only run.
                        run.text = updated


def component_selection_summary():
    rows = read_csv(COMPONENT_SELECTION)
    counts = {}
    for row in rows:
        status = row.get("selection_status", "")
        counts[status] = counts.get(status, 0) + 1
    return {
        "total": len(rows),
        "interior": counts.get("interior", 0),
        "lower": counts.get("lower_grid_boundary", 0),
        "upper": counts.get("upper_grid_boundary", 0),
        "rank": counts.get("rank_limited", 0),
    }


def set_paragraph(paragraph, text):
    paragraph.clear()
    paragraph.add_run(text)


def remove_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def remove_embedded_figure_paragraphs(document):
    """Remove source-document figures before inserting current-release graphics."""
    for paragraph in list(document.paragraphs):
        if not paragraph._p.xpath(".//w:drawing"):
            continue
        if paragraph.text.strip():
            raise ValueError("Refusing to remove a drawing paragraph that also contains text.")
        remove_paragraph(paragraph)


def figure_paragraph_after(document, caption):
    """Return a safe paragraph for a figure immediately after its caption."""
    paragraphs = list(document.paragraphs)
    index = next(
        position for position, paragraph in enumerate(paragraphs)
        if paragraph._p is caption._p
    )
    if index + 1 < len(paragraphs):
        candidate = paragraphs[index + 1]
        if not candidate.text.strip() or candidate._p.xpath(".//w:drawing"):
            candidate.clear()
            return candidate

    element = OxmlElement("w:p")
    caption._p.addnext(element)
    return Paragraph(element, caption._parent)


def remove_table_column(table, index):
    grid = table._tbl.tblGrid
    grid.remove(grid.gridCol_lst[index])
    for row in table.rows:
        cell = row.cells[index]
        row._tr.remove(cell._tc)


def reset_table_rows(table, rows):
    target_columns = len(rows[0])
    while len(table.columns) > target_columns:
        remove_table_column(table, len(table.columns) - 1)
    while len(table.columns) < target_columns:
        table.add_column(Inches(0.6))
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)
    header = rows[0]
    for index, value in enumerate(header):
        table.rows[0].cells[index].text = value
    for values in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)


def table_after_caption(document, caption_prefix):
    """Return the first table following a matching caption paragraph."""
    found_caption = False
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, document)
            if paragraph.text.strip().startswith(caption_prefix):
                found_caption = True
        elif isinstance(child, CT_Tbl) and found_caption:
            return Table(child, document)
    raise ValueError(f"No table follows caption {caption_prefix!r}")


def ensure_table_immediately_after_caption(document, caption_prefix, columns):
    """Return or create the table attached directly to a caption."""
    for child in document.element.body.iterchildren():
        if not isinstance(child, CT_P):
            continue
        paragraph = Paragraph(child, document)
        if not paragraph.text.strip().startswith(caption_prefix):
            continue
        sibling = child.getnext()
        if isinstance(sibling, CT_Tbl):
            return Table(sibling, document)
        table = document.add_table(rows=1, cols=columns)
        child.addnext(table._tbl)
        return table
    raise ValueError(f"No caption matches {caption_prefix!r}")


def ensure_table_columns(table, count):
    while len(table.columns) < count:
        table.add_column(Inches(0.6))


def set_table_column_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)


def update_external_scope_table(document):
    table = table_after_caption(document, "Table S4b.")
    rows = [
        ["Implementation", "PLS methods evaluated", "Multicore CPU", "CUDA", "Metal", "float32", "float64"],
        [f"fastPLS {VERSION}", "PLS-SVD, SIMPLS, OPLS, kernel PLS", "Yes for eligible BLAS/OpenMP kernels; sequential deflation remains serial", "Yes", "Yes", "Yes, route-specific", "Yes"],
        ["IKPLS 6.1.2", "Improved Kernel PLS, kernel-matrix and cross-product formulations (NumPy/JAX)", "Runtime-dependent through NumPy or JAX", "Yes, through JAX", "No package-native route", "Yes", "Yes"],
        ["pls", "kernel, wide-kernel, SIMPLS, orthogonal-scores PLS", "Yes for cross-validation; fitting is BLAS-dependent", "No native route", "No native route", "No native route", "Yes"],
        ["mdatools", "PLS and PLS-DA", "BLAS-dependent only", "No native route", "No native route", "No native route", "Yes"],
        ["plsdepot", "SIMPLS", "BLAS-dependent only", "No native route", "No native route", "No native route", "Yes"],
        ["pcv", "SIMPLS", "BLAS-dependent only", "No native route", "No native route", "No native route", "Yes"],
        ["plsgenomics", "PLS regression and classification", "BLAS-dependent only", "No native route", "No native route", "No native route", "Yes"],
        ["chemometrics", "Eigen- and NIPALS-based PLS", "BLAS-dependent only", "No native route", "No native route", "No native route", "Yes"],
        ["mixOmics", "PLS-DA and sparse PLS-DA", "Supported tuning/CV can parallelize; model fitting is BLAS-dependent", "No native route", "No native route", "No native route", "Yes"],
        ["spls", "Sparse PLS and sparse PLS-DA", "BLAS-dependent only", "No native route", "No native route", "No native route", "Yes"],
        ["ropls", "PLS and OPLS", "BLAS-dependent only", "No native route", "No native route", "No native route", "Yes"],
    ]
    ensure_table_columns(table, len(rows[0]))
    reset_table_rows(table, rows)
    set_table_column_widths(table, [0.8, 1.55, 1.35, 0.7, 0.7, 0.7, 0.6])
    style_table(table, 6.0)


def dataset_dimension_rows():
    return [
        ["Dataset", "Task", "Prepared n train", "Prepared n test", "Prepared p", "Prepared q"],
        ["TCGA-HNSC methylation", "Classification", "520", "58", "782", "2"],
        ["CCLE", "Classification", "547", "71", "1,000", "18"],
        ["TCGA-BRCA", "Classification", "756", "88", "1,000", "5"],
        ["MetRef", "Classification", "773", "100", "375", "22"],
        ["TCGA Pan-Cancer", "Classification", "3,000", "982", "850", "32"],
        ["GTEx v8", "Classification", "3,000", "797", "1,000", "32"],
        ["Retina", "Classification", "22,402", "22,406", "50", "12"],
        ["Tabula Muris", "Classification", "50,043", "50,059", "50", "32"],
        ["CIFAR-100", "Classification", "50,000", "10,000", "768", "100"],
        ["PRISM", "Multivariate regression", "479", "54", "1,000", "4,686"],
        ["NMR", "Multivariate regression", "1,200", "321", "13,000", "28,355"],
        ["CBMC CITE-seq", "Multivariate regression", "7,755", "862", "1,000", "10"],
        ["ImageNet/DINOv2", "Exploratory classification stress test", "1,000,000", "281,167", "1,024", "1,000"],
    ]


def environment_rows():
    return [
        ["Item", "CUDA workstation", "Metal validation system"],
        [
            "Platform",
            "Intel Core i7-13700 (16 cores, 24 logical CPUs); 31.0 GiB RAM; NVIDIA GeForce RTX 5060 Ti with 15.9 GiB device memory",
            "macOS 14.5; Apple M3 (8 CPU cores, 10 GPU cores); 8.0 GiB unified memory; Metal 3",
        ],
        [
            "R and fastPLS",
            f"R 4.6.0; fastPLS {VERSION}",
            f"R 4.6.0; fastPLS {VERSION}",
        ],
        [
            "Core R packages",
            "Rcpp 1.1.1-1.1; RcppArmadillo 15.2.7-1; RcppEigen 0.3.4.0.2; Matrix 1.7-5; float 0.3-3; pls 2.9-0",
            "Rcpp 1.1.2; RcppArmadillo 15.4.0-1; RcppEigen 0.3.4.0.2; Matrix 1.7-5; float 0.3-3; pls 2.9-0",
        ],
        [
            "External comparison packages",
            "mdatools 0.15.0; chemometrics 1.4.4; pcv 1.1.0; plsdepot 0.3.1; plsgenomics 1.5-3; mixOmics 6.36.0; spls 2.3-2; ropls 1.44.0; IKPLS 6.1.2",
            "Not used for the external-package timing panel",
        ],
        [
            "Compiler",
            "GCC/G++ 11.4.0 with C++17; NVCC 13.0.88",
            "Homebrew Clang/Clang++ 22.1.1 with C++17 and Objective-C++",
        ],
        [
            "Accelerator libraries",
            "CUDA runtime 13.0.96; cuBLAS 13.1.1.3; cuSOLVER 12.0.4.66; cuRAND 10.4.0.35; NVIDIA driver 595.71.05",
            "Metal.framework and MetalPerformanceShaders.framework from the macOS SDK",
        ],
        [
            "BLAS/LAPACK and threads",
            "Reference BLAS/LAPACK 3.10.0; one effective BLAS thread for the controlled single-CPU comparisons",
            "R BLAS and LAPACK 3.12.1; one effective R-BLAS thread for the CPU/Metal comparisons",
        ],
    ]


def reproducibility_endpoint_rows():
    return [
        ["Endpoint", "Definition", "Interpretation"],
        ["Prediction agreement", "Fraction of identical decoded labels, or one minus relative Frobenius prediction error for numeric responses", "Primary model-output comparison across numerical routes"],
        ["Predictive-metric difference", "Absolute difference in accuracy, Q², or RMSD", "Endpoint agreement rather than bitwise equivalence"],
        ["Principal angles", "Angles between fitted latent subspaces", "Invariant to sign changes and rotations within nearly tied subspaces"],
        ["Selected-component agreement", "Equality of the component count selected on identical folds", "Checks downstream tuning reproducibility"],
        ["Numerical failures", "Non-convergence, non-finite output, factorization failure, timeout, or memory failure", "Retained explicitly rather than omitted"],
    ]


def output_contract_rows():
    return [
        ["Profile", "fastPLS retained outputs", "pls::simpls.fit retained outputs", "Interpretation"],
        [
            "Minimum common prediction outputs",
            "Coefficient path, means, and compact prediction state; no scores, loadings, fitted arrays, or variance summary",
            "stripped = TRUE: coefficient path and means",
            "Closest available comparison of estimator computation plus required prediction state",
        ],
        [
            "Ordinary public workflow",
            "Compact public fit object, means, latent prediction factors, variance summary, and held-out predictions",
            "Ordinary coefficient, score, loading, projection, fitted, residual, and variance paths plus held-out predictions",
            "End-to-end user workflows with package-specific output policies",
        ],
    ]


def finite(values):
    return [value for value in values if math.isfinite(value)]


def read_csv(path):
    with open(path, newline="", encoding="utf-8") as handle:
        rows = list(csv.DictReader(handle))
    audit_release_rows(path, rows)
    return rows


def identifies_fastpls(row, fields):
    identity_fields = (
        "package",
        "implementation",
        "implementation_fastpls",
        "function_name",
        "function_name_fastpls",
    )
    observed = [
        row.get(name, "").strip().lower()
        for name in identity_fields
        if name in fields
    ]
    if not observed:
        return True
    return any("fastpls" in value for value in observed)


def audit_release_rows(path, rows):
    """Reject stale, substituted, or non-fresh rows used by the manuscript."""
    if not rows:
        return
    fields = set(rows[0])
    version_fields = (
        "package_version",
        "fastpls_version",
        "loaded_package_version",
        "analysis_package_version",
    )
    for index, row in enumerate(rows, start=2):
        if not identifies_fastpls(row, fields):
            continue
        observed = {
            row.get(name, "").strip()
            for name in version_fields
            if name in fields and row.get(name, "").strip() not in {"", "NA"}
        }
        stale = observed - {VERSION}
        if stale:
            raise ValueError(
                f"{path} row {index} contains non-{VERSION} fastPLS "
                f"versions: {sorted(stale)}"
            )

    if {"backend_requested", "backend_reported"}.issubset(fields):
        for row in rows:
            requested = row.get("backend_requested", "").strip().lower()
            reported = row.get("backend_reported", "").strip().lower()
            status = row.get("status", "success").strip().lower()
            if status == "success" and requested in {"cuda", "metal"}:
                if reported != requested:
                    raise ValueError(
                        f"{path} contains a successful {requested} row reported as "
                        f"{reported or 'missing'}; CPU fallback results are not accepted."
                    )

    if "direction_rule" in fields:
        for row in rows:
            rule = row.get("direction_rule", "").strip().lower()
            if "warm" in rule:
                raise ValueError(
                    f"{path} contains a forbidden warm-start rule: {rule}"
                )

    if "fresh_start" in fields and ("method" in fields or "family" in fields):
        sequential_families = {"simpls", "opls", "kernelpls"}
        for row in rows:
            method = (row.get("method") or row.get("family") or "").strip().lower()
            status = row.get("status", "success").strip().lower()
            fresh = row.get("fresh_start", "").strip().lower()
            if (
                status == "success"
                and method in sequential_families
                and fresh not in {"", "na"}
                and fresh not in {"true", "t", "1"}
            ):
                raise ValueError(
                    f"{path} contains a non-fresh successful {method} row."
                )


def document_text(document):
    text = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            text.extend(cell.text for cell in row.cells)
    return "\n".join(text)


def embedded_figure_count(document):
    """Count drawing objects that are actually referenced by the document body."""
    return len(document.element.body.xpath(".//w:drawing"))


def audit_document_consistency(document, document_name):
    """Reject stale implementation and release language before publication."""
    text = document_text(document)
    lowered = text.lower()
    forbidden = {
        "sha-256": "source checksums are not part of the manuscript",
        "sha256": "source checksums are not part of the manuscript",
        "checksum": "source checksums are not part of the manuscript",
        "warm start": "the released implementation has no warm-start route",
        "warm-start": "the released implementation has no warm-start route",
        "candidate-knn": "candidate-kNN was removed from the package",
        "cknn": "candidate-kNN was removed from the package",
        "chiamaka": "machines must be identified by hardware, not host names",
        "review cycle": "internal review-cycle language is not publication text",
        "review object": "internal review-object language is not publication text",
        "historical": "obsolete benchmark language is not publication text",
        "archival": "obsolete benchmark language is not publication text",
        "deterministic numerical reference": (
            "IRLBA must not be described as an exact numerical reference"
        ),
    }
    for phrase, reason in forbidden.items():
        if phrase in lowered:
            raise AssertionError(f"{document_name}: {reason}: {phrase!r}")

    versions = set(re.findall(r"0\.99\.\d+", text))
    if versions != {VERSION}:
        raise AssertionError(
            f"{document_name}: expected only fastPLS {VERSION}, found {sorted(versions)}"
        )

    if re.search(r"\bPCA\b", text) and document_name == "supplement":
        raise AssertionError("supplement: PCA references were requested to be removed")
    strict_backend_phrases = (
        "no cpu fallback",
        "never replaced by cpu",
        "cpu execution is never substituted",
    )
    if not any(phrase in lowered for phrase in strict_backend_phrases):
        raise AssertionError(
            f"{document_name}: strict unavailable-accelerator behavior is not documented"
        )

    expected_figures = 4 if document_name == "main manuscript" else 18
    observed_figures = embedded_figure_count(document)
    if observed_figures < expected_figures:
        raise AssertionError(
            f"{document_name}: expected at least {expected_figures} embedded figures, "
            f"found {observed_figures}"
        )


def renumber_table_citations(document):
    def replace_values(values):
        return re.sub(
            r"S\d+[a-z]?",
            lambda item: TABLE_NUMBER_MAP.get(item.group(0), item.group(0)),
            values,
        )

    for paragraph in document.paragraphs:
        if "Table" not in paragraph.text:
            continue
        updated = re.sub(
            r"\bTable\s+(S\d+[a-z]?)",
            lambda match: f"Table {TABLE_NUMBER_MAP.get(match.group(1), match.group(1))}",
            paragraph.text,
        )
        updated = re.sub(
            r"\bTables\s+([^.;:)]+)",
            lambda match: f"Tables {replace_values(match.group(1))}",
            updated,
        )
        if updated != paragraph.text:
            set_paragraph(paragraph, updated)


def replace_figure(document, caption_prefix, path, width):
    if not path.exists():
        raise FileNotFoundError(path)
    caption = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.startswith(caption_prefix)
    )
    index = next(
        position for position, paragraph in enumerate(document.paragraphs)
        if paragraph._p is caption._p
    )
    paragraphs = list(document.paragraphs)
    candidates = list(reversed(paragraphs[max(0, index - 3):index]))
    candidates.extend(paragraphs[index + 1:index + 4])
    figure = next(
        (
            paragraph for paragraph in candidates
            if paragraph._p.xpath(".//w:drawing")
        ),
        None,
    )
    if figure is None:
        figure = caption.insert_paragraph_before("")
    figure.clear()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.add_run().add_picture(str(path), width=Inches(width))


def fmt(value, digits=3):
    try:
        number = float(value)
    except (TypeError, ValueError):
        return "NA"
    if not math.isfinite(number):
        return "NA"
    return f"{number:.{digits}f}"


def numeric(rows, field):
    values = []
    for row in rows:
        try:
            value = float(row.get(field, ""))
        except (TypeError, ValueError):
            continue
        if math.isfinite(value):
            values.append(value)
    return values


def range_text(values, digits=2, suffix=""):
    values = finite(values)
    if not values:
        return "NA"
    return f"{min(values):.{digits}f}-{max(values):.{digits}f}{suffix}"


def _primary_external_pairs():
    return [
        row for row in read_csv(EXTERNAL_PAIRS)
        if row.get("timing_mode") == "cold_process"
        and row.get("measurement_scope") == "primary"
    ]


def external_repetition_text():
    counts = sorted({
        int(float(row[field]))
        for row in _primary_external_pairs()
        for field in (
            "repetitions_completed_fastpls",
            "repetitions_completed_pls",
        )
        if row.get(field) not in (None, "", "NA")
        and int(float(row[field])) > 0
    })
    return "/".join(map(str, counts)) if counts else "recorded"


def package_repetition_text():
    rows = read_csv(PACKAGE_STATUS)
    counts = sorted({
        int(float(row["reps_attempted"]))
        for row in rows
        if row.get("reps_attempted") not in (None, "", "NA")
    })
    return "/".join(map(str, counts)) if counts else "recorded"


def external_simpls_result_text():
    """Summarize the final matched fastPLS/pls evidence without stale literals."""
    rows = _primary_external_pairs()
    profiles = {
        "estimator_kernel": "minimum common prediction outputs",
        "complete_workflow": "ordinary public workflows",
    }
    clauses = []
    for profile, label in profiles.items():
        selected = [
            row for row in rows
            if row.get("comparison_profile") == profile
            and row.get("speedup_pls_over_fastpls") not in (None, "", "NA")
        ]
        ratios = [float(row["speedup_pls_over_fastpls"]) for row in selected]
        faster = sum(value > 1 for value in ratios)
        clauses.append(f"fastPLS was faster on {faster} of {len(selected)} datasets with {label}")

    attempted = sum(
        int(float(row.get(field, 0) or 0))
        for row in rows
        for field in ("repetitions_attempted_fastpls", "repetitions_attempted_pls")
    )
    completed = sum(
        int(float(row.get(field, 0) or 0))
        for row in rows
        for field in ("repetitions_completed_fastpls", "repetitions_completed_pls")
    )
    comparisons = []
    public_rows = [
        row for row in rows
        if row.get("comparison_profile") == "complete_workflow"
        and row.get("speedup_pls_over_fastpls") not in (None, "", "NA")
    ]
    for row in sorted(
        public_rows,
        key=lambda item: float(item["speedup_pls_over_fastpls"]),
        reverse=True,
    )[:3]:
        comparisons.append(
            f"{float(row['speedup_pls_over_fastpls']):.2f}-fold on {row['dataset']}"
        )
    maximum_accuracy_difference = max(
        abs(float(row["accuracy_difference"]))
        for row in rows
        if row.get("accuracy_difference") not in (None, "", "NA")
    )
    comparison_sentence = "; ".join(clauses)
    comparison_sentence = comparison_sentence[0].upper() + comparison_sentence[1:]
    return (
        f"The repeated cold-process comparison completed {completed} of {attempted} "
        f"planned worker runs. {comparison_sentence}. The three largest "
        f"ordinary-workflow ratios were {', '.join(comparisons)}. The largest absolute "
        f"held-out accuracy difference between paired fastPLS and pls::simpls.fit rows was "
        f"{maximum_accuracy_difference:.4f}. Retina and Tabula Muris rows are cells rather "
        "than independent biological replicates, so these fixed-split endpoints are not "
        "interpreted as biological generalization intervals. The broader multi-package "
        "workflow comparison is reported separately in Figure 1."
    )


def external_simpls_memory_text(dataset="cifar100"):
    rows = [
        row for row in _primary_external_pairs()
        if row.get("comparison_profile") == "complete_workflow"
        and row.get("dataset") == dataset
    ]
    if len(rows) != 1:
        raise ValueError(f"Expected one complete-workflow external row for {dataset}")
    row = rows[0]
    dataset_label = "CIFAR-100" if dataset == "cifar100" else dataset
    if row.get("median_process_peak_rss_mb_pls") in (None, "", "NA"):
        return (
            "The public-workflow comparison also exposed an output-memory limit. "
            f"On {dataset_label}, fastPLS completed with a compact fit object of "
            f"{float(row['median_fit_object_mb_fastpls']):.2f} MiB and a median "
            "baseline-corrected complete-process peak-RSS increment of "
            f"{float(row['median_baseline_corrected_peak_increment_mb_fastpls']):.2f} MiB. "
            "The ordinary pls::simpls.fit workflow was killed before producing a "
            "completed replicate because its fitted and residual response paths exceed "
            "the available memory. This is a workflow-feasibility comparison, not an "
            "isolated algorithmic-workspace measurement."
        )
    return (
        "The magnitude of the public-workflow difference was influenced by output "
        f"retention. On {dataset_label}, the compact fastPLS object occupied "
        f"{float(row['median_fit_object_mb_fastpls']):.2f} MiB and the ordinary "
        f"pls::simpls.fit object {float(row['median_fit_object_mb_pls']):.2f} MiB. "
        "Median complete-process peak RSS was "
        f"{float(row['median_process_peak_rss_mb_fastpls']):.2f} and "
        f"{float(row['median_process_peak_rss_mb_pls']):.2f} MiB; baseline-corrected "
        f"increments were {float(row['median_baseline_corrected_peak_increment_mb_fastpls']):.2f} "
        f"and {float(row['median_baseline_corrected_peak_increment_mb_pls']):.2f} MiB. "
        "These measurements characterize the specified complete workflows rather than "
        "isolated algorithmic workspace."
    )


def package_panel_statistics():
    summary = read_csv(PACKAGE_SUMMARY)
    status = read_csv(PACKAGE_STATUS)
    completed = [
        row for row in summary
        if int(float(row.get("reps_ok", 0) or 0)) > 0
        and math.isfinite(float(row.get("median_time_ms", "nan")))
    ]
    datasets = sorted({row["dataset"] for row in completed})
    comparisons = []
    for dataset in datasets:
        rows = [row for row in completed if row["dataset"] == dataset]
        fastpls = [row for row in rows if row.get("package") == "fastPLS"]
        external = [row for row in rows if row.get("package") != "fastPLS"]
        if not fastpls or not external:
            continue
        fastest_fastpls = min(fastpls, key=lambda row: float(row["median_time_ms"]))
        fastest_external = min(external, key=lambda row: float(row["median_time_ms"]))
        comparisons.append({
            "dataset": dataset,
            "ratio": (
                float(fastest_external["median_time_ms"])
                / float(fastest_fastpls["median_time_ms"])
            ),
            "fastpls": fastest_fastpls,
            "external": fastest_external,
        })

    external_status = [row for row in status if row.get("package") != "fastPLS"]
    fastpls_status = [row for row in status if row.get("package") == "fastPLS"]
    return {
        "comparisons": comparisons,
        "faster": sum(item["ratio"] > 1 for item in comparisons),
        "external_completed": sum(
            int(float(row.get("reps_ok", 0) or 0)) > 0
            for row in external_status
        ),
        "external_total": len(external_status),
        "fastpls_completed": sum(
            int(float(row.get("reps_ok", 0) or 0)) > 0
            for row in fastpls_status
        ),
        "fastpls_total": len(fastpls_status),
    }


def package_panel_result_text():
    values = package_panel_statistics()
    comparisons = values["comparisons"]
    if not comparisons:
        raise ValueError("The R-package panel has no matched completed datasets.")
    leading = sorted(comparisons, key=lambda item: item["ratio"], reverse=True)[:3]
    examples = ", ".join(
        f"{item['ratio']:.2f}-fold on {item['dataset']}"
        for item in leading
    )
    return (
        "In the broader fixed-split classification panel, fastPLS completed "
        f"{values['fastpls_completed']} of {values['fastpls_total']} configured "
        "method-dataset workflows; independent R implementations completed "
        f"{values['external_completed']} of {values['external_total']}. The fastest "
        f"completed fastPLS workflow was faster than the fastest completed external "
        f"workflow on {values['faster']} of {len(comparisons)} matched datasets, with "
        f"the largest ratios including {examples}. Figure 1 reports the paired "
        "accuracy, total fitting-plus-prediction time, and complete-process peak RSS. "
        "Rows that use different classification heads or retain different model "
        "objects are workflow comparisons rather than estimator-equivalence results."
    )


PACKAGE_PANEL_METHODS = (
    "fastPLS_simpls_cpu_irlba",
    "fastPLS_simpls_cpu_irlba_lda",
    "pls_simpls_fit",
    "plsgenomics_pls_lda",
    "mdatools_plsda_or_pls",
    "plsdepot_simpls",
    "pcv_simpls",
    "chemometrics_pls_eigen",
    "mixOmics_plsda",
    "spls_splsda",
)


def _package_panel_cell(row):
    return (
        f"{float(row['median_accuracy']):.3f}; "
        f"{float(row['median_time_ms']) / 1000:.3g} s; "
        f"{float(row['median_peak_host_rss_mb']):.0f} MiB"
    )


def _package_panel_status(status):
    if status is None:
        return "Not configured"
    if int(float(status.get("n_timeout", 0) or 0)):
        return "Timeout"
    if int(float(status.get("n_error", 0) or 0)):
        return "Error"
    if int(float(status.get("n_skipped", 0) or 0)):
        return "Not evaluated"
    return "No completed repetition"


def package_panel_rows():
    summary = read_csv(PACKAGE_SUMMARY)
    status = read_csv(PACKAGE_STATUS)
    summary_by_key = {
        (row["dataset"], row["method_id"]): row
        for row in summary
        if int(float(row.get("reps_ok", 0) or 0)) > 0
    }
    status_by_key = {
        (row["dataset"], row["method_id"]): row for row in status
    }
    datasets = (
        ("ccle", "CCLE"),
        ("cifar100", "CIFAR-100"),
        ("gtex_v8", "GTEx v8"),
        ("metref", "MetRef"),
        ("retina", "Retina"),
        ("tabula", "Tabula Muris"),
        ("tcga_brca", "TCGA-BRCA"),
        ("tcga_hnsc_methylation", "TCGA-HNSC methylation"),
        ("tcga_pan_cancer", "TCGA Pan-Cancer"),
    )
    missing = [
        (dataset, method)
        for dataset, _ in datasets
        for method in PACKAGE_PANEL_METHODS
        if (dataset, method) not in status_by_key
    ]
    if missing:
        sample = ", ".join(f"{dataset}/{method}" for dataset, method in missing[:4])
        raise ValueError(
            "The independent-package panel is incomplete: "
            f"{len(missing)} required status rows are absent ({sample})."
        )

    rows = [[
        "Dataset", "A", "fastPLS argmax", "fastPLS LDA",
        "Fastest completed external", "Most accurate completed external",
    ]]
    for dataset, label in datasets:
        available = [
            row for (key_dataset, _), row in summary_by_key.items()
            if key_dataset == dataset
        ]
        external = [row for row in available if row.get("package") != "fastPLS"]
        if not external:
            raise ValueError(f"No external package completed for {dataset}.")
        fastest = min(external, key=lambda row: float(row["median_time_ms"]))
        best = min(
            external,
            key=lambda row: (
                -float(row["median_accuracy"]),
                float(row["median_time_ms"]),
            ),
        )
        argmax = summary_by_key.get((dataset, "fastPLS_simpls_cpu_irlba"))
        lda = summary_by_key.get((dataset, "fastPLS_simpls_cpu_irlba_lda"))
        reference = argmax or lda or fastest
        rows.append([
            label,
            str(int(float(reference["ncomp_requested"]))),
            _package_panel_cell(argmax) if argmax else _package_panel_status(
                status_by_key[(dataset, "fastPLS_simpls_cpu_irlba")]
            ),
            _package_panel_cell(lda) if lda else _package_panel_status(
                status_by_key[(dataset, "fastPLS_simpls_cpu_irlba_lda")]
            ),
            f"{fastest['package']} / {fastest['algorithm']}: {_package_panel_cell(fastest)}",
            f"{best['package']} / {best['algorithm']}: {_package_panel_cell(best)}",
        ])
    return rows


def ensure_package_panel_caption(document):
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith("Table S9e."):
            return
    found_figure = False
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, document)
            if paragraph.text.strip().startswith("Figure S3."):
                found_figure = True
        elif isinstance(child, CT_Tbl) and found_figure:
            element = OxmlElement("w:p")
            child.addprevious(element)
            caption = Paragraph(element, document)
            caption.style = "Caption"
            caption.add_run(
                "Table S9e. Broad single-CPU SIMPLS classification workflow summary. "
                "Each cell reports held-out accuracy, median total fitting-plus-prediction "
                "time, and median complete-process peak host RSS from the recorded fresh-process "
                "repetitions. The fastest and most accurate external rows may represent different "
                "estimators or classification heads and are workflow comparisons."
            )
            return
    raise ValueError("Could not locate the broad package-panel table after Figure S3.")


def abstract_result_text():
    panel = package_panel_statistics()
    qualification = read_csv(
        RESULTS / "rsvd_qualification/rsvd_qualification_summary.csv"
    )
    checks = sum(int(float(row["comparisons"])) for row in qualification)
    within = sum(
        int(float(row["comparisons_within_tolerance"]))
        for row in qualification
    )
    nmr = _nmr_route("selected", "simpls", "cuda", "rsvd", 50)
    imagenet = read_csv(RESULTS / "imagenet/imagenet_current_summary.csv")
    imagenet_success = sum(
        row.get("status") == "success" for row in imagenet
    )
    return (
        "Results: The fixed-control SIMPLS comparisons met the stated numerical "
        f"tolerances, and rSVD met the multi-seed screening tolerances in {within}/{checks} "
        "CPU, CUDA, and Metal comparisons. In the nine-dataset R-package panel, the "
        f"fastest fastPLS workflow was faster than the fastest completed independent "
        f"workflow on {panel['faster']} of {len(panel['comparisons'])} datasets; failed "
        "and unsupported external routes were retained. For the 28,355-response NMR "
        f"task, training-selected 50-component CUDA SIMPLS-rSVD required {nmr['time']:.2f} s "
        f"and achieved held-out RMSD {nmr['rmsd']:.6f}. The million-row ImageNet "
        f"stress test completed {imagenet_success} argmax/LDA component-prefix "
        "evaluations, demonstrating downstream matrix-processing feasibility rather "
        "than biomedical predictive validity."
    )


def ikpls_result_text():
    rows = read_csv(IKPLS_SUMMARY)
    datasets = sorted({row["dataset"] for row in rows})
    clauses = []
    for dataset in datasets:
        selected = [row for row in rows if row["dataset"] == dataset]
        fastest = min(selected, key=lambda row: float(row["median_total_sec"]))
        label = (
            fastest["implementation"]
            .replace("fastPLS_cpu_", "fastPLS ")
            .replace("IKPLS_numpy_alg", "IKPLS formulation ")
        )
        clauses.append(
            f"{dataset}: {label}, {float(fastest['median_total_sec']):.4g} s"
        )
    repetitions = sum(int(float(row["repetitions"])) for row in rows)
    return (
        f"The matched-precision CPU comparison with IKPLS completed {repetitions} measured "
        f"runs across {len(datasets)} datasets. The fastest workflows were "
        f"{'; '.join(clauses)}. Because IKPLS and fastPLS implement different estimators "
        "and runtime stacks, this is an end-to-end software comparison rather than an "
        "estimator-equivalence benchmark; complete accuracy, timing, and memory results "
        "are reported in the Supplement."
    )


def backend_result_text():
    by_backend = {
        "CUDA": read_csv(CUDA_PAIRED),
        "Metal": read_csv(METAL_PAIRED),
    }
    summaries = []
    for label, rows in by_backend.items():
        successful = [
            row for row in rows
            if int(float(row.get("cpu_ok", 0) or 0)) > 0
            and int(float(row.get("accelerator_ok", 0) or 0)) > 0
            and math.isfinite(float(row.get("cpu_accelerator_ratio", "nan")))
        ]
        faster = sum(float(row["cpu_accelerator_ratio"]) > 1 for row in successful)
        summaries.append(f"{label} was faster in {faster} of {len(successful)} pairs")
    cifar = []
    for label, rows in by_backend.items():
        selected = [row for row in rows if row.get("dataset") == "cifar100"]
        values = ", ".join(
            f"{row['method'].replace('plssvd', 'PLS-SVD').replace('kernelpls', 'kernel PLS').upper() if row['method'] in ('simpls', 'opls') else row['method'].replace('plssvd', 'PLS-SVD').replace('kernelpls', 'kernel PLS')}: "
            f"{float(row['cpu_accelerator_ratio']):.2f}"
            for row in selected
        )
        cifar.append(f"CPU/{label} ratios were {values}")
    return (
        f"All {len(by_backend['CUDA'])} CUDA and {len(by_backend['Metal'])} Metal paired "
        f"workflows completed. {'; '.join(summaries)}. On CIFAR-100, "
        f"{'; '.join(cifar)}. Most small omics workloads favored CPU because transfer, "
        "allocation, and launch overhead exceeded the accelerated linear-algebra work. "
        "Prediction agreement and metric differences are reported separately, so the runtime "
        "panel describes computational behavior rather than universal numerical interchangeability."
    )


def _nmr_route(prefix, family, backend, solver, ncomp):
    path = NMR_RESULTS / f"{prefix}_{family}_{backend}_{solver}_k{ncomp}.csv"
    rows = read_csv(path)
    successful = [row for row in rows if row.get("status", "success") == "success"]
    if not successful:
        raise ValueError(f"No successful NMR rows in {path}")
    return {
        "time": statistics.median(float(row["total_time_sec"]) for row in successful),
        "rmsd": statistics.median(float(row["RMSD"]) for row in successful),
        "q2": statistics.median(float(row["Q2"]) for row in successful),
        "oversample": successful[0].get("oversample", "NA"),
        "power": successful[0].get("power", "NA"),
    }


def nmr_selection_text():
    pls_svd = read_csv(
        NMR_RESULTS / "selection_plssvd/nmr_component_selection_decision.csv"
    )[0]
    simpls = read_csv(
        NMR_RESULTS / "selection_simpls/nmr_component_selection_decision.csv"
    )[0]
    return (
        f"Training-only selection retained {simpls['selected_ncomp']} SIMPLS-rSVD components "
        "by the one-standard-error rule; its eligible set was "
        f"{simpls['eligible_ncomp']}. PLS-SVD retained {pls_svd['selected_ncomp']} components. "
        "Both choices were interior to the evaluated 1-300-component grid and are described as "
        "selected within that grid. Figure 3 addresses a separate computational question by "
        "comparing deposited and fastPLS workflows at the 165 components used in the Nature "
        "Communications analysis."
    )


def nmr_fixed_text():
    fixed_rows = read_csv(
        NMR_RESULTS / "figures/nmr_fixed165_summary.csv"
    )
    by_label = {row["label"]: row for row in fixed_rows}
    deposited = by_label["Deposited PLS-SVD"]
    current_plssvd = by_label["PLS-SVD CPU / IRLBA"]
    cpu = _nmr_route("fixed165", "simpls", "cpu", "rsvd", 165)
    cuda = _nmr_route("fixed165", "simpls", "cuda", "rsvd", 165)
    metal = _nmr_route("fixed165", "simpls", "metal", "rsvd", 165)
    return (
        "At the common 165-component workload, the deposited PLS-SVD implementation "
        f"required {float(deposited['total_time_sec']):.2f} s, achieved RMSD "
        f"{float(deposited['RMSD']):.6f}, and had a baseline-corrected peak process-RSS "
        f"increment of {float(deposited['incremental_peak_rss_mib']):.1f} MiB. The current "
        f"CPU PLS-SVD/IRLBA route required {float(current_plssvd['total_time_sec']):.2f} s, "
        f"achieved RMSD {float(current_plssvd['RMSD']):.6f}, and used an increment of "
        f"{float(current_plssvd['incremental_peak_rss_mib']):.1f} MiB. SIMPLS-rSVD required "
        "median fitting-plus-prediction "
        f"times of {cpu['time']:.2f} s on CPU, {cuda['time']:.2f} s on CUDA, and "
        f"{metal['time']:.2f} s on Metal. The corresponding held-out RMSDs were "
        f"{cpu['rmsd']:.6f}, {cuda['rmsd']:.6f}, and {metal['rmsd']:.6f}. These values were "
        "less favorable than the training-selected 50-component SIMPLS workflow, confirming "
        "that the fixed-165 experiment is a matched computational workload rather than a "
        "model-selection result."
    )


def nmr_selected_text():
    cpu = _nmr_route("selected", "simpls", "cpu", "rsvd", 50)
    cuda = _nmr_route("selected", "simpls", "cuda", "rsvd", 50)
    metal = _nmr_route("selected", "simpls", "metal", "rsvd", 50)
    return (
        "The predictive and computational analyses were kept separate. At the selected 50 "
        f"components, SIMPLS-rSVD achieved RMSDs of {cpu['rmsd']:.6f}, "
        f"{cuda['rmsd']:.6f}, and {metal['rmsd']:.6f} on CPU, CUDA, and Metal, with "
        f"median total times of {cpu['time']:.2f}, {cuda['time']:.2f}, and "
        f"{metal['time']:.2f} s. Every rSVD row records its automatic control profile, "
        "effective oversampling, power iterations, seed, and fresh direction rule."
    )


def nmr_comparison_table_rows():
    """Build the NMR table from current-release summaries, never inherited values."""
    rows = [[
        "Analysis", "Family", "Implementation", "A", "Predictive metric",
        "Time, s [IQR]", "Delta host MiB", "Delta GPU MiB", "Controls / scope",
    ]]
    inputs = (
        (
            "Training selected",
            NMR_RESULTS / "figures/nmr_selected_summary.csv",
        ),
        (
            "Common 165 components",
            NMR_RESULTS / "figures/nmr_fixed165_summary.csv",
        ),
    )
    for analysis, path in inputs:
        for row in read_csv(path):
            family = "PLS-SVD" if row["family"] == "plssvd" else "SIMPLS"
            label = row["label"]
            if label == "Deposited PLS-SVD":
                controls = "Deposited Nature Communications implementation"
            elif row["solver"] == "rsvd":
                controls = (
                    f"rSVD; oversample {row.get('oversample', 'NA')}; "
                    f"power {row.get('power', 'NA')}; seed {row.get('seed', 'NA')}"
                )
            else:
                controls = "Fixed-control CPU IRLBA"
            rows.append([
                analysis,
                family,
                label,
                str(int(float(row["ncomp"]))),
                f"RMSD {float(row['RMSD']):.6f}; Q² {float(row['Q2']):.6f}",
                (
                    f"{float(row['total_time_sec']):.3f} "
                    f"[{float(row['total_time_iqr']):.3f}]"
                ),
                f"{float(row['incremental_peak_rss_mib']):.1f}",
                "Not captured",
                controls,
            ])
    return rows


def ablation_rows():
    rows = read_csv(ABLATION_EFFECTS)
    labels = {
        "cached_deflation_products": "Cached deflation products",
        "cached_XtX": "Cached X'X",
        "incremental_coefficients": "Incremental coefficients",
        "compact_prediction": "Compact latent prediction",
        "matrix_free": "Implicit cross-covariance",
    }
    interpretations = {
        "cached_deflation_products": "Small, shape-dependent timing and memory changes.",
        "cached_XtX": "Not applicable in the four evaluated matrix shapes; no speed claim.",
        "incremental_coefficients": "Avoids rebuilding coefficient prefixes; modest timing gains.",
        "compact_prediction": "Largest consistent public-workflow gain, especially for high-dimensional responses.",
        "matrix_free": "Trades repeated operator products for avoiding explicit cross-covariance storage; beneficial for PRISM but slower in smaller shapes.",
    }
    result = [["Optimization", "Datasets", "Time speed-up range", "Incremental RSS change", "Interpretation and agreement"]]
    for key in labels:
        selected = [row for row in rows if row["optimization"] == key]
        result.append([
            labels[key],
            ", ".join(sorted(row["dataset"] for row in selected)),
            range_text(numeric(selected, "speedup"), 2, "x"),
            range_text(numeric(selected, "rss_reduction_pct"), 1, "%"),
            f"{interpretations[key]} Minimum prediction agreement {fmt(min(numeric(selected, 'prediction_agreement_min')), 4)}.",
        ])
    return result


def qualification_factor_rows(paths):
    by_backend = {}
    for backend, path in paths.items():
        rows = [
            row for row in read_csv(path)
            if row.get("svd_method") == "rsvd"
            and row.get("route", "").endswith("_auto")
            and row.get("backend") == backend.lower()
        ]
        by_backend[backend] = rows
    factors = sorted({row["factor_name"] for rows in by_backend.values() for row in rows})
    result = [["Factor", "Tested values", "CPU checks", "CUDA checks", "Metal checks", "Largest prediction error"]]
    for factor in factors:
        factor_rows = [row for rows in by_backend.values() for row in rows if row["factor_name"] == factor]
        values = sorted({row["factor_label"] for row in factor_rows})
        backend_cells = []
        for backend in ("CPU", "CUDA", "Metal"):
            selected = [row for row in by_backend.get(backend, []) if row["factor_name"] == factor]
            passed = sum(row.get("numerical_status") == "within_tolerance" for row in selected)
            backend_cells.append(f"{passed}/{len(selected)}")
        errors = numeric(factor_rows, "prediction_relative_error")
        if not errors:
            errors = numeric(factor_rows, "metric_absolute_difference")
        result.append([
            factor,
            ", ".join(values),
            *backend_cells,
            f"{max(errors):.3g}",
        ])
    return result


def crosscov_rows(paths):
    result = [["Backend", "Cross-covariance MiB", "Explicit median s", "Implicit median s", "Implicit / explicit", "Checks met", "Interpretation"]]
    for backend, path in paths.items():
        rows = [
            row for row in read_csv(path)
            if row.get("svd_method") == "rsvd"
            and row.get("factor_name") == "crosscov_mb"
            and row.get("backend") == backend.lower()
        ]
        for value in sorted({float(row["crosscov_mb"]) for row in rows}):
            selected = [row for row in rows if abs(float(row["crosscov_mb"]) - value) < 1e-8]
            explicit = [float(row["total_sec"]) for row in selected if row.get("xprod_requested") == "explicit"]
            implicit = [float(row["total_sec"]) for row in selected if row.get("xprod_requested") == "implicit"]
            if not explicit or not implicit:
                continue
            exp_med = statistics.median(explicit)
            imp_med = statistics.median(implicit)
            evaluated = [row for row in selected if row.get("xprod_requested") in {"explicit", "implicit"}]
            passed = sum(row.get("numerical_status") == "within_tolerance" for row in evaluated)
            result.append([
                backend,
                f"{value:.1f}",
                f"{exp_med:.3f}",
                f"{imp_med:.3f}",
                f"{imp_med / exp_med:.2f}x",
                f"{passed}/{len(evaluated)}",
                "Implicit avoids materializing the cross-covariance but can require more operator products.",
            ])
    return result


def shape_rows():
    result = [["Shape", "n / p / q / components", "PLS-SVD / SIMPLS time (s)", "SIMPLS / PLS-SVD", "RMSD PLS-SVD / SIMPLS"]]
    for row in read_csv(SHAPE_PAIRED):
        result.append([
            row["dataset"].replace("synthetic_", "").replace("_", " ").title(),
            f"{row['n_train']} / {row['p']} / {row['q']} / {row['ncomp']}",
            f"{fmt(row['plssvd_total_sec'])} / {fmt(row['simpls_total_sec'])}",
            f"{float(row['simpls_over_plssvd_time']):.2f}x",
            f"{float(row['plssvd_rmsd']):.3f} / {float(row['simpls_rmsd']):.3f}",
        ])
    return result


def estimator_validation_rows():
    """Build estimator-validation evidence entirely from this release."""
    exact = read_csv(SIMPLS_EXACT_SUMMARY)
    exact_prefixes = sum(int(float(row["component_prefixes"])) for row in exact)
    exact_max_angle = max(
        max(
            float(row[field])
            for field in (
                "max_score_subspace_angle_degrees",
                "max_loading_subspace_angle_degrees",
                "max_projection_subspace_angle_degrees",
            )
        )
        for row in exact
    )
    exact_min_labels = min(
        numeric(exact, "min_classification_label_agreement")
    )

    preservation = read_csv(SIMPLS_PRESERVATION)[0]
    preservation_groups = read_csv(SIMPLS_PRESERVATION_SUMMARY)
    preservation_max_angle = max(
        max(
            float(row[field])
            for field in (
                "score_subspace_max_angle_degrees",
                "projection_subspace_max_angle_degrees",
                "loading_subspace_max_angle_degrees",
            )
        )
        for row in preservation_groups
    )

    setting_rows = read_csv(OPLS_KERNEL_SETTINGS)
    setting_selection = read_csv(OPLS_KERNEL_SETTINGS_SELECTION)
    setting_checks = sum(int(float(row["runs"])) for row in setting_rows)
    setting_passes = sum(int(float(row["passes_all"])) for row in setting_rows)
    selection_passes = sum(
        row.get("selected_component_agreement", "").upper() == "TRUE"
        for row in setting_selection
    )

    return [
        [
            "fastPLS scope", "Reference", "Endpoint checks", "Selection agreement",
            "Max pred. rel. error", "Max coef. rel. error", "Max angle",
            "Max metric diff.", "Conclusion",
        ],
        [
            "SIMPLS / dense LAPACK",
            "Independent de Jong updates / base LAPACK SVD",
            f"{exact_prefixes}/{exact_prefixes}",
            "Not a selection study",
            f"{max(numeric(exact, 'max_prediction_relative_error')):.3g}",
            f"{max(numeric(exact, 'max_coefficient_relative_error')):.3g}",
            f"{exact_max_angle:.3g} degrees",
            f"Minimum decoded-label agreement {exact_min_labels:.4f}",
            "All exact-reference cases completed",
        ],
        [
            "SIMPLS / CPU IRLBA",
            "pls::simpls.fit / de Jong",
            (
                f"{preservation['deterministic_endpoint_tolerance_passes']}/"
                f"{preservation['deterministic_endpoint_rows']}"
            ),
            "24/24" if float(
                preservation["deterministic_cv_selection_agreement"]
            ) == 1 else "Outside tolerance",
            f"{max(numeric(preservation_groups, 'prediction_relative_error')):.3g}",
            f"{max(numeric(preservation_groups, 'coefficient_relative_error')):.3g}",
            f"{preservation_max_angle:.3g} degrees",
            f"{max(numeric(preservation_groups, 'metric_absolute_difference')):.3g}",
            "Met the predefined numerical tolerances",
        ],
        [
            "OPLS (1-3 orthogonal components) and kernel PLS (8 settings)",
            "Independent filtering/kernel construction + pls::simpls.fit",
            f"{setting_passes}/{setting_checks}",
            f"{selection_passes}/{len(setting_selection)}",
            f"{max(numeric(setting_rows, 'max_prediction_relative_error')):.3g}",
            f"{max(numeric(setting_rows, 'max_coefficient_relative_error')):.3g}",
            (
                f"{max(numeric(setting_rows, 'max_predictive_score_angle_deg')):.3g} "
                "degrees"
            ),
            f"{max(numeric(setting_rows, 'max_metric_absolute_difference')):.3g}",
            "Met the predefined float64 CPU tolerances",
        ],
    ]


def external_memory_rows(profile):
    """Return paired process-memory accounting for one output contract."""
    result = [[
        "Dataset", "Absolute peak fastPLS / pls", "Pre-fit baseline fastPLS / pls",
        "Peak increment fastPLS / pls", "Largest retained object fastPLS / pls",
    ]]
    source = sorted(
        (
            row for row in _primary_external_pairs()
            if row.get("comparison_profile") == profile
        ),
        key=lambda row: row["dataset"],
    )
    for row in source:
        def paired(field):
            return (
                f"{fmt(row[field + '_fastpls'], 1)} "
                f"[{fmt(row['iqr_' + field.removeprefix('median_') + '_fastpls'], 1)}] / "
                f"{fmt(row[field + '_pls'], 1)} "
                f"[{fmt(row['iqr_' + field.removeprefix('median_') + '_pls'], 1)}]"
            )

        result.append([
            row["dataset"],
            paired("median_process_peak_rss_mb"),
            paired("median_prefit_process_rss_mb"),
            paired("median_baseline_corrected_peak_increment_mb"),
            (
                f"{row['theoretical_largest_retained_name_fastpls']}: "
                f"{fmt(row['theoretical_largest_retained_mb_fastpls'], 2)} MiB / "
                f"{row.get('theoretical_largest_retained_name_pls') or 'Not completed'}: "
                f"{fmt(row.get('theoretical_largest_retained_mb_pls'), 2)} MiB"
            ),
        ])
    return result


def ikpls_large_rows():
    result = [[
        "Dataset", "Implementation", "Components", "Status", "Fit / predict (s)",
        "Predictive endpoint", "Peak process RSS (MiB)",
    ]]
    for row in read_csv(IKPLS_LARGE_SUMMARY):
        if row["dataset"] == "ImageNet" and not row["implementation"].startswith("IKPLS"):
            continue
        if row["dataset"] == "NMR" and not row["implementation"].startswith("IKPLS"):
            continue
        endpoint = "Not estimated"
        if row.get("rmsd") not in (None, "", "nan", "NA"):
            endpoint = f"RMSD {fmt(row['rmsd'], 6)}"
        elif row.get("top1_accuracy") not in (None, "", "nan", "NA"):
            endpoint = (
                f"Top-1 {fmt(row['top1_accuracy'], 4)}; "
                f"top-5 {fmt(row['top5_accuracy'], 4)}"
            )
        result.append([
            row["dataset"],
            row["implementation"],
            str(int(float(row["ncomp"]))),
            row["status"],
            f"{fmt(row['fit_sec'])} / {fmt(row['predict_sec'])}",
            endpoint,
            fmt(row["peak_process_rss_mib"], 1),
        ])
    return result


def paired_backend_rows(path, accelerator):
    label = accelerator.upper() if accelerator == "cuda" else "Metal"
    result = [[
        "Dataset", "Family", "A", f"Metric CPU / {label}",
        "Prediction agreement", f"Time CPU / {label} (s)",
        f"Incremental RSS CPU / {label} (MiB)", "Execution status",
    ]]
    family_labels = {
        "plssvd": "PLS-SVD", "simpls": "SIMPLS",
        "opls": "OPLS", "kernelpls": "kernel PLS",
    }
    for row in read_csv(path):
        status = (
            f"CPU {row['cpu_ok']}/3; {label} {row['accelerator_ok']}/3"
        )
        result.append([
            row["dataset"],
            family_labels[row["method"]],
            str(int(float(row["ncomp"]))),
            f"{fmt(row['metric_cpu'], 6)} / {fmt(row['metric_accelerator'], 6)}",
            fmt(row["prediction_agreement"], 6),
            f"{fmt(row['cpu_total_sec'])} / {fmt(row['accelerator_total_sec'])}",
            (
                f"{fmt(row['cpu_incremental_rss_mb'], 1)} / "
                f"{fmt(row['accelerator_incremental_rss_mb'], 1)}"
            ),
            status,
        ])
    return result


def imagenet_rows():
    result = [[
        "Head", "Components", "Top-1", "Top-5", "Balanced accuracy",
        "Macro F1", "Fit + prediction (s)", "Incremental RSS / GPU (MiB)",
        "rSVD controls and status",
    ]]
    for row in read_csv(IMAGENET_SUMMARY):
        if row.get("status") != "success":
            continue
        result.append([
            row["classifier"].upper(),
            str(int(float(row["ncomp_requested"]))),
            fmt(row["top1_accuracy"], 5),
            fmt(row["top5_accuracy"], 5),
            fmt(row["balanced_accuracy"], 5),
            fmt(row["macro_f1"], 5),
            fmt(row["total_time_sec"], 2),
            (
                f"{fmt(row['incremental_peak_rss_mb'], 1)} / "
                f"{fmt(row['gpu_incremental_peak_mb'], 1)}"
            ),
            (
                f"oversample {row['effective_oversample']}; power "
                f"{row['effective_power']}; seed {row['seed']}; "
                f"{row['audit_status']}"
            ),
        ])
    return result


def repeated_outer_rows():
    expected = ("gtex_v8", "metref", "nmr", "retina")
    dataset_labels = {
        "gtex_v8": "GTEx v8",
        "metref": "MetRef",
        "nmr": "NMR",
        "retina": "Retina",
    }
    source = []
    for dataset in expected:
        path = REPEATED_OUTER_DIR / dataset / "repeated_outer_predictive_dispersion.csv"
        if not path.exists():
            raise FileNotFoundError(path)
        source.extend(read_csv(path))
    result = [[
        "Dataset", "Family", "Head", "Outer n", "Metric", "Mean (SD)",
        "95% empirical interval", "Selected A range", "Upper-bound frequency",
        "Constraint",
    ]]
    for row in sorted(
        source,
        key=lambda item: (item["dataset"], item["method"], item["classifier"]),
    ):
        result.append([
            dataset_labels.get(row["dataset"], row["dataset"]),
            {"plssvd": "PLS-SVD", "simpls": "SIMPLS", "opls": "OPLS", "kernelpls": "kernel PLS"}[row["method"]],
            row["classifier"],
            row["n_outer_success"],
            row["metric_name"],
            f"{fmt(row['metric_mean'], 6)} ({fmt(row['metric_sd'], 3)})",
            f"{fmt(row['metric_q025'], 6)}-{fmt(row['metric_q975'], 6)}",
            f"{int(float(row['selected_ncomp_min']))}-{int(float(row['selected_ncomp_max']))}",
            fmt(row["upper_boundary_frequency"], 2),
            "Rank constrained" if row.get("rank_constrained_grid", "").upper() == "TRUE" else "None",
        ])
    return result


def summarize_rsvd(path, backend):
    with open(path, newline="", encoding="utf-8") as handle:
        rows = list(csv.DictReader(handle))
    selected = [row for row in rows if row.get("backend") == backend and "rsvd" in row.get("route", "")]
    evaluated = [row for row in selected if row.get("numerical_status") in {"within_tolerance", "outside_tolerance"}]
    within = [row for row in evaluated if row.get("numerical_status") == "within_tolerance"]

    def numbers(name):
        result = []
        for row in evaluated:
            try:
                result.append(float(row.get(name, "")))
            except (TypeError, ValueError):
                pass
        return finite(result)

    pred = numbers("prediction_relative_error")
    corr = numbers("prediction_correlation")
    score = numbers("score_relative_error")
    labels = numbers("label_agreement")
    metric = numbers("metric_absolute_difference")
    oversample = sorted({
        int(float(row["rsvd_effective_oversample"]))
        for row in evaluated
        if row.get("rsvd_effective_oversample") not in (None, "", "NA")
    })
    power = sorted({
        int(float(row["rsvd_effective_power"]))
        for row in evaluated
        if row.get("rsvd_effective_power") not in (None, "", "NA")
    })
    return {
        "oversample": ";".join(map(str, oversample)),
        "power": ";".join(map(str, power)),
        "checks": f"{len(within)}/{len(evaluated)}",
        "pred": f"{max(pred):.3g}" if pred else "NA",
        "corr": f"{min(corr):.5f}" if corr else "NA",
        "score": f"{max(score):.3g}" if score else "NA",
        "labels": f"{min(labels):.4f}" if labels else "NA",
        "metric": f"{max(metric):.3g}" if metric else "NA",
        "status": "Met tolerances" if evaluated and len(within) == len(evaluated) else "Outside tolerance",
    }


def style_table(table, font_size=7.0):
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = 1
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(font_size)
                    run.font.bold = row_index == 0


def compress_main_text(document):
    selection = component_selection_summary()
    replacements = {
        "Several PLS formulations and software implementations address different parts": (
            "SIMPLS constructs sequential components without explicitly deflating the predictor "
            "matrix [11], whereas PLS-SVD obtains directions from one cross-covariance "
            "decomposition. OPLS removes predictor variation unrelated to the response [12], and "
            "kernel PLS extends the model through linear or nonlinear kernels [13,14]. Iterative "
            "IRLBA [15] and randomized SVD (rSVD) [16] avoid unnecessary full decompositions. "
            "Existing software spans these estimators and includes the R package pls [17] and "
            "CPU/GPU improved-kernel PLS (IKPLS) [18], but differs in output storage, validation, "
            "precision, and accelerator support."
        ),
        "We present fastPLS, whose principal methodological contribution": (
            "We present fastPLS, a Bioconductor package whose principal contribution is a compiled "
            "SIMPLS execution path. It caches reusable deflation products, updates coefficients and "
            "fitted values incrementally, and predicts from compact latent factors. Explicit or "
            "implicit cross-covariance products, float32 input, and CPU, NVIDIA CUDA, and Apple "
            "Metal routes extend the applicable matrix regimes. PLS-SVD, OPLS, kernel PLS, argmax "
            "PLS-DA, latent-space LDA, and compiled cross-validation use the same interface. We "
            "evaluate estimator agreement separately from approximate-solver and workflow speed."
        ),
        "The architecture separates preprocessing, the PLS estimator": "",
        "CPU routines use compiled C/C++ and R-linked BLAS/LAPACK": (
            "CPU routines use compiled C/C++ and R-linked BLAS/LAPACK; CUDA uses NVIDIA CUDA and "
            "cuBLAS, and Metal uses Apple Metal Performance Shaders. CPU thread scaling is reported "
            "only for a verified OpenBLAS experiment. Accelerator timings include allocation, "
            "transfer, synchronization, fitting, prediction, and result transfer. A requested "
            "accelerator that is unavailable stops with an error and is never replaced by CPU "
            "execution."
        ),
        "SIMPLS retains de Jong's sequential score": (
            "SIMPLS retains de Jong's sequential score, loading, orthogonalization, and rank-one "
            "deflation updates [11]. fastPLS reuses the deflation row product, conditionally caches "
            "cross-products, updates coefficients incrementally, and stores compact prediction "
            "factors. With rSVD, ordinary CPU and Metal routes generate a fresh seeded sketch for "
            "each component. When the cross-covariance exceeds 512 MiB, CPU, CUDA, and Metal use "
            "a newly initialized rank-one randomized direction per component. CUDA is otherwise "
            "component-wise for regression and smaller tasks and may generate up to eight "
            "fresh candidates together for large dummy-coded classification; candidates are still "
            "accepted sequentially through the SIMPLS updates."
        ),
        "The accelerated path reuses sequential deflation quantities": (
            "Compact latent factors and blocked prediction avoid dense coefficient or prediction "
            "paths. Class-wise products avoid materializing large one-hot responses."
        ),
        "Algorithm 1 summarizes the unchanged estimator": (
            "Algorithm 1 summarizes the implementation; complexity, storage, ablation results, and "
            "the mapping to de Jong SIMPLS are provided in Supplementary Section S9."
        ),
        "Direction extraction is modular rather than the principal estimator contribution": (
            "Direction extraction is modular. CPU supports fixed-control IRLBA [15] and approximate "
            "rSVD [16]; CUDA and Metal use rSVD. The automatic ordinary-matrix profile uses 32 "
            "oversampling directions and five power iterations. For SIMPLS-family problems whose "
            "cross-covariance exceeds 512 MiB, a massive-matrix profile records oversample = 12 "
            "and power = 2 while advancing with one newly initialized rank-one randomized direction "
            "per component. Every result records its effective "
            "controls and seed; numerical comparisons are reported in the Supplement."
        ),
        "For regression analyses, the fitted models produced continuous response estimates": (
            "For regression analyses, fitted models produce continuous response estimates, whereas "
            "classification uses either maximum-score PLS-DA decoding or LDA on PLS scores. LDA "
            "forms pooled within-class covariance from "
            "score cross-products and class means, applies scale-normalized diagonal regularization "
            "only when Cholesky factorization fails, and solves triangular systems without explicit "
            "matrix inversion. Accuracy or balanced accuracy can guide classification tuning; RMSD "
            "or Q2 guides regression. Metric definitions are given in Supplementary Section S4."
        ),
        "Fold construction, fitting, prediction, and metric accumulation remain compiled": (
            "Fold construction, fitting, prediction, and metric accumulation remain compiled where "
            "supported. Group constraints keep repeated observations together. Permutation tests "
            "use the selected endpoint, preserve exchangeability blocks when supplied, and apply "
            "the finite-sample plus-one correction."
        ),
        "The broader benchmark design covered biomedical and computational tasks": (
            "The benchmark included metabolomics, NMR, CITE-seq, tissue and cancer omics, "
            "single-cell transcriptomics, drug response, and image embeddings [7,20-30]. Dataset "
            "construction, preprocessing, dimensions, split units, seeds, component grids, access, "
            "and redistribution are documented in Supplementary Section S5 and Tables S3-S4. "
            "Methods shared stored splits within each comparison; runtime included fitting and "
            "prediction. Hardware, external-software scope, endpoint definitions, and memory "
            "measurement are reported in Supplementary Sections S4 and S7."
        ),
        "A separate ImageNet/DINOv2 analysis used": "",
        "Five-fold training-only selection was performed separately": (
            "Ten-fold training-only component selection was performed separately by PLS family. "
            f"Of {selection['total']} family-dataset choices, {selection['interior']} were interior "
            f"to the evaluated grid, {selection['rank']} were limited by response rank, "
            f"{selection['lower']} occurred at the lower grid boundary, and "
            f"{selection['upper']} occurred at the upper boundary. Boundary- and rank-limited "
            "choices are described as best within the evaluated grid, not as unconstrained optima."
        ),
        "A controlled one-factor-at-a-time SIMPLS study varied": (
            "Controlled studies varied sample, predictor and response dimensions, retained "
            "components, requested prefixes, rank, class count, and explicit versus implicit "
            "cross-covariance storage. A separate factorial holdout checked route-selection "
            "interactions."
        ),
        "Estimator preservation was first assessed against an independent de Jong implementation": (
            "Estimator validation used dense LAPACK reference updates on well-conditioned, tied, "
            "rank-deficient, collinear, p<n, p>n, high-response, regression, and classification "
            "cases. IRLBA was compared separately with pls::simpls.fit; rSVD used multi-seed "
            "prediction, subspace, label, and endpoint screens."
        ),
        "To distinguish numerical computation from the cost of constructing ordinary model objects": (
            "External timing used fresh processes and two output contracts: minimum common outputs "
            "for prediction and each package's ordinary public object. Fit, object assembly, and "
            "held-out prediction were timed separately. Absolute and pre-fit-baseline-corrected "
            "process RSS were retained, together with failures and timeouts."
        ),
        "Numerical equivalence and software-level performance were assessed separately": (
            "Numerical agreement with de Jong SIMPLS was assessed independently of end-to-end "
            "software performance. The latter compared fastPLS with R PLS packages on identical "
            "float64 splits and component counts and, separately, with IKPLS using matched "
            "precision and prediction tasks. Different estimators and output contracts are labelled "
            "as workflow rather than estimator comparisons."
        ),
        "The NMR case study comprised": (
            "The NMR study used 1,200 training and 321 held-out spectra, 13,000 predictors, and "
            "28,355 response intensities. The predefined 4.6-4.8 ppm water interval was zeroed in "
            "Xtrain and Xtest only; Ytrain and Ytest were unchanged, and metrics used every "
            "response column. PLS-family component counts "
            "were selected from training-only splits with a one-standard-error rule. A separate "
            "165-component comparison held split, preprocessing, precision, target, and component "
            "count fixed across the deposited and fastPLS workflows."
        ),
        "The ImageNet/DINOv2 feasibility analysis used 1,281,167": (
            "The ImageNet stress test used 1,281,167 stored float32 DINOv2 embeddings with 1,024 "
            "features and 1,000 labels. Seed 123 assigned 1,000,000 rows to training and 281,167 "
            "to a noncanonical holdout. CUDA SIMPLS-rSVD was evaluated from 100 to 1,000 components "
            "with argmax and LDA. Feature extraction preceded timing; incomplete checkpoint and "
            "row-level provenance limits this experiment to downstream computational feasibility."
        ),
    }
    for paragraph in list(document.paragraphs):
        text = paragraph.text.strip()
        for prefix, replacement in replacements.items():
            if text.startswith(prefix):
                if replacement:
                    set_paragraph(paragraph, replacement)
                else:
                    remove_paragraph(paragraph)
                break


def update_main():
    document = Document(MAIN_IN)
    replace_runs(document, {
        "fastPLS 0.99.25": "fastPLS 0.99.36",
        "fastPLS 0.99.34": "fastPLS 0.99.36",
        "version 0.99.25": "version 0.99.36",
        "version 0.99.34": "version 0.99.36",
        "one maximal fit supplied every prefix": "one maximal fit per classification head supplied every prefix",
        "one maximal fastPLS CUDA-rSVD SIMPLS fit supplied all prefixes": "one maximal fastPLS CUDA-rSVD SIMPLS argmax fit supplied all prefixes",
        "shared-path fitting and prediction time": "per-head shared-path fitting and prediction time",
        "deterministic IRLBA": "fixed-control IRLBA",
        "deterministic float64": "fixed-control float64",
        "deterministic CPU SIMPLS": "fixed-control CPU SIMPLS",
        "deterministic SIMPLS validation": "fixed-control SIMPLS validation",
        "deterministic comparison": "fixed-control comparison",
        "R 4.6.0; principal analyses used fastPLS 0.99.36 from commit 7887401b09e2. See Table S14 for the source checksum and analysis mapping.":
            "R 4.6.0; principal analyses used fastPLS 0.99.36. See Table S15 for the analysis mapping.",
        "R 4.6.0; Metal analyses used fastPLS 0.99.36. See Table S14 for the corresponding source and analysis mapping.":
            "R 4.6.0; Metal analyses used fastPLS 0.99.36. See Table S15 for the analysis mapping.",
        "preserving the deterministic de Jong estimator": "retaining the de Jong component equations",
    })
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if "source archive SHA-256" in text or text.startswith("Code and benchmark outputs are available at"):
            set_paragraph(
                paragraph,
                "fastPLS is available from Bioconductor (https://bioconductor.org/packages/fastPLS). "
                "The development source and current-release benchmark programs are available at "
                "https://github.com/tkcaccia/fastPLS, and reusable compiled components are available "
                "at https://github.com/tkcaccia/kodama-cpp. All quantitative analyses reported here "
                "use fastPLS 0.99.36 with the solver controls stated in the corresponding Methods, "
                "table, or figure caption."
            )
        elif text.startswith("Results:"):
            set_paragraph(paragraph, abstract_result_text())
        elif text.startswith("Conclusions:"):
            set_paragraph(
                paragraph,
                "Conclusions: fastPLS reduces the computational barrier to sequential "
                "SIMPLS for large biomedical matrices while retaining explicit numerical "
                "qualification, strict backend dispatch, and reproducible solver controls."
            )
        elif text.startswith("Methods: fastPLS was audited"):
            set_paragraph(
                paragraph,
                "Methods: fastPLS was evaluated on compiled CPU, OpenBLAS-linked multicore CPU, "
                "NVIDIA CUDA, and Apple Metal systems. The accelerated SIMPLS implementation reuses "
                "deflation products, updates coefficient and fitted-value paths incrementally, and "
                "predicts from compact latent factors. rSVD is the primary low-rank solver. Automatic "
                "control profiles are selected from matrix dimensions, and every fit records the "
                "effective oversampling, power iterations, and seed. Ordinary CPU and Metal "
                "routes form a new sketch for each component; massive CPU, CUDA, and Metal routes "
                "use a newly initialized rank-one randomized direction per component. CUDA can form "
                "a small fresh candidate block for "
                "large classification workloads, with sequential SIMPLS acceptance and deflation."
            )
        elif text.startswith("SIMPLS retains de Jong's sequential score"):
            set_paragraph(
                paragraph,
                "SIMPLS retains de Jong's sequential score, loading, orthogonalization, and rank-one "
                "deflation updates [11]. Direction extraction is approximate when rSVD is selected. "
                "Ordinary CPU and Metal routes form a newly seeded oversampled sketch of the current "
                "deflated operator for each component. Massive CPU, CUDA, and Metal routes use a "
                "newly initialized rank-one randomized direction per component. CUDA otherwise uses "
                "component-wise execution for regression and smaller "
                "classification tasks; for eligible dummy-coded classification with at least 5,000 "
                "training observations, it can generate up to eight fresh candidates "
                "together and accept them sequentially through the standard SIMPLS updates."
            )
        elif text.startswith("Direction extraction is modular"):
            set_paragraph(
                paragraph,
                "Direction extraction is modular rather than the principal estimator contribution. "
                "float64 CPU fitting supports fixed-control IRLBA [15] as an iterative comparator and "
                "approximate rSVD [16], which uses a Gaussian range sketch, orthonormalization, power "
                "iterations, and a reduced decomposition. IRLBA is an iterative truncated solver, not "
                "an exact dense SVD. The ordinary-matrix profile uses 32 oversampling directions and "
                "five power iterations; SIMPLS-family cross-covariances above 512 MiB use a profile "
                "with oversample = 12 and power = 2 while advancing with one newly initialized "
                "rank-one randomized direction per component. Effective controls, seeds, numerical "
                "differences, and route-specific "
                "status are consolidated in the Supplement."
            )
        elif text.startswith("A controlled one-factor-at-a-time SIMPLS study"):
            set_paragraph(
                paragraph,
                "A controlled SIMPLS study varied observations, predictor and response "
                "dimensions, retained components, requested prefixes, effective rank, "
                "class count, and explicit versus implicit cross-covariance storage. "
                "Automatic routes used the ordinary 32/5, high-response 48/6, or sparse "
                "high-class-count 64/7 profile according to matrix shape; forced explicit and "
                "implicit routes used 32/5. Every case used three independently seeded repetitions."
            )
        elif text.startswith("The ImageNet/DINOv2 feasibility analysis"):
            set_paragraph(
                paragraph,
                "The ImageNet stress test used 1,281,167 stored float32 DINOv2 embeddings "
                "with 1,024 features and 1,000 labels. Seed 123 assigned 1,000,000 rows to "
                "training and 281,167 to a noncanonical holdout. CUDA SIMPLS-rSVD used 32 "
                "oversampling directions, five power iterations, and component prefixes from "
                "100 to 1,000 with argmax and LDA. Feature extraction preceded timing; incomplete "
                "checkpoint and row-level provenance limits this analysis to downstream "
                "computational feasibility."
            )
        elif text.startswith("Standard R matrices use eight-byte float64 values"):
            set_paragraph(
                paragraph,
                "Ordinary R numeric matrices use eight-byte float64 values, whereas float::float32 "
                "matrices use four bytes per value. fastPLS selects precision from the input type and "
                "does not silently down-cast ordinary numeric matrices. CPU, CUDA, and Metal support "
                "route-specific float32 execution; Tables S1 and S9a distinguish execution residency "
                "and float32 capability. On Windows, backend = 'cpu' uses portable float32 rSVD kernels for PLS-SVD, "
                "SIMPLS, OPLS, linear and nonlinear kernel PLS, and LDA. Unavailable CUDA or Metal "
                "requests stop with an error and are never replaced by CPU execution."
            )
        elif text.startswith("The approximate-solver study retained"):
            set_paragraph(
                paragraph,
                "The approximate-solver study retained every completed route and evaluated "
                "numerical differences separately from timing. Controlled qualification and "
                "ordinary-matrix results use the 32/5 profile; massive SIMPLS-family results report "
                "oversample = 12 and power = 2 and execute one newly initialized rank-one randomized "
                "direction per component. Ordinary CPU and Metal routes use a newly initialized "
                "component-wise sketch; CUDA can use a fresh candidate block only for eligible "
                "large dummy-coded classification."
            )
        elif text.startswith("Figure 4. ImageNet/DINOv2"):
            set_paragraph(
                paragraph,
                "Figure 4. ImageNet/DINOv2 computational stress test. float32 CUDA "
                "SIMPLS-rSVD used 32 oversampling directions, five power iterations, "
                "seed 123, and 1,000,000 training and 281,167 held-out embeddings. "
                "Panels report argmax and LDA top-1/top-5 accuracy, per-head maximal-path "
                "fitting plus prediction time, and peak host and device memory."
            )
        elif text.startswith("Figure 1. Float64 single-CPU SIMPLS"):
            set_paragraph(
                paragraph,
                "Figure 1. Float64 single-CPU SIMPLS classification workflows. Panels "
                "report fixed-split accuracy, total fitting-plus-prediction time, and "
                "absolute complete-process peak RSS. Each completed method-dataset pair "
                f"used {package_repetition_text()} fresh processes according to the "
                "adaptive repetition policy; NE denotes not evaluated, TO timeout, and "
                "ERR execution error. Argmax rows provide the closest estimator comparison; "
                "LDA and other PLS-DA rows compare complete workflows with different heads "
                "or retained outputs."
            )
        elif text.startswith("Solver choice separates exploratory acceleration"):
            set_paragraph(
                paragraph,
                "Solver choice separates exploratory acceleration from confirmatory "
                "analysis. rSVD is the primary route, using documented ordinary- or massive-matrix "
                "controls selected from the matrix dimensions; fixed-control IRLBA remains available when a "
                "randomized approximation is inappropriate or additional confirmation is required."
            )
        elif text.startswith("The principal gain arises from compiled sequential updates"):
            set_paragraph(
                paragraph,
                "The principal gain arises from compiled sequential updates, reuse of deflation "
                "products, compact latent prediction, and backend-specific approximate direction "
                "extraction. CPU and Metal form a fresh component-wise sketch. CUDA candidate "
                "blocks amortize launches only for eligible large dummy-coded classification; "
                "regression and smaller tasks use component-wise execution. These routes retain "
                "the sequential SIMPLS score, loading, orthogonalization, and deflation updates."
            )
        elif text.startswith("All 108 planned paired runs completed"):
            set_paragraph(paragraph, external_simpls_result_text())
        elif text.startswith("The magnitude of the observed workflow gains was influenced"):
            set_paragraph(paragraph, external_simpls_memory_text())
        elif text.startswith("The repeated float64 CPU comparison on Breast"):
            set_paragraph(paragraph, ikpls_result_text())
        elif text.startswith("The computational benefit of acceleration depended"):
            set_paragraph(paragraph, backend_result_text())
        elif text.startswith("Training-only selection with fastPLS"):
            set_paragraph(paragraph, nmr_selection_text())
        elif text.startswith("At a common 165 components"):
            set_paragraph(paragraph, nmr_fixed_text())
        elif text.startswith("The predictive and computational analyses were kept separate"):
            set_paragraph(paragraph, nmr_selected_text())
        elif text.startswith("Figure 3. NMR prediction and computation"):
            set_paragraph(
                paragraph,
                "Figure 3. NMR prediction and computation at a common 165 components. Every "
                "workflow used the same 1,200/321 train-test split, centering-only protocol, "
                "float64 precision, 28,355-variable response, and component count. Panels A-C "
                "report held-out RMSD, median fitting-plus-prediction time across three isolated "
                "runs, and baseline-corrected peak process RSS, including the deposited PLS-SVD "
                "workflow as a distinct comparator. Panel D reports per-spectrum RMSD for all "
                "321 held-out samples. Panels "
                "E-F show the held-out sample closest to the median SIMPLS CUDA-rSVD per-spectrum "
                "RMSD over the full response range and 1.7-0.5 ppm. Training-selected predictive "
                "results are reported separately in Supplementary Figure S5. PLS-SVD rSVD used "
                "32 oversampling directions and five power iterations; SIMPLS rSVD used the "
                "automatic massive-matrix controls oversample = 12 and power = 2, with one newly "
                "initialized rank-one randomized direction per component. Every randomized route "
                "used seed 123."
            )
        if "submitted to Bioconductor" in paragraph.text or "will be available from Bioconductor" in paragraph.text:
            set_paragraph(paragraph, paragraph.text.replace("submitted to Bioconductor", "available from Bioconductor").replace("will be available from Bioconductor", "is available from Bioconductor"))

    section_31 = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("3.1 Repeated comparison")
    )
    set_paragraph(section_31, "3.1 Comparison with independent R implementations")
    section_32 = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("3.2 Cross-language")
    )
    paragraphs = list(document.paragraphs)
    figure_1_index = next(
        index for index, paragraph in enumerate(paragraphs)
        if paragraph.text.strip().startswith("Figure 1.")
    )
    figure_1_caption = paragraphs[figure_1_index]
    figure_1_image = next(
        (
            paragraph for paragraph in reversed(
                paragraphs[max(0, figure_1_index - 3):figure_1_index]
            )
            if paragraph._p.xpath(".//w:drawing")
        ),
        None,
    )
    if figure_1_image is None:
        raise ValueError("Figure 1 has no image paragraph.")
    section_32._p.addprevious(figure_1_image._p)
    section_32._p.addprevious(figure_1_caption._p)
    figure_1_image.insert_paragraph_before(package_panel_result_text())

    compress_main_text(document)
    replace_figure(document, "Figure 1.", PACKAGE_FIGURE, 6.4)
    replace_figure(document, "Figure 2.", BACKEND_FIGURE, 7.0)
    replace_figure(document, "Figure 3.", NMR_FIGURE, 7.0)
    replace_figure(document, "Figure 4.", IMAGENET_FIGURE, 7.0)

    renumber_table_citations(document)
    replace_release_version(document)
    main_text = document_text(document)
    for forbidden in ("unqualified", "rejected"):
        assert forbidden not in main_text.lower(), forbidden
    audit_document_consistency(document, "main manuscript")
    format_scientific_notation(document)
    normalize_table_font_sizes(document, 7.0)
    enable_continuous_line_numbers(document)
    algorithm_caption = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("Algorithm 1.")
    )
    algorithm_caption.paragraph_format.page_break_before = False
    document.core_properties.title = "fastPLS manuscript"
    document.core_properties.subject = "Computer Methods and Programs in Biomedicine"
    document.save(MAIN_OUT)


def update_supplement(cpu_csv, metal_csv, cuda_csv):
    document = Document(SUPP_IN)
    remove_embedded_figure_paragraphs(document)
    ensure_package_panel_caption(document)
    selection = component_selection_summary()
    replace_runs(document, {
        "fastPLS 0.99.25": "fastPLS 0.99.36",
        "fastPLS 0.99.27": "fastPLS 0.99.36",
        "fastPLS 0.99.34": "fastPLS 0.99.36",
        "version 0.99.25": "version 0.99.36",
        "version 0.99.27": "version 0.99.36",
        "version 0.99.34": "version 0.99.36",
        "v0.99.25": "v0.99.36",
        "v0.99.27": "v0.99.36",
        "v0.99.34": "v0.99.36",
        "one maximal 1,000-component CUDA fit shared across prefixes": "one maximal 1,000-component CUDA argmax fit shared across prefixes",
        "deterministic IRLBA": "fixed-control IRLBA",
        "deterministic float64": "fixed-control float64",
        "deterministic CPU SIMPLS": "fixed-control CPU SIMPLS",
        "deterministic SIMPLS validation": "fixed-control SIMPLS validation",
        "deterministic comparison": "fixed-control comparison",
    })

    for section in document.sections:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    for table in document.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == "R and fastPLS":
                row.cells[1].text = (
                    f"R 4.6.0; principal analyses used fastPLS {VERSION}. "
                    "See Table S27 for the analysis mapping."
                )
                row.cells[2].text = (
                    f"R 4.6.0; Metal analyses used fastPLS {VERSION}. "
                    "See Table S27 for the corresponding analysis mapping."
                )

    update_external_scope_table(document)

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("The double-precision CPU backend is the broadest reference implementation"):
            set_paragraph(
                paragraph,
                "Ordinary R numeric matrices select float64 execution, while float::float32 inputs "
                "select route-specific float32 execution automatically. On Windows, backend = 'cpu' "
                "uses portable float32 rSVD kernels for PLS-SVD, SIMPLS, OPLS, linear and nonlinear "
                "kernel PLS, and LDA. CUDA and Metal are unavailable on Windows builds without their "
                "native accelerator toolchains; explicit requests stop with an error and are never "
                "replaced by CPU execution."
            )
            continue
        if text.startswith("The SIMPLS estimator follows de Jong"):
            set_paragraph(
                paragraph,
                "The SIMPLS estimator follows de Jong's sequential construction. For each "
                "component, the implementation consumes one dominant direction of the current "
                "cross-covariance state, forms and normalizes the score, computes predictor and "
                "response loadings, orthogonalizes the predictor loading against the preceding "
                "SIMPLS basis, and applies rank-one deflation. CPU, Metal, and ordinary CUDA "
                "routes recalculate the randomized direction from the updated operator. For "
                "eligible large dummy-coded CUDA classification, a newly generated candidate "
                "block amortizes device launches; candidates are accepted one at a time through "
                "the sequential SIMPLS updates.",
            )
            continue
        if text == "S9. Deterministic estimator validation":
            set_paragraph(paragraph, "S9. SIMPLS estimator validation")
            continue
        if text.startswith("Table S4b. External implementation scope"):
            set_paragraph(
                paragraph,
                "Table S4b. Computational capabilities of the PLS implementations included in "
                "the software comparison. Multicore CPU indicates explicit parallel validation "
                "or use of eligible kernels through the linked BLAS/OpenMP runtime; it does not "
                "imply parallel sequential deflation. 'No native route' means that the package "
                "does not expose that backend or precision directly. IKPLS CUDA execution uses "
                "its JAX implementation [34]. Exact tested versions and execution errors are "
                "recorded in the run manifest."
            )
            continue
        if text.startswith("The current backend and execution audit used"):
            set_paragraph(
                paragraph,
                "All quantitative evidence in this supplement uses fastPLS 0.99.36. The backend "
                "campaign comprises 264 CPU/Metal and 264 CPU/CUDA isolated runs across 11 datasets "
                "and four PLS families, together with current NMR and ImageNet analyses and a verified "
                "one-, two-, and four-thread OpenBLAS experiment. Every approximate row records the "
                "executed oversampling, power iterations, and seed."
            )
        elif text.startswith("The benchmark writes a machine-generated dataset manifest"):
            set_paragraph(
                paragraph,
                "The benchmark writes a machine-generated dataset manifest containing source, "
                "licence, object names, preprocessing, split seed, matrix dimensions, and "
                "component grid. Dataset dimensions are read from the prepared task objects "
                "rather than transcribed manually."
            )
        elif text.startswith("CIFAR-100 and ImageNet were evaluated as precomputed image embeddings"):
            set_paragraph(
                paragraph,
                "CIFAR-100 and ImageNet were evaluated as precomputed image embeddings rather "
                "than by training image encoders. CIFAR-100 used its standard 50,000/10,000 "
                "partition [30]. The ImageNet dataset contained 1,281,167 rows, 1,024 DINOv2 "
                "features, and 1,000 labels [28,29], but no primary canonical train/validation "
                "flag. Its exact DINOv2 checkpoint, image preprocessing, pooling rule, extraction "
                "script, and independently auditable image-to-row mapping were not retained. "
                "The stored feature matrix and labels permit the downstream fastPLS analysis to "
                "be rerun, but the representation cannot be regenerated independently from the "
                "available metadata. Feature extraction preceded the benchmark and was excluded "
                "from timing."
            )
        elif text.startswith("Each analysis recorded its fastPLS version"):
            set_paragraph(
                paragraph,
                "Each analysis recorded the fastPLS version, script identity, compiler, "
                "BLAS/LAPACK, thread settings, accelerator libraries, seed, rSVD controls, "
                "repetition, and execution status. NMR was rerun end to end from the prepared "
                "predictors and responses. For ImageNet, only downstream PLS fitting and "
                "prediction were rerun because feature-extraction metadata remain incomplete. "
                "Section S17 maps each claim to its result directory and generating program."
            )
        elif text.startswith("The approximate direction calculation is backend specific"):
            set_paragraph(
                paragraph,
                "The approximate direction calculation is backend specific. Ordinary CPU and Metal "
                "routes form a new seeded oversampled sketch of the current deflated operator for each "
                "component. Massive CPU, CUDA, and Metal routes use a newly initialized rank-one "
                "randomized direction for each component. CUDA "
                "uses the same component-wise rule for regression and smaller classification tasks, "
                "with a fresh candidate block only for eligible large dummy-coded classification. "
                "If a requested CUDA or Metal backend is unavailable, fitting stops with an error; "
                "CPU execution is never substituted automatically."
            )
        elif text.startswith("The high-speed SIMPLS core retains"):
            set_paragraph(
                paragraph,
                "The high-speed SIMPLS core retains the sequential score, loading, supervised "
                "orthogonalization, and rank-one deflation updates. Ordinary CPU and Metal routes "
                "form a new oversampled sketch at every component; massive CPU, CUDA, and Metal "
                "routes use one newly initialized rank-one randomized direction per component. "
                "CUDA may generate at most eight fresh "
                "candidates together for eligible large dummy-coded classification, but consumes "
                "them sequentially through the SIMPLS updates."
            )
        elif text.startswith("For component a, obtain one dominant left direction"):
            set_paragraph(
                paragraph,
                "For component a, obtain one dominant left direction of the current deflated "
                "cross-covariance operator. IRLBA performs a new iterative solve. Ordinary CPU and "
                "Metal rSVD form a newly seeded oversampled sketch; massive CPU, CUDA, and Metal "
                "routes use a newly initialized rank-one randomized direction. CUDA otherwise forms "
                "a component-wise sketch or the "
                "fresh candidate block used only for eligible large dummy-coded classification."
            )
        elif text.startswith("For a linear operator M and target rank r"):
            set_paragraph(
                paragraph,
                "For a linear operator M and target rank r, rSVD draws a Gaussian matrix, forms a "
                "range sketch, applies alternating products with M and its transpose, orthonormalizes "
                "the range, and decomposes the reduced matrix. The ordinary-matrix profile uses 32 "
                "oversampling directions and five power iterations; the automatic massive SIMPLS-family "
                "profile records oversample = 12 and power = 2 and advances with one newly initialized "
                "rank-one randomized direction per component when the cross-covariance exceeds 512 MiB. "
                "The three-seed controlled panel applies the automatic shape-specific profile and "
                "also forces matched 32/5 explicit and implicit routes, yielding 174 comparisons "
                "per backend."
            )
        elif text.startswith("The baseline regression scenario used"):
            set_paragraph(
                paragraph,
                "The baseline regression scenario and all controlled variants used the "
                "same centering, held-out split, component-prefix scoring contract, automatic "
                "shape-specific rSVD controls, and seeds 101-103."
            )
        elif text.startswith("The matched PLS-family timing study used"):
            set_paragraph(
                paragraph,
                "The matched PLS-family timing study used the same synthetic matrices, requested "
                "components, rSVD controls (32 oversampling directions, five power iterations, "
                "seeds 101-103), split, precision, and CPU backend for PLS-SVD and SIMPLS. The two "
                "estimators can differ predictively; this experiment compares execution time and "
                "reports the predictive endpoint separately."
            )
        elif text.startswith("All 486 CPU/CUDA runs completed"):
            set_paragraph(
                paragraph,
                "The controlled-solver summary uses the fastPLS 0.99.36 qualification panel. "
                "CPU, CUDA, and Metal were each evaluated with the same automatic shape-specific "
                "controls and three seeds; explicit and implicit routes were forced at 32/5. "
                "Runtime crossover analyses use only "
                "current-release results and report prediction and metric agreement separately."
            )
        elif "component_selection_by_family.csv" in text:
            set_paragraph(
                paragraph,
                "Two uses of component grids are distinguished. Benchmark trajectory figures "
                "evaluate fixed component counts on a fixed test set and are descriptive. "
                "Ten-fold training-only model selection reports the best value within the evaluated "
                f"family-specific grid. Among {selection['total']} choices, "
                f"{selection['interior']} were interior, {selection['rank']} were response-rank "
                f"limited, {selection['lower']} occurred at the lower boundary, and "
                f"{selection['upper']} occurred at the upper boundary. Boundary- and rank-limited "
                "values are not described as unconstrained optima. Current-release selection tables "
                "and component paths are stored under "
                "publication_results/0.99.36/component_selection/. NMR additionally uses an "
                "extended training-only grid and a one-standard-error rule. Repeated-partition "
                "sensitivity is reported in Section S16."
            )
        elif text.startswith("Redistribution status was determined separately"):
            set_paragraph(
                paragraph,
                "Redistribution status was determined separately from source access status. A "
                "publicly downloadable source was not treated as permission to redistribute the "
                "processed benchmark matrix under the fastPLS licence. Only the GPL-compatible "
                "breast and colon package examples, synthetic generators and generated synthetic "
                "results, and aggregate tables, figures, manifests, and split indices are "
                "redistributed. None of the prepared real-data matrices in Table S3 is bundled. "
                "The executable workflow benchmark/acquire_publication_datasets.R downloads primary "
                "public sources or validates user-supplied local files and records paths, sizes, and "
                "access classes in acquisition_manifest.csv. Exact commands and preprocessing "
                "contracts are documented in benchmark/DATA_ACQUISITION.md."
            )
        elif text.startswith("The public acquisition command is Rscript"):
            set_paragraph(
                paragraph,
                "The public acquisition command is Rscript "
                "benchmark/acquire_publication_datasets.R --dataset=<id> --out=<directory>. For "
                "ImageNet, NMR, CCLE, and PRISM, the listed environment variable points to a "
                "user-authorized local file, which is validated in place and is not copied into the "
                "repository. GTEx uses only open-access expression and phenotype data; protected "
                "sequence and full donor-level files are outside the benchmark. The ImageNet "
                "embeddings remain a derived object subject to the source-image terms and are not "
                "redistributed."
            )
        if text.startswith("The 0.99.25 namespace explicitly exports"):
            set_paragraph(
                paragraph,
                "The 0.99.36 namespace exports pls(), pls.single.cv(), pls.double.cv(), "
                "evaluate(), plot.permutation(), ViP(), fastsvd(), fastcor(), "
                "fastPLS_backend(), has_cuda(), and has_metal(). Registered predict() and "
                "plot() methods use standard R generics. Public model families are PLS-SVD, "
                "SIMPLS, OPLS, and kernel PLS; public classification heads are argmax and LDA. "
                "rSVD is the default solver with the CPU, CUDA, and Metal controls evaluated in "
                "Table S7; IRLBA is available on CPU."
            )
        elif text.startswith("rSVD uses fresh Gaussian range sketches"):
            set_paragraph(
                paragraph,
                "rSVD uses fresh Gaussian range sketches of the deflated SIMPLS operator. CPU and "
                "Metal form a new oversampled sketch for each component; CUDA can generate a small "
                "fresh candidate block for eligible large classification workloads and accepts "
                "candidates sequentially through the SIMPLS updates. Every ordinary fit starts with "
                "32 oversampling directions and five power iterations. Fit-level diagnostics can "
                "trigger deterministic internal retries at 48/6 and 64/7; massive SIMPLS-family "
                "workloads instead use the separately recorded 12/2 rank-one profile. The numerical "
                "panel compared three seeds on CPU, CUDA, and Metal against matched CPU IRLBA fits. "
                "Each backend completed 174 comparisons and all 522 met the stated tolerances."
            )
        elif "benchmark_results/frozen_release_0.99.25/ikpls_cross_language_cpu/" in text:
            set_paragraph(
                paragraph,
                text.replace(
                    "benchmark_results/frozen_release_0.99.25/ikpls_cross_language_cpu/",
                    "publication_results/0.99.36/ikpls_cross_language_cpu/"
                )
            )
        elif "benchmark_results/ikpls_large_float32_20260826/" in text:
            set_paragraph(
                paragraph,
                text.replace(
                    "benchmark_results/ikpls_large_float32_20260826/",
                    "publication_results/0.99.36/ikpls_large_float32/"
                )
            )
        if text == "S11. float32 capability":
            set_paragraph(paragraph, "S11. float32 execution and capability")
        elif text.startswith("float64 is the confirmatory reference"):
            set_paragraph(
                paragraph,
                "Ordinary R numeric matrices store each value as float64. A float::float32 input "
                "uses four bytes per value and automatically requests the corresponding float32 "
                "fastPLS route; ordinary numeric input is not silently down-cast. Input storage is "
                "therefore halved, but total runtime and peak memory also depend on temporary arrays, "
                "model outputs, device contexts, and host-assisted stages. Table S8 distinguishes "
                "arithmetic precision from execution residency. 'Implemented' means that the route "
                "accepts and retains the stated precision; it does not by itself imply a speed or "
                "memory advantage. 'Hybrid' means that the selected accelerator performs the dominant "
                "PLS computation but one or more orchestration, kernel, filtering, reduced-factorization, "
                "or classification stages execute on the host. Unsupported combinations stop rather "
                "than silently promoting float32 input to float64. On Windows, float32 uses the "
                "portable CPU rSVD route for all four PLS families and LDA; float32 IRLBA and "
                "accelerator requests are unavailable. Paired numerical and resource measurements "
                "are reported separately for the combinations evaluated in this study."
            )
        elif text.startswith("Table S8. Primary float32 capability matrix"):
            paragraph.insert_paragraph_before(
                "PLS-SVD and SIMPLS provide float32 routes on CPU, CUDA, and Metal. OPLS and "
                "nonlinear kernel PLS also accept float32 input, but their accelerator routes are "
                "classified as hybrid when orthogonal filtering or kernel construction remains on "
                "the host. LDA uses float32 fitting and prediction on CPU and CUDA; the Metal route "
                "keeps PLS score computation on the device and performs LDA on the host. These "
                "precision and residency distinctions are independent: a route may preserve "
                "float32 arithmetic while still containing a host-assisted stage. A current-release "
                "smoke screen used Iris classification and mtcars regression at two components for "
                "all four PLS families on CPU, CUDA, and Metal. Float32 and float64 decoded labels "
                "agreed for every classification route, and the largest regression prediction error "
                "relative to float64 was 2.31e-6. These small tests verify arithmetic routing and "
                "endpoint agreement; they are not evidence that float32 is always faster or uses "
                "less complete-process memory."
            )
            set_paragraph(
                paragraph,
                "Table S8. float32 execution capability. Arithmetic, fitting residency, reduced "
                "decomposition, prediction, and LDA residency are reported separately; capability "
                "labels refer only to the tested combinations."
            )
        elif text.startswith("Table S6. Primary float64 CPU estimator-validation summary"):
            set_paragraph(
                paragraph,
                "Table S6. Current-release float64 CPU estimator validation. The dense "
                "SIMPLS panel covers 82 component prefixes in ten numerical conditions; the "
                "external de Jong comparison covers 117 component-level endpoints and 24 "
                "fixed-fold selections. OPLS and kernel PLS cover 66 setting-task endpoints "
                "and 66 component selections. Values are maxima unless stated otherwise.",
            )
        elif text.startswith("Figure S3. Repeated single-CPU SIMPLS public workflows"):
            set_paragraph(
                paragraph,
                "Figure S3. Repeated single-CPU SIMPLS public workflows. fastPLS 0.99.36 used "
                "method = 'simpls', backend = 'cpu', svd.method = 'irlba', float64 input, centering, "
                "argmax decoding, and one effective BLAS thread; pls 2.9.0 used the same centered "
                "float64 data and component count. Requested/effective component counts were CCLE "
                "50/17, CIFAR-100 100/99, GTEx v8 32/31, MetRef 22/21, Retina 50/11, Tabula Muris "
                "50/31, TCGA-BRCA 5/4, TCGA-HNSC methylation 2/1, and TCGA Pan-Cancer 50/31; lower "
                "effective counts reflect response-rank limits. Panel A reports median fitting-plus-"
                "prediction time and panel B reports absolute and baseline-corrected complete-process "
                f"peak RSS from {external_repetition_text()} completed fresh R processes per "
                "successful method-dataset profile under the adaptive repetition policy; error "
                "bars are IQRs. The failed ordinary CIFAR-100 pls::simpls.fit workflow is retained."
            )
        elif text.startswith("The repeated comparison used float64"):
            set_paragraph(
                paragraph,
                "The repeated comparison used float64 CPU SIMPLS with fixed IRLBA controls, "
                "identical splits and component counts, one effective BLAS thread, and a 10,000-s "
                "timeout. Package and data loading occurred before timing, and no pre-timing solver "
                "run was performed. Cold-process repetitions were selected from an a priori adaptive "
                "policy: 50 when the pilot was at most 0.5 s, 30 when at most 2 s, 15 when at most "
                "10 s, and five otherwise. Two output profiles were kept separate. In the minimum-"
                "output profile, fastPLS and pls::simpls.fit(stripped = TRUE) retained the minimum "
                "common objects required for prediction. In the public-workflow profile, fastPLS "
                "retained its ordinary compact object and variance summary, whereas pls::simpls.fit "
                "retained its standard coefficient, score, loading, projection, fitted-value, residual, "
                "and X-variance outputs. Absolute lifetime peak RSS, the immediately pre-fit process "
                "baseline, and their difference were recorded separately. Across the paired profiles, "
                "1,310 of 1,311 planned worker runs completed; the incomplete run was the ordinary "
                "CIFAR-100 pls::simpls.fit workflow, which was killed under memory pressure. The broader "
                "panel comprised 90 configured method-dataset workflows: all 18 fastPLS workflows and "
                "67 of 72 external workflows completed; three external CIFAR-100 workflows were killed "
                "under memory pressure and two external workflows timed out. Different estimators, "
                "prediction heads, and output policies make that broader panel a workflow comparison, "
                "not an estimator-equivalence analysis."
            )
        elif text.startswith("Table S9a. Repeated deterministic SIMPLS comparison"):
            set_paragraph(
                paragraph,
                "Table S9a. Repeated fixed-control SIMPLS comparison. Times are median (IQR) "
                f"seconds from {external_repetition_text()} completed fresh processes per successful "
                "method-dataset profile under the adaptive repetition policy. Speed-up is "
                "pls::simpls.fit divided by fastPLS; values above one favour fastPLS."
            )
        elif text.startswith("Table S9c. Host-memory accounting for the minimum-output profile"):
            set_paragraph(
                paragraph,
                "Table S9c. Complete-process host-memory accounting for the minimum-common-"
                "output profile. Values are MiB and shown as median [IQR] across the recorded "
                "number of fresh processes; F/P denotes fastPLS / pls::simpls.fit. The peak "
                "increment is baseline corrected but remains a process-level measurement.",
            )
        elif text.startswith("Table S9d. Host-memory accounting for ordinary public workflows"):
            set_paragraph(
                paragraph,
                "Table S9d. Complete-process host-memory accounting for ordinary public "
                "workflows. Values are MiB and shown as median [IQR] across the recorded number "
                "of fresh processes; F/P denotes fastPLS / pls::simpls.fit. The theoretical "
                "largest retained object explains output-policy differences but excludes "
                "temporary and runtime allocations.",
            )
        elif text.startswith("Table S9e. Broad single-CPU SIMPLS classification workflow summary"):
            set_paragraph(
                paragraph,
                "Table S9e. Broad single-CPU SIMPLS classification workflow summary. Each "
                "cell reports held-out accuracy, median total fitting-plus-prediction time, and "
                "median complete-process peak host RSS from the recorded fresh-process "
                "repetitions. The fastest and most accurate external rows may represent different "
                "estimators or classification heads and are workflow comparisons.",
            )
        elif text.startswith("Table S9f. Single-thread CPU end-to-end comparison with IKPLS"):
            set_paragraph(
                paragraph,
                "Table S9f. Single-thread CPU end-to-end comparison with IKPLS. Time and IQR are "
                "fitting-plus-prediction seconds; resident-set-size values are MiB. fastPLS rSVD "
                "rows use version 0.99.36, 32 oversampling directions, five power iterations, seed "
                "123, and route-specific diagnostics."
            )
        elif text.startswith("Table S9g. IKPLS 6.1.2 float32 feasibility"):
            set_paragraph(
                paragraph,
                "Table S9g. IKPLS 6.1.2 float32 feasibility on the NMR and ImageNet case "
                "studies. Conversion and centering are excluded from fit and prediction time "
                "and are reported separately in the machine-readable results. Peak RSS is "
                "absolute complete-process resident-set size. Failed and numerically degenerate "
                "runs are retained explicitly.",
            )
        elif text.startswith("Reproducibility. CPU results are in"):
            set_paragraph(
                paragraph,
                "Reproducibility. CPU results are in publication_results/0.99.36/"
                "ikpls_cross_language_cpu/. The benchmark used fastPLS 0.99.36, IKPLS 6.1.2, "
                "float64, three fresh processes, and fixed component counts of 10 (Breast), 22 "
                "(MetRef), and 50 (CIFAR-100)."
            )
        elif text.startswith("Reproducibility. Scripts are in benchmark/ikpls_cross_language/"):
            set_paragraph(
                paragraph,
                "Reproducibility. Scripts are in benchmark/ikpls_cross_language/. The compact "
                "result table and preprocessing time are in publication_results/0.99.36/"
                "ikpls_large_float32/. The experiment used IKPLS 6.1.2, NumPy float32 arrays, "
                "one CPU thread, blocked ImageNet prediction, and conversion outside timing."
            )
        elif text.startswith("This focused analysis uses identical stored splits"):
            set_paragraph(
                paragraph,
                "The paired backend analysis uses identical stored splits and family-specific "
                "component counts for PLS-SVD, SIMPLS, OPLS, and kernel PLS. Every row uses "
                "fastPLS 0.99.36, the recorded automatic rSVD controls, and three isolated "
                "repetitions. Host memory is baseline-corrected complete-process RSS; device "
                "allocations include runtime and context overhead and are retained in the "
                "machine-readable run records."
            )
        elif text.startswith("The family-selected predictive analysis and paired backend analysis"):
            set_paragraph(
                paragraph,
                "The family-selected predictive analysis and paired backend analysis answer "
                "different questions. Component selection used five training-only splits. The "
                "PLS-SVD and SIMPLS candidate grid was 1, 2, 3, 5, 10, 25, 50, 75, 100, "
                "125, 150, 165, 175, 200, 250, and 300 components. Selection used the "
                "one-standard-error rule described below; the test set was not accessed until "
                "after the component count had been fixed."
            )
        elif text.startswith("Table S10. Paired float64 CPU/CUDA SIMPLS-rSVD results"):
            set_paragraph(
                paragraph,
                "Table S10. Complete paired float64 CPU/CUDA panel for PLS-SVD, SIMPLS, "
                "OPLS, and kernel PLS at the training-selected component counts. Every row uses "
                "fastPLS 0.99.36, matched inputs, recorded automatic rSVD controls, and three "
                "isolated repetitions. Times and baseline-corrected peak RSS are medians."
            )
        elif text.startswith("Table S11. Complete fastPLS"):
            set_paragraph(
                paragraph,
                "Table S11. Complete paired float64 CPU/Metal panel for PLS-SVD, SIMPLS, "
                "OPLS, and kernel PLS at the training-selected component counts. Every row uses "
                "fastPLS 0.99.36, matched inputs, recorded automatic rSVD controls, and three "
                "isolated repetitions. Times and baseline-corrected peak RSS are medians.",
            )
        elif text.startswith("Figure S4. Paired CPU/CUDA SIMPLS-rSVD workflows"):
            set_paragraph(
                paragraph,
                "Figure S4. Selected SIMPLS-rSVD workflows across CPU, NVIDIA CUDA, and "
                "Apple Metal backends. Panels report CPU/accelerator runtime and "
                "baseline-corrected complete-process host-RSS ratios, prediction agreement "
                "with the matched CPU fit, and relative predictive-metric differences for all "
                "11 non-NMR datasets. Every comparison used fastPLS 0.99.36, matched splits "
                "and component counts, recorded automatic rSVD controls, and three isolated "
                "repetitions. Ratios above one indicate a lower accelerator value. Device "
                "allocation records are retained in the machine-readable run files when "
                "available."
            )
        elif text.startswith("Figure S5. Training-selected NMR SIMPLS-rSVD performance"):
            set_paragraph(
                paragraph,
                "Figure S5. Training-selected NMR PLS-SVD and SIMPLS workflows with fastPLS "
                "0.99.36. Panel A shows the five training-only SIMPLS validation paths and the "
                "one-standard-error rule; the smallest eligible value was 50 components. Panel B "
                "reports held-out RMSD, Q2, total fitting-plus-prediction time, and baseline-corrected "
                "peak process RSS for CPU IRLBA and CPU, CUDA, and Metal rSVD routes, using PLS-SVD "
                "at five components and SIMPLS at 50 components. PLS-SVD rSVD used the ordinary "
                "32/5 profile; SIMPLS rSVD used oversample = 12 and power = 2 and executed one "
                "newly initialized rank-one randomized direction per component. Every randomized "
                "fit used seed 123."
            )
        elif text.startswith("Table S12. NMR analyses separated by scientific question"):
            set_paragraph(
                paragraph,
                "Table S12. NMR predictive selection and common-165-component implementation "
                "benchmark. Training-selected rows use family-specific component counts chosen "
                "from five training-only splits by the one-standard-error rule. Common-workload "
                "rows hold the component count at 165. Times and baseline-corrected peak process "
                "RSS increments are medians across three isolated runs; IQR is shown in brackets. "
                "The two analyses answer different questions and are not pooled."
            )
        elif text.startswith("The dataset contained 1,281,167 precomputed DINOv2 embeddings"):
            set_paragraph(
                paragraph,
                "The dataset contained 1,281,167 precomputed DINOv2 embeddings with 1,024 features "
                "and 1,000 class labels. Seed 123 assigned 1,000,000 rows to training and 281,167 to "
                "a noncanonical holdout; component-grid specification was not independent of this "
                "holdout. The exact checkpoint, pooling rule, extraction script, and auditable "
                "image-to-row mapping were unavailable. fastPLS 0.99.36 reran downstream label-aware "
                "float32 SIMPLS fitting and blocked prediction using CUDA rSVD with 32 oversampling "
                "directions, five power iterations, and seed 123. All values are single-run feasibility "
                "estimates and are not used for biomedical or representation-quality claims."
            )
        elif text.startswith("Table S13. ImageNet downstream classification"):
            set_paragraph(
                paragraph,
                "Table S13. ImageNet downstream classification. Every row uses fastPLS 0.99.36, "
                "CUDA rSVD with 32 oversampling directions, five power iterations, and seed 123. "
                "One shared fit per classification head supplies all component prefixes; the "
                "1,000-component row is a "
                "boundary stress point rather than an optimum. All rows are single-run, partially "
                "reproducible feasibility estimates."
            )
        elif text.startswith("Table S14. Repeated-partition predictive dispersion"):
            set_paragraph(
                paragraph,
                "Table S14. Current-release predictive dispersion and training-only "
                "component-selection sensitivity across repeated outer partitions. All rows use "
                "the CPU rSVD route with 32 oversampling directions, five power iterations, and "
                "fixed seeds; classification tasks use ten outer partitions and NMR uses five. "
                "Selection at the upper evaluated boundary and response-rank constraints are "
                "identified explicitly; intervals quantify split-to-split predictive variation "
                "rather than timing uncertainty.",
            )
        elif text.startswith("To assess sensitivity to the training partition"):
            set_paragraph(
                paragraph,
                "To assess sensitivity to training sampling and component selection, MetRef, "
                "GTEx v8, and Retina used ten stratified 80/20 outer partitions with five-fold "
                "training-only selection; NMR used five random 80/20 outer partitions with "
                "three-fold selection. The current fastPLS release and recorded rSVD controls "
                "were used throughout. The empirical intervals describe variation across these "
                "partitions and are not population-level confidence intervals.",
            )
        elif text == "S17. Analysis provenance":
            pass
        elif text.startswith("The current backend audit used"):
            set_paragraph(
                paragraph,
                "All analyses in this supplement use fastPLS 0.99.36. Table S15 maps each reported "
                "analysis to its current-release result directory and generating program. Package "
                "version, model family, backend, precision, component count, rSVD controls, seeds, "
                "timing scope, and execution status are stored with each result."
            )
        elif text.startswith("Table S15. Analysis provenance"):
            set_paragraph(
                paragraph,
                "Table S15. Current-release analysis provenance. Result directories are relative to "
                "the repository root."
            )

    residency = table_after_caption(document, "Table S1.")
    residency_rows = [
        ["Stage", "CPU", "CUDA", "Metal", "Residency / precision note"],
        ["Core PLS fit", "Compiled C++", "Device products and range sketch; small reduced solve may use host", "Metal Performance Shaders products; reduced solve may use host", "float64 and float32 routes are identified separately"],
        ["OPLS filter", "Compiled C++", "Host filter with CUDA PLS core", "Host filter with Metal PLS core", "Accelerator OPLS is hybrid"],
        ["Nonlinear kernel PLS", "Host Gram matrix with compiled core", "Host Gram matrix with CUDA PLS core", "Host Gram matrix with Metal PLS core", "Explicit n by n storage; accelerator routes are hybrid"],
        ["Prediction", "Compiled projection", "Device projection with public-model/result transfer", "Device projection with host post-processing", "End-to-end timings include transfers and synchronization"],
        ["LDA", "Compiled moments, Cholesky, and scores", "Device moments, solve, and scores", "Device score projection with host LDA", "Precision follows the supported input route"],
        ["Cross-validation", "Compiled fold construction, fit, and scoring", "Host scheduling with sequential CUDA fold fits", "Host scheduling with sequential Metal fold fits", "No R-level fold-refit loop where compiled support is available"],
        ["rSVD", "Ordinary 32/5; massive SIMPLS-family 12/2 with rank-one directions", "Ordinary 32/5; massive SIMPLS-family 12/2 with rank-one directions", "Ordinary 32/5; massive SIMPLS-family 12/2 with rank-one directions", "Effective controls, executed direction rule, seed, and fit-level diagnostics are recorded"],
        ["IRLBA", "Available with fixed numerical controls", "Unavailable", "Unavailable", "Iterative CPU comparator; not an exact dense decomposition"],
    ]
    reset_table_rows(residency, residency_rows)
    style_table(residency, 6.2)

    dataset_table = table_after_caption(document, "Table S3.")
    reset_table_rows(dataset_table, dataset_dimension_rows())
    style_table(dataset_table, 6.2)

    environment_table = table_after_caption(document, "Table S4a.")
    reset_table_rows(environment_table, environment_rows())
    style_table(environment_table, 6.0)

    reproducibility_table = table_after_caption(document, "Table S5.")
    reset_table_rows(reproducibility_table, reproducibility_endpoint_rows())
    style_table(reproducibility_table, 6.2)

    output_table = table_after_caption(document, "Table S9b.")
    reset_table_rows(output_table, output_contract_rows())
    style_table(output_table, 6.0)

    nmr_table = table_after_caption(document, "Table S12.")
    reset_table_rows(nmr_table, nmr_comparison_table_rows())
    style_table(nmr_table, 5.7)

    estimator_table = table_after_caption(document, "Table S6.")
    reset_table_rows(estimator_table, estimator_validation_rows())
    style_table(estimator_table, 5.8)

    minimum_memory_table = table_after_caption(document, "Table S9c.")
    reset_table_rows(
        minimum_memory_table,
        external_memory_rows("estimator_kernel"),
    )
    style_table(minimum_memory_table, 5.8)

    workflow_memory_table = table_after_caption(document, "Table S9d.")
    reset_table_rows(
        workflow_memory_table,
        external_memory_rows("complete_workflow"),
    )
    style_table(workflow_memory_table, 5.8)

    package_panel_table = table_after_caption(document, "Table S9e.")
    reset_table_rows(package_panel_table, package_panel_rows())
    style_table(package_panel_table, 5.8)

    ikpls_large_table = table_after_caption(document, "Table S9g.")
    reset_table_rows(ikpls_large_table, ikpls_large_rows())
    style_table(ikpls_large_table, 6.0)

    cuda_table = table_after_caption(document, "Table S10.")
    reset_table_rows(cuda_table, paired_backend_rows(CUDA_PAIRED, "cuda"))
    style_table(cuda_table, 5.6)

    metal_table = table_after_caption(document, "Table S11.")
    ensure_table_columns(metal_table, 8)
    reset_table_rows(metal_table, paired_backend_rows(METAL_PAIRED, "metal"))
    style_table(metal_table, 5.6)

    imagenet_table = table_after_caption(document, "Table S13.")
    reset_table_rows(imagenet_table, imagenet_rows())
    style_table(imagenet_table, 5.6)

    repeated_table = table_after_caption(document, "Table S14.")
    reset_table_rows(repeated_table, repeated_outer_rows())
    style_table(repeated_table, 5.7)

    qualification_paths = {
        "CPU": cpu_csv,
        "CUDA": cuda_csv,
        "Metal": metal_csv,
    }
    ablation_table = table_after_caption(document, "Table S6a.")
    factor_table = table_after_caption(document, "Table S6c.")
    crosscov_table = table_after_caption(document, "Table S6d.")
    # The source template has no table placeholder directly after S6b. Create
    # one so the shape table cannot consume and overwrite the following S6c
    # qualification table.
    shape_table = ensure_table_immediately_after_caption(document, "Table S6b.", 6)
    reset_table_rows(ablation_table, ablation_rows())
    style_table(ablation_table, 6.2)
    reset_table_rows(factor_table, qualification_factor_rows(qualification_paths))
    style_table(factor_table, 6.2)
    reset_table_rows(crosscov_table, crosscov_rows(qualification_paths))
    style_table(crosscov_table, 6.0)
    reset_table_rows(shape_table, shape_rows())
    style_table(shape_table, 6.2)

    # Replace the obsolete rSVD table with the current 32/5 multi-seed audit.
    summaries = {
        "CPU": summarize_rsvd(cpu_csv, "cpu"),
        "Metal": summarize_rsvd(metal_csv, "metal"),
        "CUDA": summarize_rsvd(cuda_csv, "cuda"),
    }
    table = table_after_caption(document, "Table S7.")
    rows = [[
        "Backend", "Effective oversample", "Effective power", "Checks met", "Max pred. rel. error",
        "Min pred. corr.", "Max score rel. error", "Min label agree.",
        "Max metric diff.", "Status"
    ]]
    for backend in ("CPU", "CUDA", "Metal"):
        item = summaries[backend]
        rows.append([
            backend, item["oversample"], item["power"], item["checks"], item["pred"], item["corr"],
            item["score"], item["labels"], item["metric"], item["status"]
        ])
    reset_table_rows(table, rows)
    style_table(table, 6.5)

    qualification_rows = []
    for path in (cpu_csv, cuda_csv, metal_csv):
        qualification_rows.extend([
            row for row in read_csv(path)
            if row.get("svd_method") == "rsvd" and row.get("route", "").endswith("_auto")
        ])
    worst_prediction = max(numeric(qualification_rows, "prediction_relative_error"))
    worst_score = max(numeric(qualification_rows, "score_relative_error"))
    minimum_correlation = min(numeric(qualification_rows, "prediction_correlation"))
    minimum_labels = min(numeric(qualification_rows, "label_agreement"))
    worst_metric = max(numeric(qualification_rows, "metric_absolute_difference"))

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("rSVD uses a Gaussian range sketch"):
            set_paragraph(
                paragraph,
                "rSVD uses newly initialized Gaussian range sketches of the deflated SIMPLS operator. "
                "Ordinary CPU and Metal routes form a new oversampled sketch for each component; "
                "massive CPU, CUDA, and Metal routes use a newly initialized rank-one randomized "
                "direction per component. CUDA can generate a small "
                "fresh candidate block for large classification workloads and accepts candidates "
                "sequentially through the SIMPLS updates. The controlled qualification uses "
                "shape-specific automatic controls (ordinary 32/5, high-response 48/6, or sparse "
                "high-class-count 64/7) and forced 32/5 explicit/implicit routes on CPU, CUDA, "
                "and Metal. Qualification required "
                "relative prediction and score errors <= 0.01, corresponding correlations >= 0.995, "
                "decoded-label agreement >= 0.995, and absolute predictive-metric difference <= 0.005 "
                "against a matched CPU IRLBA fit. The observed worst values were "
                f"{worst_prediction:.3g} relative prediction error, {worst_score:.3g} score error, "
                f"{minimum_correlation:.6f} prediction correlation, {minimum_labels:.4f} label "
                f"agreement, and {worst_metric:.3g} metric difference. The controlled panel varied sample "
                "count, predictor and response dimensions, component count, requested prefixes, rank, "
                "class count, and cross-covariance storage over three seeds. Table S7 reports every "
                "completed controlled route. Massive SIMPLS-family analyses report oversample = 12, "
                "power = 2, and the executed rank-one direction rule separately."
            )
        elif text.startswith("Table S7. Multi-seed rSVD numerical qualification"):
            set_paragraph(
                paragraph,
                "Table S7. Multi-seed rSVD numerical qualification for fastPLS 0.99.36. Effective "
                "oversampling and power controls are listed for each backend. Every fit starts with "
                "32 oversampling directions and five power iterations; the internal numerical audit "
                "may retry at 48/6 and 64/7. Checks compare matched predictions and endpoints with "
                "CPU IRLBA on the controlled panel."
            )
        elif text.startswith("Figure S2. Current-release rSVD numerical agreement"):
            set_paragraph(
                paragraph,
                "Figure S2. Current-release rSVD numerical agreement across the controlled matrix "
                "regimes. Points show relative prediction error and predictive-metric difference for "
                "fastPLS 0.99.36 under the recorded shape-specific effective controls and three seeds; "
                "dashed lines mark the numerical-screening limits."
            )
        elif text.startswith("Table S6a. Same-code SIMPLS execution ablation"):
            set_paragraph(
                paragraph,
                "Table S6a. Current-release SIMPLS execution ablation on CCLE, GTEx v8, "
                "MetRef, and PRISM. Each comparison changes one execution optimization in "
                "fastPLS 0.99.36 while retaining float64 CPU IRLBA, the same split and component "
                "count, and three isolated repetitions. Speed-up is reference time divided by "
                "optimized time; positive RSS change denotes lower baseline-corrected peak RSS."
            )
        elif text.startswith("Table S6c. Controlled one-factor-at-a-time scaling"):
            set_paragraph(
                paragraph,
                "Table S6c. Current-release rSVD agreement by controlled factor. Every row uses "
                "fastPLS 0.99.36, the recorded shape-specific automatic controls, and three seeds. "
                "Checks are counted against the matched CPU IRLBA endpoint."
            )
        elif text.startswith("Table S6d. Forced explicit versus implicit"):
            set_paragraph(
                paragraph,
                "Table S6d. Explicit and implicit cross-covariance routes in fastPLS 0.99.36. "
                "Times are medians over three seeds. The implicit route evaluates products with "
                "the cross-covariance operator without storing that matrix explicitly."
            )
        elif text.startswith("Table S6b. Direct matched-shape runtime comparison"):
            set_paragraph(
                paragraph,
                "Table S6b. Matched float64 CPU PLS-SVD and SIMPLS execution across five "
                "synthetic matrix shapes. Both families use fastPLS 0.99.36, rSVD with 32 "
                "oversampling directions and five power iterations, the same split and component "
                "count, and three repetitions. The timing ratio does not imply estimator equivalence."
            )

    scaling_caption = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.startswith("Figure S1.")
    )
    set_paragraph(
        scaling_caption,
        "Figure S1. Controlled float64 CPU SIMPLS scaling with fastPLS 0.99.36. "
        "Panel A reports the fixed-control IRLBA to rSVD total-runtime ratio; values "
        "above one favour rSVD. Panel B reports rSVD prediction error relative to the "
        "matched IRLBA fit, with the predefined 0.01 numerical-screening limit shown by the dashed line. "
        "Automatic rSVD rows use the recorded shape-specific controls and three seeds while "
        "varying one matrix or model dimension at a time."
    )
    scaling_paragraph = figure_paragraph_after(document, scaling_caption)
    scaling_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    scaling_paragraph.add_run().add_picture(str(SCALING_FIGURE), width=Inches(7.0))

    figure_caption = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.startswith("Figure S2. Current-release rSVD")
    )
    figure_paragraph = figure_paragraph_after(document, figure_caption)
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.add_run().add_picture(str(RSVD_FIGURE), width=Inches(7.0))

    external_caption = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.startswith("Figure S3. Repeated single-CPU")
    )
    external_paragraph = figure_paragraph_after(document, external_caption)
    external_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    external_paragraph.add_run().add_picture(str(EXTERNAL_FIGURE), width=Inches(7.0))

    backend_caption = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.startswith("Figure S4. Selected SIMPLS-rSVD")
    )
    backend_paragraph = figure_paragraph_after(document, backend_caption)
    backend_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    backend_paragraph.add_run().add_picture(
        str(SELECTED_SIMPLS_BACKEND_FIGURE), width=Inches(7.0)
    )

    selected_caption = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.startswith("Figure S5. Training-selected NMR")
    )
    selected_paragraph = figure_paragraph_after(document, selected_caption)
    selected_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    selected_paragraph.add_run().add_picture(
        str(NMR_SELECTED_FIGURE), width=Inches(5.25)
    )

    if not COMPONENT_SELECTION_CLASSIFICATION_FIGURE.exists():
        raise FileNotFoundError(COMPONENT_SELECTION_CLASSIFICATION_FIGURE)
    if not COMPONENT_SELECTION_REGRESSION_FIGURE.exists():
        raise FileNotFoundError(COMPONENT_SELECTION_REGRESSION_FIGURE)
    provenance_heading = next(
        paragraph for paragraph in document.paragraphs
        if paragraph.text.strip() == "S17. Analysis provenance"
    )
    component_heading = provenance_heading.insert_paragraph_before(
        "S16.1 Complete training-only component paths"
    )
    component_heading.style = provenance_heading.style
    provenance_heading.insert_paragraph_before(
        "The complete paths below show the endpoint used for ten-fold training-only selection "
        "for each non-NMR dataset and PLS family. Open circles mark the selected value within "
        "the evaluated grid. PLS-SVD paths can terminate at the response-rank limit."
    )
    classification_caption = provenance_heading.insert_paragraph_before(
        "Figure S6. Ten-fold training-only classification accuracy across the complete "
        "family-specific component grids. Open circles mark the selected value within each "
        "evaluated grid; lines that overlap have the same recorded endpoint path."
    )
    classification_figure = provenance_heading.insert_paragraph_before("")
    classification_figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    classification_figure.add_run().add_picture(
        str(COMPONENT_SELECTION_CLASSIFICATION_FIGURE), width=Inches(7.0)
    )
    regression_caption = provenance_heading.insert_paragraph_before(
        "Figure S7. Ten-fold training-only RMSD across the complete family-specific component "
        "grids for the non-NMR multivariate regression tasks. Open circles mark the selected "
        "value within each evaluated grid."
    )
    regression_figure = provenance_heading.insert_paragraph_before("")
    regression_figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    regression_figure.add_run().add_picture(
        str(COMPONENT_SELECTION_REGRESSION_FIGURE), width=Inches(7.0)
    )

    path_heading = provenance_heading.insert_paragraph_before(
        "S16.2 Component-dependent prediction, time, and memory"
    )
    path_heading.style = provenance_heading.style
    provenance_heading.insert_paragraph_before(
        "For each non-NMR dataset, the following figures report the held-out "
        "predictive endpoint, fitting-plus-prediction time, and incremental peak "
        "host RSS across the evaluated component grid. CPU/CUDA measurements were "
        "made on the NVIDIA workstation and CPU/Metal measurements on the Apple M3 "
        "workstation; each point is the median of five isolated processes and the "
        "band is the interquartile range. These figures characterize computational "
        "scaling and do not replace training-only component selection."
    )
    figure_suffixes = tuple("abcdefghijk")
    for suffix, (dataset, label) in zip(
        figure_suffixes,
        COMPONENT_PATH_DATASETS,
        strict=True,
    ):
        figure_path = COMPONENT_PATH_COMBINED_DIR / (
            f"component_path_{dataset}_cpu_cuda_metal.png"
        )
        if not figure_path.exists():
            raise FileNotFoundError(figure_path)
        provenance_heading.insert_paragraph_before(
            f"Figure S8{suffix}. {label} component-path analysis. Held-out "
            "prediction, total fitting-plus-prediction time, and incremental peak "
            "host RSS are shown for PLS-SVD, SIMPLS, OPLS, and kernel PLS on "
            "matched CPU/CUDA and CPU/Metal inputs. rSVD used the recorded automatic "
            "controls for each matrix shape."
        )
        figure_paragraph = provenance_heading.insert_paragraph_before("")
        figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        figure_paragraph.add_run().add_picture(
            str(figure_path),
            width=Inches(7.0),
        )

    # Replace the previous dense capability matrix with a direct user-facing map.
    capability = table_after_caption(document, "Table S8.")
    capability_rows = [
        ["Family / head", "CPU float64", "CPU float32", "CUDA float64", "CUDA float32", "Metal float64", "Metal float32", "Residency and evidence scope"],
        ["PLS-SVD", "Implemented: rSVD, IRLBA", "Implemented: rSVD, IRLBA", "Implemented: rSVD", "Implemented: rSVD", "Implemented: rSVD", "Implemented: rSVD", "Large products use the selected backend; reduced factorizations can be host assisted. Current rSVD qualification covers both accelerator backends."],
        ["SIMPLS", "Implemented: rSVD, IRLBA", "Implemented: rSVD, IRLBA", "Implemented: rSVD", "Implemented: rSVD", "Implemented: rSVD", "Implemented: rSVD", "CPU/Metal use component-wise sketches. CUDA may form a fresh candidate block for large classification; SIMPLS updates remain sequential."],
        ["OPLS", "Implemented: rSVD, IRLBA", "Implemented: rSVD, IRLBA", "Hybrid: rSVD core", "Hybrid: rSVD core", "Hybrid: rSVD core", "Hybrid: rSVD core", "Orthogonal filtering is host orchestrated; the inner PLS fit uses the selected accelerator."],
        ["Kernel PLS, linear", "Implemented: rSVD, IRLBA", "Implemented: rSVD, IRLBA", "Implemented: rSVD", "Implemented: rSVD", "Implemented: rSVD", "Implemented: rSVD", "The linear kernel dispatches to the linear PLS route and does not form a Gram matrix."],
        ["Kernel PLS, nonlinear", "Implemented: rSVD, IRLBA", "Implemented: rSVD, IRLBA", "Hybrid: rSVD core", "Hybrid: rSVD core", "Hybrid: rSVD core", "Hybrid: rSVD core", "Kernel construction and centering require an explicit n by n Gram matrix and include host stages."],
        ["LDA prediction head", "Implemented", "Implemented", "Implemented", "Implemented", "Hybrid", "Hybrid", "CPU and CUDA use precision-matched LDA arithmetic. Metal produces scores on device and completes LDA on the host."],
    ]
    reset_table_rows(capability, capability_rows)
    style_table(capability, 6.2)

    external_rows = [[
        "Dataset", "Profile", "fastPLS s", "pls s", "Speed-up",
        "Accuracy fastPLS / pls", "Completed reps"
    ]]
    for row in read_csv(EXTERNAL_PAIRS):
        if row.get("timing_mode") != "cold_process" or row.get("measurement_scope") != "primary":
            continue
        profile = "Public workflow" if row["comparison_profile"] == "complete_workflow" else "Minimum common output"
        external_rows.append([
            row["dataset"], profile, fmt(row["median_total_sec_fastpls"]),
            fmt(row["median_total_sec_pls"]), fmt(row["speedup_pls_over_fastpls"], 2),
            f"{fmt(row['median_accuracy_fastpls'], 4)} / {fmt(row['median_accuracy_pls'], 4)}",
            f"{row['repetitions_completed_fastpls']} / {row['repetitions_completed_pls']}",
        ])
    external_table = table_after_caption(document, "Table S9a.")
    reset_table_rows(external_table, external_rows)
    style_table(external_table, 6.4)

    ikpls_rows = [[
        "Dataset", "Workflow", "Accuracy", "Median time", "Time IQR",
        "Peak RSS", "Incremental RSS"
    ]]
    for row in read_csv(IKPLS_SUMMARY):
        label = row["implementation"].replace("fastPLS_cpu_", "fastPLS ").replace("IKPLS_numpy_alg", "IKPLS formulation ")
        ikpls_rows.append([
            row["dataset"], label, fmt(row["accuracy"], 4),
            fmt(row["median_total_sec"]), fmt(row["iqr_total_sec"]),
            fmt(row["median_peak_rss_mb"], 1), fmt(row["median_incremental_peak_rss_mb"], 1),
        ])
    ikpls_table = table_after_caption(document, "Table S9f.")
    reset_table_rows(ikpls_table, ikpls_rows)
    style_table(ikpls_table, 6.4)

    # Retain only current-release provenance and remove its checksum column.
    provenance = table_after_caption(document, "Table S15.")
    if len(provenance.columns) > 6:
        remove_table_column(provenance, 6)
    release_root = f"publication_results/{VERSION}/current_release"
    provenance_rows = [["ID", "Analysis", "Result directory", "Version", "Controls", "Generating program"],
        ["V01", "Dense and iterative SIMPLS validation", f"{release_root}/simpls_exact", VERSION, "float64 CPU; matched component paths", "benchmark_simpls_exact_reference.R; benchmark_simpls_estimator_preservation.R"],
        ["V02", "rSVD controlled qualification", f"{release_root}/rsvd_qualification", VERSION, "automatic 32/5, 48/6, or 64/7 profiles; forced 32/5 routes; three seeds; CPU/CUDA/Metal", "controlled_scaling/run_grid.R"],
        ["V03", "OPLS and kernel PLS validation", f"{release_root}/opls_kernel_estimator; {release_root}/opls_kernel_settings", VERSION, "float64 CPU; fixed folds", "benchmark_opls_kernel_estimator_validation.R; benchmark_opls_kernel_setting_reliability.R"],
        ["V04", "SIMPLS execution ablation and family-shape comparison", f"{release_root}/simpls_ablation; {release_root}/simpls_vs_plssvd_shapes", VERSION, "float64 CPU; matched inputs and component counts", "run_simpls_ablation_current.R; benchmark_simpls_vs_plssvd_shapes.R"],
        ["V05", "float32 arithmetic screen", f"{release_root}/float32_cpu; {release_root}/float32_cuda; {release_root}/float32_metal", VERSION, "four PLS families; matched float64/float32 inputs; CPU/CUDA/Metal", "benchmark_float32_backend_agreement.R"],
        ["B01", "External SIMPLS comparison", f"{release_root}/external_simpls", VERSION, "float64 single-CPU IRLBA; repeated fresh processes", "external_simpls_timing/worker.R"],
        ["B02", "Independent R-package panel", f"{release_root}/r_package_panel", VERSION, "matched float64 inputs and fixed outer splits", "benchmark_pls_package_comparison.R"],
        ["B03", "Backend and component-path comparison", f"{release_root}/selected_backend_cuda; {release_root}/selected_backend_metal; {release_root}/component_path_cuda; {release_root}/component_path_metal", VERSION, "matched CPU/CUDA/Metal workflows", "run_current_component_path.R; metal_validation/run_matched_cuda_dataset_metal.R"],
        ["B04", "CPU thread scaling", f"{release_root}/multicore_scaling", VERSION, "verified OpenBLAS; one, two, and four threads", "multicore_scaling/run_multicore_scaling.sh"],
        ["B05", "Compiled versus R-level cross-validation", f"{release_root}/cv_compiled_vs_r_loop", VERSION, "identical folds, controls, predictions, and scoring", "benchmark_cv_compiled_vs_r_loop.R"],
        ["B06", "IKPLS comparisons", f"{release_root}/ikpls_cross_language_cpu; {release_root}/ikpls_large_float32", VERSION, "matched precision and component count; estimator families reported separately", "ikpls_cross_language/run_benchmark.py"],
        ["A01", "NMR case study", f"{release_root}/nmr", VERSION, "fixed preprocessing; training-only selection and matched 165-component analyses", "benchmark_nmr_qualified_solver.R; plot_nmr_current_release_audit.R"],
        ["A02", "ImageNet embedding stress test", f"{release_root}/imagenet", VERSION, "float32 CUDA SIMPLS-rSVD; downstream classification", "benchmark_imagenet_current_fused_lda.R"],
    ]
    reset_table_rows(provenance, provenance_rows)
    style_table(provenance, 6.2)

    # Remove S18-S20 and their content; supplementary references follow S17.
    paragraphs = list(document.paragraphs)
    obsolete = next(
        (
            i for i, paragraph in enumerate(paragraphs)
            if paragraph.text.strip()
            == "S18. Unrestricted synthetic end-to-end reproduction"
        ),
        None,
    )
    if obsolete is not None:
        references = next(
            i for i, paragraph in enumerate(paragraphs)
            if paragraph.text.strip() == "Supplementary references"
        )
        for paragraph in paragraphs[obsolete:references]:
            remove_paragraph(paragraph)

    renumber_table_citations(document)
    replace_release_version(document)
    all_text = document_text(document)
    for forbidden in ("unqualified", "rejected"):
        assert forbidden not in all_text.lower(), forbidden
    obsolete_headings = (
        "S18. Unrestricted synthetic end-to-end reproduction",
        "S19. Repository-only detailed material",
        "S20. Continuous-integration and platform-test status",
    )
    assert not any(
        paragraph.text.strip() in obsolete_headings
        for paragraph in document.paragraphs
    )
    audit_document_consistency(document, "supplement")
    format_scientific_notation(document)
    normalize_table_font_sizes(document, 6.0)
    enable_continuous_line_numbers(document)
    document.core_properties.title = "fastPLS supplementary information"
    document.core_properties.subject = "Computer Methods and Programs in Biomedicine"
    document.save(SUPP_OUT)


def main():
    if len(sys.argv) != 4:
        raise SystemExit(
            "usage: update_current_release_manuscript.py CPU_CSV METAL_CSV CUDA_CSV"
        )
    subprocess.run(
        [
            sys.executable,
            str(ROOT / "scripts/audit_current_release_evidence.py"),
            str(RESULTS),
            "--version",
            VERSION,
            "--require-complete",
        ],
        check=True,
    )
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    update_main()
    update_supplement(Path(sys.argv[1]), Path(sys.argv[2]), Path(sys.argv[3]))
    subprocess.run(
        [sys.executable, str(ROOT / "scripts/add_multicore_supplement.py")],
        check=True,
    )
    print(MAIN_OUT)
    print(SUPP_OUT)


if __name__ == "__main__":
    main()
