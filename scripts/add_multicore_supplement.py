#!/usr/bin/env python3

from pathlib import Path
import re

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "artifacts/CMPB_rewrite_20260903_cycle138/fastPLS_CMPB_supplement_0.99.39.docx"
CSV = ROOT / "publication_results/0.99.39/current_release/multicore_scaling/multicore_scaling_summary.csv"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, value, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = str(value)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(6)
        run.bold = bold


def set_cell_width(cell, width_inches):
    width = int(Inches(width_inches))
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def insert_before(anchor, element):
    anchor._p.addprevious(element)


def renumber_tables_consecutively(doc):
    captions = [
        paragraph for paragraph in doc.paragraphs
        if re.match(r"^Table S\d+[a-z]?\.", paragraph.text)
    ]
    labels = [
        re.match(r"^Table (S\d+[a-z]?)\.", paragraph.text).group(1)
        for paragraph in captions
    ]
    if len(labels) != len(set(labels)):
        duplicates = sorted({label for label in labels if labels.count(label) > 1})
        raise RuntimeError(
            "Supplementary table captions are not unique: "
            + ", ".join(duplicates)
        )
    mapping = {
        label: f"S{index}"
        for index, label in enumerate(labels, start=1)
    }

    paragraphs = list(doc.paragraphs)
    paragraphs.extend(
        paragraph
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        for paragraph in cell.paragraphs
    )
    for paragraph in paragraphs:
        if not re.search(r"\bTables?\b", paragraph.text, re.IGNORECASE):
            continue
        updated = re.sub(
            r"\bS\d+[a-z]?\b",
            lambda match: mapping.get(match.group(0), match.group(0)),
            paragraph.text,
        )
        if updated != paragraph.text:
            paragraph.text = updated

    final_labels = [
        re.match(r"^Table (S\d+)\.", paragraph.text).group(1)
        for paragraph in doc.paragraphs
        if re.match(r"^Table S\d+\.", paragraph.text)
    ]
    expected = [f"S{index}" for index in range(1, len(final_labels) + 1)]
    if final_labels != expected:
        raise RuntimeError(
            "Supplementary tables are not consecutive after renumbering: "
            + ", ".join(final_labels)
        )


def main():
    data = pd.read_csv(CSV)
    doc = Document(DOCX)

    def measured_speedup(workload, cores):
        row = data[
            (data["workload"] == workload)
            & (data["requested_cores"] == cores)
        ]
        if len(row) != 1:
            raise RuntimeError(
                f"Expected one {workload!r}, {cores}-core result; found {len(row)}"
            )
        return float(row.iloc[0]["speedup"])

    intro = next(
        p for p in doc.paragraphs
        if p.text.startswith("Reproducibility experiments use identical prepared tasks")
    )
    intro.text = (
        "Reproducibility experiments use identical prepared tasks, component counts, "
        "model settings, and repetition indices within each CPU/accelerator pair. CPU "
        "builds delegate eligible dense linear algebra to the linked BLAS/LAPACK, while "
        "the sequential SIMPLS deflation steps remain serial. Multicore execution was "
        "therefore evaluated separately rather than inferred from library capability."
    )
    intro.style = "First Paragraph"

    stale_prefixes = (
        "S8.1 CPU thread-scaling benchmark",
        "fastPLS 0.99.39 was linked directly to OpenBLAS",
        "Table S4c. CPU thread-scaling of fastPLS SIMPLS/rSVD",
        "The response-wide workload benefited from parallel dense linear algebra",
    )
    for paragraph in list(doc.paragraphs):
        if paragraph.text.startswith(stale_prefixes):
            paragraph._element.getparent().remove(paragraph._element)
    for table in list(doc.tables):
        if table.rows and table.rows[0].cells[0].text.strip() == "Workload":
            table._element.getparent().remove(table._element)

    anchor = next(
        paragraph for paragraph in doc.paragraphs
        if paragraph.text.endswith("Backend reproducibility endpoints.")
    )

    heading = doc.add_paragraph("S8.1 CPU thread-scaling benchmark", style="Heading 2")
    insert_before(anchor, heading._p)

    method = doc.add_paragraph(style="First Paragraph")
    method.add_run(
        "fastPLS 0.99.39 was linked directly to OpenBLAS on the Apple M3 workstation. "
        "SIMPLS/rSVD float64 fitting and held-out prediction were run in five fresh R "
        "processes with one, two, or four OpenBLAS worker threads, corresponding to the "
        "requested CPU cores. A direct OpenBLAS runtime probe verified the active count "
        "before every fit. The controlled workloads represented sample-rich "
        "classification, predictor-wide regression, and response-wide regression. "
        "All runs used oversampling 32, five power iterations, seed 8127, and fixed "
        "component counts. Runtime includes fitting and prediction; speed-up is the "
        "one-core median divided by the corresponding multicore median."
    )
    insert_before(anchor, method._p)

    caption = doc.add_paragraph(style="Body Text")
    caption.add_run(
        "Table S4c. CPU thread-scaling of fastPLS SIMPLS/rSVD. Values are median total "
        "fit-plus-prediction time with the interquartile range in brackets across five "
        "fresh-process repetitions. Agreement indicates identical predictions and "
        "metrics across all repetitions and thread counts within a workload."
    )
    insert_before(anchor, caption._p)

    table = doc.add_table(rows=1, cols=8)
    table.style = "Table"
    table.autofit = False
    headers = [
        "Workload",
        "Shape (train/test; p; q)",
        "Components",
        "CPU cores",
        "Time, s [IQR]",
        "Speed-up",
        "Metric",
        "Agreement",
    ]
    widths = [1.25, 1.35, 0.70, 0.58, 1.05, 0.68, 0.82, 0.68]
    for index, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, header, bold=True)
        set_cell_width(cell, width)
        shade(cell, "D9EAF4")

    labels = {
        "sample-rich classification": "Sample-rich classification",
        "predictor-wide regression": "Predictor-wide regression",
        "response-wide regression": "Response-wide regression",
    }
    for _, row in data.sort_values(["workload", "requested_cores"]).iterrows():
        cells = table.add_row().cells
        values = [
            labels[row["workload"]],
            f'{int(row["n_train"])}/{int(row["n_test"])}; {int(row["p"])}; {int(row["q"])}',
            str(int(row["ncomp"])),
            str(int(row["active_openblas_threads"])),
            f'{row["median_sec"]:.3f} [{row["q1_sec"]:.3f}-{row["q3_sec"]:.3f}]',
            f'{row["speedup"]:.2f}x',
            (
                f'accuracy {row["metric_value"]:.3f}'
                if row["metric_name"] == "accuracy"
                else f'RMSD {row["metric_value"]:.3f}'
            ),
            "Identical" if bool(row["prediction_agreement"]) else "Different",
        ]
        for index, (value, width) in enumerate(zip(values, widths)):
            align = WD_ALIGN_PARAGRAPH.LEFT if index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[index], value, align=align)
            set_cell_width(cells[index], width)
        if int(row["requested_cores"]) == 4:
            for cell in cells:
                shade(cell, "F1F7FA")

    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    anchor._p.addprevious(table._tbl)

    interpretation = doc.add_paragraph(style="Body Text")
    response_two = measured_speedup("response-wide regression", 2)
    response_four = measured_speedup("response-wide regression", 4)
    sample_two = measured_speedup("sample-rich classification", 2)
    sample_four = measured_speedup("sample-rich classification", 4)
    predictor_two = measured_speedup("predictor-wide regression", 2)
    predictor_four = measured_speedup("predictor-wide regression", 4)
    interpretation.add_run(
        "The response-wide workload benefited from parallel dense linear algebra, with "
        f"{response_two:.2f}-fold and {response_four:.2f}-fold speed-ups at two and four "
        "cores, respectively. The sample-rich workload reached "
        f"{sample_two:.2f}-fold and {sample_four:.2f}-fold, and the predictor-wide "
        f"workload reached {predictor_two:.2f}-fold and {predictor_four:.2f}-fold. "
        "Values below one indicate slower execution because serial deflation and "
        "thread-management overhead dominated the shorter matrix operations. Multicore "
        "CPU execution is therefore supported, but its benefit depends on matrix shape "
        "and operation size; increasing the core count should not be assumed to "
        "accelerate every fit."
    )
    insert_before(anchor, interpretation._p)

    provenance = doc.tables[-1]
    if not any(
        "CPU thread scaling" in cell.text
        for row in provenance.rows
        for cell in row.cells
    ):
        values = [
            "B03",
            "CPU thread scaling",
            "publication_results/0.99.39/current_release/multicore_scaling",
            "0.99.39",
            "float64 CPU SIMPLS/rSVD; 1/2/4 verified OpenBLAS threads; five repetitions",
            "benchmark/multicore_scaling/run.sh",
        ]
        cells = provenance.add_row().cells
        for cell, value in zip(cells, values):
            set_cell_text(cell, value, align=WD_ALIGN_PARAGRAPH.LEFT)
        cant_split = OxmlElement("w:cantSplit")
        provenance.rows[-1]._tr.get_or_add_trPr().append(cant_split)

    renumber_tables_consecutively(doc)
    doc.save(DOCX)


if __name__ == "__main__":
    main()
